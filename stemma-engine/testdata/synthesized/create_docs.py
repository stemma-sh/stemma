# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
# ]
# ///
"""
Generate synthesized before/after DOCX pairs for testing the full
diff -> redline -> export pipeline.

Each subdirectory gets a before.docx, after.docx, and metadata.json.
"""

import json
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn


ROOT = Path(__file__).parent


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def save_pair(
    name: str,
    before: DocxDocument,
    after: DocxDocument,
    metadata: dict,
) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    before.save(str(out / "before.docx"))
    after.save(str(out / "after.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  {name}/")


def set_header_text(doc: DocxDocument, text: str) -> None:
    """Set default header text on a document."""
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()
    p.add_run(text)


def set_footer_text(doc: DocxDocument, text: str) -> None:
    """Set default footer text on a document."""
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.add_run(text)


# ---------------------------------------------------------------------------
# 1. text-substitution  (BlockModified -> ReplaceText)
# ---------------------------------------------------------------------------

def make_text_substitution() -> None:
    before = Document()
    before.add_paragraph("The quick brown fox jumps over the lazy dog.")
    before.add_paragraph("All good things must come to an end.")
    before.add_paragraph("This paragraph stays the same.")

    after = Document()
    after.add_paragraph("The quick red fox leaps over the lazy dog.")
    after.add_paragraph("All good things must come to an end.")
    after.add_paragraph("This paragraph stays the same.")

    save_pair("text-substitution", before, after, {
        "name": "text-substitution",
        "description": "Word-level text change within a paragraph",
        "pipeline": "BlockModified -> ReplaceText",
    })


# ---------------------------------------------------------------------------
# 2. paragraph-insertion  (BlockInserted -> InsertBlock)
# ---------------------------------------------------------------------------

def make_paragraph_insertion() -> None:
    before = Document()
    before.add_paragraph("First paragraph.")
    before.add_paragraph("Third paragraph.")

    after = Document()
    after.add_paragraph("First paragraph.")
    after.add_paragraph("Second paragraph, freshly inserted.")
    after.add_paragraph("Third paragraph.")

    save_pair("paragraph-insertion", before, after, {
        "name": "paragraph-insertion",
        "description": "Insert a new paragraph between existing ones",
        "pipeline": "BlockInserted -> InsertBlock",
    })


# ---------------------------------------------------------------------------
# 3. paragraph-deletion  (BlockDeleted -> ReplaceText(empty))
# ---------------------------------------------------------------------------

def make_paragraph_deletion() -> None:
    before = Document()
    before.add_paragraph("Keep this paragraph.")
    before.add_paragraph("This paragraph will be deleted.")
    before.add_paragraph("Keep this one too.")

    after = Document()
    after.add_paragraph("Keep this paragraph.")
    after.add_paragraph("Keep this one too.")

    save_pair("paragraph-deletion", before, after, {
        "name": "paragraph-deletion",
        "description": "Remove a paragraph entirely",
        "pipeline": "BlockDeleted -> ReplaceText(empty)",
    })


# ---------------------------------------------------------------------------
# 4. mixed-paragraph-changes  (all three block-level paths)
# ---------------------------------------------------------------------------

def make_mixed_paragraph_changes() -> None:
    before = Document()
    before.add_paragraph("Alpha paragraph with original text.")
    before.add_paragraph("Beta paragraph to be removed.")
    before.add_paragraph("Gamma paragraph unchanged.")
    before.add_paragraph("Delta paragraph with old wording.")

    after = Document()
    after.add_paragraph("Alpha paragraph with modified text.")
    # Beta removed
    after.add_paragraph("Gamma paragraph unchanged.")
    after.add_paragraph("Delta paragraph with new wording.")
    after.add_paragraph("Epsilon paragraph, freshly added.")

    save_pair("mixed-paragraph-changes", before, after, {
        "name": "mixed-paragraph-changes",
        "description": "Modify + insert + delete paragraphs in one document",
        "pipeline": "BlockModified + BlockDeleted + BlockInserted",
    })


# ---------------------------------------------------------------------------
# 5. table-cell-text  (BlockModified inside table cell)
# ---------------------------------------------------------------------------

def make_table_cell_text() -> None:
    before = Document()
    before.add_paragraph("Document with a table.")
    tbl = before.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "Header A"
    tbl.cell(0, 1).text = "Header B"
    tbl.cell(1, 0).text = "Original value"
    tbl.cell(1, 1).text = "Constant"

    after = Document()
    after.add_paragraph("Document with a table.")
    tbl = after.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "Header A"
    tbl.cell(0, 1).text = "Header B"
    tbl.cell(1, 0).text = "Updated value"
    tbl.cell(1, 1).text = "Constant"

    save_pair("table-cell-text", before, after, {
        "name": "table-cell-text",
        "description": "Change text inside a table cell",
        "pipeline": "BlockModified inside table cell",
    })


# ---------------------------------------------------------------------------
# 6. table-row-addition  (TableStructureChanged - skipped in tx)
# ---------------------------------------------------------------------------

def make_table_row_addition() -> None:
    before = Document()
    before.add_paragraph("Table with two rows.")
    tbl = before.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "A1"
    tbl.cell(0, 1).text = "B1"
    tbl.cell(1, 0).text = "A2"
    tbl.cell(1, 1).text = "B2"

    after = Document()
    after.add_paragraph("Table with two rows.")
    tbl = after.add_table(rows=3, cols=2)
    tbl.cell(0, 0).text = "A1"
    tbl.cell(0, 1).text = "B1"
    tbl.cell(1, 0).text = "A2"
    tbl.cell(1, 1).text = "B2"
    tbl.cell(2, 0).text = "A3"
    tbl.cell(2, 1).text = "B3"

    save_pair("table-row-addition", before, after, {
        "name": "table-row-addition",
        "description": "Add a row to a table (structure change, skipped in transaction)",
        "pipeline": "TableStructureChanged (skipped in tx, should not crash)",
    })


# ---------------------------------------------------------------------------
# 6b. table-row-deletion  (TableStructureChanged -> DeleteTableRow)
# ---------------------------------------------------------------------------

def make_table_row_deletion() -> None:
    before = Document()
    before.add_paragraph("Table with three rows.")
    tbl = before.add_table(rows=3, cols=2)
    tbl.cell(0, 0).text = "A1"
    tbl.cell(0, 1).text = "B1"
    tbl.cell(1, 0).text = "A2"
    tbl.cell(1, 1).text = "B2"
    tbl.cell(2, 0).text = "A3"
    tbl.cell(2, 1).text = "B3"

    after = Document()
    after.add_paragraph("Table with three rows.")
    tbl = after.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "A1"
    tbl.cell(0, 1).text = "B1"
    tbl.cell(1, 0).text = "A3"
    tbl.cell(1, 1).text = "B3"

    save_pair("table-row-deletion", before, after, {
        "name": "table-row-deletion",
        "description": "Delete a row from a table (middle row removed)",
        "pipeline": "TableStructureChanged -> DeleteTableRow",
    })


# ---------------------------------------------------------------------------
# 7. header-modification  (HeaderModified - skipped in tx)
# ---------------------------------------------------------------------------

def make_header_modification() -> None:
    before = Document()
    before.add_paragraph("Body text stays the same.")
    set_header_text(before, "Original Header")

    after = Document()
    after.add_paragraph("Body text stays the same.")
    set_header_text(after, "Modified Header")

    save_pair("header-modification", before, after, {
        "name": "header-modification",
        "description": "Change header text (skipped in transaction generation)",
        "pipeline": "HeaderModified (skipped in tx, should not crash)",
    })


# ---------------------------------------------------------------------------
# 8. footer-modification  (FooterModified - skipped in tx)
# ---------------------------------------------------------------------------

def make_footer_modification() -> None:
    before = Document()
    before.add_paragraph("Body text stays the same.")
    set_footer_text(before, "Original Footer")

    after = Document()
    after.add_paragraph("Body text stays the same.")
    set_footer_text(after, "Modified Footer")

    save_pair("footer-modification", before, after, {
        "name": "footer-modification",
        "description": "Change footer text (skipped in transaction generation)",
        "pipeline": "FooterModified (skipped in tx, should not crash)",
    })


# ---------------------------------------------------------------------------
# 9. opaque-redline-hyperlink  (text changes around w:hyperlink)
# ---------------------------------------------------------------------------

def _add_hyperlink(paragraph, text: str, url: str = "https://example.com") -> None:
    """Append a w:hyperlink element to a paragraph using low-level OPC."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = paragraph._element.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
    run = paragraph._element.makeelement(qn("w:r"), {})
    r_pr = paragraph._element.makeelement(qn("w:rPr"), {})
    r_style = paragraph._element.makeelement(qn("w:rStyle"), {qn("w:val"): "Hyperlink"})
    r_pr.append(r_style)
    run.append(r_pr)
    t = paragraph._element.makeelement(qn("w:t"), {})
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._element.append(hyperlink)


def make_opaque_redline_hyperlink() -> None:
    before = Document()
    p = before.add_paragraph("Text before ")
    _add_hyperlink(p, "Click here")
    run = p.add_run(" text after.")

    after = Document()
    p = after.add_paragraph("Modified before ")
    _add_hyperlink(p, "Click here")
    run = p.add_run(" modified after.")

    save_pair("opaque-redline-hyperlink", before, after, {
        "name": "opaque-redline-hyperlink",
        "description": "Text changes around a w:hyperlink — hyperlink must not end up inside w:del/w:ins",
        "pipeline": "BlockModified with OpaqueInline(Hyperlink)",
    })


# ---------------------------------------------------------------------------
# 10. opaque-redline-field  (text changes around w:fldSimple)
# ---------------------------------------------------------------------------

def _add_fld_simple(paragraph, instr: str, display: str) -> None:
    """Append a w:fldSimple element to a paragraph."""
    fld = paragraph._element.makeelement(qn("w:fldSimple"), {qn("w:instr"): instr})
    run = paragraph._element.makeelement(qn("w:r"), {})
    t = paragraph._element.makeelement(qn("w:t"), {})
    t.text = display
    run.append(t)
    fld.append(run)
    paragraph._element.append(fld)


def make_opaque_redline_field() -> None:
    before = Document()
    p = before.add_paragraph("See page ")
    _add_fld_simple(p, " PAGE ", "1")
    run = p.add_run(" for details.")

    after = Document()
    p = after.add_paragraph("Refer to page ")
    _add_fld_simple(p, " PAGE ", "1")
    run = p.add_run(" for more details.")

    save_pair("opaque-redline-field", before, after, {
        "name": "opaque-redline-field",
        "description": "Text changes around a w:fldSimple — field must not end up inside w:del/w:ins",
        "pipeline": "BlockModified with OpaqueInline(Field/Simple)",
    })


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    print("Generating synthesized test fixtures:")
    make_text_substitution()
    make_paragraph_insertion()
    make_paragraph_deletion()
    make_mixed_paragraph_changes()
    make_table_cell_text()
    make_table_row_addition()
    make_table_row_deletion()
    make_header_modification()
    make_footer_modification()
    make_opaque_redline_hyperlink()
    make_opaque_redline_field()
    print("Done.")


if __name__ == "__main__":
    main()

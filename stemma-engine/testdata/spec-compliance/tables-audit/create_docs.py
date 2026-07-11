# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for deep table property audit (ISO 29500-1 §17.4).

Each fixture exercises a specific table property from the spec that may
not be handled by our implementation.

Run:  python3 create_docs.py
"""

import json
from pathlib import Path
from lxml import etree

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Twips

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  tables-audit/{name}/")


def replace_tbl_child(tbl_pr, tag: str, new_el) -> None:
    """Remove existing child element with given tag from tblPr, then append new one."""
    existing = tbl_pr.find(qn(tag))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(new_el)


# =========================================================================
# 1) tblLayout — fixed vs autofit (§17.4.52/§17.4.53)
# =========================================================================

def make_tbl_layout():
    """Two tables: one with tblLayout type=fixed, one with type=autofit."""
    doc = Document()
    doc.add_paragraph("Table with tblLayout fixed:")

    # Table 1: fixed layout
    tbl1 = doc.add_table(rows=2, cols=2)
    tbl1.cell(0, 0).text = "Fixed A"
    tbl1.cell(0, 1).text = "Fixed B"
    tbl1.cell(1, 0).text = "Fixed C"
    tbl1.cell(1, 1).text = "Fixed D"

    tbl_pr1 = tbl1._tbl.tblPr
    layout1 = make_element("w:tblLayout", {"w:type": "fixed"})
    tbl_pr1.append(layout1)

    doc.add_paragraph("Table with tblLayout autofit:")

    # Table 2: autofit layout
    tbl2 = doc.add_table(rows=2, cols=2)
    tbl2.cell(0, 0).text = "Auto A"
    tbl2.cell(0, 1).text = "Auto B"
    tbl2.cell(1, 0).text = "Auto C"
    tbl2.cell(1, 1).text = "Auto D"

    tbl_pr2 = tbl2._tbl.tblPr
    layout2 = make_element("w:tblLayout", {"w:type": "autofit"})
    tbl_pr2.append(layout2)

    save_fixture("tbl-layout", doc, {
        "spec": "ISO 29500-1 §17.4.52/§17.4.53",
        "description": "Two tables: tblLayout fixed vs autofit",
        "tables": [
            {"layout": "fixed", "rows": 2, "cols": 2},
            {"layout": "autofit", "rows": 2, "cols": 2},
        ],
    })


# =========================================================================
# 2) tblCellSpacing (§17.4.44)
# =========================================================================

def make_tbl_cell_spacing():
    """Table with tblCellSpacing (gap between cells)."""
    doc = Document()
    doc.add_paragraph("Table with tblCellSpacing=100tw:")

    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "Spaced A"
    tbl.cell(0, 1).text = "Spaced B"
    tbl.cell(1, 0).text = "Spaced C"
    tbl.cell(1, 1).text = "Spaced D"

    tbl_pr = tbl._tbl.tblPr
    spacing = make_element("w:tblCellSpacing", {"w:w": "100", "w:type": "dxa"})
    tbl_pr.append(spacing)

    save_fixture("tbl-cell-spacing", doc, {
        "spec": "ISO 29500-1 §17.4.44",
        "description": "Table with tblCellSpacing=100 twips between cells",
        "cell_spacing_twips": 100,
    })


# =========================================================================
# 3) tblpPr — floating table positioning (§17.4.57)
# =========================================================================

def make_tblp_pr():
    """Table with floating positioning (tblpPr)."""
    doc = Document()
    doc.add_paragraph("Floating table follows:")

    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "Float A"
    tbl.cell(0, 1).text = "Float B"
    tbl.cell(1, 0).text = "Float C"
    tbl.cell(1, 1).text = "Float D"

    tbl_pr = tbl._tbl.tblPr
    tblp = make_element("w:tblpPr", {
        "w:leftFromText": "180",
        "w:rightFromText": "180",
        "w:topFromText": "0",
        "w:bottomFromText": "0",
        "w:vertAnchor": "text",
        "w:horzAnchor": "margin",
        "w:tblpXSpec": "center",
        "w:tblpY": "1440",
    })
    tbl_pr.append(tblp)

    save_fixture("tblp-pr", doc, {
        "spec": "ISO 29500-1 §17.4.57",
        "description": "Floating table with tblpPr positioning attributes",
        "positioning": {
            "leftFromText": 180,
            "rightFromText": 180,
            "vertAnchor": "text",
            "horzAnchor": "margin",
            "tblpXSpec": "center",
            "tblpY": 1440,
        },
    })


# =========================================================================
# 4) tblOverlap (§17.4.55)
# =========================================================================

def make_tbl_overlap():
    """Two floating tables with tblOverlap=never on the second."""
    doc = Document()
    doc.add_paragraph("Two overlapping floating tables:")

    # Table 1: floating, no overlap restriction
    tbl1 = doc.add_table(rows=2, cols=2)
    tbl1.cell(0, 0).text = "Overlap A"
    tbl1.cell(0, 1).text = "Overlap B"
    tbl1.cell(1, 0).text = "Overlap C"
    tbl1.cell(1, 1).text = "Overlap D"

    tbl_pr1 = tbl1._tbl.tblPr
    tblp1 = make_element("w:tblpPr", {
        "w:vertAnchor": "text",
        "w:horzAnchor": "margin",
        "w:tblpXSpec": "left",
        "w:tblpY": "0",
    })
    tbl_pr1.append(tblp1)

    # Table 2: floating, overlap=never
    tbl2 = doc.add_table(rows=2, cols=2)
    tbl2.cell(0, 0).text = "NoOverlap A"
    tbl2.cell(0, 1).text = "NoOverlap B"
    tbl2.cell(1, 0).text = "NoOverlap C"
    tbl2.cell(1, 1).text = "NoOverlap D"

    tbl_pr2 = tbl2._tbl.tblPr
    tblp2 = make_element("w:tblpPr", {
        "w:vertAnchor": "text",
        "w:horzAnchor": "margin",
        "w:tblpXSpec": "left",
        "w:tblpY": "0",
    })
    tbl_pr2.append(tblp2)
    overlap = make_element("w:tblOverlap", {"w:val": "never"})
    tbl_pr2.append(overlap)

    save_fixture("tbl-overlap", doc, {
        "spec": "ISO 29500-1 §17.4.55",
        "description": "Two floating tables, second has tblOverlap=never",
        "tables": [
            {"overlap": None},
            {"overlap": "never"},
        ],
    })


# =========================================================================
# 5) Band size > 1 (§17.7.6.5/§17.7.6.7)
# =========================================================================

def make_band_size_gt1():
    """Table style with rowBandSize=2, colBandSize=2 — verify banding groups."""
    doc = Document()

    # Create a custom table style with band sizes = 2
    styles_xml = doc.styles.element
    style_el = make_element("w:style", {"w:type": "table", "w:styleId": "BandSize2Style"})
    name_el = make_element("w:name", {"w:val": "BandSize2Style"})
    style_el.append(name_el)

    tbl_pr = make_element("w:tblPr")
    row_band = make_element("w:tblStyleRowBandSize", {"w:val": "2"})
    col_band = make_element("w:tblStyleColBandSize", {"w:val": "2"})
    tbl_pr.append(row_band)
    tbl_pr.append(col_band)
    style_el.append(tbl_pr)

    # Band1Horz (odd band): light blue shading
    band1h = make_element("w:tblStylePr", {"w:type": "band1Horz"})
    band1h_tc = make_element("w:tcPr")
    band1h_shd = make_element("w:shd", {"w:val": "clear", "w:color": "auto", "w:fill": "BDD7EE"})
    band1h_tc.append(band1h_shd)
    band1h.append(band1h_tc)
    style_el.append(band1h)

    # Band2Horz (even band): light orange shading
    band2h = make_element("w:tblStylePr", {"w:type": "band2Horz"})
    band2h_tc = make_element("w:tcPr")
    band2h_shd = make_element("w:shd", {"w:val": "clear", "w:color": "auto", "w:fill": "F8CBAD"})
    band2h_tc.append(band2h_shd)
    band2h.append(band2h_tc)
    style_el.append(band2h)

    styles_xml.append(style_el)

    doc.add_paragraph("Table with rowBandSize=2:")

    # 6-row table (no first/last row flags) with style applied
    tbl = doc.add_table(rows=6, cols=2)
    for r in range(6):
        for c in range(2):
            tbl.cell(r, c).text = f"R{r}C{c}"

    tbl_pr_el = tbl._tbl.tblPr
    style_ref = make_element("w:tblStyle", {"w:val": "BandSize2Style"})
    tbl_pr_el.insert(0, style_ref)

    # tblLook: no firstRow/lastRow, enable horizontal banding
    # Must replace the default tblLook that python-docx adds
    look = make_element("w:tblLook", {
        "w:firstRow": "0",
        "w:lastRow": "0",
        "w:firstColumn": "0",
        "w:lastColumn": "0",
        "w:noHBand": "0",
        "w:noVBand": "1",
    })
    replace_tbl_child(tbl_pr_el, "w:tblLook", look)

    save_fixture("band-size-gt1", doc, {
        "spec": "ISO 29500-1 §17.7.6.5/§17.7.6.7 (tblStyleRowBandSize)",
        "description": "Table style with rowBandSize=2. Rows 0-1=band1(blue BDD7EE), "
                       "rows 2-3=band2(orange F8CBAD), rows 4-5=band1(blue BDD7EE)",
        "row_band_size": 2,
        "expected_bands": {
            "rows_0_1": "BDD7EE",
            "rows_2_3": "F8CBAD",
            "rows_4_5": "BDD7EE",
        },
    })


# =========================================================================
# 6) Width types: pct, dxa, auto, nil (§17.4.84)
# =========================================================================

def make_width_types():
    """Four tables with different width types: dxa, pct, auto, nil."""
    doc = Document()

    # Table 1: dxa (fixed twips width)
    doc.add_paragraph("Table width type=dxa (5000tw):")
    tbl1 = doc.add_table(rows=1, cols=2)
    tbl1.cell(0, 0).text = "DXA col A"
    tbl1.cell(0, 1).text = "DXA col B"
    tbl_pr1 = tbl1._tbl.tblPr
    replace_tbl_child(tbl_pr1, "w:tblW", make_element("w:tblW", {"w:w": "5000", "w:type": "dxa"}))

    # Table 2: pct (percentage — value in 50ths of a percent, so 5000 = 100%)
    doc.add_paragraph("Table width type=pct (2500 = 50%):")
    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.cell(0, 0).text = "PCT col A"
    tbl2.cell(0, 1).text = "PCT col B"
    tbl_pr2 = tbl2._tbl.tblPr
    replace_tbl_child(tbl_pr2, "w:tblW", make_element("w:tblW", {"w:w": "2500", "w:type": "pct"}))

    # Table 3: auto
    doc.add_paragraph("Table width type=auto:")
    tbl3 = doc.add_table(rows=1, cols=2)
    tbl3.cell(0, 0).text = "Auto col A"
    tbl3.cell(0, 1).text = "Auto col B"
    tbl_pr3 = tbl3._tbl.tblPr
    replace_tbl_child(tbl_pr3, "w:tblW", make_element("w:tblW", {"w:w": "0", "w:type": "auto"}))

    # Table 4: nil (no width)
    doc.add_paragraph("Table width type=nil:")
    tbl4 = doc.add_table(rows=1, cols=2)
    tbl4.cell(0, 0).text = "Nil col A"
    tbl4.cell(0, 1).text = "Nil col B"
    tbl_pr4 = tbl4._tbl.tblPr
    replace_tbl_child(tbl_pr4, "w:tblW", make_element("w:tblW", {"w:w": "0", "w:type": "nil"}))

    save_fixture("width-types", doc, {
        "spec": "ISO 29500-1 §17.4.84",
        "description": "Four tables with different width types: dxa, pct, auto, nil",
        "tables": [
            {"width": 5000, "type": "dxa"},
            {"width": 2500, "type": "pct"},
            {"width": 0, "type": "auto"},
            {"width": 0, "type": "nil"},
        ],
    })


# =========================================================================
# 7) tblInd with explicit value (§17.4.51)
# =========================================================================

def make_tbl_ind():
    """Table with explicit tblInd=720 (half inch indent from leading margin)."""
    doc = Document()
    doc.add_paragraph("Table with tblInd=720tw:")

    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "Indent A"
    tbl.cell(0, 1).text = "Indent B"
    tbl.cell(1, 0).text = "Indent C"
    tbl.cell(1, 1).text = "Indent D"

    tbl_pr = tbl._tbl.tblPr
    tbl_ind = make_element("w:tblInd", {"w:w": "720", "w:type": "dxa"})
    tbl_pr.append(tbl_ind)

    save_fixture("tbl-ind", doc, {
        "spec": "ISO 29500-1 §17.4.51",
        "description": "Table with tblInd=720tw indent from leading margin",
        "indent_twips": 720,
    })


# =========================================================================
# 8) Cell width type=pct (§17.4.87)
# =========================================================================

def make_cell_width_pct():
    """Table with cell widths specified as percentages."""
    doc = Document()
    doc.add_paragraph("Table with cell widths in pct (50ths of percent):")

    tbl = doc.add_table(rows=2, cols=3)
    tbl.cell(0, 0).text = "25% col"
    tbl.cell(0, 1).text = "50% col"
    tbl.cell(0, 2).text = "25% col"
    tbl.cell(1, 0).text = "R1C0"
    tbl.cell(1, 1).text = "R1C1"
    tbl.cell(1, 2).text = "R1C2"

    # Set cell widths as pct: 1250 = 25%, 2500 = 50%
    for row in tbl.rows:
        cells_xml = row._tr.findall(qn("w:tc"))
        widths = [1250, 2500, 1250]
        for cell_xml, width in zip(cells_xml, widths):
            tc_pr = cell_xml.find(qn("w:tcPr"))
            if tc_pr is None:
                tc_pr = make_element("w:tcPr")
                cell_xml.insert(0, tc_pr)
            # Remove existing tcW if any
            existing = tc_pr.find(qn("w:tcW"))
            if existing is not None:
                tc_pr.remove(existing)
            tc_w = make_element("w:tcW", {"w:w": str(width), "w:type": "pct"})
            tc_pr.append(tc_w)

    save_fixture("cell-width-pct", doc, {
        "spec": "ISO 29500-1 §17.4.87",
        "description": "Table with cell widths in pct (50ths of percent): 25%, 50%, 25%",
        "cell_widths_pct50ths": [1250, 2500, 1250],
    })


# =========================================================================
# 9) Nested table (§17.4.38)
# =========================================================================

def make_nested_table():
    """Table with a nested table inside a cell."""
    doc = Document()
    doc.add_paragraph("Outer table with nested table in cell (1,0):")

    outer = doc.add_table(rows=2, cols=2)
    outer.cell(0, 0).text = "Outer R0C0"
    outer.cell(0, 1).text = "Outer R0C1"
    outer.cell(1, 1).text = "Outer R1C1"

    # Add nested table inside cell (1,0)
    cell_10 = outer.cell(1, 0)
    cell_10.text = ""  # Clear default paragraph
    p = cell_10.paragraphs[0]
    p.text = "Before nested"

    # Build nested table XML manually in the cell
    inner_tbl = OxmlElement("w:tbl")
    inner_tbl_pr = make_element("w:tblPr")
    inner_tbl_w = make_element("w:tblW", {"w:w": "0", "w:type": "auto"})
    inner_tbl_pr.append(inner_tbl_w)
    inner_borders = make_element("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = make_element(f"w:{edge}", {"w:val": "single", "w:sz": "4", "w:space": "0", "w:color": "000000"})
        inner_borders.append(b)
    inner_tbl_pr.append(inner_borders)
    inner_tbl.append(inner_tbl_pr)

    inner_grid = make_element("w:tblGrid")
    inner_grid.append(make_element("w:gridCol", {"w:w": "2000"}))
    inner_grid.append(make_element("w:gridCol", {"w:w": "2000"}))
    inner_tbl.append(inner_grid)

    for ri in range(2):
        tr = make_element("w:tr")
        for ci in range(2):
            tc = make_element("w:tc")
            tc_pr = make_element("w:tcPr")
            tc_w = make_element("w:tcW", {"w:w": "2000", "w:type": "dxa"})
            tc_pr.append(tc_w)
            tc.append(tc_pr)
            inner_p = make_element("w:p")
            inner_r = make_element("w:r")
            inner_t = make_element("w:t")
            inner_t.text = f"Inner R{ri}C{ci}"
            inner_r.append(inner_t)
            inner_p.append(inner_r)
            tc.append(inner_p)
            tr.append(tc)
        inner_tbl.append(tr)

    # Append nested table to cell, then add trailing paragraph (required by spec)
    cell_10._tc.append(inner_tbl)
    trailing_p = make_element("w:p")
    cell_10._tc.append(trailing_p)

    save_fixture("nested-table", doc, {
        "spec": "ISO 29500-1 §17.4.38",
        "description": "2x2 outer table with 2x2 nested table in cell (1,0)",
        "outer": {"rows": 2, "cols": 2},
        "inner": {"rows": 2, "cols": 2, "in_cell": [1, 0]},
    })


# =========================================================================
# 10) Cell properties: noWrap, textDirection, tcFitText (§17.4.30/72/63)
# =========================================================================

def make_cell_props() -> None:
    """Table with noWrap, textDirection, and tcFitText on individual cells.

    Cell (0,0): noWrap (no val = true)
    Cell (0,1): textDirection = tbRl
    Cell (1,0): tcFitText (no val = true)
    Cell (1,1): noWrap val=0 (false) + textDirection = btLr + tcFitText val=0 (false)
    """
    doc = Document()
    doc.add_paragraph("Before table.")

    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "NoWrap"
    tbl.cell(0, 1).text = "TextDir"
    tbl.cell(1, 0).text = "FitText"
    tbl.cell(1, 1).text = "Combined"

    tbl_element = tbl._tbl
    rows = tbl_element.findall(w("tr"))

    # Cell (0,0): noWrap (true)
    cell_00 = rows[0].findall(w("tc"))[0]
    tcPr_00 = cell_00.find(w("tcPr"))
    if tcPr_00 is None:
        tcPr_00 = make_element("w:tcPr")
        cell_00.insert(0, tcPr_00)
    tcPr_00.append(make_element("w:noWrap"))

    # Cell (0,1): textDirection = tbRl
    cell_01 = rows[0].findall(w("tc"))[1]
    tcPr_01 = cell_01.find(w("tcPr"))
    if tcPr_01 is None:
        tcPr_01 = make_element("w:tcPr")
        cell_01.insert(0, tcPr_01)
    tcPr_01.append(make_element("w:textDirection", {"w:val": "tbRl"}))

    # Cell (1,0): tcFitText (true)
    cell_10 = rows[1].findall(w("tc"))[0]
    tcPr_10 = cell_10.find(w("tcPr"))
    if tcPr_10 is None:
        tcPr_10 = make_element("w:tcPr")
        cell_10.insert(0, tcPr_10)
    tcPr_10.append(make_element("w:tcFitText"))

    # Cell (1,1): noWrap=false, textDirection=btLr, tcFitText=false
    cell_11 = rows[1].findall(w("tc"))[1]
    tcPr_11 = cell_11.find(w("tcPr"))
    if tcPr_11 is None:
        tcPr_11 = make_element("w:tcPr")
        cell_11.insert(0, tcPr_11)
    tcPr_11.append(make_element("w:noWrap", {"w:val": "0"}))
    tcPr_11.append(make_element("w:textDirection", {"w:val": "btLr"}))
    tcPr_11.append(make_element("w:tcFitText", {"w:val": "0"}))

    doc.add_paragraph("After table.")

    save_fixture("cell-props", doc, {
        "name": "cell-props",
        "spec_ref": "ISO 29500-1 §17.4.30, §17.4.63, §17.4.72",
        "description": "Table with noWrap, textDirection, tcFitText on cells",
        "expected_behavior": "All three cell properties parsed and preserved through roundtrip",
    })


# =========================================================================
# Entry point
# =========================================================================

def main():
    print("Generating tables-audit fixtures...")
    make_tbl_layout()
    make_tbl_cell_spacing()
    make_tblp_pr()
    make_tbl_overlap()
    make_band_size_gt1()
    make_width_types()
    make_tbl_ind()
    make_cell_width_pct()
    make_nested_table()
    make_cell_props()
    print("Done.")


if __name__ == "__main__":
    main()

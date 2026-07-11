# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for run formatting & font resolution deep audit.

Each fixture exercises a specific OOXML run property construct from
ISO 29500-1 section 17.3.2 (run properties) and section 17.8 (fonts).

Run:  uv run create_docs.py
"""

import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  runs-audit/{name}/")


# ── 1. bCs/iCs/szCs with cs=true ────────────────────────────────────────

def make_complex_script_properties() -> None:
    """Run with cs=true plus bCs, iCs, szCs — verify they override b/i/sz."""
    doc = Document()
    p = doc.add_paragraph()

    # Run 1: cs=true, bCs=true (should produce bold), normal sz=24, szCs=32
    run1 = p.add_run("CS bold run")
    rpr1 = run1._r.get_or_add_rPr()
    rpr1.append(make_element("w:cs"))
    rpr1.append(make_element("w:bCs"))
    rpr1.append(make_element("w:sz", {"w:val": "24"}))
    rpr1.append(make_element("w:szCs", {"w:val": "32"}))

    # Run 2: cs=true, iCs=true (should produce italic), szCs=28
    run2 = p.add_run("CS italic run")
    rpr2 = run2._r.get_or_add_rPr()
    rpr2.append(make_element("w:cs"))
    rpr2.append(make_element("w:iCs"))
    rpr2.append(make_element("w:sz", {"w:val": "24"}))
    rpr2.append(make_element("w:szCs", {"w:val": "28"}))

    # Run 3: rtl=true, bCs=true, iCs=true (both should apply via rtl)
    run3 = p.add_run("RTL bold italic run")
    rpr3 = run3._r.get_or_add_rPr()
    rpr3.append(make_element("w:rtl"))
    rpr3.append(make_element("w:bCs"))
    rpr3.append(make_element("w:iCs"))
    rpr3.append(make_element("w:szCs", {"w:val": "36"}))

    # Run 4: no cs/rtl, bCs=true (should NOT produce bold — bCs is only for CS)
    run4 = p.add_run("Non-CS bCs run")
    rpr4 = run4._r.get_or_add_rPr()
    rpr4.append(make_element("w:bCs"))

    save_fixture("complex-script-props", doc, {
        "name": "complex-script-props",
        "spec_ref": "ISO 29500-1 §17.3.2.2/17/39, MS-OI29500 §17.3.2.26(b)",
        "description": "Runs with complex script properties (bCs, iCs, szCs) "
                       "activated by w:cs and w:rtl",
        "expected_behavior": "cs/rtl causes bCs->bold, iCs->italic, szCs->fontSize",
    })


# ── 2. Theme font resolution ────────────────────────────────────────────

def make_theme_font_only() -> None:
    """Run with only asciiTheme/hAnsiTheme — no explicit font name.

    The theme should resolve to an actual font name from theme1.xml.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Run with asciiTheme="minorHAnsi" only (no w:ascii attribute)
    run = p.add_run("Theme minor font")
    rpr = run._r.get_or_add_rPr()
    rfonts = make_element("w:rFonts", {"w:asciiTheme": "minorHAnsi", "w:hAnsiTheme": "minorHAnsi"})
    rpr.append(rfonts)

    # Run with asciiTheme="majorHAnsi"
    run2 = p.add_run("Theme major font")
    rpr2 = run2._r.get_or_add_rPr()
    rfonts2 = make_element("w:rFonts", {"w:asciiTheme": "majorHAnsi", "w:hAnsiTheme": "majorHAnsi"})
    rpr2.append(rfonts2)

    # Run with all four theme slots
    run3 = p.add_run("All theme slots")
    rpr3 = run3._r.get_or_add_rPr()
    rfonts3 = make_element("w:rFonts", {
        "w:asciiTheme": "minorHAnsi",
        "w:hAnsiTheme": "minorHAnsi",
        "w:eastAsiaTheme": "minorEastAsia",
        "w:csTheme": "minorBidi",
    })
    rpr3.append(rfonts3)

    save_fixture("theme-font-only", doc, {
        "name": "theme-font-only",
        "spec_ref": "ISO 29500-1 §17.3.2.26, §17.8",
        "description": "Runs with theme font references only (no explicit font names)",
        "expected_behavior": "asciiTheme/hAnsiTheme should resolve to actual font names from theme1.xml",
    })


# ── 3. Run border (w:bdr) ───────────────────────────────────────────────

def make_run_border() -> None:
    """Run with w:bdr — character-level border.

    ISO 29500-1 §17.3.2.4: bdr specifies border around a run.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Run with a simple single-line border
    run = p.add_run("Bordered run")
    rpr = run._r.get_or_add_rPr()
    bdr = make_element("w:bdr", {
        "w:val": "single",
        "w:sz": "4",
        "w:space": "1",
        "w:color": "FF0000",
    })
    rpr.append(bdr)

    # Run without border for comparison
    p.add_run(" Normal run")

    save_fixture("run-border", doc, {
        "name": "run-border",
        "spec_ref": "ISO 29500-1 §17.3.2.4",
        "description": "Run with character-level border (w:bdr)",
        "expected_behavior": "Run border properties should be preserved in the model",
    })


# ── 4. w:position (vertical displacement) ───────────────────────────────

def make_run_position() -> None:
    """Run with w:position — vertical displacement in half-points.

    ISO 29500-1 §17.3.2.19: position raises or lowers text.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Run raised by 6 half-points (3pt)
    run_up = p.add_run("Raised text")
    rpr_up = run_up._r.get_or_add_rPr()
    rpr_up.append(make_element("w:position", {"w:val": "6"}))

    # Run lowered by 4 half-points (2pt)
    run_down = p.add_run("Lowered text")
    rpr_down = run_down._r.get_or_add_rPr()
    rpr_down.append(make_element("w:position", {"w:val": "-4"}))

    # Normal run for comparison
    p.add_run(" Normal text")

    save_fixture("run-position", doc, {
        "name": "run-position",
        "spec_ref": "ISO 29500-1 §17.3.2.19",
        "description": "Runs with vertical displacement via w:position",
        "expected_behavior": "Position offset should be preserved in the model",
    })


# ── 5. w:kern (kerning threshold) ───────────────────────────────────────

def make_run_kerning() -> None:
    """Run with w:kern — kerning threshold in half-points.

    ISO 29500-1 §17.3.2.19(a): kern specifies minimum font size for kerning.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Run with kerning threshold of 28 half-points (14pt)
    run = p.add_run("Kerned text")
    rpr = run._r.get_or_add_rPr()
    rpr.append(make_element("w:kern", {"w:val": "28"}))

    # Run without kerning for comparison
    p.add_run(" Normal text")

    save_fixture("run-kerning", doc, {
        "name": "run-kerning",
        "spec_ref": "ISO 29500-1 §17.3.2.19(a)",
        "description": "Run with kerning threshold (w:kern)",
        "expected_behavior": "Kerning threshold should be preserved in the model",
    })


# ── 6. w:w (character width scaling) ────────────────────────────────────

def make_char_width_scaling() -> None:
    """Run with w:w — character width scaling as percentage.

    ISO 29500-1 §17.3.2.43: w specifies character width scaling (100 = normal).
    """
    doc = Document()
    p = doc.add_paragraph()

    # Run at 150% width
    run_wide = p.add_run("Wide text")
    rpr_wide = run_wide._r.get_or_add_rPr()
    rpr_wide.append(make_element("w:w", {"w:val": "150"}))

    # Run at 50% width
    run_narrow = p.add_run("Narrow text")
    rpr_narrow = run_narrow._r.get_or_add_rPr()
    rpr_narrow.append(make_element("w:w", {"w:val": "50"}))

    # Normal run for comparison
    p.add_run(" Normal text")

    save_fixture("char-width-scaling", doc, {
        "name": "char-width-scaling",
        "spec_ref": "ISO 29500-1 §17.3.2.43",
        "description": "Runs with character width scaling (w:w)",
        "expected_behavior": "Character width percentage should be preserved in the model",
    })


# ── 7. w:rtl (right-to-left flag) ───────────────────────────────────────

def make_rtl_run() -> None:
    """Run with w:rtl — forces cs font slot and bidi layout.

    MS-OI29500 §17.3.2.26(b): rtl=true forces cs font slot for all chars.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Run with rtl=true and explicit cs font
    run_rtl = p.add_run("RTL run")
    rpr = run_rtl._r.get_or_add_rPr()
    rpr.append(make_element("w:rtl"))
    rfonts = make_element("w:rFonts", {
        "w:ascii": "Arial",
        "w:hAnsi": "Arial",
        "w:cs": "David",
    })
    rpr.append(rfonts)

    # Non-RTL run with same fonts
    run_ltr = p.add_run("LTR run")
    rpr2 = run_ltr._r.get_or_add_rPr()
    rfonts2 = make_element("w:rFonts", {
        "w:ascii": "Arial",
        "w:hAnsi": "Arial",
        "w:cs": "David",
    })
    rpr2.append(rfonts2)

    save_fixture("rtl-run", doc, {
        "name": "rtl-run",
        "spec_ref": "MS-OI29500 §17.3.2.26(b)",
        "description": "RTL run should force cs font slot",
        "expected_behavior": "rtl=true run uses cs font ('David') as font_family",
    })


# ── 8. Multiple font slots with explicit names ──────────────────────────

def make_multi_font_slots() -> None:
    """Run with all four font slots set to different explicit names."""
    doc = Document()
    p = doc.add_paragraph()

    run = p.add_run("Multi-font run")
    rpr = run._r.get_or_add_rPr()
    rfonts = make_element("w:rFonts", {
        "w:ascii": "Courier New",
        "w:hAnsi": "Courier New",
        "w:eastAsia": "MS Mincho",
        "w:cs": "Traditional Arabic",
    })
    rpr.append(rfonts)

    save_fixture("multi-font-slots", doc, {
        "name": "multi-font-slots",
        "spec_ref": "ISO 29500-1 §17.3.2.26",
        "description": "Run with all four rFonts slots set to different fonts",
        "expected_behavior": "Each font slot should be independently captured",
    })


# ═════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("Generating runs-audit fixtures...")
    make_complex_script_properties()
    make_theme_font_only()
    make_run_border()
    make_run_position()
    make_run_kerning()
    make_char_width_scaling()
    make_rtl_run()
    make_multi_font_slots()
    print("\nDone.")

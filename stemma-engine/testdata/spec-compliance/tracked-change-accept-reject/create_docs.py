# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for tracked change accept/reject spec-compliance tests.

These fixtures exercise the OOXML accept/reject semantics described in
ECMA-376 Part 1, Section 17.13.5 (Tracked Revisions).

Run:  uv run create_docs.py
"""

import json
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(
    name: str,
    doc: DocxDocument,
    metadata: dict,
    filename: str = "input.docx",
) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / filename))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  tracked-change-accept-reject/{name}/")


# ── Fixtures ──────────────────────────────────────────────────────────────


def make_para_mark_del_merge() -> None:
    """Two paragraphs where the first has a deleted paragraph mark.

    ECMA-376 §17.13.5.15: Accepting the deletion merges the first paragraph
    into the following paragraph, with the SECOND paragraph's properties winning.

    Structure:
      <w:p>  (style=Heading1, has w:del on para mark)
        <w:pPr>
          <w:pStyle w:val="Heading1"/>
          <w:jc w:val="center"/>
          <w:rPr><w:del .../></w:rPr>
        </w:pPr>
        <w:r><w:t>First paragraph text </w:t></w:r>
      </w:p>
      <w:p>  (alignment=left, the "following" paragraph)
        <w:pPr>
          <w:jc w:val="left"/>
        </w:pPr>
        <w:r><w:t>second paragraph text.</w:t></w:r>
      </w:p>
    """
    doc = Document()

    # First paragraph: has deleted paragraph mark, center aligned
    p1 = doc.add_paragraph("First paragraph text ")
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pPr1 = p1._p.get_or_add_pPr()
    rPr = make_element("w:rPr")
    del_el = make_element("w:del", {
        "w:id": "2000",
        "w:author": "Accept Test",
        "w:date": "2025-01-15T10:00:00Z",
    })
    rPr.append(del_el)
    pPr1.append(rPr)

    # Second paragraph: left aligned (the "following" paragraph whose props win)
    p2 = doc.add_paragraph("second paragraph text.")
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    save_fixture("para-mark-del-merge", doc, {
        "name": "para-mark-del-merge",
        "spec_ref": "ECMA-376 §17.13.5.15",
        "description": "Two paragraphs; first has deleted para mark (center), second is left-aligned. "
                       "Accept should merge into one paragraph with second's properties (left).",
    })


def make_para_mark_ins_split() -> None:
    """Two paragraphs where the first has an inserted paragraph mark.

    ECMA-376 §17.13.5.20: Accepting the insertion keeps the split.
    Rejecting it merges the paragraphs back together.

    Structure:
      <w:p>
        <w:pPr>
          <w:jc w:val="right"/>
          <w:rPr><w:ins .../></w:rPr>
        </w:pPr>
        <w:r><w:t>Before split </w:t></w:r>
      </w:p>
      <w:p>
        <w:pPr>
          <w:jc w:val="left"/>
        </w:pPr>
        <w:r><w:t>after split.</w:t></w:r>
      </w:p>
    """
    doc = Document()

    # First paragraph: has inserted paragraph mark
    p1 = doc.add_paragraph("Before split ")
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    pPr1 = p1._p.get_or_add_pPr()
    rPr = make_element("w:rPr")
    ins_el = make_element("w:ins", {
        "w:id": "2100",
        "w:author": "Accept Test",
        "w:date": "2025-01-15T11:00:00Z",
    })
    rPr.append(ins_el)
    pPr1.append(rPr)

    # Second paragraph
    p2 = doc.add_paragraph("after split.")
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    save_fixture("para-mark-ins-split", doc, {
        "name": "para-mark-ins-split",
        "spec_ref": "ECMA-376 §17.13.5.20",
        "description": "Two paragraphs; first has inserted para mark. "
                       "Accept keeps split (2 paras), reject merges back (1 para).",
    })


def make_deleted_table_row() -> None:
    """Table with one normal row and one deleted row.

    ECMA-376 §17.13.5.12: A deleted row should be removed on accept.
    """
    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "A1 normal"
    tbl.cell(0, 1).text = "B1 normal"
    tbl.cell(1, 0).text = "A2 deleted"
    tbl.cell(1, 1).text = "B2 deleted"

    tbl_element = tbl._tbl
    rows = tbl_element.findall(w("tr"))

    # Mark row 1 as deleted
    tr1 = rows[1]
    trPr = tr1.find(w("trPr"))
    if trPr is None:
        trPr = make_element("w:trPr")
        tr1.insert(0, trPr)
    del_el = make_element("w:del", {
        "w:id": "2200",
        "w:author": "Row Delete Author",
        "w:date": "2025-02-01T09:00:00Z",
    })
    trPr.append(del_el)

    save_fixture("deleted-table-row", doc, {
        "name": "deleted-table-row",
        "spec_ref": "ECMA-376 §17.13.5.12",
        "description": "Table with 2 rows; second row is marked as deleted. "
                       "Accept should remove the row, reject should keep it.",
    })


def make_ppr_change_reject() -> None:
    """Paragraph with pPrChange: current center, previous left.

    ECMA-376 §17.13.5.29: The child pPr inside pPrChange is the COMPLETE
    previous state. Rejecting should fully restore it.
    """
    doc = Document()
    p = doc.add_paragraph("Alignment was changed from left to center.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pPr = p._p.get_or_add_pPr()
    ppr_change = make_element("w:pPrChange", {
        "w:id": "2300",
        "w:author": "Format Author",
        "w:date": "2025-03-01T14:00:00Z",
    })
    prev_ppr = make_element("w:pPr")
    prev_jc = make_element("w:jc", {"w:val": "left"})
    prev_ppr.append(prev_jc)
    ppr_change.append(prev_ppr)
    pPr.append(ppr_change)

    save_fixture("ppr-change-reject", doc, {
        "name": "ppr-change-reject",
        "spec_ref": "ECMA-376 §17.13.5.29",
        "description": "Paragraph alignment changed left->center via pPrChange. "
                       "Reject should restore left alignment.",
    })


# ── Main ──────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("\n-- Tracked Change Accept/Reject Fixtures --")
    make_para_mark_del_merge()
    make_para_mark_ins_split()
    make_deleted_table_row()
    make_ppr_change_reject()
    print("\nDone.")

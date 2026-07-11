# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for contextualSpacing and spacing attribute tests.

Covers:
  - contextualSpacing between same-style paragraphs (ECMA-376 §17.3.1.9)
  - contextualSpacing explicit false overriding style's true
  - contextualSpacing false roundtrip (serializer bug)
  - beforeAutospacing / afterAutospacing (§17.3.1.33)
  - basic before/after spacing
  - line spacing with auto and exact rules

Run:  uv run create_docs.py
"""

import json
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc: DocxDocument, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  spacing-contextual/{name}/")


def save_pair(name: str, before: DocxDocument, after: DocxDocument, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    before.save(str(out / "before.docx"))
    after.save(str(out / "after.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  spacing-contextual/{name}/")


# =========================================================================
# Fixture 1: contextualSpacing suppresses between same style
# =========================================================================

def make_contextual_same_style() -> None:
    """ECMA-376 §17.3.1.9: contextualSpacing suppresses space between
    consecutive paragraphs with the same pStyle.

    Two paragraphs with the same style, both with contextualSpacing and
    spacing after="200". The spacing between them should be suppressed.
    """
    doc = Document()

    # Create style with contextualSpacing and after spacing
    styles_el = doc.styles.element
    style_el = make_element("w:style", {"w:type": "paragraph", "w:styleId": "CtxSameStyle"})
    name_el = make_element("w:name", {"w:val": "CtxSameStyle"})
    style_el.append(name_el)
    pPr_style = make_element("w:pPr")
    pPr_style.append(make_element("w:contextualSpacing"))
    pPr_style.append(make_element("w:spacing", {"w:after": "200"}))
    style_el.append(pPr_style)
    styles_el.append(style_el)

    # P0: first paragraph with CtxSameStyle
    p0 = doc.add_paragraph("P0: CtxSameStyle with contextualSpacing + after=200")
    p0.style = doc.styles["CtxSameStyle"]

    # P1: second paragraph with CtxSameStyle — spacing between P0 and P1 should be suppressed
    p1 = doc.add_paragraph("P1: CtxSameStyle with contextualSpacing + after=200")
    p1.style = doc.styles["CtxSameStyle"]

    save_fixture("contextual-same-style", doc, {
        "name": "contextual-same-style",
        "spec_ref": "ECMA-376 §17.3.1.9",
        "description": "contextualSpacing suppresses space between same-style paragraphs",
        "expected_behavior": (
            "P0 and P1 both have contextual_spacing=true (from CtxSameStyle). "
            "Space between them should be suppressed when rendered."
        ),
    })


# =========================================================================
# Fixture 2: contextualSpacing explicit false overrides style
# =========================================================================

def make_contextual_explicit_false() -> None:
    """ECMA-376 §17.3.1.9: Explicit contextualSpacing val="0" overrides
    style's true.

    Style CtxStyle defines contextualSpacing=true. P0 has direct
    contextualSpacing val="0" to override. P1 inherits from style (true).
    """
    doc = Document()

    # Create style with contextualSpacing=true
    styles_el = doc.styles.element
    style_el = make_element("w:style", {"w:type": "paragraph", "w:styleId": "CtxStyle"})
    name_el = make_element("w:name", {"w:val": "CtxStyle"})
    style_el.append(name_el)
    pPr_style = make_element("w:pPr")
    pPr_style.append(make_element("w:contextualSpacing"))
    pPr_style.append(make_element("w:spacing", {"w:before": "240", "w:after": "240"}))
    style_el.append(pPr_style)
    styles_el.append(style_el)

    # P0: explicit contextualSpacing val="0" overrides style
    p0 = doc.add_paragraph("P0: CtxStyle + direct contextualSpacing=false")
    p0.style = doc.styles["CtxStyle"]
    pPr0 = p0._p.get_or_add_pPr()
    cs0 = make_element("w:contextualSpacing", {"w:val": "0"})
    pPr0.append(cs0)

    # P1: inherits contextualSpacing=true from style
    p1 = doc.add_paragraph("P1: CtxStyle (inherits contextualSpacing=true)")
    p1.style = doc.styles["CtxStyle"]

    save_fixture("contextual-explicit-false", doc, {
        "name": "contextual-explicit-false",
        "spec_ref": "ECMA-376 §17.3.1.9",
        "description": "Direct contextualSpacing val='0' overrides style's true",
        "expected_behavior": (
            "P0: contextual_spacing=false (direct val='0' overrides style). "
            "P1: contextual_spacing=true (from style)."
        ),
    })


# =========================================================================
# Fixture 3: contextualSpacing false roundtrip (before/after pair)
# =========================================================================

def make_contextual_false_roundtrip() -> None:
    """Test that explicit contextualSpacing val="0" survives roundtrip.

    Before: style has contextualSpacing=true, paragraph has no override.
    After: paragraph gets contextualSpacing val="0" override.
    The diff_and_redline should preserve the explicit false.
    """
    # Before doc: style with contextualSpacing=true, paragraph inherits it
    before = Document()
    styles_el = before.styles.element
    style_el = make_element("w:style", {"w:type": "paragraph", "w:styleId": "CtxRT"})
    name_el = make_element("w:name", {"w:val": "CtxRT"})
    style_el.append(name_el)
    pPr_style = make_element("w:pPr")
    pPr_style.append(make_element("w:contextualSpacing"))
    pPr_style.append(make_element("w:spacing", {"w:after": "200"}))
    style_el.append(pPr_style)
    styles_el.append(style_el)

    p_before = before.add_paragraph("Paragraph with CtxRT style")
    p_before.style = before.styles["CtxRT"]

    # After doc: same style, but paragraph overrides contextualSpacing to false
    after = Document()
    styles_el2 = after.styles.element
    style_el2 = make_element("w:style", {"w:type": "paragraph", "w:styleId": "CtxRT"})
    name_el2 = make_element("w:name", {"w:val": "CtxRT"})
    style_el2.append(name_el2)
    pPr_style2 = make_element("w:pPr")
    pPr_style2.append(make_element("w:contextualSpacing"))
    pPr_style2.append(make_element("w:spacing", {"w:after": "200"}))
    style_el2.append(pPr_style2)
    styles_el2.append(style_el2)

    p_after = after.add_paragraph("Paragraph with CtxRT style")
    p_after.style = after.styles["CtxRT"]
    pPr_after = p_after._p.get_or_add_pPr()
    pPr_after.append(make_element("w:contextualSpacing", {"w:val": "0"}))

    save_pair("contextual-false-roundtrip", before, after, {
        "name": "contextual-false-roundtrip",
        "spec_ref": "ECMA-376 §17.3.1.9",
        "description": "Explicit contextualSpacing val='0' must survive diff/redline roundtrip",
        "expected_behavior": (
            "Before: contextual_spacing=true (from style). "
            "After: contextual_spacing=false (direct val='0'). "
            "Redline output should emit contextualSpacing val='0' on the paragraph."
        ),
    })


# =========================================================================
# Fixture 4: beforeAutospacing parsed
# =========================================================================

def make_before_autospacing() -> None:
    """ECMA-376 §17.3.1.33: beforeAutospacing overrides before.

    P0: spacing before="200" beforeAutospacing="1"
    Per spec, before value should be ignored when autospacing is set.
    """
    doc = Document()

    p0 = doc.add_paragraph("P0: before=200 with beforeAutospacing=1")
    pPr0 = p0._p.get_or_add_pPr()
    sp0 = make_element("w:spacing", {
        "w:before": "200",
        "w:beforeAutospacing": "1",
    })
    pPr0.append(sp0)

    # P1: control — no autospacing
    p1 = doc.add_paragraph("P1: before=200 only (control)")
    pPr1 = p1._p.get_or_add_pPr()
    sp1 = make_element("w:spacing", {"w:before": "200"})
    pPr1.append(sp1)

    save_fixture("before-autospacing", doc, {
        "name": "before-autospacing",
        "spec_ref": "ECMA-376 §17.3.1.33",
        "description": "beforeAutospacing overrides before value",
        "expected_behavior": (
            "P0: before=200 present but beforeAutospacing=1 means it should be ignored. "
            "P1: before=200 (control, no autospacing)."
        ),
    })


# =========================================================================
# Fixture 5: afterAutospacing overrides after
# =========================================================================

def make_after_autospacing() -> None:
    """ECMA-376 §17.3.1.33: afterAutospacing overrides after.

    P0: spacing after="200" afterAutospacing="1"
    Per spec, after value should be ignored when autospacing is set.
    """
    doc = Document()

    p0 = doc.add_paragraph("P0: after=200 with afterAutospacing=1")
    pPr0 = p0._p.get_or_add_pPr()
    sp0 = make_element("w:spacing", {
        "w:after": "200",
        "w:afterAutospacing": "1",
    })
    pPr0.append(sp0)

    # P1: control — no autospacing
    p1 = doc.add_paragraph("P1: after=200 only (control)")
    pPr1 = p1._p.get_or_add_pPr()
    sp1 = make_element("w:spacing", {"w:after": "200"})
    pPr1.append(sp1)

    save_fixture("after-autospacing", doc, {
        "name": "after-autospacing",
        "spec_ref": "ECMA-376 §17.3.1.33",
        "description": "afterAutospacing overrides after value",
        "expected_behavior": (
            "P0: after=200 present but afterAutospacing=1 means it should be ignored. "
            "P1: after=200 (control, no autospacing)."
        ),
    })


# =========================================================================
# Fixture 6: basic before/after spacing
# =========================================================================

def make_spacing_before_after_basic() -> None:
    """ECMA-376 §17.3.1.33: Basic before/after spacing values.

    P0: spacing before="240" after="200"
    """
    doc = Document()

    p0 = doc.add_paragraph("P0: before=240 after=200")
    pPr0 = p0._p.get_or_add_pPr()
    sp0 = make_element("w:spacing", {
        "w:before": "240",
        "w:after": "200",
    })
    pPr0.append(sp0)

    save_fixture("spacing-before-after-basic", doc, {
        "name": "spacing-before-after-basic",
        "spec_ref": "ECMA-376 §17.3.1.33",
        "description": "Basic before/after spacing values parsed correctly",
        "expected_behavior": "P0: before=240, after=200.",
    })


# =========================================================================
# Fixture 7: line spacing auto (single)
# =========================================================================

def make_spacing_line_auto() -> None:
    """ECMA-376 §17.3.1.33: line=240 lineRule=auto means single spacing.

    240 = single line spacing in auto mode (240ths of a line).
    """
    doc = Document()

    p0 = doc.add_paragraph("P0: line=240 lineRule=auto (single spacing)")
    pPr0 = p0._p.get_or_add_pPr()
    sp0 = make_element("w:spacing", {
        "w:line": "240",
        "w:lineRule": "auto",
    })
    pPr0.append(sp0)

    save_fixture("spacing-line-auto", doc, {
        "name": "spacing-line-auto",
        "spec_ref": "ECMA-376 §17.3.1.33",
        "description": "line=240 lineRule=auto is single spacing",
        "expected_behavior": "P0: line=240, lineRule=Auto.",
    })


# =========================================================================
# Fixture 8: line spacing exact
# =========================================================================

def make_spacing_line_exact() -> None:
    """ECMA-376 §17.3.1.33: line=360 lineRule=exact means exact 18pt.

    360 twips = 18pt exact line height.
    """
    doc = Document()

    p0 = doc.add_paragraph("P0: line=360 lineRule=exact (18pt exact)")
    pPr0 = p0._p.get_or_add_pPr()
    sp0 = make_element("w:spacing", {
        "w:line": "360",
        "w:lineRule": "exact",
    })
    pPr0.append(sp0)

    save_fixture("spacing-line-exact", doc, {
        "name": "spacing-line-exact",
        "spec_ref": "ECMA-376 §17.3.1.33",
        "description": "line=360 lineRule=exact is 18pt exact line height",
        "expected_behavior": "P0: line=360, lineRule=Exact.",
    })


# =========================================================================
# Main
# =========================================================================

def main() -> None:
    print("Generating spacing-contextual fixtures...")
    make_contextual_same_style()
    make_contextual_explicit_false()
    make_contextual_false_roundtrip()
    make_before_autospacing()
    make_after_autospacing()
    make_spacing_before_after_basic()
    make_spacing_line_auto()
    make_spacing_line_exact()
    print("Done.")


if __name__ == "__main__":
    main()

# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for spacing/indentation spec-compliance tests.

Tests target:
  - ECMA-376 §17.3.1.12: firstLine/hanging precedence
  - ECMA-376 §17.3.1.33: lineRule default when omitted

Run:  uv run create_docs.py
"""

import json
from pathlib import Path
from lxml import etree

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  spacing-indentation/{name}/")


# ── Fixture 1: style with both hanging and firstLine ─────────────────────

def make_hanging_wins_firstline_style() -> None:
    """Style defines ind with BOTH hanging="360" AND firstLine="720".

    ECMA-376 §17.3.1.12: "The firstLine and hanging attributes are mutually
    exclusive, if both are specified, then the firstLine value is ignored."
    hanging wins → effective_first_line_twips = -360.
    """
    doc = Document()

    styles_element = doc.styles.element
    style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "HangingWins"})
    style.append(make_element("w:name", {"w:val": "Hanging Wins"}))
    pPr = make_element("w:pPr")
    # Both hanging and firstLine on the same ind element
    ind = make_element("w:ind", {
        "w:left": "720",
        "w:hanging": "360",
        "w:firstLine": "720",
    })
    pPr.append(ind)
    style.append(pPr)
    styles_element.append(style)

    p = doc.add_paragraph(style="HangingWins")
    p.add_run("Style has both hanging=360 and firstLine=720. Hanging should win.")

    save_fixture("hanging-wins-style", doc, {
        "name": "hanging-wins-style",
        "spec_ref": "ECMA-376 §17.3.1.12",
        "description": "Style defines ind with both hanging=360 and firstLine=720; hanging should win per spec",
        "expected_behavior": "effective_first_line_twips = -360 (hanging wins, firstLine ignored)",
        "current_status": "BUG — style parser checks firstLine first, so firstLine wins incorrectly",
    })


# ── Fixture 2: numbering with both hanging and firstLine ─────────────────

def make_hanging_wins_firstline_numbering() -> None:
    """Numbering level defines ind with BOTH hanging="360" AND firstLine="720".

    Same spec rule as fixture 1 but through numbering.
    """
    doc = Document()

    # Create numbering part with an abstract numbering definition
    # We need to add numbering.xml with a level that has both hanging and firstLine
    numbering_part = doc.part.numbering_part
    numbering_element = numbering_part.numbering_definitions._numbering

    # Create abstract numbering
    abstract_num = make_element("w:abstractNum", {"w:abstractNumId": "100"})
    lvl = make_element("w:lvl", {"w:ilvl": "0"})
    lvl.append(make_element("w:start", {"w:val": "1"}))
    lvl.append(make_element("w:numFmt", {"w:val": "decimal"}))
    lvl.append(make_element("w:lvlText", {"w:val": "%1."}))
    lvl.append(make_element("w:lvlJc", {"w:val": "left"}))
    lvl_pPr = make_element("w:pPr")
    # Both hanging and firstLine on the numbering level's ind
    lvl_ind = make_element("w:ind", {
        "w:left": "720",
        "w:hanging": "360",
        "w:firstLine": "720",
    })
    lvl_pPr.append(lvl_ind)
    lvl.append(lvl_pPr)
    abstract_num.append(lvl)

    # Insert abstractNum before any existing num elements
    first_num = numbering_element.find(qn("w:num"))
    if first_num is not None:
        first_num.addprevious(abstract_num)
    else:
        numbering_element.append(abstract_num)

    # Create concrete num referencing our abstract
    num = make_element("w:num", {"w:numId": "100"})
    num.append(make_element("w:abstractNumId", {"w:val": "100"}))
    numbering_element.append(num)

    # Create paragraph using this numbering
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    numPr = make_element("w:numPr")
    numPr.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr.append(make_element("w:numId", {"w:val": "100"}))
    pPr.append(numPr)
    p.add_run("Numbering has both hanging=360 and firstLine=720. Hanging should win.")

    save_fixture("hanging-wins-numbering", doc, {
        "name": "hanging-wins-numbering",
        "spec_ref": "ECMA-376 §17.3.1.12",
        "description": "Numbering level defines ind with both hanging=360 and firstLine=720; hanging should win",
        "expected_behavior": "effective_first_line_twips = -360 (hanging wins, firstLine ignored)",
        "current_status": "BUG — numbering parser checks firstLine first, so firstLine wins incorrectly",
    })


# ── Fixture 3: direct formatting with both hanging and firstLine ─────────

def make_hanging_wins_firstline_direct() -> None:
    """Direct paragraph ind with BOTH hanging="360" AND firstLine="720".

    Control test: word_ir.rs already handles this correctly (hanging wins).
    """
    doc = Document()

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    ind = make_element("w:ind", {
        "w:left": "720",
        "w:hanging": "360",
        "w:firstLine": "720",
    })
    pPr.append(ind)
    p.add_run("Direct formatting: hanging=360 and firstLine=720. Hanging should win.")

    save_fixture("hanging-wins-direct", doc, {
        "name": "hanging-wins-direct",
        "spec_ref": "ECMA-376 §17.3.1.12",
        "description": "Direct ind with both hanging=360 and firstLine=720; hanging should win",
        "expected_behavior": "effective_first_line_twips = -360 (hanging wins)",
        "current_status": "PASS — word_ir.rs parser already checks hanging first",
    })


# ── Fixture 4: style with line but no lineRule ───────────────────────────

def make_linerule_defaults_auto_style() -> None:
    """Style defines spacing line="276" but NO lineRule.

    ECMA-376 §17.3.1.33: "If this attribute [lineRule] is omitted, then it
    shall be assumed to be of a value auto if a line attribute value is present."

    We also strip lineRule from the docDefaults pPrDefault spacing so the
    bug is exposed (otherwise python-docx's default pPrDefault has
    lineRule="auto" which masks the missing default in the style-only path).
    """
    doc = Document()

    styles_element = doc.styles.element

    # Strip lineRule from any existing docDefaults pPrDefault spacing so it
    # doesn't mask the bug via the defaults fallback.
    for dd_spacing in styles_element.findall(
        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}docDefaults"
        "//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing"
    ):
        lr_attr = qn("w:lineRule")
        if lr_attr in dd_spacing.attrib:
            del dd_spacing.attrib[lr_attr]

    style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "LineNoRule"})
    style.append(make_element("w:name", {"w:val": "Line No Rule"}))
    pPr = make_element("w:pPr")
    # line="276" but deliberately NO lineRule attribute
    sp = make_element("w:spacing", {"w:line": "276"})
    pPr.append(sp)
    style.append(pPr)
    styles_element.append(style)

    p = doc.add_paragraph(style="LineNoRule")
    p.add_run("Style has line=276 but no lineRule. Should default to auto.")

    save_fixture("linerule-default-style", doc, {
        "name": "linerule-default-style",
        "spec_ref": "ECMA-376 §17.3.1.33",
        "description": "Style defines spacing line=276 with no lineRule; should default to Auto",
        "expected_behavior": "line_rule = Some(LineSpacingRule::Auto)",
        "current_status": "BUG — style-only path does not apply lineRule default",
    })


# ── Fixture 5: direct formatting with line but no lineRule ───────────────

def make_linerule_defaults_auto_direct() -> None:
    """Direct spacing line="276" with NO lineRule.

    Control test: word_ir.rs already handles this correctly.
    """
    doc = Document()

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    sp = make_element("w:spacing", {"w:line": "276"})
    pPr.append(sp)
    p.add_run("Direct spacing: line=276, no lineRule. Should default to auto.")

    save_fixture("linerule-default-direct", doc, {
        "name": "linerule-default-direct",
        "spec_ref": "ECMA-376 §17.3.1.33",
        "description": "Direct spacing line=276 with no lineRule; should default to Auto",
        "expected_behavior": "line_rule = Some(LineSpacingRule::Auto)",
        "current_status": "PASS — word_ir.rs parser defaults lineRule to auto when line is present",
    })


# ── Fixture 6: basic left/right indentation ──────────────────────────────

def make_indent_left_right() -> None:
    """Basic paragraph with left=720 right=360 indentation."""
    doc = Document()

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    ind = make_element("w:ind", {"w:left": "720", "w:right": "360"})
    pPr.append(ind)
    p.add_run("Left=720, right=360 indentation.")

    save_fixture("indent-left-right", doc, {
        "name": "indent-left-right",
        "spec_ref": "ECMA-376 §17.3.1.12",
        "description": "Paragraph with ind left=720 right=360",
        "expected_behavior": "left=720, right=360 on domain Indentation",
    })


# ── Fixture 7: negative start indentation ────────────────────────────────

def make_indent_negative_start() -> None:
    """Paragraph with negative start indent (extends past margin)."""
    doc = Document()

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    ind = make_element("w:ind", {"w:start": "-720"})
    pPr.append(ind)
    p.add_run("Negative start=-720 extends past margin.")

    save_fixture("indent-negative-start", doc, {
        "name": "indent-negative-start",
        "spec_ref": "ECMA-376 §17.3.1.12",
        "description": "Paragraph with ind start=-720 (negative extends past margin)",
        "expected_behavior": "left=-720 on domain Indentation",
    })


# ── Main ─────────────────────────────────────────────────────────────────

def main() -> None:
    print("Generating spacing-indentation fixtures:")
    make_hanging_wins_firstline_style()
    make_hanging_wins_firstline_numbering()
    make_hanging_wins_firstline_direct()
    make_linerule_defaults_auto_style()
    make_linerule_defaults_auto_direct()
    make_indent_left_right()
    make_indent_negative_start()
    print("Done.")


if __name__ == "__main__":
    main()

# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for style cascade bug-exposure tests.

Targets four known behavioral bugs in style resolution:
  1. Toggle-docDefaults interaction (§17.7.3 + MS-OI29500 §2.1.230a)
  2. Linked style property loss (§17.7.4.6)
  3. Table conditional formatting precedence (MS-OI29500 §17.4.54a)
  4. lineRule default when direct spacing omits it (§17.3.1.33)

Run:  python create_fixtures.py
"""

import json
from pathlib import Path
from lxml import etree

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  style-cascade-bugs/{name}/")


# =========================================================================
# Bug 1: Toggle-docDefaults interaction
#
# §17.7.3 + MS-OI29500 §2.1.230a: When docDefaults sets bold=on, a
# paragraph style with bold=off should be able to override it.
#
# Scenario A: Simple case — docDefaults bold=on, para style bold=off.
# Scenario B: Char style bold=off should also override docDefaults bold=on.
# Scenario C: docDefaults italic=on, para style italic=off, char style
#   italic=on — tests multi-level toggle interaction with docDefaults.
# =========================================================================

def make_bug1_toggle_doc_defaults():
    """
    docDefaults rPr: bold=on, italic=on
    StyleNoBold: paragraph style, bold=off (explicit), italic not set
    CharNoItalic: character style, italic=off (explicit)
    CharBoldOn: character style, bold=on (explicit)

    Para 1: StyleNoBold, no char style
      → bold=OFF (style overrides docDefaults), italic=ON (from docDefaults)
    Para 2: no para style, CharNoItalic
      → bold=ON (from docDefaults), italic=OFF (char overrides docDefaults)
    Para 3: StyleNoBold, CharBoldOn
      → bold: style says OFF, char says ON — char overrides style → ON
         italic=ON (from docDefaults, nothing overrides)
    Para 4: no styles at all
      → bold=ON (docDefaults), italic=ON (docDefaults)
    """
    doc = Document()
    styles_el = doc.styles.element

    # Set docDefaults rPr: bold=on, italic=on
    doc_defaults = styles_el.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = make_element("w:docDefaults")
        styles_el.insert(0, doc_defaults)
    rpr_default_el = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default_el is None:
        rpr_default_el = make_element("w:rPrDefault")
        doc_defaults.insert(0, rpr_default_el)
    rpr = rpr_default_el.find(qn("w:rPr"))
    if rpr is None:
        rpr = make_element("w:rPr")
        rpr_default_el.append(rpr)
    for existing_b in rpr.findall(qn("w:b")):
        rpr.remove(existing_b)
    rpr.append(make_element("w:b"))
    for existing_i in rpr.findall(qn("w:i")):
        rpr.remove(existing_i)
    rpr.append(make_element("w:i"))

    # StyleNoBold: paragraph style, bold=off
    style_nobold = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "StyleNoBold",
    })
    style_nobold.append(make_element("w:name", {"w:val": "Style No Bold"}))
    rpr2 = make_element("w:rPr")
    rpr2.append(make_element("w:b", {"w:val": "0"}))
    style_nobold.append(rpr2)
    styles_el.append(style_nobold)

    # CharNoItalic: character style, italic=off
    style_noitalic = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "CharNoItalic",
    })
    style_noitalic.append(make_element("w:name", {"w:val": "Char No Italic"}))
    rpr3 = make_element("w:rPr")
    rpr3.append(make_element("w:i", {"w:val": "0"}))
    style_noitalic.append(rpr3)
    styles_el.append(style_noitalic)

    # CharBoldOn: character style, bold=on
    style_boldon = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "CharBoldOn",
    })
    style_boldon.append(make_element("w:name", {"w:val": "Char Bold On"}))
    rpr4 = make_element("w:rPr")
    rpr4.append(make_element("w:b"))
    style_boldon.append(rpr4)
    styles_el.append(style_boldon)

    # Para 1: StyleNoBold, no char style
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "StyleNoBold"}))
    p1.add_run("Para style bold off")

    # Para 2: no para style, CharNoItalic
    p2 = doc.add_paragraph()
    # Remove any auto pStyle
    p2_ppr = p2._element.find(qn("w:pPr"))
    if p2_ppr is not None:
        ps = p2_ppr.find(qn("w:pStyle"))
        if ps is not None:
            p2_ppr.remove(ps)
    r2 = p2.add_run("Char style italic off")
    r2_rpr = r2._element.get_or_add_rPr()
    r2_rpr.insert(0, make_element("w:rStyle", {"w:val": "CharNoItalic"}))

    # Para 3: StyleNoBold + CharBoldOn
    p3 = doc.add_paragraph()
    p3_ppr = p3._element.get_or_add_pPr()
    p3_ppr.append(make_element("w:pStyle", {"w:val": "StyleNoBold"}))
    r3 = p3.add_run("Para off char on")
    r3_rpr = r3._element.get_or_add_rPr()
    r3_rpr.insert(0, make_element("w:rStyle", {"w:val": "CharBoldOn"}))

    # Para 4: no styles
    p4 = doc.add_paragraph()
    p4_ppr = p4._element.find(qn("w:pPr"))
    if p4_ppr is not None:
        ps = p4_ppr.find(qn("w:pStyle"))
        if ps is not None:
            p4_ppr.remove(ps)
    p4.add_run("No styles at all")

    save_fixture("bug1-toggle-doc-defaults", doc, {
        "spec": "ISO 29500-1 §17.7.3 + MS-OI29500 §2.1.230a",
        "bug": "Toggle-docDefaults interaction",
        "description": (
            "docDefaults bold=on,italic=on with styles overriding individual toggles. "
            "Tests that styles can turn OFF toggles that docDefaults turned ON."
        ),
        "expected": {
            "para_1": "bold=OFF (style overrides), italic=ON (docDefaults)",
            "para_2": "bold=ON (docDefaults), italic=OFF (char overrides)",
            "para_3": "bold=ON (char overrides style's off), italic=ON (docDefaults)",
            "para_4": "bold=ON (docDefaults), italic=ON (docDefaults)",
        },
    })


# =========================================================================
# Bug 2: Linked style property loss
#
# §17.7.4.6: When a paragraph style is linked to a character style, runs
# inherit rPr from the linked char style. Properties NOT set by the char
# style should fall through to the paragraph style's own resolved rPr.
#
# Scenario: HeadingPara has sz=32, color=FF0000 in rPr.
#   HeadingChar (linked) has bold=on only.
#   Runs in HeadingPara should get bold=on (from char), sz=32 AND
#   color=FF0000 (from para, since char doesn't set them).
#
# Additional scenario: char style based on another char style that sets
#   a different property. Tests that basedOn chain in char style doesn't
#   incorrectly override para style properties.
# =========================================================================

def make_bug2_linked_style():
    """
    HeadingChar: character style, bold=on only, basedOn BaseChar
    BaseChar: character style, font_size=24 (different from para's 32!)
    HeadingPara: paragraph style, rPr: bold=on, sz=32, color=FF0000
      linked to HeadingChar

    Para 1: HeadingPara style, no explicit char style
      → bold=ON, font_size should be 32 (from para rPr)
        BUT if linked char's resolved chain includes BaseChar's sz=24,
        overlay_marks will overwrite para's 32 with 24 — BUG
    Para 2: HeadingPara style, explicit HeadingChar char style
      → same interaction but char style explicitly applied
    Para 3: no para style, HeadingChar char style
      → bold=ON (from HeadingChar), font_size=24 (from BaseChar via basedOn)
    """
    doc = Document()
    styles_el = doc.styles.element

    # BaseChar: character style with font_size=24
    style_base_char = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "BaseChar",
    })
    style_base_char.append(make_element("w:name", {"w:val": "Base Char"}))
    base_rpr = make_element("w:rPr")
    base_rpr.append(make_element("w:sz", {"w:val": "24"}))
    style_base_char.append(base_rpr)
    styles_el.append(style_base_char)

    # HeadingChar: character style, basedOn BaseChar, bold=on only (no sz)
    style_char = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "HeadingChar",
    })
    style_char.append(make_element("w:name", {"w:val": "Heading Char"}))
    style_char.append(make_element("w:basedOn", {"w:val": "BaseChar"}))
    style_char.append(make_element("w:link", {"w:val": "HeadingPara"}))
    char_rpr = make_element("w:rPr")
    char_rpr.append(make_element("w:b"))
    style_char.append(char_rpr)
    styles_el.append(style_char)

    # HeadingPara: paragraph style, linked to HeadingChar
    # rPr: bold=on, sz=32, color=FF0000
    style_para = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "HeadingPara",
    })
    style_para.append(make_element("w:name", {"w:val": "Heading Para"}))
    style_para.append(make_element("w:link", {"w:val": "HeadingChar"}))
    para_rpr = make_element("w:rPr")
    para_rpr.append(make_element("w:b"))
    para_rpr.append(make_element("w:sz", {"w:val": "32"}))
    para_rpr.append(make_element("w:color", {"w:val": "FF0000"}))
    style_para.append(para_rpr)
    styles_el.append(style_para)

    # Para 1: HeadingPara, no char style on run
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "HeadingPara"}))
    p1.add_run("Heading linked no char")

    # Para 2: HeadingPara + explicit HeadingChar on run
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    p2_ppr.append(make_element("w:pStyle", {"w:val": "HeadingPara"}))
    r2 = p2.add_run("Heading linked explicit char")
    r2_rpr = r2._element.get_or_add_rPr()
    r2_rpr.insert(0, make_element("w:rStyle", {"w:val": "HeadingChar"}))

    # Para 3: no para style, HeadingChar on run
    p3 = doc.add_paragraph()
    p3_ppr = p3._element.find(qn("w:pPr"))
    if p3_ppr is not None:
        ps = p3_ppr.find(qn("w:pStyle"))
        if ps is not None:
            p3_ppr.remove(ps)
    r3 = p3.add_run("Char style only")
    r3_rpr = r3._element.get_or_add_rPr()
    r3_rpr.insert(0, make_element("w:rStyle", {"w:val": "HeadingChar"}))

    save_fixture("bug2-linked-style-loss", doc, {
        "spec": "ISO 29500-1 §17.7.4.6",
        "bug": "Linked style property loss via basedOn inheritance",
        "description": (
            "HeadingPara (sz=32) linked to HeadingChar (basedOn BaseChar sz=24, bold=on). "
            "Runs in HeadingPara should get sz=32 from para rPr, not sz=24 from "
            "char style's basedOn chain. The linked char style's inherited sz=24 "
            "should NOT override the para style's explicit sz=32."
        ),
        "expected": {
            "para_1": "bold=ON, font_size=32 (para rPr wins over linked char inherited sz=24)",
            "para_2": "bold=ON, font_size=24 (explicit HeadingChar: resolved chain has sz=24)",
            "para_3": "bold=ON, font_size=24 (HeadingChar basedOn BaseChar)",
        },
    })


# =========================================================================
# Bug 3: Table conditional formatting precedence
#
# MS-OI29500 §17.4.54a + §17.7.6(c):
# Precedence order (low to high):
#   wholeTable < bands < firstCol/lastCol < firstRow/lastRow < corners
#
# The issue: conditional formatting uses HashMap for condition→properties.
# While application code uses explicit precedence ordering (matching vector),
# the style resolution merges conditionals from basedOn chains. If a parent
# style and child style both define the same condition type with different
# property subsets, the child wins entirely (no per-property merge within
# a condition type). This can lose parent properties.
#
# Scenario: Table style with firstRow (shading=blue, font_size=28) and
#   nwCell (shading=green, no font_size). The NW corner cell should get
#   shading=green (from nwCell, highest precedence) and font_size=28
#   (from firstRow, since nwCell doesn't set it).
# =========================================================================

def make_bug3_table_conditional():
    """
    Table style CondTestStyle:
      - firstRow: shading fill=4472C4 (blue), bold=true, font_size=28
      - firstCol: shading fill=70AD47 (green), bold=false
      - nwCell: shading fill=FFC000 (amber) — no bold, no font_size
      - lastRow: shading fill=ED7D31 (orange)
      - seCell: shading fill=FF0000 (red)

    Table: 3x3, tblLook enables firstRow, firstCol, lastRow, lastCol.

    Expected per MS-OI29500 §17.7.6(c) precedence:
      Cell (0,0): nwCell > firstRow > firstCol → shading=FFC000 (amber)
        bold from firstRow (since nwCell doesn't set it)
        font_size from firstRow (since nwCell doesn't set it)
      Cell (0,2): firstRow + lastCol → shading=4472C4 (blue, firstRow > lastCol), bold=true
      Cell (2,0): swCell not defined, so lastRow + firstCol → shading=ED7D31 (lastRow > firstCol)
      Cell (2,2): seCell > lastRow + lastCol → shading=FF0000 (red)
      Cell (1,0): firstCol → shading=70AD47 (green)
      Cell (1,1): no conditional → no shading
    """
    doc = Document()
    styles_el = doc.styles.element

    # Create the table style
    style = make_element("w:style", {
        "w:type": "table",
        "w:styleId": "CondTestStyle",
    })
    style.append(make_element("w:name", {"w:val": "Cond Test Style"}))

    # tblPr (required for table styles)
    tbl_pr = make_element("w:tblPr")
    style.append(tbl_pr)

    # firstRow conditional: blue shading, bold, font_size=28
    first_row = make_element("w:tblStylePr", {"w:type": "firstRow"})
    fr_tc_pr = make_element("w:tcPr")
    fr_shd = make_element("w:shd", {"w:val": "clear", "w:color": "auto", "w:fill": "4472C4"})
    fr_tc_pr.append(fr_shd)
    first_row.append(fr_tc_pr)
    fr_rpr = make_element("w:rPr")
    fr_rpr.append(make_element("w:b"))
    fr_rpr.append(make_element("w:sz", {"w:val": "28"}))
    first_row.append(fr_rpr)
    style.append(first_row)

    # firstCol conditional: green shading, bold=off
    first_col = make_element("w:tblStylePr", {"w:type": "firstCol"})
    fc_tc_pr = make_element("w:tcPr")
    fc_shd = make_element("w:shd", {"w:val": "clear", "w:color": "auto", "w:fill": "70AD47"})
    fc_tc_pr.append(fc_shd)
    first_col.append(fc_tc_pr)
    fc_rpr = make_element("w:rPr")
    fc_rpr.append(make_element("w:b", {"w:val": "0"}))
    first_col.append(fc_rpr)
    style.append(first_col)

    # nwCell conditional: amber shading only (no bold, no font_size)
    nw_cell = make_element("w:tblStylePr", {"w:type": "nwCell"})
    nw_tc_pr = make_element("w:tcPr")
    nw_shd = make_element("w:shd", {"w:val": "clear", "w:color": "auto", "w:fill": "FFC000"})
    nw_tc_pr.append(nw_shd)
    nw_cell.append(nw_tc_pr)
    style.append(nw_cell)

    # lastRow conditional: orange shading
    last_row = make_element("w:tblStylePr", {"w:type": "lastRow"})
    lr_tc_pr = make_element("w:tcPr")
    lr_shd = make_element("w:shd", {"w:val": "clear", "w:color": "auto", "w:fill": "ED7D31"})
    lr_tc_pr.append(lr_shd)
    last_row.append(lr_tc_pr)
    style.append(last_row)

    # seCell conditional: red shading
    se_cell = make_element("w:tblStylePr", {"w:type": "seCell"})
    se_tc_pr = make_element("w:tcPr")
    se_shd = make_element("w:shd", {"w:val": "clear", "w:color": "auto", "w:fill": "FF0000"})
    se_tc_pr.append(se_shd)
    se_cell.append(se_tc_pr)
    style.append(se_cell)

    styles_el.append(style)

    # Create a 3x3 table
    table = doc.add_table(rows=3, cols=3)

    # Apply our custom table style
    tbl_el = table._element
    tbl_pr_el = tbl_el.find(qn("w:tblPr"))
    if tbl_pr_el is None:
        tbl_pr_el = make_element("w:tblPr")
        tbl_el.insert(0, tbl_pr_el)

    # Remove any existing style ref
    for existing in tbl_pr_el.findall(qn("w:tblStyle")):
        tbl_pr_el.remove(existing)
    tbl_pr_el.insert(0, make_element("w:tblStyle", {"w:val": "CondTestStyle"}))

    # Set tblLook: enable firstRow, firstCol, lastRow; disable banding
    for existing in tbl_pr_el.findall(qn("w:tblLook")):
        tbl_pr_el.remove(existing)
    tbl_look = make_element("w:tblLook", {
        "w:val": "04E0",
        "w:firstRow": "1",
        "w:lastRow": "1",
        "w:firstColumn": "1",
        "w:lastColumn": "1",
        "w:noHBand": "1",
        "w:noVBand": "1",
    })
    tbl_pr_el.append(tbl_look)

    # Fill cells with identifying text
    labels = [
        ["NW corner", "Top mid", "Top right"],
        ["Mid left", "Center", "Mid right"],
        ["Bot left", "Bot mid", "Bot right"],
    ]
    for r in range(3):
        for c in range(3):
            cell = table.cell(r, c)
            cell.text = labels[r][c]

    save_fixture("bug3-table-conditional-precedence", doc, {
        "spec": "MS-OI29500 §17.4.54a + §17.7.6(c)",
        "bug": "Table conditional formatting precedence",
        "description": (
            "Table style with firstRow, firstCol, lastRow, nwCell, seCell. "
            "Tests that corner cells get highest precedence and that "
            "per-property fill-from-lower-precedence works correctly."
        ),
        "expected": {
            "cell_0_0": "shading=FFC000 (nwCell > firstRow > firstCol), bold from firstRow",
            "cell_0_2": "shading=4472C4 (firstRow), bold=true, font_size=28",
            "cell_2_0": "shading=70AD47 (firstCol, no swCell defined)",
            "cell_2_2": "shading=FF0000 (seCell > lastRow)",
            "cell_1_0": "shading=70AD47 (firstCol)",
            "cell_1_1": "no conditional shading",
        },
    })


# =========================================================================
# Bug 4: lineRule default (GAP-016)
#
# §17.3.1.33: "If [lineRule] is omitted, then it shall be assumed to be
# of a value auto if a line attribute value is present."
#
# When direct w:spacing sets w:line but omits w:lineRule, the lineRule
# should default to "auto", not inherit from the style's lineRule.
#
# Scenario: Style has lineRule=exact with line=360. Direct spacing sets
# line=480 but omits lineRule. The effective lineRule should be auto (not
# exact from the style).
# =========================================================================

def make_bug4_line_rule_default():
    """
    ExactLineStyle: paragraph style, spacing line=360 lineRule=exact
    AtLeastStyle: paragraph style, spacing line=300 lineRule=atLeast

    Para 1: ExactLineStyle, direct spacing line=480 (no lineRule)
      → lineRule should be auto (§17.3.1.33 default), NOT exact from style
    Para 2: AtLeastStyle, direct spacing line=480 (no lineRule)
      → lineRule should be auto, NOT atLeast from style
    Para 3: ExactLineStyle, no direct spacing
      → lineRule should be exact (from style), line=360
    Para 4: no style, direct spacing line=480 (no lineRule)
      → lineRule should be auto
    Para 5: no style, direct spacing line=480 lineRule=exact
      → lineRule should be exact (explicitly set in direct)
    """
    doc = Document()
    styles_el = doc.styles.element

    # ExactLineStyle: spacing line=360, lineRule=exact
    style_exact = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "ExactLineStyle",
    })
    style_exact.append(make_element("w:name", {"w:val": "Exact Line Style"}))
    ppr_exact = make_element("w:pPr")
    ppr_exact.append(make_element("w:spacing", {
        "w:line": "360",
        "w:lineRule": "exact",
    }))
    style_exact.append(ppr_exact)
    styles_el.append(style_exact)

    # AtLeastStyle: spacing line=300, lineRule=atLeast
    style_atleast = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "AtLeastStyle",
    })
    style_atleast.append(make_element("w:name", {"w:val": "At Least Style"}))
    ppr_atleast = make_element("w:pPr")
    ppr_atleast.append(make_element("w:spacing", {
        "w:line": "300",
        "w:lineRule": "atLeast",
    }))
    style_atleast.append(ppr_atleast)
    styles_el.append(style_atleast)

    # Para 1: ExactLineStyle + direct line=480 (no lineRule)
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "ExactLineStyle"}))
    p1_ppr.append(make_element("w:spacing", {"w:line": "480"}))
    p1.add_run("Exact style direct line only")

    # Para 2: AtLeastStyle + direct line=480 (no lineRule)
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    p2_ppr.append(make_element("w:pStyle", {"w:val": "AtLeastStyle"}))
    p2_ppr.append(make_element("w:spacing", {"w:line": "480"}))
    p2.add_run("AtLeast style direct line only")

    # Para 3: ExactLineStyle, no direct spacing
    p3 = doc.add_paragraph()
    p3_ppr = p3._element.get_or_add_pPr()
    p3_ppr.append(make_element("w:pStyle", {"w:val": "ExactLineStyle"}))
    p3.add_run("Exact style no direct")

    # Para 4: no style, direct line=480 (no lineRule)
    p4 = doc.add_paragraph()
    p4_ppr = p4._element.get_or_add_pPr()
    # Remove any auto pStyle
    ps = p4_ppr.find(qn("w:pStyle"))
    if ps is not None:
        p4_ppr.remove(ps)
    p4_ppr.append(make_element("w:spacing", {"w:line": "480"}))
    p4.add_run("No style direct line only")

    # Para 5: no style, direct line=480 + lineRule=exact
    p5 = doc.add_paragraph()
    p5_ppr = p5._element.get_or_add_pPr()
    ps = p5_ppr.find(qn("w:pStyle"))
    if ps is not None:
        p5_ppr.remove(ps)
    p5_ppr.append(make_element("w:spacing", {
        "w:line": "480",
        "w:lineRule": "exact",
    }))
    p5.add_run("No style direct exact")

    save_fixture("bug4-line-rule-default", doc, {
        "spec": "ISO 29500-1 §17.3.1.33",
        "bug": "lineRule default when direct spacing omits it",
        "description": (
            "Style has lineRule=exact. Direct w:spacing sets line but omits lineRule. "
            "The lineRule should default to auto per §17.3.1.33, not inherit "
            "from the style."
        ),
        "expected": {
            "para_1": "line=480, lineRule=auto (default, NOT exact from style)",
            "para_2": "line=480, lineRule=auto (default, NOT atLeast from style)",
            "para_3": "line=360, lineRule=exact (from style, no direct override)",
            "para_4": "line=480, lineRule=auto (default, no style)",
            "para_5": "line=480, lineRule=exact (explicitly set in direct)",
        },
    })


# =========================================================================
# Main
# =========================================================================

if __name__ == "__main__":
    print("Generating style-cascade-bugs fixtures:")
    make_bug1_toggle_doc_defaults()
    make_bug2_linked_style()
    make_bug3_table_conditional()
    make_bug4_line_rule_default()
    print("Done.")

# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for insideH/insideV border cascade edge cases
and cell spacing interaction.

Spec refs:
  - ISO 29500-1 §17.4.22-25 (insideH/insideV mapping)
  - ISO 29500-1 §17.4.38, §17.4.39 (tblBorders, tblPrEx)
  - ISO 29500-1 §17.4.66 (three-tier fallback)
  - MS-OI29500 §2.1.136, §2.1.138, §2.1.150 (cell spacing interaction)

Run:  cd stemma-engine/testdata/spec-compliance/table-border-inside-cascade && mise exec -- uv run create_docs.py
"""

import json
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc: DocxDocument, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  table-border-inside-cascade/{name}/")


def make_table_with_borders(doc, rows=2, cols=2, cell_texts=None):
    """Create a table, return (tbl_element, tblPr_element)."""
    tbl = doc.add_table(rows=rows, cols=cols)
    if cell_texts:
        for r in range(rows):
            for c in range(cols):
                if r < len(cell_texts) and c < len(cell_texts[r]):
                    tbl.cell(r, c).text = cell_texts[r][c]
    tbl_el = tbl._tbl
    tblPr = tbl_el.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_el.insert(0, tblPr)
    return tbl_el, tblPr


def add_distinct_tbl_borders(tblPr, border_specs):
    """Add tblBorders to tblPr with per-edge specs.

    border_specs: list of (edge, val, sz, color) tuples.
    """
    tbl_borders = make_element("w:tblBorders")
    for edge, val, sz, color in border_specs:
        border = make_element(f"w:{edge}", {
            "w:val": val, "w:sz": sz, "w:color": color, "w:space": "0",
        })
        tbl_borders.append(border)
    tblPr.append(tbl_borders)


def add_tc_borders(tbl_el, row, col, edges_spec):
    """Add tcBorders to a specific cell. edges_spec is a list of
    (edge_name, val, sz, color, space) tuples."""
    tr = tbl_el.findall(w("tr"))[row]
    tc = tr.findall(w("tc"))[col]
    tcPr = tc.find(w("tcPr"))
    if tcPr is None:
        tcPr = make_element("w:tcPr")
        tc.insert(0, tcPr)
    tc_borders = make_element("w:tcBorders")
    for edge, val, sz, color, space in edges_spec:
        border = make_element(f"w:{edge}", {
            "w:val": val, "w:sz": sz, "w:color": color, "w:space": space,
        })
        tc_borders.append(border)
    tcPr.append(tc_borders)


# =====================================================================
# Fixture 1: inside-h-mapping-3x3
# =====================================================================

def make_inside_h_mapping_3x3():
    """3x3 table with distinct table-level borders on all 6 edges.

    No cell-level borders. No tblPrEx.

    top     = single, sz=8, color="FF0000" (red)
    bottom  = single, sz=8, color="0000FF" (blue)
    left    = single, sz=8, color="00FF00" (green)
    right   = single, sz=8, color="FFFF00" (yellow)
    insideH = single, sz=8, color="FF00FF" (magenta)
    insideV = single, sz=8, color="00FFFF" (cyan)
    """
    doc = Document()
    doc.add_paragraph("Fixture 1: insideH/insideV mapping on a 3x3 table.")

    tbl_el, tblPr = make_table_with_borders(doc, 3, 3, [
        ["R0C0", "R0C1", "R0C2"],
        ["R1C0", "R1C1", "R1C2"],
        ["R2C0", "R2C1", "R2C2"],
    ])

    add_distinct_tbl_borders(tblPr, [
        ("top",     "single", "8", "FF0000"),  # red
        ("left",    "single", "8", "00FF00"),  # green
        ("bottom",  "single", "8", "0000FF"),  # blue
        ("right",   "single", "8", "FFFF00"),  # yellow
        ("insideH", "single", "8", "FF00FF"),  # magenta
        ("insideV", "single", "8", "00FFFF"),  # cyan
    ])

    save_fixture("inside-h-mapping-3x3", doc, {
        "name": "inside-h-mapping-3x3",
        "spec_ref": "ISO 29500-1 SS17.4.22, SS17.4.24, SS17.4.38",
        "description": (
            "3x3 table with distinct table-level borders: "
            "top=red, bottom=blue, left=green, right=yellow, "
            "insideH=magenta, insideV=cyan. "
            "No cell-level borders. Tests insideH/insideV mapping to cell edges."
        ),
        "expected_borders": {
            "(0,0)": {"top": "FF0000", "bottom": "FF00FF", "left": "00FF00", "right": "00FFFF"},
            "(0,1)": {"top": "FF0000", "bottom": "FF00FF", "left": "00FFFF", "right": "00FFFF"},
            "(0,2)": {"top": "FF0000", "bottom": "FF00FF", "left": "00FFFF", "right": "FFFF00"},
            "(1,0)": {"top": "FF00FF", "bottom": "FF00FF", "left": "00FF00", "right": "00FFFF"},
            "(1,1)": {"top": "FF00FF", "bottom": "FF00FF", "left": "00FFFF", "right": "00FFFF"},
            "(1,2)": {"top": "FF00FF", "bottom": "FF00FF", "left": "00FFFF", "right": "FFFF00"},
            "(2,0)": {"top": "FF00FF", "bottom": "0000FF", "left": "00FF00", "right": "00FFFF"},
            "(2,1)": {"top": "FF00FF", "bottom": "0000FF", "left": "00FFFF", "right": "00FFFF"},
            "(2,2)": {"top": "FF00FF", "bottom": "0000FF", "left": "00FFFF", "right": "FFFF00"},
        },
    })


# =====================================================================
# Fixture 2: inside-h-only-no-inside-v
# =====================================================================

def make_inside_h_only_no_inside_v():
    """2x2 table with insideH but NO insideV.

    top     = single, sz=8, color="FF0000"
    bottom  = single, sz=8, color="FF0000"
    left    = single, sz=8, color="FF0000"
    right   = single, sz=8, color="FF0000"
    insideH = dashed, sz=4, color="00FF00"
    (no insideV)

    Interior vertical edges should have NO border.
    """
    doc = Document()
    doc.add_paragraph("Fixture 2: insideH only, no insideV.")

    tbl_el, tblPr = make_table_with_borders(doc, 2, 2, [
        ["R0C0", "R0C1"],
        ["R1C0", "R1C1"],
    ])

    add_distinct_tbl_borders(tblPr, [
        ("top",     "single", "8", "FF0000"),
        ("left",    "single", "8", "FF0000"),
        ("bottom",  "single", "8", "FF0000"),
        ("right",   "single", "8", "FF0000"),
        ("insideH", "dashed", "4", "00FF00"),
        # NO insideV
    ])

    save_fixture("inside-h-only-no-inside-v", doc, {
        "name": "inside-h-only-no-inside-v",
        "spec_ref": "ISO 29500-1 SS17.4.38",
        "description": (
            "2x2 table with insideH (green dashed sz=4) but NO insideV. "
            "Interior horizontal edges get green dashed. "
            "Interior vertical edges should have NO border."
        ),
    })


# =====================================================================
# Fixture 3: cell-spacing-nonzero-borders
# =====================================================================

def make_cell_spacing_nonzero_borders():
    """2x2 table with non-zero cell spacing and mixed borders.

    All 6 table-level borders = single, sz=8, color="000000".
    tblCellSpacing = 20 (non-zero).
    Cell (0,0) has tcBorders: right = thick, sz=12, color="FF0000".
    """
    doc = Document()
    doc.add_paragraph("Fixture 3: Non-zero cell spacing with borders.")

    tbl_el, tblPr = make_table_with_borders(doc, 2, 2, [
        ["R0C0", "R0C1"],
        ["R1C0", "R1C1"],
    ])

    add_distinct_tbl_borders(tblPr, [
        ("top",     "single", "8", "000000"),
        ("left",    "single", "8", "000000"),
        ("bottom",  "single", "8", "000000"),
        ("right",   "single", "8", "000000"),
        ("insideH", "single", "8", "000000"),
        ("insideV", "single", "8", "000000"),
    ])

    # Add tblCellSpacing = 20
    tbl_cs = make_element("w:tblCellSpacing", {
        "w:w": "20", "w:type": "dxa",
    })
    tblPr.append(tbl_cs)

    # Cell (0,0) has explicit right border: thick red
    add_tc_borders(tbl_el, 0, 0, [
        ("right", "thick", "12", "FF0000", "0"),
    ])

    save_fixture("cell-spacing-nonzero-borders", doc, {
        "name": "cell-spacing-nonzero-borders",
        "spec_ref": "MS-OI29500 SS2.1.136, SS2.1.138, SS2.1.150",
        "description": (
            "2x2 table with tblCellSpacing=20 (non-zero). "
            "All table borders black single sz=8. "
            "Cell (0,0) has explicit right=thick sz=12 red. "
            "Tests cell spacing interaction with border conflict resolution."
        ),
    })


# =====================================================================
# Main
# =====================================================================

def main():
    print("\n== Table Border Inside Cascade Fixtures ==")
    make_inside_h_mapping_3x3()
    make_inside_h_only_no_inside_v()
    make_cell_spacing_nonzero_borders()
    print("\nDone.")


if __name__ == "__main__":
    main()

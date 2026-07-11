# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixture pairs (old/new) for tracked-change formatting tests.

Each pair exercises a specific formatting change scenario. The Rust tests
import both files, run diff_and_redline, and verify the pPrChange/rPrChange
content in the redline output.

IMPORTANT: The diff engine only detects formatting changes on paragraphs
whose TEXT also changed (identical-text blocks are intentionally skipped
as "formatting-only" differences). For pPrChange tests we include a small
text change alongside the formatting change so the diff engine treats the
block as BlockModified, which triggers formatting comparison in merge.

Run:  uv run create_docs.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Twips

ROOT = Path(__file__).parent


def save(doc: Document, *parts: str) -> None:
    path = ROOT / Path(*parts)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    print(f"  wrote {path}")


# ── 1. Alignment change (left -> center) with text edit ──────────────────

def create_alignment_change():
    """Old: left-aligned, text "Alignment test paragraph here."
    New: center-aligned, text "Alignment test paragraph."
    The small text difference ensures the diff engine detects a BlockModified,
    which then triggers pPrChange comparison during merge."""
    old = Document()
    p = old.add_paragraph("Alignment test paragraph here.")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    save(old, "alignment-change", "old.docx")

    new = Document()
    p = new.add_paragraph("Alignment test paragraph.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    save(new, "alignment-change", "new.docx")


# ── 2. Spacing change (before=240 -> before=480) with text edit ──────────

def create_spacing_change():
    """Old: spacing before=240, text "Spacing test paragraph here."
    New: spacing before=480, text "Spacing test paragraph."
    """
    old = Document()
    p = old.add_paragraph("Spacing test paragraph here.")
    p.paragraph_format.space_before = Twips(240)
    save(old, "spacing-change", "old.docx")

    new = Document()
    p = new.add_paragraph("Spacing test paragraph.")
    p.paragraph_format.space_before = Twips(480)
    save(new, "spacing-change", "new.docx")


# ── 3. Indentation change (add hanging indent) with text edit ────────────

def create_indentation_change():
    """Old: no indent, text "Indentation test paragraph here."
    New: left=720, hanging=360, text "Indentation test paragraph."
    """
    old = Document()
    old.add_paragraph("Indentation test paragraph here.")
    save(old, "indentation-change", "old.docx")

    new = Document()
    p = new.add_paragraph("Indentation test paragraph.")
    pf = p.paragraph_format
    pf.left_indent = Twips(720)
    pf.first_line_indent = Twips(-360)  # negative = hanging
    save(new, "indentation-change", "new.docx")


# ── 4. Bold change (normal -> bold) with text edit ───────────────────────

def create_bold_change():
    """Old: normal text "Bold test text here."
    New: bold text "Bold test text."
    """
    old = Document()
    old.add_paragraph("Bold test text here.")
    save(old, "bold-change", "old.docx")

    new = Document()
    p = new.add_paragraph()
    run = p.add_run("Bold test text.")
    run.bold = True
    save(new, "bold-change", "new.docx")


# ── 5. Complete snapshot (alignment + spacing + keepNext) ────────────────

def create_complete_snapshot():
    """Old: left + spacing-before=240 + keepNext, text "Complete snapshot test here."
    New: center (only alignment changes), text "Complete snapshot test."

    Per ECMA-376 §17.13.5.29, pPrChange should capture ALL previous properties,
    not just the one that changed.
    """
    old = Document()
    p = old.add_paragraph("Complete snapshot test here.")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Twips(240)
    # Set keepNext via low-level XML since python-docx doesn't expose it directly
    ppr = p._element.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    ppr.insert(0, keep_next)
    save(old, "complete-snapshot", "old.docx")

    new = Document()
    p = new.add_paragraph("Complete snapshot test.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Twips(240)
    ppr = p._element.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    ppr.insert(0, keep_next)
    save(new, "complete-snapshot", "new.docx")


# ── 6. Cell shading change ──────────────────────────────────────────────

def create_cell_shading_change():
    """Old: table with unshaded cell, text "Cell content here."
    New: cell gets yellow shading, text "Cell content."
    """
    old = Document()
    table = old.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Cell content here."
    save(old, "cell-shading-change", "old.docx")

    new = Document()
    table = new.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "Cell content."
    # Set cell shading via low-level XML
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "FFFF00")
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)
    save(new, "cell-shading-change", "new.docx")


if __name__ == "__main__":
    print("Creating tracked-change-formatting fixtures...")
    create_alignment_change()
    create_spacing_change()
    create_indentation_change()
    create_bold_change()
    create_complete_snapshot()
    create_cell_shading_change()
    print("Done.")

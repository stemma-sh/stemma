# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for table border conflict edge-case tests
(MS-OI29500 §2.1.169: one-sided border propagation & color brightness tiebreaker).

Run:  cd stemma-engine/testdata/spec-compliance/table-border-conflict-edge-cases && mise exec -- uv run create_docs.py
"""

import json
from pathlib import Path
from lxml import etree

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc: DocxDocument, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  table-border-conflict-edge-cases/{name}/")


def make_table_with_borders(doc, rows=2, cols=2, cell_texts=None):
    """Create a table, return (tbl_element, tblPr_element)."""
    tbl = doc.add_table(rows=rows, cols=cols)
    if cell_texts:
        for r in range(rows):
            for c in range(cols):
                if r < len(cell_texts) and c < len(cell_texts[r]):
                    tbl.cell(r, c).text = cell_texts[r][c]
    tbl_el = tbl._tbl
    tblPr = tbl_el.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_el.insert(0, tblPr)
    return tbl_el, tblPr


def add_tbl_borders(tblPr, edges, val="single", sz="12", color="FF0000", space="0"):
    """Add tblBorders to tblPr."""
    tbl_borders = make_element("w:tblBorders")
    for edge in edges:
        border = make_element(f"w:{edge}", {
            "w:val": val, "w:sz": sz, "w:color": color, "w:space": space,
        })
        tbl_borders.append(border)
    tblPr.append(tbl_borders)


def add_tc_borders(tbl_el, row, col, edges_spec):
    """Add tcBorders to a specific cell. edges_spec is a list of
    (edge_name, val, sz, color, space) tuples."""
    tr = tbl_el.findall(w("tr"))[row]
    tc = tr.findall(w("tc"))[col]
    tcPr = tc.find(w("tcPr"))
    if tcPr is None:
        tcPr = make_element("w:tcPr")
        tc.insert(0, tcPr)
    tc_borders = make_element("w:tcBorders")
    for edge, val, sz, color, space in edges_spec:
        border = make_element(f"w:{edge}", {
            "w:val": val, "w:sz": sz, "w:color": color, "w:space": space,
        })
        tc_borders.append(border)
    tcPr.append(tc_borders)


def add_tbl_grid(tbl_el, col_widths):
    """Add or replace tblGrid with specific column widths."""
    existing = tbl_el.find(w("tblGrid"))
    if existing is not None:
        tbl_el.remove(existing)
    grid = make_element("w:tblGrid")
    for cw in col_widths:
        grid.append(make_element("w:gridCol", {"w:w": str(cw)}))
    tblPr = tbl_el.find(w("tblPr"))
    if tblPr is not None:
        idx = list(tbl_el).index(tblPr) + 1
        tbl_el.insert(idx, grid)
    else:
        tbl_el.insert(0, grid)


def strip_all_tbl_borders(tblPr):
    """Remove any tblBorders that python-docx may have auto-added."""
    for existing in tblPr.findall(w("tblBorders")):
        tblPr.remove(existing)


# =====================================================================
# Fixture 1: One-sided border propagates to neighbor
# MS-OI29500 §2.1.169: When one cell has a border and the adjacent
# cell has none, the existing border should be displayed.
# =====================================================================

def make_one_sided_border_propagates():
    """1-row, 2-column table with NO table-level borders.
    Cell (0,0): right border = single, sz=12, color="FF0000" (red)
    Cell (0,1): NO tcBorders at all.

    Per MS-OI29500 §2.1.169, "none" means the opposing border wins.
    After conflict resolution, cell (0,1)'s left border should be the
    red border from cell (0,0).
    """
    doc = Document()
    doc.add_paragraph("One-sided border: cell (0,0) has right border, cell (0,1) has none.")

    tbl_el, tblPr = make_table_with_borders(doc, 1, 2, [
        ["HasRight", "NoBorders"],
    ])

    # Remove any default table borders.
    strip_all_tbl_borders(tblPr)

    add_tbl_grid(tbl_el, [3000, 3000])

    # Cell (0,0): right border = single, sz=12, color=FF0000
    add_tc_borders(tbl_el, 0, 0, [
        ("right", "single", "12", "FF0000", "0"),
    ])

    # Cell (0,1): deliberately NO tcBorders.

    save_fixture("one-sided-border-propagates", doc, {
        "name": "one-sided-border-propagates",
        "spec_ref": "MS-OI29500 SS2.1.169",
        "description": (
            "1x2 table with no table borders. "
            "Cell (0,0) has right border = single sz=12 red. "
            "Cell (0,1) has NO tcBorders. "
            "Per SS2.1.169: none means the opposing border displays. "
            "After conflict resolution, cell (0,1) left should be the red border."
        ),
    })


# =====================================================================
# Fixture 2: Table insideV fallback — both cells agree
# ISO 29500-1 §17.4.66: insideV from table-level cascades to interior
# cells and both cells at a shared edge should agree.
# =====================================================================

def make_one_sided_table_fallback():
    """2-row, 2-column table with table-level insideV border = single,
    sz=8, color="0000FF" (blue). No insideH. No cell-level borders.

    All interior vertical edges should get the insideV border.
    Cell (0,0) right and cell (0,1) left should both be blue.
    """
    doc = Document()
    doc.add_paragraph("Table insideV fallback: both cells at shared edge should agree.")

    tbl_el, tblPr = make_table_with_borders(doc, 2, 2, [
        ["R0C0", "R0C1"],
        ["R1C0", "R1C1"],
    ])

    # Remove any default table borders, then add only insideV.
    strip_all_tbl_borders(tblPr)
    add_tbl_borders(tblPr, ["insideV"], val="single", sz="8", color="0000FF")

    add_tbl_grid(tbl_el, [3000, 3000])

    save_fixture("one-sided-table-fallback", doc, {
        "name": "one-sided-table-fallback",
        "spec_ref": "ISO 29500-1 SS17.4.66",
        "description": (
            "2x2 table with only insideV border (blue single sz=8). "
            "No cell-level borders. Interior vertical edges should "
            "get the insideV border. Both cells at shared edge should agree."
        ),
    })


# =====================================================================
# Fixture 3: Color tiebreaker — darker wins
# MS-OI29500 §2.1.169: Equal weight → brightness = R + B + 2*G,
# smaller (darker) wins.
# =====================================================================

def make_color_tiebreaker_darker_wins():
    """1-row, 2-column table with NO table-level borders.
    Cell (0,0): right border = single, sz=8, color="000080"
        (dark blue: R=0, G=0, B=128, brightness=0+128+0=128)
    Cell (0,1): left border = single, sz=8, color="80FF80"
        (light green: R=128, G=255, B=128, brightness=128+128+510=766)

    Both single, same size, same weight (8*1=8).
    Tiebreaker: brightness. Dark blue (128) < light green (766).
    Dark blue wins.
    """
    doc = Document()
    doc.add_paragraph("Color tiebreaker: darker wins when weight is equal.")

    tbl_el, tblPr = make_table_with_borders(doc, 1, 2, [
        ["DarkBlue", "LightGreen"],
    ])

    strip_all_tbl_borders(tblPr)
    add_tbl_grid(tbl_el, [3000, 3000])

    # Cell (0,0): right = single, sz=8, color=000080
    add_tc_borders(tbl_el, 0, 0, [
        ("right", "single", "8", "000080", "0"),
    ])
    # Cell (0,1): left = single, sz=8, color=80FF80
    add_tc_borders(tbl_el, 0, 1, [
        ("left", "single", "8", "80FF80", "0"),
    ])

    save_fixture("color-tiebreaker-darker-wins", doc, {
        "name": "color-tiebreaker-darker-wins",
        "spec_ref": "MS-OI29500 SS2.1.169",
        "description": (
            "1x2 table with no table borders. "
            "Cell (0,0) right = single sz=8 color=000080 (dark blue, brightness=128). "
            "Cell (0,1) left = single sz=8 color=80FF80 (light green, brightness=766). "
            "Same style and weight. Per SS2.1.169, darker color wins. "
            "Dark blue (000080) should win."
        ),
    })


# =====================================================================
# Fixture 4: Color tiebreaker — secondary brightness (B + 2*G)
# MS-OI29500 §2.1.169: Equal primary brightness → B + 2*G, smaller wins.
# =====================================================================

def make_color_tiebreaker_secondary():
    """1-row, 2-column table with NO table-level borders.
    Cell (0,0): right border = single, sz=8, color="FF0000"
        (pure red: R=255, G=0, B=0, primary=R+B+2G=255+0+0=255)
        (secondary=B+2G=0+0=0)
    Cell (0,1): left border = single, sz=8, color="0000FF"
        (pure blue: R=0, G=0, B=255, primary=R+B+2G=0+255+0=255)
        (secondary=B+2G=255+0=255)

    Same style, same size, same primary brightness.
    Secondary: red (0) < blue (255). Red wins.
    """
    doc = Document()
    doc.add_paragraph("Color tiebreaker: secondary brightness (B+2G) breaks tie.")

    tbl_el, tblPr = make_table_with_borders(doc, 1, 2, [
        ["PureRed", "PureBlue"],
    ])

    strip_all_tbl_borders(tblPr)
    add_tbl_grid(tbl_el, [3000, 3000])

    # Cell (0,0): right = single, sz=8, color=FF0000
    add_tc_borders(tbl_el, 0, 0, [
        ("right", "single", "8", "FF0000", "0"),
    ])
    # Cell (0,1): left = single, sz=8, color=0000FF
    add_tc_borders(tbl_el, 0, 1, [
        ("left", "single", "8", "0000FF", "0"),
    ])

    save_fixture("color-tiebreaker-same-style-same-weight", doc, {
        "name": "color-tiebreaker-same-style-same-weight",
        "spec_ref": "MS-OI29500 SS2.1.169",
        "description": (
            "1x2 table with no table borders. "
            "Cell (0,0) right = single sz=8 color=FF0000 (red, primary=255, secondary=0). "
            "Cell (0,1) left = single sz=8 color=0000FF (blue, primary=255, secondary=255). "
            "Same style, weight, and primary brightness. "
            "Per SS2.1.169, secondary brightness (B+2G) breaks the tie. "
            "Red (secondary=0) wins over blue (secondary=255)."
        ),
    })


# =====================================================================
# Fixture 5: Color tiebreaker — darker on RIGHT cell
# Same as fixture 3 but colors swapped: darker is on cell (0,1).
# This exposes the bug where resolve_border_conflict() picks first arg.
# =====================================================================

def make_color_tiebreaker_darker_on_right():
    """1-row, 2-column table with NO table-level borders.
    Cell (0,0): right border = single, sz=8, color="80FF80"
        (light green: brightness=766)
    Cell (0,1): left border = single, sz=8, color="000080"
        (dark blue: brightness=128)

    Same style, same size, same weight. Per MS-OI29500 §2.1.169,
    dark blue (brightness=128) should win — but resolve_border_conflict()
    picks the first argument (light green), producing the wrong result.
    """
    doc = Document()
    doc.add_paragraph("Color tiebreaker reversed: darker on RIGHT cell.")

    tbl_el, tblPr = make_table_with_borders(doc, 1, 2, [
        ["LightGreen", "DarkBlue"],
    ])

    strip_all_tbl_borders(tblPr)
    add_tbl_grid(tbl_el, [3000, 3000])

    # Cell (0,0): right = single, sz=8, color=80FF80 (light green)
    add_tc_borders(tbl_el, 0, 0, [
        ("right", "single", "8", "80FF80", "0"),
    ])
    # Cell (0,1): left = single, sz=8, color=000080 (dark blue)
    add_tc_borders(tbl_el, 0, 1, [
        ("left", "single", "8", "000080", "0"),
    ])

    save_fixture("color-tiebreaker-darker-on-right", doc, {
        "name": "color-tiebreaker-darker-on-right",
        "spec_ref": "MS-OI29500 SS2.1.169",
        "description": (
            "1x2 table. Cell (0,0) right = light green (brightness=766). "
            "Cell (0,1) left = dark blue (brightness=128). Same weight. "
            "Per SS2.1.169, darker (dark blue) should win. "
            "BUG: resolve_border_conflict() picks first arg (light green)."
        ),
    })


# =====================================================================
# Fixture 6: Secondary brightness — winning color on RIGHT cell
# Same as fixture 4 but colors swapped: red (winner) is on cell (0,1).
# =====================================================================

def make_color_tiebreaker_secondary_reversed():
    """1-row, 2-column table with NO table-level borders.
    Cell (0,0): right border = single, sz=8, color="0000FF"
        (blue: primary=255, secondary=255)
    Cell (0,1): left border = single, sz=8, color="FF0000"
        (red: primary=255, secondary=0)

    Same style, weight, primary brightness.
    Secondary: red (0) < blue (255). Red should win.
    But red is on cell (0,1) (second arg), so the first-arg-wins bug
    will incorrectly pick blue.
    """
    doc = Document()
    doc.add_paragraph("Secondary brightness reversed: winning color on RIGHT cell.")

    tbl_el, tblPr = make_table_with_borders(doc, 1, 2, [
        ["PureBlue", "PureRed"],
    ])

    strip_all_tbl_borders(tblPr)
    add_tbl_grid(tbl_el, [3000, 3000])

    # Cell (0,0): right = single, sz=8, color=0000FF (blue)
    add_tc_borders(tbl_el, 0, 0, [
        ("right", "single", "8", "0000FF", "0"),
    ])
    # Cell (0,1): left = single, sz=8, color=FF0000 (red)
    add_tc_borders(tbl_el, 0, 1, [
        ("left", "single", "8", "FF0000", "0"),
    ])

    save_fixture("color-tiebreaker-secondary-reversed", doc, {
        "name": "color-tiebreaker-secondary-reversed",
        "spec_ref": "MS-OI29500 SS2.1.169",
        "description": (
            "1x2 table. Cell (0,0) right = blue (primary=255, secondary=255). "
            "Cell (0,1) left = red (primary=255, secondary=0). "
            "Same weight + primary brightness. Per SS2.1.169, red (secondary=0) wins. "
            "BUG: resolve_border_conflict() picks first arg (blue)."
        ),
    })


# =====================================================================
# Main
# =====================================================================

def main():
    print("\n== Table Border Conflict Edge-Case Fixtures ==")
    make_one_sided_border_propagates()
    make_one_sided_table_fallback()
    make_color_tiebreaker_darker_wins()
    make_color_tiebreaker_secondary()
    make_color_tiebreaker_darker_on_right()
    make_color_tiebreaker_secondary_reversed()
    print("\nDone.")


if __name__ == "__main__":
    main()

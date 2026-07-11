# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for numbering constraint spec-compliance tests.

Tests numbering resolution semantics from ISO 29500-1:
  - §17.9.18: numId=0 removes inherited numbering
  - §17.9.23: pStyle reverse binding in numbering levels
  - §17.7.4.14: direct numPr overrides style numPr

Run:  uv run create_numbering_constraints.py
"""

import json
from pathlib import Path
from lxml import etree

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

# ── XML namespace helpers ────────────────────────────────────────────────

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def w(tag: str) -> str:
    """Return a fully-qualified wordprocessingml tag."""
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    """Create an OxmlElement with optional attributes."""
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


# ── Save helpers ─────────────────────────────────────────────────────────

def save_fixture(
    area: str,
    name: str,
    doc: DocxDocument,
    metadata: dict,
    filename: str = "input.docx",
) -> None:
    """Save a single-doc fixture (for parsing/model tests)."""
    out = ROOT / area / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / filename))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  {area}/{name}/")


def _inject_numbering_xml(doc: DocxDocument, numbering_xml: str) -> None:
    """Replace the numbering.xml part with custom XML.

    This gives us full control over numbering definitions that python-docx
    doesn't expose directly.
    """
    # Ensure numbering part exists
    numbering_part = None
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        pass

    if numbering_part is None:
        # Add a dummy list to force numbering part creation
        doc.add_paragraph("dummy", style="List Bullet")
        # Remove the dummy paragraph
        body = doc.element.body
        last_p = body.findall(w("p"))[-1]
        body.remove(last_p)
        numbering_part = doc.part.numbering_part

    # Replace the numbering XML
    numbering_part._element = etree.fromstring(numbering_xml.encode("utf-8"))


def _inject_style(doc: DocxDocument, style_xml: str) -> None:
    """Append a custom style element to the styles part."""
    style_el = etree.fromstring(style_xml.encode("utf-8"))
    doc.styles.element.append(style_el)


# =========================================================================
# numId=0 removal (ISO 29500-1 §17.9.18)
# =========================================================================

def make_numid_zero_removal() -> None:
    """numId=0 removes inherited numbering from a paragraph style.

    ISO 29500-1 §17.9.18: "If this element's val attribute is 0, then the
    paragraph shall not have any numbering properties applied to it, regardless
    of any numbering properties inherited from the paragraph's style."

    Fixture:
      - Style "NumberedStyle" has numPr (numId=1, ilvl=0) pointing to bullet list
      - P1: uses "NumberedStyle" — inherits numbering from style
      - P2: uses "NumberedStyle" + direct numPr with numId="0" — numbering removed
      - P3: uses "NumberedStyle" — inherits numbering again (P2 doesn't affect P3)
    """
    doc = Document()

    # Numbering definition: bullet list
    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="bullet"/>
      <w:lvlText w:val="\u00B7"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # Create paragraph style "NumberedStyle" with numPr
    style_xml = f"""\
<w:style w:type="paragraph" w:styleId="NumberedStyle" xmlns:w="{W}">
  <w:name w:val="Numbered Style"/>
  <w:pPr>
    <w:numPr>
      <w:ilvl w:val="0"/>
      <w:numId w:val="1"/>
    </w:numPr>
  </w:pPr>
</w:style>"""
    _inject_style(doc, style_xml)

    # P1: Uses "NumberedStyle" — should inherit numbering from style
    p1 = doc.add_paragraph("Paragraph with inherited numbering from style")
    p1_pPr = p1._p.get_or_add_pPr()
    p1_pPr.append(make_element("w:pStyle", {"w:val": "NumberedStyle"}))

    # P2: Uses "NumberedStyle" + direct numPr with numId="0" — removes numbering
    p2 = doc.add_paragraph("Paragraph with numId zero removes numbering")
    p2_pPr = p2._p.get_or_add_pPr()
    p2_pPr.append(make_element("w:pStyle", {"w:val": "NumberedStyle"}))
    numPr = make_element("w:numPr")
    numPr.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr.append(make_element("w:numId", {"w:val": "0"}))
    p2_pPr.append(numPr)

    # P3: Uses "NumberedStyle" — should have numbering again
    p3 = doc.add_paragraph("Paragraph with numbering restored from style")
    p3_pPr = p3._p.get_or_add_pPr()
    p3_pPr.append(make_element("w:pStyle", {"w:val": "NumberedStyle"}))

    save_fixture("numbering-constraints", "numid-zero-removal", doc, {
        "name": "numid-zero-removal",
        "spec_ref": "ISO 29500-1 §17.9.18",
        "description": (
            "numId=0 removes inherited numbering. Style 'NumberedStyle' has "
            "numPr (numId=1). P1 and P3 use the style (get numbering). "
            "P2 uses the style + direct numId=0 (numbering removed)."
        ),
        "expected_behavior": (
            "P1: has numbering (bullet). P2: no numbering (numId=0 removes it). "
            "P3: has numbering (bullet). The removal on P2 is local."
        ),
        "current_status": "TESTING",
    })


# =========================================================================
# pStyle reverse binding (ISO 29500-1 §17.9.23)
# =========================================================================

def make_pstyle_reverse_binding() -> None:
    """pStyle in a numbering level binds a style to that level.

    ISO 29500-1 §17.9.23: "If a paragraph style includes a pStyle element
    which references a defined numbering level, then any paragraph using that
    style shall be numbered using the associated level."

    More precisely, the abstract numbering level can specify a pStyle value.
    When a paragraph has that style, even if the style itself has no numPr,
    the paragraph gets numbering from the abstract level that references it.

    Fixture:
      - abstractNum level 0 has <w:pStyle w:val="ListBullet"/>
      - Style "ListBullet" has NO numPr
      - P1: uses style "ListBullet" — gets numbering via reverse binding
      - P2: normal paragraph — no numbering
    """
    doc = Document()

    # Numbering definition with pStyle binding on level 0
    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="bullet"/>
      <w:pStyle w:val="ListBullet"/>
      <w:lvlText w:val="\u00B7"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # Create paragraph style "ListBullet" with NO numPr
    style_xml = f"""\
<w:style w:type="paragraph" w:styleId="ListBullet" xmlns:w="{W}">
  <w:name w:val="List Bullet Custom"/>
  <w:pPr/>
</w:style>"""
    _inject_style(doc, style_xml)

    # P1: Uses "ListBullet" — should get numbering from the abstract level's pStyle
    p1 = doc.add_paragraph("Paragraph with reverse binding numbering")
    p1_pPr = p1._p.get_or_add_pPr()
    p1_pPr.append(make_element("w:pStyle", {"w:val": "ListBullet"}))

    # P2: Normal paragraph — no numbering
    doc.add_paragraph("Normal paragraph without numbering")

    save_fixture("numbering-constraints", "pstyle-reverse-binding", doc, {
        "name": "pstyle-reverse-binding",
        "spec_ref": "ISO 29500-1 §17.9.23",
        "description": (
            "pStyle reverse binding: abstractNum level 0 has "
            "<w:pStyle w:val='ListBullet'/>. Style 'ListBullet' has no numPr. "
            "P1 with style 'ListBullet' should get numbering from the level."
        ),
        "expected_behavior": (
            "P1: has bullet numbering (from abstract level's pStyle binding). "
            "P2: no numbering."
        ),
        "current_status": "TESTING — pStyle reverse binding likely unimplemented",
    })


# =========================================================================
# Style numPr override (ISO 29500-1 §17.7.4.14)
# =========================================================================

def make_style_numpr_override() -> None:
    """Direct numPr overrides style-inherited numPr.

    ISO 29500-1 §17.7.4.14: "If this element is present in the paragraph
    properties of a paragraph style, then any numbering level defined using
    the numPr element on the associated abstract numbering definition shall
    be ignored, and this element shall determine the numbering which is
    applied to the paragraph."

    More practically: when a paragraph style has numPr and the paragraph
    itself also has direct numPr, the direct numPr wins.

    Fixture:
      - numbering definition 1: bullet (abstractNum 0)
      - numbering definition 2: decimal (abstractNum 1)
      - Style "NumberedStyle" has numPr pointing to numId=1 (bullets)
      - P1: uses "NumberedStyle", no direct numPr — gets bullets from style
      - P2: uses "NumberedStyle" + direct numPr pointing to numId=2 — gets decimal
    """
    doc = Document()

    # Two numbering definitions: bullets (numId=1) and decimal (numId=2)
    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="bullet"/>
      <w:lvlText w:val="\u00B7"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:abstractNum w:abstractNumId="1">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
  <w:num w:numId="2">
    <w:abstractNumId w:val="1"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # Create paragraph style "NumberedStyle" with numPr pointing to numId=1 (bullets)
    style_xml = f"""\
<w:style w:type="paragraph" w:styleId="NumberedStyle" xmlns:w="{W}">
  <w:name w:val="Numbered Style"/>
  <w:pPr>
    <w:numPr>
      <w:ilvl w:val="0"/>
      <w:numId w:val="1"/>
    </w:numPr>
  </w:pPr>
</w:style>"""
    _inject_style(doc, style_xml)

    # P1: Uses "NumberedStyle", no direct numPr — gets bullets from style
    p1 = doc.add_paragraph("Paragraph with bullets from style")
    p1_pPr = p1._p.get_or_add_pPr()
    p1_pPr.append(make_element("w:pStyle", {"w:val": "NumberedStyle"}))

    # P2: Uses "NumberedStyle" + direct numPr pointing to numId=2 (decimal)
    p2 = doc.add_paragraph("Paragraph with decimal override")
    p2_pPr = p2._p.get_or_add_pPr()
    p2_pPr.append(make_element("w:pStyle", {"w:val": "NumberedStyle"}))
    numPr = make_element("w:numPr")
    numPr.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr.append(make_element("w:numId", {"w:val": "2"}))
    p2_pPr.append(numPr)

    save_fixture("numbering-constraints", "style-numpr-override", doc, {
        "name": "style-numpr-override",
        "spec_ref": "ISO 29500-1 §17.7.4.14",
        "description": (
            "Direct numPr overrides style numPr. Style 'NumberedStyle' has "
            "numPr (numId=1, bullets). P1 uses style only (bullets). "
            "P2 uses style + direct numPr (numId=2, decimal)."
        ),
        "expected_behavior": (
            "P1: bullet numbering (from style numId=1). "
            "P2: decimal numbering '1.' (from direct numId=2, overrides style)."
        ),
        "current_status": "TESTING",
    })


# =========================================================================

def main() -> None:
    print("Generating numbering constraint fixtures:")
    make_numid_zero_removal()
    make_pstyle_reverse_binding()
    make_style_numpr_override()
    print("\nDone.")


if __name__ == "__main__":
    main()

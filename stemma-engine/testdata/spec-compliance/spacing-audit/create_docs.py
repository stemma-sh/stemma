# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for spacing & indentation audit tests.

Each fixture targets a specific behavioral constraint from ECMA-376
§17.3.1.33 (spacing) and §17.3.1.12 (ind) that our implementation
may violate.

Run:  uv run create_docs.py
  or: mise exec -- python3 create_docs.py
"""

import json
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc: DocxDocument, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  spacing-audit/{name}/")


# =========================================================================
# Fixture 1: beforeAutospacing / afterAutospacing override before/after
# =========================================================================

def make_autospacing_override() -> None:
    """§17.3.1.33: beforeAutospacing/afterAutospacing override before/after.

    When beforeAutospacing="1" is set, the before value is IGNORED and
    the consumer determines spacing automatically (like HTML). Same for
    afterAutospacing.

    P0: before=480, beforeAutospacing="1"  -> before should be ignored
    P1: after=480, afterAutospacing="1"    -> after should be ignored
    P2: before=480, after=480              -> control (no autospacing)
    P3: before=480, beforeAutospacing="1", beforeLines=200
        -> both before AND beforeLines should be ignored
    """
    doc = Document()

    # P0: beforeAutospacing overrides before
    p0 = doc.add_paragraph("P0: before=480 with beforeAutospacing=1 (before should be ignored)")
    pPr0 = p0._p.get_or_add_pPr()
    sp0 = make_element("w:spacing", {
        "w:before": "480",
        "w:beforeAutospacing": "1",
    })
    pPr0.append(sp0)

    # P1: afterAutospacing overrides after
    p1 = doc.add_paragraph("P1: after=480 with afterAutospacing=1 (after should be ignored)")
    pPr1 = p1._p.get_or_add_pPr()
    sp1 = make_element("w:spacing", {
        "w:after": "480",
        "w:afterAutospacing": "1",
    })
    pPr1.append(sp1)

    # P2: control — no autospacing, values should be preserved
    p2 = doc.add_paragraph("P2: before=480 after=480 (control, no autospacing)")
    pPr2 = p2._p.get_or_add_pPr()
    sp2 = make_element("w:spacing", {
        "w:before": "480",
        "w:after": "480",
    })
    pPr2.append(sp2)

    # P3: beforeAutospacing overrides BOTH before AND beforeLines
    p3 = doc.add_paragraph("P3: before=480, beforeLines=200, beforeAutospacing=1 (both ignored)")
    pPr3 = p3._p.get_or_add_pPr()
    sp3 = make_element("w:spacing", {
        "w:before": "480",
        "w:beforeLines": "200",
        "w:beforeAutospacing": "1",
    })
    pPr3.append(sp3)

    save_fixture("autospacing-override", doc, {
        "name": "autospacing-override",
        "spec_ref": "ISO 29500-1 §17.3.1.33",
        "description": "beforeAutospacing/afterAutospacing override before/after/beforeLines/afterLines",
        "expected_behavior": (
            "P0: before is ignored (beforeAutospacing=1). "
            "P1: after is ignored (afterAutospacing=1). "
            "P2: control — before=480, after=480. "
            "P3: both before and beforeLines are ignored (beforeAutospacing=1)."
        ),
    })


# =========================================================================
# Fixture 2: contextualSpacing explicit false vs absent
# =========================================================================

def make_contextual_spacing_explicit_false() -> None:
    """§17.3.1.9: Explicit contextualSpacing val="0" overrides style's true.

    When a style sets contextualSpacing=true but the paragraph has direct
    contextualSpacing val="0", the direct false should win.

    P0: CtxStyle (contextualSpacing=true), direct contextualSpacing val="0"
        -> should be false
    P1: CtxStyle (contextualSpacing=true), no direct override
        -> should be true (from style)
    P2: CtxStyle (contextualSpacing=true), direct contextualSpacing (no val)
        -> should be true (element present = true)
    """
    doc = Document()

    # Create style with contextualSpacing=true
    styles_el = doc.styles.element
    style_el = make_element("w:style", {"w:type": "paragraph", "w:styleId": "CtxStyle"})
    name_el = make_element("w:name", {"w:val": "CtxStyle"})
    style_el.append(name_el)
    pPr_style = make_element("w:pPr")
    pPr_style.append(make_element("w:contextualSpacing"))
    # Also add spacing so we can verify contextualSpacing matters
    pPr_style.append(make_element("w:spacing", {"w:before": "240", "w:after": "240"}))
    style_el.append(pPr_style)
    styles_el.append(style_el)

    # P0: explicit contextualSpacing val="0" should override style's true
    p0 = doc.add_paragraph("P0: CtxStyle + direct contextualSpacing=false")
    p0.style = doc.styles["CtxStyle"]
    pPr0 = p0._p.get_or_add_pPr()
    cs0 = make_element("w:contextualSpacing", {"w:val": "0"})
    pPr0.append(cs0)

    # P1: no direct contextualSpacing — inherits true from style
    p1 = doc.add_paragraph("P1: CtxStyle only (inherits contextualSpacing=true)")
    p1.style = doc.styles["CtxStyle"]

    # P2: direct contextualSpacing with no val (= true)
    p2 = doc.add_paragraph("P2: CtxStyle + direct contextualSpacing (no val = true)")
    p2.style = doc.styles["CtxStyle"]
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:contextualSpacing"))

    save_fixture("contextual-spacing-explicit-false", doc, {
        "name": "contextual-spacing-explicit-false",
        "spec_ref": "ISO 29500-1 §17.3.1.9",
        "description": "Explicit contextualSpacing val='0' should override style's true",
        "expected_behavior": (
            "P0: contextual_spacing=false (direct val='0' overrides style's true). "
            "P1: contextual_spacing=true (inherited from CtxStyle). "
            "P2: contextual_spacing=true (direct element present = true)."
        ),
    })


# =========================================================================
# Fixture 3: pPrChange spacing roundtrip (before_lines/after_lines)
# =========================================================================

def make_pprchange_spacing_roundtrip() -> None:
    """§17.13.5.29 + §17.3.1.33: pPrChange must preserve all spacing attrs.

    When a paragraph has a tracked spacing change, the previous spacing
    in w:pPrChange/w:pPr/w:spacing should preserve beforeLines, afterLines
    alongside before, after, line, lineRule.

    P0: current spacing = before=0, previous spacing = before=480, beforeLines=200,
        afterLines=100, line=360, lineRule=auto
    """
    doc = Document()

    p0 = doc.add_paragraph("P0: spacing change tracked — previous had beforeLines=200, afterLines=100")
    pPr0 = p0._p.get_or_add_pPr()

    # Current spacing
    sp_current = make_element("w:spacing", {"w:before": "0"})
    pPr0.append(sp_current)

    # pPrChange with previous spacing including beforeLines/afterLines
    ppr_change = make_element("w:pPrChange", {
        "w:id": "500",
        "w:author": "Audit",
        "w:date": "2025-01-01T00:00:00Z",
    })
    prev_ppr = make_element("w:pPr")
    prev_sp = make_element("w:spacing", {
        "w:before": "480",
        "w:beforeLines": "200",
        "w:after": "240",
        "w:afterLines": "100",
        "w:line": "360",
        "w:lineRule": "auto",
    })
    prev_ppr.append(prev_sp)
    ppr_change.append(prev_ppr)
    pPr0.append(ppr_change)

    save_fixture("pprchange-spacing-roundtrip", doc, {
        "name": "pprchange-spacing-roundtrip",
        "spec_ref": "ISO 29500-1 §17.13.5.29 + §17.3.1.33",
        "description": "pPrChange previous spacing must preserve beforeLines/afterLines",
        "expected_behavior": (
            "P0 previous_spacing should have before=480, beforeLines=200, "
            "after=240, afterLines=100, line=360, lineRule=auto. "
            "Roundtrip serialization must emit beforeLines and afterLines."
        ),
    })


# =========================================================================
# Fixture 4: pPrChange indent roundtrip (char-unit values)
# =========================================================================

def make_pprchange_indent_roundtrip() -> None:
    """§17.13.5.29 + §17.3.1.12: pPrChange must preserve char-unit indent attrs.

    When a paragraph has a tracked indent change, the previous indentation
    in w:pPrChange should preserve startChars, endChars, firstLineChars,
    hangingChars alongside left, right, firstLine, hanging.

    P0: current indent = left=0, previous = left=720, startChars=200, endChars=100
    """
    doc = Document()

    p0 = doc.add_paragraph("P0: indent change tracked — previous had startChars=200, endChars=100")
    pPr0 = p0._p.get_or_add_pPr()

    # Current indent
    ind_current = make_element("w:ind", {"w:left": "0"})
    pPr0.append(ind_current)

    # pPrChange with previous indent including char-unit values
    ppr_change = make_element("w:pPrChange", {
        "w:id": "600",
        "w:author": "Audit",
        "w:date": "2025-01-01T00:00:00Z",
    })
    prev_ppr = make_element("w:pPr")
    prev_ind = make_element("w:ind", {
        "w:left": "720",
        "w:right": "360",
        "w:startChars": "200",
        "w:endChars": "100",
    })
    prev_ppr.append(prev_ind)
    ppr_change.append(prev_ppr)
    pPr0.append(ppr_change)

    save_fixture("pprchange-indent-roundtrip", doc, {
        "name": "pprchange-indent-roundtrip",
        "spec_ref": "ISO 29500-1 §17.13.5.29 + §17.3.1.12",
        "description": "pPrChange previous indent must preserve char-unit values",
        "expected_behavior": (
            "P0 previous_indentation should have left=720, right=360, "
            "startChars=200, endChars=100. "
            "Roundtrip serialization must emit startChars and endChars."
        ),
    })


# =========================================================================
# Fixture 5: firstLineChars/hangingChars mutual exclusivity
# =========================================================================

def make_firstline_hanging_chars_exclusive() -> None:
    """§17.3.1.12: firstLineChars and hangingChars are mutually exclusive.

    Per spec, if both firstLineChars and hangingChars are specified,
    firstLineChars is ignored (hangingChars wins — opposite of twip behavior).
    Wait, let me re-read: "The firstLineChars and hangingChars attributes
    are mutually exclusive, if both are specified, then the firstLineChars
    value is ignored."

    P0: firstLineChars=200, hangingChars=100 — firstLineChars is IGNORED
    P1: firstLineChars=200 only (control)
    P2: hangingChars=100 only (control)
    P3: firstLine=360, hanging=720 — firstLine is IGNORED (hanging wins? no —
        per spec "firstLine and hanging are mutually exclusive, if both are
        specified, then the firstLine value is ignored")
    """
    doc = Document()

    # P0: both firstLineChars and hangingChars — firstLineChars is ignored
    p0 = doc.add_paragraph("P0: firstLineChars=200 + hangingChars=100 (firstLineChars ignored)")
    pPr0 = p0._p.get_or_add_pPr()
    ind0 = make_element("w:ind", {
        "w:left": "1440",
        "w:firstLineChars": "200",
        "w:hangingChars": "100",
    })
    pPr0.append(ind0)

    # P1: firstLineChars only
    p1 = doc.add_paragraph("P1: firstLineChars=200 only (control)")
    pPr1 = p1._p.get_or_add_pPr()
    ind1 = make_element("w:ind", {
        "w:left": "1440",
        "w:firstLineChars": "200",
    })
    pPr1.append(ind1)

    # P2: hangingChars only
    p2 = doc.add_paragraph("P2: hangingChars=100 only (control)")
    pPr2 = p2._p.get_or_add_pPr()
    ind2 = make_element("w:ind", {
        "w:left": "1440",
        "w:hangingChars": "100",
    })
    pPr2.append(ind2)

    # P3: both firstLine and hanging — firstLine is ignored per spec
    # NOTE: Our parser picks firstLine over hanging (checks firstLine first),
    # but the spec says "firstLine value is ignored" when both present.
    p3 = doc.add_paragraph("P3: firstLine=360 + hanging=720 (firstLine ignored per spec)")
    pPr3 = p3._p.get_or_add_pPr()
    ind3 = make_element("w:ind", {
        "w:left": "1440",
        "w:firstLine": "360",
        "w:hanging": "720",
    })
    pPr3.append(ind3)

    save_fixture("firstline-hanging-chars-exclusive", doc, {
        "name": "firstline-hanging-chars-exclusive",
        "spec_ref": "ISO 29500-1 §17.3.1.12",
        "description": "firstLineChars/hangingChars and firstLine/hanging mutual exclusivity",
        "expected_behavior": (
            "P0: firstLineChars is IGNORED, hangingChars=100 wins. "
            "P1: firstLineChars=200 (control). "
            "P2: hangingChars=100 (control). "
            "P3: firstLine is IGNORED, hanging=720 wins -> effective_first_line=-720."
        ),
    })


# =========================================================================
# Fixture 6: lineRule defaults to auto when omitted but line present
# =========================================================================

def make_line_rule_default_auto() -> None:
    """§17.3.1.33: lineRule defaults to 'auto' when omitted but line is present.

    Per spec: "If this attribute is omitted, then it shall be assumed to be
    of a value auto if a line attribute value is present."

    P0: line=276, no lineRule -> should default to auto
    P1: line=276, lineRule=auto (explicit control)
    P2: line=240, lineRule=exact (different rule)
    P3: no line, no lineRule -> no default
    """
    doc = Document()

    # P0: line without lineRule — should default to auto
    p0 = doc.add_paragraph("P0: line=276 without lineRule (should default to auto)")
    pPr0 = p0._p.get_or_add_pPr()
    sp0 = make_element("w:spacing", {"w:line": "276"})
    pPr0.append(sp0)

    # P1: explicit auto (control)
    p1 = doc.add_paragraph("P1: line=276 lineRule=auto (explicit)")
    pPr1 = p1._p.get_or_add_pPr()
    sp1 = make_element("w:spacing", {"w:line": "276", "w:lineRule": "auto"})
    pPr1.append(sp1)

    # P2: explicit exact
    p2 = doc.add_paragraph("P2: line=240 lineRule=exact")
    pPr2 = p2._p.get_or_add_pPr()
    sp2 = make_element("w:spacing", {"w:line": "240", "w:lineRule": "exact"})
    pPr2.append(sp2)

    # P3: before only, no line at all
    p3 = doc.add_paragraph("P3: before=120 only (no line, no lineRule)")
    pPr3 = p3._p.get_or_add_pPr()
    sp3 = make_element("w:spacing", {"w:before": "120"})
    pPr3.append(sp3)

    save_fixture("line-rule-default-auto", doc, {
        "name": "line-rule-default-auto",
        "spec_ref": "ISO 29500-1 §17.3.1.33",
        "description": "lineRule defaults to auto when omitted with line present",
        "expected_behavior": (
            "P0: lineRule=auto (default when line is present but lineRule omitted). "
            "P1: lineRule=auto (explicit). "
            "P2: lineRule=exact. "
            "P3: lineRule=None (no line value)."
        ),
    })


# =========================================================================
# Fixture 7: startChars supersedes start per spec
# =========================================================================

def make_start_chars_supersedes_start() -> None:
    """§17.3.1.12: When startChars is specified, start value is ignored.

    Per spec: "if the startChars attribute is specified, then this [start]
    value is ignored, and is superseded by this value."
    Also: "if the endChars attribute is specified, then this [end] value
    is ignored."

    Both values should be preserved in the model (for roundtrip), but the
    char-unit value takes behavioral precedence.

    P0: start=720, startChars=200 — both present, startChars wins
    P1: end=720, endChars=100 — both present, endChars wins
    P2: start=720 only (control)
    P3: startChars=200 only (control)
    """
    doc = Document()

    # P0: both start and startChars
    p0 = doc.add_paragraph("P0: start=720 + startChars=200 (startChars supersedes)")
    pPr0 = p0._p.get_or_add_pPr()
    ind0 = make_element("w:ind", {
        "w:start": "720",
        "w:startChars": "200",
    })
    pPr0.append(ind0)

    # P1: both end and endChars
    p1 = doc.add_paragraph("P1: end=720 + endChars=100 (endChars supersedes)")
    pPr1 = p1._p.get_or_add_pPr()
    ind1 = make_element("w:ind", {
        "w:end": "720",
        "w:endChars": "100",
    })
    pPr1.append(ind1)

    # P2: start only (control)
    p2 = doc.add_paragraph("P2: start=720 only (control)")
    pPr2 = p2._p.get_or_add_pPr()
    ind2 = make_element("w:ind", {"w:start": "720"})
    pPr2.append(ind2)

    # P3: startChars only (control)
    p3 = doc.add_paragraph("P3: startChars=200 only (control)")
    pPr3 = p3._p.get_or_add_pPr()
    ind3 = make_element("w:ind", {"w:startChars": "200"})
    pPr3.append(ind3)

    save_fixture("start-chars-supersedes-start", doc, {
        "name": "start-chars-supersedes-start",
        "spec_ref": "ISO 29500-1 §17.3.1.12",
        "description": "startChars supersedes start, endChars supersedes end",
        "expected_behavior": (
            "P0: start_chars=200 supersedes left=720 (both preserved in model). "
            "P1: end_chars=100 supersedes right=720. "
            "P2: left=720, no char-unit. "
            "P3: start_chars=200, left=None."
        ),
    })


# =========================================================================
# Fixture 8: firstLineChars supersedes firstLine
# =========================================================================

def make_firstline_chars_supersedes_firstline() -> None:
    """§17.3.1.12: When firstLineChars is specified, firstLine is ignored.

    Per spec: "If the firstLine attribute is also specified, then this
    [firstLineChars] value supersedes its other value."
    Also: "If the hanging attribute is also specified, then its
    [hangingChars] value is superseded by this value."

    P0: firstLine=360, firstLineChars=200 — firstLineChars supersedes
    P1: hanging=720, hangingChars=100 — hangingChars supersedes
    P2: firstLine=360 only (control)
    P3: firstLineChars=200 only (control)
    """
    doc = Document()

    # P0: both firstLine and firstLineChars
    p0 = doc.add_paragraph("P0: firstLine=360 + firstLineChars=200 (chars supersedes)")
    pPr0 = p0._p.get_or_add_pPr()
    ind0 = make_element("w:ind", {
        "w:left": "1440",
        "w:firstLine": "360",
        "w:firstLineChars": "200",
    })
    pPr0.append(ind0)

    # P1: both hanging and hangingChars
    p1 = doc.add_paragraph("P1: hanging=720 + hangingChars=100 (chars supersedes)")
    pPr1 = p1._p.get_or_add_pPr()
    ind1 = make_element("w:ind", {
        "w:left": "1440",
        "w:hanging": "720",
        "w:hangingChars": "100",
    })
    pPr1.append(ind1)

    # P2: firstLine only (control)
    p2 = doc.add_paragraph("P2: firstLine=360 only (control)")
    pPr2 = p2._p.get_or_add_pPr()
    ind2 = make_element("w:ind", {
        "w:left": "1440",
        "w:firstLine": "360",
    })
    pPr2.append(ind2)

    # P3: firstLineChars only (control)
    p3 = doc.add_paragraph("P3: firstLineChars=200 only (control)")
    pPr3 = p3._p.get_or_add_pPr()
    ind3 = make_element("w:ind", {
        "w:left": "1440",
        "w:firstLineChars": "200",
    })
    pPr3.append(ind3)

    save_fixture("firstline-chars-supersedes-firstline", doc, {
        "name": "firstline-chars-supersedes-firstline",
        "spec_ref": "ISO 29500-1 §17.3.1.12",
        "description": "firstLineChars supersedes firstLine, hangingChars supersedes hanging",
        "expected_behavior": (
            "P0: firstLineChars=200 supersedes firstLine=360 (both preserved). "
            "P1: hangingChars=100 supersedes hanging=720 (both preserved). "
            "P2: firstLine=360 only. "
            "P3: firstLineChars=200 only."
        ),
    })


# =========================================================================
# Fixture 9: Both beforeLines and before — both preserved for roundtrip
# =========================================================================

def make_both_before_lines_and_before() -> None:
    """§17.3.1.33: afterLines/beforeLines take precedence over after/before.

    "If the afterLines attribute or the afterAutoSpacing attribute is also
    specified, then this attribute value is ignored."

    Both values should be preserved in the model for roundtrip fidelity,
    but consumers should use beforeLines/afterLines when present.

    P0: before=480, beforeLines=100 — both preserved, beforeLines wins
    P1: after=480, afterLines=200 — both preserved, afterLines wins
    P2: before=480 only (control)
    P3: beforeLines=100 only (control)
    P4: before=0, beforeLines=100 — explicit zero before, beforeLines wins
    """
    doc = Document()

    p0 = doc.add_paragraph("P0: before=480 + beforeLines=100")
    pPr0 = p0._p.get_or_add_pPr()
    sp0 = make_element("w:spacing", {"w:before": "480", "w:beforeLines": "100"})
    pPr0.append(sp0)

    p1 = doc.add_paragraph("P1: after=480 + afterLines=200")
    pPr1 = p1._p.get_or_add_pPr()
    sp1 = make_element("w:spacing", {"w:after": "480", "w:afterLines": "200"})
    pPr1.append(sp1)

    p2 = doc.add_paragraph("P2: before=480 only (control)")
    pPr2 = p2._p.get_or_add_pPr()
    sp2 = make_element("w:spacing", {"w:before": "480"})
    pPr2.append(sp2)

    p3 = doc.add_paragraph("P3: beforeLines=100 only (control)")
    pPr3 = p3._p.get_or_add_pPr()
    sp3 = make_element("w:spacing", {"w:beforeLines": "100"})
    pPr3.append(sp3)

    p4 = doc.add_paragraph("P4: before=0 + beforeLines=100")
    pPr4 = p4._p.get_or_add_pPr()
    sp4 = make_element("w:spacing", {"w:before": "0", "w:beforeLines": "100"})
    pPr4.append(sp4)

    save_fixture("both-before-lines-and-before", doc, {
        "name": "both-before-lines-and-before",
        "spec_ref": "ISO 29500-1 §17.3.1.33",
        "description": "beforeLines/afterLines precedence over before/after with roundtrip",
        "expected_behavior": (
            "All values preserved for roundtrip. "
            "P0: before=480, beforeLines=100. "
            "P1: after=480, afterLines=200. "
            "P2: before=480. "
            "P3: beforeLines=100. "
            "P4: before=0, beforeLines=100."
        ),
    })


# =========================================================================
# Main
# =========================================================================

def main() -> None:
    print("Generating spacing-audit fixtures...")
    make_autospacing_override()
    make_contextual_spacing_explicit_false()
    make_pprchange_spacing_roundtrip()
    make_pprchange_indent_roundtrip()
    make_firstline_hanging_chars_exclusive()
    make_line_rule_default_auto()
    make_start_chars_supersedes_start()
    make_firstline_chars_supersedes_firstline()
    make_both_before_lines_and_before()
    print("Done.")


if __name__ == "__main__":
    main()

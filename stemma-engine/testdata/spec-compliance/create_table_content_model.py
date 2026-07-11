# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for table content model constraint testing.

Each fixture exercises a specific OOXML table content model constraint
that our implementation may or may not enforce.

Fixtures:
  - table-content-model/width-type-nil/       (MS-OI29500 §2.1.166)
  - table-content-model/vmerge-grid/          (§17.4.84)
  - table-content-model/grid-span-overflow/   (§17.4.17)
  - table-content-model/border-conflict/      (§17.4.66)

Run:  uv run create_table_content_model.py
"""

import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    """Return a fully-qualified wordprocessingml tag."""
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    """Create an OxmlElement with optional attributes."""
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc, metadata: dict) -> None:
    """Save a fixture DOCX under table-content-model/{name}/input.docx."""
    out = ROOT / "table-content-model" / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  table-content-model/{name}/")


def replace_tbl_child(tbl_pr, tag: str, new_el) -> None:
    """Remove existing child element with given tag from tblPr, then append new one."""
    existing = tbl_pr.find(qn(tag))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(new_el)


def set_cell_text(tc_el, text: str) -> None:
    """Replace all paragraphs in a tc element with a single paragraph containing text."""
    for p_el in tc_el.findall(w("p")):
        tc_el.remove(p_el)
    new_p = make_element("w:p")
    new_r = make_element("w:r")
    new_t = make_element("w:t")
    new_t.text = text
    new_r.append(new_t)
    new_p.append(new_r)
    tc_el.append(new_p)


def ensure_tc_pr(tc_el):
    """Return the tcPr element, creating it if absent."""
    tc_pr = tc_el.find(w("tcPr"))
    if tc_pr is None:
        tc_pr = make_element("w:tcPr")
        tc_el.insert(0, tc_pr)
    return tc_pr


# =========================================================================
# 1) Width type coercion when w=0 (MS-OI29500 §2.1.166)
# =========================================================================

def make_width_type_nil():
    """Table and cell widths with w=0 — spec says type should be coerced to auto.

    MS-OI29500 §2.1.166: If the w attribute is 0, the type attribute is
    treated as "auto" regardless of what is specified. This applies to both
    tblW and tcW elements.
    """
    doc = Document()
    doc.add_paragraph("Table with width w=0 type=pct (should coerce to auto).")

    tbl = doc.add_table(rows=1, cols=2)
    tbl.cell(0, 0).text = "Cell 1 (50% width)"
    tbl.cell(0, 1).text = "Cell 2 (w=0 dxa, should coerce to auto)"

    tbl_element = tbl._tbl

    # Set table width: w=0, type=pct — per MS-OI29500 §2.1.166, should be auto
    tbl_pr = tbl_element.tblPr
    replace_tbl_child(tbl_pr, "w:tblW", make_element("w:tblW", {
        "w:w": "0", "w:type": "pct",
    }))

    # Cell 1: normal width (50%)
    rows = tbl_element.findall(w("tr"))
    cells = rows[0].findall(w("tc"))

    tc_pr_0 = ensure_tc_pr(cells[0])
    existing = tc_pr_0.find(qn("w:tcW"))
    if existing is not None:
        tc_pr_0.remove(existing)
    tc_pr_0.append(make_element("w:tcW", {"w:w": "5000", "w:type": "pct"}))

    # Cell 2: w=0, type=dxa — per MS-OI29500 §2.1.166, should be auto
    tc_pr_1 = ensure_tc_pr(cells[1])
    existing = tc_pr_1.find(qn("w:tcW"))
    if existing is not None:
        tc_pr_1.remove(existing)
    tc_pr_1.append(make_element("w:tcW", {"w:w": "0", "w:type": "dxa"}))

    save_fixture("width-type-nil", doc, {
        "spec": "MS-OI29500 §2.1.166",
        "description": "Table w=0/type=pct and cell w=0/type=dxa — spec says treat as auto",
        "table_width": {"w": 0, "type": "pct", "expected_coercion": "auto"},
        "cell_widths": [
            {"w": 5000, "type": "pct", "expected_coercion": None},
            {"w": 0, "type": "dxa", "expected_coercion": "auto"},
        ],
    })


# =========================================================================
# 2) vMerge grid alignment (§17.4.84)
# =========================================================================

def make_vmerge_grid():
    """3x3 table with vMerge + gridSpan — tests that continuation gridSpan must match restart.

    ISO 29500-1 §17.4.84: A cell with vMerge="continue" must have the same
    gridSpan as the restart cell above it. If the gridSpan differs, the
    merge is non-conformant.

    Layout:
      Row 0: [A gridSpan=2 vMerge=restart] [B]
      Row 1: [C gridSpan=2 vMerge=continue] [D]    (matching — correct)
      Row 2: [E gridSpan=1 vMerge=continue] [F] [G] (mismatched — non-conformant)
    """
    doc = Document()
    doc.add_paragraph("Table with vMerge gridSpan mismatch in row 2.")

    tbl = doc.add_table(rows=3, cols=3)
    tbl_element = tbl._tbl
    rows = tbl_element.findall(w("tr"))

    # -- Row 0 -------------------------------------------------------
    cells0 = rows[0].findall(w("tc"))

    # Cell (0,0): gridSpan=2, vMerge=restart
    tc_pr = ensure_tc_pr(cells0[0])
    tc_pr.append(make_element("w:gridSpan", {"w:val": "2"}))
    tc_pr.append(make_element("w:vMerge", {"w:val": "restart"}))
    set_cell_text(cells0[0], "R0C0 (span=2, vMerge restart)")

    # Remove cell 1 (consumed by gridSpan)
    rows[0].remove(cells0[1])

    # Cell (0,2): normal
    set_cell_text(cells0[2], "R0C2")

    # -- Row 1 -------------------------------------------------------
    cells1 = rows[1].findall(w("tc"))

    # Cell (1,0): gridSpan=2, vMerge=continue (matching — correct)
    tc_pr = ensure_tc_pr(cells1[0])
    tc_pr.append(make_element("w:gridSpan", {"w:val": "2"}))
    tc_pr.append(make_element("w:vMerge"))  # bare = continue
    set_cell_text(cells1[0], "R1C0 (span=2, vMerge continue)")

    # Remove cell 1 (consumed by gridSpan)
    rows[1].remove(cells1[1])

    # Cell (1,2): normal
    set_cell_text(cells1[2], "R1C2")

    # -- Row 2 -------------------------------------------------------
    cells2 = rows[2].findall(w("tc"))

    # Cell (2,0): gridSpan=1 (default), vMerge=continue (MISMATCHED — non-conformant)
    tc_pr = ensure_tc_pr(cells2[0])
    tc_pr.append(make_element("w:vMerge"))  # bare = continue
    # No gridSpan attribute → defaults to 1, but restart above has gridSpan=2
    set_cell_text(cells2[0], "R2C0 (span=1, vMerge continue MISMATCH)")

    # Cells (2,1) and (2,2): normal
    set_cell_text(cells2[1], "R2C1")
    set_cell_text(cells2[2], "R2C2")

    save_fixture("vmerge-grid", doc, {
        "spec": "ISO 29500-1 §17.4.84",
        "description": "3x3 table: rows 0-1 have matching vMerge gridSpan=2, "
                       "row 2 has mismatched gridSpan=1 on vMerge continue",
        "rows": [
            {"cells": ["R0C0 (gridSpan=2, vMerge=restart)", "R0C2"]},
            {"cells": ["R1C0 (gridSpan=2, vMerge=continue)", "R1C2"]},
            {"cells": ["R2C0 (gridSpan=1, vMerge=continue MISMATCH)", "R2C1", "R2C2"]},
        ],
    })


# =========================================================================
# 3) gridSpan overflow (§17.4.17)
# =========================================================================

def make_grid_span_overflow():
    """Table with gridSpan exceeding the declared grid width.

    ISO 29500-1 §17.4.17: gridSpan specifies the number of grid columns
    the cell spans. If gridSpan exceeds the number of columns in tblGrid,
    the document is non-conformant.

    Layout: tblGrid has 3 columns, but cell (0,0) has gridSpan=5.
    """
    doc = Document()
    doc.add_paragraph("Table with gridSpan=5 but only 3 grid columns.")

    # Build table XML manually for precise control over tblGrid
    tbl_el = OxmlElement("w:tbl")

    # tblPr
    tbl_pr = make_element("w:tblPr")
    tbl_w = make_element("w:tblW", {"w:w": "0", "w:type": "auto"})
    tbl_pr.append(tbl_w)
    tbl_borders = make_element("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = make_element(f"w:{edge}", {
            "w:val": "single", "w:sz": "4", "w:space": "0", "w:color": "000000",
        })
        tbl_borders.append(b)
    tbl_pr.append(tbl_borders)
    tbl_el.append(tbl_pr)

    # tblGrid: 3 columns
    tbl_grid = make_element("w:tblGrid")
    for _ in range(3):
        tbl_grid.append(make_element("w:gridCol", {"w:w": "3000"}))
    tbl_el.append(tbl_grid)

    # Row 0: single cell with gridSpan=5 (exceeds grid width of 3)
    tr0 = make_element("w:tr")
    tc0 = make_element("w:tc")
    tc_pr0 = make_element("w:tcPr")
    tc_pr0.append(make_element("w:gridSpan", {"w:val": "5"}))
    tc_pr0.append(make_element("w:tcW", {"w:w": "9000", "w:type": "dxa"}))
    tc0.append(tc_pr0)
    p0 = make_element("w:p")
    r0 = make_element("w:r")
    t0 = make_element("w:t")
    t0.text = "R0C0 (gridSpan=5, grid has 3 cols)"
    r0.append(t0)
    p0.append(r0)
    tc0.append(p0)
    tr0.append(tc0)
    tbl_el.append(tr0)

    # Row 1: 3 normal cells
    tr1 = make_element("w:tr")
    for ci in range(3):
        tc = make_element("w:tc")
        tc_pr = make_element("w:tcPr")
        tc_pr.append(make_element("w:tcW", {"w:w": "3000", "w:type": "dxa"}))
        tc.append(tc_pr)
        p = make_element("w:p")
        r = make_element("w:r")
        t = make_element("w:t")
        t.text = f"R1C{ci}"
        r.append(t)
        p.append(r)
        tc.append(p)
        tr1.append(tc)
    tbl_el.append(tr1)

    # Insert the table into the document body
    body = doc.element.body
    body.append(tbl_el)

    # Trailing paragraph (required)
    body.append(make_element("w:p"))

    save_fixture("grid-span-overflow", doc, {
        "spec": "ISO 29500-1 §17.4.17",
        "description": "Table with tblGrid=3 columns but cell gridSpan=5 (non-conformant)",
        "grid_columns": 3,
        "overflowing_grid_span": 5,
    })


# =========================================================================
# 4) Border conflict resolution (§17.4.66)
# =========================================================================

def make_border_conflict():
    """2x2 table with conflicting adjacent cell borders and tblCellSpacing=0.

    ISO 29500-1 §17.4.66: When tblCellSpacing is 0 (collapsed borders),
    adjacent cells may specify conflicting borders. The spec defines a
    conflict resolution algorithm based on border weight and style.

    Cell (0,0) right border: red single sz=12
    Cell (0,1) left border: blue double sz=4
    """
    doc = Document()
    doc.add_paragraph("Table with conflicting adjacent cell borders (spacing=0).")

    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "R0C0 (right=red single)"
    tbl.cell(0, 1).text = "R0C1 (left=blue double)"
    tbl.cell(1, 0).text = "R1C0"
    tbl.cell(1, 1).text = "R1C1"

    tbl_element = tbl._tbl

    # Set tblCellSpacing to 0 (collapsed borders)
    tbl_pr = tbl_element.tblPr
    tbl_pr.append(make_element("w:tblCellSpacing", {"w:w": "0", "w:type": "dxa"}))

    rows = tbl_element.findall(w("tr"))
    cells0 = rows[0].findall(w("tc"))

    # Cell (0,0): right border = red single sz=12
    tc_pr_00 = ensure_tc_pr(cells0[0])
    borders_00 = make_element("w:tcBorders")
    borders_00.append(make_element("w:right", {
        "w:val": "single", "w:sz": "12", "w:space": "0", "w:color": "FF0000",
    }))
    tc_pr_00.append(borders_00)

    # Cell (0,1): left border = blue double sz=4
    tc_pr_01 = ensure_tc_pr(cells0[1])
    borders_01 = make_element("w:tcBorders")
    borders_01.append(make_element("w:left", {
        "w:val": "double", "w:sz": "4", "w:space": "0", "w:color": "0000FF",
    }))
    tc_pr_01.append(borders_01)

    save_fixture("border-conflict", doc, {
        "spec": "ISO 29500-1 §17.4.66",
        "description": "2x2 table with tblCellSpacing=0 and conflicting adjacent borders: "
                       "cell (0,0) right=red/single/12, cell (0,1) left=blue/double/4",
        "cell_spacing": 0,
        "conflicting_borders": {
            "cell_0_0_right": {"style": "single", "size": 12, "color": "FF0000"},
            "cell_0_1_left": {"style": "double", "size": 4, "color": "0000FF"},
        },
    })


# =========================================================================
# Entry point
# =========================================================================

def main():
    print("Generating table-content-model fixtures...")
    make_width_type_nil()
    make_vmerge_grid()
    make_grid_span_overflow()
    make_border_conflict()
    print("Done.")


if __name__ == "__main__":
    main()

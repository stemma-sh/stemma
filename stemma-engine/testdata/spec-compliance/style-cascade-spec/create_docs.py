# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for style cascade spec-compliance tests.

Focuses on:
1. Linked style + char basedOn chain (§17.7.4.6 + §17.7.4.3)
2. Toggle XOR across table + paragraph levels (§17.7.3)
3. lineRule absent default to "auto" (§17.3.1.33)
4. Table conditional formatting corner precedence (§17.7.6)
5. Linked style double-apply scenario (§17.7.4.6)

Run:  cd stemma-engine/testdata/spec-compliance/style-cascade-spec && uv run create_docs.py
"""

import json
from pathlib import Path
from lxml import etree

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  style-cascade-spec/{name}/")


# =========================================================================
# 1. Linked style with char style basedOn chain (§17.7.4.6 + §17.7.4.3)
#
# When a paragraph style has w:link to a character style, and that char
# style has a basedOn chain, properties inherited from the char style's
# basedOn parent should be available to runs using the paragraph style.
#
# Our code uses overlay_marks(&mut base, &linked_raw.marks) where
# linked_raw.marks is just the char style's own explicit marks — NOT its
# resolved basedOn chain. So properties the char style inherits from its
# parent are lost.
# =========================================================================

def make_linked_char_based_on():
    """
    Setup:
    - CharBase: character style with italic=on, color=FF0000
    - LinkedChar: character style, basedOn CharBase, bold=on (inherits italic, color)
    - LinkedPara: paragraph style, w:link to LinkedChar, font_size=32

    Para 1: LinkedPara style → runs should get:
      - bold=on (from LinkedChar directly)
      - italic=on (from CharBase via LinkedChar's basedOn chain)
      - color=FF0000 (from CharBase via LinkedChar's basedOn chain)
      - font_size=32 (from LinkedPara's own rPr)

    Bug hypothesis: italic and color from CharBase will be lost because
    the code only overlays LinkedChar's own raw marks (bold=on) and doesn't
    include its basedOn parent's properties.
    """
    doc = Document()
    styles_el = doc.styles.element

    # CharBase: character style with italic=on, color=FF0000
    style_char_base = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "CharBase",
    })
    style_char_base.append(make_element("w:name", {"w:val": "Char Base"}))
    base_rpr = make_element("w:rPr")
    base_rpr.append(make_element("w:i"))
    base_rpr.append(make_element("w:color", {"w:val": "FF0000"}))
    style_char_base.append(base_rpr)
    styles_el.append(style_char_base)

    # LinkedChar: character style, basedOn CharBase, bold=on
    style_linked_char = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "LinkedChar",
    })
    style_linked_char.append(make_element("w:name", {"w:val": "Linked Char"}))
    style_linked_char.append(make_element("w:basedOn", {"w:val": "CharBase"}))
    style_linked_char.append(make_element("w:link", {"w:val": "LinkedPara"}))
    lc_rpr = make_element("w:rPr")
    lc_rpr.append(make_element("w:b"))
    style_linked_char.append(lc_rpr)
    styles_el.append(style_linked_char)

    # LinkedPara: paragraph style, link to LinkedChar, font_size=32
    style_linked_para = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "LinkedPara",
    })
    style_linked_para.append(make_element("w:name", {"w:val": "Linked Para"}))
    style_linked_para.append(make_element("w:link", {"w:val": "LinkedChar"}))
    lp_rpr = make_element("w:rPr")
    lp_rpr.append(make_element("w:sz", {"w:val": "32"}))
    style_linked_para.append(lp_rpr)
    styles_el.append(style_linked_para)

    # Para 1: LinkedPara style — run should get all properties
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "LinkedPara"}))
    p1.add_run("Linked para with char basedOn chain")

    # Para 2: run with LinkedChar applied directly as char style
    # Should get bold + italic + color (from full basedOn resolution)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Run with LinkedChar directly")
    r2_rpr = r2._element.get_or_add_rPr()
    r2_rpr.insert(0, make_element("w:rStyle", {"w:val": "LinkedChar"}))

    # Para 3: run with CharBase applied — italic + color only (no bold)
    p3 = doc.add_paragraph()
    r3 = p3.add_run("Run with CharBase only")
    r3_rpr = r3._element.get_or_add_rPr()
    r3_rpr.insert(0, make_element("w:rStyle", {"w:val": "CharBase"}))

    save_fixture("linked-char-basedOn", doc, {
        "spec": "ISO 29500-1 §17.7.4.6 + §17.7.4.3",
        "description": "Linked char style with basedOn chain: inherited props must reach para style runs",
        "expected": {
            "para_1_linkedpara": {
                "bold": True,
                "italic": True,
                "color": "FF0000",
                "font_size": 32,
            },
            "para_2_linkedchar_direct": {
                "bold": True,
                "italic": True,
                "color": "FF0000",
            },
            "para_3_charbase_only": {
                "bold": False,
                "italic": True,
                "color": "FF0000",
            },
        },
    })


# =========================================================================
# 2. lineRule absent with style having lineRule=exact (§17.3.1.33)
#
# When a paragraph has a style with lineRule=exact, and a direct w:spacing
# sets line=240 but omits lineRule, the spec says lineRule should default
# to "auto" (not be inherited from the style).
#
# This fixture tests the STYLE-ONLY case: a paragraph style sets
# w:spacing line=240 but omits lineRule. Per §17.3.1.33 the default
# should be "auto".
# =========================================================================

def make_line_rule_absent():
    """
    Setup:
    - StyleNoLineRule: paragraph style with w:spacing line="360" but NO
      lineRule attribute. Per §17.3.1.33, lineRule defaults to "auto".
    - StyleExactRule: paragraph style with w:spacing line="360" lineRule="exact".
    - StyleNoLineRuleBasedOnExact: paragraph style basedOn StyleExactRule,
      sets line="240" but omits lineRule. Per §17.3.1.33, the omitted
      lineRule should reset to "auto" — NOT inherit "exact" from parent.

    Para 1: StyleNoLineRule — lineRule should be "auto" (default).
    Para 2: StyleExactRule — lineRule should be "exact" (explicit).
    Para 3: StyleNoLineRuleBasedOnExact — lineRule should be "auto"
            because the child's w:spacing sets line but omits lineRule.
    """
    doc = Document()
    styles_el = doc.styles.element

    # StyleNoLineRule: line=360, no lineRule
    style1 = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "StyleNoLineRule",
    })
    style1.append(make_element("w:name", {"w:val": "Style No LineRule"}))
    ppr1 = make_element("w:pPr")
    ppr1.append(make_element("w:spacing", {"w:line": "360"}))
    style1.append(ppr1)
    styles_el.append(style1)

    # StyleExactRule: line=360, lineRule=exact
    style2 = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "StyleExactRule",
    })
    style2.append(make_element("w:name", {"w:val": "Style Exact Rule"}))
    ppr2 = make_element("w:pPr")
    ppr2.append(make_element("w:spacing", {"w:line": "360", "w:lineRule": "exact"}))
    style2.append(ppr2)
    styles_el.append(style2)

    # StyleNoLineRuleBasedOnExact: basedOn StyleExactRule, line=240, no lineRule
    style3 = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "StyleNoLineRuleBasedOnExact",
    })
    style3.append(make_element("w:name", {"w:val": "Style No LineRule BasedOn Exact"}))
    style3.append(make_element("w:basedOn", {"w:val": "StyleExactRule"}))
    ppr3 = make_element("w:pPr")
    # Sets line but omits lineRule — per spec, should reset to "auto"
    ppr3.append(make_element("w:spacing", {"w:line": "240"}))
    style3.append(ppr3)
    styles_el.append(style3)

    # Para 1: StyleNoLineRule
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "StyleNoLineRule"}))
    p1.add_run("No lineRule style text")

    # Para 2: StyleExactRule
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    p2_ppr.append(make_element("w:pStyle", {"w:val": "StyleExactRule"}))
    p2.add_run("Exact lineRule style text")

    # Para 3: StyleNoLineRuleBasedOnExact
    p3 = doc.add_paragraph()
    p3_ppr = p3._element.get_or_add_pPr()
    p3_ppr.append(make_element("w:pStyle", {"w:val": "StyleNoLineRuleBasedOnExact"}))
    p3.add_run("BasedOn exact but no lineRule text")

    save_fixture("line-rule-absent", doc, {
        "spec": "ISO 29500-1 §17.3.1.33",
        "description": "lineRule defaults to 'auto' when omitted but line is present",
        "expected": {
            "para_1": {"line": 360, "lineRule": "auto"},
            "para_2": {"line": 360, "lineRule": "exact"},
            "para_3": {
                "line": 240,
                "lineRule": "auto (NOT exact from parent — child's w:spacing resets lineRule)",
            },
        },
    })


# =========================================================================
# 3. Table conditional formatting corner cell precedence (§17.7.6)
#
# When a cell is at the intersection of firstRow and firstCol,
# the NwCell conditional (if defined) has highest precedence.
# When NwCell is NOT defined, firstRow should win over firstCol
# per the precedence order.
# =========================================================================

def make_table_conditional_corner():
    """
    Setup:
    - TableStyle with:
      - firstRow: shading fill=4472C4 (blue)
      - firstCol: shading fill=FF0000 (red)
      - NwCell: shading fill=00FF00 (green) — corner override
    - 3x3 table with tblLook enabling firstRow + firstCol + no corners initially

    Table 1: tblLook with firstRow + firstCol enabled, uses style WITH NwCell.
      Cell(0,0) should get green (NwCell highest precedence).
      Cell(0,1) should get blue (firstRow).
      Cell(1,0) should get red (firstCol).

    Table 2: Same but style WITHOUT NwCell conditional.
      Cell(0,0) should get blue (firstRow > firstCol per §17.7.6).
    """
    doc = Document()
    styles_el = doc.styles.element

    # TableStyle1: with NwCell conditional
    ts1 = make_element("w:style", {
        "w:type": "table",
        "w:styleId": "CornerTestWithNw",
    })
    ts1.append(make_element("w:name", {"w:val": "Corner Test With Nw"}))
    # firstRow conditional: blue shading
    tsp_fr = make_element("w:tblStylePr", {"w:type": "firstRow"})
    tc_pr_fr = make_element("w:tcPr")
    tc_pr_fr.append(make_element("w:shd", {
        "w:val": "clear", "w:color": "auto", "w:fill": "4472C4",
    }))
    tsp_fr.append(tc_pr_fr)
    # Also set bold for firstRow runs
    rpr_fr = make_element("w:rPr")
    rpr_fr.append(make_element("w:b"))
    tsp_fr.append(rpr_fr)
    ts1.append(tsp_fr)

    # firstCol conditional: red shading
    tsp_fc = make_element("w:tblStylePr", {"w:type": "firstCol"})
    tc_pr_fc = make_element("w:tcPr")
    tc_pr_fc.append(make_element("w:shd", {
        "w:val": "clear", "w:color": "auto", "w:fill": "FF0000",
    }))
    tsp_fc.append(tc_pr_fc)
    ts1.append(tsp_fc)

    # NwCell conditional: green shading (highest precedence for corner)
    tsp_nw = make_element("w:tblStylePr", {"w:type": "nwCell"})
    tc_pr_nw = make_element("w:tcPr")
    tc_pr_nw.append(make_element("w:shd", {
        "w:val": "clear", "w:color": "auto", "w:fill": "00FF00",
    }))
    tsp_nw.append(tc_pr_nw)
    ts1.append(tsp_nw)

    styles_el.append(ts1)

    # TableStyle2: WITHOUT NwCell conditional
    ts2 = make_element("w:style", {
        "w:type": "table",
        "w:styleId": "CornerTestNoNw",
    })
    ts2.append(make_element("w:name", {"w:val": "Corner Test No Nw"}))
    # firstRow: blue
    tsp_fr2 = make_element("w:tblStylePr", {"w:type": "firstRow"})
    tc_pr_fr2 = make_element("w:tcPr")
    tc_pr_fr2.append(make_element("w:shd", {
        "w:val": "clear", "w:color": "auto", "w:fill": "4472C4",
    }))
    tsp_fr2.append(tc_pr_fr2)
    ts2.append(tsp_fr2)
    # firstCol: red
    tsp_fc2 = make_element("w:tblStylePr", {"w:type": "firstCol"})
    tc_pr_fc2 = make_element("w:tcPr")
    tc_pr_fc2.append(make_element("w:shd", {
        "w:val": "clear", "w:color": "auto", "w:fill": "FF0000",
    }))
    tsp_fc2.append(tc_pr_fc2)
    ts2.append(tsp_fc2)
    styles_el.append(ts2)

    # Table 1: CornerTestWithNw
    table1 = doc.add_table(rows=3, cols=3)
    tbl1_pr = table1._element.find(qn("w:tblPr"))
    if tbl1_pr is None:
        tbl1_pr = make_element("w:tblPr")
        table1._element.insert(0, tbl1_pr)
    # Set table style
    for existing in tbl1_pr.findall(qn("w:tblStyle")):
        tbl1_pr.remove(existing)
    tbl1_pr.insert(0, make_element("w:tblStyle", {"w:val": "CornerTestWithNw"}))
    # Set tblLook: firstRow=1, firstCol=1
    for existing in tbl1_pr.findall(qn("w:tblLook")):
        tbl1_pr.remove(existing)
    tbl1_pr.append(make_element("w:tblLook", {
        "w:val": "00A0",
        "w:firstRow": "1",
        "w:lastRow": "0",
        "w:firstColumn": "1",
        "w:lastColumn": "0",
        "w:noHBand": "1",
        "w:noVBand": "1",
    }))
    # Fill in cell text
    for ri, row in enumerate(table1.rows):
        for ci, cell in enumerate(row.cells):
            cell.text = f"T1 R{ri}C{ci}"

    # Add a separator paragraph
    doc.add_paragraph("---")

    # Table 2: CornerTestNoNw
    table2 = doc.add_table(rows=3, cols=3)
    tbl2_pr = table2._element.find(qn("w:tblPr"))
    if tbl2_pr is None:
        tbl2_pr = make_element("w:tblPr")
        table2._element.insert(0, tbl2_pr)
    for existing in tbl2_pr.findall(qn("w:tblStyle")):
        tbl2_pr.remove(existing)
    tbl2_pr.insert(0, make_element("w:tblStyle", {"w:val": "CornerTestNoNw"}))
    for existing in tbl2_pr.findall(qn("w:tblLook")):
        tbl2_pr.remove(existing)
    tbl2_pr.append(make_element("w:tblLook", {
        "w:val": "00A0",
        "w:firstRow": "1",
        "w:lastRow": "0",
        "w:firstColumn": "1",
        "w:lastColumn": "0",
        "w:noHBand": "1",
        "w:noVBand": "1",
    }))
    for ri, row in enumerate(table2.rows):
        for ci, cell in enumerate(row.cells):
            cell.text = f"T2 R{ri}C{ci}"

    save_fixture("table-conditional-corner", doc, {
        "spec": "ISO 29500-1 §17.7.6 + MS-OI29500 §17.4.54(a)",
        "description": "Corner cell conditional formatting precedence",
        "expected": {
            "table_1_with_nw": {
                "cell_0_0": "green (NwCell > firstRow > firstCol)",
                "cell_0_1": "blue (firstRow only)",
                "cell_1_0": "red (firstCol only)",
                "cell_1_1": "no conditional shading",
            },
            "table_2_no_nw": {
                "cell_0_0": "blue (firstRow > firstCol, no NwCell defined)",
                "cell_0_1": "blue (firstRow only)",
                "cell_1_0": "red (firstCol only)",
            },
        },
    })


# =========================================================================
# 4. Linked style double-apply: run explicitly uses linked char style
#    inside paragraph with linked para style (§17.7.4.6)
#
# When a paragraph uses LinkedPara and a run within it explicitly
# references LinkedChar (the same linked pair), the char style should
# NOT double-apply or create unexpected interactions.
# =========================================================================

def make_linked_style_double_apply():
    """
    Setup:
    - LinkedPara2: paragraph style, bold=on, font_size=28, link to LinkedChar2
    - LinkedChar2: character style, bold=on, color=0000FF, link to LinkedPara2
    - Para 1: LinkedPara2 style
      - Run A: no explicit char style → inherits from linked pair
      - Run B: explicit rStyle=LinkedChar2 → should be same as Run A

    If the resolution treats the explicit char style application differently
    from the implicit linked resolution, Run A and Run B may diverge.
    """
    doc = Document()
    styles_el = doc.styles.element

    # LinkedChar2
    style_char = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "LinkedChar2",
    })
    style_char.append(make_element("w:name", {"w:val": "Linked Char 2"}))
    style_char.append(make_element("w:link", {"w:val": "LinkedPara2"}))
    char_rpr = make_element("w:rPr")
    char_rpr.append(make_element("w:b"))
    char_rpr.append(make_element("w:color", {"w:val": "0000FF"}))
    style_char.append(char_rpr)
    styles_el.append(style_char)

    # LinkedPara2
    style_para = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "LinkedPara2",
    })
    style_para.append(make_element("w:name", {"w:val": "Linked Para 2"}))
    style_para.append(make_element("w:link", {"w:val": "LinkedChar2"}))
    para_rpr = make_element("w:rPr")
    para_rpr.append(make_element("w:b"))
    para_rpr.append(make_element("w:sz", {"w:val": "28"}))
    style_para.append(para_rpr)
    styles_el.append(style_para)

    # Para 1: LinkedPara2 style
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "LinkedPara2"}))

    # Run A: no explicit char style
    r1 = p1.add_run("Implicit linked run")

    # Run B: explicit rStyle=LinkedChar2
    r2 = p1.add_run("Explicit linked char run")
    r2_rpr = r2._element.get_or_add_rPr()
    r2_rpr.insert(0, make_element("w:rStyle", {"w:val": "LinkedChar2"}))

    save_fixture("linked-double-apply", doc, {
        "spec": "ISO 29500-1 §17.7.4.6",
        "description": "Linked style: implicit via para vs explicit char style should match",
        "expected": {
            "run_a_implicit": {
                "bold": True,
                "color": "0000FF",
                "font_size": 28,
            },
            "run_b_explicit": {
                "bold": True,
                "color": "0000FF",
                "font_size": 28,
                "note": "Should match Run A — same linked pair"
            },
        },
    })


# =========================================================================
# 5. Toggle property with paragraph style bold=off (explicit false)
#    vs bold absent (inherit) in basedOn chain (§17.7.3)
#
# When a paragraph style explicitly sets bold=false (w:b w:val="0"),
# it should be different from not setting bold at all (inheriting from
# basedOn parent). This tests whether our code distinguishes
# MarkValue::Off from MarkValue::Inherit within the basedOn chain.
# =========================================================================

def make_toggle_explicit_off_vs_inherit():
    """
    Setup:
    - BoldRoot: paragraph style, bold=on, italic=on
    - ExplicitOff: paragraph style, basedOn BoldRoot, bold=off (explicit w:val="0")
    - InheritBold: paragraph style, basedOn BoldRoot, (no bold element — inherits)

    Para 1: ExplicitOff → bold should be OFF (explicit w:val="0" overrides parent)
    Para 2: InheritBold → bold should be ON (inherited from BoldRoot)
    Both should have italic=ON (inherited from BoldRoot, not overridden).
    """
    doc = Document()
    styles_el = doc.styles.element

    # BoldRoot: bold=on, italic=on
    style_root = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "BoldRoot",
    })
    style_root.append(make_element("w:name", {"w:val": "Bold Root"}))
    root_rpr = make_element("w:rPr")
    root_rpr.append(make_element("w:b"))
    root_rpr.append(make_element("w:i"))
    style_root.append(root_rpr)
    styles_el.append(style_root)

    # ExplicitOff: basedOn BoldRoot, bold=off (val="0")
    style_off = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "ExplicitOff",
    })
    style_off.append(make_element("w:name", {"w:val": "Explicit Off"}))
    style_off.append(make_element("w:basedOn", {"w:val": "BoldRoot"}))
    off_rpr = make_element("w:rPr")
    off_rpr.append(make_element("w:b", {"w:val": "0"}))
    style_off.append(off_rpr)
    styles_el.append(style_off)

    # InheritBold: basedOn BoldRoot, no bold element (inherits)
    style_inherit = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "InheritBold",
    })
    style_inherit.append(make_element("w:name", {"w:val": "Inherit Bold"}))
    style_inherit.append(make_element("w:basedOn", {"w:val": "BoldRoot"}))
    # No rPr — inherits everything from BoldRoot
    styles_el.append(style_inherit)

    # Para 1: ExplicitOff
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "ExplicitOff"}))
    p1.add_run("Explicit bold off run")

    # Para 2: InheritBold
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    p2_ppr.append(make_element("w:pStyle", {"w:val": "InheritBold"}))
    p2.add_run("Inherited bold on run")

    # Para 3: BoldRoot directly
    p3 = doc.add_paragraph()
    p3_ppr = p3._element.get_or_add_pPr()
    p3_ppr.append(make_element("w:pStyle", {"w:val": "BoldRoot"}))
    p3.add_run("Bold root run")

    save_fixture("toggle-explicit-off-vs-inherit", doc, {
        "spec": "ISO 29500-1 §17.7.3 + §17.7.4.3",
        "description": "bold=off (explicit) vs bold absent (inherit) in basedOn chain",
        "expected": {
            "para_1_explicit_off": {"bold": False, "italic": True},
            "para_2_inherit": {"bold": True, "italic": True},
            "para_3_root": {"bold": True, "italic": True},
        },
    })


# =========================================================================
# 6. Multiple toggle properties with different patterns across hierarchy
#    levels (§17.7.3)
#
# Tests that the code handles toggle properties independently when
# different properties have different on/off patterns across the
# paragraph style and character style levels.
# =========================================================================

def make_toggle_mixed_across_levels():
    """
    Setup:
    - MixedPara: paragraph style, bold=on, italic=off, caps=on
    - MixedChar: character style, bold=off, italic=on (caps not set → inherit)

    Para 1: MixedPara + MixedChar →
      - bold: para=on, char=off → char overrides → off (MS-OI29500 §2.1.258)
      - italic: para=off, char=on → char overrides → on
      - caps: para=on, char not set → para applies → on

    This tests per-property independence of override/reset semantics.
    """
    doc = Document()
    styles_el = doc.styles.element

    # MixedPara: bold=on, italic=off, caps=on
    style_para = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "MixedPara",
    })
    style_para.append(make_element("w:name", {"w:val": "Mixed Para"}))
    para_rpr = make_element("w:rPr")
    para_rpr.append(make_element("w:b"))
    para_rpr.append(make_element("w:i", {"w:val": "0"}))
    para_rpr.append(make_element("w:caps"))
    style_para.append(para_rpr)
    styles_el.append(style_para)

    # MixedChar: bold=off, italic=on (caps not set)
    style_char = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "MixedChar",
    })
    style_char.append(make_element("w:name", {"w:val": "Mixed Char"}))
    char_rpr = make_element("w:rPr")
    char_rpr.append(make_element("w:b", {"w:val": "0"}))
    char_rpr.append(make_element("w:i"))
    style_char.append(char_rpr)
    styles_el.append(style_char)

    # Para 1: MixedPara + MixedChar
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "MixedPara"}))
    r1 = p1.add_run("Mixed toggle patterns")
    r1_rpr = r1._element.get_or_add_rPr()
    r1_rpr.insert(0, make_element("w:rStyle", {"w:val": "MixedChar"}))

    # Para 2: MixedPara only (no char style)
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    p2_ppr.append(make_element("w:pStyle", {"w:val": "MixedPara"}))
    p2.add_run("Para style only mixed")

    save_fixture("toggle-mixed-across-levels", doc, {
        "spec": "ISO 29500-1 §17.7.3 + MS-OI29500 §2.1.258",
        "description": "Multiple toggle properties with different on/off across levels",
        "expected": {
            "para_1_both_styles": {
                "bold": "OFF (char overrides para per MS-OI29500 §2.1.258)",
                "italic": "ON (char overrides para)",
                "caps": "ON (only para sets it, char inherits → para's value persists)",
            },
            "para_2_para_only": {
                "bold": "ON",
                "italic": "OFF (explicit val=0)",
                "caps": "ON",
            },
        },
    })


# =========================================================================
# Main
# =========================================================================

if __name__ == "__main__":
    print("Generating style-cascade-spec fixtures:")
    make_linked_char_based_on()
    make_line_rule_absent()
    make_table_conditional_corner()
    make_linked_style_double_apply()
    make_toggle_explicit_off_vs_inherit()
    make_toggle_mixed_across_levels()
    print("Done.")

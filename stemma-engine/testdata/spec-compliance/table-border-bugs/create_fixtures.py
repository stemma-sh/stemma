# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for table border bug tests (ISO 29500-1 SS17.4.66,
MS-OI29500 SS17.4.55).

Run:  mise exec -- python3 stemma-engine/testdata/spec-compliance/table-border-bugs/create_fixtures.py
"""

import json
from pathlib import Path
from lxml import etree

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc: DocxDocument, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  table-border-bugs/{name}/")


def make_table_with_borders(doc, rows=2, cols=2, cell_texts=None):
    """Create a table, return (tbl_element, tblPr_element)."""
    tbl = doc.add_table(rows=rows, cols=cols)
    if cell_texts:
        for r in range(rows):
            for c in range(cols):
                if r < len(cell_texts) and c < len(cell_texts[r]):
                    tbl.cell(r, c).text = cell_texts[r][c]
    tbl_el = tbl._tbl
    tblPr = tbl_el.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_el.insert(0, tblPr)
    return tbl_el, tblPr


def add_tbl_borders(tblPr, edges, val="single", sz="12", color="FF0000", space="0"):
    """Add tblBorders to tblPr."""
    tbl_borders = make_element("w:tblBorders")
    for edge in edges:
        border = make_element(f"w:{edge}", {
            "w:val": val, "w:sz": sz, "w:color": color, "w:space": space,
        })
        tbl_borders.append(border)
    tblPr.append(tbl_borders)


def add_tc_borders(tbl_el, row, col, edges_spec):
    """Add tcBorders to a specific cell. edges_spec is a list of
    (edge_name, val, sz, color, space) tuples."""
    tr = tbl_el.findall(w("tr"))[row]
    tc = tr.findall(w("tc"))[col]
    tcPr = tc.find(w("tcPr"))
    if tcPr is None:
        tcPr = make_element("w:tcPr")
        tc.insert(0, tcPr)
    tc_borders = make_element("w:tcBorders")
    for edge, val, sz, color, space in edges_spec:
        border = make_element(f"w:{edge}", {
            "w:val": val, "w:sz": sz, "w:color": color, "w:space": space,
        })
        tc_borders.append(border)
    tcPr.append(tc_borders)


def add_tbl_grid(tbl_el, col_widths):
    """Add or replace tblGrid with specific column widths."""
    existing = tbl_el.find(w("tblGrid"))
    if existing is not None:
        tbl_el.remove(existing)
    grid = make_element("w:tblGrid")
    for cw in col_widths:
        grid.append(make_element("w:gridCol", {"w:w": str(cw)}))
    tblPr = tbl_el.find(w("tblPr"))
    if tblPr is not None:
        idx = list(tbl_el).index(tblPr) + 1
        tbl_el.insert(idx, grid)
    else:
        tbl_el.insert(0, grid)


# =====================================================================
# Fixture 1: Table->cell border fallback (Bug 1)
# =====================================================================

def make_border_fallback():
    """SS17.4.66: Cells without tcBorders inherit from tblBorders.

    2x2 table. tblBorders has single red sz=12 on all 6 edges.
    NO cell has tcBorders. After import, every cell should have
    borders populated from the table-level definition.
    """
    doc = Document()
    doc.add_paragraph("Bug 1: table-to-cell border fallback.")

    tbl_el, tblPr = make_table_with_borders(doc, 2, 2, [
        ["R0C0", "R0C1"],
        ["R1C0", "R1C1"],
    ])

    add_tbl_borders(tblPr,
                    ["top", "bottom", "left", "right", "insideH", "insideV"],
                    val="single", sz="12", color="FF0000")

    save_fixture("border-fallback", doc, {
        "name": "border-fallback",
        "spec_ref": "ISO 29500-1 SS17.4.66",
        "description": (
            "2x2 table with tblBorders (red single sz=12) on all 6 edges. "
            "NO tcBorders on any cell. Cells should inherit table borders."
        ),
    })


# =====================================================================
# Fixture 2: insideV/insideH edge selection (Bug 2)
# =====================================================================

def make_inside_v_distinct():
    """SS17.4.66: Interior cells use insideV for left/right,
    insideH for top/bottom, not the table outer borders.

    2x3 table. Table borders:
      - left/right: double blue sz=12 (color=0000FF)
      - top/bottom: double blue sz=12 (color=0000FF)
      - insideV:    single green sz=4 (color=00FF00)
      - insideH:    dashed red sz=8 (color=FF0000)

    Interior cell (0,1) should have:
      - left border = insideV (green single sz=4)
      - right border = insideV (green single sz=4)
      - bottom border = insideH (red dashed sz=8)
    """
    doc = Document()
    doc.add_paragraph("Bug 2: insideV/insideH edge selection for interior cells.")

    tbl_el, tblPr = make_table_with_borders(doc, 2, 3, [
        ["R0C0", "R0C1", "R0C2"],
        ["R1C0", "R1C1", "R1C2"],
    ])

    # Build distinct border sets
    tbl_borders = make_element("w:tblBorders")
    for edge in ["top", "bottom"]:
        tbl_borders.append(make_element(f"w:{edge}", {
            "w:val": "double", "w:sz": "12", "w:color": "0000FF", "w:space": "0",
        }))
    tbl_borders.append(make_element("w:left", {
        "w:val": "double", "w:sz": "12", "w:color": "0000FF", "w:space": "0",
    }))
    tbl_borders.append(make_element("w:right", {
        "w:val": "double", "w:sz": "12", "w:color": "0000FF", "w:space": "0",
    }))
    tbl_borders.append(make_element("w:insideH", {
        "w:val": "dashed", "w:sz": "8", "w:color": "FF0000", "w:space": "0",
    }))
    tbl_borders.append(make_element("w:insideV", {
        "w:val": "single", "w:sz": "4", "w:color": "00FF00", "w:space": "0",
    }))
    tblPr.append(tbl_borders)

    save_fixture("inside-v-distinct", doc, {
        "name": "inside-v-distinct",
        "spec_ref": "ISO 29500-1 SS17.4.66, SS17.4.38",
        "description": (
            "2x3 table with distinct outer borders (double blue sz=12) "
            "vs insideV (single green sz=4) and insideH (dashed red sz=8). "
            "Interior cell (0,1) should use insideV/insideH, not outer borders."
        ),
    })


# =====================================================================
# Fixture 3: tblLook default omitted (Bug 3)
# =====================================================================

def make_tbllook_default_omitted():
    """MS-OI29500 SS17.4.55(a): When tblLook is omitted, the default
    is 0x04A0 which means firstColumn=true, noHBand=false.

    Table with a style that has firstCol conditional (yellow shading).
    tblLook is OMITTED (no element at all). The default should apply
    firstCol formatting to column 0.

    Also: noHBand=false in the default means banding should apply.
    """
    doc = Document()
    doc.add_paragraph("Bug 3: tblLook default when omitted.")

    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    styles_part_obj = doc.part.part_related_by(RT.STYLES)
    styles_el = styles_part_obj.element

    # Define table style with firstCol conditional
    style = make_element("w:style", {"w:type": "table", "w:styleId": "FirstColTestStyle"})
    style.append(make_element("w:name", {"w:val": "First Col Test Style"}))
    style.append(make_element("w:basedOn", {"w:val": "TableNormal"}))

    style_tbl_pr = make_element("w:tblPr")
    style_tbl_pr.append(make_element("w:tblStyleRowBandSize", {"w:val": "1"}))
    tbl_borders = make_element("w:tblBorders")
    for edge in ["top", "bottom", "left", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{edge}", {
            "w:val": "single", "w:sz": "4", "w:color": "000000", "w:space": "0",
        }))
    style_tbl_pr.append(tbl_borders)
    style.append(style_tbl_pr)

    # firstCol conditional: yellow shading
    fc = make_element("w:tblStylePr", {"w:type": "firstCol"})
    fc_tc = make_element("w:tcPr")
    fc_tc.append(make_element("w:shd", {
        "w:val": "clear", "w:color": "auto", "w:fill": "FFFF00",
    }))
    fc.append(fc_tc)
    style.append(fc)

    # band1Horz conditional: gray shading
    b1 = make_element("w:tblStylePr", {"w:type": "band1Horz"})
    b1_tc = make_element("w:tcPr")
    b1_tc.append(make_element("w:shd", {
        "w:val": "clear", "w:color": "auto", "w:fill": "D9D9D9",
    }))
    b1.append(b1_tc)
    style.append(b1)

    # firstRow conditional: red shading
    fr = make_element("w:tblStylePr", {"w:type": "firstRow"})
    fr_tc = make_element("w:tcPr")
    fr_tc.append(make_element("w:shd", {
        "w:val": "clear", "w:color": "auto", "w:fill": "FF0000",
    }))
    fr.append(fr_tc)
    style.append(fr)

    styles_el.append(style)

    # Create 3x2 table using this style
    tbl = doc.add_table(rows=3, cols=2)
    for r in range(3):
        for c in range(2):
            tbl.cell(r, c).text = f"R{r}C{c}"

    tbl_el = tbl._tbl
    tblPr = tbl_el.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_el.insert(0, tblPr)

    # Remove any tblLook/tblStyle python-docx may have added
    for existing in tblPr.findall(w("tblLook")):
        tblPr.remove(existing)
    for existing in tblPr.findall(w("tblStyle")):
        tblPr.remove(existing)

    # Reference the style
    tbl_style_el = make_element("w:tblStyle", {"w:val": "FirstColTestStyle"})
    tblPr.insert(0, tbl_style_el)

    # DELIBERATELY omit tblLook — default should be 0x04A0
    # which has firstColumn=true, firstRow=true, noHBand=false

    save_fixture("tbllook-default-omitted", doc, {
        "name": "tbllook-default-omitted",
        "spec_ref": "MS-OI29500 SS17.4.55(a)",
        "description": (
            "Table with style that has firstCol conditional (yellow shading) "
            "and band1Horz conditional (gray). tblLook is OMITTED. "
            "Default 0x04A0 means firstColumn=true, noHBand=false, "
            "so firstCol shading should apply to column 0, "
            "and banding should apply to data rows."
        ),
    })


# =====================================================================
# Fixture 4: Adjacent cell border conflict resolution (Bug 4)
# =====================================================================

def make_adjacent_conflict():
    """SS17.4.66 rule 3: When adjacent cells have conflicting borders
    at a shared edge, the border with higher weight wins.

    1x3 table. No table-level borders.
      - Cell (0,0): right border = thick single sz=24 red (weight=24)
      - Cell (0,1): left border = single sz=4 blue (weight=4),
                     right border = double sz=12 green (weight=36)
      - Cell (0,2): left border = single sz=8 black (weight=8)

    Per SS17.4.66 rule 3:
      - Shared edge (0,0)|(0,1): red sz=24 wins over blue sz=4
      - Shared edge (0,1)|(0,2): green double sz=12 (weight=36) wins over black sz=8
    """
    doc = Document()
    doc.add_paragraph("Bug 4: Adjacent cell border conflict resolution.")

    tbl_el, tblPr = make_table_with_borders(doc, 1, 3, [
        ["ThickRight", "Middle", "Normal"],
    ])

    add_tbl_grid(tbl_el, [2000, 2000, 2000])

    # Cell (0,0): thick right border (red single sz=24, weight=24)
    add_tc_borders(tbl_el, 0, 0, [
        ("right", "single", "24", "FF0000", "0"),
    ])
    # Cell (0,1): thin left (blue single sz=4, weight=4),
    #             double right (green double sz=12, weight=12*3=36)
    add_tc_borders(tbl_el, 0, 1, [
        ("left", "single", "4", "0000FF", "0"),
        ("right", "double", "12", "00FF00", "0"),
    ])
    # Cell (0,2): left border (black single sz=8, weight=8)
    add_tc_borders(tbl_el, 0, 2, [
        ("left", "single", "8", "000000", "0"),
    ])

    save_fixture("adjacent-conflict", doc, {
        "name": "adjacent-conflict",
        "spec_ref": "ISO 29500-1 SS17.4.66 rule 3",
        "description": (
            "1x3 table with conflicting borders between adjacent cells. "
            "Cell (0,0) right=single sz=24 red (weight=24). "
            "Cell (0,1) left=single sz=4 blue (weight=4), "
            "right=double sz=12 green (weight=36). "
            "Cell (0,2) left=single sz=8 black (weight=8). "
            "Per SS17.4.66 rule 3, higher weight wins at each shared edge."
        ),
    })


# =====================================================================
# Main
# =====================================================================

def main():
    print("\n== Table Border Bug Fixtures ==")
    make_border_fallback()
    make_inside_v_distinct()
    make_tbllook_default_omitted()
    make_adjacent_conflict()
    print("\nDone.")


if __name__ == "__main__":
    main()

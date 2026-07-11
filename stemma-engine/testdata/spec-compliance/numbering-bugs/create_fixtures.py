# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for numbering bug-exposure tests.

Each fixture exercises a specific numbering spec requirement that was
previously buggy in the implementation. Tests in
`stemma-engine/tests/spec_numbering_bugs.rs` import these fixtures and
assert the correct behavioral outcomes per ECMA-376.

Bugs targeted:
  1. suff (suffix) parsed but not used in rendered_text
  2. pStyle reverse binding not wired up in import pipeline
  3. start default value (0 vs 1 when w:start omitted)
  4. numStyleLink chain resolution not working

Run:  uv run create_fixtures.py
"""

import json
from pathlib import Path
from lxml import etree

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

# -- XML namespace helpers ---------------------------------------------------

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def w(tag: str) -> str:
    """Return a fully-qualified wordprocessingml tag."""
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    """Create an OxmlElement with optional attributes."""
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


# -- Save / injection helpers ------------------------------------------------

def save_fixture(
    name: str,
    doc: DocxDocument,
    metadata: dict,
    filename: str = "input.docx",
) -> None:
    """Save a single-doc fixture."""
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / filename))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  numbering-bugs/{name}/")


def _inject_numbering_xml(doc: DocxDocument, numbering_xml: str) -> None:
    """Replace the numbering.xml part with custom XML."""
    numbering_part = None
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        pass

    if numbering_part is None:
        doc.add_paragraph("dummy", style="List Bullet")
        body = doc.element.body
        last_p = body.findall(w("p"))[-1]
        body.remove(last_p)
        numbering_part = doc.part.numbering_part

    numbering_part._element = etree.fromstring(numbering_xml.encode("utf-8"))


def _inject_style(doc: DocxDocument, style_xml: str) -> None:
    """Append a custom style element to the styles part."""
    style_el = etree.fromstring(style_xml.encode("utf-8"))
    doc.styles.element.append(style_el)


def _add_numbered_para(
    doc: DocxDocument, text: str, num_id: str, ilvl: str = "0"
):
    """Add a paragraph with numPr referencing the given numId/ilvl."""
    p = doc.add_paragraph(text)
    pPr = p._p.get_or_add_pPr()
    numPr = make_element("w:numPr")
    numPr.append(make_element("w:ilvl", {"w:val": ilvl}))
    numPr.append(make_element("w:numId", {"w:val": num_id}))
    pPr.append(numPr)
    return p


def _add_styled_para(doc: DocxDocument, text: str, style_id: str):
    """Add a paragraph with a pStyle reference (no numPr)."""
    p = doc.add_paragraph(text)
    pPr = p._p.get_or_add_pPr()
    pStyle = make_element("w:pStyle", {"w:val": style_id})
    pPr.insert(0, pStyle)
    return p


# =========================================================================
# Bug 1: suff (suffix) parsed but not used in rendered_text
# =========================================================================

def make_suff_space() -> None:
    """ECMA-376 section 17.9.28: suff controls the separator between numbering
    text and paragraph body.

    Fixture:
      - numId=1: level 0 with suff="space" (decimal, lvlText="%1.")
      - numId=2: level 0 with suff="nothing" (decimal, lvlText="%1.")
      - numId=3: level 0 with suff="tab" (decimal, lvlText="%1.") -- control
      - Two paragraphs per numId

    Expected:
      - numId=1 rendered_text: "1. Body" (space separator)
      - numId=2 rendered_text: "1.Body" (no separator)
      - numId=3 rendered_text: "1.\tBody" (tab separator, default)

    Bug: The LevelDef.suffix field was parsed but rendered_text always used
    a tab separator regardless of the suff value.
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:suff w:val="space"/>
    </w:lvl>
  </w:abstractNum>
  <w:abstractNum w:abstractNumId="1">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:suff w:val="nothing"/>
    </w:lvl>
  </w:abstractNum>
  <w:abstractNum w:abstractNumId="2">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:suff w:val="tab"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
  <w:num w:numId="2">
    <w:abstractNumId w:val="1"/>
  </w:num>
  <w:num w:numId="3">
    <w:abstractNumId w:val="2"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    _add_numbered_para(doc, "Space suffix first", "1", "0")
    _add_numbered_para(doc, "Space suffix second", "1", "0")
    _add_numbered_para(doc, "Nothing suffix first", "2", "0")
    _add_numbered_para(doc, "Nothing suffix second", "2", "0")
    _add_numbered_para(doc, "Tab suffix first", "3", "0")
    _add_numbered_para(doc, "Tab suffix second", "3", "0")

    save_fixture("suff-space", doc, {
        "name": "suff-space",
        "spec_ref": "ECMA-376 section 17.9.28",
        "description": (
            "Three numbering definitions with suff='space', suff='nothing', "
            "and suff='tab'. The separator in rendered_text must match the "
            "suff value, not always be a tab."
        ),
        "expected_behavior": (
            "Space: '1. Body'. Nothing: '1.Body'. Tab: '1.\\tBody'."
        ),
    })


# =========================================================================
# Bug 2: pStyle reverse binding not wired up
# =========================================================================

def make_pstyle_reverse_binding() -> None:
    """ECMA-376 section 17.9.23: pStyle element on a numbering level creates
    a reverse binding from a paragraph style to that numbering level.

    Fixture:
      - abstractNum 0, level 0: decimal, pStyle="CustomListNum"
      - Style "CustomListNum" (type=paragraph): no numPr in the style itself
      - Three paragraphs using style "CustomListNum" with NO direct numPr

    Expected: paragraphs get numbering from the pStyle reverse binding:
      "1.", "2.", "3."

    Bug: build_pstyle_reverse_map() existed but was never called from the
    import pipeline. Paragraphs using a style claimed by a numbering level
    did not receive numbering.
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:pStyle w:val="CustomListNum"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # Create the paragraph style "CustomListNum" with NO numPr.
    # The numbering should come solely from the pStyle reverse binding.
    style_xml = f"""\
<w:style w:type="paragraph" w:styleId="CustomListNum" xmlns:w="{W}">
  <w:name w:val="Custom List Num"/>
  <w:pPr>
    <w:ind w:left="720" w:hanging="360"/>
  </w:pPr>
</w:style>"""
    _inject_style(doc, style_xml)

    # These paragraphs use the style but have no direct numPr.
    _add_styled_para(doc, "First via pStyle binding", "CustomListNum")
    _add_styled_para(doc, "Second via pStyle binding", "CustomListNum")
    _add_styled_para(doc, "Third via pStyle binding", "CustomListNum")

    save_fixture("pstyle-reverse-binding", doc, {
        "name": "pstyle-reverse-binding",
        "spec_ref": "ECMA-376 section 17.9.23",
        "description": (
            "Numbering level 0 has pStyle='CustomListNum'. Paragraph style "
            "'CustomListNum' has no numPr. Three paragraphs use the style. "
            "They should get numbering from the reverse binding."
        ),
        "expected_behavior": (
            "Paragraphs produce '1.', '2.', '3.' via pStyle reverse binding."
        ),
    })


# =========================================================================
# Bug 3: start default value (0 vs 1)
# =========================================================================

def make_start_omitted() -> None:
    """ECMA-376 section 17.9.25: when w:start is omitted, the starting value
    shall be zero (0).

    Fixture:
      - abstractNum 0, level 0: decimal, NO w:start element
      - Three paragraphs at level 0

    Expected: "0.", "1.", "2." (starting from 0)

    Bug: The implementation defaulted start to 1 instead of the
    spec-mandated 0, producing "1.", "2.", "3.".
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    _add_numbered_para(doc, "First item (start omitted)", "1", "0")
    _add_numbered_para(doc, "Second item", "1", "0")
    _add_numbered_para(doc, "Third item", "1", "0")

    save_fixture("start-omitted", doc, {
        "name": "start-omitted",
        "spec_ref": "ECMA-376 section 17.9.25",
        "description": (
            "Level 0 has no w:start element. Per the spec the starting "
            "value defaults to 0."
        ),
        "expected_behavior": "Items produce '0.', '1.', '2.'.",
    })


# =========================================================================
# Bug 4: numStyleLink chain resolution
# =========================================================================

def make_num_style_link_chain() -> None:
    """ECMA-376 section 17.9.21 and 17.9.27: numStyleLink and styleLink
    create bidirectional links between numbering definitions and styles.

    Fixture:
      - abstractNum 0: has numStyleLink="OutlineList" (no levels)
      - abstractNum 1: has styleLink="OutlineList" (actual levels defined)
      - numId=1 -> abstractNum 0
      - numId=2 -> abstractNum 1
      - Style "OutlineList" (type=numbering) with numPr -> numId=2
      - Three paragraphs using numId=1

    Expected: paragraphs using numId=1 follow the chain through
    numStyleLink -> abstractNum 1's levels, producing "1.", "2.", "a)".

    Bug: The numStyleLink chain was not followed in get_level(). abstractNum 0
    had no levels, so synthesis failed silently.
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:multiLevelType w:val="multilevel"/>
    <w:numStyleLink w:val="OutlineList"/>
  </w:abstractNum>
  <w:abstractNum w:abstractNumId="1">
    <w:multiLevelType w:val="multilevel"/>
    <w:styleLink w:val="OutlineList"/>
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="1">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerLetter"/>
      <w:lvlText w:val="%2)"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
  <w:num w:numId="2">
    <w:abstractNumId w:val="1"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # Create the numbering style that the chain references
    style_xml = f"""\
<w:style w:type="numbering" w:styleId="OutlineList" xmlns:w="{W}">
  <w:name w:val="Outline List"/>
  <w:pPr>
    <w:numPr>
      <w:numId w:val="2"/>
    </w:numPr>
  </w:pPr>
</w:style>"""
    _inject_style(doc, style_xml)

    # Paragraphs use numId=1, which should follow the chain
    _add_numbered_para(doc, "First via numStyleLink", "1", "0")
    _add_numbered_para(doc, "Second via numStyleLink", "1", "0")
    _add_numbered_para(doc, "Sub-item via numStyleLink", "1", "1")

    save_fixture("num-style-link-chain", doc, {
        "name": "num-style-link-chain",
        "spec_ref": "ECMA-376 section 17.9.21, section 17.9.27",
        "description": (
            "abstractNum 0 has numStyleLink='OutlineList' (no levels). "
            "abstractNum 1 has styleLink='OutlineList' with actual levels. "
            "Paragraphs using numId=1 (-> abstractNum 0) should follow "
            "the chain to get levels from abstractNum 1."
        ),
        "expected_behavior": (
            "Paragraphs produce '1.', '2.', 'a)' via the numStyleLink chain."
        ),
    })


# =========================================================================
# Main
# =========================================================================

def main() -> None:
    print("Generating numbering-bugs fixtures:")
    make_suff_space()
    make_pstyle_reverse_binding()
    make_start_omitted()
    make_num_style_link_chain()
    print("\nDone.")


if __name__ == "__main__":
    main()

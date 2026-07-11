# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for MS-OI29500 paragraph property implementation notes.

These fixtures target behaviors documented in MS-OI29500 §17.3.1 and MS-DOCX §2.3
where Microsoft Word's implementation diverges from or supplements the base ISO spec.

Run:  uv run create_ms_para_props.py
"""

import json
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    """Create an OxmlElement with optional attributes."""
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(
    area: str,
    name: str,
    doc: DocxDocument,
    metadata: dict,
    filename: str = "input.docx",
) -> None:
    """Save a single-doc fixture."""
    out = ROOT / area / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / filename))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  {area}/{name}/")


def _apply_table_style(tbl, style_id: str, look_attrs: dict) -> None:
    """Apply a table style and tblLook to an existing table element."""
    tbl_element = tbl._tbl
    tblPr = tbl_element.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_element.insert(0, tblPr)
    tblPr.insert(0, make_element("w:tblStyle", {"w:val": style_id}))
    for existing in tblPr.findall(qn("w:tblLook")):
        tblPr.remove(existing)
    tblPr.append(make_element("w:tblLook", look_attrs))


def _fill_table_cells(tbl, labels: list[list[str]]) -> None:
    """Fill table cells with labels."""
    for r_idx, row_labels in enumerate(labels):
        for c_idx, label in enumerate(row_labels):
            tbl.cell(r_idx, c_idx).text = label


# =========================================================================
# MS PARAGRAPH PROPERTIES (MS-OI29500 §17.3.1 + MS-DOCX §2.3)
# =========================================================================

def make_ms_jc_default_left() -> None:
    """MS-OI29500 2.1.45 §17.3.1.13: Word defaults jc to left when unspecified.

    The ISO spec says no alignment is applied when jc is never specified in
    the style hierarchy. Word applies left alignment as the default.
    """
    doc = Document()

    # P0: no style, no direct alignment — Word defaults to left
    doc.add_paragraph("No alignment specified anywhere.")

    # P1: style has center alignment — should resolve center
    styles_el = doc.styles.element
    style_el = make_element("w:style", {"w:type": "paragraph", "w:styleId": "CenterStyle"})
    style_el.append(make_element("w:name", {"w:val": "Center Style"}))
    pPr = make_element("w:pPr")
    pPr.append(make_element("w:jc", {"w:val": "center"}))
    style_el.append(pPr)
    styles_el.append(style_el)

    p1 = doc.add_paragraph("Center alignment from style.")
    p1._p.get_or_add_pPr().append(make_element("w:pStyle", {"w:val": "CenterStyle"}))

    # P2: Normal style, no jc in Normal — Word should still default to left
    doc.add_paragraph("Normal style, no jc defined.")

    save_fixture("ms-paragraph-props", "jc-default-left", doc, {
        "name": "jc-default-left",
        "spec_ref": "MS-OI29500 2.1.45 §17.3.1.13",
        "description": (
            "P0: no style, no direct jc — Word defaults to left. "
            "P1: CenterStyle has jc=center. "
            "P2: Normal style, no jc — Word defaults to left."
        ),
        "expected_behavior": (
            "P0: alignment=Left (Word default). "
            "P1: alignment=Center (from style). "
            "P2: alignment=Left (Word default)."
        ),
        "current_status": "GAP — our code returns None when no alignment specified",
    })


def make_ms_keeplines_in_table() -> None:
    """MS-OI29500 2.1.46 §17.3.1.14: Word ignores keepLines inside tables.

    The ISO spec says keepLines applies to all paragraphs, but Word
    ignores it for paragraphs rendered inside a table cell.
    """
    doc = Document()

    # P0: keepLines outside table (should be honored)
    p0 = doc.add_paragraph("Keep lines together — outside table.")
    p0._p.get_or_add_pPr().append(make_element("w:keepLines"))

    # Table with keepLines on paragraph inside cell
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    cell_p = cell.paragraphs[0]
    cell_p.text = "Keep lines together — inside table cell."
    cell_p._p.get_or_add_pPr().append(make_element("w:keepLines"))

    # P2: another keepLines outside table (control)
    p2 = doc.add_paragraph("Keep lines together — outside table (control).")
    p2._p.get_or_add_pPr().append(make_element("w:keepLines"))

    save_fixture("ms-paragraph-props", "keeplines-in-table", doc, {
        "name": "keeplines-in-table",
        "spec_ref": "MS-OI29500 2.1.46 §17.3.1.14",
        "description": (
            "P0: keepLines outside table. "
            "Table cell P: keepLines inside table. "
            "P2: keepLines outside table (control)."
        ),
        "expected_behavior": (
            "P0: keep_lines=true (honored). "
            "Table cell P: keep_lines=true in XML but Word ignores it in rendering. "
            "P2: keep_lines=true (honored)."
        ),
        "current_status": "GAP — we parse keepLines but don't track table context",
    })


def make_ms_widow_control_in_table() -> None:
    """MS-OI29500 2.1.66 §17.3.1.44: Word ignores widowControl inside tables.

    Same pattern as keepLines — the property is ignored
    when the paragraph is inside a table cell.
    """
    doc = Document()

    # P0: widowControl=false outside table
    p0 = doc.add_paragraph("Widow control disabled — outside table.")
    p0._p.get_or_add_pPr().append(make_element("w:widowControl", {"w:val": "0"}))

    # Table with widowControl=false on paragraph inside cell
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    cell_p = cell.paragraphs[0]
    cell_p.text = "Widow control disabled — inside table."
    cell_p._p.get_or_add_pPr().append(make_element("w:widowControl", {"w:val": "0"}))

    save_fixture("ms-paragraph-props", "widow-control-in-table", doc, {
        "name": "widow-control-in-table",
        "spec_ref": "MS-OI29500 2.1.66 §17.3.1.44",
        "description": (
            "P0: widowControl=false outside table. "
            "Table cell P: widowControl=false inside table."
        ),
        "expected_behavior": (
            "P0: widow_control=Some(false) (honored). "
            "Table cell P: widow_control=Some(false) in XML but Word ignores it."
        ),
        "current_status": "GAP — we parse widowControl but don't track table context",
    })


def make_ms_spacing_beforelines_override() -> None:
    """MS-OI29500 2.1.60 §17.3.1.33: beforeLines/afterLines override non-line-unit spacing.

    During style hierarchy application, Word overrides non-line-unit spacing
    with related non-zero line-unit spacing from earlier in the style hierarchy.

    Style defines: beforeLines=200 (line units).
    Paragraph has direct before=480 (twips).
    Per ISO spec, direct before=480 should win.
    Per MS Word, beforeLines=200 from style overrides before=480 from direct.
    """
    doc = Document()

    # Style with beforeLines and afterLines
    styles_el = doc.styles.element
    style_el = make_element("w:style", {"w:type": "paragraph", "w:styleId": "LineSpacedStyle"})
    style_el.append(make_element("w:name", {"w:val": "Line Spaced Style"}))
    pPr = make_element("w:pPr")
    pPr.append(make_element("w:spacing", {
        "w:beforeLines": "200",
        "w:afterLines": "100",
    }))
    style_el.append(pPr)
    styles_el.append(style_el)

    # P0: style only — should have beforeLines=200, afterLines=100
    p0 = doc.add_paragraph("Style only — line-unit spacing.")
    p0._p.get_or_add_pPr().append(make_element("w:pStyle", {"w:val": "LineSpacedStyle"}))

    # P1: style + direct before=480 — MS Word: beforeLines from style overrides
    p1 = doc.add_paragraph("Style + direct before=480.")
    p1_ppr = p1._p.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "LineSpacedStyle"}))
    p1_ppr.append(make_element("w:spacing", {"w:before": "480"}))

    # P2: no style, direct before=480 (control)
    p2 = doc.add_paragraph("Direct before=480 only.")
    p2._p.get_or_add_pPr().append(make_element("w:spacing", {"w:before": "480"}))

    save_fixture("ms-paragraph-props", "spacing-beforelines-override", doc, {
        "name": "spacing-beforelines-override",
        "spec_ref": "MS-OI29500 2.1.60 §17.3.1.33",
        "description": (
            "P0: LineSpacedStyle (beforeLines=200, afterLines=100). "
            "P1: LineSpacedStyle + direct before=480. "
            "P2: direct before=480 only (control)."
        ),
        "expected_behavior": (
            "P0: before_lines=200, after_lines=100. "
            "P1: Word overrides direct before=480 with style's beforeLines=200. "
            "P2: before=480."
        ),
        "current_status": "GAP — we do element-level override for direct spacing, ignoring line-unit precedence",
    })


def make_ms_ind_char_unit_zero() -> None:
    """MS-OI29500 2.1.44 §17.3.1.12(a): Word ignores zero-value character unit indents.

    Style defines startChars="100" (indent in character units).
    Paragraph has direct startChars="0" — Word ignores this and keeps
    the style's character unit indent, along with related char-unit indents.
    """
    doc = Document()

    # Style with character-unit indent
    styles_el = doc.styles.element
    style_el = make_element("w:style", {"w:type": "paragraph", "w:styleId": "CharIndentStyle"})
    style_el.append(make_element("w:name", {"w:val": "Char Indent Style"}))
    pPr = make_element("w:pPr")
    pPr.append(make_element("w:ind", {
        "w:startChars": "100",
        "w:firstLineChars": "50",
    }))
    style_el.append(pPr)
    styles_el.append(style_el)

    # P0: style only — should have character unit indents
    p0 = doc.add_paragraph("Style with startChars=100, firstLineChars=50.")
    p0._p.get_or_add_pPr().append(make_element("w:pStyle", {"w:val": "CharIndentStyle"}))

    # P1: style + direct startChars="0" — Word ignores the zero
    p1 = doc.add_paragraph("Direct startChars=0 should be ignored by Word.")
    p1_ppr = p1._p.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "CharIndentStyle"}))
    p1_ppr.append(make_element("w:ind", {"w:startChars": "0"}))

    save_fixture("ms-paragraph-props", "ind-char-unit-zero", doc, {
        "name": "ind-char-unit-zero",
        "spec_ref": "MS-OI29500 2.1.44 §17.3.1.12(a)",
        "description": (
            "P0: CharIndentStyle (startChars=100, firstLineChars=50). "
            "P1: CharIndentStyle + direct startChars=0."
        ),
        "expected_behavior": (
            "P0: character-unit indent from style. "
            "P1: Word ignores zero character unit indent, keeps style's startChars=100."
        ),
        "current_status": "GAP — we don't parse character-unit indent attributes at all",
    })


def make_ms_ind_char_unit_override() -> None:
    """MS-OI29500 2.1.44 §17.3.1.12(b): Non-zero char-unit overrides non-char-unit.

    Style defines startChars="200" (character units).
    Paragraph has direct w:start="720" (twips).
    Per MS Word: the style's non-zero startChars overrides the direct w:start.
    """
    doc = Document()

    # Style with character-unit indent
    styles_el = doc.styles.element
    style_el = make_element("w:style", {"w:type": "paragraph", "w:styleId": "CharOverrideStyle"})
    style_el.append(make_element("w:name", {"w:val": "Char Override Style"}))
    pPr = make_element("w:pPr")
    pPr.append(make_element("w:ind", {"w:startChars": "200"}))
    style_el.append(pPr)
    styles_el.append(style_el)

    # P0: style + direct w:start="720" — style's startChars should override
    p0 = doc.add_paragraph("Style startChars=200 overrides direct start=720.")
    p0_ppr = p0._p.get_or_add_pPr()
    p0_ppr.append(make_element("w:pStyle", {"w:val": "CharOverrideStyle"}))
    p0_ppr.append(make_element("w:ind", {"w:start": "720"}))

    # P1: no style, direct start=720 (control)
    p1 = doc.add_paragraph("Direct start=720 only (control).")
    p1._p.get_or_add_pPr().append(make_element("w:ind", {"w:start": "720"}))

    save_fixture("ms-paragraph-props", "ind-char-unit-override", doc, {
        "name": "ind-char-unit-override",
        "spec_ref": "MS-OI29500 2.1.44 §17.3.1.12(b)",
        "description": (
            "P0: CharOverrideStyle (startChars=200) + direct start=720. "
            "P1: direct start=720 only (control)."
        ),
        "expected_behavior": (
            "P0: Word uses startChars=200 from style, ignoring direct start=720. "
            "P1: start=720 (no char-unit interference)."
        ),
        "current_status": "GAP — we don't parse character-unit indent attributes at all",
    })


def make_ms_mirror_indents_hanging() -> None:
    """MS-OI29500 2.1.49 §17.3.1.18(a): mirrorIndents + hanging indent bug.

    Word incorrectly aligns the hanging indent to the left margin for
    odd-numbered pages when a hanging indent is used with mirrorIndents.
    """
    doc = Document()

    # P0: mirrorIndents + hanging indent
    p0 = doc.add_paragraph("Mirror indents with hanging indent.")
    p0_ppr = p0._p.get_or_add_pPr()
    p0_ppr.append(make_element("w:mirrorIndents"))
    p0_ppr.append(make_element("w:ind", {
        "w:start": "1440",
        "w:hanging": "720",
    }))

    # P1: mirrorIndents without hanging (control)
    p1 = doc.add_paragraph("Mirror indents without hanging.")
    p1_ppr = p1._p.get_or_add_pPr()
    p1_ppr.append(make_element("w:mirrorIndents"))
    p1_ppr.append(make_element("w:ind", {"w:start": "1440"}))

    # P2: hanging indent without mirrorIndents (control)
    p2 = doc.add_paragraph("Hanging indent without mirror.")
    p2_ppr = p2._p.get_or_add_pPr()
    p2_ppr.append(make_element("w:ind", {
        "w:start": "1440",
        "w:hanging": "720",
    }))

    save_fixture("ms-paragraph-props", "mirror-indents-hanging", doc, {
        "name": "mirror-indents-hanging",
        "spec_ref": "MS-OI29500 2.1.49 §17.3.1.18(a)",
        "description": (
            "P0: mirrorIndents + start=1440, hanging=720. "
            "P1: mirrorIndents + start=1440, no hanging. "
            "P2: start=1440, hanging=720, no mirrorIndents."
        ),
        "expected_behavior": (
            "P0: Word has known bug — misaligns hanging indent on odd pages. "
            "P1: mirrorIndents swaps start/end on even pages. "
            "P2: normal hanging indent behavior."
        ),
        "current_status": "GAP — we don't parse mirrorIndents at all",
    })


def make_ms_mirror_indents_in_table() -> None:
    """MS-OI29500 2.1.49 §17.3.1.18(b): Word ignores mirrorIndents in tables.

    Similar to keepLines and widowControl — the property is ignored
    when the paragraph is inside a table cell.
    """
    doc = Document()

    # P0: mirrorIndents outside table
    p0 = doc.add_paragraph("Mirror indents — outside table.")
    p0_ppr = p0._p.get_or_add_pPr()
    p0_ppr.append(make_element("w:mirrorIndents"))
    p0_ppr.append(make_element("w:ind", {"w:start": "1440", "w:end": "720"}))

    # Table with mirrorIndents on paragraph inside cell
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    cell_p = cell.paragraphs[0]
    cell_p.text = "Mirror indents — inside table (ignored by Word)."
    cell_p_ppr = cell_p._p.get_or_add_pPr()
    cell_p_ppr.append(make_element("w:mirrorIndents"))
    cell_p_ppr.append(make_element("w:ind", {"w:start": "1440", "w:end": "720"}))

    save_fixture("ms-paragraph-props", "mirror-indents-in-table", doc, {
        "name": "mirror-indents-in-table",
        "spec_ref": "MS-OI29500 2.1.49 §17.3.1.18(b)",
        "description": (
            "P0: mirrorIndents + start=1440, end=720 outside table. "
            "Table cell P: same mirrorIndents inside table."
        ),
        "expected_behavior": (
            "P0: mirrorIndents honored (swaps start/end on even pages). "
            "Table cell P: Word ignores mirrorIndents."
        ),
        "current_status": "GAP — we don't parse mirrorIndents at all",
    })


def make_ms_override_table_style_font_jc() -> None:
    """MS-DOCX 2.3.1: overrideTableStyleFontSizeAndJustification.

    When this compat setting is FALSE (the default), the default paragraph
    style's font size of 11pt/12pt does NOT override a table style's font size,
    and the default paragraph style's left justification does NOT override
    a table style's justification.

    This is a critical behavioral switch for table rendering.
    """
    doc = Document()

    # Create a table style with center alignment and 14pt font
    styles_el = doc.styles.element

    tbl_style = make_element("w:style", {"w:type": "table", "w:styleId": "FancyTableStyle"})
    tbl_style.append(make_element("w:name", {"w:val": "Fancy Table Style"}))
    tbl_pPr = make_element("w:pPr")
    tbl_pPr.append(make_element("w:jc", {"w:val": "center"}))
    tbl_style.append(tbl_pPr)
    tbl_rPr = make_element("w:rPr")
    tbl_rPr.append(make_element("w:sz", {"w:val": "28"}))  # 14pt
    tbl_style.append(tbl_rPr)
    # Add table borders for visibility
    tblPr = make_element("w:tblPr")
    tbl_borders = make_element("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "4", "w:color": "000000",
        }))
    tblPr.append(tbl_borders)
    tbl_style.append(tblPr)
    styles_el.append(tbl_style)

    doc.add_paragraph("Text outside table (Normal style).")
    tbl = doc.add_table(rows=2, cols=2)
    _apply_table_style(tbl, "FancyTableStyle", {
        "w:val": "0000",
        "w:firstRow": "0", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "1",
    })
    _fill_table_cells(tbl, [["Cell A", "Cell B"], ["Cell C", "Cell D"]])

    save_fixture("ms-paragraph-props", "override-table-style-font-jc", doc, {
        "name": "override-table-style-font-jc",
        "spec_ref": "MS-DOCX 2.3.1 overrideTableStyleFontSizeAndJustification",
        "description": (
            "FancyTableStyle has jc=center, sz=28 (14pt). "
            "Normal style has default jc=left, sz=22 (11pt). "
            "No overrideTableStyleFontSizeAndJustification compat setting."
        ),
        "expected_behavior": (
            "Default behavior (compat setting false/absent): "
            "Normal's 11pt font and left jc do NOT override table style. "
            "Table cells should render center-aligned at 14pt."
        ),
        "current_status": "GAP — we don't parse or respect compat settings",
    })


def make_ms_tab_pos_range() -> None:
    """MS-OI29500 2.1.61 §17.3.1.37: Tab stop pos restricted to [-31680, 31680].

    Word restricts tab stop positions to the range -31680 to 31680 (22 inches).
    Positions outside this range should be clamped or rejected.
    """
    doc = Document()

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    tabs_el = make_element("w:tabs")
    # Normal tab within range
    tabs_el.append(make_element("w:tab", {"w:val": "left", "w:pos": "1440"}))
    # Tab at max boundary
    tabs_el.append(make_element("w:tab", {"w:val": "right", "w:pos": "31680"}))
    # Tab beyond max — should be clamped/rejected by Word
    tabs_el.append(make_element("w:tab", {"w:val": "left", "w:pos": "32000"}))
    pPr.append(tabs_el)

    p.add_run("A\tB\tC")

    save_fixture("ms-paragraph-props", "tab-pos-range", doc, {
        "name": "tab-pos-range",
        "spec_ref": "MS-OI29500 2.1.61 §17.3.1.37",
        "description": (
            "Tab stops at: 1440 (normal), 31680 (max), 32000 (beyond max). "
            "Text has 2 tab characters."
        ),
        "expected_behavior": (
            "Tab at 1440: normal. Tab at 31680: at max boundary. "
            "Tab at 32000: Word restricts to [-31680, 31680], should be clamped."
        ),
        "current_status": "GAP — we don't validate tab stop position range",
    })


def main() -> None:
    print("Generating MS paragraph property fixtures:")
    make_ms_jc_default_left()
    make_ms_keeplines_in_table()
    make_ms_widow_control_in_table()
    make_ms_spacing_beforelines_override()
    make_ms_ind_char_unit_zero()
    make_ms_ind_char_unit_override()
    make_ms_mirror_indents_hanging()
    make_ms_mirror_indents_in_table()
    make_ms_override_table_style_font_jc()
    make_ms_tab_pos_range()
    print("\nDone.")


if __name__ == "__main__":
    main()

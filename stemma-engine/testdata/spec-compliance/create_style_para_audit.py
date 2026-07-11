# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for spec-compliance testing of paragraph property
resolution through the style chain (widowControl, keepNext, keepLines, outlineLvl).

Run:  uv run create_style_para_audit.py
"""

import json
from pathlib import Path
from lxml import etree

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def ensure_ppr(p) -> OxmlElement:
    """Return the w:pPr element for a paragraph, creating it if absent."""
    ppr = p._element.find(qn("w:pPr"))
    if ppr is None:
        ppr = make_element("w:pPr")
        p._element.insert(0, ppr)
    return ppr


def save_fixture(area: str, name: str, doc, metadata: dict) -> None:
    out = ROOT / area / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  {area}/{name}/")


# =========================================================================
# GAP-141: widowControl resolved through style chain
# =========================================================================

def make_widow_control_from_style():
    """Style defines widowControl=false; paragraph inherits it."""
    doc = Document()

    # Add a custom style "NoWidow" with widowControl=false
    styles_el = doc.styles.element
    style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "NoWidow"})
    name_el = make_element("w:name", {"w:val": "No Widow"})
    style.append(name_el)
    ppr = make_element("w:pPr")
    wc = make_element("w:widowControl", {"w:val": "0"})
    ppr.append(wc)
    style.append(ppr)
    styles_el.append(style)

    # Para 0: No style, no direct widowControl — should inherit spec default (true/None)
    p0 = doc.add_paragraph("Default paragraph")

    # Para 1: NoWidow style, no direct widowControl — should get false from style
    p1 = doc.add_paragraph("Paragraph with NoWidow style")
    ensure_ppr(p1).append(
        make_element("w:pStyle", {"w:val": "NoWidow"})
    )

    # Para 2: NoWidow style + direct widowControl=true — direct wins
    p2 = doc.add_paragraph("Paragraph with direct override")
    p2_ppr = ensure_ppr(p2)
    p2_ppr.append(make_element("w:pStyle", {"w:val": "NoWidow"}))
    p2_ppr.append(make_element("w:widowControl"))  # no val = true

    save_fixture("style-para-audit", "widow-control-from-style", doc, {
        "description": "GAP-141: widowControl resolved through style chain",
        "paragraphs": [
            {"text": "Default paragraph", "expected_widow_control": None},
            {"text": "NoWidow style", "expected_widow_control": False},
            {"text": "Direct override", "expected_widow_control": True},
        ],
    })


# =========================================================================
# GAP-142: keepNext/keepLines resolved through style chain
# =========================================================================

def make_keep_next_lines_from_style():
    """Style defines keepNext + keepLines; paragraph inherits them."""
    doc = Document()

    # Add a custom style "KeepTogether" with keepNext + keepLines
    styles_el = doc.styles.element
    style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "KeepTogether"})
    name_el = make_element("w:name", {"w:val": "Keep Together"})
    style.append(name_el)
    ppr = make_element("w:pPr")
    ppr.append(make_element("w:keepNext"))
    ppr.append(make_element("w:keepLines"))
    style.append(ppr)
    styles_el.append(style)

    # Para 0: No style — keep_next=false, keep_lines=false
    p0 = doc.add_paragraph("Default paragraph")

    # Para 1: KeepTogether style — should inherit keepNext=true, keepLines=true
    p1 = doc.add_paragraph("Paragraph with KeepTogether style")
    ensure_ppr(p1).append(
        make_element("w:pStyle", {"w:val": "KeepTogether"})
    )

    save_fixture("style-para-audit", "keep-next-lines-from-style", doc, {
        "description": "GAP-142: keepNext/keepLines resolved through style chain",
        "paragraphs": [
            {"text": "Default", "keep_next": False, "keep_lines": False},
            {"text": "KeepTogether style", "keep_next": True, "keep_lines": True},
        ],
    })


# =========================================================================
# GAP-143: outlineLvl resolved from style chain
# =========================================================================

def make_outline_lvl_from_style():
    """Custom style defines outlineLvl=2; paragraph inherits heading level 3."""
    doc = Document()

    # Add a custom style "CustomHeading" with outlineLvl=2 (heading level 3)
    styles_el = doc.styles.element
    style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "CustomHeading"})
    name_el = make_element("w:name", {"w:val": "Custom Heading"})
    style.append(name_el)
    ppr = make_element("w:pPr")
    ppr.append(make_element("w:outlineLvl", {"w:val": "2"}))
    style.append(ppr)
    styles_el.append(style)

    # Para 0: No style — no heading level
    p0 = doc.add_paragraph("Default paragraph")

    # Para 1: CustomHeading style — should get heading level 3 from outlineLvl=2
    p1 = doc.add_paragraph("Paragraph with CustomHeading style")
    ensure_ppr(p1).append(
        make_element("w:pStyle", {"w:val": "CustomHeading"})
    )

    save_fixture("style-para-audit", "outline-lvl-from-style", doc, {
        "description": "GAP-143: outlineLvl resolved from custom style chain",
        "paragraphs": [
            {"text": "Default", "heading_level": None},
            {"text": "CustomHeading style", "heading_level": 3},
        ],
    })


if __name__ == "__main__":
    print("Generating style-para-audit fixtures:")
    make_widow_control_from_style()
    make_keep_next_lines_from_style()
    make_outline_lvl_from_style()
    print("Done.")

# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for style constraint and merge behavioral tests.

Each fixture targets a specific OOXML style constraint:
  - style-constraints/linked-styles/         (§17.7.4.6 — linked paragraph/character styles)
  - style-constraints/based-on-type-mismatch/ (§17.7.4.3 — basedOn cross-type validation)
  - style-constraints/table-conditional-precedence/ (§17.7.6.6 — nwCell corner override)
  - style-constraints/spacing-line-rule-default/   (§17.3.1.33 — lineRule default behavior)

Run:  uv run create_style_constraints.py
"""

import json
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

# ── XML namespace helpers ────────────────────────────────────────────────

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    """Return a fully-qualified wordprocessingml tag."""
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    """Create an OxmlElement with optional attributes."""
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


# ── Save helpers ─────────────────────────────────────────────────────────

def save_fixture(
    area: str,
    name: str,
    doc: DocxDocument,
    metadata: dict,
    filename: str = "input.docx",
) -> None:
    """Save a single-doc fixture (for parsing/model tests)."""
    out = ROOT / area / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / filename))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  {area}/{name}/")


# ── Table style helpers ──────────────────────────────────────────────────

def _add_table_style(doc: DocxDocument, style_id: str, name: str,
                     tblPr_children=None, tcPr_children=None,
                     tblStylePrs=None) -> None:
    """Helper to inject a table style into styles.xml."""
    styles_element = doc.styles.element
    tbl_style = make_element("w:style", {"w:type": "table", "w:styleId": style_id})
    tbl_style.append(make_element("w:name", {"w:val": name}))

    if tblPr_children:
        tblPr = make_element("w:tblPr")
        for child in tblPr_children:
            tblPr.append(child)
        tbl_style.append(tblPr)

    if tcPr_children:
        tcPr = make_element("w:tcPr")
        for child in tcPr_children:
            tcPr.append(child)
        tbl_style.append(tcPr)

    if tblStylePrs:
        for spr in tblStylePrs:
            tbl_style.append(spr)

    styles_element.append(tbl_style)


def _apply_table_style(tbl, style_id: str, look_attrs: dict) -> None:
    """Apply a table style and tblLook to an existing table element."""
    tbl_element = tbl._tbl
    tblPr = tbl_element.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_element.insert(0, tblPr)
    tblPr.insert(0, make_element("w:tblStyle", {"w:val": style_id}))
    # Remove any existing tblLook before adding our own (python-docx adds a default one).
    for existing in tblPr.findall(qn("w:tblLook")):
        tblPr.remove(existing)
    tblPr.append(make_element("w:tblLook", look_attrs))


def _fill_table_cells(tbl, labels: list[list[str]]) -> None:
    """Fill table cells with labels."""
    for r_idx, row_labels in enumerate(labels):
        for c_idx, label in enumerate(row_labels):
            tbl.cell(r_idx, c_idx).text = label


# =========================================================================
# LINKED STYLES (§17.7.4.6)
# =========================================================================

def make_linked_styles() -> None:
    """Linked paragraph + character style pair with DIVERGENT formatting.

    ISO 29500-1 §17.7.4.6: The link element creates a bidirectional link
    between a paragraph style and a character style. When a paragraph uses
    a linked paragraph style, runs should inherit the linked character
    style's rPr — NOT the paragraph style's own rPr.

    To prove the linked char style wins, the two styles have different rPr:
      - HeadingLinked (paragraph): bold only
      - HeadingLinkedChar (character): italic + color=FF0000 + sz=28
    """
    doc = Document()
    styles_element = doc.styles.element

    # Paragraph style "HeadingLinked" — rPr has bold ONLY
    para_style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "HeadingLinked"})
    para_style.append(make_element("w:name", {"w:val": "Heading Linked"}))
    para_style.append(make_element("w:link", {"w:val": "HeadingLinkedChar"}))
    rpr_para = make_element("w:rPr")
    rpr_para.append(make_element("w:b"))
    para_style.append(rpr_para)
    styles_element.append(para_style)

    # Character style "HeadingLinkedChar" — rPr has italic + red + 14pt (NO bold)
    char_style = make_element("w:style", {"w:type": "character", "w:styleId": "HeadingLinkedChar"})
    char_style.append(make_element("w:name", {"w:val": "Heading Linked Char"}))
    char_style.append(make_element("w:link", {"w:val": "HeadingLinked"}))
    rpr_char = make_element("w:rPr")
    rpr_char.append(make_element("w:i"))
    rpr_char.append(make_element("w:color", {"w:val": "FF0000"}))
    rpr_char.append(make_element("w:sz", {"w:val": "28"}))
    char_style.append(rpr_char)
    styles_element.append(char_style)

    # Paragraph 1: Uses "HeadingLinked" paragraph style
    p1 = doc.add_paragraph("This paragraph uses HeadingLinked style — should be italic, red, 14pt (from linked char style).")
    pPr1 = p1._p.get_or_add_pPr()
    pPr1.append(make_element("w:pStyle", {"w:val": "HeadingLinked"}))

    # Paragraph 2: Normal paragraph for comparison
    doc.add_paragraph("This is a normal paragraph for comparison.")

    save_fixture("style-constraints", "linked-styles", doc, {
        "name": "linked-styles",
        "spec_ref": "ISO 29500-1 §17.7.4.6",
        "description": (
            "Linked paragraph style 'HeadingLinked' (bold only) with w:link "
            "to 'HeadingLinkedChar' (italic + red + 14pt). Para style and "
            "char style have DIVERGENT rPr to prove the linked char style wins."
        ),
        "expected_behavior": (
            "Para 1: italic=on, color=FF0000, size=28 from linked character "
            "style. bold=off (para style's rPr is NOT used). "
            "Para 2: normal formatting."
        ),
        "current_status": "IMPLEMENTED — w:link parsed, linked char style rPr used for paragraph runs",
    })


# =========================================================================
# BASED ON TYPE MISMATCH (§17.7.4.3)
# =========================================================================

def make_based_on_type_mismatch() -> None:
    """basedOn cross-type: character style references paragraph style.

    ISO 29500-1 §17.7.4.3: basedOn must reference a style of the same
    type. A character style with basedOn pointing to a paragraph style
    should have that basedOn ignored — only CharStyleB's own properties
    should apply.
    """
    doc = Document()
    styles_element = doc.styles.element

    # Paragraph style "ParaStyleA" with bold + color=0000FF + size=32
    para_style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "ParaStyleA"})
    para_style.append(make_element("w:name", {"w:val": "Para Style A"}))
    rpr_a = make_element("w:rPr")
    rpr_a.append(make_element("w:b"))
    rpr_a.append(make_element("w:color", {"w:val": "0000FF"}))
    rpr_a.append(make_element("w:sz", {"w:val": "32"}))
    para_style.append(rpr_a)
    styles_element.append(para_style)

    # Character style "CharStyleB" with basedOn pointing to ParaStyleA (cross-type)
    # CharStyleB's own properties: italic + color=00FF00
    char_style = make_element("w:style", {"w:type": "character", "w:styleId": "CharStyleB"})
    char_style.append(make_element("w:name", {"w:val": "Char Style B"}))
    char_style.append(make_element("w:basedOn", {"w:val": "ParaStyleA"}))
    rpr_b = make_element("w:rPr")
    rpr_b.append(make_element("w:i"))
    rpr_b.append(make_element("w:color", {"w:val": "00FF00"}))
    char_style.append(rpr_b)
    styles_element.append(char_style)

    # Paragraph with a run using "CharStyleB"
    p = doc.add_paragraph()
    run = p.add_run("This run uses CharStyleB — should be italic+green, NOT bold+blue from ParaStyleA.")
    rPr = run._r.get_or_add_rPr()
    rPr.insert(0, make_element("w:rStyle", {"w:val": "CharStyleB"}))

    save_fixture("style-constraints", "based-on-type-mismatch", doc, {
        "name": "based-on-type-mismatch",
        "spec_ref": "ISO 29500-1 §17.7.4.3",
        "description": (
            "Character style 'CharStyleB' with basedOn='ParaStyleA' "
            "(cross-type mismatch). CharStyleB defines italic + green. "
            "ParaStyleA defines bold + blue + size=32."
        ),
        "expected_behavior": (
            "Run: italic=on, color=00FF00 (from CharStyleB only). "
            "bold=off, color!=0000FF, size!=32 (basedOn ignored for cross-type)."
        ),
        "current_status": "GAP — basedOn type checking not implemented, cross-type inheritance may leak",
    })


# =========================================================================
# TABLE CONDITIONAL PRECEDENCE — nwCell (§17.7.6.6)
# =========================================================================

def make_table_conditional_precedence() -> None:
    """Table conditional formatting: nwCell overrides firstRow and firstCol.

    ISO 29500-1 §17.7.6.6: When a cell matches a corner conditional
    (nwCell), it overrides both firstRow and firstCol formatting.

    The table style defines:
      wholeTable: color=black (000000)
      firstRow:   color=blue (0000FF), bold
      firstCol:   color=green (00FF00)
      nwCell:     color=red (FF0000), bold=off (w:b w:val="0")

    Cell (0,0) is the corner — should get nwCell formatting: red, NOT bold.
    """
    doc = Document()

    # wholeTable: black text
    whole_table_spr = make_element("w:tblStylePr", {"w:type": "wholeTable"})
    wt_rPr = make_element("w:rPr")
    wt_rPr.append(make_element("w:color", {"w:val": "000000"}))
    whole_table_spr.append(wt_rPr)

    # firstRow: blue text + bold
    first_row_spr = make_element("w:tblStylePr", {"w:type": "firstRow"})
    fr_rPr = make_element("w:rPr")
    fr_rPr.append(make_element("w:color", {"w:val": "0000FF"}))
    fr_rPr.append(make_element("w:b"))
    first_row_spr.append(fr_rPr)
    fr_tcPr = make_element("w:tcPr")
    fr_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "DDDDFF"}))
    first_row_spr.append(fr_tcPr)

    # firstCol: green text
    first_col_spr = make_element("w:tblStylePr", {"w:type": "firstCol"})
    fc_rPr = make_element("w:rPr")
    fc_rPr.append(make_element("w:color", {"w:val": "00FF00"}))
    first_col_spr.append(fc_rPr)
    fc_tcPr = make_element("w:tcPr")
    fc_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "DDFFDD"}))
    first_col_spr.append(fc_tcPr)

    # nwCell: red text, bold explicitly off
    nw_cell_spr = make_element("w:tblStylePr", {"w:type": "nwCell"})
    nw_rPr = make_element("w:rPr")
    nw_rPr.append(make_element("w:color", {"w:val": "FF0000"}))
    nw_rPr.append(make_element("w:b", {"w:val": "0"}))
    nw_cell_spr.append(nw_rPr)
    nw_tcPr = make_element("w:tcPr")
    nw_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "FFDDDD"}))
    nw_cell_spr.append(nw_tcPr)

    # Base table borders
    tbl_borders = make_element("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "4", "w:color": "000000",
        }))

    _add_table_style(doc, "CornerTestTable", "Corner Test Table",
                     tblPr_children=[tbl_borders],
                     tblStylePrs=[whole_table_spr, first_row_spr,
                                  first_col_spr, nw_cell_spr])

    # 3x3 table with firstRow + firstColumn enabled
    doc.add_paragraph("Table conditional precedence: nwCell overrides firstRow+firstCol.")
    tbl = doc.add_table(rows=3, cols=3)
    _apply_table_style(tbl, "CornerTestTable", {
        "w:val": "04A0",
        "w:firstRow": "1", "w:lastRow": "0",
        "w:firstColumn": "1", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "1",
    })
    _fill_table_cells(tbl, [
        ["Corner(0,0)", "Header(0,1)", "Header(0,2)"],
        ["Col(1,0)",    "Body(1,1)",   "Body(1,2)"],
        ["Col(2,0)",    "Body(2,1)",   "Body(2,2)"],
    ])

    save_fixture("style-constraints", "table-conditional-precedence", doc, {
        "name": "table-conditional-precedence",
        "spec_ref": "ISO 29500-1 §17.7.6.6",
        "description": (
            "Table with wholeTable (black), firstRow (blue+bold), "
            "firstCol (green), nwCell (red+not bold). "
            "3x3 table with firstRow and firstColumn enabled."
        ),
        "expected_behavior": (
            "Cell (0,0): nwCell wins — color=FF0000, bold=off, fill=FFDDDD. "
            "Cell (0,1): firstRow — color=0000FF, bold=on, fill=DDDDFF. "
            "Cell (1,0): firstCol — color=00FF00, fill=DDFFDD. "
            "Cell (1,1): wholeTable — color=000000."
        ),
        "current_status": "GAP — conditional merge order uses HashMap, nwCell may not override correctly",
    })


# =========================================================================
# SPACING LINE RULE DEFAULT (§17.3.1.33)
# =========================================================================

def make_spacing_line_rule_default() -> None:
    """Spacing lineRule default when omitted with line present.

    ISO 29500-1 §17.3.1.33: "If [lineRule] is omitted, then it shall be
    assumed to be of a value auto if a line attribute value is present."

    Style "ExactSpacing": line=480, lineRule=exact
    Para 1: uses style only — should inherit exact lineRule
    Para 2: uses style + direct w:spacing w:line="360" with NO lineRule
             — per spec, lineRule defaults to "auto", NOT inherited "exact"
    """
    doc = Document()
    styles_element = doc.styles.element

    # Remove existing docDefaults
    existing_defaults = styles_element.find(w("docDefaults"))
    if existing_defaults is not None:
        styles_element.remove(existing_defaults)

    # Minimal docDefaults
    doc_defaults = make_element("w:docDefaults")
    rpr_default = make_element("w:rPrDefault")
    rpr = make_element("w:rPr")
    rpr.append(make_element("w:sz", {"w:val": "24"}))
    rpr_default.append(rpr)
    doc_defaults.append(rpr_default)
    styles_element.insert(0, doc_defaults)

    # Style "ExactSpacing" with lineRule=exact
    style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "ExactSpacing"})
    style.append(make_element("w:name", {"w:val": "Exact Spacing"}))
    ppr = make_element("w:pPr")
    ppr.append(make_element("w:spacing", {
        "w:line": "480",
        "w:lineRule": "exact",
    }))
    style.append(ppr)
    styles_element.append(style)

    # Para 1: style only — should inherit exact lineRule
    p1 = doc.add_paragraph("Style only: line=480, lineRule=exact (both from style).")
    pPr1 = p1._p.get_or_add_pPr()
    pPr1.append(make_element("w:pStyle", {"w:val": "ExactSpacing"}))

    # Para 2: style + direct spacing with line=360, NO lineRule
    # Per spec: lineRule should default to "auto" (NOT inherited "exact")
    p2 = doc.add_paragraph("Direct line=360, no lineRule: should default to auto, NOT exact from style.")
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:pStyle", {"w:val": "ExactSpacing"}))
    pPr2.append(make_element("w:spacing", {"w:line": "360"}))

    save_fixture("style-constraints", "spacing-line-rule-default", doc, {
        "name": "spacing-line-rule-default",
        "spec_ref": "ISO 29500-1 §17.3.1.33",
        "description": (
            "Style 'ExactSpacing' defines line=480 lineRule=exact. "
            "Para 1: style only (inherits exact). "
            "Para 2: direct w:spacing w:line=360 with no lineRule attr."
        ),
        "expected_behavior": (
            "Para 1: line=480, lineRule=exact (from style). "
            "Para 2: line=360 (direct), lineRule=auto (spec default when "
            "line is present but lineRule is omitted — NOT exact from style)."
        ),
        "current_status": "GAP-016 — lineRule per-field merge falls through to style's 'exact'",
    })


# =========================================================================
# MAIN
# =========================================================================

def main() -> None:
    print("Generating style constraint fixtures:")
    make_linked_styles()
    make_based_on_type_mismatch()
    make_table_conditional_precedence()
    make_spacing_line_rule_default()
    print("\nDone.")


if __name__ == "__main__":
    main()

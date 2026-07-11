# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for annotation and tracked change constraint tests
(ISO 29500-1 §17.11.2, §17.16.2, §17.13.4).

These fixtures test cross-validation and balancing constraints:
  - footnote-ref-integrity:  footnoteReference with matching footnote story
  - field-balance:           properly balanced complex field (begin/separate/end)
  - comment-ref-integrity:   comment range + reference with matching comment story

Run:  uv run create_annotation_constraints.py
"""

import json
import zipfile
import io
from pathlib import Path
from lxml import etree

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

FOOTNOTES_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
FOOTNOTES_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"

COMMENTS_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
COMMENTS_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"

AREA = "annotation-constraints"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, out_buf: io.BytesIO, metadata: dict) -> None:
    out_dir = ROOT / AREA / name
    out_dir.mkdir(parents=True, exist_ok=True)
    (out_dir / "input.docx").write_bytes(out_buf.getvalue())
    (out_dir / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  {AREA}/{name}/")


# ── ZIP manipulation helpers ─────────────────────────────────────────────

def _add_footnotes_part(doc: DocxDocument, footnotes_xml: bytes) -> io.BytesIO:
    """Add word/footnotes.xml to the docx package and wire up the relationship."""
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"

    out_buf = io.BytesIO()
    with zipfile.ZipFile(buf, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        rels_data = zin.read("word/_rels/document.xml.rels")
        rels_root = etree.fromstring(rels_data)
        existing_ids = [el.get("Id") for el in rels_root]
        max_id = 0
        for rid in existing_ids:
            if rid and rid.startswith("rId"):
                try:
                    max_id = max(max_id, int(rid[3:]))
                except ValueError:
                    pass
        new_rid = f"rId{max_id + 1}"
        rel_el = etree.SubElement(rels_root, f"{{{rels_ns}}}Relationship")
        rel_el.set("Id", new_rid)
        rel_el.set("Type", FOOTNOTES_REL_TYPE)
        rel_el.set("Target", "footnotes.xml")
        new_rels_data = etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)

        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename == "[Content_Types].xml":
                ct_root = etree.fromstring(data)
                existing = ct_root.findall(f"{{{ct_ns}}}Override[@PartName='/word/footnotes.xml']")
                if not existing:
                    ov = etree.SubElement(ct_root, f"{{{ct_ns}}}Override")
                    ov.set("PartName", "/word/footnotes.xml")
                    ov.set("ContentType", FOOTNOTES_CONTENT_TYPE)
                data = etree.tostring(ct_root, xml_declaration=True, encoding="UTF-8", standalone=True)
            elif item.filename == "word/_rels/document.xml.rels":
                data = new_rels_data

            zout.writestr(item, data)

        zout.writestr("word/footnotes.xml", footnotes_xml)

    out_buf.seek(0)
    return out_buf


def _add_comments_part(doc: DocxDocument, comments_xml: bytes) -> io.BytesIO:
    """Add word/comments.xml to the docx package and wire up the relationship."""
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"

    out_buf = io.BytesIO()
    with zipfile.ZipFile(buf, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        rels_data = zin.read("word/_rels/document.xml.rels")
        rels_root = etree.fromstring(rels_data)
        existing_ids = [el.get("Id") for el in rels_root]
        max_id = 0
        for rid in existing_ids:
            if rid and rid.startswith("rId"):
                try:
                    max_id = max(max_id, int(rid[3:]))
                except ValueError:
                    pass
        new_rid = f"rId{max_id + 1}"
        rel_el = etree.SubElement(rels_root, f"{{{rels_ns}}}Relationship")
        rel_el.set("Id", new_rid)
        rel_el.set("Type", COMMENTS_REL_TYPE)
        rel_el.set("Target", "comments.xml")
        new_rels_data = etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)

        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename == "[Content_Types].xml":
                ct_root = etree.fromstring(data)
                existing = ct_root.findall(f"{{{ct_ns}}}Override[@PartName='/word/comments.xml']")
                if not existing:
                    ov = etree.SubElement(ct_root, f"{{{ct_ns}}}Override")
                    ov.set("PartName", "/word/comments.xml")
                    ov.set("ContentType", COMMENTS_CONTENT_TYPE)
                data = etree.tostring(ct_root, xml_declaration=True, encoding="UTF-8", standalone=True)
            elif item.filename == "word/_rels/document.xml.rels":
                data = new_rels_data

            zout.writestr(item, data)

        zout.writestr("word/comments.xml", comments_xml)

    out_buf.seek(0)
    return out_buf


# =========================================================================
# FOOTNOTE REFERENCE INTEGRITY (ISO 29500-1 §17.11.2, §17.3.3.9)
# =========================================================================

def make_footnote_ref_integrity() -> None:
    """Footnote reference with matching footnote story — happy path.

    ISO 29500-1 §17.11.2/§17.3.3.9: A footnoteReference in the body
    must have a corresponding footnote in footnotes.xml with the same ID.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Clear the paragraph
    for child in list(p._p):
        p._p.remove(child)

    # "Text with footnote" run
    r1 = make_element("w:r")
    t1 = make_element("w:t")
    t1.text = "Text with footnote"
    t1.set(qn("xml:space"), "preserve")
    r1.append(t1)
    p._p.append(r1)

    # footnoteReference run
    r2 = make_element("w:r")
    fn_ref = make_element("w:footnoteReference", {"w:id": "1"})
    r2.append(fn_ref)
    p._p.append(r2)

    # Build footnotes.xml with separator (id=0) and matching footnote (id=1)
    root = etree.Element(f"{{{W}}}footnotes", nsmap={"w": W, "r": R})

    fn0 = etree.SubElement(root, f"{{{W}}}footnote", attrib={
        f"{{{W}}}id": "0",
        f"{{{W}}}type": "separator",
    })
    p0 = etree.SubElement(fn0, f"{{{W}}}p")
    r0 = etree.SubElement(p0, f"{{{W}}}r")
    t0 = etree.SubElement(r0, f"{{{W}}}t")
    t0.text = ""

    fn1 = etree.SubElement(root, f"{{{W}}}footnote", attrib={f"{{{W}}}id": "1"})
    p1 = etree.SubElement(fn1, f"{{{W}}}p")
    r_fn = etree.SubElement(p1, f"{{{W}}}r")
    t_fn = etree.SubElement(r_fn, f"{{{W}}}t")
    t_fn.text = "This is the footnote"

    footnotes_xml = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)

    out_buf = _add_footnotes_part(doc, footnotes_xml)

    save_fixture("footnote-ref-integrity", out_buf, {
        "name": "footnote-ref-integrity",
        "spec_ref": "ISO 29500-1 §17.11.2, §17.3.3.9",
        "description": "Body paragraph with footnoteReference id=1, matching footnote in footnotes.xml",
        "expected_behavior": "FootnoteReference(id=1) in body, FootnoteStory(id=1) in CanonDoc.footnotes",
        "current_status": "SUPPORTED — happy path footnote cross-reference",
    })


# =========================================================================
# FIELD BALANCE (ISO 29500-1 §17.16.2)
# =========================================================================

def make_field_balance() -> None:
    """Properly balanced complex field: begin -> instrText -> separate -> result -> end.

    ISO 29500-1 §17.16.2: Complex fields must be balanced — every fldChar(begin)
    must be matched by a fldChar(separate) and fldChar(end).
    """
    doc = Document()
    p = doc.add_paragraph()

    # Clear the paragraph
    for child in list(p._p):
        p._p.remove(child)

    # fldChar begin
    r_begin = make_element("w:r")
    fld_begin = make_element("w:fldChar", {"w:fldCharType": "begin"})
    r_begin.append(fld_begin)
    p._p.append(r_begin)

    # instrText
    r_instr = make_element("w:r")
    instr_text = make_element("w:instrText")
    instr_text.text = " DATE "
    instr_text.set(qn("xml:space"), "preserve")
    r_instr.append(instr_text)
    p._p.append(r_instr)

    # fldChar separate
    r_sep = make_element("w:r")
    fld_sep = make_element("w:fldChar", {"w:fldCharType": "separate"})
    r_sep.append(fld_sep)
    p._p.append(r_sep)

    # result text
    r_result = make_element("w:r")
    t_result = make_element("w:t")
    t_result.text = "2025-01-01"
    r_result.append(t_result)
    p._p.append(r_result)

    # fldChar end
    r_end = make_element("w:r")
    fld_end = make_element("w:fldChar", {"w:fldCharType": "end"})
    r_end.append(fld_end)
    p._p.append(r_end)

    # Second paragraph: normal text (no fields)
    doc.add_paragraph("Normal text")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    save_fixture("field-balance", buf, {
        "name": "field-balance",
        "spec_ref": "ISO 29500-1 §17.16.2",
        "description": "Paragraph with balanced complex field (begin/instrText/separate/result/end) and a plain paragraph",
        "expected_behavior": "Field opaques in order: Begin, Instruction, Separate, End",
        "current_status": "SUPPORTED — complex field elements parsed as individual OpaqueInlines",
    })


# =========================================================================
# COMMENT REFERENCE INTEGRITY (ISO 29500-1 §17.13.4)
# =========================================================================

def make_comment_ref_integrity() -> None:
    """Comment with range markers and reference, matching comment story — happy path.

    ISO 29500-1 §17.13.4: commentRangeStart/End define the commented range,
    commentReference links the body location to the comment story.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Clear the paragraph
    for child in list(p._p):
        p._p.remove(child)

    # commentRangeStart
    range_start = make_element("w:commentRangeStart", {"w:id": "1"})
    p._p.append(range_start)

    # commented text
    run = make_element("w:r")
    t = make_element("w:t")
    t.text = "Commented text"
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    p._p.append(run)

    # commentRangeEnd
    range_end = make_element("w:commentRangeEnd", {"w:id": "1"})
    p._p.append(range_end)

    # commentReference
    ref_run = make_element("w:r")
    ref_el = make_element("w:commentReference", {"w:id": "1"})
    ref_run.append(ref_el)
    p._p.append(ref_run)

    # Build comments.xml
    comments_root = etree.Element(f"{{{W}}}comments", nsmap={"w": W, "r": R})
    comment_el = etree.SubElement(comments_root, f"{{{W}}}comment", attrib={
        f"{{{W}}}id": "1",
        f"{{{W}}}author": "Test",
        f"{{{W}}}date": "2025-01-01T00:00:00Z",
        f"{{{W}}}initials": "TE",
    })
    cp = etree.SubElement(comment_el, f"{{{W}}}p")
    cr = etree.SubElement(cp, f"{{{W}}}r")
    ct = etree.SubElement(cr, f"{{{W}}}t")
    ct.text = "My comment"
    comments_xml = etree.tostring(comments_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    out_buf = _add_comments_part(doc, comments_xml)

    save_fixture("comment-ref-integrity", out_buf, {
        "name": "comment-ref-integrity",
        "spec_ref": "ISO 29500-1 §17.13.4",
        "description": "Paragraph with commentRangeStart/End/Reference id=1 and matching comment in comments.xml",
        "expected_behavior": "CommentRangeStart, CommentRangeEnd, CommentReference(id=1) in body; CommentStory(id=1) in CanonDoc.comments",
        "current_status": "SUPPORTED — happy path comment cross-reference",
    })


# =========================================================================

def main() -> None:
    print("Generating annotation constraint spec-compliance fixtures:")
    make_footnote_ref_integrity()
    make_field_balance()
    make_comment_ref_integrity()
    print("Done.")


if __name__ == "__main__":
    main()

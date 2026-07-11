# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for move tracking and tracked change import fidelity tests.

These fixtures support tests in `stemma-engine/tests/spec_tracked_change_moves.rs`.

Fixtures:
  - reordered-paragraphs/    before+after pair with paragraphs B/C swapped
  - short-paragraphs/        before+after pair with short text (below move threshold)
  - multi-paragraph-move/    before+after pair with multiple paragraphs moved as a group
  - preexisting-insertions/  single DOCX with pre-existing w:ins tracked changes
  - preexisting-deletions/   single DOCX with pre-existing w:del tracked changes

Run:  uv run create_docs.py
"""

import json
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_pair(
    name: str,
    before: DocxDocument,
    after: DocxDocument,
    metadata: dict,
) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    before.save(str(out / "before.docx"))
    after.save(str(out / "after.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  {name}/")


def save_fixture(
    name: str,
    doc: DocxDocument,
    metadata: dict,
    filename: str = "input.docx",
) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / filename))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  {name}/")


# =========================================================================
# 1. reordered-paragraphs — paragraphs B/C swapped (text > 20 chars)
# =========================================================================

def make_reordered_paragraphs() -> None:
    """Before: A, B, C. After: A, C, B. Paragraphs have >20 chars for move detection."""
    before = Document()
    before.add_paragraph("This is paragraph Alpha which stays in place at the top of the document.")
    before.add_paragraph("This is paragraph Bravo and it contains enough text for the move detector.")
    before.add_paragraph("This is paragraph Charlie and it also has sufficient text for detection.")

    after = Document()
    after.add_paragraph("This is paragraph Alpha which stays in place at the top of the document.")
    after.add_paragraph("This is paragraph Charlie and it also has sufficient text for detection.")
    after.add_paragraph("This is paragraph Bravo and it contains enough text for the move detector.")

    save_pair("reordered-paragraphs", before, after, {
        "name": "reordered-paragraphs",
        "spec_ref": "ECMA-376 §17.13.5.22-28",
        "description": "Paragraphs B and C swapped; text exceeds 20 char move threshold",
        "expected_behavior": "diff_and_redline should produce moveFrom/moveTo (not del+ins)",
    })


# =========================================================================
# 2. short-paragraphs — text below 20-char move threshold
# =========================================================================

def make_short_paragraphs() -> None:
    """Before: AB, CD. After: CD, AB. Paragraphs have <20 chars — no move detection."""
    before = Document()
    before.add_paragraph("AB")
    before.add_paragraph("CD")

    after = Document()
    after.add_paragraph("CD")
    after.add_paragraph("AB")

    save_pair("short-paragraphs", before, after, {
        "name": "short-paragraphs",
        "spec_ref": "Implementation detail: MOVE_MIN_TEXT_LENGTH = 20",
        "description": "Paragraphs swapped but text is too short for move detection",
        "expected_behavior": "diff_and_redline should produce del+ins (not moveFrom/moveTo)",
    })


# =========================================================================
# 3. multi-paragraph-move — multiple paragraphs moved as a group
# =========================================================================

def make_multi_paragraph_move() -> None:
    """Move the conclusion paragraph from position 5 to position 2.

    Before: Intro, Detail-A, Detail-B, Detail-C, Conclusion
    After:  Intro, Conclusion, Detail-A, Detail-B, Detail-C

    The conclusion paragraph is moved (>20 chars). Detail paragraphs
    each have unique text exceeding the move detection threshold.
    """
    before = Document()
    before.add_paragraph("Introduction paragraph that sets up the document context and remains in place.")
    before.add_paragraph("Detail paragraph Alpha with unique content that exceeds twenty characters easily.")
    before.add_paragraph("Detail paragraph Bravo with distinct content that also exceeds the threshold.")
    before.add_paragraph("Detail paragraph Charlie with separate content that surpasses the limit too.")
    before.add_paragraph("Conclusion paragraph that wraps up the document and will be repositioned.")

    after = Document()
    after.add_paragraph("Introduction paragraph that sets up the document context and remains in place.")
    after.add_paragraph("Conclusion paragraph that wraps up the document and will be repositioned.")
    after.add_paragraph("Detail paragraph Alpha with unique content that exceeds twenty characters easily.")
    after.add_paragraph("Detail paragraph Bravo with distinct content that also exceeds the threshold.")
    after.add_paragraph("Detail paragraph Charlie with separate content that surpasses the limit too.")

    save_pair("multi-paragraph-move", before, after, {
        "name": "multi-paragraph-move",
        "spec_ref": "ECMA-376 §17.13.5.22-28",
        "description": "Multiple paragraphs moved as a group; each paragraph exceeds 20-char threshold",
        "expected_behavior": "All moved paragraphs should have matching moveFrom/moveTo range markers linked by name",
    })


# =========================================================================
# 4. preexisting-insertions — DOCX with pre-existing w:ins
# =========================================================================

def make_preexisting_insertions() -> None:
    """A DOCX with pre-existing w:ins around some text.

    Paragraph 1: "Hello " + [inserted: "wonderful "] + "world."
    The w:ins wrapper has author="Insert Author" and date="2025-03-01T12:00:00Z".
    """
    doc = Document()
    p = doc.add_paragraph()

    # Normal run: "Hello "
    p.add_run("Hello ")

    # Inserted run: "wonderful "
    ins_wrapper = make_element("w:ins", {
        "w:id": "100",
        "w:author": "Insert Author",
        "w:date": "2025-03-01T12:00:00Z",
    })
    run_ins = OxmlElement("w:r")
    t = make_element("w:t")
    t.text = "wonderful "
    t.set(qn("xml:space"), "preserve")
    run_ins.append(t)
    ins_wrapper.append(run_ins)
    p._p.append(ins_wrapper)

    # Normal run: "world."
    p.add_run("world.")

    save_fixture("preexisting-insertions", doc, {
        "name": "preexisting-insertions",
        "spec_ref": "ISO 29500-1 §17.13.5.18",
        "description": "Pre-existing w:ins with author/date around 'wonderful '",
        "expected_behavior": "import_docx().canonical should show TrackedSegment with Inserted status and RevisionInfo",
    })


# =========================================================================
# 5. preexisting-deletions — DOCX with pre-existing w:del
# =========================================================================

def make_preexisting_deletions() -> None:
    """A DOCX with pre-existing w:del around some text.

    Paragraph 1: "The " + [deleted: "original "] + "revised document."
    The w:del wrapper has author="Delete Author" and date="2025-04-15T09:30:00Z".
    """
    doc = Document()
    p = doc.add_paragraph()

    # Normal run: "The "
    p.add_run("The ")

    # Deleted run: "original "
    del_wrapper = make_element("w:del", {
        "w:id": "200",
        "w:author": "Delete Author",
        "w:date": "2025-04-15T09:30:00Z",
    })
    run_del = OxmlElement("w:r")
    del_t = make_element("w:delText")
    del_t.text = "original "
    del_t.set(qn("xml:space"), "preserve")
    run_del.append(del_t)
    del_wrapper.append(run_del)
    p._p.append(del_wrapper)

    # Normal run: "revised document."
    p.add_run("revised document.")

    save_fixture("preexisting-deletions", doc, {
        "name": "preexisting-deletions",
        "spec_ref": "ISO 29500-1 §17.13.5.4",
        "description": "Pre-existing w:del with w:delText and author/date around 'original '",
        "expected_behavior": "import_docx().canonical should show TrackedSegment with Deleted status and RevisionInfo",
    })


# =========================================================================
# 6. preexisting-moves — DOCX with pre-existing w:moveFrom/w:moveTo
# =========================================================================

def make_preexisting_moves() -> None:
    """A DOCX with pre-existing move tracking markup.

    Paragraph layout:
      P1: "Paragraph Alpha stays in place."  (normal)
      P2: moveTo[name="move_0"]: "Paragraph Charlie was moved here."  (inserted)
      P3: "Paragraph Bravo stays in place."  (normal)
      P4: moveFrom[name="move_0"]: "Paragraph Charlie was moved here."  (deleted)

    The moveFrom and moveTo are linked by w:name="move_0", and wrapped in
    moveFromRangeStart/End and moveToRangeStart/End markers respectively.
    """
    from lxml import etree

    doc = Document()
    body = doc.element.body

    # Remove default empty paragraph
    for p in body.findall(qn("w:p")):
        body.remove(p)

    ns = {"w": W}

    # P1: normal paragraph
    p1 = OxmlElement("w:p")
    r1 = OxmlElement("w:r")
    t1 = make_element("w:t")
    t1.text = "Paragraph Alpha stays in place."
    r1.append(t1)
    p1.append(r1)
    body.append(p1)

    # moveToRangeStart
    mtr_start = make_element("w:moveToRangeStart", {
        "w:id": "10",
        "w:name": "move_0",
        "w:author": "Move Author",
        "w:date": "2025-06-01T12:00:00Z",
    })
    body.append(mtr_start)

    # P2: moveTo container with paragraph
    move_to = make_element("w:moveTo", {
        "w:id": "200",
        "w:author": "Move Author",
        "w:date": "2025-06-01T12:00:00Z",
    })
    p2 = OxmlElement("w:p")
    r2 = OxmlElement("w:r")
    t2 = make_element("w:t")
    t2.text = "Paragraph Charlie was moved here."
    r2.append(t2)
    p2.append(r2)
    move_to.append(p2)
    body.append(move_to)

    # moveToRangeEnd
    mtr_end = make_element("w:moveToRangeEnd", {
        "w:id": "10",
    })
    body.append(mtr_end)

    # P3: normal paragraph
    p3 = OxmlElement("w:p")
    r3 = OxmlElement("w:r")
    t3 = make_element("w:t")
    t3.text = "Paragraph Bravo stays in place."
    r3.append(t3)
    p3.append(r3)
    body.append(p3)

    # moveFromRangeStart
    mfr_start = make_element("w:moveFromRangeStart", {
        "w:id": "11",
        "w:name": "move_0",
        "w:author": "Move Author",
        "w:date": "2025-06-01T12:00:00Z",
    })
    body.append(mfr_start)

    # P4: moveFrom container with paragraph
    move_from = make_element("w:moveFrom", {
        "w:id": "201",
        "w:author": "Move Author",
        "w:date": "2025-06-01T12:00:00Z",
    })
    p4 = OxmlElement("w:p")
    r4 = OxmlElement("w:r")
    t4 = make_element("w:delText")
    t4.text = "Paragraph Charlie was moved here."
    r4.append(t4)
    p4.append(r4)
    move_from.append(p4)
    body.append(move_from)

    # moveFromRangeEnd
    mfr_end = make_element("w:moveFromRangeEnd", {
        "w:id": "11",
    })
    body.append(mfr_end)

    save_fixture("preexisting-moves", doc, {
        "name": "preexisting-moves",
        "spec_ref": "ECMA-376 §17.13.5.21-28",
        "description": "Pre-existing moveFrom/moveTo with range markers linked by name='move_0'",
        "expected_behavior": "import_docx().canonical should preserve move_id on TrackedBlocks linking source and destination",
    })


# =========================================================================
# Main
# =========================================================================

if __name__ == "__main__":
    print("\n── Tracked Change Moves ──")
    make_reordered_paragraphs()
    make_short_paragraphs()
    make_multi_paragraph_move()
    make_preexisting_insertions()
    make_preexisting_deletions()
    make_preexisting_moves()
    print("\nDone.")

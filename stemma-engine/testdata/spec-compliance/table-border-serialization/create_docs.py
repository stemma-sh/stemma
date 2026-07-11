# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for paragraph border serialization ordering and
conditional formatting application order tests.

Run:  cd stemma-engine/testdata/spec-compliance/table-border-serialization && mise exec -- uv run create_docs.py
"""

import json
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc: DocxDocument, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  table-border-serialization/{name}/")


# =====================================================================
# Fixture 1: Paragraph borders (all 6 edges)
# =====================================================================

def make_paragraph_borders():
    """ECMA-376 Annex A, CT_PBdr: paragraph with all 6 border edges set.

    CT_PBdr sequence must be: top, left, bottom, right, between, bar.
    Each edge has a distinct color for easy identification.
    """
    doc = Document()

    para = doc.add_paragraph("Paragraph with all border edges.")
    pPr = para._p.get_or_add_pPr()

    pbdr = make_element("w:pBdr")
    # Emit in correct schema order: top, left, bottom, right, between, bar
    pbdr.append(make_element("w:top", {
        "w:val": "single", "w:sz": "4", "w:color": "FF0000", "w:space": "0",
    }))
    pbdr.append(make_element("w:left", {
        "w:val": "single", "w:sz": "4", "w:color": "00FF00", "w:space": "0",
    }))
    pbdr.append(make_element("w:bottom", {
        "w:val": "single", "w:sz": "4", "w:color": "0000FF", "w:space": "0",
    }))
    pbdr.append(make_element("w:right", {
        "w:val": "single", "w:sz": "4", "w:color": "FFFF00", "w:space": "0",
    }))
    pbdr.append(make_element("w:between", {
        "w:val": "single", "w:sz": "4", "w:color": "FF00FF", "w:space": "0",
    }))
    pbdr.append(make_element("w:bar", {
        "w:val": "single", "w:sz": "4", "w:color": "00FFFF", "w:space": "0",
    }))
    pPr.append(pbdr)

    save_fixture("paragraph-borders", doc, {
        "name": "paragraph-borders",
        "spec_ref": "ECMA-376 Annex A, CT_PBdr",
        "description": (
            "Single paragraph with all 6 border edges (top, left, bottom, "
            "right, between, bar) each with a distinct color. "
            "Used to verify serialization emits children in schema order: "
            "top, left, bottom, right, between, bar."
        ),
    })


# =====================================================================
# Fixture 2: Conditional formatting — firstRow beats firstCol at corner
# =====================================================================

def make_conditional_row_beats_column():
    """MS-OI29500 §2.1.246: Word applies conditional formatting in order:
    bands < firstCol/lastCol < firstRow/lastRow < corners.

    At cell (0,0), both firstRow AND firstCol apply.
    firstRow is applied AFTER firstCol, so firstRow shading wins.

    Table style:
      - firstRow: shading fill=FF0000 (red), borders top=thick sz=12 red
      - firstCol: shading fill=0000FF (blue), borders left=thick sz=12 blue

    tblLook: firstRow=1, firstColumn=1, lastRow=0, lastColumn=0,
             noHBand=1, noVBand=1
    """
    doc = Document()
    doc.add_paragraph("Conditional formatting: firstRow beats firstCol at corner.")

    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    styles_part_obj = doc.part.part_related_by(RT.STYLES)
    styles_el = styles_part_obj.element

    # Define a table style with firstRow and firstCol conditionals
    style = make_element("w:style", {"w:type": "table", "w:styleId": "RowBeatsColStyle"})
    style.append(make_element("w:name", {"w:val": "Row Beats Col Style"}))
    style.append(make_element("w:basedOn", {"w:val": "TableNormal"}))

    # Base table properties
    style_tbl_pr = make_element("w:tblPr")
    tbl_borders = make_element("w:tblBorders")
    for edge in ["top", "bottom", "left", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{edge}", {
            "w:val": "single", "w:sz": "4", "w:color": "000000", "w:space": "0",
        }))
    style_tbl_pr.append(tbl_borders)
    style.append(style_tbl_pr)

    # firstRow conditional: red shading + thick red top border
    fr = make_element("w:tblStylePr", {"w:type": "firstRow"})
    fr_tc = make_element("w:tcPr")
    fr_tc.append(make_element("w:shd", {
        "w:val": "clear", "w:color": "auto", "w:fill": "FF0000",
    }))
    fr_borders = make_element("w:tcBorders")
    fr_borders.append(make_element("w:top", {
        "w:val": "thick", "w:sz": "12", "w:color": "FF0000", "w:space": "0",
    }))
    fr_tc.append(fr_borders)
    fr.append(fr_tc)
    style.append(fr)

    # firstCol conditional: blue shading + thick blue left border
    fc = make_element("w:tblStylePr", {"w:type": "firstCol"})
    fc_tc = make_element("w:tcPr")
    fc_tc.append(make_element("w:shd", {
        "w:val": "clear", "w:color": "auto", "w:fill": "0000FF",
    }))
    fc_borders = make_element("w:tcBorders")
    fc_borders.append(make_element("w:left", {
        "w:val": "thick", "w:sz": "12", "w:color": "0000FF", "w:space": "0",
    }))
    fc_tc.append(fc_borders)
    fc.append(fc_tc)
    style.append(fc)

    styles_el.append(style)

    # Create 3x3 table using this style
    tbl = doc.add_table(rows=3, cols=3)
    for r in range(3):
        for c in range(3):
            tbl.cell(r, c).text = f"R{r}C{c}"

    tbl_el = tbl._tbl
    tblPr = tbl_el.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_el.insert(0, tblPr)

    # Remove any tblLook/tblStyle python-docx may have added
    for existing in tblPr.findall(w("tblLook")):
        tblPr.remove(existing)
    for existing in tblPr.findall(w("tblStyle")):
        tblPr.remove(existing)

    # Reference the style
    tbl_style_el = make_element("w:tblStyle", {"w:val": "RowBeatsColStyle"})
    tblPr.insert(0, tbl_style_el)

    # tblLook: firstRow=1, firstColumn=1, noHBand=1, noVBand=1
    tbl_look = make_element("w:tblLook", {
        "w:val": "00A0",
        "w:firstRow": "1",
        "w:lastRow": "0",
        "w:firstColumn": "1",
        "w:lastColumn": "0",
        "w:noHBand": "1",
        "w:noVBand": "1",
    })
    tblPr.append(tbl_look)

    save_fixture("conditional-row-beats-column", doc, {
        "name": "conditional-row-beats-column",
        "spec_ref": "MS-OI29500 §2.1.246",
        "description": (
            "3x3 table with style that has firstRow conditional (red shading FF0000) "
            "and firstCol conditional (blue shading 0000FF). "
            "tblLook enables firstRow and firstColumn. "
            "Per MS-OI29500 §2.1.246, firstRow is applied AFTER firstCol, "
            "so at corner cell (0,0) firstRow shading (red) should win."
        ),
    })


# =====================================================================
# Main
# =====================================================================

def main():
    print("\n== Table Border Serialization Fixtures ==")
    make_paragraph_borders()
    make_conditional_row_beats_column()
    print("\nDone.")


if __name__ == "__main__":
    main()

# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for underline style roundtrip fidelity tests.

Roundtrip fixtures (tests 1-4):
  Each has a before.docx and after.docx pair.  Both share the same
  underlined paragraph; the second paragraph differs to force
  diff_and_redline through the serialize path.

Import fixture (test 5):
  A single input.docx with multiple underline styles in one paragraph.

Run:  uv run create_docs.py
"""

import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    """Create an OxmlElement with optional attributes."""
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc, metadata: dict, filename: str = "input.docx") -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / filename))
    if filename == "input.docx" or filename == "before.docx":
        (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  run-underline-fidelity/{name}/{filename}")


def make_roundtrip_pair(style: str) -> None:
    """before/after pair: underlined text stays, plain text changes."""
    # before.docx
    before = Document()
    p1 = before.add_paragraph()
    run = p1.add_run(f"{style}-underline text")
    rPr = run._r.get_or_add_rPr()
    rPr.append(make_element("w:u", {"w:val": style}))
    before.add_paragraph("Base paragraph for diff anchor.")

    save_fixture(f"roundtrip-{style}", before, {
        "name": f"roundtrip-{style}",
        "spec_ref": "ISO 29500-1 section 17.3.2.40 / section 17.18.99 ST_Underline",
        "description": (
            f"Before/after pair: paragraph with w:u w:val=\"{style}\" is "
            "unchanged; second paragraph differs to trigger serialization."
        ),
        "expected_behavior": (
            f"After diff_and_redline + export, the serialized XML must "
            f"contain w:u w:val=\"{style}\", not \"single\"."
        ),
    }, filename="before.docx")

    # after.docx — same underlined paragraph, different second paragraph
    after = Document()
    p1a = after.add_paragraph()
    run_a = p1a.add_run(f"{style}-underline text")
    rPr_a = run_a._r.get_or_add_rPr()
    rPr_a.append(make_element("w:u", {"w:val": style}))
    after.add_paragraph("Modified paragraph for diff anchor.")

    save_fixture(f"roundtrip-{style}", after, {}, filename="after.docx")


def make_multi_style_fixture() -> None:
    """One paragraph with runs for several underline styles."""
    doc = Document()
    p = doc.add_paragraph()

    styles = ["single", "double", "dotted", "dash", "wave",
              "thick", "dashLong", "dotDash", "wavyDouble"]
    for style in styles:
        run = p.add_run(f"{style}-underline ")
        rPr = run._r.get_or_add_rPr()
        rPr.append(make_element("w:u", {"w:val": style}))

    save_fixture("multi-style", doc, {
        "name": "multi-style",
        "spec_ref": "ISO 29500-1 section 17.3.2.40 / section 17.18.99 ST_Underline",
        "description": "Paragraph with runs using 9 different underline styles",
        "expected_behavior": (
            "Each run's style_props.underline_style must match the w:val used. "
            "All runs must carry Mark::Underline."
        ),
    })


if __name__ == "__main__":
    print("Generating underline fidelity fixtures:")

    for style in ["double", "dotted", "wave", "dash"]:
        make_roundtrip_pair(style)

    make_multi_style_fixture()

    print("Done.")

# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for section constraint / behavioral tests.

These fixtures exercise header/footer inheritance across sections,
titlePg interaction, and continuous section property inheritance
(ISO 29500-1 sections 17.6.17, 17.10.2, 17.10.5, 17.10.6).

Run:  uv run create_section_constraints.py
"""

import json
import zipfile
import io
from pathlib import Path
from lxml import etree

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

HEADER_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"
HEADER_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/header"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    """Create an OxmlElement with optional attributes."""
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def _build_header_xml(paragraphs: list[str]) -> bytes:
    """Build a w:hdr XML part with the given paragraph texts."""
    root = etree.Element(f"{{{W}}}hdr", nsmap={"w": W, "r": R})
    for text in paragraphs:
        p = etree.SubElement(root, f"{{{W}}}p")
        r = etree.SubElement(p, f"{{{W}}}r")
        t = etree.SubElement(r, f"{{{W}}}t")
        t.text = text
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def save_fixture(
    area: str,
    name: str,
    data: bytes | io.BytesIO,
    metadata: dict,
    filename: str = "input.docx",
) -> None:
    """Save a fixture file and its metadata."""
    out = ROOT / area / name
    out.mkdir(parents=True, exist_ok=True)
    if isinstance(data, io.BytesIO):
        (out / filename).write_bytes(data.getvalue())
    else:
        (out / filename).write_bytes(data)
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  {area}/{name}/")


# =========================================================================
# HEADER INHERITANCE (ISO 29500-1 sections 17.10.2, 17.10.5)
# =========================================================================

def make_header_inheritance() -> None:
    """Two sections: Section 1 defines a default header, Section 2 does not.

    ISO 29500-1 sections 17.10.2 / 17.10.5: If a section does not define a header
    reference, it inherits the header from the previous section. This fixture
    tests that behavioral constraint.
    """
    doc = Document()
    doc.add_paragraph("Section 1 content")

    hdr1 = _build_header_xml(["Section 1 Header"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    out_buf = io.BytesIO()
    rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"

    with zipfile.ZipFile(buf, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        rels_data = zin.read("word/_rels/document.xml.rels")
        rels_root = etree.fromstring(rels_data)
        existing_ids = [el.get("Id") for el in rels_root]
        max_id = 0
        for rid in existing_ids:
            if rid and rid.startswith("rId"):
                try:
                    max_id = max(max_id, int(rid[3:]))
                except ValueError:
                    pass

        rid_hdr1 = f"rId{max_id + 1}"
        rel_el = etree.SubElement(rels_root, f"{{{rels_ns}}}Relationship")
        rel_el.set("Id", rid_hdr1)
        rel_el.set("Type", HEADER_REL_TYPE)
        rel_el.set("Target", "header1.xml")
        new_rels_data = etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)

        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename == "[Content_Types].xml":
                ct_root = etree.fromstring(data)
                pn = "/word/header1.xml"
                existing = ct_root.findall(f"{{{ct_ns}}}Override[@PartName='{pn}']")
                if not existing:
                    ov = etree.SubElement(ct_root, f"{{{ct_ns}}}Override")
                    ov.set("PartName", pn)
                    ov.set("ContentType", HEADER_CONTENT_TYPE)
                data = etree.tostring(ct_root, xml_declaration=True, encoding="UTF-8", standalone=True)

            elif item.filename == "word/_rels/document.xml.rels":
                data = new_rels_data

            elif item.filename == "word/document.xml":
                doc_root = etree.fromstring(data)
                body = doc_root.find(f"{{{W}}}body")
                # Remove existing paragraphs and sectPr
                for child in list(body):
                    body.remove(child)

                # Section 1: paragraph + sectPr with header reference and nextPage break
                p1 = etree.SubElement(body, f"{{{W}}}p")
                r1 = etree.SubElement(p1, f"{{{W}}}r")
                t1 = etree.SubElement(r1, f"{{{W}}}t")
                t1.text = "Section 1 content"

                p1b = etree.SubElement(body, f"{{{W}}}p")
                pPr1 = etree.SubElement(p1b, f"{{{W}}}pPr")
                s1 = etree.SubElement(pPr1, f"{{{W}}}sectPr")
                etree.SubElement(s1, f"{{{W}}}type", attrib={f"{{{W}}}val": "nextPage"})
                h1_ref = etree.SubElement(s1, f"{{{W}}}headerReference")
                h1_ref.set(f"{{{R}}}id", rid_hdr1)
                h1_ref.set(f"{{{W}}}type", "default")

                # Section 2: paragraph, NO sectPr with header reference
                # (should inherit header from Section 1 per spec)
                p2 = etree.SubElement(body, f"{{{W}}}p")
                r2 = etree.SubElement(p2, f"{{{W}}}r")
                t2 = etree.SubElement(r2, f"{{{W}}}t")
                t2.text = "Section 2 content"

                # Body-level sectPr (final section) — no header reference
                body_sect = etree.SubElement(body, f"{{{W}}}sectPr")
                etree.SubElement(body_sect, f"{{{W}}}pgSz",
                                 attrib={f"{{{W}}}w": "12240", f"{{{W}}}h": "15840"})

                data = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)

            zout.writestr(item, data)
        zout.writestr("word/header1.xml", hdr1)

    out_buf.seek(0)
    save_fixture("section-constraints", "header-inheritance", out_buf, {
        "name": "header-inheritance",
        "spec_ref": "ISO 29500-1 sections 17.10.2, 17.10.5",
        "description": "Section 1 defines a default header 'Section 1 Header'; Section 2 has no header reference (should inherit from Section 1 per spec)",
        "expected_behavior": "Section 2 should inherit the header from Section 1. The document should have exactly 1 header part.",
        "current_status": "PARTIAL - header part is parsed, but cross-section inheritance is not computed in the model",
    })


# =========================================================================
# TITLE PAGE INTERACTION (ISO 29500-1 sections 17.10.6, 17.6.18)
# =========================================================================

def make_title_page_interaction() -> None:
    """Section with titlePg=true, both first-page and default headers.

    ISO 29500-1 sections 17.10.6 / 17.6.18: When w:titlePg is present in the
    sectPr, the first page of the section uses the first-page header instead
    of the default header. Subsequent pages use the default header.
    """
    doc = Document()
    doc.add_paragraph("Title page test")

    default_hdr = _build_header_xml(["Default Header"])
    first_hdr = _build_header_xml(["First Page Header"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    out_buf = io.BytesIO()
    rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"

    with zipfile.ZipFile(buf, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        rels_data = zin.read("word/_rels/document.xml.rels")
        rels_root = etree.fromstring(rels_data)
        existing_ids = [el.get("Id") for el in rels_root]
        max_id = 0
        for rid in existing_ids:
            if rid and rid.startswith("rId"):
                try:
                    max_id = max(max_id, int(rid[3:]))
                except ValueError:
                    pass
        rid_default = f"rId{max_id + 1}"
        rid_first = f"rId{max_id + 2}"
        for new_rid, fname in [(rid_default, "header1.xml"), (rid_first, "header2.xml")]:
            rel_el = etree.SubElement(rels_root, f"{{{rels_ns}}}Relationship")
            rel_el.set("Id", new_rid)
            rel_el.set("Type", HEADER_REL_TYPE)
            rel_el.set("Target", fname)
        new_rels_data = etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)

        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename == "[Content_Types].xml":
                ct_root = etree.fromstring(data)
                for i in range(1, 3):
                    pn = f"/word/header{i}.xml"
                    existing = ct_root.findall(f"{{{ct_ns}}}Override[@PartName='{pn}']")
                    if not existing:
                        ov = etree.SubElement(ct_root, f"{{{ct_ns}}}Override")
                        ov.set("PartName", pn)
                        ov.set("ContentType", HEADER_CONTENT_TYPE)
                data = etree.tostring(ct_root, xml_declaration=True, encoding="UTF-8", standalone=True)

            elif item.filename == "word/_rels/document.xml.rels":
                data = new_rels_data

            elif item.filename == "word/document.xml":
                doc_root = etree.fromstring(data)
                # Add titlePg + header references to the body-level sectPr
                for sect_pr in doc_root.iter(f"{{{W}}}sectPr"):
                    etree.SubElement(sect_pr, f"{{{W}}}titlePg")
                    h_def = etree.SubElement(sect_pr, f"{{{W}}}headerReference")
                    h_def.set(f"{{{R}}}id", rid_default)
                    h_def.set(f"{{{W}}}type", "default")
                    h_first = etree.SubElement(sect_pr, f"{{{W}}}headerReference")
                    h_first.set(f"{{{R}}}id", rid_first)
                    h_first.set(f"{{{W}}}type", "first")
                data = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)

            zout.writestr(item, data)
        zout.writestr("word/header1.xml", default_hdr)
        zout.writestr("word/header2.xml", first_hdr)

    out_buf.seek(0)
    save_fixture("section-constraints", "title-page-interaction", out_buf, {
        "name": "title-page-interaction",
        "spec_ref": "ISO 29500-1 sections 17.10.6, 17.6.18",
        "description": "Section with titlePg=true, default header 'Default Header' and first-page header 'First Page Header'",
        "expected_behavior": "First page shows 'First Page Header', subsequent pages show 'Default Header'. body_section_properties.title_page should be Some(true).",
        "current_status": "PARTIAL - headers are parsed, titlePg flag is parsed, but renderer does not select headers based on titlePg",
    })


# =========================================================================
# CONTINUOUS SECTION INHERITANCE (ISO 29500-1 sections 17.6.17)
# =========================================================================

def make_continuous_section() -> None:
    """Two sections: Section 1 with explicit margins, Section 2 continuous with no margins.

    ISO 29500-1 sections 17.6.17: A continuous section break means the new section
    continues on the same page. If the continuous section does not specify its
    own page properties (margins, page size), it inherits them from the
    previous section.
    """
    doc = Document()

    # Build document XML manually for precise control
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    out_buf = io.BytesIO()

    with zipfile.ZipFile(buf, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename == "word/document.xml":
                doc_root = etree.fromstring(data)
                body = doc_root.find(f"{{{W}}}body")
                # Remove existing content
                for child in list(body):
                    body.remove(child)

                # Section 1: paragraph + sectPr with explicit margins and nextPage break
                p1 = etree.SubElement(body, f"{{{W}}}p")
                r1 = etree.SubElement(p1, f"{{{W}}}r")
                t1 = etree.SubElement(r1, f"{{{W}}}t")
                t1.text = "Wide margins"

                p1b = etree.SubElement(body, f"{{{W}}}p")
                pPr1 = etree.SubElement(p1b, f"{{{W}}}pPr")
                s1 = etree.SubElement(pPr1, f"{{{W}}}sectPr")
                etree.SubElement(s1, f"{{{W}}}pgSz",
                                 attrib={f"{{{W}}}w": "12240", f"{{{W}}}h": "15840"})
                etree.SubElement(s1, f"{{{W}}}pgMar", attrib={
                    f"{{{W}}}top": "1440",
                    f"{{{W}}}bottom": "1440",
                    f"{{{W}}}left": "1800",
                    f"{{{W}}}right": "1800",
                    f"{{{W}}}header": "720",
                    f"{{{W}}}footer": "720",
                    f"{{{W}}}gutter": "0",
                })
                etree.SubElement(s1, f"{{{W}}}type", attrib={f"{{{W}}}val": "nextPage"})

                # Section 2: paragraph + body-level sectPr with continuous break, NO margins
                p2 = etree.SubElement(body, f"{{{W}}}p")
                r2 = etree.SubElement(p2, f"{{{W}}}r")
                t2 = etree.SubElement(r2, f"{{{W}}}t")
                t2.text = "Inherited margins"

                # Body-level sectPr — continuous type, no pgMar (should inherit)
                body_sect = etree.SubElement(body, f"{{{W}}}sectPr")
                etree.SubElement(body_sect, f"{{{W}}}pgSz",
                                 attrib={f"{{{W}}}w": "12240", f"{{{W}}}h": "15840"})
                etree.SubElement(body_sect, f"{{{W}}}type",
                                 attrib={f"{{{W}}}val": "continuous"})

                data = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)

            zout.writestr(item, data)

    out_buf.seek(0)
    save_fixture("section-constraints", "continuous-section", out_buf, {
        "name": "continuous-section",
        "spec_ref": "ISO 29500-1 sections 17.6.17",
        "description": "Section 1 has margins (top=1440, bottom=1440, left=1800, right=1800); Section 2 is continuous with no margins specified (should inherit from Section 1)",
        "expected_behavior": "Section 2 (continuous) should inherit page margins from Section 1. margin_top should be 1440, margin_left should be 1800.",
        "current_status": "GAP - continuous section property inheritance is not implemented; Section 2 margins will be None",
    })


# =========================================================================
# PARTIAL HEADER INHERITANCE (ISO 29500-1 sections 17.10.2)
# =========================================================================

def make_partial_header_inheritance() -> None:
    """Three sections with partial header declarations to test per-kind inheritance.

    - S1 sectPr: declares default (header1.xml) + first (header2.xml)
    - S2 sectPr: declares only first (header3.xml) — should inherit default from S1
    - S3 (body sectPr): declares nothing — should inherit both from S2
    """
    doc = Document()
    doc.add_paragraph("placeholder")

    hdr1 = _build_header_xml(["Default Header S1"])
    hdr2 = _build_header_xml(["First Header S1"])
    hdr3 = _build_header_xml(["First Header S2"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    out_buf = io.BytesIO()
    rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"

    with zipfile.ZipFile(buf, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        rels_data = zin.read("word/_rels/document.xml.rels")
        rels_root = etree.fromstring(rels_data)
        existing_ids = [el.get("Id") for el in rels_root]
        max_id = 0
        for rid in existing_ids:
            if rid and rid.startswith("rId"):
                try:
                    max_id = max(max_id, int(rid[3:]))
                except ValueError:
                    pass

        rid_hdr1 = f"rId{max_id + 1}"
        rid_hdr2 = f"rId{max_id + 2}"
        rid_hdr3 = f"rId{max_id + 3}"
        for new_rid, fname in [
            (rid_hdr1, "header1.xml"),
            (rid_hdr2, "header2.xml"),
            (rid_hdr3, "header3.xml"),
        ]:
            rel_el = etree.SubElement(rels_root, f"{{{rels_ns}}}Relationship")
            rel_el.set("Id", new_rid)
            rel_el.set("Type", HEADER_REL_TYPE)
            rel_el.set("Target", fname)
        new_rels_data = etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)

        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename == "[Content_Types].xml":
                ct_root = etree.fromstring(data)
                for i in range(1, 4):
                    pn = f"/word/header{i}.xml"
                    existing = ct_root.findall(f"{{{ct_ns}}}Override[@PartName='{pn}']")
                    if not existing:
                        ov = etree.SubElement(ct_root, f"{{{ct_ns}}}Override")
                        ov.set("PartName", pn)
                        ov.set("ContentType", HEADER_CONTENT_TYPE)
                data = etree.tostring(ct_root, xml_declaration=True, encoding="UTF-8", standalone=True)

            elif item.filename == "word/_rels/document.xml.rels":
                data = new_rels_data

            elif item.filename == "word/document.xml":
                doc_root = etree.fromstring(data)
                body = doc_root.find(f"{{{W}}}body")
                for child in list(body):
                    body.remove(child)

                # Section 1: paragraph + sectPr with default + first headers + titlePg
                p1 = etree.SubElement(body, f"{{{W}}}p")
                r1 = etree.SubElement(p1, f"{{{W}}}r")
                t1 = etree.SubElement(r1, f"{{{W}}}t")
                t1.text = "Section 1"

                p1b = etree.SubElement(body, f"{{{W}}}p")
                pPr1 = etree.SubElement(p1b, f"{{{W}}}pPr")
                s1 = etree.SubElement(pPr1, f"{{{W}}}sectPr")
                etree.SubElement(s1, f"{{{W}}}titlePg")
                etree.SubElement(s1, f"{{{W}}}type", attrib={f"{{{W}}}val": "nextPage"})
                h1_ref = etree.SubElement(s1, f"{{{W}}}headerReference")
                h1_ref.set(f"{{{R}}}id", rid_hdr1)
                h1_ref.set(f"{{{W}}}type", "default")
                h2_ref = etree.SubElement(s1, f"{{{W}}}headerReference")
                h2_ref.set(f"{{{R}}}id", rid_hdr2)
                h2_ref.set(f"{{{W}}}type", "first")

                # Section 2: paragraph + sectPr with only first header (overrides S1's first) + titlePg
                p2 = etree.SubElement(body, f"{{{W}}}p")
                r2 = etree.SubElement(p2, f"{{{W}}}r")
                t2 = etree.SubElement(r2, f"{{{W}}}t")
                t2.text = "Section 2"

                p2b = etree.SubElement(body, f"{{{W}}}p")
                pPr2 = etree.SubElement(p2b, f"{{{W}}}pPr")
                s2 = etree.SubElement(pPr2, f"{{{W}}}sectPr")
                etree.SubElement(s2, f"{{{W}}}titlePg")
                etree.SubElement(s2, f"{{{W}}}type", attrib={f"{{{W}}}val": "nextPage"})
                h3_ref = etree.SubElement(s2, f"{{{W}}}headerReference")
                h3_ref.set(f"{{{R}}}id", rid_hdr3)
                h3_ref.set(f"{{{W}}}type", "first")

                # Section 3: paragraph + body-level sectPr with NO headers + titlePg
                # (titlePg needed so inherited First refs are preserved)
                p3 = etree.SubElement(body, f"{{{W}}}p")
                r3 = etree.SubElement(p3, f"{{{W}}}r")
                t3 = etree.SubElement(r3, f"{{{W}}}t")
                t3.text = "Section 3"

                body_sect = etree.SubElement(body, f"{{{W}}}sectPr")
                etree.SubElement(body_sect, f"{{{W}}}titlePg")
                etree.SubElement(body_sect, f"{{{W}}}pgSz",
                                 attrib={f"{{{W}}}w": "12240", f"{{{W}}}h": "15840"})

                data = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)

            zout.writestr(item, data)
        zout.writestr("word/header1.xml", hdr1)
        zout.writestr("word/header2.xml", hdr2)
        zout.writestr("word/header3.xml", hdr3)

    out_buf.seek(0)
    save_fixture("section-constraints", "partial-header-inheritance", out_buf, {
        "name": "partial-header-inheritance",
        "spec_ref": "ISO 29500-1 sections 17.10.2",
        "description": "S1 declares default+first headers; S2 declares only first (inherits default from S1); S3 declares nothing (inherits both from S2)",
        "expected_behavior": "Per-kind inheritance: S2 inherits default from S1, S3 inherits both from S2. S2's first != S1's first (overridden).",
    })


# =========================================================================
# TITLE PAGE ABSENT / FALSE (ISO 29500-1 §17.10.6)
# =========================================================================

def make_title_page_absent() -> None:
    """Section with default + first-page headers but NO titlePg element.

    ISO 29500-1 §17.10.6: When w:titlePg is absent, first-page headers/footers
    SHALL NOT be shown. The import pipeline should filter out First-kind refs
    and stories.
    """
    doc = Document()
    doc.add_paragraph("No title page flag")

    default_hdr = _build_header_xml(["Default Header"])
    first_hdr = _build_header_xml(["First Page Header"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    out_buf = io.BytesIO()
    rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"

    with zipfile.ZipFile(buf, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        rels_data = zin.read("word/_rels/document.xml.rels")
        rels_root = etree.fromstring(rels_data)
        existing_ids = [el.get("Id") for el in rels_root]
        max_id = 0
        for rid in existing_ids:
            if rid and rid.startswith("rId"):
                try:
                    max_id = max(max_id, int(rid[3:]))
                except ValueError:
                    pass
        rid_default = f"rId{max_id + 1}"
        rid_first = f"rId{max_id + 2}"
        for new_rid, fname in [(rid_default, "header1.xml"), (rid_first, "header2.xml")]:
            rel_el = etree.SubElement(rels_root, f"{{{rels_ns}}}Relationship")
            rel_el.set("Id", new_rid)
            rel_el.set("Type", HEADER_REL_TYPE)
            rel_el.set("Target", fname)
        new_rels_data = etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)

        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename == "[Content_Types].xml":
                ct_root = etree.fromstring(data)
                for i in range(1, 3):
                    pn = f"/word/header{i}.xml"
                    existing = ct_root.findall(f"{{{ct_ns}}}Override[@PartName='{pn}']")
                    if not existing:
                        ov = etree.SubElement(ct_root, f"{{{ct_ns}}}Override")
                        ov.set("PartName", pn)
                        ov.set("ContentType", HEADER_CONTENT_TYPE)
                data = etree.tostring(ct_root, xml_declaration=True, encoding="UTF-8", standalone=True)

            elif item.filename == "word/_rels/document.xml.rels":
                data = new_rels_data

            elif item.filename == "word/document.xml":
                doc_root = etree.fromstring(data)
                # Add header references to the body-level sectPr but NO titlePg
                for sect_pr in doc_root.iter(f"{{{W}}}sectPr"):
                    h_def = etree.SubElement(sect_pr, f"{{{W}}}headerReference")
                    h_def.set(f"{{{R}}}id", rid_default)
                    h_def.set(f"{{{W}}}type", "default")
                    h_first = etree.SubElement(sect_pr, f"{{{W}}}headerReference")
                    h_first.set(f"{{{R}}}id", rid_first)
                    h_first.set(f"{{{W}}}type", "first")
                data = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)

            zout.writestr(item, data)
        zout.writestr("word/header1.xml", default_hdr)
        zout.writestr("word/header2.xml", first_hdr)

    out_buf.seek(0)
    save_fixture("section-constraints", "title-page-absent", out_buf, {
        "name": "title-page-absent",
        "spec_ref": "ISO 29500-1 §17.10.6",
        "description": "Section with default + first-page headers but NO titlePg element. First-page header should be filtered out.",
        "expected_behavior": "First-kind header refs and stories should be filtered from the model since titlePg is absent.",
    })


def make_title_page_false() -> None:
    """Section with default + first-page headers and titlePg val='false'.

    ISO 29500-1 §17.10.6: When w:titlePg is explicitly false, first-page
    headers/footers SHALL NOT be shown. Same behavior as absent.
    """
    doc = Document()
    doc.add_paragraph("Explicit false title page")

    default_hdr = _build_header_xml(["Default Header"])
    first_hdr = _build_header_xml(["First Page Header"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    out_buf = io.BytesIO()
    rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"

    with zipfile.ZipFile(buf, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        rels_data = zin.read("word/_rels/document.xml.rels")
        rels_root = etree.fromstring(rels_data)
        existing_ids = [el.get("Id") for el in rels_root]
        max_id = 0
        for rid in existing_ids:
            if rid and rid.startswith("rId"):
                try:
                    max_id = max(max_id, int(rid[3:]))
                except ValueError:
                    pass
        rid_default = f"rId{max_id + 1}"
        rid_first = f"rId{max_id + 2}"
        for new_rid, fname in [(rid_default, "header1.xml"), (rid_first, "header2.xml")]:
            rel_el = etree.SubElement(rels_root, f"{{{rels_ns}}}Relationship")
            rel_el.set("Id", new_rid)
            rel_el.set("Type", HEADER_REL_TYPE)
            rel_el.set("Target", fname)
        new_rels_data = etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)

        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename == "[Content_Types].xml":
                ct_root = etree.fromstring(data)
                for i in range(1, 3):
                    pn = f"/word/header{i}.xml"
                    existing = ct_root.findall(f"{{{ct_ns}}}Override[@PartName='{pn}']")
                    if not existing:
                        ov = etree.SubElement(ct_root, f"{{{ct_ns}}}Override")
                        ov.set("PartName", pn)
                        ov.set("ContentType", HEADER_CONTENT_TYPE)
                data = etree.tostring(ct_root, xml_declaration=True, encoding="UTF-8", standalone=True)

            elif item.filename == "word/_rels/document.xml.rels":
                data = new_rels_data

            elif item.filename == "word/document.xml":
                doc_root = etree.fromstring(data)
                for sect_pr in doc_root.iter(f"{{{W}}}sectPr"):
                    # titlePg with val="false" (explicit false)
                    tp = etree.SubElement(sect_pr, f"{{{W}}}titlePg")
                    tp.set(f"{{{W}}}val", "false")
                    h_def = etree.SubElement(sect_pr, f"{{{W}}}headerReference")
                    h_def.set(f"{{{R}}}id", rid_default)
                    h_def.set(f"{{{W}}}type", "default")
                    h_first = etree.SubElement(sect_pr, f"{{{W}}}headerReference")
                    h_first.set(f"{{{R}}}id", rid_first)
                    h_first.set(f"{{{W}}}type", "first")
                data = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)

            zout.writestr(item, data)
        zout.writestr("word/header1.xml", default_hdr)
        zout.writestr("word/header2.xml", first_hdr)

    out_buf.seek(0)
    save_fixture("section-constraints", "title-page-false", out_buf, {
        "name": "title-page-false",
        "spec_ref": "ISO 29500-1 §17.10.6",
        "description": "Section with default + first-page headers and titlePg val='false'. First-page header should be filtered out.",
        "expected_behavior": "First-kind header refs and stories should be filtered from the model since titlePg is explicitly false.",
    })


# =========================================================================

def main() -> None:
    print("Generating section constraint fixtures:")
    make_header_inheritance()
    make_title_page_interaction()
    make_continuous_section()
    make_partial_header_inheritance()
    make_title_page_absent()
    make_title_page_false()
    print("Done.")


if __name__ == "__main__":
    main()

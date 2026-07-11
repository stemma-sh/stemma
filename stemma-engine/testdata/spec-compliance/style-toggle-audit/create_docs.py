# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for style toggle property & cascade audit.

Exercises toggle property XOR semantics (ISO 29500-1 §17.7.3), linked style
property resolution (§17.7.4.6), implicit Normal style application (§17.7.4.17),
and doc defaults interactions.

Run:  python create_docs.py
"""

import json
from pathlib import Path
from lxml import etree

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  style-toggle-audit/{name}/")


# =========================================================================
# 1. Para style + char style toggle XOR (§17.7.3 + MS-OI29500 §2.1.258)
# =========================================================================

def make_toggle_xor_para_char():
    """
    §17.7.3: When para style sets bold=on and char style sets bold=on,
    the ECMA spec says XOR → off.
    MS-OI29500 §2.1.258: Word RESETS (override), so bold=on.

    Setup:
    - ParaBold: paragraph style with bold=on
    - CharBold: character style with bold=on
    - Para 1: has ParaBold, run has CharBold → test XOR interaction
    - Para 2: has ParaBold, run has NO char style → bold from para only
    - Para 3: no para style, run has CharBold → bold from char only
    """
    doc = Document()
    styles_el = doc.styles.element

    # Paragraph style: ParaBold — bold=on
    style_para = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "ParaBold",
    })
    style_para.append(make_element("w:name", {"w:val": "Para Bold"}))
    rpr = make_element("w:rPr")
    rpr.append(make_element("w:b"))
    style_para.append(rpr)
    styles_el.append(style_para)

    # Character style: CharBold — bold=on
    style_char = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "CharBold",
    })
    style_char.append(make_element("w:name", {"w:val": "Char Bold"}))
    rpr2 = make_element("w:rPr")
    rpr2.append(make_element("w:b"))
    style_char.append(rpr2)
    styles_el.append(style_char)

    # Para 1: ParaBold + CharBold
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "ParaBold"}))
    r1 = p1.add_run("Both para and char bold")
    r1_rpr = r1._element.get_or_add_rPr()
    r1_rpr.insert(0, make_element("w:rStyle", {"w:val": "CharBold"}))

    # Para 2: ParaBold only, no char style
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    p2_ppr.append(make_element("w:pStyle", {"w:val": "ParaBold"}))
    r2 = p2.add_run("Para bold only")
    # Ensure no rStyle
    r2_rpr = r2._element.find(qn("w:rPr"))
    if r2_rpr is not None:
        rs = r2_rpr.find(qn("w:rStyle"))
        if rs is not None:
            r2_rpr.remove(rs)

    # Para 3: no pStyle, CharBold only
    p3 = doc.add_paragraph()
    r3 = p3.add_run("Char bold only")
    r3_rpr = r3._element.get_or_add_rPr()
    r3_rpr.insert(0, make_element("w:rStyle", {"w:val": "CharBold"}))

    save_fixture("toggle-xor-para-char", doc, {
        "spec": "ISO 29500-1 §17.7.3 + MS-OI29500 §2.1.258",
        "description": "Para style bold + char style bold: XOR vs reset",
        "expected": {
            "para_1": "bold=ON per Word (reset), bold=OFF per ECMA (XOR on^on=off)",
            "para_2": "bold=ON (para style only, no XOR needed)",
            "para_3": "bold=ON (char style only, no XOR needed)",
        },
    })


# =========================================================================
# 2. Doc defaults bold=on + toggle shortcut (§17.7.3 + MS-OI29500 §2.1.230)
# =========================================================================

def make_toggle_doc_defaults_on():
    """
    §17.7.3: If doc defaults sets a toggle to true, the effective value is
    always true. But MS-OI29500 §2.1.230a says Word's behavior is different.

    Setup:
    - docDefaults rPr: bold=on
    - StyleNoBold: paragraph style with bold=off (explicitly off)
    - Para 1: uses StyleNoBold → test whether bold=off is honored
    - Para 2: unstyled → should be bold from doc defaults
    """
    doc = Document()
    styles_el = doc.styles.element

    # Set docDefaults rPr to bold=on
    doc_defaults = styles_el.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = make_element("w:docDefaults")
        styles_el.insert(0, doc_defaults)
    rpr_default_el = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default_el is None:
        rpr_default_el = make_element("w:rPrDefault")
        doc_defaults.insert(0, rpr_default_el)
    rpr = rpr_default_el.find(qn("w:rPr"))
    if rpr is None:
        rpr = make_element("w:rPr")
        rpr_default_el.append(rpr)
    # Remove any existing bold
    for existing_b in rpr.findall(qn("w:b")):
        rpr.remove(existing_b)
    rpr.append(make_element("w:b"))

    # StyleNoBold: paragraph style with explicit bold=off
    style_nobold = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "StyleNoBold",
    })
    style_nobold.append(make_element("w:name", {"w:val": "Style No Bold"}))
    rpr2 = make_element("w:rPr")
    rpr2.append(make_element("w:b", {"w:val": "0"}))
    style_nobold.append(rpr2)
    styles_el.append(style_nobold)

    # Para 1: StyleNoBold — explicit bold=off should turn off doc defaults bold
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "StyleNoBold"}))
    p1.add_run("Explicitly not bold")

    # Para 2: no pStyle — doc defaults bold=on applies
    p2 = doc.add_paragraph()
    p2.add_run("Doc defaults bold")

    save_fixture("toggle-doc-defaults-on", doc, {
        "spec": "ISO 29500-1 §17.7.3 + MS-OI29500 §2.1.230",
        "description": "docDefaults bold=on: para style bold=off should turn it off",
        "expected": {
            "para_1": "bold=OFF (para style explicitly turns off bold; MS behavior)",
            "para_2": "bold=ON (doc defaults bold applies to unstyled paragraph)",
        },
    })


# =========================================================================
# 3. Linked style missing properties (§17.7.4.6)
# =========================================================================

def make_linked_style_missing_props():
    """
    §17.7.4.6: When a paragraph style has w:link to a character style, runs
    inherit the linked character style's rPr. But what about properties the
    linked char style doesn't set?

    Setup:
    - HeadingPara: paragraph style, bold=on, font_size=32, basedOn Normal
    - HeadingChar: character style, bold=on (NO font_size), linked from HeadingPara
    - Para 1: uses HeadingPara → run should get bold=on AND font_size=32
      But if linked style replaces para chain entirely, font_size is lost.
    """
    doc = Document()
    styles_el = doc.styles.element

    # HeadingChar: character style with bold=on only (no font size)
    style_char = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "HeadingChar",
    })
    style_char.append(make_element("w:name", {"w:val": "Heading Char"}))
    link_to_para = make_element("w:link", {"w:val": "HeadingPara"})
    style_char.append(link_to_para)
    char_rpr = make_element("w:rPr")
    char_rpr.append(make_element("w:b"))
    style_char.append(char_rpr)
    styles_el.append(style_char)

    # HeadingPara: paragraph style with bold=on + font_size=32, linked to HeadingChar
    style_para = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "HeadingPara",
    })
    style_para.append(make_element("w:name", {"w:val": "Heading Para"}))
    link_to_char = make_element("w:link", {"w:val": "HeadingChar"})
    style_para.append(link_to_char)
    para_rpr = make_element("w:rPr")
    para_rpr.append(make_element("w:b"))
    para_rpr.append(make_element("w:sz", {"w:val": "32"}))
    style_para.append(para_rpr)
    styles_el.append(style_para)

    # Para 1: HeadingPara style
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "HeadingPara"}))
    p1.add_run("Heading with linked style")

    # Para 2: no pStyle, run with HeadingChar applied as character style
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Run with HeadingChar")
    r2_rpr = r2._element.get_or_add_rPr()
    r2_rpr.insert(0, make_element("w:rStyle", {"w:val": "HeadingChar"}))

    save_fixture("linked-style-missing-props", doc, {
        "spec": "ISO 29500-1 §17.7.4.6",
        "description": "Linked style should use char style rPr but not lose para style properties",
        "expected": {
            "para_1": "bold=ON, font_size=32 (from para style; linked char has no sz)",
            "para_2": "bold=ON (from HeadingChar), font_size from doc defaults/Normal",
        },
    })


# =========================================================================
# 4. 3-level basedOn toggle chain (§17.7.3 first-value-encountered)
# =========================================================================

def make_toggle_3level_based_on():
    """
    §17.7.3: Within a single hierarchy level, the first value encountered in
    the basedOn chain is used. This tests a 3-level chain where each level
    sets bold.

    Setup:
    - StyleRoot: character style, bold=on
    - StyleMid: character style, basedOn StyleRoot, bold=off
    - StyleLeaf: character style, basedOn StyleMid, bold=on

    The "first value encountered" from StyleLeaf is bold=on (from StyleLeaf itself).
    Result for the character level should be bold=on.
    """
    doc = Document()
    styles_el = doc.styles.element

    # StyleRoot: bold=on
    style_root = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "StyleRoot",
    })
    style_root.append(make_element("w:name", {"w:val": "Style Root"}))
    root_rpr = make_element("w:rPr")
    root_rpr.append(make_element("w:b"))
    style_root.append(root_rpr)
    styles_el.append(style_root)

    # StyleMid: basedOn StyleRoot, bold=off
    style_mid = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "StyleMid",
    })
    style_mid.append(make_element("w:name", {"w:val": "Style Mid"}))
    style_mid.append(make_element("w:basedOn", {"w:val": "StyleRoot"}))
    mid_rpr = make_element("w:rPr")
    mid_rpr.append(make_element("w:b", {"w:val": "0"}))
    style_mid.append(mid_rpr)
    styles_el.append(style_mid)

    # StyleLeaf: basedOn StyleMid, bold=on
    style_leaf = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "StyleLeaf",
    })
    style_leaf.append(make_element("w:name", {"w:val": "Style Leaf"}))
    style_leaf.append(make_element("w:basedOn", {"w:val": "StyleMid"}))
    leaf_rpr = make_element("w:rPr")
    leaf_rpr.append(make_element("w:b"))
    style_leaf.append(leaf_rpr)
    styles_el.append(style_leaf)

    # Also test italic: only set at root level, not overridden
    # StyleRoot already has bold, let's add italic=on
    root_rpr.append(make_element("w:i"))
    # StyleMid: bold=off, italic not set (should inherit from root)
    # StyleLeaf: bold=on, italic not set (should inherit from root)

    # Para 1: run with StyleLeaf → bold=on (leaf overrides), italic=on (from root)
    p1 = doc.add_paragraph()
    r1 = p1.add_run("Leaf style run")
    r1_rpr = r1._element.get_or_add_rPr()
    r1_rpr.insert(0, make_element("w:rStyle", {"w:val": "StyleLeaf"}))

    # Para 2: run with StyleMid → bold=off, italic=on (from root)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Mid style run")
    r2_rpr = r2._element.get_or_add_rPr()
    r2_rpr.insert(0, make_element("w:rStyle", {"w:val": "StyleMid"}))

    # Para 3: run with StyleRoot → bold=on, italic=on
    p3 = doc.add_paragraph()
    r3 = p3.add_run("Root style run")
    r3_rpr = r3._element.get_or_add_rPr()
    r3_rpr.insert(0, make_element("w:rStyle", {"w:val": "StyleRoot"}))

    save_fixture("toggle-3level-basedOn", doc, {
        "spec": "ISO 29500-1 §17.7.3 (first-value-encountered in basedOn chain)",
        "description": "3-level char style basedOn chain with toggle overrides",
        "expected": {
            "para_1_leaf": "bold=ON (leaf sets it), italic=ON (inherited from root)",
            "para_2_mid": "bold=OFF (mid explicitly off), italic=ON (inherited from root)",
            "para_3_root": "bold=ON (root sets it), italic=ON (root sets it)",
        },
    })


# =========================================================================
# 5. Direct formatting overrides toggle (§17.7.3 direct=non-XOR)
# =========================================================================

def make_toggle_direct_override():
    """
    §17.7.3: Direct formatting is NOT XOR'd with style toggles.
    If direct says bold=off, it wins regardless.

    Setup:
    - BoldPara: paragraph style, bold=on
    - BoldChar: character style, bold=on
    - Para 1: BoldPara + BoldChar + direct bold=off → bold=off
    - Para 2: BoldPara + BoldChar + direct bold=on → bold=on
    - Para 3: BoldPara + direct bold=off (no char style) → bold=off
    """
    doc = Document()
    styles_el = doc.styles.element

    # BoldPara: paragraph style with bold=on
    style_para = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "BoldPara",
    })
    style_para.append(make_element("w:name", {"w:val": "Bold Para"}))
    rpr1 = make_element("w:rPr")
    rpr1.append(make_element("w:b"))
    style_para.append(rpr1)
    styles_el.append(style_para)

    # BoldChar: character style with bold=on
    style_char = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "BoldChar",
    })
    style_char.append(make_element("w:name", {"w:val": "Bold Char"}))
    rpr2 = make_element("w:rPr")
    rpr2.append(make_element("w:b"))
    style_char.append(rpr2)
    styles_el.append(style_char)

    # Para 1: BoldPara + BoldChar + direct bold=off
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "BoldPara"}))
    r1 = p1.add_run("Direct bold off")
    r1_rpr = r1._element.get_or_add_rPr()
    r1_rpr.insert(0, make_element("w:rStyle", {"w:val": "BoldChar"}))
    r1_rpr.append(make_element("w:b", {"w:val": "0"}))

    # Para 2: BoldPara + BoldChar + direct bold=on
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    p2_ppr.append(make_element("w:pStyle", {"w:val": "BoldPara"}))
    r2 = p2.add_run("Direct bold on")
    r2_rpr = r2._element.get_or_add_rPr()
    r2_rpr.insert(0, make_element("w:rStyle", {"w:val": "BoldChar"}))
    r2_rpr.append(make_element("w:b"))

    # Para 3: BoldPara + direct bold=off (no char style)
    p3 = doc.add_paragraph()
    p3_ppr = p3._element.get_or_add_pPr()
    p3_ppr.append(make_element("w:pStyle", {"w:val": "BoldPara"}))
    r3 = p3.add_run("Direct off no char")
    r3_rpr = r3._element.get_or_add_rPr()
    r3_rpr.append(make_element("w:b", {"w:val": "0"}))

    save_fixture("toggle-direct-override", doc, {
        "spec": "ISO 29500-1 §17.7.3 (direct formatting always wins)",
        "description": "Direct bold overrides style toggles without XOR",
        "expected": {
            "para_1": "bold=OFF (direct off wins over all styles)",
            "para_2": "bold=ON (direct on wins)",
            "para_3": "bold=OFF (direct off wins over para style)",
        },
    })


# =========================================================================
# 6. Non-toggle per-property override in deep chain (§17.7.1)
# =========================================================================

def make_non_toggle_deep_chain():
    """
    §17.7.1: Non-toggle properties use simple last-wins (child overrides parent).
    Properties NOT set by child inherit from parent.

    Setup:
    - BaseStyle: character style, font_size=24, color=FF0000, italic=on
    - ChildStyle: basedOn BaseStyle, font_size=28 (overrides), color not set (inherits)
    - GrandchildStyle: basedOn ChildStyle, color=0000FF (overrides), font_size not set (inherits 28)

    Para 1: run with GrandchildStyle → font_size=28 (from child), color=0000FF (from grandchild), italic=on (from base)
    """
    doc = Document()
    styles_el = doc.styles.element

    # BaseStyle
    style_base = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "BaseStyle",
    })
    style_base.append(make_element("w:name", {"w:val": "Base Style"}))
    base_rpr = make_element("w:rPr")
    base_rpr.append(make_element("w:sz", {"w:val": "24"}))
    base_rpr.append(make_element("w:color", {"w:val": "FF0000"}))
    base_rpr.append(make_element("w:i"))
    style_base.append(base_rpr)
    styles_el.append(style_base)

    # ChildStyle: basedOn BaseStyle, overrides font_size
    style_child = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "ChildStyle",
    })
    style_child.append(make_element("w:name", {"w:val": "Child Style"}))
    style_child.append(make_element("w:basedOn", {"w:val": "BaseStyle"}))
    child_rpr = make_element("w:rPr")
    child_rpr.append(make_element("w:sz", {"w:val": "28"}))
    style_child.append(child_rpr)
    styles_el.append(style_child)

    # GrandchildStyle: basedOn ChildStyle, overrides color
    style_gc = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "GrandchildStyle",
    })
    style_gc.append(make_element("w:name", {"w:val": "Grandchild Style"}))
    style_gc.append(make_element("w:basedOn", {"w:val": "ChildStyle"}))
    gc_rpr = make_element("w:rPr")
    gc_rpr.append(make_element("w:color", {"w:val": "0000FF"}))
    style_gc.append(gc_rpr)
    styles_el.append(style_gc)

    # Para 1: GrandchildStyle
    p1 = doc.add_paragraph()
    r1 = p1.add_run("Grandchild style run")
    r1_rpr = r1._element.get_or_add_rPr()
    r1_rpr.insert(0, make_element("w:rStyle", {"w:val": "GrandchildStyle"}))

    # Para 2: ChildStyle
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Child style run")
    r2_rpr = r2._element.get_or_add_rPr()
    r2_rpr.insert(0, make_element("w:rStyle", {"w:val": "ChildStyle"}))

    # Para 3: BaseStyle
    p3 = doc.add_paragraph()
    r3 = p3.add_run("Base style run")
    r3_rpr = r3._element.get_or_add_rPr()
    r3_rpr.insert(0, make_element("w:rStyle", {"w:val": "BaseStyle"}))

    save_fixture("non-toggle-deep-chain", doc, {
        "spec": "ISO 29500-1 §17.7.1 (per-property override)",
        "description": "Non-toggle properties: child overrides per-property, inherits rest",
        "expected": {
            "para_1_grandchild": {
                "font_size": 28, "color": "0000FF", "italic": True,
            },
            "para_2_child": {
                "font_size": 28, "color": "FF0000", "italic": True,
            },
            "para_3_base": {
                "font_size": 24, "color": "FF0000", "italic": True,
            },
        },
    })


# =========================================================================
# 7. Implicit Normal style (§17.7.4.17)
# =========================================================================

def make_implicit_normal_style():
    """
    §17.7.4.17: When a paragraph has no explicit pStyle, the default paragraph
    style (typically "Normal") is implied. The Normal style's rPr should apply
    to unstyled runs in that paragraph.

    Setup:
    - Modify Normal style to have color=008000 (green), font_size=28
    - Para 1: no pStyle, no rStyle → should get Normal's color and font_size
    - Para 2: explicit pStyle=Normal → same result
    - Para 3: different pStyle (CustomPara) → should NOT get Normal's properties
    """
    doc = Document()
    styles_el = doc.styles.element

    # Find the existing Normal style and modify it
    normal_style = None
    for style in styles_el.findall(qn("w:style")):
        if style.get(qn("w:styleId")) == "Normal":
            normal_style = style
            break

    if normal_style is not None:
        # Get or add rPr
        rpr = normal_style.find(qn("w:rPr"))
        if rpr is None:
            rpr = make_element("w:rPr")
            normal_style.append(rpr)
        # Set color and font_size
        for existing in rpr.findall(qn("w:color")):
            rpr.remove(existing)
        rpr.append(make_element("w:color", {"w:val": "008000"}))
        for existing in rpr.findall(qn("w:sz")):
            rpr.remove(existing)
        rpr.append(make_element("w:sz", {"w:val": "28"}))

    # CustomPara: a different paragraph style with color=0000FF
    style_custom = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "CustomPara",
    })
    style_custom.append(make_element("w:name", {"w:val": "Custom Para"}))
    custom_rpr = make_element("w:rPr")
    custom_rpr.append(make_element("w:color", {"w:val": "0000FF"}))
    style_custom.append(custom_rpr)
    styles_el.append(style_custom)

    # Para 1: no pStyle → implicit Normal
    p1 = doc.add_paragraph()
    # Remove any pStyle that python-docx may add
    p1_ppr = p1._element.find(qn("w:pPr"))
    if p1_ppr is not None:
        ps = p1_ppr.find(qn("w:pStyle"))
        if ps is not None:
            p1_ppr.remove(ps)
    p1.add_run("Implicit Normal style")

    # Para 2: explicit pStyle=Normal
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    # Remove any auto pStyle first
    ps = p2_ppr.find(qn("w:pStyle"))
    if ps is not None:
        p2_ppr.remove(ps)
    p2_ppr.insert(0, make_element("w:pStyle", {"w:val": "Normal"}))
    p2.add_run("Explicit Normal style")

    # Para 3: CustomPara style
    p3 = doc.add_paragraph()
    p3_ppr = p3._element.get_or_add_pPr()
    ps = p3_ppr.find(qn("w:pStyle"))
    if ps is not None:
        p3_ppr.remove(ps)
    p3_ppr.insert(0, make_element("w:pStyle", {"w:val": "CustomPara"}))
    p3.add_run("Custom para style")

    save_fixture("implicit-normal-style", doc, {
        "spec": "ISO 29500-1 §17.7.4.17",
        "description": "Unstyled paragraphs should apply the default paragraph style (Normal)",
        "expected": {
            "para_1_implicit": {"color": "008000", "font_size": 28},
            "para_2_explicit": {"color": "008000", "font_size": 28},
            "para_3_custom": {"color": "0000FF", "font_size": None},
        },
    })


# =========================================================================
# 8. Char style overriding para style rPr (§17.7.2)
# =========================================================================

def make_char_style_overrides_para_rpr():
    """
    §17.7.2: The cascade is docDefaults → table → para → char → direct.
    Character style rPr overrides paragraph style rPr for runs that have both.

    Setup:
    - ColorPara: paragraph style with color=FF0000 (red), font_size=24
    - SizesChar: character style with font_size=32 (overrides para's 24)
    - Para 1: ColorPara + SizesChar → color=FF0000 (para, not overridden),
      font_size=32 (char overrides para)
    - Para 2: ColorPara, no char style → color=FF0000, font_size=24
    """
    doc = Document()
    styles_el = doc.styles.element

    # ColorPara: paragraph style
    style_para = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "ColorPara",
    })
    style_para.append(make_element("w:name", {"w:val": "Color Para"}))
    para_rpr = make_element("w:rPr")
    para_rpr.append(make_element("w:color", {"w:val": "FF0000"}))
    para_rpr.append(make_element("w:sz", {"w:val": "24"}))
    style_para.append(para_rpr)
    styles_el.append(style_para)

    # SizesChar: character style
    style_char = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "SizesChar",
    })
    style_char.append(make_element("w:name", {"w:val": "Sizes Char"}))
    char_rpr = make_element("w:rPr")
    char_rpr.append(make_element("w:sz", {"w:val": "32"}))
    style_char.append(char_rpr)
    styles_el.append(style_char)

    # Para 1: ColorPara + SizesChar
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "ColorPara"}))
    r1 = p1.add_run("Char overrides para size")
    r1_rpr = r1._element.get_or_add_rPr()
    r1_rpr.insert(0, make_element("w:rStyle", {"w:val": "SizesChar"}))

    # Para 2: ColorPara only
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    p2_ppr.append(make_element("w:pStyle", {"w:val": "ColorPara"}))
    p2.add_run("Para style only")

    save_fixture("char-overrides-para-rpr", doc, {
        "spec": "ISO 29500-1 §17.7.2",
        "description": "Character style overrides paragraph style rPr per-property",
        "expected": {
            "para_1": {"color": "FF0000", "font_size": 32},
            "para_2": {"color": "FF0000", "font_size": 24},
        },
    })


# =========================================================================
# 9. docDefaults rPr + style override interaction (§17.7.5)
# =========================================================================

def make_doc_defaults_rpr_cascade():
    """
    §17.7.5: docDefaults rPrDefault provides base formatting for ALL runs.
    Style properties override docDefaults per-property.

    Setup:
    - docDefaults: font_size=20, color=808080 (gray)
    - OverrideStyle: paragraph style with font_size=28 (overrides default 20),
      color not set (should inherit gray from docDefaults)
    - Para 1: OverrideStyle → font_size=28, color=808080
    - Para 2: no style → font_size=20, color=808080
    """
    doc = Document()
    styles_el = doc.styles.element

    # Set docDefaults rPr
    doc_defaults = styles_el.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = make_element("w:docDefaults")
        styles_el.insert(0, doc_defaults)
    rpr_default_el = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default_el is None:
        rpr_default_el = make_element("w:rPrDefault")
        doc_defaults.insert(0, rpr_default_el)
    rpr = rpr_default_el.find(qn("w:rPr"))
    if rpr is None:
        rpr = make_element("w:rPr")
        rpr_default_el.append(rpr)
    # Set font_size and color in docDefaults
    for existing in rpr.findall(qn("w:sz")):
        rpr.remove(existing)
    rpr.append(make_element("w:sz", {"w:val": "20"}))
    for existing in rpr.findall(qn("w:color")):
        rpr.remove(existing)
    rpr.append(make_element("w:color", {"w:val": "808080"}))

    # OverrideStyle: font_size=28, no color
    style_override = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "OverrideStyle",
    })
    style_override.append(make_element("w:name", {"w:val": "Override Style"}))
    override_rpr = make_element("w:rPr")
    override_rpr.append(make_element("w:sz", {"w:val": "28"}))
    style_override.append(override_rpr)
    styles_el.append(style_override)

    # Para 1: OverrideStyle
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "OverrideStyle"}))
    p1.add_run("Override style run")

    # Para 2: no style
    p2 = doc.add_paragraph()
    # Remove any auto-added pStyle
    p2_ppr = p2._element.find(qn("w:pPr"))
    if p2_ppr is not None:
        ps = p2_ppr.find(qn("w:pStyle"))
        if ps is not None:
            p2_ppr.remove(ps)
    p2.add_run("Unstyled run")

    save_fixture("doc-defaults-rpr-cascade", doc, {
        "spec": "ISO 29500-1 §17.7.5",
        "description": "docDefaults rPr provides base; style overrides per-property",
        "expected": {
            "para_1": {"font_size": 28, "color": "808080"},
            "para_2": {"font_size": 20, "color": "808080"},
        },
    })


# =========================================================================
# 10. Multiple toggle properties mixed (§17.7.3)
# =========================================================================

def make_toggle_multi_property():
    """
    Test that different toggle properties can behave independently.

    Setup:
    - ParaToggle: paragraph style with bold=on, italic=on
    - CharToggle: character style with bold=on (XOR with para → off), italic not set
    - Para 1: ParaToggle + CharToggle → bold should be XOR'd (off per spec),
      italic from para only (on)
    """
    doc = Document()
    styles_el = doc.styles.element

    # ParaToggle: bold=on, italic=on
    style_para = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "ParaToggle",
    })
    style_para.append(make_element("w:name", {"w:val": "Para Toggle"}))
    rpr1 = make_element("w:rPr")
    rpr1.append(make_element("w:b"))
    rpr1.append(make_element("w:i"))
    style_para.append(rpr1)
    styles_el.append(style_para)

    # CharToggle: bold=on only
    style_char = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "CharToggle",
    })
    style_char.append(make_element("w:name", {"w:val": "Char Toggle"}))
    rpr2 = make_element("w:rPr")
    rpr2.append(make_element("w:b"))
    style_char.append(rpr2)
    styles_el.append(style_char)

    # Para 1: ParaToggle + CharToggle
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "ParaToggle"}))
    r1 = p1.add_run("Multi toggle test")
    r1_rpr = r1._element.get_or_add_rPr()
    r1_rpr.insert(0, make_element("w:rStyle", {"w:val": "CharToggle"}))

    save_fixture("toggle-multi-property", doc, {
        "spec": "ISO 29500-1 §17.7.3",
        "description": "Different toggle properties behave independently across levels",
        "expected": {
            "para_1": {
                "bold": "OFF per ECMA (XOR on^on=off), ON per Word (reset)",
                "italic": "ON (only para style sets it, no XOR needed)",
            },
        },
    })


# =========================================================================
# Main
# =========================================================================

if __name__ == "__main__":
    print("Generating style-toggle-audit fixtures:")
    make_toggle_xor_para_char()
    make_toggle_doc_defaults_on()
    make_linked_style_missing_props()
    make_toggle_3level_based_on()
    make_toggle_direct_override()
    make_non_toggle_deep_chain()
    make_implicit_normal_style()
    make_char_style_overrides_para_rpr()
    make_doc_defaults_rpr_cascade()
    make_toggle_multi_property()
    print("Done.")

# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for spec-compliance testing of stories
(headers, footers, footnotes, endnotes) against ISO 29500-1 §17.10, §17.11.

Run:  uv run create_stories.py
"""

import json
import zipfile
import io
from pathlib import Path
from lxml import etree

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

HEADER_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml"
FOOTER_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml"
FOOTNOTES_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
ENDNOTES_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml"

HEADER_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/header"
FOOTER_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer"
FOOTNOTES_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
ENDNOTES_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def _build_header_xml(paragraphs: list[str]) -> bytes:
    root = etree.Element(f"{{{W}}}hdr", nsmap={"w": W, "r": R})
    for text in paragraphs:
        p = etree.SubElement(root, f"{{{W}}}p")
        r = etree.SubElement(p, f"{{{W}}}r")
        t = etree.SubElement(r, f"{{{W}}}t")
        t.text = text
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _build_header_xml_with_formatting(paragraphs: list[dict]) -> bytes:
    root = etree.Element(f"{{{W}}}hdr", nsmap={"w": W, "r": R})
    for para in paragraphs:
        p = etree.SubElement(root, f"{{{W}}}p")
        r = etree.SubElement(p, f"{{{W}}}r")
        rPr = etree.SubElement(r, f"{{{W}}}rPr")
        if para.get("bold"):
            etree.SubElement(rPr, f"{{{W}}}b")
        if para.get("italic"):
            etree.SubElement(rPr, f"{{{W}}}i")
        t = etree.SubElement(r, f"{{{W}}}t")
        t.text = para["text"]
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _build_footer_xml(paragraphs: list[str]) -> bytes:
    root = etree.Element(f"{{{W}}}ftr", nsmap={"w": W, "r": R})
    for text in paragraphs:
        p = etree.SubElement(root, f"{{{W}}}p")
        r = etree.SubElement(p, f"{{{W}}}r")
        t = etree.SubElement(r, f"{{{W}}}t")
        t.text = text
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _build_footnotes_xml(footnotes: list[dict]) -> bytes:
    root = etree.Element(f"{{{W}}}footnotes", nsmap={"w": W, "r": R})
    for fn in footnotes:
        attrs = {f"{{{W}}}id": fn["id"]}
        if fn.get("type"):
            attrs[f"{{{W}}}type"] = fn["type"]
        fn_el = etree.SubElement(root, f"{{{W}}}footnote", attrib=attrs)
        for text in fn.get("paragraphs", []):
            p = etree.SubElement(fn_el, f"{{{W}}}p")
            r = etree.SubElement(p, f"{{{W}}}r")
            t = etree.SubElement(r, f"{{{W}}}t")
            t.text = text
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _build_endnotes_xml(endnotes: list[dict]) -> bytes:
    root = etree.Element(f"{{{W}}}endnotes", nsmap={"w": W, "r": R})
    for en in endnotes:
        attrs = {f"{{{W}}}id": en["id"]}
        if en.get("type"):
            attrs[f"{{{W}}}type"] = en["type"]
        en_el = etree.SubElement(root, f"{{{W}}}endnote", attrib=attrs)
        for text in en.get("paragraphs", []):
            p = etree.SubElement(en_el, f"{{{W}}}p")
            r = etree.SubElement(p, f"{{{W}}}r")
            t = etree.SubElement(r, f"{{{W}}}t")
            t.text = text
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _add_story_parts(doc: DocxDocument, parts: dict, sect_pr_additions: list[tuple] | None = None) -> io.BytesIO:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    rid_map = {}
    out_buf = io.BytesIO()
    rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"

    with zipfile.ZipFile(buf, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        rels_data = zin.read("word/_rels/document.xml.rels")
        rels_root = etree.fromstring(rels_data)
        existing_ids = [el.get("Id") for el in rels_root]
        max_id = 0
        for rid in existing_ids:
            if rid and rid.startswith("rId"):
                try:
                    max_id = max(max_id, int(rid[3:]))
                except ValueError:
                    pass

        for part_path in parts:
            new_rid = f"rId{max_id + 1}"
            max_id += 1
            rid_map[part_path] = new_rid
            rel_el = etree.SubElement(rels_root, f"{{{rels_ns}}}Relationship")
            rel_el.set("Id", new_rid)
            rel_el.set("Type", parts[part_path][2])
            rel_el.set("Target", part_path)

        new_rels_data = etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)

        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename == "[Content_Types].xml":
                ct_root = etree.fromstring(data)
                for part_path, (_, content_type, _) in parts.items():
                    pn = f"/word/{part_path}"
                    existing = ct_root.findall(f"{{{ct_ns}}}Override[@PartName='{pn}']")
                    if not existing:
                        ov = etree.SubElement(ct_root, f"{{{ct_ns}}}Override")
                        ov.set("PartName", pn)
                        ov.set("ContentType", content_type)
                data = etree.tostring(ct_root, xml_declaration=True, encoding="UTF-8", standalone=True)

            elif item.filename == "word/_rels/document.xml.rels":
                data = new_rels_data

            elif item.filename == "word/document.xml" and sect_pr_additions:
                doc_root = etree.fromstring(data)
                has_first = any(t[2] == "first" for t in sect_pr_additions)
                for sect_pr in doc_root.iter(f"{{{W}}}sectPr"):
                    # Add titlePg when first-page headers/footers are present
                    if has_first:
                        etree.SubElement(sect_pr, f"{{{W}}}titlePg")
                    for tag, part_path, type_attr in sect_pr_additions:
                        rid = rid_map.get(part_path)
                        if rid:
                            ref_el = etree.SubElement(sect_pr, f"{{{W}}}{tag}")
                            ref_el.set(f"{{{R}}}id", rid)
                            ref_el.set(f"{{{W}}}type", type_attr)
                data = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)

            elif item.filename == "word/settings.xml" and sect_pr_additions:
                has_even = any(t[2] == "even" for t in sect_pr_additions)
                if has_even:
                    settings_root = etree.fromstring(data)
                    etree.SubElement(settings_root, f"{{{W}}}evenAndOddHeaders")
                    data = etree.tostring(settings_root, xml_declaration=True, encoding="UTF-8", standalone=True)

            zout.writestr(item, data)

        for part_path, (content, _, _) in parts.items():
            zout.writestr(f"word/{part_path}", content)

    out_buf.seek(0)
    return out_buf


def save_story_fixture(name: str, out_buf: io.BytesIO, metadata: dict) -> None:
    out_dir = ROOT / "stories" / name
    out_dir.mkdir(parents=True, exist_ok=True)
    (out_dir / "input.docx").write_bytes(out_buf.getvalue())
    (out_dir / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  stories/{name}/")


def make_header_types() -> None:
    doc = Document()
    doc.add_paragraph("Body content for header types test.")
    parts = {
        "header1.xml": (_build_header_xml(["Default Header"]), HEADER_CONTENT_TYPE, HEADER_REL_TYPE),
        "header2.xml": (_build_header_xml(["First Page Header"]), HEADER_CONTENT_TYPE, HEADER_REL_TYPE),
        "header3.xml": (_build_header_xml(["Even Page Header"]), HEADER_CONTENT_TYPE, HEADER_REL_TYPE),
    }
    sect_pr_additions = [
        ("headerReference", "header1.xml", "default"),
        ("headerReference", "header2.xml", "first"),
        ("headerReference", "header3.xml", "even"),
    ]
    out_buf = _add_story_parts(doc, parts, sect_pr_additions)
    save_story_fixture("header-types", out_buf, {
        "name": "header-types",
        "spec_ref": "ISO 29500-1 §17.10.5",
        "description": "Document with three header types: default, first page, even page",
    })


def make_footer_types() -> None:
    doc = Document()
    doc.add_paragraph("Body content for footer types test.")
    parts = {
        "footer1.xml": (_build_footer_xml(["Default Footer"]), FOOTER_CONTENT_TYPE, FOOTER_REL_TYPE),
        "footer2.xml": (_build_footer_xml(["First Page Footer"]), FOOTER_CONTENT_TYPE, FOOTER_REL_TYPE),
        "footer3.xml": (_build_footer_xml(["Even Page Footer"]), FOOTER_CONTENT_TYPE, FOOTER_REL_TYPE),
    }
    sect_pr_additions = [
        ("footerReference", "footer1.xml", "default"),
        ("footerReference", "footer2.xml", "first"),
        ("footerReference", "footer3.xml", "even"),
    ]
    out_buf = _add_story_parts(doc, parts, sect_pr_additions)
    save_story_fixture("footer-types", out_buf, {
        "name": "footer-types",
        "spec_ref": "ISO 29500-1 §17.10.3",
        "description": "Document with three footer types: default, first page, even page",
    })


def make_footnote_references() -> None:
    doc = Document()
    p = doc.add_paragraph()
    for child in list(p._p):
        p._p.remove(child)

    r1 = make_element("w:r")
    t1 = make_element("w:t")
    t1.text = "See footnote"
    t1.set(qn("xml:space"), "preserve")
    r1.append(t1)
    p._p.append(r1)

    r2 = make_element("w:r")
    fn_ref = make_element("w:footnoteReference", {"w:id": "1"})
    r2.append(fn_ref)
    p._p.append(r2)

    footnotes_xml = _build_footnotes_xml([
        {"id": "0", "type": "separator", "paragraphs": [""]},
        {"id": "1", "type": None, "paragraphs": ["This is footnote one."]},
    ])
    parts = {"footnotes.xml": (footnotes_xml, FOOTNOTES_CONTENT_TYPE, FOOTNOTES_REL_TYPE)}
    out_buf = _add_story_parts(doc, parts)
    save_story_fixture("footnote-references", out_buf, {
        "name": "footnote-references",
        "spec_ref": "ISO 29500-1 §17.11.13",
        "description": "Body contains footnoteReference linking to FootnoteStory id=1",
    })


def make_endnote_references() -> None:
    doc = Document()
    p = doc.add_paragraph()
    for child in list(p._p):
        p._p.remove(child)

    r1 = make_element("w:r")
    t1 = make_element("w:t")
    t1.text = "See endnote"
    t1.set(qn("xml:space"), "preserve")
    r1.append(t1)
    p._p.append(r1)

    r2 = make_element("w:r")
    en_ref = make_element("w:endnoteReference", {"w:id": "1"})
    r2.append(en_ref)
    p._p.append(r2)

    endnotes_xml = _build_endnotes_xml([
        {"id": "0", "type": "separator", "paragraphs": [""]},
        {"id": "1", "type": None, "paragraphs": ["This is endnote one."]},
    ])
    parts = {"endnotes.xml": (endnotes_xml, ENDNOTES_CONTENT_TYPE, ENDNOTES_REL_TYPE)}
    out_buf = _add_story_parts(doc, parts)
    save_story_fixture("endnote-references", out_buf, {
        "name": "endnote-references",
        "spec_ref": "ISO 29500-1 §17.11.2",
        "description": "Body contains endnoteReference linking to EndnoteStory id=1",
    })


def make_footnote_types() -> None:
    doc = Document()
    doc.add_paragraph("Body with footnotes of different types.")
    footnotes_xml = _build_footnotes_xml([
        {"id": "0", "type": "separator", "paragraphs": [""]},
        {"id": "1", "type": "continuationSeparator", "paragraphs": [""]},
        {"id": "2", "type": None, "paragraphs": ["Normal footnote content."]},
    ])
    parts = {"footnotes.xml": (footnotes_xml, FOOTNOTES_CONTENT_TYPE, FOOTNOTES_REL_TYPE)}
    out_buf = _add_story_parts(doc, parts)
    save_story_fixture("footnote-types", out_buf, {
        "name": "footnote-types",
        "spec_ref": "ISO 29500-1 §17.11.20",
        "description": "Footnotes with type=separator, continuationSeparator, and normal",
    })


def make_multi_section_headers() -> None:
    doc = Document()
    doc.add_paragraph("Section 1 content.")

    hdr1 = _build_header_xml(["Header for Section 1"])
    hdr2 = _build_header_xml(["Header for Section 2"])
    hdr3 = _build_header_xml(["Header for Section 3"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    out_buf = io.BytesIO()
    rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"

    with zipfile.ZipFile(buf, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        # Pre-compute rIds from rels file
        rels_data = zin.read("word/_rels/document.xml.rels")
        rels_root = etree.fromstring(rels_data)
        existing_ids = [el.get("Id") for el in rels_root]
        max_id = 0
        for rid in existing_ids:
            if rid and rid.startswith("rId"):
                try:
                    max_id = max(max_id, int(rid[3:]))
                except ValueError:
                    pass
        rids = [f"rId{max_id + i}" for i in range(1, 4)]
        for i, new_rid in enumerate(rids, 1):
            rel_el = etree.SubElement(rels_root, f"{{{rels_ns}}}Relationship")
            rel_el.set("Id", new_rid)
            rel_el.set("Type", HEADER_REL_TYPE)
            rel_el.set("Target", f"header{i}.xml")
        new_rels_data = etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)

        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename == "[Content_Types].xml":
                ct_root = etree.fromstring(data)
                for i in range(1, 4):
                    pn = f"/word/header{i}.xml"
                    existing = ct_root.findall(f"{{{ct_ns}}}Override[@PartName='{pn}']")
                    if not existing:
                        ov = etree.SubElement(ct_root, f"{{{ct_ns}}}Override")
                        ov.set("PartName", pn)
                        ov.set("ContentType", HEADER_CONTENT_TYPE)
                data = etree.tostring(ct_root, xml_declaration=True, encoding="UTF-8", standalone=True)

            elif item.filename == "word/_rels/document.xml.rels":
                data = new_rels_data

            elif item.filename == "word/document.xml":
                doc_root = etree.fromstring(data)
                body = doc_root.find(f"{{{W}}}body")
                for p in body.findall(f"{{{W}}}p"):
                    body.remove(p)
                for sp in body.findall(f"{{{W}}}sectPr"):
                    body.remove(sp)

                # Section 1
                p1 = etree.SubElement(body, f"{{{W}}}p")
                r1 = etree.SubElement(p1, f"{{{W}}}r")
                t1 = etree.SubElement(r1, f"{{{W}}}t")
                t1.text = "Section 1 content."
                p1b = etree.SubElement(body, f"{{{W}}}p")
                pPr1 = etree.SubElement(p1b, f"{{{W}}}pPr")
                s1 = etree.SubElement(pPr1, f"{{{W}}}sectPr")
                h1 = etree.SubElement(s1, f"{{{W}}}headerReference")
                h1.set(f"{{{R}}}id", rids[0])
                h1.set(f"{{{W}}}type", "default")

                # Section 2
                p2 = etree.SubElement(body, f"{{{W}}}p")
                r2 = etree.SubElement(p2, f"{{{W}}}r")
                t2 = etree.SubElement(r2, f"{{{W}}}t")
                t2.text = "Section 2 content."
                p2b = etree.SubElement(body, f"{{{W}}}p")
                pPr2 = etree.SubElement(p2b, f"{{{W}}}pPr")
                s2 = etree.SubElement(pPr2, f"{{{W}}}sectPr")
                h2 = etree.SubElement(s2, f"{{{W}}}headerReference")
                h2.set(f"{{{R}}}id", rids[1])
                h2.set(f"{{{W}}}type", "default")

                # Section 3
                p3 = etree.SubElement(body, f"{{{W}}}p")
                r3 = etree.SubElement(p3, f"{{{W}}}r")
                t3 = etree.SubElement(r3, f"{{{W}}}t")
                t3.text = "Section 3 content."
                s3 = etree.SubElement(body, f"{{{W}}}sectPr")
                h3 = etree.SubElement(s3, f"{{{W}}}headerReference")
                h3.set(f"{{{R}}}id", rids[2])
                h3.set(f"{{{W}}}type", "default")

                data = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)

            zout.writestr(item, data)
        zout.writestr("word/header1.xml", hdr1)
        zout.writestr("word/header2.xml", hdr2)
        zout.writestr("word/header3.xml", hdr3)

    out_buf.seek(0)
    save_story_fixture("multi-section-headers", out_buf, {
        "name": "multi-section-headers",
        "spec_ref": "ISO 29500-1 §17.10.5",
        "description": "Three sections each with a distinct default header",
    })


def make_first_page_header() -> None:
    doc = Document()
    doc.add_paragraph("Body content with title page header.")

    default_hdr = _build_header_xml(["Default Header"])
    first_hdr = _build_header_xml(["Title Page Header"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    out_buf = io.BytesIO()
    rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"

    with zipfile.ZipFile(buf, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        # Pre-compute rIds from rels file
        rels_data = zin.read("word/_rels/document.xml.rels")
        rels_root = etree.fromstring(rels_data)
        existing_ids = [el.get("Id") for el in rels_root]
        max_id = 0
        for rid in existing_ids:
            if rid and rid.startswith("rId"):
                try:
                    max_id = max(max_id, int(rid[3:]))
                except ValueError:
                    pass
        rid_default = f"rId{max_id + 1}"
        rid_first = f"rId{max_id + 2}"
        for new_rid, fname in [(rid_default, "header1.xml"), (rid_first, "header2.xml")]:
            rel_el = etree.SubElement(rels_root, f"{{{rels_ns}}}Relationship")
            rel_el.set("Id", new_rid)
            rel_el.set("Type", HEADER_REL_TYPE)
            rel_el.set("Target", fname)
        new_rels_data = etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)

        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename == "[Content_Types].xml":
                ct_root = etree.fromstring(data)
                for i in range(1, 3):
                    pn = f"/word/header{i}.xml"
                    existing = ct_root.findall(f"{{{ct_ns}}}Override[@PartName='{pn}']")
                    if not existing:
                        ov = etree.SubElement(ct_root, f"{{{ct_ns}}}Override")
                        ov.set("PartName", pn)
                        ov.set("ContentType", HEADER_CONTENT_TYPE)
                data = etree.tostring(ct_root, xml_declaration=True, encoding="UTF-8", standalone=True)

            elif item.filename == "word/_rels/document.xml.rels":
                data = new_rels_data

            elif item.filename == "word/document.xml":
                doc_root = etree.fromstring(data)
                for sect_pr in doc_root.iter(f"{{{W}}}sectPr"):
                    etree.SubElement(sect_pr, f"{{{W}}}titlePg")
                    h_def = etree.SubElement(sect_pr, f"{{{W}}}headerReference")
                    h_def.set(f"{{{R}}}id", rid_default)
                    h_def.set(f"{{{W}}}type", "default")
                    h_first = etree.SubElement(sect_pr, f"{{{W}}}headerReference")
                    h_first.set(f"{{{R}}}id", rid_first)
                    h_first.set(f"{{{W}}}type", "first")
                data = etree.tostring(doc_root, xml_declaration=True, encoding="UTF-8", standalone=True)

            zout.writestr(item, data)
        zout.writestr("word/header1.xml", default_hdr)
        zout.writestr("word/header2.xml", first_hdr)

    out_buf.seek(0)
    save_story_fixture("first-page-header", out_buf, {
        "name": "first-page-header",
        "spec_ref": "ISO 29500-1 §17.10.6",
        "description": "Section with titlePg flag, default + first-page headers",
    })


def make_footnote_content() -> None:
    doc = Document()
    p = doc.add_paragraph()
    for child in list(p._p):
        p._p.remove(child)

    r1 = make_element("w:r")
    t1 = make_element("w:t")
    t1.text = "Text with footnote"
    t1.set(qn("xml:space"), "preserve")
    r1.append(t1)
    p._p.append(r1)

    r2 = make_element("w:r")
    fn_ref = make_element("w:footnoteReference", {"w:id": "1"})
    r2.append(fn_ref)
    p._p.append(r2)

    root = etree.Element(f"{{{W}}}footnotes", nsmap={"w": W, "r": R})
    fn0 = etree.SubElement(root, f"{{{W}}}footnote", attrib={f"{{{W}}}id": "0", f"{{{W}}}type": "separator"})
    p0 = etree.SubElement(fn0, f"{{{W}}}p")
    r0 = etree.SubElement(p0, f"{{{W}}}r")
    t0 = etree.SubElement(r0, f"{{{W}}}t")
    t0.text = ""

    fn1 = etree.SubElement(root, f"{{{W}}}footnote", attrib={f"{{{W}}}id": "1"})
    p1_el = etree.SubElement(fn1, f"{{{W}}}p")
    r_fn = etree.SubElement(p1_el, f"{{{W}}}r")
    rPr_fn = etree.SubElement(r_fn, f"{{{W}}}rPr")
    etree.SubElement(rPr_fn, f"{{{W}}}b")
    t_fn = etree.SubElement(r_fn, f"{{{W}}}t")
    t_fn.text = "Bold footnote content."

    footnotes_xml = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    parts = {"footnotes.xml": (footnotes_xml, FOOTNOTES_CONTENT_TYPE, FOOTNOTES_REL_TYPE)}
    out_buf = _add_story_parts(doc, parts)
    save_story_fixture("footnote-content", out_buf, {
        "name": "footnote-content",
        "spec_ref": "ISO 29500-1 §17.11",
        "description": "Footnote with bold-formatted text content",
    })


def make_header_footer_content() -> None:
    doc = Document()
    doc.add_paragraph("Body content with formatted header/footer.")
    header_xml = _build_header_xml_with_formatting([
        {"text": "Bold Header Text", "bold": True, "italic": False},
        {"text": "Italic Header Line", "bold": False, "italic": True},
    ])
    footer_xml = _build_footer_xml(["Page Footer"])
    parts = {
        "header1.xml": (header_xml, HEADER_CONTENT_TYPE, HEADER_REL_TYPE),
        "footer1.xml": (footer_xml, FOOTER_CONTENT_TYPE, FOOTER_REL_TYPE),
    }
    sect_pr_additions = [
        ("headerReference", "header1.xml", "default"),
        ("footerReference", "footer1.xml", "default"),
    ]
    out_buf = _add_story_parts(doc, parts, sect_pr_additions)
    save_story_fixture("header-footer-content", out_buf, {
        "name": "header-footer-content",
        "spec_ref": "ISO 29500-1 §17.10",
        "description": "Header with bold and italic paragraphs; footer with plain text",
    })


def make_footnote_multi_paragraph() -> None:
    doc = Document()
    p = doc.add_paragraph()
    for child in list(p._p):
        p._p.remove(child)

    r1 = make_element("w:r")
    t1 = make_element("w:t")
    t1.text = "Text with multi-paragraph footnote"
    t1.set(qn("xml:space"), "preserve")
    r1.append(t1)
    p._p.append(r1)

    r2 = make_element("w:r")
    fn_ref = make_element("w:footnoteReference", {"w:id": "1"})
    r2.append(fn_ref)
    p._p.append(r2)

    footnotes_xml = _build_footnotes_xml([
        {"id": "0", "type": "separator", "paragraphs": [""]},
        {"id": "1", "type": None, "paragraphs": [
            "First paragraph of the footnote.",
            "Second paragraph of the footnote.",
            "Third paragraph of the footnote.",
        ]},
    ])
    parts = {"footnotes.xml": (footnotes_xml, FOOTNOTES_CONTENT_TYPE, FOOTNOTES_REL_TYPE)}
    out_buf = _add_story_parts(doc, parts)
    save_story_fixture("footnote-multi-paragraph", out_buf, {
        "name": "footnote-multi-paragraph",
        "spec_ref": "ISO 29500-1 §17.11",
        "description": "Footnote with 3 paragraphs of content",
    })


def main() -> None:
    print("Generating story spec-compliance fixtures:")
    make_header_types()
    make_footer_types()
    make_footnote_references()
    make_endnote_references()
    make_footnote_types()
    make_multi_section_headers()
    make_first_page_header()
    make_footnote_content()
    make_header_footer_content()
    make_footnote_multi_paragraph()
    print("Done.")


if __name__ == "__main__":
    main()

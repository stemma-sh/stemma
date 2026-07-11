# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for deep style resolution audit (ISO 29500-1 §17.7).

Exercises known gaps in style cascade, default character style application,
deep basedOn numPr inheritance, wholeTable conditional formatting,
conditional precedence, and toggle property XOR through basedOn chains.

Run:  python create_docs.py
"""

import json
from pathlib import Path
from lxml import etree

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  styles-audit/{name}/")


# =========================================================================
# 1. Default character style (§17.7.4.17)
# =========================================================================

def make_default_char_style():
    """
    ISO 29500-1 §17.7.4.17: When a run has no explicit rStyle, the default
    character style (w:type="character" w:default="1") should apply.

    We create a document with:
    - A default character style that sets font size to 28 half-points (14pt)
      and color to FF0000 (red).
    - A paragraph with an unstyled run. The run should inherit 28/FF0000.
    - A paragraph with a run using an explicit character style (overrides).
    """
    doc = Document()

    # Access styles.xml
    styles_part = doc.part.element.find(qn("w:body")).getparent()
    styles_el = doc.styles.element

    # Add a default character style "DefaultFont" with sz=28, color=FF0000
    style_el = make_element("w:style", {
        "w:type": "character",
        "w:default": "1",
        "w:styleId": "DefaultParagraphFont",
    })
    # MS-OI29500 §17.7.4.17a: Word ignores child elements of
    # "DefaultParagraphFont" style. So we use a DIFFERENT styleId.
    # Actually, let's use a custom name to avoid the special-case.
    style_el = make_element("w:style", {
        "w:type": "character",
        "w:default": "1",
        "w:styleId": "CustomDefaultChar",
    })
    name_el = make_element("w:name", {"w:val": "Custom Default Char"})
    style_el.append(name_el)
    rpr = make_element("w:rPr")
    sz = make_element("w:sz", {"w:val": "28"})
    color = make_element("w:color", {"w:val": "FF0000"})
    rpr.append(sz)
    rpr.append(color)
    style_el.append(rpr)
    styles_el.append(style_el)

    # Remove any existing DefaultParagraphFont to avoid conflict
    for existing in styles_el.findall(qn("w:style")):
        sid = existing.get(qn("w:styleId"))
        if sid == "DefaultParagraphFont":
            styles_el.remove(existing)

    # Paragraph 1: unstyled run — should get default char style properties
    p1 = doc.add_paragraph()
    p1_el = p1._element
    # Clear any auto-generated pPr style
    run1 = p1.add_run("Unstyled run")
    # Make sure no rStyle is set on the run
    rpr1 = run1._element.find(qn("w:rPr"))
    if rpr1 is not None:
        rstyle = rpr1.find(qn("w:rStyle"))
        if rstyle is not None:
            rpr1.remove(rstyle)

    # Paragraph 2: run with explicit char style that sets color=0000FF
    char_override_style = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "BlueOverride",
    })
    char_name = make_element("w:name", {"w:val": "Blue Override"})
    char_override_style.append(char_name)
    char_rpr = make_element("w:rPr")
    char_color = make_element("w:color", {"w:val": "0000FF"})
    char_rpr.append(char_color)
    char_override_style.append(char_rpr)
    styles_el.append(char_override_style)

    p2 = doc.add_paragraph()
    run2 = p2.add_run("Blue override run")
    run2_rpr = run2._element.get_or_add_rPr()
    rstyle = make_element("w:rStyle", {"w:val": "BlueOverride"})
    run2_rpr.insert(0, rstyle)

    save_fixture("default-char-style", doc, {
        "spec": "ISO 29500-1 §17.7.4.17",
        "description": "Default character style applies to unstyled runs",
        "expected": {
            "para_1_run": {"font_size": 28, "color": "FF0000"},
            "para_2_run": {"color": "0000FF", "font_size": 28},
        },
    })


# =========================================================================
# 2. Deep basedOn chain for numPr (§17.7.4.3 + §17.7.4.14)
# =========================================================================

def make_deep_based_on_numpr():
    """
    Three-level basedOn chain where the ROOT style defines numPr.
    StyleA (root): numPr with numId=1, ilvl=0
    StyleB (basedOn A): adds left indent
    StyleC (basedOn B): adds bold rPr

    Paragraph using StyleC should inherit numPr from StyleA through the chain.
    """
    doc = Document()
    styles_el = doc.styles.element

    # Create a simple numbering definition in numbering.xml
    numbering_part = doc.part.numbering_part
    numbering_el = numbering_part.element

    # Add abstract num with a decimal format
    abstract_num = make_element("w:abstractNum", {"w:abstractNumId": "1"})
    lvl = make_element("w:lvl", {"w:ilvl": "0"})
    start = make_element("w:start", {"w:val": "1"})
    num_fmt = make_element("w:numFmt", {"w:val": "decimal"})
    lvl_text = make_element("w:lvlText", {"w:val": "%1."})
    lvl.append(start)
    lvl.append(num_fmt)
    lvl.append(lvl_text)
    abstract_num.append(lvl)
    numbering_el.append(abstract_num)

    num = make_element("w:num", {"w:numId": "1"})
    abstract_ref = make_element("w:abstractNumId", {"w:val": "1"})
    num.append(abstract_ref)
    numbering_el.append(num)

    # StyleA: root style with numPr
    style_a = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "NumStyleA",
    })
    style_a.append(make_element("w:name", {"w:val": "Num Style A"}))
    ppr_a = make_element("w:pPr")
    numpr_a = make_element("w:numPr")
    numpr_a.append(make_element("w:numId", {"w:val": "1"}))
    numpr_a.append(make_element("w:ilvl", {"w:val": "0"}))
    ppr_a.append(numpr_a)
    style_a.append(ppr_a)
    styles_el.append(style_a)

    # StyleB: basedOn A, adds indent
    style_b = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "NumStyleB",
    })
    style_b.append(make_element("w:name", {"w:val": "Num Style B"}))
    style_b.append(make_element("w:basedOn", {"w:val": "NumStyleA"}))
    ppr_b = make_element("w:pPr")
    ind_b = make_element("w:ind", {"w:left": "720"})
    ppr_b.append(ind_b)
    style_b.append(ppr_b)
    styles_el.append(style_b)

    # StyleC: basedOn B, adds bold
    style_c = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "NumStyleC",
    })
    style_c.append(make_element("w:name", {"w:val": "Num Style C"}))
    style_c.append(make_element("w:basedOn", {"w:val": "NumStyleB"}))
    rpr_c = make_element("w:rPr")
    rpr_c.append(make_element("w:b"))
    style_c.append(rpr_c)
    styles_el.append(style_c)

    # Paragraph 1: uses StyleC — should inherit numPr from A
    p1 = doc.add_paragraph("Item via StyleC")
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.insert(0, make_element("w:pStyle", {"w:val": "NumStyleC"}))

    # Paragraph 2: uses StyleA directly — should have numPr
    p2 = doc.add_paragraph("Item via StyleA")
    p2_ppr = p2._element.get_or_add_pPr()
    p2_ppr.insert(0, make_element("w:pStyle", {"w:val": "NumStyleA"}))

    # Paragraph 3: uses StyleB — should inherit numPr from A
    p3 = doc.add_paragraph("Item via StyleB")
    p3_ppr = p3._element.get_or_add_pPr()
    p3_ppr.insert(0, make_element("w:pStyle", {"w:val": "NumStyleB"}))

    save_fixture("deep-numpr-chain", doc, {
        "spec": "ISO 29500-1 §17.7.4.3 + §17.7.4.14",
        "description": "numPr inheritance through 3-level basedOn chain",
        "expected": {
            "para_1_numpr": "inherited from NumStyleA via chain",
            "para_2_numpr": "direct from NumStyleA",
            "para_3_numpr": "inherited from NumStyleA via NumStyleB",
        },
    })


# =========================================================================
# 3. wholeTable conditional (§17.7.6)
# =========================================================================

def make_whole_table_conditional():
    """
    Table style with wholeTable tblStylePr that sets cell shading to AABBCC
    and bold text. All cells in the table should receive these properties
    (unless overridden by higher-precedence conditionals).
    """
    doc = Document()
    styles_el = doc.styles.element

    # Table style with wholeTable conditional
    tbl_style = make_element("w:style", {
        "w:type": "table",
        "w:styleId": "WholeTableStyle",
    })
    tbl_style.append(make_element("w:name", {"w:val": "Whole Table Style"}))

    # Base table borders
    tblpr = make_element("w:tblPr")
    tbl_borders = make_element("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = make_element(f"w:{edge}", {
            "w:val": "single",
            "w:sz": "4",
            "w:space": "0",
            "w:color": "000000",
        })
        tbl_borders.append(border)
    tblpr.append(tbl_borders)
    tbl_style.append(tblpr)

    # wholeTable conditional: shading=AABBCC + bold
    whole_tbl_pr = make_element("w:tblStylePr", {"w:type": "wholeTable"})
    tc_pr = make_element("w:tcPr")
    shd = make_element("w:shd", {
        "w:val": "clear",
        "w:color": "auto",
        "w:fill": "AABBCC",
    })
    tc_pr.append(shd)
    whole_tbl_pr.append(tc_pr)
    rpr_whole = make_element("w:rPr")
    rpr_whole.append(make_element("w:b"))
    whole_tbl_pr.append(rpr_whole)
    tbl_style.append(whole_tbl_pr)

    styles_el.append(tbl_style)

    # Create a 2x2 table using the style
    table = doc.add_table(rows=2, cols=2)

    # Set table style reference
    tbl_el = table._element
    tbl_pr_el = tbl_el.find(qn("w:tblPr"))
    if tbl_pr_el is None:
        tbl_pr_el = make_element("w:tblPr")
        tbl_el.insert(0, tbl_pr_el)
    tbl_style_ref = make_element("w:tblStyle", {"w:val": "WholeTableStyle"})
    tbl_pr_el.insert(0, tbl_style_ref)

    # tblLook: disable firstRow/lastRow/firstCol/lastCol so only wholeTable applies
    tbl_look = make_element("w:tblLook", {
        "w:val": "0000",
        "w:firstRow": "0",
        "w:lastRow": "0",
        "w:firstColumn": "0",
        "w:lastColumn": "0",
        "w:noHBand": "1",
        "w:noVBand": "1",
    })
    tbl_pr_el.append(tbl_look)

    # Add text to cells
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"Cell {i},{j}"

    save_fixture("whole-table-conditional", doc, {
        "spec": "ISO 29500-1 §17.7.6 + §17.7.6.1",
        "description": "wholeTable tblStylePr applies shading and bold to all cells",
        "expected": {
            "all_cells_shading": "AABBCC",
            "all_cells_bold": True,
        },
    })


# =========================================================================
# 4. Conditional precedence: firstRow + firstCol intersection (§17.7.6)
# =========================================================================

def make_conditional_precedence():
    """
    Per MS-OI29500 §17.7.6(c), conditional formatting precedence is:
    wholeTable < bands < firstCol/lastCol < firstRow/lastRow < corners.

    We create a table style with:
    - firstRow: shading=FF0000 (red)
    - firstCol: shading=0000FF (blue)
    Cell at (0,0) is both firstRow AND firstCol. Per spec, firstRow has
    higher precedence than firstCol, so cell (0,0) should be red.
    """
    doc = Document()
    styles_el = doc.styles.element

    tbl_style = make_element("w:style", {
        "w:type": "table",
        "w:styleId": "PrecedenceTable",
    })
    tbl_style.append(make_element("w:name", {"w:val": "Precedence Table"}))

    # Base table borders
    tblpr = make_element("w:tblPr")
    tbl_borders = make_element("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = make_element(f"w:{edge}", {
            "w:val": "single",
            "w:sz": "4",
            "w:space": "0",
            "w:color": "000000",
        })
        tbl_borders.append(border)
    tblpr.append(tbl_borders)
    tbl_style.append(tblpr)

    # firstRow conditional: red shading
    first_row_pr = make_element("w:tblStylePr", {"w:type": "firstRow"})
    tc_pr_fr = make_element("w:tcPr")
    shd_fr = make_element("w:shd", {
        "w:val": "clear", "w:color": "auto", "w:fill": "FF0000",
    })
    tc_pr_fr.append(shd_fr)
    first_row_pr.append(tc_pr_fr)
    tbl_style.append(first_row_pr)

    # firstCol conditional: blue shading
    first_col_pr = make_element("w:tblStylePr", {"w:type": "firstCol"})
    tc_pr_fc = make_element("w:tcPr")
    shd_fc = make_element("w:shd", {
        "w:val": "clear", "w:color": "auto", "w:fill": "0000FF",
    })
    tc_pr_fc.append(shd_fc)
    first_col_pr.append(tc_pr_fc)
    tbl_style.append(first_col_pr)

    styles_el.append(tbl_style)

    # 3x3 table
    table = doc.add_table(rows=3, cols=3)
    tbl_el = table._element
    tbl_pr_el = tbl_el.find(qn("w:tblPr"))
    if tbl_pr_el is None:
        tbl_pr_el = make_element("w:tblPr")
        tbl_el.insert(0, tbl_pr_el)
    tbl_pr_el.insert(0, make_element("w:tblStyle", {"w:val": "PrecedenceTable"}))

    # tblLook: enable firstRow and firstCol
    tbl_look = make_element("w:tblLook", {
        "w:val": "00A0",
        "w:firstRow": "1",
        "w:lastRow": "0",
        "w:firstColumn": "1",
        "w:lastColumn": "0",
        "w:noHBand": "1",
        "w:noVBand": "1",
    })
    tbl_pr_el.append(tbl_look)

    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"R{i}C{j}"

    save_fixture("conditional-precedence", doc, {
        "spec": "MS-OI29500 §17.7.6(c)",
        "description": "firstRow vs firstCol precedence at intersection cell",
        "expected": {
            "cell_0_0": "FF0000 (firstRow wins over firstCol)",
            "cell_0_1": "FF0000 (firstRow)",
            "cell_1_0": "0000FF (firstCol)",
            "cell_1_1": "no conditional shading",
        },
    })


# =========================================================================
# 5. Toggle property XOR through basedOn chain (§17.7.3)
# =========================================================================

def make_toggle_xor_based_on():
    """
    ISO 29500-1 §17.7.3: Toggle properties use XOR across hierarchy levels.

    basedOn chain for paragraph styles:
    - StyleRoot: bold=on
    - StyleChild (basedOn Root): bold=on
    Since both levels set bold=on, XOR → bold should be OFF at the style level.
    Then at the run level, the para style contributes this resolved value.

    Also test:
    - StyleOffChild (basedOn Root): bold=off (explicit val="0")
    XOR of on ^ off = on.
    """
    doc = Document()
    styles_el = doc.styles.element

    # StyleRoot: paragraph style with bold=on in rPr
    style_root = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "ToggleRoot",
    })
    style_root.append(make_element("w:name", {"w:val": "Toggle Root"}))
    rpr_root = make_element("w:rPr")
    rpr_root.append(make_element("w:b"))
    style_root.append(rpr_root)
    styles_el.append(style_root)

    # StyleChild: basedOn Root, bold=on
    style_child = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "ToggleChild",
    })
    style_child.append(make_element("w:name", {"w:val": "Toggle Child"}))
    style_child.append(make_element("w:basedOn", {"w:val": "ToggleRoot"}))
    rpr_child = make_element("w:rPr")
    rpr_child.append(make_element("w:b"))
    style_child.append(rpr_child)
    styles_el.append(style_child)

    # StyleOffChild: basedOn Root, bold=off
    style_off = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "ToggleOffChild",
    })
    style_off.append(make_element("w:name", {"w:val": "Toggle Off Child"}))
    style_off.append(make_element("w:basedOn", {"w:val": "ToggleRoot"}))
    rpr_off = make_element("w:rPr")
    rpr_off.append(make_element("w:b", {"w:val": "0"}))
    style_off.append(rpr_off)
    styles_el.append(style_off)

    # Paragraph 1: ToggleRoot — bold=on
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.insert(0, make_element("w:pStyle", {"w:val": "ToggleRoot"}))
    p1.add_run("Root bold on")

    # Paragraph 2: ToggleChild — both layers bold=on, XOR → off
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    p2_ppr.insert(0, make_element("w:pStyle", {"w:val": "ToggleChild"}))
    p2.add_run("Child XOR off")

    # Paragraph 3: ToggleOffChild — root=on child=off, XOR → on
    p3 = doc.add_paragraph()
    p3_ppr = p3._element.get_or_add_pPr()
    p3_ppr.insert(0, make_element("w:pStyle", {"w:val": "ToggleOffChild"}))
    p3.add_run("Off child XOR on")

    save_fixture("toggle-xor-based-on", doc, {
        "spec": "ISO 29500-1 §17.7.3 + §17.7.4.3",
        "description": "Toggle XOR through basedOn chain (bold on/on, on/off)",
        "expected": {
            "para_1": "bold ON (single level)",
            "para_2": "bold OFF (on XOR on = off)",
            "para_3": "bold ON (on XOR off = on)",
        },
    })


# =========================================================================
# Main
# =========================================================================

if __name__ == "__main__":
    print("\n── Styles Audit Fixtures ──")
    make_default_char_style()
    make_deep_based_on_numpr()
    make_whole_table_conditional()
    make_conditional_precedence()
    make_toggle_xor_based_on()
    print("\nDone.")

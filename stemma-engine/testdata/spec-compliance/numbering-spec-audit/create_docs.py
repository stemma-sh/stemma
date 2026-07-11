# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for the numbering spec-audit round.

Tests numbering behavioral constraints from ISO 29500-1 §17.9 and MS-OI29500
that are not yet covered by existing spec_numbering* tests.

Run:  uv run create_docs.py
"""

import json
from pathlib import Path
from lxml import etree

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

# ── XML namespace helpers ────────────────────────────────────────────────

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def w(tag: str) -> str:
    """Return a fully-qualified wordprocessingml tag."""
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    """Create an OxmlElement with optional attributes."""
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


# ── Save helpers ─────────────────────────────────────────────────────────

def save_fixture(
    name: str,
    doc: DocxDocument,
    metadata: dict,
    filename: str = "input.docx",
) -> None:
    """Save a single-doc fixture."""
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / filename))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  numbering-spec-audit/{name}/")


def _inject_numbering_xml(doc: DocxDocument, numbering_xml: str) -> None:
    """Replace the numbering.xml part with custom XML."""
    numbering_part = None
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        pass

    if numbering_part is None:
        doc.add_paragraph("dummy", style="List Bullet")
        body = doc.element.body
        last_p = body.findall(w("p"))[-1]
        body.remove(last_p)
        numbering_part = doc.part.numbering_part

    numbering_part._element = etree.fromstring(numbering_xml.encode("utf-8"))


def _add_numbered_para(doc: DocxDocument, text: str, num_id: str, ilvl: str = "0"):
    """Add a paragraph with numPr referencing the given numId/ilvl."""
    p = doc.add_paragraph(text)
    pPr = p._p.get_or_add_pPr()
    numPr = make_element("w:numPr")
    numPr.append(make_element("w:ilvl", {"w:val": ilvl}))
    numPr.append(make_element("w:numId", {"w:val": num_id}))
    pPr.append(numPr)
    return p


# =========================================================================
# 1. Interleaved lists — counter continuity (§17.9.11)
# =========================================================================

def make_interleaved_lists() -> None:
    """§17.9.11: Counter increments per paragraph at that level, sequential or not.

    When two different numId lists are interleaved, switching back to a
    previous numId should CONTINUE its counter, not restart it.

    Fixture:
      - numId=1: decimal list (abstractNum 0)
      - numId=2: lowerLetter list (abstractNum 1)
      - Sequence: numId=1 x2, numId=2 x2, numId=1 x2

    Expected: 1., 2., a), b), 3., 4.
    The critical assertion: the 5th paragraph (numId=1 again) should be "3."
    not "1." — the counter for numId=1 must survive the switch to numId=2.
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:abstractNum w:abstractNumId="1">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerLetter"/>
      <w:lvlText w:val="%1)"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
  <w:num w:numId="2">
    <w:abstractNumId w:val="1"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    _add_numbered_para(doc, "Decimal first", "1", "0")    # 1.
    _add_numbered_para(doc, "Decimal second", "1", "0")   # 2.
    _add_numbered_para(doc, "Letter first", "2", "0")     # a)
    _add_numbered_para(doc, "Letter second", "2", "0")    # b)
    _add_numbered_para(doc, "Decimal continues", "1", "0")  # 3. (NOT 1.)
    _add_numbered_para(doc, "Decimal continues", "1", "0")  # 4. (NOT 2.)

    save_fixture("interleaved-lists", doc, {
        "name": "interleaved-lists",
        "spec_ref": "ISO 29500-1 §17.9.11",
        "description": (
            "Two lists (numId=1 decimal, numId=2 lowerLetter) interleaved. "
            "Counter for numId=1 should survive the switch to numId=2 and back."
        ),
        "expected_behavior": (
            "Sequence: '1.', '2.', 'a)', 'b)', '3.', '4.'. "
            "numId=1 continues at 3 after returning from numId=2."
        ),
    })


# =========================================================================
# 2. lvlText forward reference — entire text ignored (MS-OI29500 §17.9.11c)
# =========================================================================

def make_lvltext_forward_ref() -> None:
    """MS-OI29500 §17.9.11c: Forward %N references invalidate entire lvlText.

    When lvlText at level 0 contains %3 (referencing level 2, which is
    higher than current level 0), Word ignores the ENTIRE lvlText and
    produces empty text.

    Fixture:
      - Level 0: decimal, lvlText="%1.%3" — %3 is a forward reference
      - Level 1: decimal, lvlText="%1.%2" — normal (control)

    Expected:
      - Level 0: empty string (entire lvlText ignored per MS-OI29500)
      - Level 1: "1.1" (normal multi-level)
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1.%3"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="1">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1.%2"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="2">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1.%2.%3"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    _add_numbered_para(doc, "Forward ref (should be empty)", "1", "0")
    _add_numbered_para(doc, "Normal multi-level", "1", "1")
    _add_numbered_para(doc, "Forward ref second", "1", "0")

    save_fixture("lvltext-forward-ref", doc, {
        "name": "lvltext-forward-ref",
        "spec_ref": "MS-OI29500 §17.9.11c",
        "description": (
            "Level 0 has lvlText='%1.%3' where %3 references level 2 "
            "(higher than current level 0). Per MS-OI29500, the entire "
            "lvlText is ignored."
        ),
        "expected_behavior": (
            "Level 0: empty synthesized text (forward reference invalidates). "
            "Level 1: '1.1' (normal)."
        ),
    })


# =========================================================================
# 3. numFmt omitted defaults to decimal (§17.9.17)
# =========================================================================

def make_numfmt_default_decimal() -> None:
    """§17.9.17: If numFmt is omitted, default is decimal.

    Fixture:
      - Level 0: NO w:numFmt element (should default to decimal)
      - Level 1: explicit lowerLetter (control)

    Expected: Level 0 produces "1.", "2." (decimal default).
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="1">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerLetter"/>
      <w:lvlText w:val="%2)"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    _add_numbered_para(doc, "No numFmt (default decimal)", "1", "0")
    _add_numbered_para(doc, "Second (should be 2.)", "1", "0")
    _add_numbered_para(doc, "Control lowerLetter", "1", "1")

    save_fixture("numfmt-default-decimal", doc, {
        "name": "numfmt-default-decimal",
        "spec_ref": "ISO 29500-1 §17.9.17",
        "description": (
            "Level 0 has no w:numFmt element. Per §17.9.17, "
            "the default is decimal."
        ),
        "expected_behavior": (
            "Level 0: '1.', '2.' (decimal). Level 1: 'a)' (explicit lowerLetter)."
        ),
    })


# =========================================================================
# 4. lvlRestart default vs explicit — 3-level cascade (§17.9.10)
# =========================================================================

def make_lvl_restart_default_cascade() -> None:
    """§17.9.10: Default lvlRestart (omitted) restarts when any higher level appears.

    Fixture: 3-level list, all levels use default restart behavior.
      - Level 0: decimal
      - Level 1: lowerLetter (default restart)
      - Level 2: lowerRoman (default restart)

    Sequence: 1. / a) / i) / ii) / b) / 2. / a) / i)

    Key assertions:
      - After "b)" at level 1, level 2 restarts to "i)" (default: level 1 triggers)
      - After "2." at level 0, both level 1 and level 2 restart
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="1">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerLetter"/>
      <w:lvlText w:val="%2)"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="2">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerRoman"/>
      <w:lvlText w:val="%3)"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    _add_numbered_para(doc, "Top level", "1", "0")          # 1.
    _add_numbered_para(doc, "Sub a", "1", "1")              # a)
    _add_numbered_para(doc, "Sub-sub first", "1", "2")      # i)
    _add_numbered_para(doc, "Sub-sub second", "1", "2")     # ii)
    # Back to level 1 — level 2 should restart (default behavior)
    _add_numbered_para(doc, "Sub b", "1", "1")              # b)
    _add_numbered_para(doc, "Sub-sub restarted", "1", "2")  # i) (restarted)
    # Back to level 0 — both levels 1 and 2 restart
    _add_numbered_para(doc, "Top level 2", "1", "0")        # 2.
    _add_numbered_para(doc, "Sub a again", "1", "1")        # a)
    _add_numbered_para(doc, "Sub-sub first again", "1", "2")  # i)

    save_fixture("lvl-restart-default-cascade", doc, {
        "name": "lvl-restart-default-cascade",
        "spec_ref": "ISO 29500-1 §17.9.10",
        "description": (
            "3-level list with default restart behavior (no lvlRestart elements). "
            "All deeper levels restart when any higher level appears."
        ),
        "expected_behavior": (
            "Sequence: 1. / a) / i) / ii) / b) / i) / 2. / a) / i). "
            "Level 2 restarts when level 1 appears (default behavior)."
        ),
    })


# =========================================================================
# 5. Interleaved lists with multi-level (§17.9.11)
# =========================================================================

def make_interleaved_multi_level() -> None:
    """§17.9.11: Interleaved lists preserve per-level counters.

    This is the common legal doc pattern: clause numbering interleaved
    with sub-lists. When returning to the outer list, it continues.

    Fixture:
      - numId=1: 2-level outline (decimal / lowerLetter)
      - numId=2: bullets (separate list)
      - Sequence: numId=1 at levels 0,1 → numId=2 → numId=1 at levels 0,1

    Expected: 1. / a) / b) / • / • / 2. / c) / d)
    Level 1 of numId=1 should continue at c) after returning, because
    the counter survived the interruption by numId=2 and also because
    level 0 hit "2." which restarts level 1 — so this is actually:
    1. / a) / b) / • / • / 2. / a) / b)

    Wait — level 0 "2." triggers restart of level 1, so level 1 restarts.
    Let me redesign: the interesting test is whether level 0 continues.
    """
    doc = Document()

    # U+F0B7 is the PUA bullet character used by Word's Symbol font.
    # Our code maps PUA chars >= U+F000 to "•".
    bullet_char = "\uF0B7"

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="1">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerLetter"/>
      <w:lvlText w:val="%2)"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:abstractNum w:abstractNumId="1">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="bullet"/>
      <w:lvlText w:val="{bullet_char}"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
  <w:num w:numId="2">
    <w:abstractNumId w:val="1"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # First run of numId=1
    _add_numbered_para(doc, "Clause one", "1", "0")       # 1.
    _add_numbered_para(doc, "Clause two", "1", "0")       # 2.
    _add_numbered_para(doc, "Sub a under 2", "1", "1")    # a)
    # Interruption by bullets
    _add_numbered_para(doc, "Bullet note", "2", "0")      # •
    _add_numbered_para(doc, "Another bullet", "2", "0")   # •
    # Return to numId=1 at level 0 — should continue at 3.
    _add_numbered_para(doc, "Clause three", "1", "0")     # 3. (NOT 1.)
    _add_numbered_para(doc, "Clause four", "1", "0")      # 4. (NOT 2.)

    save_fixture("interleaved-multi-level", doc, {
        "name": "interleaved-multi-level",
        "spec_ref": "ISO 29500-1 §17.9.11",
        "description": (
            "Legal doc pattern: numbered clauses interrupted by bullet list, "
            "then clauses continue. Counter must survive the interruption."
        ),
        "expected_behavior": (
            "Sequence: '1.', '2.', 'a)', bullet, bullet, '3.', '4.'. "
            "numId=1 level 0 continues at 3 after bullet interruption."
        ),
    })


# =========================================================================
# 6. numFmt=bullet uses literal lvlText (§17.9.17)
# =========================================================================

def make_bullet_literal_text() -> None:
    """§17.9.17: When numFmt is bullet, lvlText is used literally.

    The %N expansion should NOT happen for bullet format. The literal
    text of lvlText is used as the bullet character.

    Fixture:
      - Level 0: numFmt=bullet, lvlText="→" (literal arrow)
      - Level 1: numFmt=bullet, lvlText="–" (literal em-dash)

    Expected: "→" and "–" displayed literally, not expanded as %N.
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="bullet"/>
      <w:lvlText w:val="→"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="1">
      <w:start w:val="1"/>
      <w:numFmt w:val="bullet"/>
      <w:lvlText w:val="–"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    _add_numbered_para(doc, "Arrow bullet", "1", "0")
    _add_numbered_para(doc, "Another arrow", "1", "0")
    _add_numbered_para(doc, "Dash sub-bullet", "1", "1")

    save_fixture("bullet-literal-text", doc, {
        "name": "bullet-literal-text",
        "spec_ref": "ISO 29500-1 §17.9.17",
        "description": (
            "Bullet format uses literal lvlText. "
            "Level 0: '→', Level 1: '–'."
        ),
        "expected_behavior": (
            "'→', '→', '–' — literal bullet characters, not %N expanded."
        ),
    })


# =========================================================================

def main() -> None:
    print("Generating numbering spec-audit fixtures:")
    make_interleaved_lists()
    make_lvltext_forward_ref()
    make_numfmt_default_decimal()
    make_lvl_restart_default_cascade()
    make_interleaved_multi_level()
    make_bullet_literal_text()
    print("\nDone.")


if __name__ == "__main__":
    main()

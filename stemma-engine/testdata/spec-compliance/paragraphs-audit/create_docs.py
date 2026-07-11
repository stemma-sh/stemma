# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for paragraph properties deep audit (ISO 29500-1 §17.3.1).

These fixtures target paragraph properties that are likely missing or incomplete
in our domain model: bidi, textAlignment, suppressAutoHyphens, snapToGrid,
overflowPunct, adjustRightInd, wordWrap, framePr.

Run:  uv run create_docs.py
"""

import json
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    """Create an OxmlElement with optional attributes."""
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(
    name: str,
    doc: DocxDocument,
    metadata: dict,
    filename: str = "input.docx",
) -> None:
    """Save a single-doc fixture."""
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / filename))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  paragraphs-audit/{name}/")


# =========================================================================
# bidi (§17.3.1.6) — Right to Left Paragraph Layout
# =========================================================================

def make_bidi() -> None:
    """ISO 29500-1 §17.3.1.6: bidi sets right-to-left paragraph direction.

    P0: bidi=true (RTL paragraph)
    P1: no bidi (LTR, default)
    """
    doc = Document()

    # P0: bidi paragraph
    p0 = doc.add_paragraph("RTL paragraph with bidi enabled.")
    p0._p.get_or_add_pPr().append(make_element("w:bidi"))

    # P1: normal LTR paragraph (control)
    doc.add_paragraph("Normal LTR paragraph (no bidi).")

    save_fixture("bidi", doc, {
        "name": "bidi",
        "spec_ref": "ISO 29500-1 §17.3.1.6",
        "description": (
            "P0: bidi=true (RTL paragraph direction). "
            "P1: no bidi (LTR, default)."
        ),
        "expected_behavior": (
            "P0: paragraph direction is RTL — affects ind, jc, tab, textDirection. "
            "P1: normal LTR paragraph."
        ),
    })


# =========================================================================
# textAlignment (§17.3.1.39) — Vertical Character Alignment on Line
# =========================================================================

def make_text_alignment() -> None:
    """ISO 29500-1 §17.3.1.39: textAlignment controls vertical char alignment.

    P0: textAlignment=top
    P1: textAlignment=center
    P2: textAlignment=baseline
    P3: textAlignment=bottom
    P4: textAlignment=auto
    P5: no textAlignment (control)
    """
    doc = Document()

    for val in ["top", "center", "baseline", "bottom", "auto"]:
        p = doc.add_paragraph(f"Text alignment: {val}.")
        p._p.get_or_add_pPr().append(
            make_element("w:textAlignment", {"w:val": val})
        )

    # P5: no textAlignment (control)
    doc.add_paragraph("No textAlignment (default).")

    save_fixture("text-alignment", doc, {
        "name": "text-alignment",
        "spec_ref": "ISO 29500-1 §17.3.1.39",
        "description": (
            "P0-P4: textAlignment with values top/center/baseline/bottom/auto. "
            "P5: no textAlignment (control)."
        ),
        "expected_behavior": (
            "Each paragraph should preserve its textAlignment value in the model. "
            "P5 should have no textAlignment."
        ),
    })


# =========================================================================
# suppressAutoHyphens (§17.3.1.34) — Suppress Hyphenation for Paragraph
# =========================================================================

def make_suppress_auto_hyphens() -> None:
    """ISO 29500-1 §17.3.1.34: suppressAutoHyphens exempts paragraph from hyphenation.

    P0: suppressAutoHyphens=true
    P1: suppressAutoHyphens=false (explicit off)
    P2: no suppressAutoHyphens (control)
    """
    doc = Document()

    # P0: suppress hyphenation
    p0 = doc.add_paragraph("Hyphenation suppressed on this paragraph.")
    p0._p.get_or_add_pPr().append(make_element("w:suppressAutoHyphens"))

    # P1: explicit false
    p1 = doc.add_paragraph("Hyphenation explicitly enabled.")
    p1._p.get_or_add_pPr().append(
        make_element("w:suppressAutoHyphens", {"w:val": "0"})
    )

    # P2: control
    doc.add_paragraph("No suppressAutoHyphens (default).")

    save_fixture("suppress-auto-hyphens", doc, {
        "name": "suppress-auto-hyphens",
        "spec_ref": "ISO 29500-1 §17.3.1.34",
        "description": (
            "P0: suppressAutoHyphens=true. "
            "P1: suppressAutoHyphens=false (explicit). "
            "P2: no suppressAutoHyphens (control)."
        ),
        "expected_behavior": (
            "P0: hyphenation suppressed. "
            "P1: hyphenation explicitly not suppressed. "
            "P2: default (inherit from style hierarchy)."
        ),
    })


# =========================================================================
# snapToGrid (§17.3.1.32) — Use Document Grid for Inter-Line Spacing
# =========================================================================

def make_snap_to_grid() -> None:
    """ISO 29500-1 §17.3.1.32: snapToGrid controls document grid alignment.

    P0: snapToGrid=false (opt out of grid)
    P1: snapToGrid=true (explicit)
    P2: no snapToGrid (control — default is true per spec)
    """
    doc = Document()

    # P0: opt out of grid
    p0 = doc.add_paragraph("Snap to grid disabled.")
    p0._p.get_or_add_pPr().append(
        make_element("w:snapToGrid", {"w:val": "0"})
    )

    # P1: explicit true
    p1 = doc.add_paragraph("Snap to grid explicitly enabled.")
    p1._p.get_or_add_pPr().append(make_element("w:snapToGrid"))

    # P2: control
    doc.add_paragraph("No snapToGrid (default true per spec).")

    save_fixture("snap-to-grid", doc, {
        "name": "snap-to-grid",
        "spec_ref": "ISO 29500-1 §17.3.1.32",
        "description": (
            "P0: snapToGrid=false (opt out). "
            "P1: snapToGrid=true (explicit). "
            "P2: no snapToGrid (control)."
        ),
        "expected_behavior": (
            "P0: paragraph opts out of document grid spacing. "
            "P1: paragraph uses document grid. "
            "P2: default (true per spec)."
        ),
    })


# =========================================================================
# overflowPunct (§17.3.1.21) — Allow Punctuation Overflow
# =========================================================================

def make_overflow_punct() -> None:
    """ISO 29500-1 §17.3.1.21: overflowPunct controls punctuation overflow.

    P0: overflowPunct=false (disallow overflow)
    P1: overflowPunct=true (explicit — same as default)
    P2: no overflowPunct (control — default is true)
    """
    doc = Document()

    # P0: disable punctuation overflow
    p0 = doc.add_paragraph("Punctuation overflow disabled.")
    p0._p.get_or_add_pPr().append(
        make_element("w:overflowPunct", {"w:val": "0"})
    )

    # P1: explicit true
    p1 = doc.add_paragraph("Punctuation overflow explicitly enabled.")
    p1._p.get_or_add_pPr().append(make_element("w:overflowPunct"))

    # P2: control
    doc.add_paragraph("No overflowPunct (default true).")

    save_fixture("overflow-punct", doc, {
        "name": "overflow-punct",
        "spec_ref": "ISO 29500-1 §17.3.1.21",
        "description": (
            "P0: overflowPunct=false (disallow). "
            "P1: overflowPunct=true (explicit). "
            "P2: no overflowPunct (control)."
        ),
        "expected_behavior": (
            "P0: punctuation may not extend past text extents. "
            "P1: punctuation allowed to overflow. "
            "P2: default (true per spec)."
        ),
    })


# =========================================================================
# adjustRightInd (§17.3.1.1) — Auto-Adjust Right Indent for Doc Grid
# =========================================================================

def make_adjust_right_ind() -> None:
    """ISO 29500-1 §17.3.1.1: adjustRightInd controls right indent auto-adjustment.

    P0: adjustRightInd=false
    P1: adjustRightInd=true (explicit — same as default)
    P2: no adjustRightInd (control — default is true)
    """
    doc = Document()

    # P0: disable auto-adjust
    p0 = doc.add_paragraph("Auto-adjust right indent disabled.")
    p0._p.get_or_add_pPr().append(
        make_element("w:adjustRightInd", {"w:val": "0"})
    )

    # P1: explicit true
    p1 = doc.add_paragraph("Auto-adjust right indent explicitly enabled.")
    p1._p.get_or_add_pPr().append(make_element("w:adjustRightInd"))

    # P2: control
    doc.add_paragraph("No adjustRightInd (default true).")

    save_fixture("adjust-right-ind", doc, {
        "name": "adjust-right-ind",
        "spec_ref": "ISO 29500-1 §17.3.1.1",
        "description": (
            "P0: adjustRightInd=false. "
            "P1: adjustRightInd=true (explicit). "
            "P2: no adjustRightInd (control)."
        ),
        "expected_behavior": (
            "P0: right indent not auto-adjusted for document grid. "
            "P1: right indent auto-adjusted. "
            "P2: default (true per spec)."
        ),
    })


# =========================================================================
# wordWrap (§17.3.1.45) — Allow Line Breaking At Character Level
# =========================================================================

def make_word_wrap() -> None:
    """ISO 29500-1 §17.3.1.45: wordWrap controls word vs character line breaking.

    P0: wordWrap=false (break at character level)
    P1: wordWrap=true (explicit — break at word level)
    P2: no wordWrap (control — default is word-level)
    """
    doc = Document()

    # P0: character-level breaking
    p0 = doc.add_paragraph("Word wrap off — character-level breaking.")
    p0._p.get_or_add_pPr().append(
        make_element("w:wordWrap", {"w:val": "0"})
    )

    # P1: explicit word-level
    p1 = doc.add_paragraph("Word wrap on — word-level breaking.")
    p1._p.get_or_add_pPr().append(make_element("w:wordWrap"))

    # P2: control
    doc.add_paragraph("No wordWrap (default word-level).")

    save_fixture("word-wrap", doc, {
        "name": "word-wrap",
        "spec_ref": "ISO 29500-1 §17.3.1.45",
        "description": (
            "P0: wordWrap=false (character-level breaking). "
            "P1: wordWrap=true (word-level). "
            "P2: no wordWrap (control)."
        ),
        "expected_behavior": (
            "P0: break at character level. "
            "P1: break at word level. "
            "P2: default (word-level)."
        ),
    })


# =========================================================================
# framePr (§17.3.1.11) — Text Frame Properties
# =========================================================================

def make_frame_pr() -> None:
    """ISO 29500-1 §17.3.1.11: framePr defines text frame properties.

    P0: framePr with width, height, anchors, wrap
    P1: normal paragraph (control)
    """
    doc = Document()

    # P0: paragraph in a text frame
    p0 = doc.add_paragraph("Text frame paragraph with specific dimensions.")
    p0._p.get_or_add_pPr().append(make_element("w:framePr", {
        "w:w": "2191",
        "w:h": "811",
        "w:hRule": "exact",
        "w:hSpace": "180",
        "w:wrap": "around",
        "w:vAnchor": "text",
        "w:hAnchor": "page",
        "w:x": "1921",
    }))

    # P1: normal paragraph (control)
    doc.add_paragraph("Normal paragraph (no frame).")

    save_fixture("frame-pr", doc, {
        "name": "frame-pr",
        "spec_ref": "ISO 29500-1 §17.3.1.11",
        "description": (
            "P0: framePr with w=2191, h=811, hRule=exact, hSpace=180, "
            "wrap=around, vAnchor=text, hAnchor=page, x=1921. "
            "P1: normal paragraph (control)."
        ),
        "expected_behavior": (
            "P0: paragraph is part of a text frame with all attributes preserved. "
            "P1: no frame properties."
        ),
    })


# =========================================================================
# bidi + mirrorIndents interaction
# =========================================================================

def make_bidi_mirror_indents() -> None:
    """Interaction: bidi + mirrorIndents together.

    P0: bidi + mirrorIndents + asymmetric indent (start=1440, end=720)
    P1: bidi only + same indent (control)
    P2: mirrorIndents only + same indent (control)
    """
    doc = Document()

    # P0: bidi + mirrorIndents
    p0 = doc.add_paragraph("Bidi + mirrorIndents + asymmetric indent.")
    p0_ppr = p0._p.get_or_add_pPr()
    p0_ppr.append(make_element("w:bidi"))
    p0_ppr.append(make_element("w:mirrorIndents"))
    p0_ppr.append(make_element("w:ind", {"w:start": "1440", "w:end": "720"}))

    # P1: bidi only
    p1 = doc.add_paragraph("Bidi only + asymmetric indent.")
    p1_ppr = p1._p.get_or_add_pPr()
    p1_ppr.append(make_element("w:bidi"))
    p1_ppr.append(make_element("w:ind", {"w:start": "1440", "w:end": "720"}))

    # P2: mirrorIndents only
    p2 = doc.add_paragraph("MirrorIndents only + asymmetric indent.")
    p2_ppr = p2._p.get_or_add_pPr()
    p2_ppr.append(make_element("w:mirrorIndents"))
    p2_ppr.append(make_element("w:ind", {"w:start": "1440", "w:end": "720"}))

    save_fixture("bidi-mirror-indents", doc, {
        "name": "bidi-mirror-indents",
        "spec_ref": "ISO 29500-1 §17.3.1.6 + §17.3.1.18",
        "description": (
            "P0: bidi + mirrorIndents + start=1440, end=720. "
            "P1: bidi only + start=1440, end=720. "
            "P2: mirrorIndents only + start=1440, end=720."
        ),
        "expected_behavior": (
            "P0: both bidi and mirrorIndents should be in the model. "
            "P1: bidi only. P2: mirrorIndents only."
        ),
    })


# =========================================================================
# Multiple boolean properties combined
# =========================================================================

def make_combined_booleans() -> None:
    """Multiple boolean paragraph properties combined on one paragraph.

    P0: suppressAutoHyphens + snapToGrid=false + overflowPunct=false + adjustRightInd=false + wordWrap=false
    P1: all defaults (control)
    """
    doc = Document()

    # P0: all non-default boolean properties
    p0 = doc.add_paragraph("All boolean properties set to non-default values.")
    p0_ppr = p0._p.get_or_add_pPr()
    p0_ppr.append(make_element("w:suppressAutoHyphens"))
    p0_ppr.append(make_element("w:snapToGrid", {"w:val": "0"}))
    p0_ppr.append(make_element("w:overflowPunct", {"w:val": "0"}))
    p0_ppr.append(make_element("w:adjustRightInd", {"w:val": "0"}))
    p0_ppr.append(make_element("w:wordWrap", {"w:val": "0"}))

    # P1: all defaults
    doc.add_paragraph("All defaults (control).")

    save_fixture("combined-booleans", doc, {
        "name": "combined-booleans",
        "spec_ref": "ISO 29500-1 §17.3.1",
        "description": (
            "P0: suppressAutoHyphens=true, snapToGrid=false, overflowPunct=false, "
            "adjustRightInd=false, wordWrap=false. "
            "P1: all defaults (control)."
        ),
        "expected_behavior": (
            "P0: all five boolean properties should be reflected in the model. "
            "P1: no special properties."
        ),
    })


def main() -> None:
    print("Generating paragraph audit fixtures:")
    make_bidi()
    make_text_alignment()
    make_suppress_auto_hyphens()
    make_snap_to_grid()
    make_overflow_punct()
    make_adjust_right_ind()
    make_word_wrap()
    make_frame_pr()
    make_bidi_mirror_indents()
    make_combined_booleans()
    print("\nDone.")


if __name__ == "__main__":
    main()

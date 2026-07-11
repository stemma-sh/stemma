# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for run serialization spec-compliance tests.

Each fixture exercises a specific serialization concern:
  - Theme font references (asciiTheme/hAnsiTheme/eastAsiaTheme) in rFonts
  - Explicit font names in rFonts
  - rPr element ordering (bdr position, full ordering)

Tests need before/after pairs with actual text changes so the diff engine
produces a non-empty diff, forcing the code through serialize_canonical_docx
(which calls build_rpr). An identity diff takes an early-return path that
preserves the original XML bytes and bypasses the serializer.

Run:  uv run create_docs.py
"""

import json
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from lxml import etree

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ── helpers ───────────────────────────────────────────────────────────────

def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    """Create an OxmlElement with optional attributes."""
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_pair(
    name: str,
    before: DocxDocument,
    after: DocxDocument,
    metadata: dict,
) -> None:
    """Save a before/after DOCX pair."""
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    before.save(str(out / "before.docx"))
    after.save(str(out / "after.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  {name}/")


def save_fixture(
    name: str,
    doc: DocxDocument,
    metadata: dict,
    filename: str = "input.docx",
) -> None:
    """Save a single-doc fixture."""
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / filename))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  {name}/")


def _add_formatted_run(paragraph, text: str, attribs: list[tuple[str, dict | None]]) -> None:
    """Add a run with given text and rPr child elements."""
    run = paragraph.add_run(text)
    rPr = run._r.get_or_add_rPr()
    for tag, attrs in attribs:
        rPr.append(make_element(tag, attrs))


# ---------------------------------------------------------------------------
# 1. theme-font-ascii — before/after pair with asciiTheme on surviving text
# ---------------------------------------------------------------------------

def make_theme_font_ascii() -> None:
    """Before/after pair where a run with asciiTheme="minorHAnsi" survives.

    ISO 29500-1 §17.3.2.26: rFonts can reference theme fonts via
    asciiTheme/hAnsiTheme attributes. After diff_and_redline, the surviving
    run should preserve its theme font reference.

    The text change is in a DIFFERENT paragraph so the theme-font paragraph
    goes through the serializer without being modified.
    """
    before = Document()
    p1 = before.add_paragraph()
    run = p1.add_run("Theme font text")
    rPr = run._r.get_or_add_rPr()
    fonts = make_element("w:rFonts", {
        "w:asciiTheme": "minorHAnsi",
        "w:hAnsiTheme": "minorHAnsi",
    })
    rPr.append(fonts)
    before.add_paragraph("This text will change.")

    after = Document()
    p1 = after.add_paragraph()
    run = p1.add_run("Theme font text")
    rPr = run._r.get_or_add_rPr()
    fonts = make_element("w:rFonts", {
        "w:asciiTheme": "minorHAnsi",
        "w:hAnsiTheme": "minorHAnsi",
    })
    rPr.append(fonts)
    after.add_paragraph("This text has been modified.")

    save_pair("theme-font-ascii", before, after, {
        "name": "theme-font-ascii",
        "spec_ref": "ISO 29500-1 §17.3.2.26",
        "description": "Diff pair with asciiTheme/hAnsiTheme='minorHAnsi' on surviving run, text change in second paragraph",
        "expected_behavior": "asciiTheme/hAnsiTheme attributes should survive roundtrip on the re-serialized run",
        "current_status": "BUG — theme references resolved to font names on import, lost on export",
    })


# ---------------------------------------------------------------------------
# 2. theme-font-eastasia — before/after pair with eastAsiaTheme
# ---------------------------------------------------------------------------

def make_theme_font_eastasia() -> None:
    """Before/after pair where a run with eastAsiaTheme="minorEastAsia" survives.

    ISO 29500-1 §17.3.2.26: eastAsiaTheme references should survive roundtrip.
    """
    before = Document()
    p1 = before.add_paragraph()
    run = p1.add_run("East Asia theme font text")
    rPr = run._r.get_or_add_rPr()
    fonts = make_element("w:rFonts", {
        "w:eastAsiaTheme": "minorEastAsia",
    })
    rPr.append(fonts)
    before.add_paragraph("This paragraph will be edited.")

    after = Document()
    p1 = after.add_paragraph()
    run = p1.add_run("East Asia theme font text")
    rPr = run._r.get_or_add_rPr()
    fonts = make_element("w:rFonts", {
        "w:eastAsiaTheme": "minorEastAsia",
    })
    rPr.append(fonts)
    after.add_paragraph("This paragraph was edited.")

    save_pair("theme-font-eastasia", before, after, {
        "name": "theme-font-eastasia",
        "spec_ref": "ISO 29500-1 §17.3.2.26",
        "description": "Diff pair with eastAsiaTheme='minorEastAsia', text change in second paragraph",
        "expected_behavior": "eastAsiaTheme attribute should survive roundtrip",
        "current_status": "BUG — theme references resolved to font names on import, lost on export",
    })


# ---------------------------------------------------------------------------
# 3. explicit-font — before/after pair with explicit ascii="Courier New"
# ---------------------------------------------------------------------------

def make_explicit_font() -> None:
    """Before/after pair where a run with explicit w:ascii="Courier New" survives.

    ISO 29500-1 §17.3.2.26: Explicit font names should survive roundtrip.
    """
    before = Document()
    p1 = before.add_paragraph()
    run = p1.add_run("Explicit font text")
    rPr = run._r.get_or_add_rPr()
    fonts = make_element("w:rFonts", {
        "w:ascii": "Courier New",
        "w:hAnsi": "Courier New",
    })
    rPr.append(fonts)
    before.add_paragraph("Original second paragraph.")

    after = Document()
    p1 = after.add_paragraph()
    run = p1.add_run("Explicit font text")
    rPr = run._r.get_or_add_rPr()
    fonts = make_element("w:rFonts", {
        "w:ascii": "Courier New",
        "w:hAnsi": "Courier New",
    })
    rPr.append(fonts)
    after.add_paragraph("Changed second paragraph.")

    save_pair("explicit-font", before, after, {
        "name": "explicit-font",
        "spec_ref": "ISO 29500-1 §17.3.2.26",
        "description": "Diff pair with explicit ascii/hAnsi='Courier New', text change in second paragraph",
        "expected_behavior": "Explicit font name should survive roundtrip unchanged",
        "current_status": "OK — explicit fonts are preserved correctly",
    })


# ---------------------------------------------------------------------------
# 4. rpr-bdr-ordering — before/after pair with border + other rPr properties
# ---------------------------------------------------------------------------

def make_rpr_bdr_ordering() -> None:
    """Before/after pair where a run with bdr + bold + sz + underline + vertAlign survives.

    ECMA-376 Annex A CT_RPr: bdr is at position 28, between u (27) and
    vertAlign (31). The serializer must emit bdr at the correct position.
    """
    before = Document()
    p1 = before.add_paragraph()
    _add_formatted_run(p1, "Bordered text", [
        ("w:b", None),
        ("w:sz", {"w:val": "28"}),
        ("w:u", {"w:val": "single"}),
        ("w:bdr", {"w:val": "single", "w:sz": "4", "w:space": "1", "w:color": "FF0000"}),
        ("w:vertAlign", {"w:val": "superscript"}),
    ])
    before.add_paragraph("Text before edit.")

    after = Document()
    p1 = after.add_paragraph()
    _add_formatted_run(p1, "Bordered text", [
        ("w:b", None),
        ("w:sz", {"w:val": "28"}),
        ("w:u", {"w:val": "single"}),
        ("w:bdr", {"w:val": "single", "w:sz": "4", "w:space": "1", "w:color": "FF0000"}),
        ("w:vertAlign", {"w:val": "superscript"}),
    ])
    after.add_paragraph("Text after edit.")

    save_pair("rpr-bdr-ordering", before, after, {
        "name": "rpr-bdr-ordering",
        "spec_ref": "ECMA-376 Annex A CT_RPr",
        "description": "Diff pair with bold/sz/underline/bdr/vertAlign run, text change in second paragraph",
        "expected_behavior": "bdr should appear at position 28 (after u, before vertAlign) in re-serialized rPr",
        "current_status": "BUG — bdr emitted after cs (position 33) instead of position 28",
    })


# ---------------------------------------------------------------------------
# 5. rpr-full-ordering — before/after pair with many rPr properties
# ---------------------------------------------------------------------------

def make_rpr_full_ordering() -> None:
    """Before/after pair where a run with many rPr properties survives.

    ECMA-376 Annex A CT_RPr ordering — verify all elements in correct order.
    """
    before = Document()
    p1 = before.add_paragraph()
    _add_formatted_run(p1, "Fully formatted text", [
        ("w:rFonts", {"w:ascii": "Arial", "w:hAnsi": "Arial"}),
        ("w:b", None),
        ("w:i", None),
        ("w:color", {"w:val": "0000FF"}),
        ("w:sz", {"w:val": "24"}),
        ("w:szCs", {"w:val": "24"}),
        ("w:highlight", {"w:val": "yellow"}),
        ("w:u", {"w:val": "single"}),
        ("w:vertAlign", {"w:val": "superscript"}),
        ("w:rtl", None),
        ("w:cs", None),
        ("w:lang", {"w:val": "en-US"}),
    ])
    before.add_paragraph("Unchanged paragraph for anchor.")
    before.add_paragraph("This paragraph has original content.")

    after = Document()
    p1 = after.add_paragraph()
    _add_formatted_run(p1, "Fully formatted text", [
        ("w:rFonts", {"w:ascii": "Arial", "w:hAnsi": "Arial"}),
        ("w:b", None),
        ("w:i", None),
        ("w:color", {"w:val": "0000FF"}),
        ("w:sz", {"w:val": "24"}),
        ("w:szCs", {"w:val": "24"}),
        ("w:highlight", {"w:val": "yellow"}),
        ("w:u", {"w:val": "single"}),
        ("w:vertAlign", {"w:val": "superscript"}),
        ("w:rtl", None),
        ("w:cs", None),
        ("w:lang", {"w:val": "en-US"}),
    ])
    after.add_paragraph("Unchanged paragraph for anchor.")
    after.add_paragraph("This paragraph has updated content.")

    save_pair("rpr-full-ordering", before, after, {
        "name": "rpr-full-ordering",
        "spec_ref": "ECMA-376 Annex A CT_RPr",
        "description": "Diff pair with many rPr properties on surviving run, text change in third paragraph",
        "expected_behavior": "All rPr child elements in ECMA-376 Annex A order after re-serialization",
        "current_status": "BUG — bdr misordered (emitted after cs instead of position 28)",
    })


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    print("Generating run-serialization test fixtures:")
    make_theme_font_ascii()
    make_theme_font_eastasia()
    make_explicit_font()
    make_rpr_bdr_ordering()
    make_rpr_full_ordering()
    print("Done.")


if __name__ == "__main__":
    main()

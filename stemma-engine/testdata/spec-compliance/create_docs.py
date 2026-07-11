# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for spec-compliance testing against ISO 29500-1.

Each fixture exercises a specific OOXML construct from the spec that our
implementation may or may not handle. The spec-compliance tests in `stemma-engine/tests/`
import these fixtures and assert the parsed domain model correctly represents
the spec-defined semantics.

Fixtures are organized by spec area:
  - tracked-changes/  (ISO 29500-1 §17.13)
  - numbering/        (ISO 29500-1 §17.9)
  - formatting/       (ISO 29500-1 §17.2, §17.3)
  - tables/           (ISO 29500-1 §17.4)

Run:  uv run create_docs.py
"""

import json
import copy
from pathlib import Path
from lxml import etree

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, Twips, RGBColor, Emu

ROOT = Path(__file__).parent

# ── XML namespace helpers ────────────────────────────────────────────────

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def w(tag: str) -> str:
    """Return a fully-qualified wordprocessingml tag."""
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    """Create an OxmlElement with optional attributes."""
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


# ── Save helpers ─────────────────────────────────────────────────────────

def save_fixture(
    area: str,
    name: str,
    doc: DocxDocument,
    metadata: dict,
    filename: str = "input.docx",
) -> None:
    """Save a single-doc fixture (for parsing/model tests)."""
    out = ROOT / area / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / filename))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  {area}/{name}/")


def save_pair(
    area: str,
    name: str,
    before: DocxDocument,
    after: DocxDocument,
    metadata: dict,
) -> None:
    """Save a before/after pair (for diff/redline tests)."""
    out = ROOT / area / name
    out.mkdir(parents=True, exist_ok=True)
    before.save(str(out / "before.docx"))
    after.save(str(out / "after.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  {area}/{name}/")


# =========================================================================
# TRACKED CHANGES (ISO 29500-1 §17.13)
# =========================================================================

def make_tracked_changes_fixtures() -> None:
    print("\n── Tracked Changes ──")
    make_ppr_change()
    make_sect_pr_change()
    make_tbl_pr_change()
    make_cell_ins_del()
    make_preexisting_ins_del_roundtrip()
    make_rpr_change()
    make_row_ins_del()
    make_tracked_moves()
    make_nested_tracked_changes()
    make_para_mark_del()
    make_block_level_ins()
    make_block_level_del()
    make_multi_author_tracked()
    make_rpr_change_multi_props()


def make_ppr_change() -> None:
    """w:pPrChange — tracked paragraph property change (alignment left→center).

    ISO 29500-1 §17.13.5.29: pPrChange stores previous paragraph properties
    with author/date. Our implementation does not parse this.
    """
    doc = Document()
    p = doc.add_paragraph("This paragraph had its alignment changed from left to center.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Inject w:pPrChange into pPr
    pPr = p._p.get_or_add_pPr()
    ppr_change = make_element("w:pPrChange", {
        "w:id": "100",
        "w:author": "Spec Test",
        "w:date": "2025-01-01T00:00:00Z",
    })
    prev_ppr = make_element("w:pPr")
    prev_jc = make_element("w:jc", {"w:val": "left"})
    prev_ppr.append(prev_jc)
    ppr_change.append(prev_ppr)
    pPr.append(ppr_change)

    save_fixture("tracked-changes", "ppr-change", doc, {
        "name": "ppr-change",
        "spec_ref": "ISO 29500-1 §17.13.5.29",
        "description": "Paragraph with tracked alignment change (left → center) via w:pPrChange",
        "expected_behavior": "Parser should expose the pPrChange with previous alignment=left, current=center",
        "current_status": "NOT_PARSED — pPrChange is silently ignored",
    })


def make_sect_pr_change() -> None:
    """w:sectPrChange — tracked section property change.

    ISO 29500-1 §17.13.5.32: sectPrChange stores previous section properties.
    """
    doc = Document()
    doc.add_paragraph("Content before section change.")

    # Access the section and inject sectPrChange
    section = doc.sections[0]
    sect_pr = section._sectPr

    sect_pr_change = make_element("w:sectPrChange", {
        "w:id": "200",
        "w:author": "Spec Test",
        "w:date": "2025-01-01T00:00:00Z",
    })
    prev_sect_pr = make_element("w:sectPr")
    # Previous page size was different
    prev_pg_sz = make_element("w:pgSz", {
        "w:w": "12240",  # letter width
        "w:h": "15840",  # letter height
    })
    prev_sect_pr.append(prev_pg_sz)
    sect_pr_change.append(prev_sect_pr)
    sect_pr.append(sect_pr_change)

    save_fixture("tracked-changes", "sect-pr-change", doc, {
        "name": "sect-pr-change",
        "spec_ref": "ISO 29500-1 §17.13.5.32",
        "description": "Section with tracked page size change via w:sectPrChange",
        "expected_behavior": "Parser should expose sectPrChange with previous page dimensions",
        "current_status": "NOT_PARSED — sectPr is preserved opaquely but changes not tracked",
    })


def make_tbl_pr_change() -> None:
    """w:tblPrChange, w:trPrChange, w:tcPrChange — tracked table property changes.

    ISO 29500-1 §17.13.5.34/36/37: Store previous table/row/cell properties.
    """
    doc = Document()
    doc.add_paragraph("Table with tracked property changes.")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "A1"
    tbl.cell(0, 1).text = "B1"
    tbl.cell(1, 0).text = "A2"
    tbl.cell(1, 1).text = "B2"

    # Inject tblPrChange into tblPr
    tbl_element = tbl._tbl
    tblPr = tbl_element.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_element.insert(0, tblPr)

    tbl_pr_change = make_element("w:tblPrChange", {
        "w:id": "300",
        "w:author": "Spec Test",
        "w:date": "2025-01-01T00:00:00Z",
    })
    prev_tbl_pr = make_element("w:tblPr")
    # Previous table had different width
    prev_width = make_element("w:tblW", {"w:w": "5000", "w:type": "dxa"})
    prev_tbl_pr.append(prev_width)
    tbl_pr_change.append(prev_tbl_pr)
    tblPr.append(tbl_pr_change)

    # Inject trPrChange into first row
    tr = tbl_element.findall(w("tr"))[0]
    trPr = tr.find(w("trPr"))
    if trPr is None:
        trPr = make_element("w:trPr")
        tr.insert(0, trPr)

    tr_pr_change = make_element("w:trPrChange", {
        "w:id": "301",
        "w:author": "Spec Test",
        "w:date": "2025-01-01T00:00:00Z",
    })
    prev_tr_pr = make_element("w:trPr")
    prev_tr_height = make_element("w:trHeight", {"w:val": "400", "w:hRule": "exact"})
    prev_tr_pr.append(prev_tr_height)
    tr_pr_change.append(prev_tr_pr)
    trPr.append(tr_pr_change)

    # Inject tcPrChange into first cell
    tc = tr.findall(w("tc"))[0]
    tcPr = tc.find(w("tcPr"))
    if tcPr is None:
        tcPr = make_element("w:tcPr")
        tc.insert(0, tcPr)

    tc_pr_change = make_element("w:tcPrChange", {
        "w:id": "302",
        "w:author": "Spec Test",
        "w:date": "2025-01-01T00:00:00Z",
    })
    prev_tc_pr = make_element("w:tcPr")
    prev_tc_width = make_element("w:tcW", {"w:w": "2000", "w:type": "dxa"})
    prev_tc_pr.append(prev_tc_width)
    tc_pr_change.append(prev_tc_pr)
    tcPr.append(tc_pr_change)

    save_fixture("tracked-changes", "tbl-pr-change", doc, {
        "name": "tbl-pr-change",
        "spec_ref": "ISO 29500-1 §17.13.5.34/36/37",
        "description": "Table with tracked property changes on table, row, and cell levels",
        "expected_behavior": "Parser should expose tblPrChange/trPrChange/tcPrChange with previous values",
        "current_status": "NOT_PARSED — table property changes silently ignored",
    })


def make_cell_ins_del() -> None:
    """w:cellIns / w:cellDel — tracked cell insertion and deletion.

    ISO 29500-1 §17.13.5.1/2: cellIns/cellDel in tcPr mark cells as
    inserted/deleted as part of a tracked change.
    """
    doc = Document()
    doc.add_paragraph("Table with tracked cell insertion and deletion.")
    tbl = doc.add_table(rows=2, cols=3)
    tbl.cell(0, 0).text = "A1"
    tbl.cell(0, 1).text = "B1 (inserted)"
    tbl.cell(0, 2).text = "C1"
    tbl.cell(1, 0).text = "A2"
    tbl.cell(1, 1).text = "B2 (deleted)"
    tbl.cell(1, 2).text = "C2"

    tbl_element = tbl._tbl

    # Mark B1 cell as inserted
    row0 = tbl_element.findall(w("tr"))[0]
    cell_b1 = row0.findall(w("tc"))[1]
    tcPr_b1 = cell_b1.find(w("tcPr"))
    if tcPr_b1 is None:
        tcPr_b1 = make_element("w:tcPr")
        cell_b1.insert(0, tcPr_b1)
    cell_ins = make_element("w:cellIns", {
        "w:id": "400",
        "w:author": "Spec Test",
        "w:date": "2025-01-01T00:00:00Z",
    })
    tcPr_b1.append(cell_ins)

    # Mark B2 cell as deleted
    row1 = tbl_element.findall(w("tr"))[1]
    cell_b2 = row1.findall(w("tc"))[1]
    tcPr_b2 = cell_b2.find(w("tcPr"))
    if tcPr_b2 is None:
        tcPr_b2 = make_element("w:tcPr")
        cell_b2.insert(0, tcPr_b2)
    cell_del = make_element("w:cellDel", {
        "w:id": "401",
        "w:author": "Spec Test",
        "w:date": "2025-01-01T00:00:00Z",
    })
    tcPr_b2.append(cell_del)

    save_fixture("tracked-changes", "cell-ins-del", doc, {
        "name": "cell-ins-del",
        "spec_ref": "ISO 29500-1 §17.13.5.1/2",
        "description": "Table with tracked cell insertion (B1) and deletion (B2)",
        "expected_behavior": "Parser should expose cellIns/cellDel status on affected cells",
        "current_status": "NOT_PARSED — cellIns/cellDel silently ignored",
    })


def make_preexisting_ins_del_roundtrip() -> None:
    """Pre-existing w:ins/w:del with w:rPrChange — roundtrip fidelity.

    Verifies that a DOCX with native tracked changes (not computed by our
    diff engine) roundtrips correctly through import → export.
    """
    doc = Document()
    p = doc.add_paragraph()

    # "Hello " as normal text
    run_normal = p.add_run("Hello ")

    # "world" as an insertion with formatting change
    run_ins = OxmlElement("w:r")
    rPr = make_element("w:rPr")
    bold = make_element("w:b")
    rPr.append(bold)

    rpr_change = make_element("w:rPrChange", {
        "w:id": "500",
        "w:author": "Spec Test",
        "w:date": "2025-01-01T00:00:00Z",
    })
    prev_rpr = make_element("w:rPr")
    # Previously was not bold
    rpr_change.append(prev_rpr)
    rPr.append(rpr_change)
    run_ins.append(rPr)

    t = make_element("w:t")
    t.text = "world"
    run_ins.append(t)

    ins_wrapper = make_element("w:ins", {
        "w:id": "501",
        "w:author": "Spec Test",
        "w:date": "2025-01-01T00:00:00Z",
    })
    ins_wrapper.append(run_ins)
    p._p.append(ins_wrapper)

    # "old text" as a deletion
    run_del = OxmlElement("w:r")
    del_t = make_element("w:delText")
    del_t.text = "old text"
    del_t.set(qn("xml:space"), "preserve")
    run_del.append(del_t)

    del_wrapper = make_element("w:del", {
        "w:id": "502",
        "w:author": "Spec Test",
        "w:date": "2025-01-01T00:00:00Z",
    })
    del_wrapper.append(run_del)
    p._p.append(del_wrapper)

    save_fixture("tracked-changes", "preexisting-ins-del", doc, {
        "name": "preexisting-ins-del",
        "spec_ref": "ISO 29500-1 §17.13.5.14/18/31",
        "description": "Pre-existing tracked changes: w:ins (bold, with rPrChange) + w:del",
        "expected_behavior": "Import should parse all tracked change markup; export should preserve it",
        "current_status": "SUPPORTED — basic ins/del roundtrip works",
    })


def make_rpr_change() -> None:
    """w:rPrChange — tracked run formatting change (text made bold).

    ISO 29500-1 §17.13.5.30: rPrChange stores previous run properties
    with author/date. The run's current formatting is bold; rPrChange
    records that it previously was not bold.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Build a run with bold formatting and rPrChange
    run_el = OxmlElement("w:r")
    rPr = make_element("w:rPr")
    bold = make_element("w:b")
    rPr.append(bold)

    rpr_change = make_element("w:rPrChange", {
        "w:id": "600",
        "w:author": "Spec Test",
        "w:date": "2025-01-01T00:00:00Z",
    })
    prev_rpr = make_element("w:rPr")
    # Previously was not bold (empty rPr means no formatting)
    rpr_change.append(prev_rpr)
    rPr.append(rpr_change)
    run_el.append(rPr)

    t = make_element("w:t")
    t.text = "This text was made bold."
    run_el.append(t)
    p._p.append(run_el)

    save_fixture("tracked-changes", "rpr-change", doc, {
        "name": "rpr-change",
        "spec_ref": "ISO 29500-1 §17.13.5.30",
        "description": "Run with tracked formatting change (not bold → bold) via w:rPrChange",
        "expected_behavior": "TextNode.formatting_change should have author, date, and empty previous_marks",
    })


def make_row_ins_del() -> None:
    """w:ins / w:del on w:trPr — tracked row insertion/deletion.

    ISO 29500-1 §17.13.5.14/17.13.5.3: ins/del inside trPr marks entire
    rows as tracked insertions or deletions.
    """
    doc = Document()
    doc.add_paragraph("Table with tracked row insertion and deletion.")
    tbl = doc.add_table(rows=3, cols=2)
    tbl.cell(0, 0).text = "A1"
    tbl.cell(0, 1).text = "B1"
    tbl.cell(1, 0).text = "A2 (inserted row)"
    tbl.cell(1, 1).text = "B2 (inserted row)"
    tbl.cell(2, 0).text = "A3 (deleted row)"
    tbl.cell(2, 1).text = "B3 (deleted row)"

    tbl_element = tbl._tbl
    rows = tbl_element.findall(w("tr"))

    # Mark row 1 as inserted
    tr1 = rows[1]
    trPr1 = tr1.find(w("trPr"))
    if trPr1 is None:
        trPr1 = make_element("w:trPr")
        tr1.insert(0, trPr1)
    ins = make_element("w:ins", {
        "w:id": "700",
        "w:author": "Row Author",
        "w:date": "2025-06-15T10:30:00Z",
    })
    trPr1.append(ins)

    # Mark row 2 as deleted
    tr2 = rows[2]
    trPr2 = tr2.find(w("trPr"))
    if trPr2 is None:
        trPr2 = make_element("w:trPr")
        tr2.insert(0, trPr2)
    delete = make_element("w:del", {
        "w:id": "701",
        "w:author": "Row Author",
        "w:date": "2025-06-15T10:31:00Z",
    })
    trPr2.append(delete)

    save_fixture("tracked-changes", "row-ins-del", doc, {
        "name": "row-ins-del",
        "spec_ref": "ISO 29500-1 §17.13.5.14/17.13.5.3",
        "description": "Table with tracked row insertion (row 1) and deletion (row 2)",
        "expected_behavior": "TableRowNode.tracking_status is Inserted/Deleted with author/date",
    })


def make_tracked_moves() -> None:
    """w:moveTo / w:moveFrom — tracked move ranges.

    ISO 29500-1 §17.13.5.22/25: Content moved from one location to another
    is wrapped in moveFrom/moveTo with corresponding range markers.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Normal text
    p.add_run("Before move. ")

    # moveFrom range start marker
    move_from_start = make_element("w:moveFromRangeStart", {
        "w:id": "800",
        "w:author": "Move Author",
        "w:date": "2025-03-01T12:00:00Z",
        "w:name": "move1",
    })
    p._p.append(move_from_start)

    # moveFrom container with the moved text
    move_from = make_element("w:moveFrom", {
        "w:id": "801",
        "w:author": "Move Author",
        "w:date": "2025-03-01T12:00:00Z",
    })
    mf_run = OxmlElement("w:r")
    mf_t = make_element("w:t")
    mf_t.text = "moved text"
    mf_run.append(mf_t)
    move_from.append(mf_run)
    p._p.append(move_from)

    # moveFrom range end marker
    move_from_end = make_element("w:moveFromRangeEnd", {"w:id": "800"})
    p._p.append(move_from_end)

    # Normal text between
    p.add_run(" Middle. ")

    # moveTo range start marker
    move_to_start = make_element("w:moveToRangeStart", {
        "w:id": "802",
        "w:author": "Move Author",
        "w:date": "2025-03-01T12:00:00Z",
        "w:name": "move1",
    })
    p._p.append(move_to_start)

    # moveTo container
    move_to = make_element("w:moveTo", {
        "w:id": "803",
        "w:author": "Move Author",
        "w:date": "2025-03-01T12:00:00Z",
    })
    mt_run = OxmlElement("w:r")
    mt_t = make_element("w:t")
    mt_t.text = "moved text"
    mt_run.append(mt_t)
    move_to.append(mt_run)
    p._p.append(move_to)

    # moveTo range end marker
    move_to_end = make_element("w:moveToRangeEnd", {"w:id": "802"})
    p._p.append(move_to_end)

    p.add_run(" After move.")

    save_fixture("tracked-changes", "tracked-moves", doc, {
        "name": "tracked-moves",
        "spec_ref": "ISO 29500-1 §17.13.5.22/25",
        "description": "Paragraph with tracked move (moveFrom + moveTo with range markers)",
        "expected_behavior": "Move range markers preserved as decorations; moved text visible in atoms",
    })


def make_nested_tracked_changes() -> None:
    """Nested tracked changes: w:ins containing a run with w:rPrChange.

    Tests that both the insertion tracking and the formatting change
    are captured on the same text content.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Normal text first
    p.add_run("Normal text. ")

    # Create an insertion that also has a formatting change
    ins_wrapper = make_element("w:ins", {
        "w:id": "900",
        "w:author": "Insert Author",
        "w:date": "2025-04-01T08:00:00Z",
    })
    run_el = OxmlElement("w:r")
    rPr = make_element("w:rPr")
    bold = make_element("w:b")
    rPr.append(bold)
    italic = make_element("w:i")
    rPr.append(italic)

    # rPrChange: was previously just bold (not italic)
    rpr_change = make_element("w:rPrChange", {
        "w:id": "901",
        "w:author": "Format Author",
        "w:date": "2025-04-01T09:00:00Z",
    })
    prev_rpr = make_element("w:rPr")
    prev_bold = make_element("w:b")
    prev_rpr.append(prev_bold)
    rpr_change.append(prev_rpr)
    rPr.append(rpr_change)
    run_el.append(rPr)

    t = make_element("w:t")
    t.text = "inserted and reformatted"
    run_el.append(t)
    ins_wrapper.append(run_el)
    p._p.append(ins_wrapper)

    save_fixture("tracked-changes", "nested-tracked", doc, {
        "name": "nested-tracked",
        "spec_ref": "ISO 29500-1 §17.13.5.14/30",
        "description": "Paragraph with w:ins containing a run with w:rPrChange (nested tracked changes)",
        "expected_behavior": "Text is within ins wrapper; rPrChange is captured on TextNode.formatting_change",
    })


def make_para_mark_del() -> None:
    """w:del in w:rPr inside w:pPr — paragraph mark deletion tracking.

    ISO 29500-1 §17.13.5.28: When a paragraph mark is deleted (typically
    to merge paragraphs), w:del appears in the paragraph properties' run
    properties (w:pPr/w:rPr/w:del).
    """
    doc = Document()
    p = doc.add_paragraph("Paragraph with deleted paragraph mark.")

    # Inject w:del into w:pPr/w:rPr
    pPr = p._p.get_or_add_pPr()
    rPr = make_element("w:rPr")
    del_el = make_element("w:del", {
        "w:id": "1000",
        "w:author": "Para Author",
        "w:date": "2025-05-01T14:00:00Z",
    })
    rPr.append(del_el)
    pPr.append(rPr)

    save_fixture("tracked-changes", "para-mark-del", doc, {
        "name": "para-mark-del",
        "spec_ref": "ISO 29500-1 §17.13.5.28",
        "description": "Paragraph with deleted paragraph mark (w:del in w:pPr/w:rPr)",
        "expected_behavior": "ParagraphNode.para_mark_status should be Deleted with author/date",
    })


def make_block_level_ins() -> None:
    """Block-level w:ins — entire paragraph wrapped in tracked insertion.

    ISO 29500-1 §17.13.5.18: A w:ins element at body level wrapping
    an entire w:p means the whole paragraph was inserted.
    """
    doc = Document()
    doc.add_paragraph("Normal paragraph before insertion.")

    # Inject a block-level w:ins wrapping a full paragraph
    body = doc.element.body
    ins_wrapper = make_element("w:ins", {
        "w:id": "1100",
        "w:author": "Block Author",
        "w:date": "2025-07-01T16:00:00Z",
    })
    inner_p = OxmlElement("w:p")
    inner_r = OxmlElement("w:r")
    inner_t = make_element("w:t")
    inner_t.text = "This entire paragraph was inserted."
    inner_r.append(inner_t)
    inner_p.append(inner_r)
    ins_wrapper.append(inner_p)

    # Insert before the final sectPr
    sect_pr = body.find(w("sectPr"))
    if sect_pr is not None:
        sect_pr.addprevious(ins_wrapper)
    else:
        body.append(ins_wrapper)

    save_fixture("tracked-changes", "block-level-ins", doc, {
        "name": "block-level-ins",
        "spec_ref": "ISO 29500-1 §17.13.5.18",
        "description": "Entire paragraph wrapped in block-level w:ins",
        "expected_behavior": "TrackedBlock.status should be Inserted for the wrapped paragraph",
    })


def make_block_level_del() -> None:
    """Block-level w:del — entire paragraph wrapped in tracked deletion.

    ISO 29500-1 §17.13.5.4: A w:del element at body level wrapping
    an entire w:p means the whole paragraph was deleted.
    """
    doc = Document()
    doc.add_paragraph("Normal paragraph before deletion.")

    # Inject a block-level w:del wrapping a full paragraph
    body = doc.element.body
    del_wrapper = make_element("w:del", {
        "w:id": "1200",
        "w:author": "Delete Author",
        "w:date": "2025-08-15T10:30:00Z",
    })
    inner_p = OxmlElement("w:p")
    inner_r = OxmlElement("w:r")
    inner_t = make_element("w:t")
    inner_t.text = "This entire paragraph was deleted."
    inner_r.append(inner_t)
    inner_p.append(inner_r)
    del_wrapper.append(inner_p)

    # Insert before the final sectPr
    sect_pr = body.find(w("sectPr"))
    if sect_pr is not None:
        sect_pr.addprevious(del_wrapper)
    else:
        body.append(del_wrapper)

    save_fixture("tracked-changes", "block-level-del", doc, {
        "name": "block-level-del",
        "spec_ref": "ISO 29500-1 §17.13.5.4",
        "description": "Entire paragraph wrapped in block-level w:del",
        "expected_behavior": "TrackedBlock.status should be Deleted for the wrapped paragraph",
    })


def make_multi_author_tracked() -> None:
    """Multiple tracked change authors in one paragraph.

    Different changes by different authors: one insertion and one deletion
    in the same paragraph, by different people.
    """
    doc = Document()
    p = doc.add_paragraph()

    p.add_run("Shared text. ")

    # Insertion by Author A
    ins_a = make_element("w:ins", {
        "w:id": "1200",
        "w:author": "Author A",
        "w:date": "2025-08-01T10:00:00Z",
    })
    run_a = OxmlElement("w:r")
    t_a = make_element("w:t")
    t_a.text = "Added by A. "
    run_a.append(t_a)
    ins_a.append(run_a)
    p._p.append(ins_a)

    # Deletion by Author B
    del_b = make_element("w:del", {
        "w:id": "1201",
        "w:author": "Author B",
        "w:date": "2025-08-02T11:00:00Z",
    })
    run_b = OxmlElement("w:r")
    t_b = make_element("w:delText")
    t_b.text = "Removed by B."
    t_b.set(qn("xml:space"), "preserve")
    run_b.append(t_b)
    del_b.append(run_b)
    p._p.append(del_b)

    save_fixture("tracked-changes", "multi-author", doc, {
        "name": "multi-author",
        "spec_ref": "ISO 29500-1 §17.13.5.14/18",
        "description": "Paragraph with tracked changes by two different authors (ins by A, del by B)",
        "expected_behavior": "redline_extract finds both authors' changes with correct attribution",
    })


def make_rpr_change_multi_props() -> None:
    """w:rPrChange with multiple property changes — bold + color.

    ISO 29500-1 §17.13.5.30: rPrChange can record changes to multiple
    properties. Here the text was previously not bold and had color FF0000;
    now it is bold with no color.
    """
    doc = Document()
    p = doc.add_paragraph()

    run_el = OxmlElement("w:r")
    rPr = make_element("w:rPr")
    bold = make_element("w:b")
    rPr.append(bold)
    # Current: bold, no color

    rpr_change = make_element("w:rPrChange", {
        "w:id": "1300",
        "w:author": "Multi Prop Author",
        "w:date": "2025-09-01T15:00:00Z",
    })
    prev_rpr = make_element("w:rPr")
    # Previous: not bold, had red color
    prev_color = make_element("w:color", {"w:val": "FF0000"})
    prev_rpr.append(prev_color)
    rpr_change.append(prev_rpr)
    rPr.append(rpr_change)
    run_el.append(rPr)

    t = make_element("w:t")
    t.text = "Text with multiple formatting changes."
    run_el.append(t)
    p._p.append(run_el)

    save_fixture("tracked-changes", "rpr-change-multi", doc, {
        "name": "rpr-change-multi",
        "spec_ref": "ISO 29500-1 §17.13.5.30",
        "description": "Run with rPrChange tracking multiple property changes (no bold+red → bold+no color)",
        "expected_behavior": "FormattingChange captures both previous color and absence of bold",
    })


# =========================================================================
# NUMBERING (ISO 29500-1 §17.9)
# =========================================================================

def make_numbering_fixtures() -> None:
    print("\n── Numbering ──")
    make_lvl_override()
    make_is_lgl()
    make_lvl_restart()
    make_multi_level_numbering()


def _inject_numbering_xml(doc: DocxDocument, numbering_xml: str) -> None:
    """Replace the numbering.xml part with custom XML.

    This gives us full control over numbering definitions that python-docx
    doesn't expose directly.
    """
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    # Ensure numbering part exists
    numbering_part = None
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        pass

    if numbering_part is None:
        # Add a dummy list to force numbering part creation
        doc.add_paragraph("dummy", style="List Bullet")
        # Remove the dummy paragraph
        body = doc.element.body
        last_p = body.findall(w("p"))[-1]
        body.remove(last_p)
        numbering_part = doc.part.numbering_part

    # Replace the numbering XML
    numbering_part._element = etree.fromstring(numbering_xml.encode("utf-8"))


def make_lvl_override() -> None:
    """w:lvlOverride / w:startOverride — numbering restart.

    ISO 29500-1 §17.9.8/9: A w:num can contain w:lvlOverride children
    that override the abstract numbering definition. w:startOverride
    sets the counter to a specific value.

    This is critical for legal docs where numbered lists restart (e.g.,
    clause (a) restarting in a new section).
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}"
             xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerLetter"/>
      <w:lvlText w:val="(%1)"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
  <w:num w:numId="2">
    <w:abstractNumId w:val="0"/>
    <w:lvlOverride w:ilvl="0">
      <w:startOverride w:val="1"/>
    </w:lvlOverride>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # First list: items (a), (b), (c) using numId=1
    for text in ["First item in first list", "Second item in first list", "Third item in first list"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "0"}))
        numPr.append(make_element("w:numId", {"w:val": "1"}))
        pPr.append(numPr)

    # Regular paragraph break
    doc.add_paragraph("Section break — new clause starts here.")

    # Second list: should restart at (a) via numId=2 with lvlOverride
    for text in ["First item in second list (should be (a))", "Second item in second list (should be (b))"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "0"}))
        numPr.append(make_element("w:numId", {"w:val": "2"}))
        pPr.append(numPr)

    save_fixture("numbering", "lvl-override", doc, {
        "name": "lvl-override",
        "spec_ref": "ISO 29500-1 §17.9.8/9",
        "description": "Two numbered lists sharing an abstract definition; second list restarts via lvlOverride/startOverride",
        "expected_behavior": "First list: (a),(b),(c). Second list restarts: (a),(b). The startOverride resets the counter.",
        "current_status": "NOT_PARSED — lvlOverride/startOverride silently ignored; second list continues as (d),(e)",
    })


def make_is_lgl() -> None:
    """w:isLgl — legal numbering style.

    ISO 29500-1 §17.9.5: When isLgl is true on a level, all inherited
    level values are displayed as decimal regardless of their numFmt.
    E.g., level text "%1.%2.%3" with formats (upperRoman, lowerLetter, decimal)
    displays as "1.1.1" instead of "I.a.1".
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}"
             xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="upperRoman"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="1">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerLetter"/>
      <w:lvlText w:val="%1.%2."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="2">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:isLgl/>
      <w:lvlText w:val="%1.%2.%3"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # Level 0 item
    p0 = doc.add_paragraph("Top-level item")
    pPr0 = p0._p.get_or_add_pPr()
    numPr0 = make_element("w:numPr")
    numPr0.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr0.append(make_element("w:numId", {"w:val": "1"}))
    pPr0.append(numPr0)

    # Level 1 item
    p1 = doc.add_paragraph("Sub-item")
    pPr1 = p1._p.get_or_add_pPr()
    numPr1 = make_element("w:numPr")
    numPr1.append(make_element("w:ilvl", {"w:val": "1"}))
    numPr1.append(make_element("w:numId", {"w:val": "1"}))
    pPr1.append(numPr1)

    # Level 2 item (isLgl) — should display "1.1.1" not "I.a.1"
    p2 = doc.add_paragraph("Legal-numbered sub-sub-item (should be 1.1.1)")
    pPr2 = p2._p.get_or_add_pPr()
    numPr2 = make_element("w:numPr")
    numPr2.append(make_element("w:ilvl", {"w:val": "2"}))
    numPr2.append(make_element("w:numId", {"w:val": "1"}))
    pPr2.append(numPr2)

    save_fixture("numbering", "is-lgl", doc, {
        "name": "is-lgl",
        "spec_ref": "ISO 29500-1 §17.9.5",
        "description": "Multi-level list with isLgl on level 2 — inherited levels should render as decimal",
        "expected_behavior": "Level 0: 'I.', Level 1: 'I.a.', Level 2: '1.1.1' (not 'I.a.1')",
        "current_status": "NOT_PARSED — isLgl ignored; level 2 renders as 'I.a.1'",
    })


def make_lvl_restart() -> None:
    """w:lvlRestart — level restart control.

    ISO 29500-1 §17.9.10: lvlRestart controls when a numbering level
    resets its counter. val="0" means never restart. Otherwise, the value
    specifies which higher level triggers a restart.
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}"
             xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="1">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerLetter"/>
      <w:lvlText w:val="%2)"/>
      <w:lvlJc w:val="left"/>
      <w:lvlRestart w:val="0"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # Level 0: "1."
    p = doc.add_paragraph("First top-level item")
    pPr = p._p.get_or_add_pPr()
    numPr = make_element("w:numPr")
    numPr.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr.append(make_element("w:numId", {"w:val": "1"}))
    pPr.append(numPr)

    # Level 1: "a)" — starts counter
    for text in ["Sub item a", "Sub item b"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "1"}))
        numPr.append(make_element("w:numId", {"w:val": "1"}))
        pPr.append(numPr)

    # Level 0 again: "2."
    p = doc.add_paragraph("Second top-level item")
    pPr = p._p.get_or_add_pPr()
    numPr = make_element("w:numPr")
    numPr.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr.append(make_element("w:numId", {"w:val": "1"}))
    pPr.append(numPr)

    # Level 1 again: should NOT restart because lvlRestart=0
    for text in ["Sub item c (should be c, not a)", "Sub item d (should be d, not b)"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "1"}))
        numPr.append(make_element("w:numId", {"w:val": "1"}))
        pPr.append(numPr)

    save_fixture("numbering", "lvl-restart", doc, {
        "name": "lvl-restart",
        "spec_ref": "ISO 29500-1 §17.9.10",
        "description": "Two-level list where level 1 has lvlRestart=0 (never restart)",
        "expected_behavior": "Level 1 counter continues across level 0 items: a), b), then c), d) — NOT restarting to a), b)",
        "current_status": "NOT_PARSED — lvlRestart ignored; counter incorrectly resets to a), b) after each level 0 item",
    })


def make_multi_level_numbering() -> None:
    """Multi-level numbering with %N cross-references and various formats.

    Tests that less common number formats are handled.
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}"
             xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="upperRoman"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="1">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1.%2"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="2">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerLetter"/>
      <w:lvlText w:val="%1.%2.%3)"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    items = [
        (0, "Article One"),
        (1, "Section one-one"),
        (2, "Clause a"),
        (2, "Clause b"),
        (1, "Section one-two"),
        (2, "Clause a of section two"),
        (0, "Article Two"),
        (1, "Section two-one"),
    ]

    for level, text in items:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": str(level)}))
        numPr.append(make_element("w:numId", {"w:val": "1"}))
        pPr.append(numPr)

    save_fixture("numbering", "multi-level", doc, {
        "name": "multi-level",
        "spec_ref": "ISO 29500-1 §17.9.11",
        "description": "Three-level outline numbering with cross-level references in lvlText",
        "expected_behavior": "I. / I.1 / I.1.a) / I.1.b) / I.2 / I.2.a) / II. / II.1",
        "current_status": "SUPPORTED — basic multi-level numbering works",
    })


# =========================================================================
# PARAGRAPH & RUN FORMATTING (ISO 29500-1 §17.2, §17.3)
# =========================================================================

def make_formatting_fixtures() -> None:
    print("\n── Formatting ──")
    make_paragraph_spacing()
    make_paragraph_borders()
    make_vanish()
    make_paragraph_shading()
    make_toggle_xor()
    make_sect_pr_in_ppr()
    make_run_spacing()


def make_paragraph_spacing() -> None:
    """w:spacing — paragraph spacing (before/after/line).

    ISO 29500-1 §17.3.1.33: spacing element in pPr controls vertical
    spacing before and after paragraphs, plus line spacing.
    """
    doc = Document()

    p1 = doc.add_paragraph("Paragraph with 24pt spacing before.")
    pPr1 = p1._p.get_or_add_pPr()
    spacing1 = make_element("w:spacing", {
        "w:before": "480",  # 480 twips = 24pt
        "w:after": "0",
    })
    pPr1.append(spacing1)

    p2 = doc.add_paragraph("Paragraph with 12pt spacing after and 1.5 line spacing.")
    pPr2 = p2._p.get_or_add_pPr()
    spacing2 = make_element("w:spacing", {
        "w:before": "0",
        "w:after": "240",   # 240 twips = 12pt
        "w:line": "360",    # 360 = 1.5 lines (240 = single)
        "w:lineRule": "auto",
    })
    pPr2.append(spacing2)

    p3 = doc.add_paragraph("Paragraph with exact 18pt line spacing.")
    pPr3 = p3._p.get_or_add_pPr()
    spacing3 = make_element("w:spacing", {
        "w:line": "360",     # 360 twips = 18pt
        "w:lineRule": "exact",
    })
    pPr3.append(spacing3)

    save_fixture("formatting", "paragraph-spacing", doc, {
        "name": "paragraph-spacing",
        "spec_ref": "ISO 29500-1 §17.3.1.33",
        "description": "Paragraphs with various spacing configurations (before, after, line, lineRule)",
        "expected_behavior": "Parser should expose spacing.before, spacing.after, spacing.line, spacing.line_rule on ParagraphNode",
        "current_status": "NOT_PARSED — w:spacing silently ignored (P0 gap)",
    })


def make_paragraph_borders() -> None:
    """w:pBdr — paragraph borders.

    ISO 29500-1 §17.3.1.24: pBdr defines borders around a paragraph.
    Critical for legal docs (signature blocks, clause separators).
    """
    doc = Document()

    p1 = doc.add_paragraph("Paragraph with bottom border (signature line style).")
    pPr1 = p1._p.get_or_add_pPr()
    pBdr = make_element("w:pBdr")
    bottom = make_element("w:bottom", {
        "w:val": "single",
        "w:sz": "12",
        "w:space": "1",
        "w:color": "000000",
    })
    pBdr.append(bottom)
    pPr1.append(pBdr)

    p2 = doc.add_paragraph("Paragraph with box border (all sides).")
    pPr2 = p2._p.get_or_add_pPr()
    pBdr2 = make_element("w:pBdr")
    for side in ["top", "left", "bottom", "right"]:
        border = make_element(f"w:{side}", {
            "w:val": "single",
            "w:sz": "4",
            "w:space": "4",
            "w:color": "FF0000",
        })
        pBdr2.append(border)
    pPr2.append(pBdr2)

    save_fixture("formatting", "paragraph-borders", doc, {
        "name": "paragraph-borders",
        "spec_ref": "ISO 29500-1 §17.3.1.24",
        "description": "Paragraphs with bottom border and box border",
        "expected_behavior": "Parser should expose paragraph borders (at minimum, presence of border on each side)",
        "current_status": "NOT_PARSED — w:pBdr silently ignored (P0 gap)",
    })


def make_vanish() -> None:
    """w:vanish — hidden text.

    ISO 29500-1 §17.3.2.41: When true, text is hidden and should not be
    displayed (but may be printed with a setting). Important for legal
    documents that use hidden text for internal notes.
    """
    doc = Document()
    p = doc.add_paragraph()
    run_visible = p.add_run("Visible text. ")
    run_hidden = p.add_run("This text is hidden.")
    # Set vanish on the hidden run
    rPr = run_hidden._r.get_or_add_rPr()
    vanish = make_element("w:vanish")
    rPr.append(vanish)
    run_after = p.add_run(" More visible text.")

    save_fixture("formatting", "vanish", doc, {
        "name": "vanish",
        "spec_ref": "ISO 29500-1 §17.3.2.41",
        "description": "Paragraph with a hidden text run (w:vanish)",
        "expected_behavior": "Parser should flag the hidden run or exclude it from visible text projection",
        "current_status": "NOT_PARSED — w:vanish silently ignored; hidden text appears as normal text",
    })


def make_paragraph_shading() -> None:
    """w:shd — paragraph shading/background.

    ISO 29500-1 §17.3.1.31: shd in pPr sets the paragraph background.
    Used for highlighting clauses or callout sections.
    """
    doc = Document()

    p1 = doc.add_paragraph("Paragraph with yellow background shading.")
    pPr1 = p1._p.get_or_add_pPr()
    shd = make_element("w:shd", {
        "w:val": "clear",
        "w:color": "auto",
        "w:fill": "FFFF00",  # yellow
    })
    pPr1.append(shd)

    p2 = doc.add_paragraph("Paragraph with no shading (control).")

    save_fixture("formatting", "paragraph-shading", doc, {
        "name": "paragraph-shading",
        "spec_ref": "ISO 29500-1 §17.3.1.31",
        "description": "Paragraph with yellow background shading",
        "expected_behavior": "Parser should expose paragraph shading fill color",
        "current_status": "NOT_PARSED — w:shd in pPr silently ignored",
    })


def make_toggle_xor() -> None:
    """Toggle property XOR semantics — ISO 29500-1 §17.7.3.

    The spec requires toggle properties (bold, italic, caps, etc.) to use
    XOR when combining across style hierarchy levels. Our implementation
    uses simple override (last-wins).

    Test case: paragraph style sets bold=true, character style also sets bold=true.
    XOR result: bold=false. Override result: bold=true.
    """
    doc = Document()

    # Create a paragraph style with bold
    from docx.oxml.ns import qn as _qn
    styles_element = doc.styles.element

    # Add paragraph style "BoldPara"
    bold_para_style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "BoldPara"})
    name_el = make_element("w:name", {"w:val": "Bold Para"})
    bold_para_style.append(name_el)
    rPr_style = make_element("w:rPr")
    bold_style = make_element("w:b")
    rPr_style.append(bold_style)
    bold_para_style.append(rPr_style)
    styles_element.append(bold_para_style)

    # Add character style "BoldChar"
    bold_char_style = make_element("w:style", {"w:type": "character", "w:styleId": "BoldChar"})
    name_el2 = make_element("w:name", {"w:val": "Bold Char"})
    bold_char_style.append(name_el2)
    rPr_char = make_element("w:rPr")
    bold_char = make_element("w:b")
    rPr_char.append(bold_char)
    bold_char_style.append(rPr_char)
    styles_element.append(bold_char_style)

    # Create paragraph with BoldPara style
    p = doc.add_paragraph(style="BoldPara")
    run = p.add_run("This text has bold from paragraph style AND character style.")
    # Apply BoldChar character style
    rPr = run._r.get_or_add_rPr()
    r_style = make_element("w:rStyle", {"w:val": "BoldChar"})
    rPr.insert(0, r_style)

    # Control paragraph — just paragraph style bold, no character style
    p2 = doc.add_paragraph(style="BoldPara")
    p2.add_run("This text has bold only from paragraph style.")

    save_fixture("formatting", "toggle-xor", doc, {
        "name": "toggle-xor",
        "spec_ref": "ISO 29500-1 §17.7.3",
        "description": "Paragraph style bold=true + character style bold=true. XOR spec says bold=false.",
        "expected_behavior": "First paragraph's run should resolve to bold=false (XOR of two trues). Second paragraph's run should be bold=true.",
        "current_status": "NON_COMPLIANT — toggle properties use override (last-wins), not XOR. First paragraph resolves to bold=true (incorrect per spec).",
    })


def make_sect_pr_in_ppr() -> None:
    """sectPr in pPr — section properties embedded in paragraph.

    ISO 29500-1 §17.6.18: sectPr can appear as a child of pPr to mark
    section breaks. These must survive the redline pipeline.
    """
    doc = Document()

    doc.add_paragraph("Content in section one.")

    # Add a section break by inserting sectPr in the paragraph's pPr
    p_break = doc.add_paragraph("Last paragraph of section one.")
    pPr = p_break._p.get_or_add_pPr()
    sect_pr = make_element("w:sectPr")
    pg_sz = make_element("w:pgSz", {"w:w": "12240", "w:h": "15840"})
    sect_pr.append(pg_sz)
    pg_mar = make_element("w:pgMar", {
        "w:top": "1440", "w:right": "1440",
        "w:bottom": "1440", "w:left": "1440",
        "w:header": "720", "w:footer": "720",
        "w:gutter": "0",
    })
    sect_pr.append(pg_mar)
    pPr.append(sect_pr)

    doc.add_paragraph("Content in section two (different section).")

    save_fixture("formatting", "sect-pr-in-ppr", doc, {
        "name": "sect-pr-in-ppr",
        "spec_ref": "ISO 29500-1 §17.6.18",
        "description": "Section break via sectPr in pPr — must survive redline pipeline",
        "expected_behavior": "section_properties_raw on the paragraph should contain the sectPr XML blob",
        "current_status": "PARTIALLY_SUPPORTED — parsed but may be lost during redline (P0 gap)",
    })


def make_run_spacing() -> None:
    """w:spacing in rPr — character spacing.

    ISO 29500-1 §17.3.2.33: w:spacing in rPr controls inter-character
    spacing (condensed/expanded). Separate from w:spacing in pPr.
    """
    doc = Document()
    p = doc.add_paragraph()
    run_expanded = p.add_run("Expanded spacing text. ")
    rPr1 = run_expanded._r.get_or_add_rPr()
    spacing1 = make_element("w:spacing", {"w:val": "60"})  # 60 twips expanded
    rPr1.append(spacing1)

    run_condensed = p.add_run("Condensed spacing text.")
    rPr2 = run_condensed._r.get_or_add_rPr()
    spacing2 = make_element("w:spacing", {"w:val": "-20"})  # condensed
    rPr2.append(spacing2)

    save_fixture("formatting", "run-spacing", doc, {
        "name": "run-spacing",
        "spec_ref": "ISO 29500-1 §17.3.2.33",
        "description": "Runs with expanded and condensed character spacing",
        "expected_behavior": "Parser should expose character spacing value on TextNode/StyleProps",
        "current_status": "NOT_PARSED — character spacing silently ignored",
    })


# =========================================================================
# TABLES (ISO 29500-1 §17.4)
# =========================================================================

def make_table_fixtures() -> None:
    print("\n── Tables ──")
    make_tbl_cell_margins()
    make_tr_height()
    make_tbl_style()
    make_tbl_header()
    make_complex_merge()
    make_grid_before_after()


def make_tbl_cell_margins() -> None:
    """tblCellMar / tcMar — cell margins.

    ISO 29500-1 §17.4.42/43: Table-level default cell margins and
    per-cell margin overrides.
    """
    doc = Document()
    doc.add_paragraph("Table with custom cell margins.")

    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "Default margins"
    tbl.cell(0, 1).text = "Default margins"
    tbl.cell(1, 0).text = "Custom margins"
    tbl.cell(1, 1).text = "Default margins"

    # Set table-level cell margins
    tbl_element = tbl._tbl
    tblPr = tbl_element.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_element.insert(0, tblPr)

    tbl_cell_mar = make_element("w:tblCellMar")
    for side, val in [("top", "100"), ("left", "200"), ("bottom", "100"), ("right", "200")]:
        margin = make_element(f"w:{side}", {"w:w": val, "w:type": "dxa"})
        tbl_cell_mar.append(margin)
    tblPr.append(tbl_cell_mar)

    # Set per-cell margins on cell (1,0)
    row1 = tbl_element.findall(w("tr"))[1]
    cell_10 = row1.findall(w("tc"))[0]
    tcPr = cell_10.find(w("tcPr"))
    if tcPr is None:
        tcPr = make_element("w:tcPr")
        cell_10.insert(0, tcPr)

    tc_mar = make_element("w:tcMar")
    for side, val in [("top", "200"), ("left", "400"), ("bottom", "200"), ("right", "400")]:
        margin = make_element(f"w:{side}", {"w:w": val, "w:type": "dxa"})
        tc_mar.append(margin)
    tcPr.append(tc_mar)

    save_fixture("tables", "cell-margins", doc, {
        "name": "cell-margins",
        "spec_ref": "ISO 29500-1 §17.4.42/43",
        "description": "Table with global cell margins and per-cell margin override",
        "expected_behavior": "Parser should expose table-level and cell-level margin values",
        "current_status": "NOT_PARSED — cell margins silently ignored",
    })


def make_tr_height() -> None:
    """w:trHeight — row height with height rules.

    ISO 29500-1 §17.4.81: trHeight in trPr sets row height with rule
    (auto/atLeast/exact).
    """
    doc = Document()
    doc.add_paragraph("Table with specified row heights.")

    tbl = doc.add_table(rows=3, cols=2)
    tbl.cell(0, 0).text = "Auto height"
    tbl.cell(0, 1).text = "Row 1"
    tbl.cell(1, 0).text = "Exact 1 inch"
    tbl.cell(1, 1).text = "Row 2"
    tbl.cell(2, 0).text = "At least 0.5 inch"
    tbl.cell(2, 1).text = "Row 3"

    tbl_element = tbl._tbl
    rows = tbl_element.findall(w("tr"))

    # Row 0: auto (no explicit height)

    # Row 1: exact 1 inch = 1440 twips
    trPr1 = rows[1].find(w("trPr"))
    if trPr1 is None:
        trPr1 = make_element("w:trPr")
        rows[1].insert(0, trPr1)
    tr_height1 = make_element("w:trHeight", {"w:val": "1440", "w:hRule": "exact"})
    trPr1.append(tr_height1)

    # Row 2: atLeast 0.5 inch = 720 twips
    trPr2 = rows[2].find(w("trPr"))
    if trPr2 is None:
        trPr2 = make_element("w:trPr")
        rows[2].insert(0, trPr2)
    tr_height2 = make_element("w:trHeight", {"w:val": "720", "w:hRule": "atLeast"})
    trPr2.append(tr_height2)

    save_fixture("tables", "row-height", doc, {
        "name": "row-height",
        "spec_ref": "ISO 29500-1 §17.4.81",
        "description": "Table with auto, exact, and atLeast row heights",
        "expected_behavior": "Parser should expose trHeight value and hRule on rows",
        "current_status": "NOT_PARSED — trHeight silently ignored",
    })


def make_tbl_style() -> None:
    """w:tblStyle — table style reference.

    ISO 29500-1 §17.4.63: tblStyle in tblPr references a named style
    that provides default formatting for the table.
    """
    doc = Document()

    # Create a table style in styles.xml
    styles_element = doc.styles.element
    tbl_style = make_element("w:style", {"w:type": "table", "w:styleId": "LegalTable"})
    name_el = make_element("w:name", {"w:val": "Legal Table"})
    tbl_style.append(name_el)

    # Table formatting in the style
    tblPr_style = make_element("w:tblPr")
    tbl_borders = make_element("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = make_element(f"w:{side}", {
            "w:val": "single",
            "w:sz": "4",
            "w:space": "0",
            "w:color": "000000",
        })
        tbl_borders.append(border)
    tblPr_style.append(tbl_borders)
    tbl_style.append(tblPr_style)

    styles_element.append(tbl_style)

    # Create table referencing the style
    doc.add_paragraph("Table using a named table style.")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "A1"
    tbl.cell(0, 1).text = "B1"
    tbl.cell(1, 0).text = "A2"
    tbl.cell(1, 1).text = "B2"

    # Set tblStyle reference
    tbl_element = tbl._tbl
    tblPr = tbl_element.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_element.insert(0, tblPr)
    tbl_style_ref = make_element("w:tblStyle", {"w:val": "LegalTable"})
    tblPr.insert(0, tbl_style_ref)

    save_fixture("tables", "tbl-style", doc, {
        "name": "tbl-style",
        "spec_ref": "ISO 29500-1 §17.4.63",
        "description": "Table referencing a named table style for border formatting",
        "expected_behavior": "Parser should resolve tblStyle and apply table-level borders from the style",
        "current_status": "NOT_PARSED — tblStyle reference ignored; table uses no style-inherited formatting",
    })


def make_tbl_header() -> None:
    """w:tblHeader — header row repeat.

    ISO 29500-1 §17.4.50: tblHeader in trPr marks a row as a header row
    that should repeat at the top of each page when the table spans pages.
    """
    doc = Document()
    doc.add_paragraph("Table with header row repeat.")

    tbl = doc.add_table(rows=4, cols=2)
    tbl.cell(0, 0).text = "Header A"
    tbl.cell(0, 1).text = "Header B"
    for i in range(1, 4):
        tbl.cell(i, 0).text = f"Data A{i}"
        tbl.cell(i, 1).text = f"Data B{i}"

    # Mark first row as header
    tbl_element = tbl._tbl
    row0 = tbl_element.findall(w("tr"))[0]
    trPr0 = row0.find(w("trPr"))
    if trPr0 is None:
        trPr0 = make_element("w:trPr")
        row0.insert(0, trPr0)
    tbl_header = make_element("w:tblHeader")
    trPr0.append(tbl_header)

    save_fixture("tables", "tbl-header", doc, {
        "name": "tbl-header",
        "spec_ref": "ISO 29500-1 §17.4.50",
        "description": "Table with first row marked as repeating header",
        "expected_behavior": "Parser should expose the tblHeader flag on the header row",
        "current_status": "NOT_PARSED — tblHeader silently ignored",
    })


def make_complex_merge() -> None:
    """Complex cell merges — vMerge with non-rectangular patterns.

    Tests that the canonicalization correctly handles vertical merges
    spanning multiple rows with horizontal gridSpan overlaps.
    """
    doc = Document()
    doc.add_paragraph("Table with complex merge pattern.")

    # Create a 4x3 table with complex merges:
    # Row 0: [A (span 2 cols, vmerge start)] [C]
    # Row 1: [A (vmerge continue)]            [D]
    # Row 2: [E] [F]                          [G]
    tbl = doc.add_table(rows=3, cols=3)
    tbl_element = tbl._tbl

    # Row 0
    row0 = tbl_element.findall(w("tr"))[0]
    cells0 = row0.findall(w("tc"))
    # Cell A: gridSpan=2, vMerge=restart
    tcPr_a = cells0[0].find(w("tcPr"))
    if tcPr_a is None:
        tcPr_a = make_element("w:tcPr")
        cells0[0].insert(0, tcPr_a)
    tcPr_a.append(make_element("w:gridSpan", {"w:val": "2"}))
    tcPr_a.append(make_element("w:vMerge", {"w:val": "restart"}))
    # Set text
    for p_el in cells0[0].findall(w("p")):
        cells0[0].remove(p_el)
    new_p = make_element("w:p")
    new_r = make_element("w:r")
    new_t = make_element("w:t")
    new_t.text = "A (merged 2x2)"
    new_r.append(new_t)
    new_p.append(new_r)
    cells0[0].append(new_p)

    # Remove second cell (consumed by gridSpan)
    row0.remove(cells0[1])

    # Cell C: normal
    for p_el in cells0[2].findall(w("p")):
        cells0[2].remove(p_el)
    new_p = make_element("w:p")
    new_r = make_element("w:r")
    new_t = make_element("w:t")
    new_t.text = "C"
    new_r.append(new_t)
    new_p.append(new_r)
    cells0[2].append(new_p)

    # Row 1
    row1 = tbl_element.findall(w("tr"))[1]
    cells1 = row1.findall(w("tc"))
    # Cell A continue: gridSpan=2, vMerge=continue
    tcPr_a1 = cells1[0].find(w("tcPr"))
    if tcPr_a1 is None:
        tcPr_a1 = make_element("w:tcPr")
        cells1[0].insert(0, tcPr_a1)
    tcPr_a1.append(make_element("w:gridSpan", {"w:val": "2"}))
    tcPr_a1.append(make_element("w:vMerge"))  # continue (no val)
    # Remove text from continuation cell
    for p_el in cells1[0].findall(w("p")):
        cells1[0].remove(p_el)
    cells1[0].append(make_element("w:p"))

    row1.remove(cells1[1])

    # Cell D: normal
    for p_el in cells1[2].findall(w("p")):
        cells1[2].remove(p_el)
    new_p = make_element("w:p")
    new_r = make_element("w:r")
    new_t = make_element("w:t")
    new_t.text = "D"
    new_r.append(new_t)
    new_p.append(new_r)
    cells1[2].append(new_p)

    # Row 2: normal cells
    row2 = tbl_element.findall(w("tr"))[2]
    cells2 = row2.findall(w("tc"))
    for i, label in enumerate(["E", "F", "G"]):
        for p_el in cells2[i].findall(w("p")):
            cells2[i].remove(p_el)
        new_p = make_element("w:p")
        new_r = make_element("w:r")
        new_t = make_element("w:t")
        new_t.text = label
        new_r.append(new_t)
        new_p.append(new_r)
        cells2[i].append(new_p)

    save_fixture("tables", "complex-merge", doc, {
        "name": "complex-merge",
        "spec_ref": "ISO 29500-1 §17.4.17/84",
        "description": "3x3 table with 2x2 merged cell (gridSpan + vMerge) in top-left",
        "expected_behavior": "Canonical table should show correct owner_grid with A owning [0,0],[0,1],[1,0],[1,1]",
        "current_status": "SUPPORTED — basic vMerge/gridSpan canonicalization works; testing complex pattern",
    })


def make_grid_before_after() -> None:
    """gridBefore / gridAfter — staggered table layouts.

    ISO 29500-1 §17.4.14/3: gridBefore/gridAfter in trPr add empty grid
    units before/after the row's cells, creating staggered layouts.
    """
    doc = Document()
    doc.add_paragraph("Table with staggered rows (gridBefore/gridAfter).")

    # Create a table with 4 grid columns
    tbl = doc.add_table(rows=3, cols=4)
    tbl_element = tbl._tbl

    # Set up tblGrid with 4 columns
    tbl_grid = tbl_element.find(w("tblGrid"))

    rows = tbl_element.findall(w("tr"))

    # Row 0: normal, all 4 cells
    cells0 = rows[0].findall(w("tc"))
    for i, label in enumerate(["A1", "B1", "C1", "D1"]):
        for p_el in cells0[i].findall(w("p")):
            cells0[i].remove(p_el)
        new_p = make_element("w:p")
        new_r = make_element("w:r")
        new_t = make_element("w:t")
        new_t.text = label
        new_r.append(new_t)
        new_p.append(new_r)
        cells0[i].append(new_p)

    # Row 1: gridBefore=1, only 3 cells (offset by 1)
    trPr1 = rows[1].find(w("trPr"))
    if trPr1 is None:
        trPr1 = make_element("w:trPr")
        rows[1].insert(0, trPr1)
    trPr1.append(make_element("w:gridBefore", {"w:val": "1"}))
    wBefore = make_element("w:wBefore", {"w:w": "2000", "w:type": "dxa"})
    trPr1.append(wBefore)

    # Remove the 4th cell since we only need 3 with gridBefore=1
    cells1 = rows[1].findall(w("tc"))
    rows[1].remove(cells1[3])
    for i, label in enumerate(["B2", "C2", "D2"]):
        for p_el in cells1[i].findall(w("p")):
            cells1[i].remove(p_el)
        new_p = make_element("w:p")
        new_r = make_element("w:r")
        new_t = make_element("w:t")
        new_t.text = label
        new_r.append(new_t)
        new_p.append(new_r)
        cells1[i].append(new_p)

    # Row 2: gridAfter=1, only 3 cells (1 empty at end)
    trPr2 = rows[2].find(w("trPr"))
    if trPr2 is None:
        trPr2 = make_element("w:trPr")
        rows[2].insert(0, trPr2)
    trPr2.append(make_element("w:gridAfter", {"w:val": "1"}))
    wAfter = make_element("w:wAfter", {"w:w": "2000", "w:type": "dxa"})
    trPr2.append(wAfter)

    cells2 = rows[2].findall(w("tc"))
    rows[2].remove(cells2[3])
    for i, label in enumerate(["A3", "B3", "C3"]):
        for p_el in cells2[i].findall(w("p")):
            cells2[i].remove(p_el)
        new_p = make_element("w:p")
        new_r = make_element("w:r")
        new_t = make_element("w:t")
        new_t.text = label
        new_r.append(new_t)
        new_p.append(new_r)
        cells2[i].append(new_p)

    save_fixture("tables", "grid-before-after", doc, {
        "name": "grid-before-after",
        "spec_ref": "ISO 29500-1 §17.4.14/3",
        "description": "Table with staggered rows using gridBefore and gridAfter",
        "expected_behavior": "Canonical table should correctly position cells with grid offsets",
        "current_status": "PARTIALLY_SUPPORTED — gridBefore/gridAfter parsed but positioning may be incorrect",
    })


# =========================================================================
# RUN FORMATTING (ISO 29500-1 §17.3.2)
# =========================================================================

def make_run_formatting_fixtures() -> None:
    print("\n── Run Formatting ──")
    make_underline_styles()
    make_strikethrough_types()
    make_font_slots()
    make_text_effects()
    make_language_tags()
    make_highlight_colors()
    make_caps_smallcaps()


def make_underline_styles() -> None:
    """w:u — underline with various styles.

    ISO 29500-1 §17.3.2.40: The w:u element specifies an underline with
    a style attribute (single, double, dotted, dashed, wave, etc.).
    """
    doc = Document()
    p = doc.add_paragraph()

    styles = ["single", "double", "dotted", "dash", "wave"]
    for style in styles:
        run = p.add_run(f"{style}-underline ")
        rPr = run._r.get_or_add_rPr()
        u = make_element("w:u", {"w:val": style})
        rPr.append(u)

    save_fixture("run-formatting", "underline-styles", doc, {
        "name": "underline-styles",
        "spec_ref": "ISO 29500-1 §17.3.2.40",
        "description": "Paragraph with runs using different underline styles: single, double, dotted, dashed, wave",
        "expected_behavior": "Each run should carry Mark::Underline; style detail (single vs double etc.) should be distinguishable",
        "current_status": "PARTIALLY_SUPPORTED — underline presence detected, but style detail collapsed to binary On/Off",
    })


def make_strikethrough_types() -> None:
    """w:strike / w:dstrike — single and double strikethrough.

    ISO 29500-1 §17.3.2.37/6: strike for single, dstrike for double.
    """
    doc = Document()
    p = doc.add_paragraph()

    run_single = p.add_run("single-strike ")
    rPr1 = run_single._r.get_or_add_rPr()
    rPr1.append(make_element("w:strike"))

    run_double = p.add_run("double-strike")
    rPr2 = run_double._r.get_or_add_rPr()
    rPr2.append(make_element("w:dstrike"))

    save_fixture("run-formatting", "strikethrough-types", doc, {
        "name": "strikethrough-types",
        "spec_ref": "ISO 29500-1 §17.3.2.37/6",
        "description": "Paragraph with single strikethrough (w:strike) and double strikethrough (w:dstrike)",
        "expected_behavior": "Both runs should carry Mark::Strike; double-strike should be distinguishable from single",
        "current_status": "PARTIALLY_SUPPORTED — both map to Mark::Strike, type detail lost",
    })


def make_font_slots() -> None:
    """w:rFonts — font family with all four slots.

    ISO 29500-1 §17.3.2.26: rFonts specifies fonts for ascii, hAnsi,
    eastAsia, and cs (complex script) character ranges.
    """
    doc = Document()
    p = doc.add_paragraph()

    run = p.add_run("Multi-font-slot text")
    rPr = run._r.get_or_add_rPr()
    fonts = make_element("w:rFonts", {
        "w:ascii": "Arial",
        "w:hAnsi": "Arial",
        "w:eastAsia": "MS Mincho",
        "w:cs": "Arial",
    })
    rPr.append(fonts)

    save_fixture("run-formatting", "font-slots", doc, {
        "name": "font-slots",
        "spec_ref": "ISO 29500-1 §17.3.2.26",
        "description": "Run with all four rFonts slots: ascii=Arial, hAnsi=Arial, eastAsia=MS Mincho, cs=Arial",
        "expected_behavior": "All four font slots should be accessible; font_family should capture at least ascii/hAnsi",
        "current_status": "PARTIALLY_SUPPORTED — font_family captures ascii/hAnsi; eastAsia and cs slots not exposed",
    })


def make_text_effects() -> None:
    """w:emboss, w:imprint, w:outline, w:shadow — text effects.

    ISO 29500-1 §17.3.2.8/12/21/31: Various text effects as boolean
    run properties.
    """
    doc = Document()
    p = doc.add_paragraph()

    effects = [
        ("emboss-text ", "w:emboss"),
        ("imprint-text ", "w:imprint"),
        ("outline-text ", "w:outline"),
        ("shadow-text", "w:shadow"),
    ]
    for text, tag in effects:
        run = p.add_run(text)
        rPr = run._r.get_or_add_rPr()
        rPr.append(make_element(tag))

    save_fixture("run-formatting", "text-effects", doc, {
        "name": "text-effects",
        "spec_ref": "ISO 29500-1 §17.3.2.8/12/21/31",
        "description": "Runs with emboss, imprint, outline, and shadow text effects",
        "expected_behavior": "Each effect should be exposed as a mark or property on the text node",
        "current_status": "NOT_PARSED — text effects silently ignored",
    })


def make_language_tags() -> None:
    """w:lang — language tags on runs.

    ISO 29500-1 §17.3.2.20: The w:lang element specifies language for
    proofing tools via val (Latin), eastAsia, and bidi attributes.
    """
    doc = Document()
    p = doc.add_paragraph()

    run_en = p.add_run("English text ")
    rPr1 = run_en._r.get_or_add_rPr()
    rPr1.append(make_element("w:lang", {"w:val": "en-US"}))

    run_fr = p.add_run("Texte francais ")
    rPr2 = run_fr._r.get_or_add_rPr()
    rPr2.append(make_element("w:lang", {"w:val": "fr-FR"}))

    run_ja = p.add_run("Japanese context")
    rPr3 = run_ja._r.get_or_add_rPr()
    rPr3.append(make_element("w:lang", {"w:eastAsia": "ja-JP"}))

    save_fixture("run-formatting", "language-tags", doc, {
        "name": "language-tags",
        "spec_ref": "ISO 29500-1 §17.3.2.20",
        "description": "Runs with different language tags: en-US, fr-FR, ja-JP (eastAsia)",
        "expected_behavior": "Language tags should be accessible on the text node or style props",
        "current_status": "NOT_PARSED — w:lang silently dropped",
    })


def make_highlight_colors() -> None:
    """w:highlight — text highlight colors.

    ISO 29500-1 §17.3.2.13: Named highlight colors on runs.
    """
    doc = Document()
    p = doc.add_paragraph()

    colors = ["yellow", "green", "cyan", "magenta"]
    for color in colors:
        run = p.add_run(f"{color}-highlight ")
        rPr = run._r.get_or_add_rPr()
        rPr.append(make_element("w:highlight", {"w:val": color}))

    save_fixture("run-formatting", "highlight-colors", doc, {
        "name": "highlight-colors",
        "spec_ref": "ISO 29500-1 §17.3.2.13",
        "description": "Runs with different highlight colors: yellow, green, cyan, magenta",
        "expected_behavior": "style_props.highlight should capture the named color for each run",
        "current_status": "SUPPORTED — highlight colors parsed into StyleProps.highlight",
    })


def make_caps_smallcaps() -> None:
    """w:caps / w:smallCaps — capitalization properties.

    ISO 29500-1 §17.3.2.5/30: caps renders all text as uppercase,
    smallCaps renders lowercase as small capitals.
    """
    doc = Document()
    p = doc.add_paragraph()

    run_caps = p.add_run("all-caps text ")
    rPr1 = run_caps._r.get_or_add_rPr()
    rPr1.append(make_element("w:caps"))

    run_sc = p.add_run("small-caps text")
    rPr2 = run_sc._r.get_or_add_rPr()
    rPr2.append(make_element("w:smallCaps"))

    save_fixture("run-formatting", "caps-smallcaps", doc, {
        "name": "caps-smallcaps",
        "spec_ref": "ISO 29500-1 §17.3.2.5/30",
        "description": "Runs with w:caps and w:smallCaps",
        "expected_behavior": "First run should have Mark::Caps, second should have Mark::SmallCaps",
        "current_status": "SUPPORTED — caps and smallCaps parsed into marks",
    })


# =========================================================================
# TABS (ISO 29500-1 §17.3.1.38)
# =========================================================================

def make_tab_fixtures() -> None:
    print("\n── Tabs ──")
    make_direct_tabs()
    make_style_inherited_tabs()
    make_clear_tab_stop()
    make_default_tab_interval()
    make_ptab_widget()


def make_direct_tabs() -> None:
    """Explicit tab stops at known positions with different alignments.

    ISO 29500-1 §17.3.1.38: w:tabs/w:tab defines custom tab stop positions
    with alignment (left, center, right) and optional leaders.
    """
    doc = Document()

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()

    # Add w:tabs with three stops
    tabs_el = make_element("w:tabs")
    tabs_el.append(make_element("w:tab", {"w:val": "left", "w:pos": "1440"}))
    tabs_el.append(make_element("w:tab", {"w:val": "center", "w:pos": "4320"}))
    tabs_el.append(make_element("w:tab", {"w:val": "right", "w:pos": "7200", "w:leader": "dot"}))
    pPr.append(tabs_el)

    # Text with tab characters between words
    run = p.add_run("Left\tCenter\tRight")

    save_fixture("tabs", "direct-tabs", doc, {
        "name": "direct-tabs",
        "spec_ref": "ISO 29500-1 §17.3.1.38",
        "description": "Paragraph with explicit tab stops: left@1440, center@4320, right@7200 (dot leader)",
        "expected_behavior": "Parser should expose tab_stops with correct positions, alignments, and leaders",
        "current_status": "SUPPORTED — tab stops parsed into ParagraphNode.tab_stops",
    })


def make_style_inherited_tabs() -> None:
    """Paragraph style defines tab stops; paragraph inherits them.

    ISO 29500-1 §17.3.1.38 + §17.7: Tab stops defined in a paragraph style
    should be inherited by paragraphs using that style.
    """
    doc = Document()

    # Create a paragraph style with tab stops
    styles_element = doc.styles.element
    tab_style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "TabStyle"})
    name_el = make_element("w:name", {"w:val": "Tab Style"})
    tab_style.append(name_el)
    pPr_style = make_element("w:pPr")
    tabs_el = make_element("w:tabs")
    tabs_el.append(make_element("w:tab", {"w:val": "left", "w:pos": "720"}))
    tabs_el.append(make_element("w:tab", {"w:val": "center", "w:pos": "3600"}))
    pPr_style.append(tabs_el)
    tab_style.append(pPr_style)
    styles_element.append(tab_style)

    # Create a paragraph using that style
    p = doc.add_paragraph(style="TabStyle")
    p.add_run("Col1\tCol2")

    save_fixture("tabs", "style-inherited-tabs", doc, {
        "name": "style-inherited-tabs",
        "spec_ref": "ISO 29500-1 §17.3.1.38, §17.7",
        "description": "Paragraph style defines tab stops at 720 (left) and 3600 (center); paragraph uses that style",
        "expected_behavior": "Paragraph should inherit tab stops from its style",
        "current_status": "SUPPORTED — style tab inheritance is implemented",
    })


def make_clear_tab_stop() -> None:
    """Style defines tabs; paragraph clears one with val='clear'.

    ISO 29500-1 §17.3.1.38: A tab with val='clear' removes an inherited
    tab stop at the same position.
    """
    doc = Document()

    # Style with two tab stops
    styles_element = doc.styles.element
    clear_style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "ClearTabStyle"})
    name_el = make_element("w:name", {"w:val": "Clear Tab Style"})
    clear_style.append(name_el)
    pPr_style = make_element("w:pPr")
    tabs_el = make_element("w:tabs")
    tabs_el.append(make_element("w:tab", {"w:val": "left", "w:pos": "1440"}))
    tabs_el.append(make_element("w:tab", {"w:val": "right", "w:pos": "2880"}))
    pPr_style.append(tabs_el)
    clear_style.append(pPr_style)
    styles_element.append(clear_style)

    # Paragraph uses the style but clears the stop at 1440
    p = doc.add_paragraph(style="ClearTabStyle")
    pPr = p._p.get_or_add_pPr()
    direct_tabs = make_element("w:tabs")
    direct_tabs.append(make_element("w:tab", {"w:val": "clear", "w:pos": "1440"}))
    pPr.append(direct_tabs)

    p.add_run("Before\tAfter")

    save_fixture("tabs", "clear-tab-stop", doc, {
        "name": "clear-tab-stop",
        "spec_ref": "ISO 29500-1 §17.3.1.38",
        "description": "Style defines tabs at 1440 and 2880; paragraph clears the stop at 1440",
        "expected_behavior": "Effective tab stops should only include 2880 (1440 removed by clear)",
        "current_status": "SUPPORTED — clear semantics implemented in overlay_tab_stops",
    })


def make_default_tab_interval() -> None:
    """Custom w:defaultTabStop in settings.xml.

    ISO 29500-1 §17.15.1.25: defaultTabStop sets the interval for automatic
    tab stops. Default is 720 twips (0.5 inch).

    python-docx inserts its own defaultTabStop=720 on save, so we patch the
    saved ZIP afterwards to replace the value with 360.
    """
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("A\tB\tC\tD")

    out = ROOT / "tabs" / "default-tab-interval"
    out.mkdir(parents=True, exist_ok=True)
    docx_path = str(out / "input.docx")
    doc.save(docx_path)

    # Patch settings.xml inside the saved DOCX to set defaultTabStop=360.
    # python-docx always writes defaultTabStop=720; we do a targeted replacement.
    import zipfile, io
    buf = io.BytesIO(Path(docx_path).read_bytes())
    with zipfile.ZipFile(buf, "r") as zin:
        entries = {name: zin.read(name) for name in zin.namelist()}
    settings = entries["word/settings.xml"].decode("utf-8")
    settings = settings.replace(
        'w:val="720"',
        'w:val="360"',
        1,  # only replace the first occurrence (the defaultTabStop)
    )
    entries["word/settings.xml"] = settings.encode("utf-8")
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in entries.items():
            zout.writestr(name, data)

    metadata = {
        "name": "default-tab-interval",
        "spec_ref": "ISO 29500-1 §17.15.1.25",
        "description": "Document with w:defaultTabStop set to 360 twips (0.25 inch) in settings.xml",
        "expected_behavior": "Tab stops should be synthesized at 360-twip intervals instead of default 720",
        "current_status": "SUPPORTED — defaultTabStop parsed from settings.xml",
    }
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  tabs/default-tab-interval/")


def make_ptab_widget() -> None:
    """w:ptab — positioned tab element.

    ISO 29500-1 §17.3.3.22: ptab is an absolute position tab character that
    specifies alignment and leader separately from tab stop definitions.
    Should be treated as an opaque inline widget.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Normal text before ptab
    p.add_run("Before ")

    # Inject w:ptab into a run
    run_el = OxmlElement("w:r")
    ptab = make_element("w:ptab", {
        "w:alignment": "right",
        "w:relativeTo": "margin",
        "w:leader": "dot",
    })
    run_el.append(ptab)
    p._p.append(run_el)

    # Normal text after ptab
    p.add_run(" After")

    save_fixture("tabs", "ptab-widget", doc, {
        "name": "ptab-widget",
        "spec_ref": "ISO 29500-1 §17.3.3.22",
        "description": "Paragraph containing a w:ptab element (positioned tab)",
        "expected_behavior": "w:ptab should produce an OpaqueInline widget in the paragraph segments",
        "current_status": "SUPPORTED — ptab handled as opaque widget",
    })


# =========================================================================
# DOCUMENT STRUCTURE (ISO 29500-1 §17.3, §17.6)
# =========================================================================

def make_structure_fixtures() -> None:
    print("\n── Structure ──")
    make_flow_control_keep_next()
    make_flow_control_page_break_before()
    make_widow_control()
    make_section_page_size_orientation()
    make_section_columns()
    make_indentation_types()


def make_flow_control_keep_next() -> None:
    """w:keepNext + w:keepLines — paragraph flow control.

    ISO 29500-1 §17.3.1.14/15: keepNext keeps paragraph with next;
    keepLines keeps all lines of paragraph on same page.
    """
    doc = Document()

    p = doc.add_paragraph("This paragraph has keepNext and keepLines enabled.")
    pPr = p._p.get_or_add_pPr()
    pPr.append(make_element("w:keepNext"))
    pPr.append(make_element("w:keepLines"))

    doc.add_paragraph("This is the next paragraph that should stay with the previous one.")

    save_fixture("structure", "flow-control-keep-next", doc, {
        "name": "flow-control-keep-next",
        "spec_ref": "ISO 29500-1 §17.3.1.14/15",
        "description": "Paragraph with w:keepNext and w:keepLines in pPr",
        "expected_behavior": "Parser should expose keepNext and keepLines flags on ParagraphNode",
        "current_status": "NOT_PARSED — keepNext/keepLines silently ignored",
    })


def make_flow_control_page_break_before() -> None:
    """w:pageBreakBefore — force page break before paragraph.

    ISO 29500-1 §17.3.1.23: pageBreakBefore forces a page break before
    the paragraph.
    """
    doc = Document()

    doc.add_paragraph("Content before the page break.")

    p = doc.add_paragraph("This paragraph has pageBreakBefore enabled.")
    pPr = p._p.get_or_add_pPr()
    pPr.append(make_element("w:pageBreakBefore"))

    save_fixture("structure", "flow-control-page-break-before", doc, {
        "name": "flow-control-page-break-before",
        "spec_ref": "ISO 29500-1 §17.3.1.23",
        "description": "Paragraph with w:pageBreakBefore in pPr",
        "expected_behavior": "Parser should expose pageBreakBefore flag on ParagraphNode",
        "current_status": "NOT_PARSED — pageBreakBefore silently ignored",
    })


def make_widow_control() -> None:
    """w:widowControl — disable widow/orphan control.

    ISO 29500-1 §17.3.1.44: widowControl defaults to true per spec.
    Setting w:val="0" explicitly disables it.
    """
    doc = Document()

    p = doc.add_paragraph("This paragraph has widow/orphan control explicitly disabled.")
    pPr = p._p.get_or_add_pPr()
    pPr.append(make_element("w:widowControl", {"w:val": "0"}))

    doc.add_paragraph("This paragraph uses default widow control (enabled).")

    save_fixture("structure", "widow-control", doc, {
        "name": "widow-control",
        "spec_ref": "ISO 29500-1 §17.3.1.44",
        "description": "Paragraph with w:widowControl val=0 (disabled)",
        "expected_behavior": "Parser should expose widowControl=false on ParagraphNode",
        "current_status": "NOT_PARSED — widowControl silently ignored",
    })


def make_section_page_size_orientation() -> None:
    """sectPr with pgSz — A4 landscape orientation.

    ISO 29500-1 §17.6.14: pgSz defines page dimensions and orientation.
    A4 = w:w=11906 w:h=16838. Landscape adds w:orient="landscape".
    """
    doc = Document()

    doc.add_paragraph("Content in a section with A4 landscape orientation.")

    # Put sectPr in the last paragraph's pPr
    p = doc.add_paragraph("Last paragraph of the landscape section.")
    pPr = p._p.get_or_add_pPr()
    sect_pr = make_element("w:sectPr")
    pg_sz = make_element("w:pgSz", {
        "w:w": "11906",
        "w:h": "16838",
        "w:orient": "landscape",
    })
    sect_pr.append(pg_sz)
    pPr.append(sect_pr)

    save_fixture("structure", "section-page-size-orientation", doc, {
        "name": "section-page-size-orientation",
        "spec_ref": "ISO 29500-1 §17.6.14",
        "description": "Section with A4 page size and landscape orientation via sectPr in pPr",
        "expected_behavior": "Parser should expose page size and orientation as structured data",
        "current_status": "PARTIALLY_SUPPORTED — sectPr preserved as raw XML, not parsed into structured fields",
    })


def make_section_columns() -> None:
    """sectPr with cols — 2-column layout.

    ISO 29500-1 §17.6.4: cols in sectPr defines column layout.
    """
    doc = Document()

    doc.add_paragraph("Content in a 2-column section.")

    p = doc.add_paragraph("Last paragraph of the 2-column section.")
    pPr = p._p.get_or_add_pPr()
    sect_pr = make_element("w:sectPr")
    cols = make_element("w:cols", {"w:num": "2"})
    sect_pr.append(cols)
    pPr.append(sect_pr)

    save_fixture("structure", "section-columns", doc, {
        "name": "section-columns",
        "spec_ref": "ISO 29500-1 §17.6.4",
        "description": "Section with 2-column layout via w:cols in sectPr",
        "expected_behavior": "Parser should expose column count as structured data",
        "current_status": "PARTIALLY_SUPPORTED — sectPr preserved as raw XML, not parsed into structured fields",
    })


def make_indentation_types() -> None:
    """w:ind — various indentation types.

    ISO 29500-1 §17.3.1.12: ind in pPr defines paragraph indentation.
    Supports left, right, firstLine, and hanging indent.
    """
    doc = Document()

    # Left indent: 720 twips = 0.5 inch
    p1 = doc.add_paragraph("Paragraph with left indent (720 twips = 0.5 inch).")
    pPr1 = p1._p.get_or_add_pPr()
    pPr1.append(make_element("w:ind", {"w:left": "720"}))

    # Right indent: 720 twips
    p2 = doc.add_paragraph("Paragraph with right indent (720 twips = 0.5 inch).")
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:ind", {"w:right": "720"}))

    # First-line indent: 360 twips = 0.25 inch
    p3 = doc.add_paragraph("Paragraph with first-line indent (360 twips = 0.25 inch).")
    pPr3 = p3._p.get_or_add_pPr()
    pPr3.append(make_element("w:ind", {"w:firstLine": "360"}))

    # Hanging indent: 360 twips (stored as w:hanging)
    p4 = doc.add_paragraph("Paragraph with hanging indent (360 twips = 0.25 inch).")
    pPr4 = p4._p.get_or_add_pPr()
    pPr4.append(make_element("w:ind", {"w:left": "720", "w:hanging": "360"}))

    # Combined left + right + firstLine
    p5 = doc.add_paragraph("Paragraph with left=720, right=360, firstLine=360.")
    pPr5 = p5._p.get_or_add_pPr()
    pPr5.append(make_element("w:ind", {"w:left": "720", "w:right": "360", "w:firstLine": "360"}))

    save_fixture("structure", "indentation-types", doc, {
        "name": "indentation-types",
        "spec_ref": "ISO 29500-1 §17.3.1.12",
        "description": "Paragraphs with various indentation types: left, right, firstLine, hanging, combined",
        "expected_behavior": "Parser should expose all indentation values. Hanging stored as negative first_line.",
        "current_status": "SUPPORTED — indentation is fully parsed",
    })


# =========================================================================
# COMMENTS (ISO 29500-1 §17.13.4)
# =========================================================================

WPC = "http://schemas.openxmlformats.org/markup-compatibility/2006"
W14 = "http://schemas.microsoft.com/office/word/2010/wordml"
W15 = "http://schemas.microsoft.com/office/word/2012/wordml"

COMMENTS_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
COMMENTS_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"


def _inject_comment_range(p_element, comment_id: str, text_before: str, commented_text: str, text_after: str) -> None:
    """Inject commentRangeStart, commented text run, commentRangeEnd, and commentReference into a paragraph.

    Replaces all existing content in the paragraph.
    """
    # Clear existing children (runs, etc.)
    for child in list(p_element):
        p_element.remove(child)

    # Add pPr back if needed (preserve paragraph properties)
    # Not needed for our test fixtures

    # Text before the comment range
    if text_before:
        run = make_element("w:r")
        t = make_element("w:t")
        t.text = text_before
        t.set(qn("xml:space"), "preserve")
        run.append(t)
        p_element.append(run)

    # commentRangeStart
    range_start = make_element("w:commentRangeStart", {"w:id": comment_id})
    p_element.append(range_start)

    # The commented text
    run = make_element("w:r")
    t = make_element("w:t")
    t.text = commented_text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    p_element.append(run)

    # commentRangeEnd
    range_end = make_element("w:commentRangeEnd", {"w:id": comment_id})
    p_element.append(range_end)

    # commentReference run
    ref_run = make_element("w:r")
    ref_el = make_element("w:commentReference", {"w:id": comment_id})
    ref_run.append(ref_el)
    p_element.append(ref_run)

    # Text after the comment range
    if text_after:
        run = make_element("w:r")
        t = make_element("w:t")
        t.text = text_after
        t.set(qn("xml:space"), "preserve")
        run.append(t)
        p_element.append(run)


def _build_comments_xml(comments: list[dict]) -> bytes:
    """Build word/comments.xml content from a list of comment dicts.

    Each dict: { "id": str, "author": str, "date": str, "text": str, "parent_id": str|None }
    """
    root = etree.Element(
        f"{{{W}}}comments",
        nsmap={"w": W, "r": R},
    )

    for c in comments:
        attrs = {
            f"{{{W}}}id": c["id"],
            f"{{{W}}}author": c["author"],
            f"{{{W}}}date": c["date"],
            f"{{{W}}}initials": c.get("initials", c["author"][:2].upper()),
        }
        comment_el = etree.SubElement(root, f"{{{W}}}comment", attrib=attrs)

        # Add a paragraph with the comment text
        p = etree.SubElement(comment_el, f"{{{W}}}p")
        r = etree.SubElement(p, f"{{{W}}}r")
        t = etree.SubElement(r, f"{{{W}}}t")
        t.text = c["text"]

    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _add_comments_part(doc: DocxDocument, comments_xml: bytes) -> None:
    """Add word/comments.xml to the docx package and wire up the relationship."""
    import zipfile
    import io
    import shutil
    from copy import deepcopy

    # Save to a temporary buffer, then manipulate the zip
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    out_buf = io.BytesIO()
    with zipfile.ZipFile(buf, "r") as zin, zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename == "[Content_Types].xml":
                # Add comments content type
                ct_root = etree.fromstring(data)
                ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
                # Check if already exists
                existing = ct_root.findall(f"{{{ct_ns}}}Override[@PartName='/word/comments.xml']")
                if not existing:
                    override_el = etree.SubElement(ct_root, f"{{{ct_ns}}}Override")
                    override_el.set("PartName", "/word/comments.xml")
                    override_el.set("ContentType", COMMENTS_CONTENT_TYPE)
                data = etree.tostring(ct_root, xml_declaration=True, encoding="UTF-8", standalone=True)

            elif item.filename == "word/_rels/document.xml.rels":
                # Add relationship to comments.xml
                rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
                rels_root = etree.fromstring(data)
                # Find next rId
                existing_ids = [el.get("Id") for el in rels_root]
                max_id = 0
                for rid in existing_ids:
                    if rid and rid.startswith("rId"):
                        try:
                            max_id = max(max_id, int(rid[3:]))
                        except ValueError:
                            pass
                new_rid = f"rId{max_id + 1}"
                rel_el = etree.SubElement(rels_root, f"{{{rels_ns}}}Relationship")
                rel_el.set("Id", new_rid)
                rel_el.set("Type", COMMENTS_REL_TYPE)
                rel_el.set("Target", "comments.xml")
                data = etree.tostring(rels_root, xml_declaration=True, encoding="UTF-8", standalone=True)

            zout.writestr(item, data)

        # Add the comments.xml part
        zout.writestr("word/comments.xml", comments_xml)

    # Now re-load the doc from the manipulated zip
    out_buf.seek(0)
    return out_buf


def make_comment_fixtures() -> None:
    print("\n── Comments ──")
    make_basic_comment()
    make_multi_paragraph_comment()
    make_multiple_comments()
    make_comment_with_reply()


def make_basic_comment() -> None:
    """Basic comment anchored to a text range.

    ISO 29500-1 §17.13.4.2/3/5: commentRangeStart/End mark the commented
    range in the document body; commentReference links to the comment story.
    """
    doc = Document()
    p = doc.add_paragraph("Some text before. This is commented. Some text after.")

    # Inject comment range markers
    _inject_comment_range(
        p._p,
        comment_id="1",
        text_before="Some text before. ",
        commented_text="This is commented.",
        text_after=" Some text after.",
    )

    # Build comments.xml
    comments_xml = _build_comments_xml([{
        "id": "1",
        "author": "Test Author",
        "date": "2025-01-01T00:00:00Z",
        "text": "This is a comment.",
    }])

    # Save and inject comments part
    out_buf = _add_comments_part(doc, comments_xml)

    # Write directly
    out_dir = ROOT / "comments" / "basic-comment"
    out_dir.mkdir(parents=True, exist_ok=True)
    (out_dir / "input.docx").write_bytes(out_buf.getvalue())
    (out_dir / "metadata.json").write_text(json.dumps({
        "name": "basic-comment",
        "spec_ref": "ISO 29500-1 §17.13.4.2/3/5",
        "description": "Paragraph with a comment anchored to a text range via commentRangeStart/End/Reference",
        "expected_behavior": "CommentStory with id=1, author=Test Author parsed; range markers in body inlines",
        "current_status": "SUPPORTED — comments are parsed into CommentStory",
    }, indent=2) + "\n")
    print(f"  comments/basic-comment/")


def make_multi_paragraph_comment() -> None:
    """Comment range spanning multiple paragraphs.

    ISO 29500-1 §17.13.4.2/3: commentRangeStart and commentRangeEnd can
    be in different paragraphs, creating a cross-block range.
    """
    doc = Document()
    p1 = doc.add_paragraph("First paragraph with comment start.")
    p2 = doc.add_paragraph("Second paragraph inside comment range.")
    p3 = doc.add_paragraph("Third paragraph with comment end.")

    # Clear p1 and inject commentRangeStart + text
    for child in list(p1._p):
        p1._p.remove(child)
    range_start = make_element("w:commentRangeStart", {"w:id": "1"})
    p1._p.append(range_start)
    run = make_element("w:r")
    t = make_element("w:t")
    t.text = "First paragraph with comment start."
    run.append(t)
    p1._p.append(run)

    # p2 is just normal text inside the range (no markers needed)

    # Clear p3 and inject text + commentRangeEnd + commentReference
    for child in list(p3._p):
        p3._p.remove(child)
    run = make_element("w:r")
    t = make_element("w:t")
    t.text = "Third paragraph with comment end."
    run.append(t)
    p3._p.append(run)
    range_end = make_element("w:commentRangeEnd", {"w:id": "1"})
    p3._p.append(range_end)
    ref_run = make_element("w:r")
    ref_el = make_element("w:commentReference", {"w:id": "1"})
    ref_run.append(ref_el)
    p3._p.append(ref_run)

    comments_xml = _build_comments_xml([{
        "id": "1",
        "author": "Test Author",
        "date": "2025-01-01T00:00:00Z",
        "text": "Multi-paragraph comment.",
    }])

    out_buf = _add_comments_part(doc, comments_xml)

    out_dir = ROOT / "comments" / "multi-paragraph-comment"
    out_dir.mkdir(parents=True, exist_ok=True)
    (out_dir / "input.docx").write_bytes(out_buf.getvalue())
    (out_dir / "metadata.json").write_text(json.dumps({
        "name": "multi-paragraph-comment",
        "spec_ref": "ISO 29500-1 §17.13.4.2/3",
        "description": "Comment range spanning three paragraphs — commentRangeStart in p1, commentRangeEnd in p3",
        "expected_behavior": "CommentRangeStart in first paragraph, CommentRangeEnd in third paragraph",
        "current_status": "SUPPORTED — cross-block comment ranges parsed",
    }, indent=2) + "\n")
    print(f"  comments/multi-paragraph-comment/")


def make_multiple_comments() -> None:
    """Multiple independent comments on different text ranges.

    ISO 29500-1 §17.13.4: Multiple comment anchors with distinct IDs
    can coexist in the same paragraph.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Clear paragraph
    for child in list(p._p):
        p._p.remove(child)

    # First comment range (id=1)
    range_start1 = make_element("w:commentRangeStart", {"w:id": "1"})
    p._p.append(range_start1)
    run1 = make_element("w:r")
    t1 = make_element("w:t")
    t1.text = "First commented text"
    t1.set(qn("xml:space"), "preserve")
    run1.append(t1)
    p._p.append(run1)
    range_end1 = make_element("w:commentRangeEnd", {"w:id": "1"})
    p._p.append(range_end1)
    ref_run1 = make_element("w:r")
    ref_run1.append(make_element("w:commentReference", {"w:id": "1"}))
    p._p.append(ref_run1)

    # Separator text
    sep_run = make_element("w:r")
    sep_t = make_element("w:t")
    sep_t.text = " and "
    sep_t.set(qn("xml:space"), "preserve")
    sep_run.append(sep_t)
    p._p.append(sep_run)

    # Second comment range (id=2)
    range_start2 = make_element("w:commentRangeStart", {"w:id": "2"})
    p._p.append(range_start2)
    run2 = make_element("w:r")
    t2 = make_element("w:t")
    t2.text = "second commented text"
    t2.set(qn("xml:space"), "preserve")
    run2.append(t2)
    p._p.append(run2)
    range_end2 = make_element("w:commentRangeEnd", {"w:id": "2"})
    p._p.append(range_end2)
    ref_run2 = make_element("w:r")
    ref_run2.append(make_element("w:commentReference", {"w:id": "2"}))
    p._p.append(ref_run2)

    comments_xml = _build_comments_xml([
        {
            "id": "1",
            "author": "Author One",
            "date": "2025-01-01T00:00:00Z",
            "text": "Comment on first range.",
        },
        {
            "id": "2",
            "author": "Author Two",
            "date": "2025-01-02T00:00:00Z",
            "text": "Comment on second range.",
        },
    ])

    out_buf = _add_comments_part(doc, comments_xml)

    out_dir = ROOT / "comments" / "multiple-comments"
    out_dir.mkdir(parents=True, exist_ok=True)
    (out_dir / "input.docx").write_bytes(out_buf.getvalue())
    (out_dir / "metadata.json").write_text(json.dumps({
        "name": "multiple-comments",
        "spec_ref": "ISO 29500-1 §17.13.4",
        "description": "Paragraph with two separate comments on different text ranges (id=1, id=2)",
        "expected_behavior": "Two CommentStory entries with distinct IDs, authors, and content",
        "current_status": "SUPPORTED — multiple comments independently parsed",
    }, indent=2) + "\n")
    print(f"  comments/multiple-comments/")


def make_comment_with_reply() -> None:
    """Comment with a reply (threaded comments).

    ISO 29500-1 §17.13.4: A reply comment references its parent via
    a custom attribute. Word uses w16:paraId or similar, but the common
    pattern is that the reply w:comment has no range in the body — only
    the parent comment has range markers.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Inject comment range for parent comment (id=1)
    _inject_comment_range(
        p._p,
        comment_id="1",
        text_before="",
        commented_text="Text with a comment and reply.",
        text_after="",
    )

    # Build comments.xml with parent and reply
    # The reply comment (id=2) references parent (id=1).
    # We put the parent ID link in the comment element itself.
    root = etree.Element(
        f"{{{W}}}comments",
        nsmap={"w": W, "r": R},
    )

    # Parent comment
    parent_attrs = {
        f"{{{W}}}id": "1",
        f"{{{W}}}author": "Test Author",
        f"{{{W}}}date": "2025-01-01T00:00:00Z",
        f"{{{W}}}initials": "TA",
    }
    parent_el = etree.SubElement(root, f"{{{W}}}comment", attrib=parent_attrs)
    pp = etree.SubElement(parent_el, f"{{{W}}}p")
    pr = etree.SubElement(pp, f"{{{W}}}r")
    pt = etree.SubElement(pr, f"{{{W}}}t")
    pt.text = "Parent comment text."

    # Reply comment — no range markers in body, just a comment entry
    reply_attrs = {
        f"{{{W}}}id": "2",
        f"{{{W}}}author": "Reply Author",
        f"{{{W}}}date": "2025-01-02T00:00:00Z",
        f"{{{W}}}initials": "RA",
    }
    reply_el = etree.SubElement(root, f"{{{W}}}comment", attrib=reply_attrs)
    rp = etree.SubElement(reply_el, f"{{{W}}}p")
    rr = etree.SubElement(rp, f"{{{W}}}r")
    rt = etree.SubElement(rr, f"{{{W}}}t")
    rt.text = "Reply to parent comment."

    comments_xml = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)

    out_buf = _add_comments_part(doc, comments_xml)

    out_dir = ROOT / "comments" / "comment-with-reply"
    out_dir.mkdir(parents=True, exist_ok=True)
    (out_dir / "input.docx").write_bytes(out_buf.getvalue())
    (out_dir / "metadata.json").write_text(json.dumps({
        "name": "comment-with-reply",
        "spec_ref": "ISO 29500-1 §17.13.4",
        "description": "Parent comment (id=1) and reply comment (id=2) — tests comment threading",
        "expected_behavior": "Both comments parsed as separate CommentStory entries",
        "current_status": "SUPPORTED — reply comments parsed (threading metadata not modeled)",
    }, indent=2) + "\n")
    print(f"  comments/comment-with-reply/")


# =========================================================================
# BOOKMARKS (ISO 29500-1 §17.13.6)
# =========================================================================

def make_bookmark_fixtures() -> None:
    print("\n── Bookmarks ──")
    make_basic_bookmark()
    make_multi_paragraph_bookmark()
    make_nested_bookmarks()
    make_named_bookmark()
    make_bookmark_roundtrip()


def make_basic_bookmark() -> None:
    """w:bookmarkStart/w:bookmarkEnd — basic bookmark wrapping text.

    ISO 29500-1 §17.13.6.2/1: bookmarkStart/bookmarkEnd define a bookmark
    range around text content within a paragraph.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Inject bookmarkStart
    bm_start = make_element("w:bookmarkStart", {
        "w:id": "0",
        "w:name": "my_bookmark",
    })
    p._p.append(bm_start)

    # Add a run with text
    p.add_run("Bookmarked text")

    # Inject bookmarkEnd
    bm_end = make_element("w:bookmarkEnd", {"w:id": "0"})
    p._p.append(bm_end)

    save_fixture("bookmarks", "basic-bookmark", doc, {
        "name": "basic-bookmark",
        "spec_ref": "ISO 29500-1 §17.13.6.2/1",
        "description": "Paragraph with a bookmark range wrapping text",
        "expected_behavior": "bookmarkStart and bookmarkEnd should produce Decoration inlines with DecorationType::Bookmark",
        "current_status": "SUPPORTED — bookmarks parsed as zero-width Decoration nodes",
    })


def make_multi_paragraph_bookmark() -> None:
    """Bookmark range spanning two paragraphs.

    ISO 29500-1 §17.13.6.2/1: A bookmark range can span across paragraph
    boundaries. bookmarkStart appears in one paragraph, bookmarkEnd in another.
    """
    doc = Document()

    # Paragraph 1 with bookmarkStart
    p1 = doc.add_paragraph()
    bm_start = make_element("w:bookmarkStart", {
        "w:id": "1",
        "w:name": "cross_para_bookmark",
    })
    p1._p.append(bm_start)
    p1.add_run("Start of bookmarked range")

    # Paragraph 2 with bookmarkEnd
    p2 = doc.add_paragraph()
    p2.add_run("End of bookmarked range")
    bm_end = make_element("w:bookmarkEnd", {"w:id": "1"})
    p2._p.append(bm_end)

    save_fixture("bookmarks", "multi-paragraph-bookmark", doc, {
        "name": "multi-paragraph-bookmark",
        "spec_ref": "ISO 29500-1 §17.13.6.2/1",
        "description": "Bookmark range spanning two paragraphs (bookmarkStart in para 1, bookmarkEnd in para 2)",
        "expected_behavior": "bookmarkStart decoration in paragraph 1, bookmarkEnd decoration in paragraph 2",
        "current_status": "SUPPORTED — cross-paragraph bookmarks parsed as decorations in respective paragraphs",
    })


def make_nested_bookmarks() -> None:
    """Overlapping/nested bookmark ranges.

    ISO 29500-1 §17.13.6: Bookmark ranges may overlap — one bookmark
    can start before another ends.
    """
    doc = Document()
    p = doc.add_paragraph()

    # [bookmark A start]
    bm_a_start = make_element("w:bookmarkStart", {
        "w:id": "2",
        "w:name": "bookmark_a",
    })
    p._p.append(bm_a_start)

    p.add_run("text1 ")

    # [bookmark B start]
    bm_b_start = make_element("w:bookmarkStart", {
        "w:id": "3",
        "w:name": "bookmark_b",
    })
    p._p.append(bm_b_start)

    p.add_run("text2 ")

    # [bookmark A end]
    bm_a_end = make_element("w:bookmarkEnd", {"w:id": "2"})
    p._p.append(bm_a_end)

    p.add_run("text3")

    # [bookmark B end]
    bm_b_end = make_element("w:bookmarkEnd", {"w:id": "3"})
    p._p.append(bm_b_end)

    save_fixture("bookmarks", "nested-bookmarks", doc, {
        "name": "nested-bookmarks",
        "spec_ref": "ISO 29500-1 §17.13.6.2/1",
        "description": "Two overlapping bookmark ranges: A wraps text1+text2, B wraps text2+text3",
        "expected_behavior": "Four Decoration nodes (two starts, two ends) with correct pairing by id",
        "current_status": "SUPPORTED — overlapping bookmarks parsed as independent decorations",
    })


def make_named_bookmark() -> None:
    """Bookmark with a semantically meaningful name for hyperlink anchoring.

    ISO 29500-1 §17.13.6.2: The w:name attribute on bookmarkStart provides
    a unique name that can be referenced by hyperlinks (w:anchor).
    """
    doc = Document()
    p = doc.add_paragraph()

    bm_start = make_element("w:bookmarkStart", {
        "w:id": "4",
        "w:name": "clause_1_start",
    })
    p._p.append(bm_start)

    p.add_run("Clause 1: Definitions and Interpretation")

    bm_end = make_element("w:bookmarkEnd", {"w:id": "4"})
    p._p.append(bm_end)

    save_fixture("bookmarks", "named-bookmark", doc, {
        "name": "named-bookmark",
        "spec_ref": "ISO 29500-1 §17.13.6.2",
        "description": "Bookmark with semantically meaningful name 'clause_1_start' for hyperlink anchor targets",
        "expected_behavior": "Bookmark decorations should be present; name should be preserved for cross-reference resolution",
        "current_status": "SUPPORTED — bookmark name is preserved in the opaque element XML",
    })


def make_bookmark_roundtrip() -> None:
    """Bookmark that must survive import -> export -> reimport.

    Same structure as basic-bookmark but specifically tests that bookmarks
    survive the redline pipeline.
    """
    doc = Document()
    p = doc.add_paragraph()

    bm_start = make_element("w:bookmarkStart", {
        "w:id": "5",
        "w:name": "roundtrip_bookmark",
    })
    p._p.append(bm_start)

    p.add_run("Text with bookmark for roundtrip testing")

    bm_end = make_element("w:bookmarkEnd", {"w:id": "5"})
    p._p.append(bm_end)

    save_fixture("bookmarks", "bookmark-roundtrip", doc, {
        "name": "bookmark-roundtrip",
        "spec_ref": "ISO 29500-1 §17.13.6.2/1",
        "description": "Bookmark for roundtrip fidelity testing (import -> redline export -> reimport)",
        "expected_behavior": "Bookmark decorations must survive the full redline pipeline",
        "current_status": "SUPPORTED — bookmarks survive as Decoration nodes through redline",
    })


# =========================================================================
# HYPERLINKS (ISO 29500-1 §17.16.22)
# =========================================================================

def make_hyperlink_fixtures() -> None:
    print("\n── Hyperlinks ──")
    make_external_url_hyperlink()
    make_internal_anchor_hyperlink()
    make_hyperlink_with_formatting()
    make_hyperlink_no_text()


def make_external_url_hyperlink() -> None:
    """w:hyperlink with r:id — external URL hyperlink.

    ISO 29500-1 §17.16.22: A w:hyperlink element with r:id pointing to
    a relationship entry that has TargetMode="External" and the URL as Target.
    """
    doc = Document()
    p = doc.add_paragraph("Text before link. ")

    # Add external hyperlink relationship
    r_id = doc.part.relate_to(
        'https://example.com',
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )

    # Build w:hyperlink element with r:id
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = make_element("w:rPr")
    r_style = make_element("w:rStyle", {"w:val": "Hyperlink"})
    rPr.append(r_style)
    run.append(rPr)
    t = make_element("w:t")
    t.text = "Click here"
    run.append(t)
    hyperlink.append(run)

    p._p.append(hyperlink)
    p.add_run(" Text after link.")

    save_fixture("hyperlinks", "external-url", doc, {
        "name": "external-url",
        "spec_ref": "ISO 29500-1 §17.16.22",
        "description": "Paragraph with an external hyperlink to https://example.com via r:id relationship",
        "expected_behavior": "HyperlinkData.url should be Some('https://example.com'), text should be 'Click here'",
        "current_status": "PARTIAL — r:id relationship resolution not implemented; url is always None",
    })


def make_internal_anchor_hyperlink() -> None:
    """w:hyperlink with w:anchor — internal bookmark link.

    ISO 29500-1 §17.16.22: A w:hyperlink with w:anchor attribute links to
    a bookmark within the same document. No r:id is needed.
    """
    doc = Document()

    # Create a bookmark target first
    p_target = doc.add_paragraph("This is the bookmark target paragraph.")
    # Inject bookmarkStart and bookmarkEnd
    bookmark_start = make_element("w:bookmarkStart", {
        "w:id": "0",
        "w:name": "bookmark_name",
    })
    bookmark_end = make_element("w:bookmarkEnd", {"w:id": "0"})
    p_target._p.insert(0, bookmark_start)
    p_target._p.append(bookmark_end)

    # Create the hyperlink paragraph
    p = doc.add_paragraph("Click ")

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), "bookmark_name")

    run = OxmlElement("w:r")
    rPr = make_element("w:rPr")
    r_style = make_element("w:rStyle", {"w:val": "Hyperlink"})
    rPr.append(r_style)
    run.append(rPr)
    t = make_element("w:t")
    t.text = "jump to bookmark"
    run.append(t)
    hyperlink.append(run)

    p._p.append(hyperlink)
    p.add_run(" to navigate.")

    save_fixture("hyperlinks", "internal-anchor", doc, {
        "name": "internal-anchor",
        "spec_ref": "ISO 29500-1 §17.16.22",
        "description": "Paragraph with an internal hyperlink using w:anchor='bookmark_name'",
        "expected_behavior": "HyperlinkData.anchor should be Some('bookmark_name'), url should be None",
        "current_status": "SUPPORTED — anchor attribute is parsed from the element",
    })


def make_hyperlink_with_formatting() -> None:
    """w:hyperlink with multiple formatted runs.

    ISO 29500-1 §17.16.22: A hyperlink can contain multiple runs with
    different formatting. The display text is the concatenation of all runs.
    """
    doc = Document()
    p = doc.add_paragraph("Link with formatting: ")

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), "some_target")

    # First run: bold
    run1 = OxmlElement("w:r")
    rPr1 = make_element("w:rPr")
    rPr1.append(make_element("w:b"))
    r_style1 = make_element("w:rStyle", {"w:val": "Hyperlink"})
    rPr1.append(r_style1)
    run1.append(rPr1)
    t1 = make_element("w:t")
    t1.text = "bold part"
    run1.append(t1)
    hyperlink.append(run1)

    # Second run: not bold
    run2 = OxmlElement("w:r")
    rPr2 = make_element("w:rPr")
    r_style2 = make_element("w:rStyle", {"w:val": "Hyperlink"})
    rPr2.append(r_style2)
    run2.append(rPr2)
    t2 = make_element("w:t")
    t2.set(qn("xml:space"), "preserve")
    t2.text = " and normal part"
    run2.append(t2)
    hyperlink.append(run2)

    p._p.append(hyperlink)

    save_fixture("hyperlinks", "hyperlink-with-formatting", doc, {
        "name": "hyperlink-with-formatting",
        "spec_ref": "ISO 29500-1 §17.16.22",
        "description": "Hyperlink with two runs: first bold, second normal",
        "expected_behavior": "HyperlinkData.text should be 'bold part and normal part' (concatenation of all runs)",
        "current_status": "SUPPORTED — text extraction recursively collects from all runs",
    })


def make_hyperlink_no_text() -> None:
    """w:hyperlink with no nested runs — empty hyperlink.

    Edge case: a hyperlink element with no content. Should parse without
    crashing and produce an empty or absent hyperlink node.
    """
    doc = Document()
    p = doc.add_paragraph("Before empty link.")

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), "nowhere")
    # No runs inside — empty hyperlink

    p._p.append(hyperlink)
    p.add_run(" After empty link.")

    save_fixture("hyperlinks", "hyperlink-no-text", doc, {
        "name": "hyperlink-no-text",
        "spec_ref": "ISO 29500-1 §17.16.22",
        "description": "Hyperlink element with no nested runs (empty hyperlink)",
        "expected_behavior": "Should not crash; may produce an OpaqueInline with empty text or be filtered out",
        "current_status": "SUPPORTED — parses without crash, produces hyperlink with empty text",
    })


# =========================================================================
# FIELDS (ISO 29500-1 §17.16)
# =========================================================================

def make_field_fixtures() -> None:
    print("\n── Fields ──")
    make_simple_field()
    make_complex_field()
    make_toc_field()
    make_ref_field()
    make_multiple_fields()


def make_simple_field() -> None:
    """w:fldSimple — self-contained simple field (PAGE).

    ISO 29500-1 §17.16.19: fldSimple wraps the field instruction and
    result in a single element with a w:instr attribute.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Inject fldSimple directly into the paragraph
    fld_simple = make_element("w:fldSimple", {"w:instr": " PAGE "})
    run = make_element("w:r")
    t = make_element("w:t")
    t.text = "1"
    run.append(t)
    fld_simple.append(run)
    p._p.append(fld_simple)

    save_fixture("fields", "simple-field", doc, {
        "name": "simple-field",
        "spec_ref": "ISO 29500-1 §17.16.19",
        "description": "Paragraph with a w:fldSimple field (PAGE) containing result text '1'",
        "expected_behavior": "Parser should produce OpaqueInline with FieldKind::Simple, instruction_text=' PAGE ', result_text='1'",
        "current_status": "SUPPORTED — fldSimple parsed into OpaqueKind::Field(FieldData)",
    })


def make_complex_field() -> None:
    """Complex field using fldChar begin/separate/end with instrText.

    ISO 29500-1 §17.16.18: Complex fields use fldChar elements to mark
    begin, separate, and end boundaries. instrText between begin and
    separate holds the field instruction.
    """
    doc = Document()
    p = doc.add_paragraph()

    # begin
    r_begin = make_element("w:r")
    fld_begin = make_element("w:fldChar", {"w:fldCharType": "begin"})
    r_begin.append(fld_begin)
    p._p.append(r_begin)

    # instrText
    r_instr = make_element("w:r")
    instr_text = make_element("w:instrText")
    instr_text.text = ' DATE \\@ "yyyy-MM-dd" '
    instr_text.set(qn("xml:space"), "preserve")
    r_instr.append(instr_text)
    p._p.append(r_instr)

    # separate
    r_sep = make_element("w:r")
    fld_sep = make_element("w:fldChar", {"w:fldCharType": "separate"})
    r_sep.append(fld_sep)
    p._p.append(r_sep)

    # result text
    r_result = make_element("w:r")
    t_result = make_element("w:t")
    t_result.text = "2025-01-01"
    r_result.append(t_result)
    p._p.append(r_result)

    # end
    r_end = make_element("w:r")
    fld_end = make_element("w:fldChar", {"w:fldCharType": "end"})
    r_end.append(fld_end)
    p._p.append(r_end)

    save_fixture("fields", "complex-field", doc, {
        "name": "complex-field",
        "spec_ref": "ISO 29500-1 §17.16.18",
        "description": "Paragraph with a complex field (DATE) using fldChar begin/separate/end",
        "expected_behavior": "Parser should produce OpaqueInlines for Begin, Instruction, Separate, End",
        "current_status": "SUPPORTED — complex field elements parsed into individual OpaqueInlines",
    })


def make_toc_field() -> None:
    """w:fldSimple with TOC instruction.

    ISO 29500-1 §17.16.36: TOC field generates a table of contents
    from heading levels specified by \\o switch.
    """
    doc = Document()
    p = doc.add_paragraph()

    fld_simple = make_element("w:fldSimple", {"w:instr": ' TOC \\o "1-3" '})
    run = make_element("w:r")
    t = make_element("w:t")
    t.text = "Table of Contents placeholder"
    run.append(t)
    fld_simple.append(run)
    p._p.append(fld_simple)

    save_fixture("fields", "toc-field", doc, {
        "name": "toc-field",
        "spec_ref": "ISO 29500-1 §17.16.36",
        "description": "Paragraph with a fldSimple TOC field with \\o switch",
        "expected_behavior": "Parser should produce OpaqueInline with FieldKind::Simple and instruction containing 'TOC'",
        "current_status": "SUPPORTED — fldSimple parsed regardless of instruction type",
    })


def make_ref_field() -> None:
    """w:fldSimple with REF instruction and a bookmark target.

    ISO 29500-1 §17.16.16: REF field displays the content of a
    specified bookmark.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Add a bookmark named "my_bookmark"
    bookmark_start = make_element("w:bookmarkStart", {
        "w:id": "0",
        "w:name": "my_bookmark",
    })
    p._p.append(bookmark_start)

    run_bookmarked = p.add_run("Bookmarked content in Section 1")

    bookmark_end = make_element("w:bookmarkEnd", {"w:id": "0"})
    p._p.append(bookmark_end)

    # Second paragraph with the REF field
    p2 = doc.add_paragraph()
    fld_simple = make_element("w:fldSimple", {"w:instr": " REF my_bookmark "})
    run = make_element("w:r")
    t = make_element("w:t")
    t.text = "See Section 1"
    run.append(t)
    fld_simple.append(run)
    p2._p.append(fld_simple)

    save_fixture("fields", "ref-field", doc, {
        "name": "ref-field",
        "spec_ref": "ISO 29500-1 §17.16.16",
        "description": "Document with a bookmark and a REF field referencing it",
        "expected_behavior": "Parser should produce OpaqueInline with FieldKind::Simple and instruction containing 'REF my_bookmark'",
        "current_status": "SUPPORTED — fldSimple parsed; bookmark presence verifiable separately",
    })


def make_multiple_fields() -> None:
    """Multiple fldSimple fields in a single paragraph.

    Tests that two independent fields in the same paragraph are both
    parsed as separate OpaqueInline nodes.
    """
    doc = Document()
    p = doc.add_paragraph()

    p.add_run("Page ")

    # First field: PAGE
    fld1 = make_element("w:fldSimple", {"w:instr": " PAGE "})
    r1 = make_element("w:r")
    t1 = make_element("w:t")
    t1.text = "1"
    r1.append(t1)
    fld1.append(r1)
    p._p.append(fld1)

    p.add_run(" of ")

    # Second field: NUMPAGES
    fld2 = make_element("w:fldSimple", {"w:instr": " NUMPAGES "})
    r2 = make_element("w:r")
    t2 = make_element("w:t")
    t2.text = "10"
    r2.append(t2)
    fld2.append(r2)
    p._p.append(fld2)

    save_fixture("fields", "multiple-fields", doc, {
        "name": "multiple-fields",
        "spec_ref": "ISO 29500-1 §17.16.19",
        "description": "Paragraph with two fldSimple fields: PAGE and NUMPAGES",
        "expected_behavior": "Both fields should produce separate OpaqueInline nodes with FieldKind::Simple",
        "current_status": "SUPPORTED — each fldSimple parsed independently",
    })


# =========================================================================
# FIELDS — DEEP (ISO 29500-1 §17.16, §17.3.3)
# =========================================================================

def make_fields_deep_fixtures() -> None:
    print("\n── Fields Deep ──")
    make_page_number_field()
    make_date_field_simple()
    make_seq_field()
    make_complex_field_nested()
    make_special_characters()
    make_styleref_field()


def make_page_number_field() -> None:
    """Header with PAGE and NUMPAGES fields as fldSimple.

    ISO 29500-1 §17.16.19: fldSimple fields in a header paragraph.
    """
    doc = Document()
    p = doc.add_paragraph()

    p.add_run("Page ")

    fld_page = make_element("w:fldSimple", {"w:instr": " PAGE "})
    r_page = make_element("w:r")
    t_page = make_element("w:t")
    t_page.text = "3"
    r_page.append(t_page)
    fld_page.append(r_page)
    p._p.append(fld_page)

    p.add_run(" of ")

    fld_numpages = make_element("w:fldSimple", {"w:instr": " NUMPAGES "})
    r_num = make_element("w:r")
    t_num = make_element("w:t")
    t_num.text = "15"
    r_num.append(t_num)
    fld_numpages.append(r_num)
    p._p.append(fld_numpages)

    save_fixture("fields-deep", "page-number-field", doc, {
        "name": "page-number-field",
        "spec_ref": "ISO 29500-1 §17.16.19",
        "description": "Paragraph with PAGE and NUMPAGES fldSimple fields",
        "expected_behavior": "Both fields produce FieldKind::Simple with their respective instructions",
        "current_status": "SUPPORTED — fldSimple parsed for any instruction type",
    })


def make_date_field_simple() -> None:
    """Paragraph with DATE field in fldSimple format.

    ISO 29500-1 §17.16.6: DATE field displays the current date.
    """
    doc = Document()
    p = doc.add_paragraph()

    fld_date = make_element("w:fldSimple", {"w:instr": r' DATE \@ "yyyy-MM-dd" '})
    r = make_element("w:r")
    t = make_element("w:t")
    t.text = "2025-01-15"
    r.append(t)
    fld_date.append(r)
    p._p.append(fld_date)

    save_fixture("fields-deep", "date-field", doc, {
        "name": "date-field",
        "spec_ref": "ISO 29500-1 §17.16.6",
        "description": "Paragraph with a DATE fldSimple field with format switch",
        "expected_behavior": "FieldKind::Simple with instruction containing 'DATE'",
        "current_status": "SUPPORTED — fldSimple parsed for any instruction type",
    })


def make_seq_field() -> None:
    """Paragraphs with SEQ fields for figure numbering.

    ISO 29500-1 §17.16.30: SEQ field inserts a sequence number.
    """
    doc = Document()

    # First SEQ field: "Figure 1"
    p1 = doc.add_paragraph()
    p1.add_run("Figure ")
    fld1 = make_element("w:fldSimple", {"w:instr": " SEQ Figure "})
    r1 = make_element("w:r")
    t1 = make_element("w:t")
    t1.text = "1"
    r1.append(t1)
    fld1.append(r1)
    p1._p.append(fld1)
    p1.add_run(": First diagram")

    # Second SEQ field: "Figure 2"
    p2 = doc.add_paragraph()
    p2.add_run("Figure ")
    fld2 = make_element("w:fldSimple", {"w:instr": " SEQ Figure "})
    r2 = make_element("w:r")
    t2 = make_element("w:t")
    t2.text = "2"
    r2.append(t2)
    fld2.append(r2)
    p2._p.append(fld2)
    p2.add_run(": Second diagram")

    save_fixture("fields-deep", "seq-field", doc, {
        "name": "seq-field",
        "spec_ref": "ISO 29500-1 §17.16.30",
        "description": "Two paragraphs with SEQ Figure fields (cached values 1 and 2)",
        "expected_behavior": "Each SEQ field produces FieldKind::Simple with instruction containing 'SEQ'",
        "current_status": "SUPPORTED — fldSimple parsed for any instruction type",
    })


def make_complex_field_nested() -> None:
    """Nested complex field: IF field containing a REF field.

    ISO 29500-1 §17.16.18: Complex fields can be nested — a field
    between begin/end of an outer field can itself be a complete
    begin/instrText/separate/end sequence.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Outer IF field: begin
    r = make_element("w:r")
    r.append(make_element("w:fldChar", {"w:fldCharType": "begin"}))
    p._p.append(r)

    # Outer instrText: IF condition
    r = make_element("w:r")
    instr = make_element("w:instrText")
    instr.text = ' IF '
    instr.set(qn("xml:space"), "preserve")
    r.append(instr)
    p._p.append(r)

    # Inner REF field (nested): begin
    r = make_element("w:r")
    r.append(make_element("w:fldChar", {"w:fldCharType": "begin"}))
    p._p.append(r)

    # Inner instrText: REF
    r = make_element("w:r")
    instr = make_element("w:instrText")
    instr.text = ' REF bookmark1 '
    instr.set(qn("xml:space"), "preserve")
    r.append(instr)
    p._p.append(r)

    # Inner: separate
    r = make_element("w:r")
    r.append(make_element("w:fldChar", {"w:fldCharType": "separate"}))
    p._p.append(r)

    # Inner: result text
    r = make_element("w:r")
    t = make_element("w:t")
    t.text = "value1"
    r.append(t)
    p._p.append(r)

    # Inner: end
    r = make_element("w:r")
    r.append(make_element("w:fldChar", {"w:fldCharType": "end"}))
    p._p.append(r)

    # Continue outer instrText: = "yes" "True" "False"
    r = make_element("w:r")
    instr = make_element("w:instrText")
    instr.text = ' = "yes" "True" "False" '
    instr.set(qn("xml:space"), "preserve")
    r.append(instr)
    p._p.append(r)

    # Outer: separate
    r = make_element("w:r")
    r.append(make_element("w:fldChar", {"w:fldCharType": "separate"}))
    p._p.append(r)

    # Outer: result text
    r = make_element("w:r")
    t = make_element("w:t")
    t.text = "False"
    r.append(t)
    p._p.append(r)

    # Outer: end
    r = make_element("w:r")
    r.append(make_element("w:fldChar", {"w:fldCharType": "end"}))
    p._p.append(r)

    save_fixture("fields-deep", "complex-field-nested", doc, {
        "name": "complex-field-nested",
        "spec_ref": "ISO 29500-1 §17.16.18",
        "description": "Nested complex fields: IF field with inner REF field",
        "expected_behavior": "Outer and inner begin/separate/end should all be captured as field opaques",
        "current_status": "SUPPORTED — each fldChar/instrText is independently parsed",
    })


def make_special_characters() -> None:
    """Paragraph with w:sym, w:noBreakHyphen, w:softHyphen, w:cr elements.

    ISO 29500-1 §17.3.3: Special run content elements.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Regular text before special chars
    p.add_run("Before")

    # w:sym — symbol character (Wingdings char F0FC = checkmark)
    r_sym = make_element("w:r")
    sym = make_element("w:sym", {"w:font": "Wingdings", "w:char": "F0FC"})
    r_sym.append(sym)
    p._p.append(r_sym)

    # w:noBreakHyphen
    r_nbh = make_element("w:r")
    nbh = make_element("w:noBreakHyphen")
    r_nbh.append(nbh)
    p._p.append(r_nbh)

    # w:softHyphen
    r_sh = make_element("w:r")
    sh = make_element("w:softHyphen")
    r_sh.append(sh)
    p._p.append(r_sh)

    # w:cr — carriage return
    r_cr = make_element("w:r")
    cr = make_element("w:cr")
    r_cr.append(cr)
    p._p.append(r_cr)

    # w:lastRenderedPageBreak
    r_lrpb = make_element("w:r")
    lrpb = make_element("w:lastRenderedPageBreak")
    r_lrpb.append(lrpb)
    p._p.append(r_lrpb)

    p.add_run("After")

    save_fixture("fields-deep", "special-characters", doc, {
        "name": "special-characters",
        "spec_ref": "ISO 29500-1 §17.3.3",
        "description": "Paragraph with w:sym, w:noBreakHyphen, w:softHyphen, w:cr, w:lastRenderedPageBreak",
        "expected_behavior": "sym -> OpaqueInline (Unknown), noBreakHyphen/softHyphen/lastRenderedPageBreak -> Decoration, cr -> HardBreak",
        "current_status": "SUPPORTED — each element classified per word_ir.rs rules",
    })


def make_styleref_field() -> None:
    """STYLEREF field referencing Heading1 style.

    ISO 29500-1 §17.16.36: STYLEREF retrieves text formatted with
    a specified style.
    """
    doc = Document()

    # A heading paragraph
    doc.add_heading("Chapter One Title", level=1)
    doc.add_paragraph("Some body text under the heading.")

    # STYLEREF field
    p = doc.add_paragraph()
    p.add_run("Current heading: ")
    fld = make_element("w:fldSimple", {"w:instr": ' STYLEREF "Heading 1" '})
    r = make_element("w:r")
    t = make_element("w:t")
    t.text = "Chapter One Title"
    r.append(t)
    fld.append(r)
    p._p.append(fld)

    save_fixture("fields-deep", "styleref-field", doc, {
        "name": "styleref-field",
        "spec_ref": "ISO 29500-1 §17.16.36",
        "description": "STYLEREF fldSimple referencing Heading 1 style",
        "expected_behavior": "FieldKind::Simple with instruction containing 'STYLEREF'",
        "current_status": "SUPPORTED — fldSimple parsed for any instruction type",
    })


# =========================================================================
# STYLES (ISO 29500-1 §17.7)
# =========================================================================

def make_styles_fixtures() -> None:
    print("\n── Styles ──")
    make_style_doc_defaults()
    make_style_based_on_chain()
    make_style_char_overrides_para()
    make_style_toggle_xor_hierarchy()
    make_style_normal_default()
    make_style_table_conditional()
    make_style_numbering_from_style()
    make_style_linked_styles()


def make_style_doc_defaults() -> None:
    """docDefaults -- document-default run and paragraph properties.

    ISO 29500-1 §17.7.4.17: docDefaults/rPrDefault/rPr establishes the
    base run formatting for the entire document. When no style overrides
    a property, docDefaults values should be used.
    """
    doc = Document()

    # Set document defaults via styles.xml: font=Georgia, size=28 (14pt)
    styles_element = doc.styles.element

    # Remove existing docDefaults if any and replace
    existing_defaults = styles_element.find(w("docDefaults"))
    if existing_defaults is not None:
        styles_element.remove(existing_defaults)

    doc_defaults = make_element("w:docDefaults")
    rpr_default = make_element("w:rPrDefault")
    rpr = make_element("w:rPr")
    rpr.append(make_element("w:rFonts", {"w:ascii": "Georgia", "w:hAnsi": "Georgia"}))
    rpr.append(make_element("w:sz", {"w:val": "28"}))  # 14pt
    rpr.append(make_element("w:color", {"w:val": "333333"}))
    rpr_default.append(rpr)
    doc_defaults.append(rpr_default)

    ppr_default = make_element("w:pPrDefault")
    ppr = make_element("w:pPr")
    ppr.append(make_element("w:spacing", {
        "w:after": "200",
        "w:line": "276",
        "w:lineRule": "auto",
    }))
    ppr_default.append(ppr)
    doc_defaults.append(ppr_default)

    # Insert docDefaults as the first child of styles
    styles_element.insert(0, doc_defaults)

    # Add a plain paragraph with no style override -- should get docDefaults
    p1 = doc.add_paragraph("This paragraph should get Georgia 14pt and color #333333 from docDefaults.")

    # Add a paragraph with a style that overrides font but not size
    override_style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "FontOverride"})
    override_style.append(make_element("w:name", {"w:val": "Font Override"}))
    style_rpr = make_element("w:rPr")
    style_rpr.append(make_element("w:rFonts", {"w:ascii": "Courier New", "w:hAnsi": "Courier New"}))
    override_style.append(style_rpr)
    styles_element.append(override_style)

    p2 = doc.add_paragraph("This paragraph should get Courier New (from style) but 14pt and #333333 from docDefaults.")
    p2_el = p2._p
    pPr2 = p2_el.get_or_add_pPr()
    pPr2.append(make_element("w:pStyle", {"w:val": "FontOverride"}))

    save_fixture("styles", "doc-defaults", doc, {
        "name": "doc-defaults",
        "spec_ref": "ISO 29500-1 §17.7.4.17",
        "description": "Document with docDefaults setting Georgia 14pt #333333. Second paragraph overrides font via style.",
        "expected_behavior": "Para 1 runs: font=Georgia, size=28, color=333333. Para 2 runs: font=Courier New, size=28, color=333333.",
        "current_status": "TESTING",
    })


def make_style_based_on_chain() -> None:
    """basedOn chain: C -> B -> A. Each adds a run property.

    ISO 29500-1 §17.7.4.3: A style can inherit from another via basedOn.
    Properties cascade from ancestor to descendant: the most-derived style's
    explicit properties win, and unset properties inherit from the parent.
    """
    doc = Document()
    styles_element = doc.styles.element

    # Style A: font=Times New Roman, size=24 (12pt)
    style_a = make_element("w:style", {"w:type": "paragraph", "w:styleId": "StyleA"})
    style_a.append(make_element("w:name", {"w:val": "Style A"}))
    rpr_a = make_element("w:rPr")
    rpr_a.append(make_element("w:rFonts", {"w:ascii": "Times New Roman", "w:hAnsi": "Times New Roman"}))
    rpr_a.append(make_element("w:sz", {"w:val": "24"}))
    style_a.append(rpr_a)
    styles_element.append(style_a)

    # Style B: basedOn A, adds bold, overrides size to 28
    style_b = make_element("w:style", {"w:type": "paragraph", "w:styleId": "StyleB"})
    style_b.append(make_element("w:name", {"w:val": "Style B"}))
    style_b.append(make_element("w:basedOn", {"w:val": "StyleA"}))
    rpr_b = make_element("w:rPr")
    rpr_b.append(make_element("w:b"))
    rpr_b.append(make_element("w:sz", {"w:val": "28"}))
    style_b.append(rpr_b)
    styles_element.append(style_b)

    # Style C: basedOn B, adds italic, overrides font to Arial
    style_c = make_element("w:style", {"w:type": "paragraph", "w:styleId": "StyleC"})
    style_c.append(make_element("w:name", {"w:val": "Style C"}))
    style_c.append(make_element("w:basedOn", {"w:val": "StyleB"}))
    rpr_c = make_element("w:rPr")
    rpr_c.append(make_element("w:i"))
    rpr_c.append(make_element("w:rFonts", {"w:ascii": "Arial", "w:hAnsi": "Arial"}))
    style_c.append(rpr_c)
    styles_element.append(style_c)

    # Paragraph using Style C -- should get: Arial, 28, bold+italic
    p = doc.add_paragraph("Style C: Arial (from C), 14pt (from B), bold (from B), italic (from C).")
    pPr = p._p.get_or_add_pPr()
    pPr.append(make_element("w:pStyle", {"w:val": "StyleC"}))

    # Paragraph using Style A -- baseline
    p2 = doc.add_paragraph("Style A: Times New Roman 12pt, no bold, no italic.")
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:pStyle", {"w:val": "StyleA"}))

    save_fixture("styles", "based-on-chain", doc, {
        "name": "based-on-chain",
        "spec_ref": "ISO 29500-1 §17.7.4.3",
        "description": "Three-level basedOn chain: C->B->A. Each level adds/overrides run properties.",
        "expected_behavior": "StyleC paragraph: font=Arial, size=28, bold=on, italic=on. StyleA paragraph: font=Times New Roman, size=24.",
        "current_status": "TESTING",
    })


def make_style_char_overrides_para() -> None:
    """Character style overrides paragraph style run properties.

    ISO 29500-1 §17.7.2: The formatting cascade for runs is:
    direct formatting > character style > paragraph style > docDefaults.
    """
    doc = Document()
    styles_element = doc.styles.element

    # Paragraph style: sets italic and color=FF0000
    para_style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "RedItalicPara"})
    para_style.append(make_element("w:name", {"w:val": "Red Italic Para"}))
    rpr_para = make_element("w:rPr")
    rpr_para.append(make_element("w:i"))
    rpr_para.append(make_element("w:color", {"w:val": "FF0000"}))
    rpr_para.append(make_element("w:sz", {"w:val": "24"}))
    para_style.append(rpr_para)
    styles_element.append(para_style)

    # Character style: overrides color to 0000FF, adds bold
    char_style = make_element("w:style", {"w:type": "character", "w:styleId": "BlueBoldChar"})
    char_style.append(make_element("w:name", {"w:val": "Blue Bold Char"}))
    rpr_char = make_element("w:rPr")
    rpr_char.append(make_element("w:b"))
    rpr_char.append(make_element("w:color", {"w:val": "0000FF"}))
    char_style.append(rpr_char)
    styles_element.append(char_style)

    # Paragraph with RedItalicPara style
    p = doc.add_paragraph(style="RedItalicPara")

    # Run 1: no char style -- gets para style (italic, red, 12pt)
    run1 = p.add_run("Para style only (italic, red). ")

    # Run 2: with BlueBoldChar -- should override color to blue, add bold,
    # but italic still comes from para style
    run2 = p.add_run("Char style overrides (bold, blue, italic from para).")
    rPr2 = run2._r.get_or_add_rPr()
    rPr2.insert(0, make_element("w:rStyle", {"w:val": "BlueBoldChar"}))

    # Run 3: direct formatting overrides everything -- green color
    run3 = p.add_run(" Direct override (green color).")
    rPr3 = run3._r.get_or_add_rPr()
    rPr3.append(make_element("w:color", {"w:val": "00FF00"}))

    save_fixture("styles", "char-overrides-para", doc, {
        "name": "char-overrides-para",
        "spec_ref": "ISO 29500-1 §17.7.2",
        "description": "Formatting cascade: direct > char style > para style. Tests all three levels.",
        "expected_behavior": "Run1: italic=on, color=FF0000. Run2: bold=on (char), color=0000FF (char), italic=on (para via XOR). Run3: color=00FF00 (direct).",
        "current_status": "TESTING",
    })


def make_style_toggle_xor_hierarchy() -> None:
    """Toggle property XOR across style hierarchy levels.

    ISO 29500-1 §17.7.3: Toggle properties (bold, italic, etc.) use XOR
    across hierarchy levels. If para style sets bold=true and char style
    also sets bold=true, the effective value is bold=false (XOR).
    """
    doc = Document()
    styles_element = doc.styles.element

    # Paragraph style: bold=true
    bold_para = make_element("w:style", {"w:type": "paragraph", "w:styleId": "BoldPara2"})
    bold_para.append(make_element("w:name", {"w:val": "Bold Para 2"}))
    rpr_bp = make_element("w:rPr")
    rpr_bp.append(make_element("w:b"))
    bold_para.append(rpr_bp)
    styles_element.append(bold_para)

    # Character style: bold=true
    bold_char = make_element("w:style", {"w:type": "character", "w:styleId": "BoldChar2"})
    bold_char.append(make_element("w:name", {"w:val": "Bold Char 2"}))
    rpr_bc = make_element("w:rPr")
    rpr_bc.append(make_element("w:b"))
    bold_char.append(rpr_bc)
    styles_element.append(bold_char)

    # Character style: bold=false (explicit off)
    unbold_char = make_element("w:style", {"w:type": "character", "w:styleId": "UnboldChar"})
    unbold_char.append(make_element("w:name", {"w:val": "Unbold Char"}))
    rpr_ub = make_element("w:rPr")
    rpr_ub.append(make_element("w:b", {"w:val": "0"}))
    unbold_char.append(rpr_ub)
    styles_element.append(unbold_char)

    # Paragraph with BoldPara2 style
    p = doc.add_paragraph(style="BoldPara2")

    # Run 1: para=bold, char=bold -- XOR -> effective false
    run1 = p.add_run("Para bold + char bold = XOR false. ")
    rPr1 = run1._r.get_or_add_rPr()
    rPr1.insert(0, make_element("w:rStyle", {"w:val": "BoldChar2"}))

    # Run 2: para=bold, no char style -- should be bold
    run2 = p.add_run("Para bold only = true. ")

    # Run 3: para=bold, char=bold-off -- XOR -> on XOR off = on
    run3 = p.add_run("Para bold + char unbold = XOR true.")
    rPr3 = run3._r.get_or_add_rPr()
    rPr3.insert(0, make_element("w:rStyle", {"w:val": "UnboldChar"}))

    save_fixture("styles", "toggle-xor-hierarchy", doc, {
        "name": "toggle-xor-hierarchy",
        "spec_ref": "ISO 29500-1 §17.7.3",
        "description": "Toggle XOR across hierarchy: para bold + char bold = XOR false; para bold + char unbold = XOR true.",
        "expected_behavior": "Run1: bold=off (XOR true^true=false). Run2: bold=on. Run3: bold=on (XOR true^false=true).",
        "current_status": "TESTING",
    })


def make_style_normal_default() -> None:
    """Normal style as the implicit default paragraph style.

    ISO 29500-1 §17.7.4.17: When a paragraph has no pStyle, it implicitly
    uses the default paragraph style (typically 'Normal'). Properties from
    Normal should apply.
    """
    doc = Document()
    styles_element = doc.styles.element

    # Ensure Normal style has specific properties
    normal_style = None
    for child in styles_element:
        if child.tag == w("style"):
            sid = child.get(w("styleId"))
            if sid == "Normal":
                normal_style = child
                break

    if normal_style is None:
        normal_style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "Normal", "w:default": "1"})
        normal_style.append(make_element("w:name", {"w:val": "Normal"}))
        styles_element.append(normal_style)

    # Set paragraph properties on Normal: alignment=both (justify)
    existing_ppr = normal_style.find(w("pPr"))
    if existing_ppr is not None:
        normal_style.remove(existing_ppr)
    ppr_normal = make_element("w:pPr")
    ppr_normal.append(make_element("w:jc", {"w:val": "both"}))
    normal_style.append(ppr_normal)

    # Set run properties on Normal: font=Calibri, size=22 (11pt)
    existing_rpr = normal_style.find(w("rPr"))
    if existing_rpr is not None:
        normal_style.remove(existing_rpr)
    rpr_normal = make_element("w:rPr")
    rpr_normal.append(make_element("w:rFonts", {"w:ascii": "Calibri", "w:hAnsi": "Calibri"}))
    rpr_normal.append(make_element("w:sz", {"w:val": "22"}))
    normal_style.append(rpr_normal)

    # Paragraph 1: unstyled (should get Normal properties)
    p1 = doc.add_paragraph("Unstyled paragraph -- should inherit Normal style: Calibri 11pt, justify.")

    # Paragraph 2: explicit Normal
    p2 = doc.add_paragraph("Explicit Normal style.")
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:pStyle", {"w:val": "Normal"}))

    save_fixture("styles", "normal-default", doc, {
        "name": "normal-default",
        "spec_ref": "ISO 29500-1 §17.7.4.17",
        "description": "Normal style defines Calibri 11pt with justify alignment. Unstyled paragraphs should inherit.",
        "expected_behavior": "Both paragraphs: font=Calibri, size=22, alignment=justify (both).",
        "current_status": "TESTING",
    })


def make_style_table_conditional() -> None:
    """Table style with conditional formatting (firstRow, lastRow, band).

    ISO 29500-1 §17.7.6: Table styles can define conditional formatting
    for first row, last row, banded rows, etc. via tblStylePr elements.
    """
    doc = Document()
    styles_element = doc.styles.element

    # Create a table style with conditional formatting
    tbl_style = make_element("w:style", {"w:type": "table", "w:styleId": "ConditionalTable"})
    tbl_style.append(make_element("w:name", {"w:val": "Conditional Table"}))

    # Base table formatting: thin borders
    tblPr = make_element("w:tblPr")
    tbl_borders = make_element("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "4", "w:color": "999999",
        }))
    tblPr.append(tbl_borders)
    tbl_style.append(tblPr)

    # Default cell shading: no fill
    tcPr_base = make_element("w:tcPr")
    tcPr_base.append(make_element("w:shd", {"w:val": "clear", "w:fill": "auto"}))
    tbl_style.append(tcPr_base)

    # Conditional: firstRow -- bold text, blue shading
    first_row = make_element("w:tblStylePr", {"w:type": "firstRow"})
    fr_rpr = make_element("w:rPr")
    fr_rpr.append(make_element("w:b"))
    first_row.append(fr_rpr)
    fr_tcPr = make_element("w:tcPr")
    fr_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "4472C4"}))
    first_row.append(fr_tcPr)
    tbl_style.append(first_row)

    # Conditional: band1Horz -- light gray shading
    band1 = make_element("w:tblStylePr", {"w:type": "band1Horz"})
    b1_tcPr = make_element("w:tcPr")
    b1_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "D9E2F3"}))
    band1.append(b1_tcPr)
    tbl_style.append(band1)

    styles_element.append(tbl_style)

    # Create table referencing the style
    doc.add_paragraph("Table with conditional formatting style.")
    tbl = doc.add_table(rows=4, cols=3)
    tbl_element = tbl._tbl

    # Apply style
    tblPr_ref = tbl_element.tblPr
    if tblPr_ref is None:
        tblPr_ref = make_element("w:tblPr")
        tbl_element.insert(0, tblPr_ref)
    tblPr_ref.insert(0, make_element("w:tblStyle", {"w:val": "ConditionalTable"}))
    # Enable table look flags
    tblPr_ref.append(make_element("w:tblLook", {
        "w:val": "04A0",
        "w:firstRow": "1",
        "w:lastRow": "0",
        "w:firstColumn": "0",
        "w:lastColumn": "0",
        "w:noHBand": "0",
        "w:noVBand": "1",
    }))

    labels = [
        ["Header A", "Header B", "Header C"],
        ["Data A1", "Data B1", "Data C1"],
        ["Data A2", "Data B2", "Data C2"],
        ["Data A3", "Data B3", "Data C3"],
    ]
    for r_idx, row_labels in enumerate(labels):
        for c_idx, label in enumerate(row_labels):
            tbl.cell(r_idx, c_idx).text = label

    save_fixture("styles", "table-conditional", doc, {
        "name": "table-conditional",
        "spec_ref": "ISO 29500-1 §17.7.6",
        "description": "Table style with firstRow (bold+blue) and band1Horz (light gray) conditional formatting.",
        "expected_behavior": "First row cells: shading fill=4472C4. Band rows: shading fill=D9E2F3. Table borders from style.",
        "current_status": "TESTING",
    })


def make_style_numbering_from_style() -> None:
    """Paragraph style with numPr -- style-based numbering.

    ISO 29500-1 §17.7.4.14: A paragraph style can include numPr in its
    pPr, associating paragraphs of that style with a numbering definition.
    """
    doc = Document()
    styles_element = doc.styles.element

    # Create a numbering definition
    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}"
             xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # Create paragraph style with numPr
    num_style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "NumberedClause"})
    num_style.append(make_element("w:name", {"w:val": "Numbered Clause"}))
    ppr_ns = make_element("w:pPr")
    numPr = make_element("w:numPr")
    numPr.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr.append(make_element("w:numId", {"w:val": "1"}))
    ppr_ns.append(numPr)
    num_style.append(ppr_ns)
    styles_element.append(num_style)

    # Add paragraphs using the style -- they should auto-number
    for text in ["First clause", "Second clause", "Third clause"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        pPr.append(make_element("w:pStyle", {"w:val": "NumberedClause"}))

    save_fixture("styles", "numbering-from-style", doc, {
        "name": "numbering-from-style",
        "spec_ref": "ISO 29500-1 §17.7.4.14",
        "description": "Paragraph style with numPr. Paragraphs using this style should auto-number.",
        "expected_behavior": "Paragraphs should have numbering info: 1. First clause, 2. Second clause, 3. Third clause.",
        "current_status": "TESTING",
    })


def make_style_linked_styles() -> None:
    """Linked styles -- paragraph + character style pair.

    ISO 29500-1 §17.7.4.6: The link element in a style definition creates
    a bidirectional link between a paragraph style and a character style.
    When a linked character style is applied to a full paragraph, it should
    activate the paragraph style. When applied to a run, the char style applies.
    """
    doc = Document()
    styles_element = doc.styles.element

    # Paragraph style: Heading-like, with link to a character style
    linked_para = make_element("w:style", {"w:type": "paragraph", "w:styleId": "LinkedPara"})
    linked_para.append(make_element("w:name", {"w:val": "Linked Para"}))
    linked_para.append(make_element("w:link", {"w:val": "LinkedChar"}))
    rpr_lp = make_element("w:rPr")
    rpr_lp.append(make_element("w:b"))
    rpr_lp.append(make_element("w:sz", {"w:val": "32"}))
    rpr_lp.append(make_element("w:color", {"w:val": "2E74B5"}))
    linked_para.append(rpr_lp)
    ppr_lp = make_element("w:pPr")
    ppr_lp.append(make_element("w:jc", {"w:val": "center"}))
    linked_para.append(ppr_lp)
    styles_element.append(linked_para)

    # Character style: linked to the paragraph style
    linked_char = make_element("w:style", {"w:type": "character", "w:styleId": "LinkedChar"})
    linked_char.append(make_element("w:name", {"w:val": "Linked Char"}))
    linked_char.append(make_element("w:link", {"w:val": "LinkedPara"}))
    rpr_lc = make_element("w:rPr")
    rpr_lc.append(make_element("w:b"))
    rpr_lc.append(make_element("w:sz", {"w:val": "32"}))
    rpr_lc.append(make_element("w:color", {"w:val": "2E74B5"}))
    linked_char.append(rpr_lc)
    styles_element.append(linked_char)

    # Paragraph using LinkedPara style
    p1 = doc.add_paragraph("Full paragraph with LinkedPara style: bold, 16pt, blue, centered.")
    pPr1 = p1._p.get_or_add_pPr()
    pPr1.append(make_element("w:pStyle", {"w:val": "LinkedPara"}))

    # Paragraph with a run using the linked character style
    p2 = doc.add_paragraph()
    p2.add_run("Normal text. ")
    run_linked = p2.add_run("This run uses LinkedChar style.")
    rPr_linked = run_linked._r.get_or_add_rPr()
    rPr_linked.insert(0, make_element("w:rStyle", {"w:val": "LinkedChar"}))

    save_fixture("styles", "linked-styles", doc, {
        "name": "linked-styles",
        "spec_ref": "ISO 29500-1 §17.7.4.6",
        "description": "Linked paragraph+character style pair. Tests both full-paragraph and run-level application.",
        "expected_behavior": "Para1: bold=on, size=32, color=2E74B5, align=center. Para2 run2: bold=on, size=32, color=2E74B5 (char style).",
        "current_status": "TESTING",
    })


# =========================================================================
# CONTENT TYPES (ISO 29500-1 §17.3.3, §17.5.2, §22)
# =========================================================================

M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
MC_NS_CONTENT = "http://schemas.openxmlformats.org/markup-compatibility/2006"


def _make_1px_png() -> bytes:
    """Create a minimal 1x1 red PNG (valid)."""
    import struct, zlib
    width, height = 1, 1
    raw_data = b'\x00\xff\x00\x00\xff'
    compressed = zlib.compress(raw_data)

    def chunk(chunk_type: bytes, data: bytes) -> bytes:
        c = chunk_type + data
        crc = struct.pack(">I", zlib.crc32(c) & 0xffffffff)
        return struct.pack(">I", len(data)) + c + crc

    sig = b'\x89PNG\r\n\x1a\n'
    ihdr_data = struct.pack(">IIBBBBB", width, height, 8, 6, 0, 0, 0)
    return sig + chunk(b'IHDR', ihdr_data) + chunk(b'IDAT', compressed) + chunk(b'IEND', b'')


def make_content_fixtures() -> None:
    print("\n── Content Types ──")
    make_inline_image()
    make_anchored_image()
    make_block_sdt()
    make_inline_sdt()
    make_inline_math()
    make_display_math()
    make_ruby_text()
    make_alternate_content()


def make_inline_image() -> None:
    """Inline image (§20.4.2.8 wp:inline)."""
    doc = Document()
    import io
    png_bytes = _make_1px_png()
    p = doc.add_paragraph("Text before image.")
    run = p.add_run()
    run.add_picture(io.BytesIO(png_bytes), width=Inches(0.5))
    p.add_run(" Text after image.")

    save_fixture("content", "inline-image", doc, {
        "name": "inline-image",
        "spec_ref": "ISO 29500-1 §20.4.2.8",
        "description": "Paragraph with an inline image (wp:inline inside w:drawing)",
        "expected_behavior": "OpaqueInline with kind Drawing should be present",
        "current_status": "SUPPORTED",
    })


def make_anchored_image() -> None:
    """Anchored/floating image (§20.4.2.3 wp:anchor)."""
    doc = Document()
    import io
    png_bytes = _make_1px_png()
    p = doc.add_paragraph("Paragraph with an anchored (floating) image.")
    run = p.add_run()
    run.add_picture(io.BytesIO(png_bytes), width=Inches(0.5))

    save_fixture("content", "anchored-image", doc, {
        "name": "anchored-image",
        "spec_ref": "ISO 29500-1 §20.4.2.3",
        "description": "Paragraph with an image (verifies Drawing opaque kind)",
        "expected_behavior": "OpaqueInline with kind Drawing should be present",
        "current_status": "SUPPORTED",
    })


def make_block_sdt() -> None:
    """Block-level SDT (§17.5.2.32)."""
    doc = Document()
    doc.add_paragraph("Paragraph before block SDT.")

    body = doc.element.body
    sdt = make_element("w:sdt")
    sdt_pr = make_element("w:sdtPr")
    sdt_alias = make_element("w:alias", {"w:val": "BlockControl"})
    sdt_pr.append(sdt_alias)
    sdt.append(sdt_pr)

    sdt_content = make_element("w:sdtContent")
    inner_p = make_element("w:p")
    inner_r = make_element("w:r")
    inner_t = make_element("w:t")
    inner_t.text = "Content inside block SDT"
    inner_r.append(inner_t)
    inner_p.append(inner_r)
    sdt_content.append(inner_p)
    sdt.append(sdt_content)

    body.append(sdt)
    doc.add_paragraph("Paragraph after block SDT.")

    save_fixture("content", "block-sdt", doc, {
        "name": "block-sdt",
        "spec_ref": "ISO 29500-1 §17.5.2.32",
        "description": "Block-level SDT wrapping a paragraph",
        "expected_behavior": "Content within the SDT should be accessible",
        "current_status": "SUPPORTED",
    })


def make_inline_sdt() -> None:
    """Inline SDT (§17.5.2.31)."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Text before ")

    sdt = OxmlElement("w:sdt")
    sdt_pr = make_element("w:sdtPr")
    sdt_alias = make_element("w:alias", {"w:val": "InlineControl"})
    sdt_pr.append(sdt_alias)
    sdt.append(sdt_pr)

    sdt_content = make_element("w:sdtContent")
    inner_r = make_element("w:r")
    inner_t = make_element("w:t")
    inner_t.text = "SDT content"
    inner_r.append(inner_t)
    sdt_content.append(inner_r)
    sdt.append(sdt_content)

    p._p.append(sdt)
    p.add_run(" text after SDT.")

    save_fixture("content", "inline-sdt", doc, {
        "name": "inline-sdt",
        "spec_ref": "ISO 29500-1 §17.5.2.31",
        "description": "Inline SDT (content control) within a paragraph",
        "expected_behavior": "OpaqueInline with kind Sdt should be present",
        "current_status": "SUPPORTED",
    })


def make_inline_math() -> None:
    """Math equation -- inline (§22): m:oMath inside a paragraph."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("The equation is ")

    omath = OxmlElement("m:oMath")

    mr = OxmlElement("m:r")
    mt = OxmlElement("m:t")
    mt.text = "x+y=z"
    mr.append(mt)
    omath.append(mr)
    p._p.append(omath)

    p.add_run(" in the text.")

    save_fixture("content", "inline-math", doc, {
        "name": "inline-math",
        "spec_ref": "ISO 29500-1 §22",
        "description": "Paragraph with an inline math equation (m:oMath)",
        "expected_behavior": "OpaqueInline with kind Omml should be present",
        "current_status": "SUPPORTED",
    })


def make_display_math() -> None:
    """Math equation -- display (§22): m:oMathPara."""
    doc = Document()
    doc.add_paragraph("Text before display math.")

    p = doc.add_paragraph()
    omath_para = OxmlElement("m:oMathPara")

    omath = OxmlElement("m:oMath")
    mr = OxmlElement("m:r")
    mt = OxmlElement("m:t")
    mt.text = "a^2+b^2=c^2"
    mr.append(mt)
    omath.append(mr)
    omath_para.append(omath)
    p._p.append(omath_para)

    doc.add_paragraph("Text after display math.")

    save_fixture("content", "display-math", doc, {
        "name": "display-math",
        "spec_ref": "ISO 29500-1 §22",
        "description": "Display math equation using m:oMathPara",
        "expected_behavior": "Paragraph should contain OpaqueInline with kind Omml",
        "current_status": "SUPPORTED",
    })


def make_ruby_text() -> None:
    """Ruby text (§17.3.3.25): East Asian annotation."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Text with ruby: ")

    ruby_run = OxmlElement("w:r")
    ruby = OxmlElement("w:ruby")
    ruby_pr = make_element("w:rubyPr")
    ruby_align = make_element("w:rubyAlign", {"w:val": "center"})
    ruby_pr.append(ruby_align)
    ruby.append(ruby_pr)

    rt = make_element("w:rt")
    rt_r = make_element("w:r")
    rt_t = make_element("w:t")
    rt_t.text = "かん"
    rt_r.append(rt_t)
    rt.append(rt_r)
    ruby.append(rt)

    ruby_base = make_element("w:rubyBase")
    base_r = make_element("w:r")
    base_t = make_element("w:t")
    base_t.text = "漢"
    base_r.append(base_t)
    ruby_base.append(base_r)
    ruby.append(ruby_base)

    ruby_run.append(ruby)
    p._p.append(ruby_run)

    p.add_run(" end.")

    save_fixture("content", "ruby-text", doc, {
        "name": "ruby-text",
        "spec_ref": "ISO 29500-1 §17.3.3.25",
        "description": "East Asian ruby annotation (w:ruby)",
        "expected_behavior": "OpaqueInline with kind Ruby should be present",
        "current_status": "SUPPORTED",
    })


def make_alternate_content() -> None:
    """AlternateContent (mc:AlternateContent): Choice/Fallback."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Before MC block. ")

    _MC = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    _WPS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
    mc_nsmap = {"mc": _MC, "w": W, "wps": _WPS}

    run = OxmlElement("w:r")
    mc_ac = etree.SubElement(run, f"{{{_MC}}}AlternateContent", nsmap=mc_nsmap)

    mc_choice = etree.SubElement(mc_ac, f"{{{_MC}}}Choice")
    mc_choice.set("Requires", "wps")
    choice_drawing = OxmlElement("w:drawing")
    mc_choice.append(choice_drawing)

    mc_fallback = etree.SubElement(mc_ac, f"{{{_MC}}}Fallback")
    fallback_pict = OxmlElement("w:pict")
    mc_fallback.append(fallback_pict)

    p._p.append(run)

    p.add_run(" After MC block.")

    save_fixture("content", "alternate-content", doc, {
        "name": "alternate-content",
        "spec_ref": "ECMA-376 Part 3 (MCE)",
        "description": "mc:AlternateContent with Choice (drawing) and Fallback (pict)",
        "expected_behavior": "Should be handled as OpaqueInline with Drawing kind",
        "current_status": "SUPPORTED",
    })


# =========================================================================
# ADDITIONAL NUMBERING (ISO 29500-1 §17.9) -- Deep levels & formats
# =========================================================================

def make_additional_numbering_fixtures() -> None:
    print("\n── Additional Numbering ──")
    make_all_nine_levels()
    make_number_format_types()
    make_start_at_override()
    make_shared_abstract_num()


def make_all_nine_levels() -> None:
    """All 9 numbering levels (0-8)."""
    doc = Document()

    levels_xml = ""
    for i in range(9):
        pct_refs = ".".join(f"%{j+1}" for j in range(i+1))
        levels_xml += f"""
    <w:lvl w:ilvl="{i}">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="{pct_refs}."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>"""

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    {levels_xml}
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    for level in range(9):
        p = doc.add_paragraph(f"Level {level} item")
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": str(level)}))
        numPr.append(make_element("w:numId", {"w:val": "1"}))
        pPr.append(numPr)

    save_fixture("numbering", "all-nine-levels", doc, {
        "name": "all-nine-levels",
        "spec_ref": "ISO 29500-1 §17.9.6",
        "description": "Multi-level list exercising all 9 levels (0-8)",
        "expected_behavior": "Each level synthesizes correct prefix: 1. / 1.1. / ... / 1.1.1.1.1.1.1.1.1.",
        "current_status": "SUPPORTED",
    })


def make_number_format_types() -> None:
    """Various numFmt types: bullet, lowerRoman, upperRoman, upperLetter."""
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="bullet"/>
      <w:lvlText w:val="&#xF0B7;"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:abstractNum w:abstractNumId="1">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerRoman"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:abstractNum w:abstractNumId="2">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="upperRoman"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:abstractNum w:abstractNumId="3">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="upperLetter"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>
  <w:num w:numId="2"><w:abstractNumId w:val="1"/></w:num>
  <w:num w:numId="3"><w:abstractNumId w:val="2"/></w:num>
  <w:num w:numId="4"><w:abstractNumId w:val="3"/></w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # Bullet items (numId=1)
    for text in ["Bullet item one", "Bullet item two"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "0"}))
        numPr.append(make_element("w:numId", {"w:val": "1"}))
        pPr.append(numPr)

    doc.add_paragraph("---")

    # Lower roman items (numId=2)
    for text in ["Lower roman one", "Lower roman two", "Lower roman three"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "0"}))
        numPr.append(make_element("w:numId", {"w:val": "2"}))
        pPr.append(numPr)

    doc.add_paragraph("---")

    # Upper roman items (numId=3)
    for text in ["Upper roman one", "Upper roman two"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "0"}))
        numPr.append(make_element("w:numId", {"w:val": "3"}))
        pPr.append(numPr)

    doc.add_paragraph("---")

    # Upper letter items (numId=4)
    for text in ["Upper letter one", "Upper letter two"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "0"}))
        numPr.append(make_element("w:numId", {"w:val": "4"}))
        pPr.append(numPr)

    save_fixture("numbering", "number-format-types", doc, {
        "name": "number-format-types",
        "spec_ref": "ISO 29500-1 §17.9.17",
        "description": "Lists with bullet, lowerRoman, upperRoman, and upperLetter formats",
        "expected_behavior": "bullet=bullet, lowerRoman=i./ii./iii., upperRoman=I./II., upperLetter=A./B.",
        "current_status": "SUPPORTED",
    })


def make_start_at_override() -> None:
    """Start-at override (§17.9.15): List starting at 5."""
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="5"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    for text in ["Starting at five", "Next item", "Third item"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "0"}))
        numPr.append(make_element("w:numId", {"w:val": "1"}))
        pPr.append(numPr)

    save_fixture("numbering", "start-at-override", doc, {
        "name": "start-at-override",
        "spec_ref": "ISO 29500-1 §17.9.15",
        "description": "List with w:start val=5",
        "expected_behavior": "Synthesized text should be 5., 6., 7.",
        "current_status": "SUPPORTED",
    })


def make_shared_abstract_num() -> None:
    """Two lists sharing abstractNumId but different numId overrides."""
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
  <w:num w:numId="2">
    <w:abstractNumId w:val="0"/>
    <w:lvlOverride w:ilvl="0">
      <w:startOverride w:val="10"/>
    </w:lvlOverride>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    for text in ["List A item one", "List A item two"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "0"}))
        numPr.append(make_element("w:numId", {"w:val": "1"}))
        pPr.append(numPr)

    doc.add_paragraph("Separator paragraph.")

    for text in ["List B item one (should be 10)", "List B item two (should be 11)"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "0"}))
        numPr.append(make_element("w:numId", {"w:val": "2"}))
        pPr.append(numPr)

    save_fixture("numbering", "shared-abstract-num", doc, {
        "name": "shared-abstract-num",
        "spec_ref": "ISO 29500-1 §17.9.8",
        "description": "Two lists sharing abstractNumId=0; numId=2 has startOverride=10",
        "expected_behavior": "First list: 1., 2. Second list: 10., 11.",
        "current_status": "SUPPORTED",
    })


# =========================================================================
# Main
# =========================================================================

# =========================================================================
# EDGE CASES — Spacing & Indentation (ISO 29500-1 §17.3.1)
# =========================================================================

def make_edge_case_spacing_fixtures() -> None:
    print("\n── Edge Cases: Spacing & Indentation ──")
    make_contextual_spacing()
    make_spacing_line_rules()
    make_spacing_before_after_zero_vs_absent()
    make_before_after_lines_precedence()
    make_indent_negative_left()


def make_contextual_spacing() -> None:
    """contextualSpacing suppresses spacing between same-style paragraphs.

    ISO 29500-1 §17.3.1.9: When contextualSpacing is true and adjacent
    paragraphs share the same style, before/after spacing is suppressed
    between them. Paragraphs with different styles keep spacing.
    """
    doc = Document()

    # Create two custom paragraph styles
    styles_element = doc.styles.element

    style_a = make_element("w:style", {"w:type": "paragraph", "w:styleId": "StyleA"})
    style_a.append(make_element("w:name", {"w:val": "Style A"}))
    pPr_a = make_element("w:pPr")
    pPr_a.append(make_element("w:spacing", {"w:before": "240", "w:after": "240"}))
    pPr_a.append(make_element("w:contextualSpacing"))
    style_a.append(pPr_a)
    styles_element.append(style_a)

    style_b = make_element("w:style", {"w:type": "paragraph", "w:styleId": "StyleB"})
    style_b.append(make_element("w:name", {"w:val": "Style B"}))
    pPr_b = make_element("w:pPr")
    pPr_b.append(make_element("w:spacing", {"w:before": "480", "w:after": "480"}))
    style_b.append(pPr_b)
    styles_element.append(style_b)

    # P1: StyleA with contextualSpacing (set on style)
    p1 = doc.add_paragraph("P1: StyleA with contextualSpacing", style="StyleA")

    # P2: StyleA with contextualSpacing (same style as P1 — spacing between should be suppressed)
    p2 = doc.add_paragraph("P2: StyleA with contextualSpacing", style="StyleA")

    # P3: StyleB (different style — spacing between P2/P3 should be kept)
    p3 = doc.add_paragraph("P3: StyleB (different style)", style="StyleB")

    save_fixture("edge-cases", "contextual-spacing", doc, {
        "name": "contextual-spacing",
        "spec_ref": "ISO 29500-1 §17.3.1.9",
        "description": "contextualSpacing: P1+P2 (same style, spacing suppressed), P3 (different style, spacing kept)",
        "expected_behavior": "Spacing between P1/P2 logically suppressed. P2/P3 keeps spacing. Domain model may or may not model contextualSpacing directly.",
    })


def make_spacing_line_rules() -> None:
    """Line spacing with different lineRule values.

    ISO 29500-1 §17.3.1.33: The lineRule attribute controls interpretation
    of the line value:
    - auto: value in 240ths of a line (240=single, 360=1.5, 480=double)
    - exact: value in twips (exact height, text may clip)
    - atLeast: value in twips (minimum height, expands if needed)
    """
    doc = Document()

    # Single spacing: line=240, lineRule=auto
    p1 = doc.add_paragraph("Single spaced (auto 240)")
    pPr1 = p1._p.get_or_add_pPr()
    pPr1.append(make_element("w:spacing", {
        "w:line": "240",
        "w:lineRule": "auto",
    }))

    # 1.5 spacing: line=360, lineRule=auto
    p2 = doc.add_paragraph("1.5 spaced (auto 360)")
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:spacing", {
        "w:line": "360",
        "w:lineRule": "auto",
    }))

    # Double spacing: line=480, lineRule=auto
    p3 = doc.add_paragraph("Double spaced (auto 480)")
    pPr3 = p3._p.get_or_add_pPr()
    pPr3.append(make_element("w:spacing", {
        "w:line": "480",
        "w:lineRule": "auto",
    }))

    # Exact 240 twips (12pt) — NOT single spacing
    p4 = doc.add_paragraph("Exact 240 twips (12pt line height)")
    pPr4 = p4._p.get_or_add_pPr()
    pPr4.append(make_element("w:spacing", {
        "w:line": "240",
        "w:lineRule": "exact",
    }))

    # At least 240 twips
    p5 = doc.add_paragraph("At least 240 twips")
    pPr5 = p5._p.get_or_add_pPr()
    pPr5.append(make_element("w:spacing", {
        "w:line": "240",
        "w:lineRule": "atLeast",
    }))

    save_fixture("edge-cases", "spacing-line-rules", doc, {
        "name": "spacing-line-rules",
        "spec_ref": "ISO 29500-1 §17.3.1.33",
        "description": "Paragraphs with various line spacing rules: auto (240/360/480), exact (240), atLeast (240)",
        "expected_behavior": "Domain model should preserve both the line value AND the lineRule correctly",
    })


def make_spacing_before_after_zero_vs_absent() -> None:
    """Spacing before/after: explicit zero vs absent (inherit).

    ISO 29500-1 §17.3.1.33: before="0" explicitly sets zero spacing,
    overriding any style value. Absent before attribute means inherit
    from style (could be any value).
    """
    doc = Document()

    # Create a style with before=240, after=240
    styles_element = doc.styles.element
    style_spaced = make_element("w:style", {"w:type": "paragraph", "w:styleId": "SpacedStyle"})
    style_spaced.append(make_element("w:name", {"w:val": "Spaced Style"}))
    pPr_s = make_element("w:pPr")
    pPr_s.append(make_element("w:spacing", {"w:before": "240", "w:after": "240"}))
    style_spaced.append(pPr_s)
    styles_element.append(style_spaced)

    # P1: SpacedStyle, no direct spacing override — should inherit before=240, after=240
    p1 = doc.add_paragraph("P1: SpacedStyle, no direct override (inherits before=240, after=240)", style="SpacedStyle")

    # P2: SpacedStyle, direct before="0" — should have before=0 (overrides style), after=240 (inherit)
    p2 = doc.add_paragraph("P2: SpacedStyle, direct before=0 (overrides to 0)", style="SpacedStyle")
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:spacing", {"w:before": "0"}))

    # P3: SpacedStyle, direct after="0" — should have before=240 (inherit), after=0 (overrides)
    p3 = doc.add_paragraph("P3: SpacedStyle, direct after=0 (overrides to 0)", style="SpacedStyle")
    pPr3 = p3._p.get_or_add_pPr()
    pPr3.append(make_element("w:spacing", {"w:after": "0"}))

    # P4: SpacedStyle, direct before="0", after="0" — both overridden to 0
    p4 = doc.add_paragraph("P4: SpacedStyle, direct before=0 AND after=0", style="SpacedStyle")
    pPr4 = p4._p.get_or_add_pPr()
    pPr4.append(make_element("w:spacing", {"w:before": "0", "w:after": "0"}))

    save_fixture("edge-cases", "spacing-zero-vs-absent", doc, {
        "name": "spacing-zero-vs-absent",
        "spec_ref": "ISO 29500-1 §17.3.1.33",
        "description": "Spacing before/after: explicit 0 vs absent (inherit from style)",
        "expected_behavior": "before=0 overrides style's before=240. Absent before inherits style value.",
    })


def make_before_after_lines_precedence() -> None:
    """beforeLines/afterLines precedence over before/after.

    ISO 29500-1 §17.3.1.33: If beforeLines is specified alongside before,
    the beforeLines takes precedence. beforeLines is in hundredths of a line
    (100 = one line).
    """
    doc = Document()

    # P1: both before and beforeLines — beforeLines should take precedence
    p1 = doc.add_paragraph("P1: before=480 AND beforeLines=100 (beforeLines wins)")
    pPr1 = p1._p.get_or_add_pPr()
    pPr1.append(make_element("w:spacing", {
        "w:before": "480",
        "w:beforeLines": "100",
    }))

    # P2: only before — straightforward
    p2 = doc.add_paragraph("P2: before=480 only")
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:spacing", {
        "w:before": "480",
    }))

    # P3: both after and afterLines
    p3 = doc.add_paragraph("P3: after=480 AND afterLines=100 (afterLines wins)")
    pPr3 = p3._p.get_or_add_pPr()
    pPr3.append(make_element("w:spacing", {
        "w:after": "480",
        "w:afterLines": "100",
    }))

    # P4: only beforeLines (no before) — still valid
    p4 = doc.add_paragraph("P4: beforeLines=200 only (2 lines)")
    pPr4 = p4._p.get_or_add_pPr()
    pPr4.append(make_element("w:spacing", {
        "w:beforeLines": "200",
    }))

    save_fixture("edge-cases", "spacing-before-after-lines", doc, {
        "name": "spacing-before-after-lines",
        "spec_ref": "ISO 29500-1 §17.3.1.33",
        "description": "beforeLines/afterLines precedence over before/after",
        "expected_behavior": "If beforeLines is present alongside before, beforeLines takes precedence per spec. We may not implement beforeLines at all.",
    })


def make_indent_negative_left() -> None:
    """Negative left indent (outdent) and explicit zero vs absent.

    ISO 29500-1 §17.3.1.12: left indent can be negative (outdent), zero
    (explicit override), or absent (inherit from style).
    """
    doc = Document()

    # Create a style with left indent
    styles_element = doc.styles.element
    style_indented = make_element("w:style", {"w:type": "paragraph", "w:styleId": "IndentedStyle"})
    style_indented.append(make_element("w:name", {"w:val": "Indented Style"}))
    pPr_i = make_element("w:pPr")
    pPr_i.append(make_element("w:ind", {"w:left": "720"}))
    style_indented.append(pPr_i)
    styles_element.append(style_indented)

    # P1: Negative left indent (outdent) — valid per spec
    p1 = doc.add_paragraph("P1: Negative left indent (-720 twips = -0.5 inch outdent)")
    pPr1 = p1._p.get_or_add_pPr()
    pPr1.append(make_element("w:ind", {"w:left": "-720"}))

    # P2: Explicit zero left indent (overrides style)
    p2 = doc.add_paragraph("P2: Explicit zero left indent (overrides style)", style="IndentedStyle")
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:ind", {"w:left": "0"}))

    # P3: No direct indent — inherit from IndentedStyle (left=720)
    p3 = doc.add_paragraph("P3: No direct indent (inherits left=720 from style)", style="IndentedStyle")

    # P4: Negative left with positive first-line (outdent body, indent first line)
    p4 = doc.add_paragraph("P4: Negative left (-360) with firstLine=720")
    pPr4 = p4._p.get_or_add_pPr()
    pPr4.append(make_element("w:ind", {"w:left": "-360", "w:firstLine": "720"}))

    save_fixture("edge-cases", "indent-negative-left", doc, {
        "name": "indent-negative-left",
        "spec_ref": "ISO 29500-1 §17.3.1.12",
        "description": "Negative left indent, explicit zero override, and inherit from style",
        "expected_behavior": "Negative indent values are valid signed integers. Explicit 0 overrides style. Absent inherits.",
    })


# =========================================================================
# EDGE CASES — XML Parsing (ISO 29500-1 §22.9.2, §17.3)
# =========================================================================

def make_edge_case_parsing_fixtures() -> None:
    print("\n── Edge Cases: XML Parsing ──")
    make_st_onoff_boolean_forms()
    make_firstline_vs_hanging()
    make_vmerge_absent_value()
    make_del_text_variants()


def make_st_onoff_boolean_forms() -> None:
    """ST_OnOff boolean value forms — all valid representations.

    ISO 29500-1 §22.9.2.7: The ST_OnOff simple type accepts:
      ON:  "true", "on", "1", or absent val attribute (bare element)
      OFF: "false", "off", "0"
      Absent element entirely = inherit (not present in marks)

    Tests bold, italic, and strike with every form to ensure consistency.
    """
    doc = Document()

    # --- Bold tests ---

    # 1. Bare <w:b/> (no val attribute) = bold ON
    p_bare = doc.add_paragraph("Bold bare element")
    r_bare = p_bare.runs[0]
    rPr = r_bare._r.get_or_add_rPr()
    # Clear any existing bold python-docx may have set
    for existing_b in rPr.findall(w("b")):
        rPr.remove(existing_b)
    b = make_element("w:b")
    rPr.append(b)

    # 2. <w:b w:val="1"/> = bold ON
    p_val1 = doc.add_paragraph("Bold val=1")
    r_val1 = p_val1.runs[0]
    rPr = r_val1._r.get_or_add_rPr()
    for existing_b in rPr.findall(w("b")):
        rPr.remove(existing_b)
    b = make_element("w:b", {"w:val": "1"})
    rPr.append(b)

    # 3. <w:b w:val="true"/> = bold ON
    p_true = doc.add_paragraph("Bold val=true")
    r_true = p_true.runs[0]
    rPr = r_true._r.get_or_add_rPr()
    for existing_b in rPr.findall(w("b")):
        rPr.remove(existing_b)
    b = make_element("w:b", {"w:val": "true"})
    rPr.append(b)

    # 4. <w:b w:val="on"/> = bold ON
    p_on = doc.add_paragraph("Bold val=on")
    r_on = p_on.runs[0]
    rPr = r_on._r.get_or_add_rPr()
    for existing_b in rPr.findall(w("b")):
        rPr.remove(existing_b)
    b = make_element("w:b", {"w:val": "on"})
    rPr.append(b)

    # 5. <w:b w:val="0"/> = bold OFF
    p_val0 = doc.add_paragraph("Bold val=0")
    r_val0 = p_val0.runs[0]
    rPr = r_val0._r.get_or_add_rPr()
    for existing_b in rPr.findall(w("b")):
        rPr.remove(existing_b)
    b = make_element("w:b", {"w:val": "0"})
    rPr.append(b)

    # 6. <w:b w:val="false"/> = bold OFF
    p_false = doc.add_paragraph("Bold val=false")
    r_false = p_false.runs[0]
    rPr = r_false._r.get_or_add_rPr()
    for existing_b in rPr.findall(w("b")):
        rPr.remove(existing_b)
    b = make_element("w:b", {"w:val": "false"})
    rPr.append(b)

    # 7. <w:b w:val="off"/> = bold OFF
    p_off = doc.add_paragraph("Bold val=off")
    r_off = p_off.runs[0]
    rPr = r_off._r.get_or_add_rPr()
    for existing_b in rPr.findall(w("b")):
        rPr.remove(existing_b)
    b = make_element("w:b", {"w:val": "off"})
    rPr.append(b)

    # 8. No <w:b> element at all = inherit (no bold mark present)
    p_absent = doc.add_paragraph("Bold absent (inherit)")
    # Don't add any bold element - just plain text

    # --- Italic tests ---

    # 9. Bare <w:i/> = italic ON
    p_i_bare = doc.add_paragraph("Italic bare element")
    r_i_bare = p_i_bare.runs[0]
    rPr = r_i_bare._r.get_or_add_rPr()
    for existing in rPr.findall(w("i")):
        rPr.remove(existing)
    i_el = make_element("w:i")
    rPr.append(i_el)

    # 10. <w:i w:val="0"/> = italic OFF
    p_i_off = doc.add_paragraph("Italic val=0")
    r_i_off = p_i_off.runs[0]
    rPr = r_i_off._r.get_or_add_rPr()
    for existing in rPr.findall(w("i")):
        rPr.remove(existing)
    i_el = make_element("w:i", {"w:val": "0"})
    rPr.append(i_el)

    # 11. <w:i w:val="off"/> = italic OFF
    p_i_off2 = doc.add_paragraph("Italic val=off")
    r_i_off2 = p_i_off2.runs[0]
    rPr = r_i_off2._r.get_or_add_rPr()
    for existing in rPr.findall(w("i")):
        rPr.remove(existing)
    i_el = make_element("w:i", {"w:val": "off"})
    rPr.append(i_el)

    # --- Strike tests ---

    # 12. Bare <w:strike/> = strike ON
    p_s_bare = doc.add_paragraph("Strike bare element")
    r_s_bare = p_s_bare.runs[0]
    rPr = r_s_bare._r.get_or_add_rPr()
    for existing in rPr.findall(w("strike")):
        rPr.remove(existing)
    s_el = make_element("w:strike")
    rPr.append(s_el)

    # 13. <w:strike w:val="off"/> = strike OFF
    p_s_off = doc.add_paragraph("Strike val=off")
    r_s_off = p_s_off.runs[0]
    rPr = r_s_off._r.get_or_add_rPr()
    for existing in rPr.findall(w("strike")):
        rPr.remove(existing)
    s_el = make_element("w:strike", {"w:val": "off"})
    rPr.append(s_el)

    save_fixture("edge-cases", "st-onoff-boolean-forms", doc, {
        "name": "st-onoff-boolean-forms",
        "spec_ref": "ISO 29500-1 §22.9.2.7",
        "description": "All valid ST_OnOff representations for bold/italic/strike",
        "expected_behavior": (
            "bare element=ON, val=1/true/on=ON, val=0/false/off=OFF, "
            "absent element=inherit (no mark)"
        ),
        "paragraphs": [
            "Bold bare element (ON)",
            "Bold val=1 (ON)",
            "Bold val=true (ON)",
            "Bold val=on (ON)",
            "Bold val=0 (OFF)",
            "Bold val=false (OFF)",
            "Bold val=off (OFF)",
            "Bold absent (inherit)",
            "Italic bare (ON)",
            "Italic val=0 (OFF)",
            "Italic val=off (OFF)",
            "Strike bare (ON)",
            "Strike val=off (OFF)",
        ],
    })


def make_firstline_vs_hanging() -> None:
    """firstLine vs hanging indent mutual exclusion.

    ISO 29500-1 §17.3.1.12: When both w:firstLine and w:hanging are present
    on the same w:ind element, w:hanging is ignored and w:firstLine takes
    precedence.
    """
    doc = Document()

    # Paragraph 1: only firstLine
    p1 = doc.add_paragraph("Only firstLine indent (720 twips = 0.5 inch)")
    pPr1 = p1._p.get_or_add_pPr()
    ind1 = make_element("w:ind", {"w:firstLine": "720"})
    pPr1.append(ind1)

    # Paragraph 2: only hanging
    p2 = doc.add_paragraph("Only hanging indent (360 twips)")
    pPr2 = p2._p.get_or_add_pPr()
    ind2 = make_element("w:ind", {"w:hanging": "360"})
    pPr2.append(ind2)

    # Paragraph 3: both firstLine and hanging — firstLine should win
    p3 = doc.add_paragraph("Both firstLine=720 and hanging=360 — firstLine wins")
    pPr3 = p3._p.get_or_add_pPr()
    ind3 = make_element("w:ind", {"w:firstLine": "720", "w:hanging": "360"})
    pPr3.append(ind3)

    save_fixture("edge-cases", "firstline-vs-hanging", doc, {
        "name": "firstline-vs-hanging",
        "spec_ref": "ISO 29500-1 §17.3.1.12",
        "description": "firstLine vs hanging indent mutual exclusion",
        "expected_behavior": (
            "When both firstLine and hanging are on the same w:ind, "
            "firstLine takes precedence (hanging is ignored). "
            "Para 1: first_line=+720, Para 2: first_line=-360, "
            "Para 3: first_line=+720 (hanging ignored)."
        ),
    })


def make_vmerge_absent_value() -> None:
    """vMerge absent value means 'continue'.

    ISO 29500-1 §17.4.84: The w:vMerge element controls vertical merge.
    - <w:vMerge w:val="restart"/> = start of new merge group
    - <w:vMerge/> (no val) = continue merge from cell above
    - <w:vMerge w:val="continue"/> = continue (explicit)
    - No vMerge element = no merge (cell is independent)

    The subtle case: bare <w:vMerge/> is NOT "restart", it's "continue".
    """
    doc = Document()
    doc.add_paragraph("Table with vertical merge variants.")

    tbl = doc.add_table(rows=4, cols=2)
    tbl.cell(0, 0).text = "Merge start (restart)"
    tbl.cell(0, 1).text = "B1"
    tbl.cell(1, 0).text = "Bare vMerge (continue)"
    tbl.cell(1, 1).text = "B2"
    tbl.cell(2, 0).text = "Explicit continue"
    tbl.cell(2, 1).text = "B3"
    tbl.cell(3, 0).text = "No vMerge (independent)"
    tbl.cell(3, 1).text = "B4"

    tbl_element = tbl._tbl
    rows = tbl_element.findall(w("tr"))

    # Row 0, Col 0: vMerge restart
    tc00 = rows[0].findall(w("tc"))[0]
    tcPr00 = tc00.find(w("tcPr"))
    if tcPr00 is None:
        tcPr00 = make_element("w:tcPr")
        tc00.insert(0, tcPr00)
    vmerge_restart = make_element("w:vMerge", {"w:val": "restart"})
    tcPr00.append(vmerge_restart)

    # Row 1, Col 0: bare vMerge (no val = continue)
    tc10 = rows[1].findall(w("tc"))[0]
    tcPr10 = tc10.find(w("tcPr"))
    if tcPr10 is None:
        tcPr10 = make_element("w:tcPr")
        tc10.insert(0, tcPr10)
    vmerge_bare = make_element("w:vMerge")
    tcPr10.append(vmerge_bare)

    # Row 2, Col 0: explicit vMerge continue
    tc20 = rows[2].findall(w("tc"))[0]
    tcPr20 = tc20.find(w("tcPr"))
    if tcPr20 is None:
        tcPr20 = make_element("w:tcPr")
        tc20.insert(0, tcPr20)
    vmerge_continue = make_element("w:vMerge", {"w:val": "continue"})
    tcPr20.append(vmerge_continue)

    # Row 3, Col 0: no vMerge (independent cell)
    # Nothing to inject

    save_fixture("edge-cases", "vmerge-absent-value", doc, {
        "name": "vmerge-absent-value",
        "spec_ref": "ISO 29500-1 §17.4.84",
        "description": "vMerge with restart, bare (continue), explicit continue, and absent",
        "expected_behavior": (
            "Row 0: Restart, Row 1: Continue (bare), "
            "Row 2: Continue (explicit), Row 3: None"
        ),
    })


def make_del_text_variants() -> None:
    """delText vs t for deleted run content.

    ISO 29500-1 §17.3.3.7: In w:del tracked changes, the text element
    should be w:delText (not w:t). However, Word sometimes emits w:t
    inside w:del, so parsers should handle both.
    """
    doc = Document()

    # Paragraph 1: proper w:del with w:delText
    p1 = doc.add_paragraph()
    p1.add_run("Before deletion. ")
    del_wrapper = make_element("w:del", {
        "w:id": "5000",
        "w:author": "Del Author",
        "w:date": "2025-06-01T10:00:00Z",
    })
    del_run = OxmlElement("w:r")
    del_text = make_element("w:delText")
    del_text.text = "properly deleted text"
    del_text.set(qn("xml:space"), "preserve")
    del_run.append(del_text)
    del_wrapper.append(del_run)
    p1._p.append(del_wrapper)

    # Paragraph 2: w:del with w:t (technically wrong but Word emits this)
    p2 = doc.add_paragraph()
    p2.add_run("Before wrong deletion. ")
    del_wrapper2 = make_element("w:del", {
        "w:id": "5001",
        "w:author": "Del Author",
        "w:date": "2025-06-01T10:00:00Z",
    })
    del_run2 = OxmlElement("w:r")
    del_t = make_element("w:t")
    del_t.text = "deleted with w:t"
    del_t.set(qn("xml:space"), "preserve")
    del_run2.append(del_t)
    del_wrapper2.append(del_run2)
    p2._p.append(del_wrapper2)

    # Paragraph 3: w:ins with w:t (normal insertion, for comparison)
    p3 = doc.add_paragraph()
    p3.add_run("Before insertion. ")
    ins_wrapper = make_element("w:ins", {
        "w:id": "5002",
        "w:author": "Ins Author",
        "w:date": "2025-06-01T11:00:00Z",
    })
    ins_run = OxmlElement("w:r")
    ins_t = make_element("w:t")
    ins_t.text = "inserted text"
    ins_t.set(qn("xml:space"), "preserve")
    ins_run.append(ins_t)
    ins_wrapper.append(ins_run)
    p3._p.append(ins_wrapper)

    save_fixture("edge-cases", "del-text-variants", doc, {
        "name": "del-text-variants",
        "spec_ref": "ISO 29500-1 §17.3.3.7",
        "description": "Deleted text using w:delText (correct) and w:t (Word quirk)",
        "expected_behavior": (
            "Both w:delText and w:t inside w:del should produce visible text "
            "in the deleted segment. w:ins with w:t is the normal insertion form."
        ),
    })


# =========================================================================
# EDGE CASES — Tables (ISO 29500-1 §17.4)
# =========================================================================

def make_edge_case_table_fixtures() -> None:
    print("\n── Edge Cases: Tables ──")
    make_border_conflict_resolution()
    make_gridspan_explicit_vs_absent()
    make_empty_table_cells()
    make_table_width_types()
    make_grid_before_after_edge()
    make_nested_table()


def make_border_conflict_resolution() -> None:
    """EC-14: Table border conflict — cell borders override table borders.

    Per ISO 29500-1 §17.4.38, when cell-level borders are specified they
    take precedence over table-level borders (including insideH/insideV).
    """
    doc = Document()
    doc.add_paragraph("Table with border conflict resolution test.")

    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "Cell borders"
    tbl.cell(0, 1).text = "Table borders"
    tbl.cell(1, 0).text = "Cell borders"
    tbl.cell(1, 1).text = "Table borders"

    tbl_el = tbl._tbl
    tblPr = tbl_el.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_el.insert(0, tblPr)

    # Set thick red table-level borders (sz=24 = 3pt)
    tblBorders = make_element("w:tblBorders")
    for edge in ["top", "bottom", "left", "right", "insideH", "insideV"]:
        border = make_element(f"w:{edge}", {
            "w:val": "single",
            "w:sz": "24",
            "w:color": "FF0000",
            "w:space": "0",
        })
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Override cell (0,0) with thin blue borders (sz=4 = 0.5pt)
    row0 = tbl_el.findall(w("tr"))[0]
    cell_00 = row0.findall(w("tc"))[0]
    tcPr = cell_00.find(w("tcPr"))
    if tcPr is None:
        tcPr = make_element("w:tcPr")
        cell_00.insert(0, tcPr)

    tcBorders = make_element("w:tcBorders")
    for edge in ["top", "bottom", "left", "right"]:
        border = make_element(f"w:{edge}", {
            "w:val": "single",
            "w:sz": "4",
            "w:color": "0000FF",
            "w:space": "0",
        })
        tcBorders.append(border)
    tcPr.append(tcBorders)

    save_fixture("edge-cases", "border-conflict", doc, {
        "name": "border-conflict",
        "spec_ref": "ISO 29500-1 §17.4.38",
        "description": "Table with thick red table-level borders; cell (0,0) overrides with thin blue borders",
        "expected_behavior": "Cell (0,0) should have blue/sz=4 borders; cell (0,1) should have red/sz=24 borders from table-level",
        "current_status": "SUPPORTED — testing conflict resolution priority",
    })


def make_gridspan_explicit_vs_absent() -> None:
    """EC-16: gridSpan='1' explicit vs absent gridSpan.

    Per spec, gridSpan defaults to 1 when absent. An explicit
    <w:gridSpan w:val='1'/> should produce the same domain model.
    """
    doc = Document()
    doc.add_paragraph("Two tables: one with explicit gridSpan=1, one without.")

    # Table 1: no gridSpan at all (default)
    tbl1 = doc.add_table(rows=2, cols=2)
    tbl1.cell(0, 0).text = "A1"
    tbl1.cell(0, 1).text = "B1"
    tbl1.cell(1, 0).text = "A2"
    tbl1.cell(1, 1).text = "B2"

    doc.add_paragraph("Separator.")

    # Table 2: explicit gridSpan="1" on every cell
    tbl2 = doc.add_table(rows=2, cols=2)
    tbl2.cell(0, 0).text = "A1"
    tbl2.cell(0, 1).text = "B1"
    tbl2.cell(1, 0).text = "A2"
    tbl2.cell(1, 1).text = "B2"

    tbl2_el = tbl2._tbl
    for tr in tbl2_el.findall(w("tr")):
        for tc in tr.findall(w("tc")):
            tcPr = tc.find(w("tcPr"))
            if tcPr is None:
                tcPr = make_element("w:tcPr")
                tc.insert(0, tcPr)
            tcPr.append(make_element("w:gridSpan", {"w:val": "1"}))

    save_fixture("edge-cases", "gridspan-explicit-vs-absent", doc, {
        "name": "gridspan-explicit-vs-absent",
        "spec_ref": "ISO 29500-1 §17.4.17",
        "description": "Two identical 2x2 tables: one with no gridSpan, one with explicit gridSpan=1 on all cells",
        "expected_behavior": "Both tables should produce grid_span=1 on every cell (identical domain model)",
        "current_status": "SUPPORTED — testing default equivalence",
    })


def make_empty_table_cells() -> None:
    """EC-17: Empty table cells must contain at least one paragraph.

    Per spec, every cell must contain at least one block-level element
    (typically a paragraph). Even empty cells get a <w:p/>.
    """
    doc = Document()
    doc.add_paragraph("Table with empty cells.")

    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "Has content"
    # Cell (0,1): leave with default empty paragraph from python-docx
    # Cell (1,0): explicitly clear all content and add empty <w:p/>
    tbl_el = tbl._tbl
    row1 = tbl_el.findall(w("tr"))[1]
    cell_10 = row1.findall(w("tc"))[0]
    # Remove all paragraphs
    for p_el in cell_10.findall(w("p")):
        cell_10.remove(p_el)
    # Add bare <w:p/> (empty paragraph)
    cell_10.append(make_element("w:p"))

    # Cell (1,1): completely empty (no child elements except tcPr)
    cell_11 = row1.findall(w("tc"))[1]
    for p_el in cell_11.findall(w("p")):
        cell_11.remove(p_el)
    # Leave with no paragraph at all — spec says there should be one,
    # but real DOCX files sometimes omit it

    save_fixture("edge-cases", "empty-table-cells", doc, {
        "name": "empty-table-cells",
        "spec_ref": "ISO 29500-1 §17.4.66",
        "description": "Table with cells: normal content, default empty, explicit empty <w:p/>, no paragraph at all",
        "expected_behavior": "Every cell should have at least one (possibly empty) paragraph in the domain model",
        "current_status": "SUPPORTED — testing empty cell paragraph invariant",
    })


def make_table_width_types() -> None:
    """EC-18: Table width types — auto, pct, dxa, nil.

    Tests that all four width types are preserved in the domain model.
    """
    doc = Document()
    doc.add_paragraph("Tables with different width types.")

    # Table 1: pct (50% = 5000 fiftieths-of-a-percent)
    tbl1 = doc.add_table(rows=1, cols=2)
    tbl1.cell(0, 0).text = "pct width"
    tbl1.cell(0, 1).text = "50%"
    tblPr1 = tbl1._tbl.tblPr
    if tblPr1 is None:
        tblPr1 = make_element("w:tblPr")
        tbl1._tbl.insert(0, tblPr1)
    # Remove any existing tblW
    existing = tblPr1.find(w("tblW"))
    if existing is not None:
        tblPr1.remove(existing)
    tblPr1.append(make_element("w:tblW", {"w:w": "5000", "w:type": "pct"}))

    doc.add_paragraph("---")

    # Table 2: auto
    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.cell(0, 0).text = "auto width"
    tbl2.cell(0, 1).text = "auto"
    tblPr2 = tbl2._tbl.tblPr
    if tblPr2 is None:
        tblPr2 = make_element("w:tblPr")
        tbl2._tbl.insert(0, tblPr2)
    existing = tblPr2.find(w("tblW"))
    if existing is not None:
        tblPr2.remove(existing)
    tblPr2.append(make_element("w:tblW", {"w:w": "0", "w:type": "auto"}))

    doc.add_paragraph("---")

    # Table 3: dxa (9360 twips = 6.5 inches)
    tbl3 = doc.add_table(rows=1, cols=2)
    tbl3.cell(0, 0).text = "dxa width"
    tbl3.cell(0, 1).text = "9360 twips"
    tblPr3 = tbl3._tbl.tblPr
    if tblPr3 is None:
        tblPr3 = make_element("w:tblPr")
        tbl3._tbl.insert(0, tblPr3)
    existing = tblPr3.find(w("tblW"))
    if existing is not None:
        tblPr3.remove(existing)
    tblPr3.append(make_element("w:tblW", {"w:w": "9360", "w:type": "dxa"}))

    doc.add_paragraph("---")

    # Table 4: nil
    tbl4 = doc.add_table(rows=1, cols=2)
    tbl4.cell(0, 0).text = "nil width"
    tbl4.cell(0, 1).text = "no width"
    tblPr4 = tbl4._tbl.tblPr
    if tblPr4 is None:
        tblPr4 = make_element("w:tblPr")
        tbl4._tbl.insert(0, tblPr4)
    existing = tblPr4.find(w("tblW"))
    if existing is not None:
        tblPr4.remove(existing)
    tblPr4.append(make_element("w:tblW", {"w:w": "0", "w:type": "nil"}))

    save_fixture("edge-cases", "table-width-types", doc, {
        "name": "table-width-types",
        "spec_ref": "ISO 29500-1 §17.4.84",
        "description": "Four tables with width types: pct (5000=50%), auto (0), dxa (9360 twips), nil (0)",
        "expected_behavior": "Each table's width measurement should preserve both w and width_type values",
        "current_status": "SUPPORTED — testing width type preservation",
    })


def make_grid_before_after_edge() -> None:
    """EC-19: Row with gridBefore/gridAfter.

    First row has 3 cells, second row has gridBefore=1 and only 2 cells.
    Tests that grid_before is exposed on the row.
    """
    doc = Document()
    doc.add_paragraph("Table with gridBefore on second row.")

    tbl = doc.add_table(rows=2, cols=3)
    tbl_el = tbl._tbl

    rows = tbl_el.findall(w("tr"))

    # Row 0: 3 normal cells
    cells0 = rows[0].findall(w("tc"))
    for i, label in enumerate(["X1", "Y1", "Z1"]):
        for p_el in cells0[i].findall(w("p")):
            cells0[i].remove(p_el)
        new_p = make_element("w:p")
        new_r = make_element("w:r")
        new_t = make_element("w:t")
        new_t.text = label
        new_r.append(new_t)
        new_p.append(new_r)
        cells0[i].append(new_p)

    # Row 1: gridBefore=1, only 2 cells
    trPr1 = rows[1].find(w("trPr"))
    if trPr1 is None:
        trPr1 = make_element("w:trPr")
        rows[1].insert(0, trPr1)
    trPr1.append(make_element("w:gridBefore", {"w:val": "1"}))
    wBefore = make_element("w:wBefore", {"w:w": "2000", "w:type": "dxa"})
    trPr1.append(wBefore)

    # Remove the 3rd cell
    cells1 = rows[1].findall(w("tc"))
    rows[1].remove(cells1[2])
    for i, label in enumerate(["Y2", "Z2"]):
        for p_el in cells1[i].findall(w("p")):
            cells1[i].remove(p_el)
        new_p = make_element("w:p")
        new_r = make_element("w:r")
        new_t = make_element("w:t")
        new_t.text = label
        new_r.append(new_t)
        new_p.append(new_r)
        cells1[i].append(new_p)

    save_fixture("edge-cases", "grid-before-edge", doc, {
        "name": "grid-before-edge",
        "spec_ref": "ISO 29500-1 §17.4.14",
        "description": "Row 0 has 3 cells; row 1 has gridBefore=1 and 2 cells",
        "expected_behavior": "Row 1 should expose grid_before=1 and have 2 cells",
        "current_status": "SUPPORTED — testing gridBefore exposure",
    })


def make_nested_table() -> None:
    """EC-20: Nested tables — a table cell containing another table.

    Outer 2x2 table with cell (0,0) containing a nested 2x1 table.
    """
    doc = Document()
    doc.add_paragraph("Document with nested table.")

    # Create outer table
    outer = doc.add_table(rows=2, cols=2)
    outer_el = outer._tbl

    # Fill non-nested cells
    outer.cell(0, 1).text = "Outer B1"
    outer.cell(1, 0).text = "Outer A2"
    outer.cell(1, 1).text = "Outer B2"

    # Build nested table XML in cell (0,0)
    row0 = outer_el.findall(w("tr"))[0]
    cell_00 = row0.findall(w("tc"))[0]
    # Clear existing paragraphs
    for p_el in cell_00.findall(w("p")):
        cell_00.remove(p_el)

    # Add a paragraph before the nested table
    p_before = make_element("w:p")
    r_before = make_element("w:r")
    t_before = make_element("w:t")
    t_before.text = "Before nested"
    r_before.append(t_before)
    p_before.append(r_before)
    cell_00.append(p_before)

    # Build the nested table: 2 rows x 1 col
    nested_tbl = make_element("w:tbl")
    nested_tblPr = make_element("w:tblPr")
    nested_tblW = make_element("w:tblW", {"w:w": "0", "w:type": "auto"})
    nested_tblPr.append(nested_tblW)
    nested_tbl.append(nested_tblPr)

    nested_grid = make_element("w:tblGrid")
    nested_grid.append(make_element("w:gridCol", {"w:w": "4000"}))
    nested_tbl.append(nested_grid)

    for label in ["Nested R1", "Nested R2"]:
        tr = make_element("w:tr")
        tc = make_element("w:tc")
        tc_pr = make_element("w:tcPr")
        tc_w = make_element("w:tcW", {"w:w": "4000", "w:type": "dxa"})
        tc_pr.append(tc_w)
        tc.append(tc_pr)
        p = make_element("w:p")
        r = make_element("w:r")
        t = make_element("w:t")
        t.text = label
        r.append(t)
        p.append(r)
        tc.append(p)
        tr.append(tc)
        nested_tbl.append(tr)

    cell_00.append(nested_tbl)

    # Per spec, a cell must end with a paragraph (after the nested table)
    p_after = make_element("w:p")
    cell_00.append(p_after)

    save_fixture("edge-cases", "nested-table", doc, {
        "name": "nested-table",
        "spec_ref": "ISO 29500-1 §17.4.38",
        "description": "Outer 2x2 table with cell (0,0) containing a nested 2x1 table",
        "expected_behavior": "Nested table should appear in cell (0,0).blocks as a BlockNode::Table",
        "current_status": "SUPPORTED — testing nested table access",
    })


# =========================================================================
# RUN RESOLUTION (ISO 29500-1 §17.3.2, §17.7.3, §17.8)
# =========================================================================


def make_run_resolution_fixtures() -> None:
    print("\n── Run Resolution ──")
    make_theme_font_references()
    make_toggle_xor_3level_chain()
    make_default_char_style()
    make_linked_style_cross_para()
    make_font_slot_style_resolution()
    make_para_mark_rpr()


def make_theme_font_references() -> None:
    """rFonts with asciiTheme/hAnsiTheme references instead of direct font names.

    ISO 29500-1 §17.3.2.26: rFonts can specify fonts via theme references
    (asciiTheme, hAnsiTheme, eastAsiaTheme, csTheme) that point to
    majorFont/minorFont entries in theme1.xml.  The resolved font name
    should come from the theme, not the literal attribute value.

    GAP-022: Our code only reads w:ascii/w:hAnsi/w:eastAsia/w:cs -- the
    theme reference attributes are silently ignored.
    """
    doc = Document()

    p = doc.add_paragraph()

    # Run A: only theme reference -- should resolve via theme but our code
    # will likely produce font_family=None because it only reads w:ascii/w:hAnsi.
    run_a = p.add_run("Theme-only font reference ")
    rPr_a = run_a._r.get_or_add_rPr()
    fonts_a = make_element("w:rFonts", {
        "w:asciiTheme": "majorHAnsi",
        "w:hAnsiTheme": "majorHAnsi",
    })
    rPr_a.append(fonts_a)

    # Run B: explicit + theme -- explicit should be used
    run_b = p.add_run("Explicit plus theme font reference")
    rPr_b = run_b._r.get_or_add_rPr()
    fonts_b = make_element("w:rFonts", {
        "w:ascii": "Courier New",
        "w:hAnsi": "Courier New",
        "w:asciiTheme": "majorHAnsi",
        "w:hAnsiTheme": "majorHAnsi",
    })
    rPr_b.append(fonts_b)

    save_fixture("run-resolution", "theme-font-references", doc, {
        "name": "theme-font-references",
        "spec_ref": "ISO 29500-1 §17.3.2.26",
        "description": "rFonts with asciiTheme/hAnsiTheme references. Run A has only theme refs, Run B has both explicit and theme.",
        "expected_behavior": "Run A: font_family should resolve from theme (e.g. 'Calibri Light' for majorHAnsi). Run B: font_family='Courier New' (explicit wins).",
        "current_status": "GAP -- theme font attributes (asciiTheme, hAnsiTheme) are not read; Run A gets font_family=None",
        "gap_id": "GAP-022",
    })


def make_toggle_xor_3level_chain() -> None:
    """Toggle XOR across a 3-level character style basedOn chain.

    ISO 29500-1 §17.7.3: Toggle properties use XOR across hierarchy levels.
    Within the same hierarchy level (e.g., character styles), the basedOn chain
    should be walked to find the first explicit value -- the resolved value for
    the level is then XORed with the other levels.

    Setup:
    - ParaStyle: bold=on, italic=on
    - CharStyleA: bold=on
    - CharStyleB basedOn CharStyleA: bold=on, italic=on
    - CharStyleC basedOn CharStyleB: no bold (inherit), italic=on
    """
    doc = Document()
    styles_element = doc.styles.element

    # Paragraph style: bold=on, italic=on
    para_s = make_element("w:style", {"w:type": "paragraph", "w:styleId": "TogglePara3"})
    para_s.append(make_element("w:name", {"w:val": "Toggle Para 3"}))
    rpr_p = make_element("w:rPr")
    rpr_p.append(make_element("w:b"))
    rpr_p.append(make_element("w:i"))
    para_s.append(rpr_p)
    styles_element.append(para_s)

    # CharStyleA: bold=on
    char_a = make_element("w:style", {"w:type": "character", "w:styleId": "CharA3"})
    char_a.append(make_element("w:name", {"w:val": "Char A 3"}))
    rpr_a = make_element("w:rPr")
    rpr_a.append(make_element("w:b"))
    char_a.append(rpr_a)
    styles_element.append(char_a)

    # CharStyleB basedOn CharStyleA: bold=on, italic=on
    char_b = make_element("w:style", {"w:type": "character", "w:styleId": "CharB3"})
    char_b.append(make_element("w:name", {"w:val": "Char B 3"}))
    char_b.append(make_element("w:basedOn", {"w:val": "CharA3"}))
    rpr_b = make_element("w:rPr")
    rpr_b.append(make_element("w:b"))
    rpr_b.append(make_element("w:i"))
    char_b.append(rpr_b)
    styles_element.append(char_b)

    # CharStyleC basedOn CharStyleB: no bold (inherit), italic=on
    char_c = make_element("w:style", {"w:type": "character", "w:styleId": "CharC3"})
    char_c.append(make_element("w:name", {"w:val": "Char C 3"}))
    char_c.append(make_element("w:basedOn", {"w:val": "CharB3"}))
    rpr_c = make_element("w:rPr")
    rpr_c.append(make_element("w:i"))
    char_c.append(rpr_c)
    styles_element.append(char_c)

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pPr.append(make_element("w:pStyle", {"w:val": "TogglePara3"}))

    # Run 1: CharA3 (bold=on), para bold=on -> XOR -> effective bold=off
    run1 = p.add_run("CharA3: bold XOR para bold = off. ")
    rPr1 = run1._r.get_or_add_rPr()
    rPr1.insert(0, make_element("w:rStyle", {"w:val": "CharA3"}))

    # Run 2: CharB3 (basedOn A, bold=on), para bold=on -> XOR -> effective bold=off
    run2 = p.add_run("CharB3: char bold XOR para bold = off. ")
    rPr2 = run2._r.get_or_add_rPr()
    rPr2.insert(0, make_element("w:rStyle", {"w:val": "CharB3"}))

    # Run 3: CharC3 (basedOn B, no bold). Walk: C has no bold, go to B -> bold=on.
    # Char level = on. XOR with para(on) = off.
    run3 = p.add_run("CharC3: inherited char bold XOR para bold = off. ")
    rPr3 = run3._r.get_or_add_rPr()
    rPr3.insert(0, make_element("w:rStyle", {"w:val": "CharC3"}))

    # Run 4: no char style, para bold=on -> just bold=on (no XOR partner)
    run4 = p.add_run("No char style: just para bold = on.")

    save_fixture("run-resolution", "toggle-xor-3level-chain", doc, {
        "name": "toggle-xor-3level-chain",
        "spec_ref": "ISO 29500-1 §17.7.3",
        "description": "3-level char basedOn chain + paragraph style, testing toggle XOR accumulation.",
        "expected_behavior": (
            "Run1(CharA3): bold=off (on XOR on). "
            "Run2(CharB3): bold=off (on XOR on). "
            "Run3(CharC3): bold=off (inherited on XOR on). "
            "Run4(no char): bold=on. "
            "All runs: italic resolved through XOR as well."
        ),
        "current_status": "GAP-023 -- overlay_marks flattens chain correctly for this case, but toggle XOR within basedOn levels is not spec-compliant for all scenarios",
        "gap_id": "GAP-023",
    })


def make_default_char_style() -> None:
    """Default character style (w:type="character" w:default="1").

    ISO 29500-1 §17.7.4.17 / §17.7.2: A character style with w:default="1"
    should be applied to all runs that have no explicit rStyle, similar to
    how the default paragraph style is applied to unstyled paragraphs.

    GAP-024: Our code tracks default_para_style_id but has no corresponding
    default_char_style_id. Runs without rStyle get None for char_style,
    skipping the default character style entirely.
    """
    doc = Document()
    styles_element = doc.styles.element

    # Define a default character style that sets color=0000FF (blue)
    default_char = make_element("w:style", {
        "w:type": "character",
        "w:default": "1",
        "w:styleId": "DefaultParagraphFont",
    })
    default_char.append(make_element("w:name", {"w:val": "Default Paragraph Font"}))
    rpr_dc = make_element("w:rPr")
    rpr_dc.append(make_element("w:color", {"w:val": "0000FF"}))
    default_char.append(rpr_dc)
    styles_element.append(default_char)

    # Paragraph with no explicit styles -- run should get default char style color
    p1 = doc.add_paragraph()
    run1 = p1.add_run("No rStyle -- default char style should apply color=blue. ")

    # Run with explicit rStyle -- should NOT get default char style
    explicit_char = make_element("w:style", {"w:type": "character", "w:styleId": "RedChar"})
    explicit_char.append(make_element("w:name", {"w:val": "Red Char"}))
    rpr_ec = make_element("w:rPr")
    rpr_ec.append(make_element("w:color", {"w:val": "FF0000"}))
    explicit_char.append(rpr_ec)
    styles_element.append(explicit_char)

    run2 = p1.add_run("Explicit rStyle=RedChar -- color should be red.")
    rPr2 = run2._r.get_or_add_rPr()
    rPr2.insert(0, make_element("w:rStyle", {"w:val": "RedChar"}))

    save_fixture("run-resolution", "default-char-style", doc, {
        "name": "default-char-style",
        "spec_ref": "ISO 29500-1 §17.7.4.17",
        "description": "Default character style (w:default='1') sets color=blue. Tests that unstyled runs inherit it.",
        "expected_behavior": "Run1 (no rStyle): color=0000FF from default char style. Run2 (rStyle=RedChar): color=FF0000.",
        "current_status": "GAP-024 -- no default_char_style_id; unstyled runs skip char style level entirely",
        "gap_id": "GAP-024",
    })


def make_linked_style_cross_para() -> None:
    """Linked char style applied to a run in a paragraph with a different para style.

    ISO 29500-1 §17.7.4.6: When a linked character style is applied to a run,
    the run gets the rPr from the character style definition. The linked
    paragraph style's rPr is part of the character style's definition (they
    share formatting).

    This tests: LinkedChar (linked to LinkedPara with bold, color=2E74B5)
    applied to a run inside a paragraph styled with a DIFFERENT para style.
    The run should still get the LinkedChar formatting.
    """
    doc = Document()
    styles_element = doc.styles.element

    # "Other" paragraph style with italic, color=red
    other_para = make_element("w:style", {"w:type": "paragraph", "w:styleId": "OtherPara"})
    other_para.append(make_element("w:name", {"w:val": "Other Para"}))
    rpr_op = make_element("w:rPr")
    rpr_op.append(make_element("w:i"))
    rpr_op.append(make_element("w:color", {"w:val": "FF0000"}))
    other_para.append(rpr_op)
    styles_element.append(other_para)

    # LinkedPara paragraph style with bold, color=2E74B5
    linked_para = make_element("w:style", {"w:type": "paragraph", "w:styleId": "LinkedPara2"})
    linked_para.append(make_element("w:name", {"w:val": "Linked Para 2"}))
    linked_para.append(make_element("w:link", {"w:val": "LinkedChar2"}))
    rpr_lp = make_element("w:rPr")
    rpr_lp.append(make_element("w:b"))
    rpr_lp.append(make_element("w:color", {"w:val": "2E74B5"}))
    linked_para.append(rpr_lp)
    styles_element.append(linked_para)

    # LinkedChar character style linked to LinkedPara
    linked_char = make_element("w:style", {"w:type": "character", "w:styleId": "LinkedChar2"})
    linked_char.append(make_element("w:name", {"w:val": "Linked Char 2"}))
    linked_char.append(make_element("w:link", {"w:val": "LinkedPara2"}))
    rpr_lc = make_element("w:rPr")
    rpr_lc.append(make_element("w:b"))
    rpr_lc.append(make_element("w:color", {"w:val": "2E74B5"}))
    linked_char.append(rpr_lc)
    styles_element.append(linked_char)

    # Paragraph with OtherPara style
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pPr.append(make_element("w:pStyle", {"w:val": "OtherPara"}))

    # Run 1: no char style -- gets OtherPara rPr (italic, red)
    run1 = p.add_run("OtherPara only: italic red. ")

    # Run 2: LinkedChar2 applied -- should get bold + blue from the char style
    run2 = p.add_run("LinkedChar2: bold blue, italic from para.")
    rPr2 = run2._r.get_or_add_rPr()
    rPr2.insert(0, make_element("w:rStyle", {"w:val": "LinkedChar2"}))

    save_fixture("run-resolution", "linked-style-cross-para", doc, {
        "name": "linked-style-cross-para",
        "spec_ref": "ISO 29500-1 §17.7.4.6",
        "description": "Linked char style applied to run inside paragraph with different para style.",
        "expected_behavior": "Run1: italic=on, color=FF0000 (from OtherPara). Run2: bold=on (char), color=2E74B5 (char), italic from para via XOR.",
        "current_status": "TESTING -- linked styles are treated as regular char styles (link element is ignored)",
    })


def make_font_slot_style_resolution() -> None:
    """Font slot resolution through the style chain.

    ISO 29500-1 §17.3.2.26: rFonts has four slots (ascii, hAnsi, eastAsia, cs).
    Each slot should independently cascade through the style hierarchy.

    Setup: paragraph style sets eastAsia="MS Mincho", char style sets
    ascii/hAnsi="Courier New", direct formatting sets cs="David".
    The resolved run should have all three from different levels.
    """
    doc = Document()
    styles_element = doc.styles.element

    # Paragraph style: sets eastAsia font
    font_para = make_element("w:style", {"w:type": "paragraph", "w:styleId": "FontSlotPara"})
    font_para.append(make_element("w:name", {"w:val": "Font Slot Para"}))
    rpr_fp = make_element("w:rPr")
    rpr_fp.append(make_element("w:rFonts", {"w:eastAsia": "MS Mincho"}))
    font_para.append(rpr_fp)
    styles_element.append(font_para)

    # Character style: sets ascii/hAnsi font
    font_char = make_element("w:style", {"w:type": "character", "w:styleId": "FontSlotChar"})
    font_char.append(make_element("w:name", {"w:val": "Font Slot Char"}))
    rpr_fc = make_element("w:rPr")
    rpr_fc.append(make_element("w:rFonts", {"w:ascii": "Courier New", "w:hAnsi": "Courier New"}))
    font_char.append(rpr_fc)
    styles_element.append(font_char)

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pPr.append(make_element("w:pStyle", {"w:val": "FontSlotPara"}))

    # Run with char style + direct cs font
    run = p.add_run("Font slots from 3 levels")
    rPr = run._r.get_or_add_rPr()
    rPr.insert(0, make_element("w:rStyle", {"w:val": "FontSlotChar"}))
    rPr.append(make_element("w:rFonts", {"w:cs": "David"}))

    save_fixture("run-resolution", "font-slot-style-resolution", doc, {
        "name": "font-slot-style-resolution",
        "spec_ref": "ISO 29500-1 §17.3.2.26",
        "description": "Font slots cascade independently through style chain: para=eastAsia, char=ascii/hAnsi, direct=cs.",
        "expected_behavior": "font_family='Courier New' (char), font_east_asia='MS Mincho' (para), font_cs='David' (direct).",
        "current_status": "TESTING -- font slots cascade independently via resolve_option",
    })


def make_para_mark_rpr() -> None:
    """Paragraph mark formatting (w:pPr/w:rPr) separate from run formatting.

    ISO 29500-1 §17.3.1.29: The rPr inside pPr specifies formatting for the
    paragraph mark (the pilcrow character). This is distinct from the
    formatting of text runs within the paragraph.

    Our code extracts para_mark_status (ins/del tracking) from pPr/rPr but
    does NOT extract the formatting marks (bold, color, etc.) from it.
    """
    doc = Document()

    p = doc.add_paragraph()

    # Add a run with normal formatting
    run = p.add_run("Normal text in paragraph.")

    # Set paragraph mark formatting: bold + color=FF0000
    pPr = p._p.get_or_add_pPr()
    rPr_mark = make_element("w:rPr")
    rPr_mark.append(make_element("w:b"))
    rPr_mark.append(make_element("w:color", {"w:val": "FF0000"}))
    pPr.append(rPr_mark)

    save_fixture("run-resolution", "para-mark-rpr", doc, {
        "name": "para-mark-rpr",
        "spec_ref": "ISO 29500-1 §17.3.1.29",
        "description": "Paragraph with pPr/rPr setting paragraph mark to bold+red, distinct from run formatting.",
        "expected_behavior": "Text run: no bold, no color override. Paragraph mark: bold=on, color=FF0000.",
        "current_status": "INFORMATIONAL -- para mark formatting not exposed (only tracking status extracted)",
    })


# =========================================================================
# NUMBERING + INDENT/SPACING INTERACTION (ISO 29500-1 §17.9, §17.3.1.12)
# =========================================================================

def make_numbering_indent_fixtures() -> None:
    print("\n── Numbering + Indent/Spacing ──")
    make_ni_direct_partial_override()
    make_ni_numbering_overrides_style()
    make_ni_numpr_from_deep_basedon()
    make_ni_lvl_override_partial()
    make_ni_direct_zero_overrides_numbering()
    make_ni_numbering_spacing_interaction()


def make_ni_direct_partial_override() -> None:
    """Direct w:ind partial override with numbering — the key element-level bug.

    ISO 29500-1 §17.3.1.12 + §17.9.22: When a paragraph has numbering with
    level indent (left=720, hanging=360), and direct w:ind sets only w:left="1440",
    the spec says direct paragraph properties override numbering level properties.

    Per §17.3.1.12: "Indentation settings are overridden on an individual basis."
    Per §17.9.22: "Paragraph properties specified on the numbered paragraph itself
    override the paragraph properties specified by pPr elements within a numbering
    lvl element."

    The question is whether "override" means per-field or whole-element.
    If per-field: left=1440 (direct), firstLine=-360 (from numbering hanging).
    If whole-element: left=1440 (direct), firstLine=0 (direct w:ind replaces all).

    This test documents the whole-element interpretation (GAP-025).
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:pPr>
        <w:ind w:left="720" w:hanging="360"/>
      </w:pPr>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # Paragraph with numbering AND direct w:ind that only sets w:left
    p = doc.add_paragraph("Numbered paragraph with direct left override")
    pPr = p._p.get_or_add_pPr()
    numPr = make_element("w:numPr")
    numPr.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr.append(make_element("w:numId", {"w:val": "1"}))
    pPr.append(numPr)
    pPr.append(make_element("w:ind", {"w:left": "1440"}))

    # Control: paragraph with numbering only (no direct indent)
    p2 = doc.add_paragraph("Numbered paragraph with no direct indent")
    pPr2 = p2._p.get_or_add_pPr()
    numPr2 = make_element("w:numPr")
    numPr2.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr2.append(make_element("w:numId", {"w:val": "1"}))
    pPr2.append(numPr2)

    save_fixture("numbering-indent", "direct-partial-override", doc, {
        "name": "direct-partial-override",
        "spec_ref": "ISO 29500-1 §17.3.1.12, §17.9.22",
        "description": "Numbering level indent (left=720, hanging=360) with direct w:ind w:left=1440.",
        "expected_behavior": "P0: left=1440, firstLine=0 (direct w:ind replaces numbering indent). P1: left=720, firstLine=-360 (numbering level only).",
        "current_status": "GAP-025 — per-field merge incorrectly mixes direct + numbering fields",
    })


def make_ni_numbering_overrides_style() -> None:
    """Numbering level indent overrides style indent.

    ISO 29500-1 §17.9.22: Numbering level's pPr takes priority over the
    paragraph style's indent when no direct indent is specified.
    """
    doc = Document()
    styles_element = doc.styles.element

    # Create a paragraph style with indent left=360
    style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "IndentedStyle"})
    style.append(make_element("w:name", {"w:val": "Indented Style"}))
    ppr_s = make_element("w:pPr")
    ppr_s.append(make_element("w:ind", {"w:left": "360"}))
    style.append(ppr_s)
    styles_element.append(style)

    # Create numbering with level indent left=720, hanging=360
    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:pPr>
        <w:ind w:left="720" w:hanging="360"/>
      </w:pPr>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # Paragraph with style + numbering, no direct indent
    p = doc.add_paragraph("Style indent=360, numbering indent left=720 hanging=360")
    pPr = p._p.get_or_add_pPr()
    pPr.append(make_element("w:pStyle", {"w:val": "IndentedStyle"}))
    numPr = make_element("w:numPr")
    numPr.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr.append(make_element("w:numId", {"w:val": "1"}))
    pPr.append(numPr)

    # Control: paragraph with style only, no numbering
    p2 = doc.add_paragraph("Style indent=360, no numbering")
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:pStyle", {"w:val": "IndentedStyle"}))

    save_fixture("numbering-indent", "numbering-overrides-style", doc, {
        "name": "numbering-overrides-style",
        "spec_ref": "ISO 29500-1 §17.9.22",
        "description": "Numbering level indent (left=720, hanging=360) should override style indent (left=360).",
        "expected_behavior": "P0: left=720, firstLine=-360 (numbering wins). P1: left=360 (style only).",
        "current_status": "TESTING",
    })


def make_ni_numpr_from_deep_basedon() -> None:
    """numPr inherited through deep basedOn style chain.

    ISO 29500-1 §17.7.4.3 + §17.7.4.14: A paragraph style with numPr is
    inherited through basedOn. StyleA defines numPr, StyleB basedOn StyleA,
    StyleC basedOn StyleB. Paragraph with StyleC should get numbering.
    """
    doc = Document()
    styles_element = doc.styles.element

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:pPr>
        <w:ind w:left="720" w:hanging="360"/>
      </w:pPr>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # StyleA: defines numPr
    style_a = make_element("w:style", {"w:type": "paragraph", "w:styleId": "NumStyleA"})
    style_a.append(make_element("w:name", {"w:val": "Num Style A"}))
    ppr_a = make_element("w:pPr")
    numPr_a = make_element("w:numPr")
    numPr_a.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr_a.append(make_element("w:numId", {"w:val": "1"}))
    ppr_a.append(numPr_a)
    style_a.append(ppr_a)
    styles_element.append(style_a)

    # StyleB: basedOn StyleA (inherits numPr)
    style_b = make_element("w:style", {"w:type": "paragraph", "w:styleId": "NumStyleB"})
    style_b.append(make_element("w:name", {"w:val": "Num Style B"}))
    style_b.append(make_element("w:basedOn", {"w:val": "NumStyleA"}))
    styles_element.append(style_b)

    # StyleC: basedOn StyleB (inherits numPr through chain)
    style_c = make_element("w:style", {"w:type": "paragraph", "w:styleId": "NumStyleC"})
    style_c.append(make_element("w:name", {"w:val": "Num Style C"}))
    style_c.append(make_element("w:basedOn", {"w:val": "NumStyleB"}))
    styles_element.append(style_c)

    # Paragraphs using StyleC — should inherit numbering from chain
    for text in ["First via chain", "Second via chain", "Third via chain"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        pPr.append(make_element("w:pStyle", {"w:val": "NumStyleC"}))

    save_fixture("numbering-indent", "numpr-deep-basedon", doc, {
        "name": "numpr-deep-basedon",
        "spec_ref": "ISO 29500-1 §17.7.4.3, §17.7.4.14",
        "description": "numPr inherited through StyleC -> StyleB -> StyleA chain.",
        "expected_behavior": "All three paragraphs should have numbering: 1., 2., 3. with indent left=720, firstLine=-360.",
        "current_status": "TESTING",
    })


def make_ni_lvl_override_partial() -> None:
    """lvlOverride with startOverride only — indent preserved from abstract level.

    ISO 29500-1 §17.9.8: A lvlOverride with only a startOverride should not
    affect the format or indent of the level — those come from the abstract
    numbering definition.
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:pPr>
        <w:ind w:left="720" w:hanging="360"/>
      </w:pPr>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
  <w:num w:numId="2">
    <w:abstractNumId w:val="0"/>
    <w:lvlOverride w:ilvl="0">
      <w:startOverride w:val="5"/>
    </w:lvlOverride>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # Paragraph with numId=1 (original)
    p1 = doc.add_paragraph("Original numbering, item 1")
    pPr1 = p1._p.get_or_add_pPr()
    numPr1 = make_element("w:numPr")
    numPr1.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr1.append(make_element("w:numId", {"w:val": "1"}))
    pPr1.append(numPr1)

    # Paragraph with numId=2 (startOverride=5, indent should be same)
    p2 = doc.add_paragraph("Override start=5, same indent expected")
    pPr2 = p2._p.get_or_add_pPr()
    numPr2 = make_element("w:numPr")
    numPr2.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr2.append(make_element("w:numId", {"w:val": "2"}))
    pPr2.append(numPr2)

    p3 = doc.add_paragraph("Override continuation, item 6")
    pPr3 = p3._p.get_or_add_pPr()
    numPr3 = make_element("w:numPr")
    numPr3.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr3.append(make_element("w:numId", {"w:val": "2"}))
    pPr3.append(numPr3)

    save_fixture("numbering-indent", "lvl-override-partial", doc, {
        "name": "lvl-override-partial",
        "spec_ref": "ISO 29500-1 §17.9.8, §17.9.26",
        "description": "lvlOverride with only startOverride=5; format and indent come from abstract level.",
        "expected_behavior": "P0: prefix='1.', indent left=720, firstLine=-360. P1: prefix='5.', same indent. P2: prefix='6.', same indent.",
        "current_status": "TESTING",
    })


def make_ni_direct_zero_overrides_numbering() -> None:
    """Direct w:ind w:left="0" zeroing numbering indent.

    ISO 29500-1 §17.9.22 + §17.3.1.12: Direct paragraph properties override
    numbering level properties. When direct w:ind w:left="0" is present, the
    numbering level's indent should be completely ignored.
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:pPr>
        <w:ind w:left="720" w:hanging="360"/>
      </w:pPr>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # Paragraph with numbering + direct w:ind w:left="0"
    p = doc.add_paragraph("Numbered paragraph with left zeroed out")
    pPr = p._p.get_or_add_pPr()
    numPr = make_element("w:numPr")
    numPr.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr.append(make_element("w:numId", {"w:val": "1"}))
    pPr.append(numPr)
    pPr.append(make_element("w:ind", {"w:left": "0"}))

    # Control: numbered paragraph without direct indent
    p2 = doc.add_paragraph("Numbered paragraph with default indent")
    pPr2 = p2._p.get_or_add_pPr()
    numPr2 = make_element("w:numPr")
    numPr2.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr2.append(make_element("w:numId", {"w:val": "1"}))
    pPr2.append(numPr2)

    save_fixture("numbering-indent", "direct-zero-overrides", doc, {
        "name": "direct-zero-overrides",
        "spec_ref": "ISO 29500-1 §17.9.22, §17.3.1.12",
        "description": "Direct w:ind w:left='0' on a numbered paragraph should zero out the numbering indent.",
        "expected_behavior": "P0: left=0, firstLine=0 (direct replaces). P1: left=720, firstLine=-360 (numbering only).",
        "current_status": "GAP-025 — per-field merge lets firstLine leak from numbering",
    })


def make_ni_numbering_spacing_interaction() -> None:
    """Numbering + style spacing interaction.

    Numbering levels don't define spacing. The paragraph style's spacing should
    apply normally. Direct w:spacing then overrides the style spacing.
    """
    doc = Document()
    styles_element = doc.styles.element

    # Create a paragraph style with spacing before=120, after=120
    style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "SpacedStyle"})
    style.append(make_element("w:name", {"w:val": "Spaced Style"}))
    ppr_s = make_element("w:pPr")
    ppr_s.append(make_element("w:spacing", {"w:before": "120", "w:after": "120"}))
    style.append(ppr_s)
    styles_element.append(style)

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:pPr>
        <w:ind w:left="720" w:hanging="360"/>
      </w:pPr>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # P0: Style spacing + numbering, no direct spacing
    p = doc.add_paragraph("Numbered with style spacing before=120, after=120")
    pPr = p._p.get_or_add_pPr()
    pPr.append(make_element("w:pStyle", {"w:val": "SpacedStyle"}))
    numPr = make_element("w:numPr")
    numPr.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr.append(make_element("w:numId", {"w:val": "1"}))
    pPr.append(numPr)

    # P1: Style spacing + numbering + direct spacing w:before="0"
    p2 = doc.add_paragraph("Numbered with direct spacing before=0 override")
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:pStyle", {"w:val": "SpacedStyle"}))
    numPr2 = make_element("w:numPr")
    numPr2.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr2.append(make_element("w:numId", {"w:val": "1"}))
    pPr2.append(numPr2)
    pPr2.append(make_element("w:spacing", {"w:before": "0"}))

    save_fixture("numbering-indent", "numbering-spacing-interaction", doc, {
        "name": "numbering-spacing-interaction",
        "spec_ref": "ISO 29500-1 §17.9.22, §17.3.1.33",
        "description": "Numbering doesn't define spacing; style spacing applies. Direct spacing overrides style.",
        "expected_behavior": "P0: spacing before=120, after=120 (from style). P1: spacing before=0, after=120 (direct before overrides).",
        "current_status": "TESTING",
    })


# =========================================================================
# STYLE CASCADE (ISO 29500-1 §17.3.1.12, §17.3.1.33, §17.7.2)
# =========================================================================


def make_style_cascade_fixtures() -> None:
    print("\n── Style Cascade ──")
    make_cascade_indent_partial_direct()
    make_cascade_spacing_partial_direct()
    make_cascade_deep_based_on_indent()
    make_cascade_doc_defaults_spacing()
    make_cascade_normal_style_indent()


def make_cascade_indent_partial_direct() -> None:
    """Direct w:ind with partial attributes -- per-field merge with style.

    ISO 29500-1 §17.3.1.12: "Indentation settings are overridden on an
    individual basis - if any single attribute on this element is omitted on
    a given paragraph, its value is determined by the setting previously set
    at any level of the style hierarchy."

    Style: left=720, right=360, firstLine=240
    Direct: left=1440 only
    Expected: left=1440 (direct), right=360 (style), firstLine=240 (style)
    """
    doc = Document()
    styles_element = doc.styles.element

    # Define a paragraph style with full indentation
    style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "IndentedStyle"})
    style.append(make_element("w:name", {"w:val": "Indented Style"}))
    ppr = make_element("w:pPr")
    ppr.append(make_element("w:ind", {
        "w:left": "720",
        "w:right": "360",
        "w:firstLine": "240",
    }))
    style.append(ppr)
    styles_element.append(style)

    # Para 1: style only, no direct formatting
    p1 = doc.add_paragraph("Style only: left=720, right=360, firstLine=240")
    pPr1 = p1._p.get_or_add_pPr()
    pPr1.append(make_element("w:pStyle", {"w:val": "IndentedStyle"}))

    # Para 2: style + direct w:ind with only left attribute
    p2 = doc.add_paragraph("Direct left=1440, style right=360, style firstLine=240")
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:pStyle", {"w:val": "IndentedStyle"}))
    pPr2.append(make_element("w:ind", {"w:left": "1440"}))

    # Para 3: direct w:ind overriding all three fields
    p3 = doc.add_paragraph("Direct left=0, right=0, firstLine=0")
    pPr3 = p3._p.get_or_add_pPr()
    pPr3.append(make_element("w:pStyle", {"w:val": "IndentedStyle"}))
    pPr3.append(make_element("w:ind", {
        "w:left": "0",
        "w:right": "0",
        "w:firstLine": "0",
    }))

    save_fixture("style-cascade", "indent-partial-direct", doc, {
        "name": "indent-partial-direct",
        "spec_ref": "ISO 29500-1 §17.3.1.12",
        "description": (
            "Style defines w:ind left=720 right=360 firstLine=240. "
            "Para 2 has direct w:ind left=1440 only. Per spec, absent "
            "attributes inherit from style individually."
        ),
        "expected_behavior": (
            "Para 1: left=720, right=360, firstLine=240 (all from style). "
            "Para 2: left=1440 (direct), right=360 (style), firstLine=240 (style). "
            "Para 3: left=0, right=0, firstLine=0 (all direct)."
        ),
        "current_status": "TESTING",
    })


def make_cascade_spacing_partial_direct() -> None:
    """Direct w:spacing with partial attributes -- per-field merge with style.

    ISO 29500-1 §17.3.1.33: Each attribute on w:spacing falls through to
    the style hierarchy independently when omitted.

    Style: before=240, after=120, line=276, lineRule=auto
    Direct: before=0 only
    Expected: before=0 (direct), after=120 (style), line=276 (style)

    Also tests lineRule defaulting: per spec, if lineRule is omitted but
    line is present, lineRule defaults to "auto" (NOT inherited from style).
    """
    doc = Document()
    styles_element = doc.styles.element

    # Remove existing docDefaults
    existing_defaults = styles_element.find(w("docDefaults"))
    if existing_defaults is not None:
        styles_element.remove(existing_defaults)

    # Minimal docDefaults
    doc_defaults = make_element("w:docDefaults")
    rpr_default = make_element("w:rPrDefault")
    rpr = make_element("w:rPr")
    rpr.append(make_element("w:sz", {"w:val": "24"}))
    rpr_default.append(rpr)
    doc_defaults.append(rpr_default)
    styles_element.insert(0, doc_defaults)

    # Style with full spacing
    style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "SpacedStyle"})
    style.append(make_element("w:name", {"w:val": "Spaced Style"}))
    ppr = make_element("w:pPr")
    ppr.append(make_element("w:spacing", {
        "w:before": "240",
        "w:after": "120",
        "w:line": "276",
        "w:lineRule": "auto",
    }))
    style.append(ppr)
    styles_element.append(style)

    # Para 1: style only
    p1 = doc.add_paragraph("Style only: before=240, after=120, line=276")
    pPr1 = p1._p.get_or_add_pPr()
    pPr1.append(make_element("w:pStyle", {"w:val": "SpacedStyle"}))

    # Para 2: direct w:spacing with only before=0
    p2 = doc.add_paragraph("Direct before=0, style after=120, style line=276")
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:pStyle", {"w:val": "SpacedStyle"}))
    pPr2.append(make_element("w:spacing", {"w:before": "0"}))

    # Para 3: direct w:spacing with line=480 but no lineRule
    # Per spec: lineRule defaults to "auto" when line is present
    p3 = doc.add_paragraph("Direct line=480, no lineRule attr, should default to auto")
    pPr3 = p3._p.get_or_add_pPr()
    pPr3.append(make_element("w:pStyle", {"w:val": "SpacedStyle"}))
    pPr3.append(make_element("w:spacing", {"w:line": "480"}))

    # Second style with lineRule=exact (to expose GAP-016)
    style2 = make_element("w:style", {"w:type": "paragraph", "w:styleId": "ExactSpacedStyle"})
    style2.append(make_element("w:name", {"w:val": "Exact Spaced Style"}))
    ppr2 = make_element("w:pPr")
    ppr2.append(make_element("w:spacing", {
        "w:before": "240",
        "w:after": "120",
        "w:line": "400",
        "w:lineRule": "exact",
    }))
    style2.append(ppr2)
    styles_element.append(style2)

    # Para 4: ExactSpacedStyle + direct line=480 without lineRule
    # GAP-016: Per spec, lineRule should default to "auto" since line is
    # present. But per-field merge will fall through to style's "exact".
    p4 = doc.add_paragraph("ExactStyle + direct line=480, no lineRule: should be auto, bug gives exact")
    pPr4 = p4._p.get_or_add_pPr()
    pPr4.append(make_element("w:pStyle", {"w:val": "ExactSpacedStyle"}))
    pPr4.append(make_element("w:spacing", {"w:line": "480"}))

    save_fixture("style-cascade", "spacing-partial-direct", doc, {
        "name": "spacing-partial-direct",
        "spec_ref": "ISO 29500-1 §17.3.1.33",
        "description": (
            "Style defines w:spacing before=240 after=120 line=276 lineRule=auto. "
            "Para 2 has direct w:spacing before=0 only. "
            "Para 3 has direct w:spacing line=480 with no lineRule attribute. "
            "Para 4 uses ExactSpacedStyle (lineRule=exact) with direct line=480 "
            "and no lineRule attr, exposing GAP-016."
        ),
        "expected_behavior": (
            "Para 1: before=240, after=120, line=276, lineRule=auto (all style). "
            "Para 2: before=0 (direct), after=120 (style), line=276 (style). "
            "Para 3: line=480 (direct), lineRule=auto (spec default). "
            "Para 4: line=480 (direct), lineRule=auto (spec default, NOT exact "
            "from style -- GAP-016)."
        ),
        "current_status": "TESTING -- GAP-016 exposed by para 4",
    })


def make_cascade_deep_based_on_indent() -> None:
    """Deep basedOn chain (3 levels) with per-field indent inheritance.

    ISO 29500-1 §17.7.4.3 + §17.3.1.12: Within the style chain, per-field
    merge is correct. Then direct formatting is also merged per-field.

    StyleA: left=360
    StyleB basedOn A: right=360
    StyleC basedOn B: firstLine=240
    Resolved for C: left=360 (A), right=360 (B), firstLine=240 (C)
    Direct on paragraph: left=720
    Expected: left=720 (direct), right=360 (B), firstLine=240 (C)
    """
    doc = Document()
    styles_element = doc.styles.element

    # StyleA: only left indent
    style_a = make_element("w:style", {"w:type": "paragraph", "w:styleId": "IndentA"})
    style_a.append(make_element("w:name", {"w:val": "Indent A"}))
    ppr_a = make_element("w:pPr")
    ppr_a.append(make_element("w:ind", {"w:left": "360"}))
    style_a.append(ppr_a)
    styles_element.append(style_a)

    # StyleB basedOn A: only right indent
    style_b = make_element("w:style", {"w:type": "paragraph", "w:styleId": "IndentB"})
    style_b.append(make_element("w:name", {"w:val": "Indent B"}))
    style_b.append(make_element("w:basedOn", {"w:val": "IndentA"}))
    ppr_b = make_element("w:pPr")
    ppr_b.append(make_element("w:ind", {"w:right": "360"}))
    style_b.append(ppr_b)
    styles_element.append(style_b)

    # StyleC basedOn B: only firstLine indent
    style_c = make_element("w:style", {"w:type": "paragraph", "w:styleId": "IndentC"})
    style_c.append(make_element("w:name", {"w:val": "Indent C"}))
    style_c.append(make_element("w:basedOn", {"w:val": "IndentB"}))
    ppr_c = make_element("w:pPr")
    ppr_c.append(make_element("w:ind", {"w:firstLine": "240"}))
    style_c.append(ppr_c)
    styles_element.append(style_c)

    # Para 1: StyleC only
    p1 = doc.add_paragraph("StyleC: left=360 (A), right=360 (B), firstLine=240 (C)")
    pPr1 = p1._p.get_or_add_pPr()
    pPr1.append(make_element("w:pStyle", {"w:val": "IndentC"}))

    # Para 2: StyleC + direct left=720
    p2 = doc.add_paragraph("StyleC + direct left=720: right=360 (B), firstLine=240 (C)")
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:pStyle", {"w:val": "IndentC"}))
    pPr2.append(make_element("w:ind", {"w:left": "720"}))

    # Para 3: StyleB only
    p3 = doc.add_paragraph("StyleB: left=360 (A), right=360 (B), no firstLine")
    pPr3 = p3._p.get_or_add_pPr()
    pPr3.append(make_element("w:pStyle", {"w:val": "IndentB"}))

    save_fixture("style-cascade", "deep-based-on-indent", doc, {
        "name": "deep-based-on-indent",
        "spec_ref": "ISO 29500-1 §17.7.4.3, §17.3.1.12",
        "description": (
            "Three-level basedOn chain with each level contributing a different "
            "indent field. Tests per-field inheritance within style chain and "
            "per-field merge with direct formatting."
        ),
        "expected_behavior": (
            "Para 1 (StyleC): left=360, right=360, firstLine=240. "
            "Para 2 (StyleC+direct): left=720, right=360, firstLine=240. "
            "Para 3 (StyleB): left=360, right=360, no firstLine."
        ),
        "current_status": "TESTING",
    })


def make_cascade_doc_defaults_spacing() -> None:
    """docDefaults spacing interaction with style and direct formatting.

    ISO 29500-1 §17.7.4.17 + §17.3.1.33: docDefaults provides fallback.
    Per-field: direct > style > docDefaults.

    docDefaults: after=200, line=276
    Style: before=120 (only before)
    Expected without direct: before=120 (style), after=200 (docDefaults), line=276 (docDefaults)
    """
    doc = Document()
    styles_element = doc.styles.element

    # Set docDefaults with spacing
    existing_defaults = styles_element.find(w("docDefaults"))
    if existing_defaults is not None:
        styles_element.remove(existing_defaults)

    doc_defaults = make_element("w:docDefaults")
    rpr_default = make_element("w:rPrDefault")
    rpr = make_element("w:rPr")
    rpr.append(make_element("w:sz", {"w:val": "24"}))
    rpr_default.append(rpr)
    doc_defaults.append(rpr_default)
    ppr_default = make_element("w:pPrDefault")
    ppr = make_element("w:pPr")
    ppr.append(make_element("w:spacing", {
        "w:after": "200",
        "w:line": "276",
        "w:lineRule": "auto",
    }))
    ppr_default.append(ppr)
    doc_defaults.append(ppr_default)
    styles_element.insert(0, doc_defaults)

    # Style with only before=120
    style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "BeforeOnlyStyle"})
    style.append(make_element("w:name", {"w:val": "Before Only Style"}))
    ppr_s = make_element("w:pPr")
    ppr_s.append(make_element("w:spacing", {"w:before": "120"}))
    style.append(ppr_s)
    styles_element.append(style)

    # Para 1: no style, no direct
    p1 = doc.add_paragraph("No style: after=200, line=276 from docDefaults")

    # Para 2: style only
    p2 = doc.add_paragraph("Style: before=120, after=200 (defaults), line=276 (defaults)")
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:pStyle", {"w:val": "BeforeOnlyStyle"}))

    # Para 3: style + direct before=0
    p3 = doc.add_paragraph("Style + direct before=0: after=200 (defaults), line=276 (defaults)")
    pPr3 = p3._p.get_or_add_pPr()
    pPr3.append(make_element("w:pStyle", {"w:val": "BeforeOnlyStyle"}))
    pPr3.append(make_element("w:spacing", {"w:before": "0"}))

    save_fixture("style-cascade", "doc-defaults-spacing", doc, {
        "name": "doc-defaults-spacing",
        "spec_ref": "ISO 29500-1 §17.7.4.17, §17.3.1.33",
        "description": (
            "docDefaults sets spacing after=200 line=276. Style sets only before=120. "
            "Tests three-level fallback: direct > style > docDefaults."
        ),
        "expected_behavior": (
            "Para 1: after=200, line=276 (docDefaults). "
            "Para 2: before=120 (style), after=200 (docDefaults), line=276 (docDefaults). "
            "Para 3: before=0 (direct), after=200 (docDefaults), line=276 (docDefaults)."
        ),
        "current_status": "TESTING",
    })


def make_cascade_normal_style_indent() -> None:
    """Normal (default) paragraph style indent with direct override.

    ISO 29500-1 §17.7.4.17 + §17.3.1.12: When no explicit pStyle is set,
    the Normal style applies. Direct w:ind per-field overrides Normal.

    Normal: left=720, firstLine=360
    Direct: left=0 (only left)
    Expected: left=0 (direct), firstLine=360 (Normal)
    """
    doc = Document()
    styles_element = doc.styles.element

    # Modify the Normal style to have indentation
    normal_style = None
    for style_el in styles_element.findall(w("style")):
        style_id = style_el.get(qn("w:styleId"))
        if style_id == "Normal":
            normal_style = style_el
            break

    if normal_style is None:
        normal_style = make_element("w:style", {
            "w:type": "paragraph",
            "w:styleId": "Normal",
            "w:default": "1",
        })
        normal_style.append(make_element("w:name", {"w:val": "Normal"}))
        styles_element.append(normal_style)

    # Add/replace pPr with indentation
    existing_ppr = normal_style.find(w("pPr"))
    if existing_ppr is not None:
        normal_style.remove(existing_ppr)
    ppr = make_element("w:pPr")
    ppr.append(make_element("w:ind", {
        "w:left": "720",
        "w:firstLine": "360",
    }))
    normal_style.append(ppr)

    # Para 1: no explicit pStyle
    p1 = doc.add_paragraph("Normal style: left=720, firstLine=360")

    # Para 2: direct w:ind left=0
    p2 = doc.add_paragraph("Direct left=0, Normal firstLine=360")
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:ind", {"w:left": "0"}))

    # Para 3: explicit pStyle=Normal + direct w:ind left=0
    p3 = doc.add_paragraph("Explicit Normal + direct left=0, firstLine=360")
    pPr3 = p3._p.get_or_add_pPr()
    pPr3.append(make_element("w:pStyle", {"w:val": "Normal"}))
    pPr3.append(make_element("w:ind", {"w:left": "0"}))

    save_fixture("style-cascade", "normal-style-indent", doc, {
        "name": "normal-style-indent",
        "spec_ref": "ISO 29500-1 §17.3.1.12, §17.7.4.17",
        "description": (
            "Normal style defines left=720 firstLine=360. "
            "Tests that direct w:ind left=0 overrides left but firstLine "
            "falls through to Normal per per-field merge."
        ),
        "expected_behavior": (
            "Para 1: left=720, firstLine=360 (Normal). "
            "Para 2: left=0 (direct), firstLine=360 (Normal). "
            "Para 3: left=0 (direct), firstLine=360 (Normal)."
        ),
        "current_status": "TESTING",
    })


# =========================================================================
# TABLE CASCADE (ISO 29500-1 §17.7.6, §17.4.43, §17.4.63)
# =========================================================================

def make_table_cascade_fixtures() -> None:
    print("\n── Table Cascade ──")
    make_tc_conditional_precedence()
    make_tc_style_paragraph_props()
    make_tc_style_run_props()
    make_tc_cell_margin_chain()
    make_tc_band_interaction()
    make_tc_style_table_level_props()


def _add_table_style(doc: DocxDocument, style_id: str, name: str,
                     tblPr_children=None, tcPr_children=None,
                     tblStylePrs=None) -> None:
    """Helper to inject a table style into styles.xml."""
    styles_element = doc.styles.element
    tbl_style = make_element("w:style", {"w:type": "table", "w:styleId": style_id})
    tbl_style.append(make_element("w:name", {"w:val": name}))

    if tblPr_children:
        tblPr = make_element("w:tblPr")
        for child in tblPr_children:
            tblPr.append(child)
        tbl_style.append(tblPr)

    if tcPr_children:
        tcPr = make_element("w:tcPr")
        for child in tcPr_children:
            tcPr.append(child)
        tbl_style.append(tcPr)

    if tblStylePrs:
        for spr in tblStylePrs:
            tbl_style.append(spr)

    styles_element.append(tbl_style)


def _apply_table_style(tbl, style_id: str, look_attrs: dict) -> None:
    """Apply a table style and tblLook to an existing table element."""
    tbl_element = tbl._tbl
    tblPr = tbl_element.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_element.insert(0, tblPr)
    tblPr.insert(0, make_element("w:tblStyle", {"w:val": style_id}))
    # Remove any existing tblLook before adding our own (python-docx adds a default one).
    for existing in tblPr.findall(qn("w:tblLook")):
        tblPr.remove(existing)
    tblPr.append(make_element("w:tblLook", look_attrs))


def _fill_table_cells(tbl, labels: list[list[str]]) -> None:
    """Fill table cells with labels."""
    for r_idx, row_labels in enumerate(labels):
        for c_idx, label in enumerate(row_labels):
            tbl.cell(r_idx, c_idx).text = label


def make_tc_conditional_precedence() -> None:
    """Conditional formatting precedence: firstCol > firstRow at corner cell.

    ISO 29500-1 §17.7.6: When multiple conditional formats apply, they are
    applied in order: whole-table, banded cols, banded rows, firstRow/lastRow,
    firstCol/lastCol, corner cells. Subsequent formats override previous ones.

    So firstCol overrides firstRow for a corner cell (0,0).
    """
    doc = Document()

    # Build tblStylePr elements
    # firstRow: blue shading (4472C4)
    first_row_spr = make_element("w:tblStylePr", {"w:type": "firstRow"})
    fr_tcPr = make_element("w:tcPr")
    fr_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "4472C4"}))
    first_row_spr.append(fr_tcPr)

    # lastRow: yellow shading (FFC000)
    last_row_spr = make_element("w:tblStylePr", {"w:type": "lastRow"})
    lr_tcPr = make_element("w:tcPr")
    lr_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "FFC000"}))
    last_row_spr.append(lr_tcPr)

    # firstCol: green shading (70AD47)
    first_col_spr = make_element("w:tblStylePr", {"w:type": "firstCol"})
    fc_tcPr = make_element("w:tcPr")
    fc_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "70AD47"}))
    first_col_spr.append(fc_tcPr)

    # lastCol: red shading (FF0000)
    last_col_spr = make_element("w:tblStylePr", {"w:type": "lastCol"})
    lc_tcPr = make_element("w:tcPr")
    lc_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "FF0000"}))
    last_col_spr.append(lc_tcPr)

    # Base table borders
    tbl_borders = make_element("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "4", "w:color": "000000",
        }))

    _add_table_style(doc, "PrecedenceTable", "Precedence Table",
                     tblPr_children=[tbl_borders],
                     tblStylePrs=[first_row_spr, last_row_spr,
                                  first_col_spr, last_col_spr])

    # 4x3 table with all conditional flags enabled
    doc.add_paragraph("Conditional precedence test.")
    tbl = doc.add_table(rows=4, cols=3)
    _apply_table_style(tbl, "PrecedenceTable", {
        "w:val": "04A0",
        "w:firstRow": "1", "w:lastRow": "1",
        "w:firstColumn": "1", "w:lastColumn": "1",
        "w:noHBand": "1", "w:noVBand": "1",
    })
    _fill_table_cells(tbl, [
        ["R0C0", "R0C1", "R0C2"],
        ["R1C0", "R1C1", "R1C2"],
        ["R2C0", "R2C1", "R2C2"],
        ["R3C0", "R3C1", "R3C2"],
    ])

    save_fixture("table-cascade", "conditional-precedence", doc, {
        "name": "conditional-precedence",
        "spec_ref": "ISO 29500-1 §17.7.6",
        "description": (
            "Table with firstRow (blue), lastRow (yellow), firstCol (green), "
            "lastCol (red) conditionals. All tblLook flags enabled."
        ),
        "expected_behavior": (
            "Per §17.7.6 precedence: firstCol/lastCol > firstRow/lastRow. "
            "Cell (0,0): green (firstCol > firstRow). "
            "Cell (0,2): red (lastCol > firstRow). "
            "Cell (3,0): green (firstCol > lastRow). "
            "Cell (3,2): red (lastCol > lastRow). "
            "Cell (0,1): blue (firstRow only). "
            "Cell (3,1): yellow (lastRow only). "
            "Cell (1,0): green (firstCol only). "
            "Cell (1,2): red (lastCol only)."
        ),
        "current_status": "GAP — precedence order in apply_conditional_formatting is wrong",
    })


def make_tc_style_paragraph_props() -> None:
    """Table style conditional formatting with paragraph properties.

    ISO 29500-1 §17.7.6.1: tblStylePr can contain pPr (paragraph properties)
    that should be applied to paragraphs within matching cells.
    """
    doc = Document()

    # firstRow conditional with center alignment
    first_row_spr = make_element("w:tblStylePr", {"w:type": "firstRow"})
    fr_pPr = make_element("w:pPr")
    fr_pPr.append(make_element("w:jc", {"w:val": "center"}))
    first_row_spr.append(fr_pPr)
    fr_tcPr = make_element("w:tcPr")
    fr_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "DDDDDD"}))
    first_row_spr.append(fr_tcPr)

    _add_table_style(doc, "ParaPropTable", "Paragraph Props Table",
                     tblStylePrs=[first_row_spr])

    doc.add_paragraph("Table style with conditional pPr.")
    tbl = doc.add_table(rows=3, cols=2)
    _apply_table_style(tbl, "ParaPropTable", {
        "w:val": "04A0",
        "w:firstRow": "1", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "1",
    })
    _fill_table_cells(tbl, [
        ["Header 1", "Header 2"],
        ["Data A", "Data B"],
        ["Data C", "Data D"],
    ])

    save_fixture("table-cascade", "style-para-props", doc, {
        "name": "style-para-props",
        "spec_ref": "ISO 29500-1 §17.7.6.1",
        "description": "Table style with firstRow conditional setting jc=center in pPr.",
        "expected_behavior": (
            "Paragraphs in first row cells should have alignment=center from "
            "the conditional pPr. Data rows should have no explicit alignment."
        ),
        "current_status": "GAP — ConditionalCellProps does not carry pPr or rPr",
    })


def make_tc_style_run_props() -> None:
    """Table style conditional formatting with run properties.

    ISO 29500-1 §17.7.6.2: tblStylePr can contain rPr (run properties)
    that should be applied to runs within matching cells.
    """
    doc = Document()

    # firstRow conditional with bold
    first_row_spr = make_element("w:tblStylePr", {"w:type": "firstRow"})
    fr_rPr = make_element("w:rPr")
    fr_rPr.append(make_element("w:b"))
    fr_rPr.append(make_element("w:color", {"w:val": "FFFFFF"}))
    first_row_spr.append(fr_rPr)
    fr_tcPr = make_element("w:tcPr")
    fr_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "4472C4"}))
    first_row_spr.append(fr_tcPr)

    _add_table_style(doc, "RunPropTable", "Run Props Table",
                     tblStylePrs=[first_row_spr])

    doc.add_paragraph("Table style with conditional rPr.")
    tbl = doc.add_table(rows=3, cols=2)
    _apply_table_style(tbl, "RunPropTable", {
        "w:val": "04A0",
        "w:firstRow": "1", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "1",
    })
    _fill_table_cells(tbl, [
        ["Bold Header 1", "Bold Header 2"],
        ["Normal A", "Normal B"],
        ["Normal C", "Normal D"],
    ])

    save_fixture("table-cascade", "style-run-props", doc, {
        "name": "style-run-props",
        "spec_ref": "ISO 29500-1 §17.7.6.2",
        "description": "Table style with firstRow conditional setting bold in rPr.",
        "expected_behavior": (
            "Text in first row cells should be bold and white (color=FFFFFF) from "
            "the conditional rPr. Data rows should have default formatting."
        ),
        "current_status": "GAP — ConditionalCellProps does not carry rPr",
    })


def make_tc_cell_margin_chain() -> None:
    """Cell margin inheritance: direct cell > table tblCellMar > table style.

    ISO 29500-1 §17.4.43: Cell margins follow a three-level chain:
    1. Direct cell margins (w:tcMar in w:tcPr) — highest priority
    2. Table default cell margins (w:tblCellMar in w:tblPr) — direct table
    3. Table style default cell margins (w:tblCellMar in style's w:tblPr) — lowest
    """
    doc = Document()

    # Table style with default cell margins (top=50, bottom=50, left=50, right=50)
    style_cell_mar = make_element("w:tblCellMar")
    for side in ["top", "bottom", "left", "right"]:
        style_cell_mar.append(make_element(f"w:{side}", {
            "w:w": "50", "w:type": "dxa",
        }))

    _add_table_style(doc, "MarginTable", "Margin Table",
                     tblPr_children=[style_cell_mar])

    doc.add_paragraph("Cell margin inheritance chain.")
    tbl = doc.add_table(rows=3, cols=2)
    _apply_table_style(tbl, "MarginTable", {
        "w:val": "04A0",
        "w:firstRow": "0", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "1",
    })
    _fill_table_cells(tbl, [
        ["Style margins", "Style margins"],
        ["Table margins", "Table margins"],
        ["Direct margins", "Direct margins"],
    ])

    # Table-level override: set tblCellMar on the table itself (100 twips each)
    tbl_element = tbl._tbl
    tblPr = tbl_element.tblPr
    tbl_cell_mar = make_element("w:tblCellMar")
    for side in ["top", "bottom", "left", "right"]:
        tbl_cell_mar.append(make_element(f"w:{side}", {
            "w:w": "100", "w:type": "dxa",
        }))
    tblPr.append(tbl_cell_mar)

    # Direct cell override on cell (2,0): tcMar with 200tw each
    cell_20 = tbl.cell(2, 0)
    tc_el = cell_20._tc
    tcPr = tc_el.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = make_element("w:tcPr")
        tc_el.insert(0, tcPr)
    tc_mar = make_element("w:tcMar")
    tc_mar.append(make_element("w:top", {"w:w": "200", "w:type": "dxa"}))
    tc_mar.append(make_element("w:bottom", {"w:w": "200", "w:type": "dxa"}))
    tc_mar.append(make_element("w:left", {"w:w": "200", "w:type": "dxa"}))
    tc_mar.append(make_element("w:right", {"w:w": "200", "w:type": "dxa"}))
    tcPr.append(tc_mar)

    save_fixture("table-cascade", "cell-margin-chain", doc, {
        "name": "cell-margin-chain",
        "spec_ref": "ISO 29500-1 §17.4.43",
        "description": (
            "Table style margins=50tw, table direct margins=100tw, "
            "cell (2,0) direct margins=200tw."
        ),
        "expected_behavior": (
            "Row 0-1 cells: margins from table direct (100tw each). "
            "Cell (2,0): direct margins (200tw each). "
            "If table direct were absent, style margins (50tw) would apply."
        ),
        "current_status": "TESTING — verifying three-level chain",
    })


def make_tc_band_interaction() -> None:
    """Band row + band column interaction at intersection cells.

    ISO 29500-1 §17.7.6: Banded rows are higher precedence than banded columns.
    When both band1Horz and band1Vert apply to a cell, band row wins.
    """
    doc = Document()

    # band1Horz: light gray shading (D9D9D9)
    band1h_spr = make_element("w:tblStylePr", {"w:type": "band1Horz"})
    b1h_tcPr = make_element("w:tcPr")
    b1h_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "D9D9D9"}))
    band1h_spr.append(b1h_tcPr)

    # band2Horz: white (FFFFFF)
    band2h_spr = make_element("w:tblStylePr", {"w:type": "band2Horz"})
    b2h_tcPr = make_element("w:tcPr")
    b2h_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "FFFFFF"}))
    band2h_spr.append(b2h_tcPr)

    # band1Vert: light blue (BDD7EE)
    band1v_spr = make_element("w:tblStylePr", {"w:type": "band1Vert"})
    b1v_tcPr = make_element("w:tcPr")
    b1v_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "BDD7EE"}))
    band1v_spr.append(b1v_tcPr)

    # band2Vert: white (FFFFFF)
    band2v_spr = make_element("w:tblStylePr", {"w:type": "band2Vert"})
    b2v_tcPr = make_element("w:tcPr")
    b2v_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "FFFFFF"}))
    band2v_spr.append(b2v_tcPr)

    _add_table_style(doc, "BandInteraction", "Band Interaction",
                     tblStylePrs=[band1h_spr, band2h_spr,
                                  band1v_spr, band2v_spr])

    doc.add_paragraph("Band row + band column interaction.")
    tbl = doc.add_table(rows=4, cols=4)
    _apply_table_style(tbl, "BandInteraction", {
        "w:val": "0000",
        "w:firstRow": "0", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "0", "w:noVBand": "0",
    })
    labels = [[f"R{r}C{c}" for c in range(4)] for r in range(4)]
    _fill_table_cells(tbl, labels)

    save_fixture("table-cascade", "band-interaction", doc, {
        "name": "band-interaction",
        "spec_ref": "ISO 29500-1 §17.7.6",
        "description": (
            "Table with band1Horz (gray D9D9D9), band2Horz (white), "
            "band1Vert (blue BDD7EE), band2Vert (white). "
            "Both horizontal and vertical banding enabled."
        ),
        "expected_behavior": (
            "Per 17.7.6: banded rows override banded columns. "
            "Cell (0,0): band1Horz row + band1Vert col -> gray (D9D9D9) from band row. "
            "Cell (0,1): band1Horz row + band2Vert col -> gray (D9D9D9) from band row. "
            "Cell (1,0): band2Horz row + band1Vert col -> white from band row. "
            "Band rows always win over band columns at intersection."
        ),
        "current_status": "GAP — current code applies band columns after band rows (wrong precedence)",
    })


def make_tc_style_table_level_props() -> None:
    """Table style with table-level properties (alignment, indent).

    ISO 29500-1 §17.4.63: Table style can define tblPr with alignment (jc),
    indent (tblInd), and other table-level properties that should propagate.
    """
    doc = Document()

    # Table style with center alignment and indent
    tblPr_children = []
    tblPr_children.append(make_element("w:jc", {"w:val": "center"}))
    tblPr_children.append(make_element("w:tblInd", {
        "w:w": "720", "w:type": "dxa",
    }))

    # Add borders for visibility
    tbl_borders = make_element("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "4", "w:color": "000000",
        }))
    tblPr_children.append(tbl_borders)

    _add_table_style(doc, "AlignedTable", "Aligned Table",
                     tblPr_children=tblPr_children)

    doc.add_paragraph("Table style with alignment and indent.")
    tbl = doc.add_table(rows=2, cols=2)
    _apply_table_style(tbl, "AlignedTable", {
        "w:val": "04A0",
        "w:firstRow": "0", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "1",
    })
    _fill_table_cells(tbl, [
        ["Cell A", "Cell B"],
        ["Cell C", "Cell D"],
    ])

    save_fixture("table-cascade", "style-table-props", doc, {
        "name": "style-table-props",
        "spec_ref": "ISO 29500-1 §17.4.63",
        "description": "Table style sets jc=center and tblInd=720. Table has no direct alignment.",
        "expected_behavior": (
            "Table should inherit center alignment and 720tw indent from style. "
            "Currently TableStyleProps does not carry alignment or indent."
        ),
        "current_status": "GAP — TableStyleProps missing alignment and indent fields",
    })


# =========================================================================
# INDENT & SPACING EDGE CASES (ISO 29500-1 §17.3.1.12, §17.3.1.33)
# =========================================================================

def make_indent_edge_case_fixtures() -> None:
    print("\n── Indent & Spacing Edge Cases ──")
    make_iec_hanging_with_tabs()
    make_iec_mirror_indent()
    make_iec_numbering_indent_partial_override()
    make_iec_spacing_element_override()
    make_iec_zero_indent_explicit()
    make_iec_firstline_and_hanging_both_present()
    make_iec_indent_with_numbering_tab()


def make_iec_hanging_with_tabs() -> None:
    """Hanging indent with tab characters in body text.

    ISO 29500-1 §17.3.1.12: w:left="1440" w:hanging="720" means:
      left margin = 1440 twips, first line pulled back by 720 twips.
      Effective first-line edge = 720 twips.
    The hanging indent should be preserved even when body text contains tabs.
    Tab characters should NOT absorb or cancel the hanging indent value.
    """
    doc = Document()

    p = doc.add_paragraph("")
    # Clear default run, add text with embedded tab
    p.clear()
    run = p.add_run("Header\tValue")
    pPr = p._p.get_or_add_pPr()
    pPr.append(make_element("w:ind", {
        "w:left": "1440",
        "w:hanging": "720",
    }))

    # Control: same indent, no tabs
    p2 = doc.add_paragraph("No tabs, same hanging indent")
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:ind", {
        "w:left": "1440",
        "w:hanging": "720",
    }))

    save_fixture("indent-edge-cases", "hanging-with-tabs", doc, {
        "name": "hanging-with-tabs",
        "spec_ref": "ISO 29500-1 §17.3.1.12",
        "description": (
            "Paragraph with w:left=1440 w:hanging=720 and body text containing "
            "a tab character. Hanging indent should be preserved."
        ),
        "expected_behavior": (
            "P0: left=1440, first_line=-720 (hanging preserved despite tabs). "
            "P1: left=1440, first_line=-720 (control, no tabs)."
        ),
        "current_status": "TESTING",
    })


def make_iec_mirror_indent() -> None:
    """Mirror indent using w:start / w:end instead of w:left / w:right.

    ISO 29500-1 §17.3.1.12: w:start and w:end are logical-direction
    equivalents of w:left and w:right. For LTR text, w:start maps to
    left and w:end maps to right.
    """
    doc = Document()

    # Para with w:start and w:end
    p = doc.add_paragraph("Mirror indent: start=720, end=360")
    pPr = p._p.get_or_add_pPr()
    pPr.append(make_element("w:ind", {
        "w:start": "720",
        "w:end": "360",
    }))

    # Control: same values using w:left and w:right
    p2 = doc.add_paragraph("Standard indent: left=720, right=360")
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:ind", {
        "w:left": "720",
        "w:right": "360",
    }))

    save_fixture("indent-edge-cases", "mirror-indent", doc, {
        "name": "mirror-indent",
        "spec_ref": "ISO 29500-1 §17.3.1.12",
        "description": (
            "Paragraph using w:start/w:end (logical direction) instead of "
            "w:left/w:right. For LTR text, start=left and end=right."
        ),
        "expected_behavior": (
            "P0: left=720, right=360 (from w:start/w:end). "
            "P1: left=720, right=360 (from w:left/w:right — same result)."
        ),
        "current_status": "TESTING",
    })


def make_iec_numbering_indent_partial_override() -> None:
    """Numbering indent with partial direct override — no hanging leak.

    ISO 29500-1 §17.9.22 + §17.3.1.12: Numbering level defines
    left=1440, hanging=360. Paragraph has direct w:ind with ONLY
    w:left="720" (no firstLine or hanging). Per element-level override,
    the entire w:ind is replaced: left=720, first_line=None.
    The numbering hanging should NOT leak through.
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:pPr>
        <w:ind w:left="1440" w:hanging="360"/>
      </w:pPr>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # Paragraph with numbering + direct w:ind w:left="720" only
    p = doc.add_paragraph("Numbered with direct left=720 only, no firstLine")
    pPr = p._p.get_or_add_pPr()
    numPr = make_element("w:numPr")
    numPr.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr.append(make_element("w:numId", {"w:val": "1"}))
    pPr.append(numPr)
    pPr.append(make_element("w:ind", {"w:left": "720"}))

    # Control: numbering only
    p2 = doc.add_paragraph("Numbered with no direct indent")
    pPr2 = p2._p.get_or_add_pPr()
    numPr2 = make_element("w:numPr")
    numPr2.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr2.append(make_element("w:numId", {"w:val": "1"}))
    pPr2.append(numPr2)

    save_fixture("indent-edge-cases", "numbering-indent-partial-override", doc, {
        "name": "numbering-indent-partial-override",
        "spec_ref": "ISO 29500-1 §17.9.22, §17.3.1.12",
        "description": (
            "Numbering level has left=1440, hanging=360. Paragraph has direct "
            "w:ind with only w:left=720. Element-level override means the "
            "numbering hanging should NOT leak through."
        ),
        "expected_behavior": (
            "P0: left=720, first_line=None (element-level override). "
            "P1: left=1440, first_line=-360 (numbering only, control)."
        ),
        "current_status": "TESTING",
    })


def make_iec_spacing_element_override() -> None:
    """Spacing element-level override question.

    Style defines w:spacing before=240, after=120, line=276.
    Paragraph has direct w:spacing with ONLY before="0".

    Question: Does w:spacing use element-level override (like w:ind) or
    per-field merge? This test documents the element-level override
    interpretation for consistency with w:ind behavior.

    Note: The existing style-cascade tests already test per-field merge for
    spacing. This fixture specifically tests whether element-level override
    applies when direct w:spacing is present.
    """
    doc = Document()
    styles_element = doc.styles.element

    # Remove existing docDefaults
    existing_defaults = styles_element.find(w("docDefaults"))
    if existing_defaults is not None:
        styles_element.remove(existing_defaults)

    # Minimal docDefaults
    doc_defaults = make_element("w:docDefaults")
    rpr_default = make_element("w:rPrDefault")
    rpr = make_element("w:rPr")
    rpr.append(make_element("w:sz", {"w:val": "24"}))
    rpr_default.append(rpr)
    doc_defaults.append(rpr_default)
    styles_element.insert(0, doc_defaults)

    # Style with full spacing
    style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "FullSpacedStyle"})
    style.append(make_element("w:name", {"w:val": "Full Spaced Style"}))
    ppr = make_element("w:pPr")
    ppr.append(make_element("w:spacing", {
        "w:before": "240",
        "w:after": "120",
        "w:line": "276",
        "w:lineRule": "auto",
    }))
    style.append(ppr)
    styles_element.append(style)

    # Para 0: style only (control)
    p0 = doc.add_paragraph("Style only: before=240, after=120, line=276")
    pPr0 = p0._p.get_or_add_pPr()
    pPr0.append(make_element("w:pStyle", {"w:val": "FullSpacedStyle"}))

    # Para 1: style + direct w:spacing with only before=0
    # Element-level override: after and line become None (not inherited)
    p1 = doc.add_paragraph("Direct before=0: element-level override means after/line are None")
    pPr1 = p1._p.get_or_add_pPr()
    pPr1.append(make_element("w:pStyle", {"w:val": "FullSpacedStyle"}))
    pPr1.append(make_element("w:spacing", {"w:before": "0"}))

    save_fixture("indent-edge-cases", "spacing-element-override", doc, {
        "name": "spacing-element-override",
        "spec_ref": "ISO 29500-1 §17.3.1.33",
        "description": (
            "Style defines spacing before=240, after=120, line=276. "
            "Paragraph has direct w:spacing with only before=0. "
            "Tests element-level override vs per-field merge."
        ),
        "expected_behavior": (
            "If element-level override: P1 has before=0, after=None, line=None. "
            "If per-field merge: P1 has before=0, after=120, line=276. "
            "This test asserts element-level override for consistency with w:ind."
        ),
        "current_status": "TESTING — open question per GAP-027",
    })


def make_iec_zero_indent_explicit() -> None:
    """Explicit w:left="0" vs absent w:ind.

    A style defines w:left="720". A paragraph with explicit w:ind w:left="0"
    should override the style indent to 0. This is different from a paragraph
    with no w:ind at all (which inherits from style).
    """
    doc = Document()
    styles_element = doc.styles.element

    # Style with left=720
    style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "Indented720"})
    style.append(make_element("w:name", {"w:val": "Indented 720"}))
    ppr = make_element("w:pPr")
    ppr.append(make_element("w:ind", {"w:left": "720"}))
    style.append(ppr)
    styles_element.append(style)

    # Para 0: style only — should get left=720 from style
    p0 = doc.add_paragraph("Style only: left=720 from style")
    pPr0 = p0._p.get_or_add_pPr()
    pPr0.append(make_element("w:pStyle", {"w:val": "Indented720"}))

    # Para 1: style + explicit w:left="0" — should override to left=0
    p1 = doc.add_paragraph("Explicit left=0: overrides style to 0")
    pPr1 = p1._p.get_or_add_pPr()
    pPr1.append(make_element("w:pStyle", {"w:val": "Indented720"}))
    pPr1.append(make_element("w:ind", {"w:left": "0"}))

    # Para 2: no style, no indent — should have no indent
    p2 = doc.add_paragraph("No style, no indent: indent should be None")

    save_fixture("indent-edge-cases", "zero-indent-explicit", doc, {
        "name": "zero-indent-explicit",
        "spec_ref": "ISO 29500-1 §17.3.1.12",
        "description": (
            "Style defines w:left=720. Paragraph with explicit w:left=0 should "
            "override to 0. Paragraph with no w:ind inherits from style."
        ),
        "expected_behavior": (
            "P0: left=720 (from style). "
            "P1: left=0 (explicit override). "
            "P2: indent=None (no style, no direct)."
        ),
        "current_status": "TESTING",
    })


def make_iec_firstline_and_hanging_both_present() -> None:
    """Both firstLine and hanging on same w:ind element.

    ISO 29500-1 §17.3.1.12: firstLine and hanging are mutually exclusive.
    If both are present, the behavior is technically undefined/invalid.
    Test what our parser does with this edge case.
    """
    doc = Document()

    # Para with both firstLine and hanging
    p = doc.add_paragraph("Both firstLine=360 and hanging=720 — invalid per spec")
    pPr = p._p.get_or_add_pPr()
    pPr.append(make_element("w:ind", {
        "w:left": "1440",
        "w:firstLine": "360",
        "w:hanging": "720",
    }))

    # Control: only firstLine
    p2 = doc.add_paragraph("Only firstLine=360")
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:ind", {
        "w:left": "1440",
        "w:firstLine": "360",
    }))

    # Control: only hanging
    p3 = doc.add_paragraph("Only hanging=720")
    pPr3 = p3._p.get_or_add_pPr()
    pPr3.append(make_element("w:ind", {
        "w:left": "1440",
        "w:hanging": "720",
    }))

    save_fixture("indent-edge-cases", "firstline-and-hanging-both-present", doc, {
        "name": "firstline-and-hanging-both-present",
        "spec_ref": "ISO 29500-1 §17.3.1.12",
        "description": (
            "w:ind with both w:firstLine and w:hanging attributes. "
            "Per spec, these are mutually exclusive."
        ),
        "expected_behavior": (
            "P0: Undefined — test documents actual parser behavior. "
            "P1: left=1440, first_line=360 (firstLine only). "
            "P2: left=1440, first_line=-720 (hanging only)."
        ),
        "current_status": "TESTING — spec-invalid input",
    })


def make_iec_indent_with_numbering_tab() -> None:
    """Numbering level with indent and tab stop — structural numbering.

    ISO 29500-1 §17.9.22 + §17.3.1.21: Numbering level defines
    indent w:left="720" w:hanging="360" AND a tab stop at 720.
    The numbering tab positions the number text, and the hanging
    indent at 360 offsets the first line.
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:pPr>
        <w:ind w:left="720" w:hanging="360"/>
        <w:tabs>
          <w:tab w:val="num" w:pos="720"/>
        </w:tabs>
      </w:pPr>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # Auto-numbered paragraph
    p = doc.add_paragraph("Auto-numbered with tab stop at 720")
    pPr = p._p.get_or_add_pPr()
    numPr = make_element("w:numPr")
    numPr.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr.append(make_element("w:numId", {"w:val": "1"}))
    pPr.append(numPr)

    # Second numbered paragraph
    p2 = doc.add_paragraph("Second auto-numbered paragraph")
    pPr2 = p2._p.get_or_add_pPr()
    numPr2 = make_element("w:numPr")
    numPr2.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr2.append(make_element("w:numId", {"w:val": "1"}))
    pPr2.append(numPr2)

    save_fixture("indent-edge-cases", "indent-with-numbering-tab", doc, {
        "name": "indent-with-numbering-tab",
        "spec_ref": "ISO 29500-1 §17.9.22, §17.3.1.21",
        "description": (
            "Numbering level with indent left=720, hanging=360 and tab stop "
            "at pos=720. Tests structural auto-numbering indent behavior."
        ),
        "expected_behavior": (
            "P0: left=720, first_line=-360 (from numbering level). "
            "P1: left=720, first_line=-360 (second auto-numbered item)."
        ),
        "current_status": "TESTING",
    })


# =========================================================================
# TAB STOP / INDENTATION INTERACTION (ISO 29500-1 §17.3.1.12 + §17.3.1.38)
# =========================================================================

def make_tab_indent_interaction_fixtures() -> None:
    print("\n── Tab + Indent Interaction ──")
    make_tii_prefix_negative_indent()
    make_tii_prefix_positive_indent()
    make_tii_prefix_hanging_indent()
    make_tii_tab_absorption_zero_firstline()
    make_tii_prefix_space_separator()
    make_tii_negative_indent_no_prefix()


def make_tii_prefix_negative_indent() -> None:
    """Prefix with negative indent: continuation lines should use w:left, not body_left.

    ISO 29500-1 §17.3.1.12: w:left specifies the indentation for ALL lines
    except the first (which adds w:firstLine on top). When a prefix like "(a)\\t"
    is stripped, the first line starts at the tab stop position (360), but
    continuation lines must still use the resolved w:left (-720).

    The correct first_line is: tab_stop_pos - w:left = 360 - (-720) = 1080.
    """
    doc = Document()

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()

    # Indentation: left=-720, right=-360, no hanging/firstLine
    pPr.append(make_element("w:ind", {"w:left": "-720", "w:right": "-360"}))

    # Tab stop at position 360
    tabs_el = make_element("w:tabs")
    tabs_el.append(make_element("w:tab", {"w:val": "left", "w:pos": "360"}))
    pPr.append(tabs_el)

    # Text with prefix pattern "(a)\t" followed by body text
    run = p.add_run("(a)\tEquity Financing. The Company shall issue shares...")

    save_fixture("tab-indent-interaction", "prefix-negative-indent", doc, {
        "name": "prefix-negative-indent",
        "spec_ref": "ISO 29500-1 §17.3.1.12, §17.3.1.38",
        "description": (
            "Paragraph with w:left=-720, w:right=-360, tab at 360, "
            "text '(a)\\tEquity Financing...'. Tests that continuation lines "
            "use w:left, not the tab stop position."
        ),
        "expected_behavior": (
            "After prefix stripping: left=-720 (continuation indent), "
            "first_line=1080 (=360-(-720), offset from continuation to first-line edge). "
            "Current bug: left=360, first_line=None."
        ),
        "current_status": "GAP — literal_prefix branch forces left=body_left, drops first_line",
    })


def make_tii_prefix_positive_indent() -> None:
    """Prefix with positive indent: same class of bug but with positive w:left.

    ISO 29500-1 §17.3.1.12: w:left=720, tab at 1440. After stripping "(a)\\t",
    continuation lines should be at w:left=720. First line at tab stop 1440.
    first_line = 1440 - 720 = 720.
    """
    doc = Document()

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()

    # Indentation: left=720, no hanging/firstLine
    pPr.append(make_element("w:ind", {"w:left": "720"}))

    # Tab stop at position 1440
    tabs_el = make_element("w:tabs")
    tabs_el.append(make_element("w:tab", {"w:val": "left", "w:pos": "1440"}))
    pPr.append(tabs_el)

    # Text with prefix pattern
    run = p.add_run("(a)\tBody text that wraps to test continuation indent.")

    save_fixture("tab-indent-interaction", "prefix-positive-indent", doc, {
        "name": "prefix-positive-indent",
        "spec_ref": "ISO 29500-1 §17.3.1.12, §17.3.1.38",
        "description": (
            "Paragraph with w:left=720, tab at 1440, text '(a)\\tBody text...'. "
            "Same bug class as negative indent but with positive values."
        ),
        "expected_behavior": (
            "After prefix stripping: left=720 (continuation indent), "
            "first_line=720 (=1440-720). "
            "Current bug: left=1440, first_line=None."
        ),
        "current_status": "GAP — literal_prefix branch forces left=body_left",
    })


def make_tii_prefix_hanging_indent() -> None:
    """Prefix with hanging indent: w:hanging should interact correctly with tab position.

    ISO 29500-1 §17.3.1.12: w:left=720, w:hanging=360 means continuation at 720,
    first line at 720-360=360. With a tab stop at 720 and prefix "1.\\t", after
    stripping the prefix the body starts at the tab stop position (720).
    first_line = 720 - 720 = 0 (first line and continuation are at the same position).
    """
    doc = Document()

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()

    # Indentation: left=720, hanging=360
    pPr.append(make_element("w:ind", {"w:left": "720", "w:hanging": "360"}))

    # Tab stop at position 720
    tabs_el = make_element("w:tabs")
    tabs_el.append(make_element("w:tab", {"w:val": "left", "w:pos": "720"}))
    pPr.append(tabs_el)

    # Text with prefix pattern
    run = p.add_run("1.\tBody text that wraps to test hanging indent.")

    save_fixture("tab-indent-interaction", "prefix-hanging-indent", doc, {
        "name": "prefix-hanging-indent",
        "spec_ref": "ISO 29500-1 §17.3.1.12, §17.3.1.38",
        "description": (
            "Paragraph with w:left=720, w:hanging=360, tab at 720, "
            "text '1.\\tBody text...'. Tests hanging indent preserved after prefix stripping."
        ),
        "expected_behavior": (
            "After prefix stripping: left=720 (continuation indent), "
            "first_line=0 (=720-720, body starts at same position as continuation). "
            "Current bug: left=720, first_line=None (coincidentally correct left, wrong first_line logic)."
        ),
        "current_status": "GAP — literal_prefix branch uses body_left instead of resolved w:left",
    })


def make_tii_tab_absorption_zero_firstline() -> None:
    """Tab absorption when firstLine is 0: left indent should not be modified.

    ISO 29500-1 §17.3.1.12: When w:left=720 and no firstLine (or firstLine=0),
    and body text contains tabs, the tab absorption branch (runtime.rs:3519)
    should not fire because first_line_twips==0. The left indent should pass through.
    """
    doc = Document()

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()

    # Indentation: left=720, no firstLine
    pPr.append(make_element("w:ind", {"w:left": "720"}))

    # Tab stop at 1440
    tabs_el = make_element("w:tabs")
    tabs_el.append(make_element("w:tab", {"w:val": "left", "w:pos": "1440"}))
    pPr.append(tabs_el)

    # Body text with tabs (no prefix pattern)
    run = p.add_run("A\tB")

    save_fixture("tab-indent-interaction", "tab-absorption-zero-firstline", doc, {
        "name": "tab-absorption-zero-firstline",
        "spec_ref": "ISO 29500-1 §17.3.1.12, §17.3.1.38",
        "description": (
            "Paragraph with w:left=720, no firstLine, text 'A\\tB'. "
            "Tests that tab absorption doesn't fire when firstLine is 0."
        ),
        "expected_behavior": (
            "left=720, first_line=None (passthrough). "
            "Tab absorption only applies when first_line_twips!=0."
        ),
        "current_status": "TESTING — this branch may already work correctly",
    })


def make_tii_prefix_space_separator() -> None:
    """Space-separated prefix should not trigger tab-stop-based body_left computation.

    ISO 29500-1 §17.3.1.12: When a prefix is separated by a space (not a tab),
    there is no tab consumed by the prefix, so body_left should not be set from
    a tab stop position. The resolved indentation should pass through unchanged.
    """
    doc = Document()

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()

    # Indentation: left=360
    pPr.append(make_element("w:ind", {"w:left": "360"}))

    # Tab stop (should not be consumed by the space-separated prefix)
    tabs_el = make_element("w:tabs")
    tabs_el.append(make_element("w:tab", {"w:val": "left", "w:pos": "720"}))
    pPr.append(tabs_el)

    # Text with space-separated prefix (not tab)
    run = p.add_run("(a) Body text with space separator, not tab.")

    save_fixture("tab-indent-interaction", "prefix-space-separator", doc, {
        "name": "prefix-space-separator",
        "spec_ref": "ISO 29500-1 §17.3.1.12, §17.3.1.38",
        "description": (
            "Paragraph with w:left=360, tab at 720, text '(a) Body text...' "
            "(space separator, not tab). Tests that space-separated prefix "
            "doesn't incorrectly use tab stop logic."
        ),
        "expected_behavior": (
            "left=360 (from resolved indent, unchanged by prefix stripping). "
            "Space-separated prefix does not consume a tab stop."
        ),
        "current_status": "TESTING — space prefix may or may not trigger the bug",
    })


def make_tii_negative_indent_no_prefix() -> None:
    """Negative indent with tab but no prefix pattern: tab absorption without prefix stripping.

    ISO 29500-1 §17.3.1.12: w:left=-360, tab at 720, text starts with tab
    but no prefix pattern like "(a)". The tab absorption branch (runtime.rs:3519)
    should fire if first_line_twips != 0, but since no firstLine is set, this
    should fall through to passthrough.
    """
    doc = Document()

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()

    # Indentation: left=-360, no firstLine
    pPr.append(make_element("w:ind", {"w:left": "-360"}))

    # Tab stop at 720
    tabs_el = make_element("w:tabs")
    tabs_el.append(make_element("w:tab", {"w:val": "left", "w:pos": "720"}))
    pPr.append(tabs_el)

    # Text with tab but no prefix pattern
    run = p.add_run("\tBody text starting with tab, no prefix pattern.")

    save_fixture("tab-indent-interaction", "negative-indent-no-prefix", doc, {
        "name": "negative-indent-no-prefix",
        "spec_ref": "ISO 29500-1 §17.3.1.12, §17.3.1.38",
        "description": (
            "Paragraph with w:left=-360, tab at 720, text '\\tBody text...' "
            "(tab but no prefix). Tests tab absorption without prefix stripping."
        ),
        "expected_behavior": (
            "left=-360, first_line=None (no firstLine set, tab absorption should not fire). "
            "The tab positions the text but doesn't change the indent model."
        ),
        "current_status": "TESTING — tab absorption branch gated on first_line_twips!=0",
    })


# =========================================================================
# TABLE STYLE RESOLUTION (ISO 29500-1 §17.4, §17.7.6)
# =========================================================================

def make_table_style_resolution_fixtures() -> None:
    print("\n── Table Style Resolution ──")
    make_tsr_border_conflict_cell_vs_table()
    make_tsr_conditional_font_formatting()
    make_tsr_conditional_paragraph_alignment()
    make_tsr_nested_table_style_inherit()
    make_tsr_table_style_indent_alignment()
    make_tsr_row_banding_alternation()


def make_tsr_border_conflict_cell_vs_table() -> None:
    """Cell borders vs table-level borders from table style.

    ISO 29500-1 §17.4.39-44: When cell-level borders conflict with
    table-level borders (from the style), the cell-level borders win.
    This tests the case where the TABLE STYLE provides the table borders
    (unlike the edge-case border-conflict fixture which uses direct table borders).
    """
    doc = Document()

    # Table style with thick red borders on all edges
    tbl_borders = make_element("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "18", "w:color": "FF0000",
        }))

    _add_table_style(doc, "BorderConflictStyle", "Border Conflict Style",
                     tblPr_children=[tbl_borders])

    doc.add_paragraph("Border conflict: cell vs table style borders.")
    tbl = doc.add_table(rows=2, cols=2)
    _apply_table_style(tbl, "BorderConflictStyle", {
        "w:val": "04A0",
        "w:firstRow": "0", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "1",
    })
    _fill_table_cells(tbl, [
        ["Styled borders", "Styled borders"],
        ["Cell override", "Styled borders"],
    ])

    # Cell (1,0) gets explicit thin blue borders
    cell_10 = tbl.cell(1, 0)
    tc_el = cell_10._tc
    tcPr = tc_el.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = make_element("w:tcPr")
        tc_el.insert(0, tcPr)
    tc_borders = make_element("w:tcBorders")
    for side in ["top", "bottom", "left", "right"]:
        tc_borders.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "4", "w:color": "0000FF",
        }))
    tcPr.append(tc_borders)

    save_fixture("table-style-resolution", "border-conflict-cell-vs-table", doc, {
        "name": "border-conflict-cell-vs-table",
        "spec_ref": "ISO 29500-1 §17.4.39-44",
        "description": (
            "Table style provides thick red (sz=18) borders on all edges. "
            "Cell (1,0) has explicit thin blue (sz=4) tcBorders."
        ),
        "expected_behavior": (
            "Cell (1,0): blue sz=4 borders (cell overrides style table borders). "
            "Cell (0,0), (0,1), (1,1): no cell-level borders (inherit from style)."
        ),
        "current_status": "GAP — testing cell vs style-provided table border precedence",
    })


def make_tsr_conditional_font_formatting() -> None:
    """Conditional firstRow with run formatting (bold + font size).

    ISO 29500-1 §17.7.6.2: tblStylePr can contain rPr (run properties).
    This tests bold AND font size together (unlike the table-cascade fixture
    which only tests bold + color).
    """
    doc = Document()

    # firstRow conditional: bold + 14pt font
    first_row_spr = make_element("w:tblStylePr", {"w:type": "firstRow"})
    fr_rPr = make_element("w:rPr")
    fr_rPr.append(make_element("w:b"))
    fr_rPr.append(make_element("w:sz", {"w:val": "28"}))  # 14pt = 28 half-points
    first_row_spr.append(fr_rPr)
    fr_tcPr = make_element("w:tcPr")
    fr_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "2E75B6"}))
    first_row_spr.append(fr_tcPr)

    # Base table borders
    tbl_borders = make_element("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "4", "w:color": "000000",
        }))

    _add_table_style(doc, "FontCondStyle", "Font Conditional Style",
                     tblPr_children=[tbl_borders],
                     tblStylePrs=[first_row_spr])

    doc.add_paragraph("Conditional font formatting test.")
    tbl = doc.add_table(rows=3, cols=2)
    _apply_table_style(tbl, "FontCondStyle", {
        "w:val": "04A0",
        "w:firstRow": "1", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "1",
    })
    _fill_table_cells(tbl, [
        ["Header A", "Header B"],
        ["Data 1", "Data 2"],
        ["Data 3", "Data 4"],
    ])

    save_fixture("table-style-resolution", "conditional-font-formatting", doc, {
        "name": "conditional-font-formatting",
        "spec_ref": "ISO 29500-1 §17.7.6.2",
        "description": (
            "Table style with firstRow conditional: bold + sz=28 (14pt) in rPr, "
            "blue shading (2E75B6) in tcPr."
        ),
        "expected_behavior": (
            "First row cells: text should be bold and 14pt from conditional rPr. "
            "Data rows: default font formatting."
        ),
        "current_status": "GAP-020 — conditional rPr not propagated to runs",
    })


def make_tsr_conditional_paragraph_alignment() -> None:
    """Conditional firstRow with paragraph alignment (center).

    ISO 29500-1 §17.7.6.1: tblStylePr can contain pPr that applies to
    paragraphs within matching cells. This tests center alignment on firstRow
    combined with right alignment on lastRow to verify both conditions work.
    """
    doc = Document()

    # firstRow conditional: center alignment
    first_row_spr = make_element("w:tblStylePr", {"w:type": "firstRow"})
    fr_pPr = make_element("w:pPr")
    fr_pPr.append(make_element("w:jc", {"w:val": "center"}))
    first_row_spr.append(fr_pPr)

    # lastRow conditional: right alignment
    last_row_spr = make_element("w:tblStylePr", {"w:type": "lastRow"})
    lr_pPr = make_element("w:pPr")
    lr_pPr.append(make_element("w:jc", {"w:val": "right"}))
    last_row_spr.append(lr_pPr)

    _add_table_style(doc, "ParaAlignStyle", "Para Alignment Style",
                     tblStylePrs=[first_row_spr, last_row_spr])

    doc.add_paragraph("Conditional paragraph alignment test.")
    tbl = doc.add_table(rows=4, cols=2)
    _apply_table_style(tbl, "ParaAlignStyle", {
        "w:val": "04A0",
        "w:firstRow": "1", "w:lastRow": "1",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "1",
    })
    _fill_table_cells(tbl, [
        ["Center 1", "Center 2"],
        ["Normal A", "Normal B"],
        ["Normal C", "Normal D"],
        ["Right 1", "Right 2"],
    ])

    save_fixture("table-style-resolution", "conditional-paragraph-alignment", doc, {
        "name": "conditional-paragraph-alignment",
        "spec_ref": "ISO 29500-1 §17.7.6.1",
        "description": (
            "Table style with firstRow conditional (jc=center in pPr) and "
            "lastRow conditional (jc=right in pPr)."
        ),
        "expected_behavior": (
            "First row paragraphs: center alignment from conditional pPr. "
            "Last row paragraphs: right alignment from conditional pPr. "
            "Middle rows: no alignment (default/left)."
        ),
        "current_status": "GAP-020 — conditional pPr not propagated to paragraphs",
    })


def make_tsr_nested_table_style_inherit() -> None:
    """Nested table with its own style resolves independently.

    ISO 29500-1 §17.4.38: A table cell can contain a nested table.
    The inner table's style should resolve independently — it should NOT
    inherit conditional formatting or shading from the outer table's style.
    """
    doc = Document()

    # Outer table style: green shading on all cells
    outer_whole_spr = make_element("w:tblStylePr", {"w:type": "wholeTable"})
    ow_tcPr = make_element("w:tcPr")
    ow_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "C6EFCE"}))
    outer_whole_spr.append(ow_tcPr)

    tbl_borders_outer = make_element("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tbl_borders_outer.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "8", "w:color": "00B050",
        }))

    _add_table_style(doc, "OuterGreen", "Outer Green Style",
                     tblPr_children=[tbl_borders_outer],
                     tblStylePrs=[outer_whole_spr])

    # Inner table style: orange shading on all cells
    inner_whole_spr = make_element("w:tblStylePr", {"w:type": "wholeTable"})
    iw_tcPr = make_element("w:tcPr")
    iw_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "FBE5D6"}))
    inner_whole_spr.append(iw_tcPr)

    tbl_borders_inner = make_element("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tbl_borders_inner.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "4", "w:color": "ED7D31",
        }))

    _add_table_style(doc, "InnerOrange", "Inner Orange Style",
                     tblPr_children=[tbl_borders_inner],
                     tblStylePrs=[inner_whole_spr])

    doc.add_paragraph("Nested table with independent style.")
    # Create outer 2x2 table
    outer_tbl = doc.add_table(rows=2, cols=2)
    _apply_table_style(outer_tbl, "OuterGreen", {
        "w:val": "04A0",
        "w:firstRow": "0", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "1",
    })
    _fill_table_cells(outer_tbl, [
        ["Outer A", "Outer B"],
        ["", "Outer D"],
    ])

    # Insert inner 2x1 table into cell (1,0) using lxml
    cell_10 = outer_tbl.cell(1, 0)
    tc_el = cell_10._tc
    # Remove existing paragraphs
    for p in tc_el.findall(qn("w:p")):
        tc_el.remove(p)

    # Build inner table XML manually
    inner_tbl = make_element("w:tbl")
    inner_tblPr = make_element("w:tblPr")
    inner_tblPr.append(make_element("w:tblStyle", {"w:val": "InnerOrange"}))
    inner_tblPr.append(make_element("w:tblW", {"w:w": "0", "w:type": "auto"}))
    inner_tblPr.append(make_element("w:tblLook", {
        "w:val": "04A0",
        "w:firstRow": "0", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "1",
    }))
    inner_tbl.append(inner_tblPr)

    inner_grid = make_element("w:tblGrid")
    inner_grid.append(make_element("w:gridCol", {"w:w": "2000"}))
    inner_tbl.append(inner_grid)

    for label in ["Inner X", "Inner Y"]:
        tr = make_element("w:tr")
        tc = make_element("w:tc")
        tc_p = make_element("w:p")
        tc_r = make_element("w:r")
        tc_t = make_element("w:t")
        tc_t.text = label
        tc_r.append(tc_t)
        tc_p.append(tc_r)
        tc.append(tc_p)
        tr.append(tc)
        inner_tbl.append(tr)

    tc_el.append(inner_tbl)
    # Must have a trailing paragraph per spec
    tc_el.append(make_element("w:p"))

    save_fixture("table-style-resolution", "nested-table-style-inherit", doc, {
        "name": "nested-table-style-inherit",
        "spec_ref": "ISO 29500-1 §17.4.38",
        "description": (
            "Outer table with green style (C6EFCE shading, green borders). "
            "Inner table in cell (1,0) with orange style (FBE5D6 shading, orange borders)."
        ),
        "expected_behavior": (
            "Outer table cells: green shading (C6EFCE). "
            "Inner table cells: orange shading (FBE5D6), resolved independently. "
            "Inner table should NOT inherit outer table's green shading."
        ),
        "current_status": "GAP-003 — testing nested table style independence",
    })


def make_tsr_table_style_indent_alignment() -> None:
    """Table-level properties from style: tblInd and jc.

    ISO 29500-1 §17.4.63: Table style tblPr can define table alignment (jc)
    and table indent (tblInd). These should propagate when the table has no
    direct overrides. Unlike style-table-props in table-cascade, this fixture
    also includes a second table with direct overrides to test precedence.
    """
    doc = Document()

    # Table style with center alignment + indent
    tblPr_children = []
    tblPr_children.append(make_element("w:jc", {"w:val": "center"}))
    tblPr_children.append(make_element("w:tblInd", {
        "w:w": "360", "w:type": "dxa",
    }))

    tbl_borders = make_element("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "4", "w:color": "000000",
        }))
    tblPr_children.append(tbl_borders)

    _add_table_style(doc, "IndentAlignStyle", "Indent Align Style",
                     tblPr_children=tblPr_children)

    # Table 1: no direct alignment/indent — should inherit from style
    doc.add_paragraph("Table 1: style provides alignment and indent.")
    tbl1 = doc.add_table(rows=2, cols=2)
    _apply_table_style(tbl1, "IndentAlignStyle", {
        "w:val": "04A0",
        "w:firstRow": "0", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "1",
    })
    _fill_table_cells(tbl1, [
        ["Style Center", "Style Center"],
        ["Indent 360", "Indent 360"],
    ])

    # Table 2: direct override — right alignment, no indent
    doc.add_paragraph("Table 2: direct override of alignment.")
    tbl2 = doc.add_table(rows=2, cols=2)
    _apply_table_style(tbl2, "IndentAlignStyle", {
        "w:val": "04A0",
        "w:firstRow": "0", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "1",
    })
    _fill_table_cells(tbl2, [
        ["Direct Right", "Direct Right"],
        ["No Indent", "No Indent"],
    ])
    # Add direct jc override on table 2
    tbl2_el = tbl2._tbl
    tbl2_Pr = tbl2_el.tblPr
    tbl2_Pr.append(make_element("w:jc", {"w:val": "right"}))

    save_fixture("table-style-resolution", "table-style-indent-alignment", doc, {
        "name": "table-style-indent-alignment",
        "spec_ref": "ISO 29500-1 §17.4.63",
        "description": (
            "Style defines jc=center and tblInd=360. "
            "Table 1 has no direct overrides. "
            "Table 2 has direct jc=right override."
        ),
        "expected_behavior": (
            "Table 1: center alignment and 360tw indent from style. "
            "Table 2: right alignment from direct override, "
            "360tw indent from style (only alignment is overridden)."
        ),
        "current_status": "GAP-021 — table alignment/indent not resolved from style",
    })


def make_tsr_row_banding_alternation() -> None:
    """Row banding alternation with band1Horz and band2Horz.

    ISO 29500-1 §17.7.6: band1Horz applies to odd rows (0, 2, 4...),
    band2Horz applies to even rows (1, 3, 5...) when banding is enabled.
    This tests a 6-row table to verify the banding cycles correctly beyond
    just two rows (unlike band-interaction which tests row vs column precedence).
    """
    doc = Document()

    # band1Horz: light blue (DAEEF3)
    band1h_spr = make_element("w:tblStylePr", {"w:type": "band1Horz"})
    b1h_tcPr = make_element("w:tcPr")
    b1h_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "DAEEF3"}))
    band1h_spr.append(b1h_tcPr)

    # band2Horz: light yellow (FFFFCC)
    band2h_spr = make_element("w:tblStylePr", {"w:type": "band2Horz"})
    b2h_tcPr = make_element("w:tcPr")
    b2h_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "FFFFCC"}))
    band2h_spr.append(b2h_tcPr)

    tbl_borders = make_element("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "4", "w:color": "999999",
        }))

    _add_table_style(doc, "BandingStyle", "Row Banding Style",
                     tblPr_children=[tbl_borders],
                     tblStylePrs=[band1h_spr, band2h_spr])

    doc.add_paragraph("Row banding alternation test.")
    tbl = doc.add_table(rows=6, cols=2)
    _apply_table_style(tbl, "BandingStyle", {
        "w:val": "0000",
        "w:firstRow": "0", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "0", "w:noVBand": "1",
    })
    labels = [[f"R{r}C{c}" for c in range(2)] for r in range(6)]
    _fill_table_cells(tbl, labels)

    save_fixture("table-style-resolution", "row-banding-alternation", doc, {
        "name": "row-banding-alternation",
        "spec_ref": "ISO 29500-1 §17.7.6",
        "description": (
            "Table with band1Horz (light blue DAEEF3), band2Horz (light yellow FFFFCC). "
            "6 rows, horizontal banding enabled, no firstRow/lastRow/col conditionals."
        ),
        "expected_behavior": (
            "Row 0: blue (DAEEF3, band1Horz). Row 1: yellow (FFFFCC, band2Horz). "
            "Row 2: blue. Row 3: yellow. Row 4: blue. Row 5: yellow. "
            "Banding cycles every 2 rows."
        ),
        "current_status": "GAP-003 — testing banding cycle behavior",
    })


# =========================================================================
# SECTIONS (ISO 29500-1 §17.6)
# =========================================================================

def make_section_fixtures() -> None:
    print("\n── Sections ──")
    make_page_margins()
    make_section_types()
    make_multi_section_margins()
    make_page_borders()
    make_line_numbering()
    make_vertical_alignment()
    make_page_number_format()


def make_page_margins() -> None:
    """pgMar — custom page margins.

    ISO 29500-1 §17.6.11: pgMar defines page margins (top, bottom, left,
    right, header, footer, gutter). All values in twips.
    """
    doc = Document()

    doc.add_paragraph("Content in a section with custom page margins.")

    p = doc.add_paragraph("Last paragraph of the section.")
    pPr = p._p.get_or_add_pPr()
    sect_pr = make_element("w:sectPr")
    pg_mar = make_element("w:pgMar", {
        "w:top": "1440",
        "w:bottom": "1440",
        "w:left": "1800",
        "w:right": "1800",
        "w:header": "720",
        "w:footer": "720",
        "w:gutter": "360",
    })
    sect_pr.append(pg_mar)
    pPr.append(sect_pr)

    save_fixture("sections", "page-margins", doc, {
        "name": "page-margins",
        "spec_ref": "ISO 29500-1 §17.6.11",
        "description": "Section with custom pgMar (top=1440, bottom=1440, left=1800, right=1800, header=720, footer=720, gutter=360)",
        "expected_behavior": "Parser should expose all margin values in SectionProperties",
        "current_status": "GAP — pgMar not yet parsed into SectionProperties",
    })


def make_section_types() -> None:
    """Section types — nextPage, continuous, evenPage.

    ISO 29500-1 §17.6.17: w:type in sectPr defines the section break type.
    """
    doc = Document()

    # Section 1: nextPage break
    doc.add_paragraph("Content in section 1 (nextPage break).")
    p1 = doc.add_paragraph("End of section 1.")
    pPr1 = p1._p.get_or_add_pPr()
    sect_pr1 = make_element("w:sectPr")
    sect_pr1.append(make_element("w:type", {"w:val": "nextPage"}))
    sect_pr1.append(make_element("w:pgSz", {"w:w": "12240", "w:h": "15840"}))
    pPr1.append(sect_pr1)

    # Section 2: continuous break
    doc.add_paragraph("Content in section 2 (continuous break).")
    p2 = doc.add_paragraph("End of section 2.")
    pPr2 = p2._p.get_or_add_pPr()
    sect_pr2 = make_element("w:sectPr")
    sect_pr2.append(make_element("w:type", {"w:val": "continuous"}))
    sect_pr2.append(make_element("w:pgSz", {"w:w": "12240", "w:h": "15840"}))
    pPr2.append(sect_pr2)

    # Section 3: evenPage break
    doc.add_paragraph("Content in section 3 (evenPage break).")

    save_fixture("sections", "section-types", doc, {
        "name": "section-types",
        "spec_ref": "ISO 29500-1 §17.6.17",
        "description": "Three sections: nextPage, continuous, evenPage break types",
        "expected_behavior": "Parser should expose section type in SectionProperties",
        "current_status": "GAP — section type not yet parsed into SectionProperties",
    })


def make_multi_section_margins() -> None:
    """Multiple sections with different margins.

    Tests that section-level margins are preserved per section independently.
    """
    doc = Document()

    # Section 1: narrow margins
    doc.add_paragraph("Content in section 1 with narrow margins.")
    p1 = doc.add_paragraph("End of section 1.")
    pPr1 = p1._p.get_or_add_pPr()
    sect_pr1 = make_element("w:sectPr")
    sect_pr1.append(make_element("w:pgMar", {
        "w:top": "720",
        "w:bottom": "720",
        "w:left": "720",
        "w:right": "720",
        "w:header": "360",
        "w:footer": "360",
        "w:gutter": "0",
    }))
    pPr1.append(sect_pr1)

    # Section 2: wide margins
    doc.add_paragraph("Content in section 2 with wide margins.")
    p2 = doc.add_paragraph("End of section 2.")
    pPr2 = p2._p.get_or_add_pPr()
    sect_pr2 = make_element("w:sectPr")
    sect_pr2.append(make_element("w:pgMar", {
        "w:top": "2880",
        "w:bottom": "2880",
        "w:left": "2880",
        "w:right": "2880",
        "w:header": "1440",
        "w:footer": "1440",
        "w:gutter": "720",
    }))
    pPr2.append(sect_pr2)

    # Final content in last section (margins set in body sectPr, which we don't test here)
    doc.add_paragraph("Content in section 3 (body-level section).")

    save_fixture("sections", "multi-section-margins", doc, {
        "name": "multi-section-margins",
        "spec_ref": "ISO 29500-1 §17.6.11",
        "description": "Two mid-document sections with different margins (narrow vs wide)",
        "expected_behavior": "Each section preserves its own margin values independently",
        "current_status": "GAP — pgMar not yet parsed into SectionProperties",
    })


def make_page_borders() -> None:
    """pgBorders — page borders.

    ISO 29500-1 §17.6.7: pgBorders in sectPr defines page-level borders.
    """
    doc = Document()

    doc.add_paragraph("Content in a section with page borders.")

    p = doc.add_paragraph("Last paragraph.")
    pPr = p._p.get_or_add_pPr()
    sect_pr = make_element("w:sectPr")

    pg_borders = make_element("w:pgBorders", {"w:offsetFrom": "page"})
    for edge in ["top", "left", "bottom", "right"]:
        border_el = make_element(f"w:{edge}", {
            "w:val": "single",
            "w:sz": "12",
            "w:space": "24",
            "w:color": "0000FF",
        })
        pg_borders.append(border_el)
    sect_pr.append(pg_borders)
    pPr.append(sect_pr)

    save_fixture("sections", "page-borders", doc, {
        "name": "page-borders",
        "spec_ref": "ISO 29500-1 §17.6.7",
        "description": "Section with solid blue page borders on all four edges",
        "expected_behavior": "Parser should expose pgBorders in SectionProperties",
        "current_status": "GAP — pgBorders not yet parsed into SectionProperties",
    })


def make_line_numbering() -> None:
    """lnNumType — line numbering.

    ISO 29500-1 §17.6.8: lnNumType in sectPr defines line numbering settings.
    """
    doc = Document()

    doc.add_paragraph("Content in a section with line numbering.")

    p = doc.add_paragraph("Last paragraph.")
    pPr = p._p.get_or_add_pPr()
    sect_pr = make_element("w:sectPr")
    ln_num = make_element("w:lnNumType", {
        "w:countBy": "5",
        "w:start": "1",
        "w:restart": "newPage",
    })
    sect_pr.append(ln_num)
    pPr.append(sect_pr)

    save_fixture("sections", "line-numbering", doc, {
        "name": "line-numbering",
        "spec_ref": "ISO 29500-1 §17.6.8",
        "description": "Section with line numbering: countBy=5, start=1, restart=newPage",
        "expected_behavior": "Parser should expose lnNumType attributes in SectionProperties",
        "current_status": "GAP — lnNumType not yet parsed into SectionProperties",
    })


def make_vertical_alignment() -> None:
    """vAlign — vertical text alignment on page.

    ISO 29500-1 §17.6.20: vAlign in sectPr defines vertical text alignment.
    """
    doc = Document()

    doc.add_paragraph("Content in a vertically centered section.")

    p = doc.add_paragraph("Last paragraph.")
    pPr = p._p.get_or_add_pPr()
    sect_pr = make_element("w:sectPr")
    v_align = make_element("w:vAlign", {"w:val": "center"})
    sect_pr.append(v_align)
    pPr.append(sect_pr)

    save_fixture("sections", "vertical-alignment", doc, {
        "name": "vertical-alignment",
        "spec_ref": "ISO 29500-1 §17.6.20",
        "description": "Section with vAlign=center (text centered vertically on page)",
        "expected_behavior": "Parser should expose vAlign in SectionProperties",
        "current_status": "GAP — vAlign not yet parsed into SectionProperties",
    })


def make_page_number_format() -> None:
    """pgNumType — page number format and starting value.

    ISO 29500-1 §17.6.12: pgNumType in sectPr defines page number formatting.
    """
    doc = Document()

    doc.add_paragraph("Content in a section with page number format.")

    p = doc.add_paragraph("Last paragraph.")
    pPr = p._p.get_or_add_pPr()
    sect_pr = make_element("w:sectPr")
    pg_num = make_element("w:pgNumType", {
        "w:fmt": "lowerRoman",
        "w:start": "3",
    })
    sect_pr.append(pg_num)
    pPr.append(sect_pr)

    save_fixture("sections", "page-number-format", doc, {
        "name": "page-number-format",
        "spec_ref": "ISO 29500-1 §17.6.12",
        "description": "Section with pgNumType: fmt=lowerRoman, start=3",
        "expected_behavior": "Parser should expose pgNumType attributes in SectionProperties",
        "current_status": "GAP — pgNumType not yet parsed into SectionProperties",
    })


# =========================================================================
# TABLE CONDITIONAL DEEP (ISO 29500-1 §17.7.6) — deep probing tests
# =========================================================================

def make_table_conditional_deep_fixtures() -> None:
    print("\n── Table Conditional Deep ──")
    make_tcd_tbllook_disables_firstrow()
    make_tcd_conditional_cell_margins()
    make_tcd_direct_shading_overrides_conditional()
    make_tcd_conditional_firstrow_borders()
    make_tcd_whole_table_conditional()
    make_tcd_conditional_merge_different_properties()
    make_tcd_band_column_shading()
    make_tcd_direct_para_alignment_overrides_conditional()
    make_tcd_direct_font_size_overrides_conditional()


def make_tcd_tbllook_disables_firstrow() -> None:
    """tblLook w:firstRow="0" should suppress firstRow conditional.

    ISO 29500-1 §17.7.6: Conditional formatting only applies when the
    corresponding tblLook flag is enabled. When w:firstRow="0", the
    firstRow tblStylePr should be ignored even though it exists in the style.
    """
    doc = Document()

    # firstRow conditional: blue shading
    first_row_spr = make_element("w:tblStylePr", {"w:type": "firstRow"})
    fr_tcPr = make_element("w:tcPr")
    fr_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "4472C4"}))
    first_row_spr.append(fr_tcPr)

    # band1Horz: light gray
    band1h_spr = make_element("w:tblStylePr", {"w:type": "band1Horz"})
    b1h_tcPr = make_element("w:tcPr")
    b1h_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "D9D9D9"}))
    band1h_spr.append(b1h_tcPr)

    tbl_borders = make_element("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "4", "w:color": "000000",
        }))

    _add_table_style(doc, "TblLookTest", "TblLook Test Style",
                     tblPr_children=[tbl_borders],
                     tblStylePrs=[first_row_spr, band1h_spr])

    doc.add_paragraph("tblLook disables firstRow conditional.")
    tbl = doc.add_table(rows=3, cols=2)
    # Key: w:firstRow="0" — disables firstRow conditional
    _apply_table_style(tbl, "TblLookTest", {
        "w:val": "0000",
        "w:firstRow": "0", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "0", "w:noVBand": "1",
    })
    _fill_table_cells(tbl, [
        ["R0C0", "R0C1"],
        ["R1C0", "R1C1"],
        ["R2C0", "R2C1"],
    ])

    save_fixture("table-conditional-deep", "tbllook-disables-firstrow", doc, {
        "name": "tbllook-disables-firstrow",
        "spec_ref": "ISO 29500-1 §17.7.6",
        "description": (
            "Table style has firstRow (blue 4472C4) and band1Horz (gray D9D9D9) "
            "conditionals. tblLook has w:firstRow='0' — disabling firstRow conditional."
        ),
        "expected_behavior": (
            "Row 0: should get band1Horz shading (D9D9D9), NOT firstRow blue. "
            "Row 1: band2Horz (no shading — band2Horz not defined). "
            "Row 2: band1Horz shading (D9D9D9)."
        ),
        "current_status": "testing tblLook flag suppression of firstRow",
    })


def make_tcd_conditional_cell_margins() -> None:
    """firstRow conditional with cell margins (tcMar).

    ISO 29500-1 §17.7.6: tblStylePr tcPr can contain tblCellMar for
    cell margins. When firstRow conditional matches, these margins
    should apply to cells in the first row.
    """
    doc = Document()

    # firstRow conditional: cell margins 150tw + blue shading
    first_row_spr = make_element("w:tblStylePr", {"w:type": "firstRow"})
    fr_tcPr = make_element("w:tcPr")
    fr_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "4472C4"}))
    fr_cellMar = make_element("w:tblCellMar")
    for side in ["top", "bottom", "left", "right"]:
        fr_cellMar.append(make_element(f"w:{side}", {
            "w:w": "150", "w:type": "dxa",
        }))
    fr_tcPr.append(fr_cellMar)
    first_row_spr.append(fr_tcPr)

    tbl_borders = make_element("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "4", "w:color": "000000",
        }))

    _add_table_style(doc, "CellMarginCondStyle", "Cell Margin Conditional Style",
                     tblPr_children=[tbl_borders],
                     tblStylePrs=[first_row_spr])

    doc.add_paragraph("Conditional cell margins from firstRow.")
    tbl = doc.add_table(rows=3, cols=2)
    _apply_table_style(tbl, "CellMarginCondStyle", {
        "w:val": "04A0",
        "w:firstRow": "1", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "1",
    })
    _fill_table_cells(tbl, [
        ["Header 1", "Header 2"],
        ["Data A", "Data B"],
        ["Data C", "Data D"],
    ])

    save_fixture("table-conditional-deep", "conditional-cell-margins", doc, {
        "name": "conditional-cell-margins",
        "spec_ref": "ISO 29500-1 §17.7.6",
        "description": (
            "Table style with firstRow conditional: shading (4472C4) and "
            "tblCellMar (150tw all sides) in tcPr."
        ),
        "expected_behavior": (
            "First row cells: shading 4472C4 AND margins 150tw on all sides. "
            "Data row cells: no shading, no conditional margins."
        ),
        "current_status": "testing conditional cell margin propagation",
    })


def make_tcd_direct_shading_overrides_conditional() -> None:
    """Direct cell shading should override conditional shading.

    ISO 29500-1 §17.7.6: Direct formatting > conditional > style default.
    When a cell has direct w:shd in tcPr AND a conditional sets shading,
    the direct formatting wins.
    """
    doc = Document()

    # firstRow conditional: blue shading
    first_row_spr = make_element("w:tblStylePr", {"w:type": "firstRow"})
    fr_tcPr = make_element("w:tcPr")
    fr_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "4472C4"}))
    first_row_spr.append(fr_tcPr)

    tbl_borders = make_element("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "4", "w:color": "000000",
        }))

    _add_table_style(doc, "DirectOverrideStyle", "Direct Override Style",
                     tblPr_children=[tbl_borders],
                     tblStylePrs=[first_row_spr])

    doc.add_paragraph("Direct cell shading overrides conditional.")
    tbl = doc.add_table(rows=2, cols=2)
    _apply_table_style(tbl, "DirectOverrideStyle", {
        "w:val": "04A0",
        "w:firstRow": "1", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "1",
    })
    _fill_table_cells(tbl, [
        ["Direct Green", "Conditional Blue"],
        ["No shading", "No shading"],
    ])

    # Cell (0,0): direct green shading that should override conditional blue
    cell_00 = tbl.cell(0, 0)
    tc_el = cell_00._tc
    tcPr = tc_el.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = make_element("w:tcPr")
        tc_el.insert(0, tcPr)
    tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "00B050"}))

    save_fixture("table-conditional-deep", "direct-shading-overrides-conditional", doc, {
        "name": "direct-shading-overrides-conditional",
        "spec_ref": "ISO 29500-1 §17.7.6",
        "description": (
            "Table style with firstRow conditional (blue 4472C4 shading). "
            "Cell (0,0) has direct green (00B050) shading in tcPr."
        ),
        "expected_behavior": (
            "Cell (0,0): green (00B050) from direct, NOT blue from conditional. "
            "Cell (0,1): blue (4472C4) from conditional (no direct shading). "
            "Cell (1,0), (1,1): no shading (not in first row)."
        ),
        "current_status": "testing direct > conditional precedence",
    })


def make_tcd_conditional_firstrow_borders() -> None:
    """firstRow conditional with cell borders.

    ISO 29500-1 §17.7.6: tblStylePr tcPr can contain tcBorders.
    When firstRow conditional matches, these borders should apply
    to cells in the first row.
    """
    doc = Document()

    # firstRow conditional: thick red borders + blue shading
    first_row_spr = make_element("w:tblStylePr", {"w:type": "firstRow"})
    fr_tcPr = make_element("w:tcPr")
    fr_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "4472C4"}))
    fr_borders = make_element("w:tcBorders")
    for side in ["top", "bottom", "left", "right"]:
        fr_borders.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "12", "w:color": "FF0000",
        }))
    fr_tcPr.append(fr_borders)
    first_row_spr.append(fr_tcPr)

    # Base table borders: thin black
    tbl_borders = make_element("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "4", "w:color": "000000",
        }))

    _add_table_style(doc, "FirstRowBordersStyle", "First Row Borders Style",
                     tblPr_children=[tbl_borders],
                     tblStylePrs=[first_row_spr])

    doc.add_paragraph("Conditional firstRow with cell borders.")
    tbl = doc.add_table(rows=3, cols=2)
    _apply_table_style(tbl, "FirstRowBordersStyle", {
        "w:val": "04A0",
        "w:firstRow": "1", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "1",
    })
    _fill_table_cells(tbl, [
        ["Header 1", "Header 2"],
        ["Data A", "Data B"],
        ["Data C", "Data D"],
    ])

    save_fixture("table-conditional-deep", "conditional-firstrow-borders", doc, {
        "name": "conditional-firstrow-borders",
        "spec_ref": "ISO 29500-1 §17.7.6",
        "description": (
            "Table style with firstRow conditional: blue shading (4472C4) and "
            "thick red borders (sz=12, FF0000) in tcPr. Base table has thin black borders."
        ),
        "expected_behavior": (
            "First row cells: blue shading, thick red borders (sz=12, FF0000). "
            "Data row cells: no cell-level borders (only table-level thin black)."
        ),
        "current_status": "testing conditional cell border propagation",
    })


def make_tcd_whole_table_conditional() -> None:
    """wholeTable conditional type (tblStylePr type='wholeTable').

    ISO 29500-1 §17.7.6: The wholeTable conditional applies to all cells
    as a baseline before more specific conditionals are layered on.
    """
    doc = Document()

    # wholeTable conditional: light gray shading
    whole_spr = make_element("w:tblStylePr", {"w:type": "wholeTable"})
    wt_tcPr = make_element("w:tcPr")
    wt_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "F2F2F2"}))
    whole_spr.append(wt_tcPr)

    # firstRow conditional: blue shading (should override wholeTable on row 0)
    first_row_spr = make_element("w:tblStylePr", {"w:type": "firstRow"})
    fr_tcPr = make_element("w:tcPr")
    fr_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "4472C4"}))
    first_row_spr.append(fr_tcPr)

    tbl_borders = make_element("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "4", "w:color": "000000",
        }))

    _add_table_style(doc, "WholeTableStyle", "Whole Table Style",
                     tblPr_children=[tbl_borders],
                     tblStylePrs=[whole_spr, first_row_spr])

    doc.add_paragraph("wholeTable conditional with firstRow override.")
    tbl = doc.add_table(rows=3, cols=2)
    _apply_table_style(tbl, "WholeTableStyle", {
        "w:val": "04A0",
        "w:firstRow": "1", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "1",
    })
    _fill_table_cells(tbl, [
        ["Header 1", "Header 2"],
        ["Data A", "Data B"],
        ["Data C", "Data D"],
    ])

    save_fixture("table-conditional-deep", "whole-table-conditional", doc, {
        "name": "whole-table-conditional",
        "spec_ref": "ISO 29500-1 §17.7.6",
        "description": (
            "Table style with wholeTable conditional (gray F2F2F2 shading) and "
            "firstRow conditional (blue 4472C4 shading). firstRow enabled in tblLook."
        ),
        "expected_behavior": (
            "Row 0: blue (4472C4) from firstRow (overrides wholeTable). "
            "Rows 1-2: gray (F2F2F2) from wholeTable. "
            "wholeTable is the lowest-precedence conditional, applied to all cells."
        ),
        "current_status": "GAP — TblStylePrType enum does not include WholeTable",
    })


def make_tcd_conditional_merge_different_properties() -> None:
    """Multiple conditionals setting different property types should merge.

    ISO 29500-1 §17.7.6: When firstRow sets borders and firstCol sets shading,
    cell (0,0) should get BOTH — borders from firstRow and shading from firstCol.
    Non-overlapping properties merge across conditional layers.
    """
    doc = Document()

    # firstRow conditional: thick red borders only (no shading)
    first_row_spr = make_element("w:tblStylePr", {"w:type": "firstRow"})
    fr_tcPr = make_element("w:tcPr")
    fr_borders = make_element("w:tcBorders")
    for side in ["top", "bottom", "left", "right"]:
        fr_borders.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "12", "w:color": "FF0000",
        }))
    fr_tcPr.append(fr_borders)
    first_row_spr.append(fr_tcPr)

    # firstCol conditional: green shading only (no borders)
    first_col_spr = make_element("w:tblStylePr", {"w:type": "firstCol"})
    fc_tcPr = make_element("w:tcPr")
    fc_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "70AD47"}))
    first_col_spr.append(fc_tcPr)

    tbl_borders = make_element("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "4", "w:color": "000000",
        }))

    _add_table_style(doc, "MergePropsStyle", "Merge Props Style",
                     tblPr_children=[tbl_borders],
                     tblStylePrs=[first_row_spr, first_col_spr])

    doc.add_paragraph("Conditional merge: firstRow borders + firstCol shading.")
    tbl = doc.add_table(rows=3, cols=3)
    _apply_table_style(tbl, "MergePropsStyle", {
        "w:val": "04A0",
        "w:firstRow": "1", "w:lastRow": "0",
        "w:firstColumn": "1", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "1",
    })
    _fill_table_cells(tbl, [
        ["R0C0", "R0C1", "R0C2"],
        ["R1C0", "R1C1", "R1C2"],
        ["R2C0", "R2C1", "R2C2"],
    ])

    save_fixture("table-conditional-deep", "conditional-merge-properties", doc, {
        "name": "conditional-merge-properties",
        "spec_ref": "ISO 29500-1 §17.7.6",
        "description": (
            "firstRow conditional sets thick red borders (no shading). "
            "firstCol conditional sets green shading (no borders). "
            "Cell (0,0) matches both."
        ),
        "expected_behavior": (
            "Cell (0,0): thick red borders FROM firstRow + green shading FROM firstCol. "
            "Cell (0,1): thick red borders from firstRow, no shading. "
            "Cell (1,0): green shading from firstCol, no cell borders. "
            "Cell (1,1): no conditional formatting (interior cell)."
        ),
        "current_status": "testing non-overlapping property merge across conditionals",
    })


def make_tcd_band_column_shading() -> None:
    """Band column (vertical) shading via band1Vert/band2Vert.

    ISO 29500-1 §17.7.6: band1Vert applies to odd-indexed columns (0, 2, ...),
    band2Vert applies to even-indexed columns (1, 3, ...).
    noVBand must be "0" for vertical banding to be active.
    """
    doc = Document()

    # band1Vert: light green
    band1v_spr = make_element("w:tblStylePr", {"w:type": "band1Vert"})
    b1v_tcPr = make_element("w:tcPr")
    b1v_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "C6EFCE"}))
    band1v_spr.append(b1v_tcPr)

    # band2Vert: light orange
    band2v_spr = make_element("w:tblStylePr", {"w:type": "band2Vert"})
    b2v_tcPr = make_element("w:tcPr")
    b2v_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "FBE5D6"}))
    band2v_spr.append(b2v_tcPr)

    tbl_borders = make_element("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "4", "w:color": "000000",
        }))

    _add_table_style(doc, "BandColStyle", "Band Column Style",
                     tblPr_children=[tbl_borders],
                     tblStylePrs=[band1v_spr, band2v_spr])

    doc.add_paragraph("Band column (vertical) shading test.")
    tbl = doc.add_table(rows=3, cols=4)
    # noVBand="0" to enable vertical banding; noHBand="1" to disable horizontal
    _apply_table_style(tbl, "BandColStyle", {
        "w:val": "0000",
        "w:firstRow": "0", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "0",
    })
    labels = [[f"R{r}C{c}" for c in range(4)] for r in range(3)]
    _fill_table_cells(tbl, labels)

    save_fixture("table-conditional-deep", "band-column-shading", doc, {
        "name": "band-column-shading",
        "spec_ref": "ISO 29500-1 §17.7.6",
        "description": (
            "Table with band1Vert (green C6EFCE) and band2Vert (orange FBE5D6). "
            "4 columns, vertical banding enabled (noVBand=0), no row/col conditionals."
        ),
        "expected_behavior": (
            "Col 0: green (C6EFCE, band1Vert). Col 1: orange (FBE5D6, band2Vert). "
            "Col 2: green. Col 3: orange. Same pattern across all rows."
        ),
        "current_status": "testing vertical banding shading",
    })


# =========================================================================
# INDENT INTERACTION AUDIT (GAP-101, GAP-102, GAP-103)
# =========================================================================

def make_indent_interaction_audit_fixtures() -> None:
    print("\n── Indent Interaction Audit ──")
    make_iia_gap101_char_unit_prefix_stripping()
    make_iia_gap102_tab_absorption_positive_firstline()
    make_iia_gap103_numbering_blocks_style_firstline()


def make_iia_gap101_char_unit_prefix_stripping() -> None:
    """GAP-101: Character-unit indent survives prefix stripping.

    When a paragraph has w:ind with character-unit values (w:startChars, w:endChars)
    and the runtime strips a literal prefix like "(a)\\t", the character-unit values
    must be propagated to the output Indentation — not hardcoded to None.

    Fixture:
      P0: w:left=720, w:startChars="200", w:endChars="100",
          tab at 1440, text "(a)\\tBody text..." — triggers prefix stripping.
      P1: w:left=720, w:startChars="200", w:endChars="100",
          no prefix — control (passthrough).
    """
    doc = Document()

    # P0: prefix stripped — char-unit values must survive
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pPr.append(make_element("w:ind", {
        "w:left": "720",
        "w:startChars": "200",
        "w:endChars": "100",
    }))
    tabs_el = make_element("w:tabs")
    tabs_el.append(make_element("w:tab", {"w:val": "left", "w:pos": "1440"}))
    pPr.append(tabs_el)
    p.add_run("(a)\tBody text with char-unit indent and prefix.")

    # P1: control — no prefix, char-unit values pass through
    p2 = doc.add_paragraph()
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:ind", {
        "w:left": "720",
        "w:startChars": "200",
        "w:endChars": "100",
    }))
    p2.add_run("Body text with char-unit indent, no prefix.")

    save_fixture("indent-interaction-audit", "char-unit-prefix-stripping", doc, {
        "name": "char-unit-prefix-stripping",
        "spec_ref": "ISO 29500-1 §17.3.1.12, MS-OI29500 2.1.44",
        "description": (
            "P0: w:left=720, w:startChars=200, w:endChars=100, tab at 1440, "
            "text '(a)\\tBody text...'. P1: same indent, no prefix (control)."
        ),
        "expected_behavior": (
            "P0 after prefix stripping: start_chars=200, end_chars=100 must survive. "
            "GAP-101 bug: prefix stripping branch hardcodes start_chars=None, end_chars=None."
        ),
        "current_status": "GAP-101",
    })


def make_iia_gap102_tab_absorption_positive_firstline() -> None:
    """GAP-102: Tab absorption must NOT fire for positive firstLine.

    When a paragraph has positive w:firstLine (not hanging) and body text
    contains tabs, the tab absorption branch should not fire. Tab absorption
    is only correct for hanging indent (negative firstLine).

    Fixture:
      P0: w:left=720, w:firstLine=360, text "A\\tB" — has tabs + positive firstLine.
      P1: w:left=720, w:firstLine=360, text "AB" — no tabs (control).
    """
    doc = Document()

    # P0: positive firstLine + tabs — tab absorption must NOT fire
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pPr.append(make_element("w:ind", {"w:left": "720", "w:firstLine": "360"}))
    tabs_el = make_element("w:tabs")
    tabs_el.append(make_element("w:tab", {"w:val": "left", "w:pos": "1440"}))
    pPr.append(tabs_el)
    p.add_run("A\tB")

    # P1: control — same indent, no tabs
    p2 = doc.add_paragraph()
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:ind", {"w:left": "720", "w:firstLine": "360"}))
    p2.add_run("AB")

    save_fixture("indent-interaction-audit", "tab-absorption-positive-firstline", doc, {
        "name": "tab-absorption-positive-firstline",
        "spec_ref": "ISO 29500-1 §17.3.1.12",
        "description": (
            "P0: w:left=720, w:firstLine=360, text 'A\\tB'. "
            "P1: same indent, text 'AB' (control, no tabs)."
        ),
        "expected_behavior": (
            "P0: left=720, first_line=360 (same as P1 — tab absorption must NOT fire "
            "for positive firstLine). GAP-102 bug: condition first_line_twips!=0 "
            "incorrectly fires for positive firstLine, absorbing it into left."
        ),
        "current_status": "GAP-102",
    })


def make_iia_gap103_numbering_blocks_style_firstline() -> None:
    """GAP-103: Numbering indent replaces style indent as a whole element.

    When numbering provides w:ind (e.g., left=720 but no firstLine), the
    numbering w:ind should replace the style's w:ind entirely. The style's
    firstLine must NOT leak through via per-field merge.

    Fixture:
      Style "FirstLineStyle": w:ind w:left=360, w:firstLine=240
      Numbering level 0: w:ind w:left=720 (no firstLine, no hanging)
      P0: Style + numbering — numbering w:ind replaces style w:ind entirely.
      P1: Style only (control) — should have left=360, firstLine=240.
    """
    doc = Document()
    styles_element = doc.styles.element

    # Create a paragraph style with firstLine indent
    style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "FirstLineStyle"})
    style.append(make_element("w:name", {"w:val": "First Line Style"}))
    ppr_s = make_element("w:pPr")
    ppr_s.append(make_element("w:ind", {"w:left": "360", "w:firstLine": "240"}))
    style.append(ppr_s)
    styles_element.append(style)

    # Numbering with left=720, NO firstLine/hanging
    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:pPr>
        <w:ind w:left="720"/>
      </w:pPr>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # P0: style + numbering — numbering w:ind should replace style w:ind entirely
    p = doc.add_paragraph("Numbered paragraph with style firstLine")
    pPr = p._p.get_or_add_pPr()
    pPr.append(make_element("w:pStyle", {"w:val": "FirstLineStyle"}))
    numPr = make_element("w:numPr")
    numPr.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr.append(make_element("w:numId", {"w:val": "1"}))
    pPr.append(numPr)

    # P1: style only (control)
    p2 = doc.add_paragraph("Non-numbered paragraph with style firstLine")
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:pStyle", {"w:val": "FirstLineStyle"}))

    save_fixture("indent-interaction-audit", "numbering-blocks-style-firstline", doc, {
        "name": "numbering-blocks-style-firstline",
        "spec_ref": "ISO 29500-1 §17.9.22, §17.3.1.12",
        "description": (
            "Style 'FirstLineStyle': w:ind left=360, firstLine=240. "
            "Numbering level 0: w:ind left=720 (no firstLine). "
            "P0: style + numbering. P1: style only (control)."
        ),
        "expected_behavior": (
            "P0: left=720, firstLine=None (numbering w:ind replaces style w:ind entirely). "
            "P1: left=360, firstLine=240. "
            "GAP-103 bug: per-field merge leaks style firstLine=240 into P0."
        ),
        "current_status": "GAP-103",
    })


def make_tcd_direct_para_alignment_overrides_conditional() -> None:
    """Direct paragraph alignment should override conditional alignment.

    ISO 29500-1 §17.7.6: Direct formatting > conditional > style default.
    When a paragraph has direct w:jc AND a conditional sets jc, direct wins.
    GAP-110.
    """
    doc = Document()

    # firstRow conditional with center alignment in pPr
    first_row_spr = make_element("w:tblStylePr", {"w:type": "firstRow"})
    fr_pPr = make_element("w:pPr")
    fr_pPr.append(make_element("w:jc", {"w:val": "center"}))
    first_row_spr.append(fr_pPr)
    fr_tcPr = make_element("w:tcPr")
    fr_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "DDDDDD"}))
    first_row_spr.append(fr_tcPr)

    tbl_borders = make_element("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "4", "w:color": "000000",
        }))

    _add_table_style(doc, "DirectAlignOverride", "Direct Align Override",
                     tblPr_children=[tbl_borders],
                     tblStylePrs=[first_row_spr])

    doc.add_paragraph("Direct paragraph alignment overrides conditional.")
    tbl = doc.add_table(rows=2, cols=2)
    _apply_table_style(tbl, "DirectAlignOverride", {
        "w:val": "04A0",
        "w:firstRow": "1", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "1",
    })
    _fill_table_cells(tbl, [
        ["Direct Right", "Conditional Center"],
        ["No alignment", "No alignment"],
    ])

    # Cell (0,0): direct right alignment that should override conditional center
    cell_00 = tbl.cell(0, 0)
    tc_el = cell_00._tc
    p_el = tc_el.findall(qn("w:p"))[0]
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = make_element("w:pPr")
        p_el.insert(0, pPr)
    pPr.append(make_element("w:jc", {"w:val": "right"}))

    save_fixture("table-conditional-deep", "direct-para-alignment-overrides-conditional", doc, {
        "name": "direct-para-alignment-overrides-conditional",
        "spec_ref": "ISO 29500-1 §17.7.6",
        "description": (
            "Table style with firstRow conditional (center alignment in pPr). "
            "Cell (0,0) paragraph has direct jc=right."
        ),
        "expected_behavior": (
            "Cell (0,0): right alignment (direct overrides conditional center). "
            "Cell (0,1): center alignment (from conditional, no direct). "
            "Data rows: default alignment."
        ),
        "current_status": "GAP-110 — testing direct pPr > conditional pPr precedence",
    })


def make_tcd_direct_font_size_overrides_conditional() -> None:
    """Direct run font size should override conditional font size.

    ISO 29500-1 §17.7.6: Direct formatting > conditional > style default.
    When a run has direct w:sz AND a conditional sets sz, direct wins.
    GAP-111.
    """
    doc = Document()

    # firstRow conditional with 28pt (56 half-points) font size in rPr
    first_row_spr = make_element("w:tblStylePr", {"w:type": "firstRow"})
    fr_rPr = make_element("w:rPr")
    fr_rPr.append(make_element("w:sz", {"w:val": "56"}))
    first_row_spr.append(fr_rPr)
    fr_tcPr = make_element("w:tcPr")
    fr_tcPr.append(make_element("w:shd", {"w:val": "clear", "w:fill": "DDDDDD"}))
    first_row_spr.append(fr_tcPr)

    tbl_borders = make_element("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{side}", {
            "w:val": "single", "w:sz": "4", "w:color": "000000",
        }))

    _add_table_style(doc, "DirectSzOverride", "Direct Sz Override",
                     tblPr_children=[tbl_borders],
                     tblStylePrs=[first_row_spr])

    doc.add_paragraph("Direct run font size overrides conditional.")
    tbl = doc.add_table(rows=2, cols=2)
    _apply_table_style(tbl, "DirectSzOverride", {
        "w:val": "04A0",
        "w:firstRow": "1", "w:lastRow": "0",
        "w:firstColumn": "0", "w:lastColumn": "0",
        "w:noHBand": "1", "w:noVBand": "1",
    })
    _fill_table_cells(tbl, [
        ["Direct 16pt", "Conditional 28pt"],
        ["No size", "No size"],
    ])

    # Cell (0,0): direct 16pt (32 half-points) font size on the run
    cell_00 = tbl.cell(0, 0)
    tc_el = cell_00._tc
    p_el = tc_el.findall(qn("w:p"))[0]
    r_el = p_el.findall(qn("w:r"))[0]
    rPr = r_el.find(qn("w:rPr"))
    if rPr is None:
        rPr = make_element("w:rPr")
        r_el.insert(0, rPr)
    rPr.append(make_element("w:sz", {"w:val": "32"}))

    save_fixture("table-conditional-deep", "direct-font-size-overrides-conditional", doc, {
        "name": "direct-font-size-overrides-conditional",
        "spec_ref": "ISO 29500-1 §17.7.6",
        "description": (
            "Table style with firstRow conditional (28pt/56hp font size in rPr). "
            "Cell (0,0) run has direct sz=32 (16pt)."
        ),
        "expected_behavior": (
            "Cell (0,0): 16pt/32hp font size (direct overrides conditional 28pt). "
            "Cell (0,1): 28pt/56hp font size (from conditional, no direct). "
            "Data rows: default font size."
        ),
        "current_status": "GAP-111 — testing direct rPr > conditional rPr precedence",
    })


def make_numbering_spec_fixtures() -> None:
    """Numbering behavior tests per ECMA-376 §17.9."""
    print("\n── Numbering Spec ──")
    make_ns_suff_nothing()
    make_ns_suff_space()
    make_ns_pstyle_reverse_binding()
    make_ns_numstylelink_chain()
    make_ns_start_omitted_default_zero()
    make_ns_numid_counter_resume()
    make_ns_start_override_vs_absent()


def make_ns_suff_nothing() -> None:
    """suff="nothing" — no separator between number text and paragraph text.

    ISO 29500-1 §17.9.28: When suff val="nothing", there shall be no content
    (no tab, no space) between the numbering text and the paragraph body.
    Default (omitted) is "tab".
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}"
             xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:suff w:val="nothing"/>
    make_table_borders_spec_fixtures()
    </w:lvl>
  </w:abstractNum>
  <w:abstractNum w:abstractNumId="1">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
  <w:num w:numId="2">
    <w:abstractNumId w:val="1"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # Two paragraphs with suff="nothing" (numId=1)
    for text in ["No separator after number", "Second item no separator"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "0"}))
        numPr.append(make_element("w:numId", {"w:val": "1"}))
        pPr.append(numPr)

    # Two paragraphs with default suff (tab, numId=2) as control
    for text in ["Tab separator (default)", "Second tab separator"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "0"}))
        numPr.append(make_element("w:numId", {"w:val": "2"}))
        pPr.append(numPr)

    save_fixture("numbering-spec", "suff-nothing", doc, {
        "name": "suff-nothing",
        "spec_ref": "ISO 29500-1 §17.9.28",
        "description": "Numbering level with suff='nothing' — no separator between number and text",
        "expected_behavior": (
            "numId=1 (suff=nothing): rendered_text = '1.No separator after number' (no tab/space). "
            "numId=2 (default tab): rendered_text = '1.\\tTab separator (default)'."
        ),
        "current_status": "TESTING",
    })


def make_ns_suff_space() -> None:
    """suff="space" — space separator between number text and paragraph text.

    ISO 29500-1 §17.9.28: When suff val="space", a single space character
    shall be placed between the numbering text and the paragraph body.
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}"
             xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:suff w:val="space"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    for text in ["Space after number", "Second space item"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "0"}))
        numPr.append(make_element("w:numId", {"w:val": "1"}))
        pPr.append(numPr)

    save_fixture("numbering-spec", "suff-space", doc, {
        "name": "suff-space",
        "spec_ref": "ISO 29500-1 §17.9.28",
        "description": "Numbering level with suff='space' — space separator between number and text",
        "expected_behavior": "rendered_text = '1. Space after number' (single space, not tab).",
        "current_status": "TESTING",
    })


def make_ns_pstyle_reverse_binding() -> None:
    """pStyle reverse binding — paragraph style triggers numbering from abstractNum.

    ISO 29500-1 §17.9.23: When a lvl element contains <w:pStyle w:val="X"/>,
    any paragraph with style "X" gets numbering from that level, even if the
    paragraph itself has no numPr. The style may or may not have a numPr element.

    This fixture tests three cases:
      P0: paragraph with style "ListNum" and NO explicit numPr — should get
          numbering via the pStyle reverse binding.
      P1: same style, second paragraph — counter should increment.
      P2: paragraph with explicit numPr pointing to numId=1 ilvl=0 — control,
          should work normally.
    """
    doc = Document()
    styles_element = doc.styles.element

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}"
             xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:pStyle w:val="ListNum"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # Create the "ListNum" paragraph style — note: NO numPr in the style itself.
    # The numbering comes purely from the abstractNum's pStyle binding.
    list_num_style = make_element("w:style", {"w:type": "paragraph", "w:styleId": "ListNum"})
    list_num_style.append(make_element("w:name", {"w:val": "List Num"}))
    ppr_ln = make_element("w:pPr")
    ppr_ln.append(make_element("w:ind", {"w:left": "720", "w:hanging": "360"}))
    list_num_style.append(ppr_ln)
    styles_element.append(list_num_style)

    # P0: Style-only numbering (via reverse binding)
    p0 = doc.add_paragraph("First item via pStyle binding")
    pPr0 = p0._p.get_or_add_pPr()
    pPr0.append(make_element("w:pStyle", {"w:val": "ListNum"}))

    # P1: Same style, should increment
    p1 = doc.add_paragraph("Second item via pStyle binding")
    pPr1 = p1._p.get_or_add_pPr()
    pPr1.append(make_element("w:pStyle", {"w:val": "ListNum"}))

    # P2: Explicit numPr (control)
    p2 = doc.add_paragraph("Third item via explicit numPr")
    pPr2 = p2._p.get_or_add_pPr()
    numPr = make_element("w:numPr")
    numPr.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr.append(make_element("w:numId", {"w:val": "1"}))
    pPr2.append(numPr)

    save_fixture("numbering-spec", "pstyle-reverse-binding", doc, {
        "name": "pstyle-reverse-binding",
        "spec_ref": "ISO 29500-1 §17.9.23",
        "description": (
            "abstractNum level 0 has pStyle='ListNum'. Paragraphs with style "
            "'ListNum' (and no explicit numPr) should get numbering via the "
            "reverse binding."
        ),
        "expected_behavior": (
            "P0: synthesized '1.' via pStyle binding. "
            "P1: synthesized '2.' via pStyle binding. "
            "P2: synthesized '3.' via explicit numPr."
        ),
        "current_status": "TESTING",
    })


def make_ns_numstylelink_chain() -> None:
    """numStyleLink chain — abstractNum references a numbering style that
    points to the actual definition.

    ISO 29500-1 §17.9.21 + §17.9.27: An abstractNum with numStyleLink
    references a numbering style. That style's numPr points to a num instance,
    which references the abstractNum that has styleLink declaring it as the
    source for that style.

    Chain: abstractNum 0 --numStyleLink--> style "NumStyle"
           style "NumStyle" --numPr--> numId 2
           numId 2 --> abstractNum 1
           abstractNum 1 --styleLink--> "NumStyle"

    When a paragraph uses numId=1 (which references abstractNum 0), the
    implementation must follow the chain to find levels from abstractNum 1.
    """
    doc = Document()
    styles_element = doc.styles.element

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}"
             xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:multiLevelType w:val="multilevel"/>
    <w:numStyleLink w:val="NumStyle"/>
  </w:abstractNum>
  <w:abstractNum w:abstractNumId="1">
    <w:multiLevelType w:val="multilevel"/>
    <w:styleLink w:val="NumStyle"/>
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="upperRoman"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="1">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerLetter"/>
      <w:lvlText w:val="%2)"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
  <w:num w:numId="2">
    <w:abstractNumId w:val="1"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # Create the numbering style (required for the chain, though our code
    # currently only uses the abstractNum-level numStyleLink/styleLink matching)
    num_style = make_element("w:style", {"w:type": "numbering", "w:styleId": "NumStyle"})
    num_style.append(make_element("w:name", {"w:val": "NumStyle"}))
    ppr_ns = make_element("w:pPr")
    numPr_ns = make_element("w:numPr")
    numPr_ns.append(make_element("w:numId", {"w:val": "2"}))
    ppr_ns.append(numPr_ns)
    num_style.append(ppr_ns)
    styles_element.append(num_style)

    # Paragraphs using numId=1 (which has numStyleLink → must resolve via chain)
    p0 = doc.add_paragraph("Level 0 via numStyleLink chain")
    pPr0 = p0._p.get_or_add_pPr()
    numPr0 = make_element("w:numPr")
    numPr0.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr0.append(make_element("w:numId", {"w:val": "1"}))
    pPr0.append(numPr0)

    p1 = doc.add_paragraph("Level 1 via numStyleLink chain")
    pPr1 = p1._p.get_or_add_pPr()
    numPr1 = make_element("w:numPr")
    numPr1.append(make_element("w:ilvl", {"w:val": "1"}))
    numPr1.append(make_element("w:numId", {"w:val": "1"}))
    pPr1.append(numPr1)

    p2 = doc.add_paragraph("Level 0 second item via numStyleLink chain")
    pPr2 = p2._p.get_or_add_pPr()
    numPr2 = make_element("w:numPr")
    numPr2.append(make_element("w:ilvl", {"w:val": "0"}))
    numPr2.append(make_element("w:numId", {"w:val": "1"}))
    pPr2.append(numPr2)

    save_fixture("numbering-spec", "numstylelink-chain", doc, {
        "name": "numstylelink-chain",
        "spec_ref": "ISO 29500-1 §17.9.21 + §17.9.27",
        "description": (
            "abstractNum 0 has numStyleLink='NumStyle'. abstractNum 1 has "
            "styleLink='NumStyle' and actual levels. Paragraphs reference numId=1 "
            "(abstractNum 0) — levels must be resolved via the numStyleLink chain."
        ),
        "expected_behavior": (
            "P0: 'I.' (upperRoman). P1: 'a)' (lowerLetter). P2: 'II.' (upperRoman, incremented)."
        ),
        "current_status": "TESTING",
    })


def make_ns_start_omitted_default_zero() -> None:
    """start element omitted — default is 0 per spec.

    ISO 29500-1 §17.9.25: "If this element is omitted, then the starting
    value shall be zero (0)."

    This is counter-intuitive — most users expect 1. When numFmt is decimal
    and start is omitted, the first item should display "0." not "1.".
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}"
             xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:abstractNum w:abstractNumId="1">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
  <w:num w:numId="2">
    <w:abstractNumId w:val="1"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # numId=1: start omitted (should default to 0)
    for text in ["Start-omitted first", "Start-omitted second", "Start-omitted third"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "0"}))
        numPr.append(make_element("w:numId", {"w:val": "1"}))
        pPr.append(numPr)

    # numId=2: start=1 (control, should be 1., 2., 3.)
    for text in ["Start-one first", "Start-one second", "Start-one third"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "0"}))
        numPr.append(make_element("w:numId", {"w:val": "2"}))
        pPr.append(numPr)

    save_fixture("numbering-spec", "start-omitted-default-zero", doc, {
        "name": "start-omitted-default-zero",
        "spec_ref": "ISO 29500-1 §17.9.25",
        "description": (
            "Level with no w:start element. Per spec, default is 0. "
            "numId=1 (start omitted): should be 0., 1., 2. "
            "numId=2 (start=1, control): should be 1., 2., 3."
        ),
        "expected_behavior": (
            "numId=1: '0.', '1.', '2.' (start defaults to 0). "
            "numId=2: '1.', '2.', '3.' (explicit start=1)."
        ),
        "current_status": "TESTING",
    })


def make_ns_numid_counter_resume() -> None:
    """Switching back to a previous numId should RESUME its counter, not restart.

    Per ISO 29500-1 §17.9.15: Each numbering definition instance (w:num) is
    independent. When a document interleaves two numId references, the counter
    for each should be maintained independently.

    This tests: numId=1 (items 1,2), then numId=2 (items 1,2), then back to
    numId=1 which should continue at 3 (not restart at 1).
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}"
             xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:abstractNum w:abstractNumId="1">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerLetter"/>
      <w:lvlText w:val="(%1)"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
  <w:num w:numId="2">
    <w:abstractNumId w:val="1"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # First batch: numId=1 -> 1., 2.
    for text in ["Decimal list first", "Decimal list second"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "0"}))
        numPr.append(make_element("w:numId", {"w:val": "1"}))
        pPr.append(numPr)

    # Interleave: numId=2 -> (a), (b)
    for text in ["Letter list first", "Letter list second"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "0"}))
        numPr.append(make_element("w:numId", {"w:val": "2"}))
        pPr.append(numPr)

    # Resume: numId=1 -> should be 3., 4. (NOT restart at 1.)
    for text in ["Decimal list third (should be 3)", "Decimal list fourth (should be 4)"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "0"}))
        numPr.append(make_element("w:numId", {"w:val": "1"}))
        pPr.append(numPr)

    save_fixture("numbering-spec", "numid-counter-resume", doc, {
        "name": "numid-counter-resume",
        "spec_ref": "ISO 29500-1 §17.9.15",
        "description": (
            "Two numId interleaved: numId=1 (1.,2.), then numId=2 ((a),(b)), "
            "then back to numId=1 which should RESUME at 3., not restart at 1."
        ),
        "expected_behavior": (
            "numId=1: '1.', '2.', then after interleave '3.', '4.'. "
            "numId=2: '(a)', '(b)'."
        ),
        "current_status": "TESTING",
    })


def make_ns_start_override_vs_absent() -> None:
    """startOverride=1 vs absent startOverride — verifying restart semantics.

    ISO 29500-1 §17.9.26: startOverride resets the counter on first encounter
    of a particular (numId, ilvl). When absent, the counter inherits from the
    abstract definition's start value and continues normally.

    This fixture uses two num instances sharing the same abstractNum:
      numId=1: no overrides — used for first batch
      numId=2: startOverride=1 on ilvl=0 — used for second batch
      numId=3: no overrides — used for third batch (should continue from numId=1's last value)

    The key question: does startOverride=1 on numId=2 force a restart to 1,
    and does numId=3 (with no override) correctly continue from where numId=1
    left off (if they share the abstract definition)?
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}"
             xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
  <w:num w:numId="2">
    <w:abstractNumId w:val="0"/>
    <w:lvlOverride w:ilvl="0">
      <w:startOverride w:val="1"/>
    </w:lvlOverride>
  </w:num>
  <w:num w:numId="3">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # First batch: numId=1, should be 1., 2., 3.
    for text in ["First batch item one", "First batch item two", "First batch item three"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "0"}))
        numPr.append(make_element("w:numId", {"w:val": "1"}))
        pPr.append(numPr)

    # Second batch: numId=2 with startOverride=1, should restart: 1., 2.
    for text in ["Second batch restarted one", "Second batch restarted two"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "0"}))
        numPr.append(make_element("w:numId", {"w:val": "2"}))
        pPr.append(numPr)

    # Third batch: numId=3, no override — per spec, each numId has its own counter
    # so this should start fresh at 1., 2. (separate num instance)
    for text in ["Third batch item one", "Third batch item two"]:
        p = doc.add_paragraph(text)
        pPr = p._p.get_or_add_pPr()
        numPr = make_element("w:numPr")
        numPr.append(make_element("w:ilvl", {"w:val": "0"}))
        numPr.append(make_element("w:numId", {"w:val": "3"}))
        pPr.append(numPr)

    save_fixture("numbering-spec", "start-override-vs-absent", doc, {
        "name": "start-override-vs-absent",
        "spec_ref": "ISO 29500-1 §17.9.26",
        "description": (
            "Three num instances sharing an abstractNum. numId=1 (no override), "
            "numId=2 (startOverride=1), numId=3 (no override). Tests that "
            "startOverride correctly resets, and separate num instances have "
            "independent counters."
        ),
        "expected_behavior": (
            "numId=1: 1., 2., 3. "
            "numId=2: 1., 2. (restarted by startOverride). "
            "numId=3: 1., 2. (fresh counter for different numId)."
        ),
        "current_status": "TESTING",
    })


def main() -> None:
    print("Generating spec-compliance test fixtures:")
    make_tracked_changes_fixtures()
    make_numbering_fixtures()
    make_formatting_fixtures()
    make_table_fixtures()
    make_run_formatting_fixtures()
    make_tab_fixtures()
    make_structure_fixtures()
    make_comment_fixtures()
    make_bookmark_fixtures()
    make_hyperlink_fixtures()
    make_field_fixtures()
    make_fields_deep_fixtures()
    make_styles_fixtures()
    make_content_fixtures()
    make_additional_numbering_fixtures()
    make_edge_case_spacing_fixtures()
    make_edge_case_parsing_fixtures()
    make_edge_case_table_fixtures()
    make_run_resolution_fixtures()
    make_numbering_indent_fixtures()
    make_style_cascade_fixtures()
    make_table_cascade_fixtures()
    make_indent_edge_case_fixtures()
    make_tab_indent_interaction_fixtures()
    make_table_style_resolution_fixtures()
    make_section_fixtures()
    make_table_conditional_deep_fixtures()
    make_indent_interaction_audit_fixtures()
    make_indent_firstline_cascade_fixtures()
    make_redline_compat_fixtures()
    make_numbering_spec_fixtures()
    print("\nDone.")


def make_indent_firstline_cascade_fixtures() -> None:
    """Isolate the firstLine cascade question: does Word inherit firstLine
    from the Normal style when a paragraph has direct w:ind with left only?

    ISO 29500-1 §17.3.1.12 says per-attribute cascade, but Word may behave
    differently. This fixture creates a minimal DOCX to test the exact scenario.
    """
    print("\n── Indent firstLine Cascade ──")
    make_ifc_direct_left_only()


def make_ifc_direct_left_only() -> None:
    """Minimal reproduction: Normal has firstLine=720, paragraph has direct
    w:ind w:left="-720" but omits firstLine/hanging.

    Question: does the paragraph inherit firstLine=720 from Normal?
    Open in Word and check if the first line is indented relative to continuation.

    Paragraphs:
      P0: No direct w:ind — pure style inheritance (control, should show firstLine=720)
      P1: Direct w:ind w:left="-720" w:right="-360" — the SAFE scenario
      P2: Direct w:ind w:left="-720" w:right="-360" w:firstLine="0" — explicit zero (control)
      P3: Direct w:ind w:left="-720" w:right="-360" w:firstLine="720" — explicit 720 (control)
    """
    doc = Document()

    # Modify Normal style to add firstLine=720 (matching the SAFE document)
    normal_style = doc.styles["Normal"]
    normal_pPr = normal_style.element.get_or_add_pPr()
    # Remove any existing ind element
    for existing_ind in normal_pPr.findall(qn("w:ind")):
        normal_pPr.remove(existing_ind)
    normal_pPr.append(make_element("w:ind", {"w:firstLine": "720"}))

    # P0: No direct w:ind — pure style inheritance
    p0 = doc.add_paragraph(
        "P0 (control — no direct w:ind): This paragraph has no direct "
        "indentation override. It should inherit firstLine=720 from the "
        "Normal style. If this text wraps, the first line should be indented "
        "36pt (720 twips) relative to continuation lines. "
        "Extra words to force wrapping on a standard page width."
    )

    # P1: Direct w:ind w:left="-720" w:right="-360" — the SAFE scenario
    p1 = doc.add_paragraph(
        "P1 (SAFE scenario — direct left only): This paragraph has direct "
        "w:ind w:left=\"-720\" w:right=\"-360\" but omits firstLine/hanging. "
        "The question: does firstLine=720 inherit from Normal, or default to 0? "
        "If inherited, the first line is at -720+720=0 and continuation at -720. "
        "If zero, all lines are at -720. Check Word rendering. "
        "Extra words to force wrapping on a standard page width."
    )
    pPr1 = p1._p.get_or_add_pPr()
    pPr1.append(make_element("w:ind", {"w:left": "-720", "w:right": "-360"}))

    # P2: Explicit firstLine=0 (control)
    p2 = doc.add_paragraph(
        "P2 (control — explicit firstLine=0): This paragraph has direct "
        "w:ind w:left=\"-720\" w:right=\"-360\" w:firstLine=\"0\". "
        "All lines should start at -720 twips. No first-line offset. "
        "Extra words to force wrapping on a standard page width."
    )
    pPr2 = p2._p.get_or_add_pPr()
    pPr2.append(make_element("w:ind", {"w:left": "-720", "w:right": "-360", "w:firstLine": "0"}))

    # P3: Explicit firstLine=720 (control)
    p3 = doc.add_paragraph(
        "P3 (control — explicit firstLine=720): This paragraph has direct "
        "w:ind w:left=\"-720\" w:right=\"-360\" w:firstLine=\"720\". "
        "First line at -720+720=0, continuation at -720. This SHOULD show "
        "a visible first-line indent of 36pt. "
        "Extra words to force wrapping on a standard page width."
    )
    pPr3 = p3._p.get_or_add_pPr()
    pPr3.append(make_element("w:ind", {"w:left": "-720", "w:right": "-360", "w:firstLine": "720"}))

    save_fixture("formatting", "indent-firstline-cascade", doc, {
        "name": "indent-firstline-cascade",
        "spec_ref": "ISO 29500-1 §17.3.1.12",
        "description": (
            "Normal style has firstLine=720. "
            "P0: no direct w:ind (pure style inheritance). "
            "P1: direct w:ind left=-720, right=-360, no firstLine (the SAFE scenario). "
            "P2: direct w:ind left=-720, right=-360, firstLine=0 (explicit zero). "
            "P3: direct w:ind left=-720, right=-360, firstLine=720 (explicit 720)."
        ),
        "expected_behavior": (
            "P0: left=0, firstLine=720 (inherited from Normal). "
            "P1: QUESTION — does firstLine inherit from Normal (720) or default to 0? "
            "P2: left=-720, firstLine=0 (explicit, no first-line offset). "
            "P3: left=-720, firstLine=720 (explicit, first line at 0)."
        ),
        "open_in_word": (
            "Open this DOCX in Word. Compare P1 vs P2 and P1 vs P3. "
            "If P1 looks like P2 (no first-line offset), Word treats absent firstLine as 0. "
            "If P1 looks like P3 (first line indented), Word cascades firstLine from Normal."
        ),
    })


# =========================================================================
# REDLINE COMPAT (roundtrip verification for Word/GDocs compatibility)
# =========================================================================

def make_redline_compat_fixtures() -> None:
    print("\n── Redline Compat ──")
    make_table_props_roundtrip()
    make_para_id_roundtrip()


def make_table_props_roundtrip() -> None:
    """Table with explicit formatting for roundtrip verification.

    Verifies that tblPr (borders, width, indent, alignment),
    tblGrid (column widths), and tcPr (cell width, borders, shading,
    vAlign, margins) survive import → redline → export.
    """
    doc = Document()
    doc.add_paragraph("Before table.")

    tbl = doc.add_table(rows=2, cols=3)
    tbl.cell(0, 0).text = "A1"
    tbl.cell(0, 1).text = "B1"
    tbl.cell(0, 2).text = "C1"
    tbl.cell(1, 0).text = "A2"
    tbl.cell(1, 1).text = "B2"
    tbl.cell(1, 2).text = "C2"

    tbl_element = tbl._tbl
    tblPr = tbl_element.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_element.insert(0, tblPr)

    # Remove python-docx default tblW (type=auto, w=0) before adding ours
    for existing_w in tblPr.findall(w("tblW")):
        tblPr.remove(existing_w)
    tbl_w = make_element("w:tblW", {"w:w": "5000", "w:type": "dxa"})
    tblPr.append(tbl_w)

    # Table alignment: center
    tbl_jc = make_element("w:jc", {"w:val": "center"})
    tblPr.append(tbl_jc)

    # Table indent: 360 twips
    tbl_ind = make_element("w:tblInd", {"w:w": "360", "w:type": "dxa"})
    tblPr.append(tbl_ind)

    # Table borders
    tbl_borders = make_element("w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = make_element(f"w:{edge}", {
            "w:val": "single",
            "w:sz": "4",
            "w:space": "0",
            "w:color": "000000",
        })
        tbl_borders.append(border)
    tblPr.append(tbl_borders)

    # Default cell margins
    tbl_cell_mar = make_element("w:tblCellMar")
    for side, val in [("top", "72"), ("left", "115"), ("bottom", "72"), ("right", "115")]:
        margin = make_element(f"w:{side}", {"w:w": val, "w:type": "dxa"})
        tbl_cell_mar.append(margin)
    tblPr.append(tbl_cell_mar)

    # Grid columns
    existing_grid = tbl_element.find(w("tblGrid"))
    if existing_grid is not None:
        tbl_element.remove(existing_grid)
    tbl_grid = make_element("w:tblGrid")
    for col_w in ["1500", "2000", "1500"]:
        grid_col = make_element("w:gridCol", {"w:w": col_w})
        tbl_grid.append(grid_col)
    # Insert after tblPr
    tblPr_idx = list(tbl_element).index(tblPr)
    tbl_element.insert(tblPr_idx + 1, tbl_grid)

    # Cell (0,0): explicit tcPr with shading and vAlign
    rows = tbl_element.findall(w("tr"))
    cell_00 = rows[0].findall(w("tc"))[0]
    tcPr = cell_00.find(w("tcPr"))
    if tcPr is None:
        tcPr = make_element("w:tcPr")
        cell_00.insert(0, tcPr)
    # Remove python-docx default tcW before adding ours
    for existing_w in tcPr.findall(w("tcW")):
        tcPr.remove(existing_w)
    tc_w = make_element("w:tcW", {"w:w": "1500", "w:type": "dxa"})
    tcPr.append(tc_w)
    shd = make_element("w:shd", {"w:val": "clear", "w:color": "auto", "w:fill": "FFFF00"})
    tcPr.append(shd)
    v_align = make_element("w:vAlign", {"w:val": "center"})
    tcPr.append(v_align)

    doc.add_paragraph("After table.")

    save_fixture("redline-compat", "table-props-roundtrip", doc, {
        "name": "table-props-roundtrip",
        "spec_ref": "ISO 29500-1 §17.4",
        "description": "Table with tblPr, tblGrid, tcPr for roundtrip verification",
        "expected_behavior": "All table/cell properties survive redline export",
    })


def make_para_id_roundtrip() -> None:
    """Paragraphs with w14:paraId and w14:textId attributes.

    MS-DOCX §2.6.2.3/2.6.2.4: w14:paraId and w14:textId are 8-char
    hex strings on w:p elements used for paragraph identity tracking.
    """
    W14 = "http://schemas.microsoft.com/office/word/2010/wordml"

    doc = Document()

    # Ensure w14 namespace is declared on the document root
    root = doc.element
    root.set(f"{{{W14}}}Ignorable", "")  # forces namespace declaration
    # Remove the dummy attribute — we just needed the namespace registered
    if f"{{{W14}}}Ignorable" in root.attrib:
        del root.attrib[f"{{{W14}}}Ignorable"]

    p1 = doc.add_paragraph("First paragraph with paraId.")
    p1._p.set(f"{{{W14}}}paraId", "3B2A1C4D")
    p1._p.set(f"{{{W14}}}textId", "77A80E01")

    p2 = doc.add_paragraph("Second paragraph with paraId.")
    p2._p.set(f"{{{W14}}}paraId", "5F6E7D8C")
    p2._p.set(f"{{{W14}}}textId", "AA990B02")

    p3 = doc.add_paragraph("Third paragraph without paraId.")

    save_fixture("redline-compat", "para-id-roundtrip", doc, {
        "name": "para-id-roundtrip",
        "spec_ref": "MS-DOCX §2.6.2.3/2.6.2.4",
        "description": "Paragraphs with w14:paraId and w14:textId hex attributes",
        "expected_behavior": "paraId/textId survive import → redline → export roundtrip",
    })


# =========================================================================
# TABLE BORDERS SPEC COMPLIANCE (ISO 29500-1 §17.4.66, MS-OI29500 §17.4.66(a))
# =========================================================================

def make_table_borders_spec_fixtures() -> None:
    print("\n── Table Borders Spec ──")
    make_tbs_3x3_inside_borders()
    make_tbs_nil_border_suppresses_table_border()
    make_tbs_none_border_yields_to_table_border()
    make_tbs_adjacent_conflict_dashed_weight()
    make_tbs_cell_override_despite_lower_weight()
    make_tbs_empty_tbl_borders_element()


def make_tbs_3x3_inside_borders() -> None:
    """ISO 29500-1 §17.4.66 + §17.4.38: 3x3 table with distinct inside borders.

    tblBorders:
      - top/bottom/left/right: thick double blue sz=24 (outer)
      - insideH: single red sz=12
      - insideV: single green sz=8

    NO cells have tcBorders. After import:
      - Corner cell (0,0): top=outer, left=outer, bottom=insideH, right=insideV
      - Interior cell (1,1): top=insideH, bottom=insideH, left=insideV, right=insideV
      - Edge cell (0,1): top=outer, bottom=insideH, left=insideV, right=insideV
      - Edge cell (2,2): top=insideH, bottom=outer, left=insideV, right=outer
    """
    doc = Document()
    doc.add_paragraph("3x3 table with distinct insideH/insideV for border cascade tests.")

    tbl = doc.add_table(rows=3, cols=3)
    for r in range(3):
        for c in range(3):
            tbl.cell(r, c).text = f"R{r}C{c}"

    tbl_el = tbl._tbl
    tblPr = tbl_el.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_el.insert(0, tblPr)

    tbl_borders = make_element("w:tblBorders")
    # Outer borders: double blue sz=24
    for edge in ["top", "bottom", "left", "right"]:
        tbl_borders.append(make_element(f"w:{edge}", {
            "w:val": "double", "w:sz": "24", "w:color": "0000FF", "w:space": "0",
        }))
    # insideH: single red sz=12
    tbl_borders.append(make_element("w:insideH", {
        "w:val": "single", "w:sz": "12", "w:color": "FF0000", "w:space": "0",
    }))
    # insideV: single green sz=8
    tbl_borders.append(make_element("w:insideV", {
        "w:val": "single", "w:sz": "8", "w:color": "00FF00", "w:space": "0",
    }))
    tblPr.append(tbl_borders)

    save_fixture("table-borders-spec", "3x3-inside-borders", doc, {
        "name": "3x3-inside-borders",
        "spec_ref": "ISO 29500-1 §17.4.66, §17.4.38",
        "description": (
            "3x3 table with distinct outer (double blue sz=24), "
            "insideH (single red sz=12), insideV (single green sz=8). "
            "No tcBorders on any cell. Tests cascade into all 9 cell positions."
        ),
    })


def make_tbs_nil_border_suppresses_table_border() -> None:
    """MS-OI29500 §17.4.66(a): A cell border with w:val='nil' should suppress
    the border entirely, even when the table defines a border on that edge.

    2x2 table. tblBorders: single red sz=12 on all edges.
    Cell (0,0) has tcBorders with top=nil (all others absent).
    Per MS-OI29500: 'If the conflicting table cell border is nil,
    then no border shall be displayed.'

    So cell (0,0) top should be nil (no border displayed), while
    cell (0,0) other edges should inherit from the table.
    """
    doc = Document()
    doc.add_paragraph("nil border suppresses table border test.")

    tbl = doc.add_table(rows=2, cols=2)
    for r in range(2):
        for c in range(2):
            tbl.cell(r, c).text = f"R{r}C{c}"

    tbl_el = tbl._tbl
    tblPr = tbl_el.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_el.insert(0, tblPr)

    # Table borders: single red sz=12 on all edges
    tbl_borders = make_element("w:tblBorders")
    for edge in ["top", "bottom", "left", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{edge}", {
            "w:val": "single", "w:sz": "12", "w:color": "FF0000", "w:space": "0",
        }))
    tblPr.append(tbl_borders)

    # Cell (0,0): tcBorders with top=nil only
    row0 = tbl_el.findall(w("tr"))[0]
    cell_00 = row0.findall(w("tc"))[0]
    tcPr = cell_00.find(w("tcPr"))
    if tcPr is None:
        tcPr = make_element("w:tcPr")
        cell_00.insert(0, tcPr)
    tc_borders = make_element("w:tcBorders")
    tc_borders.append(make_element("w:top", {
        "w:val": "nil",
    }))
    tcPr.append(tc_borders)

    save_fixture("table-borders-spec", "nil-border-suppresses-table", doc, {
        "name": "nil-border-suppresses-table",
        "spec_ref": "MS-OI29500 §17.4.66(a)",
        "description": (
            "2x2 table with red single sz=12 table borders. "
            "Cell (0,0) has tcBorders with top=nil. "
            "Per MS-OI29500: nil means 'no border shall be displayed'. "
            "Cell (0,0) top should be nil/suppressed."
        ),
    })


def make_tbs_none_border_yields_to_table_border() -> None:
    """MS-OI29500 §17.4.66(a): A cell border with w:val='none' means
    'no border on this cell', but the opposing (table) border should
    be displayed.

    2x2 table. tblBorders: single red sz=12 on all edges.
    Cell (0,0) has tcBorders with top=none.
    Per MS-OI29500: 'If the conflicting table cell border is none
    (no border), then the opposing border shall be displayed.'

    After adjacent conflict resolution at the table-vs-cell level,
    the table border should win because 'none' yields.
    """
    doc = Document()
    doc.add_paragraph("none border yields to opposing table border test.")

    tbl = doc.add_table(rows=2, cols=2)
    for r in range(2):
        for c in range(2):
            tbl.cell(r, c).text = f"R{r}C{c}"

    tbl_el = tbl._tbl
    tblPr = tbl_el.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_el.insert(0, tblPr)

    # Table borders: single red sz=12 on all edges
    tbl_borders = make_element("w:tblBorders")
    for edge in ["top", "bottom", "left", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{edge}", {
            "w:val": "single", "w:sz": "12", "w:color": "FF0000", "w:space": "0",
        }))
    tblPr.append(tbl_borders)

    # Cell (0,0): tcBorders with top=none only
    row0 = tbl_el.findall(w("tr"))[0]
    cell_00 = row0.findall(w("tc"))[0]
    tcPr = cell_00.find(w("tcPr"))
    if tcPr is None:
        tcPr = make_element("w:tcPr")
        cell_00.insert(0, tcPr)
    tc_borders = make_element("w:tcBorders")
    tc_borders.append(make_element("w:top", {
        "w:val": "none", "w:sz": "0", "w:color": "auto", "w:space": "0",
    }))
    tcPr.append(tc_borders)

    save_fixture("table-borders-spec", "none-border-yields-to-table", doc, {
        "name": "none-border-yields-to-table",
        "spec_ref": "MS-OI29500 §17.4.66(a)",
        "description": (
            "2x2 table with red single sz=12 table borders. "
            "Cell (0,0) has tcBorders with top=none. "
            "Per MS-OI29500: 'none' means 'no cell border' so the opposing "
            "(table) border should be displayed."
        ),
    })


def make_tbs_adjacent_conflict_dashed_weight() -> None:
    """MS-OI29500 §17.4.66(a): Dashed/dotted borders get weight=1
    regardless of border width and number.

    1x2 table (no table borders).
    Cell (0,0): right = dashed sz=48 (per MS, weight should be 1)
    Cell (0,1): left = single sz=4 (weight = 4 * 1 = 4)

    Per MS-OI29500: 'The borders with dotted and dashed styles shall
    be assigned the weight 1 regardless of the border width and number.'

    So the single sz=4 (weight=4) should beat dashed sz=48 (weight=1).
    """
    doc = Document()
    doc.add_paragraph("Dashed border weight = 1 regardless of size.")

    tbl = doc.add_table(rows=1, cols=2)
    tbl.cell(0, 0).text = "Dashed thick"
    tbl.cell(0, 1).text = "Single thin"

    tbl_el = tbl._tbl
    tblPr = tbl_el.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_el.insert(0, tblPr)

    # Cell (0,0): right = dashed sz=48 red
    row0 = tbl_el.findall(w("tr"))[0]
    cell_00 = row0.findall(w("tc"))[0]
    tcPr_00 = cell_00.find(w("tcPr"))
    if tcPr_00 is None:
        tcPr_00 = make_element("w:tcPr")
        cell_00.insert(0, tcPr_00)
    tc_borders_00 = make_element("w:tcBorders")
    tc_borders_00.append(make_element("w:right", {
        "w:val": "dashed", "w:sz": "48", "w:color": "FF0000", "w:space": "0",
    }))
    tcPr_00.append(tc_borders_00)

    # Cell (0,1): left = single sz=4 blue
    cell_01 = row0.findall(w("tc"))[1]
    tcPr_01 = cell_01.find(w("tcPr"))
    if tcPr_01 is None:
        tcPr_01 = make_element("w:tcPr")
        cell_01.insert(0, tcPr_01)
    tc_borders_01 = make_element("w:tcBorders")
    tc_borders_01.append(make_element("w:left", {
        "w:val": "single", "w:sz": "4", "w:color": "0000FF", "w:space": "0",
    }))
    tcPr_01.append(tc_borders_01)

    save_fixture("table-borders-spec", "adjacent-dashed-weight", doc, {
        "name": "adjacent-dashed-weight",
        "spec_ref": "MS-OI29500 §17.4.66(a)",
        "description": (
            "1x2 table. Cell (0,0) right=dashed sz=48 red. "
            "Cell (0,1) left=single sz=4 blue. "
            "Per MS-OI29500, dashed styles get weight=1 regardless of size. "
            "Single sz=4 (weight=4) should beat dashed sz=48 (weight=1)."
        ),
    })


def make_tbs_cell_override_despite_lower_weight() -> None:
    """ISO 29500-1 §17.4.38 + §17.4.66: Cell borders always override
    table borders, regardless of weight.

    Per §17.4.38: 'If there is a cell border, then the cell border
    shall be displayed.'

    2x2 table. tblBorders: thick double blue sz=48 on all edges.
    Cell (0,0) has tcBorders: single green sz=4 on all edges.

    Despite the table border having higher weight, the cell border
    wins because cell borders take precedence over table borders.
    """
    doc = Document()
    doc.add_paragraph("Cell border overrides table border regardless of weight.")

    tbl = doc.add_table(rows=2, cols=2)
    for r in range(2):
        for c in range(2):
            tbl.cell(r, c).text = f"R{r}C{c}"

    tbl_el = tbl._tbl
    tblPr = tbl_el.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_el.insert(0, tblPr)

    # Table borders: thick double blue sz=48
    tbl_borders = make_element("w:tblBorders")
    for edge in ["top", "bottom", "left", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{edge}", {
            "w:val": "double", "w:sz": "48", "w:color": "0000FF", "w:space": "0",
        }))
    tblPr.append(tbl_borders)

    # Cell (0,0): single green sz=4 on all edges
    row0 = tbl_el.findall(w("tr"))[0]
    cell_00 = row0.findall(w("tc"))[0]
    tcPr = cell_00.find(w("tcPr"))
    if tcPr is None:
        tcPr = make_element("w:tcPr")
        cell_00.insert(0, tcPr)
    tc_borders = make_element("w:tcBorders")
    for edge in ["top", "bottom", "left", "right"]:
        tc_borders.append(make_element(f"w:{edge}", {
            "w:val": "single", "w:sz": "4", "w:color": "00FF00", "w:space": "0",
        }))
    tcPr.append(tc_borders)

    save_fixture("table-borders-spec", "cell-override-lower-weight", doc, {
        "name": "cell-override-lower-weight",
        "spec_ref": "ISO 29500-1 §17.4.38, §17.4.66",
        "description": (
            "2x2 table with heavy double blue sz=48 table borders. "
            "Cell (0,0) has thin single green sz=4 cell borders. "
            "Cell border wins over table border regardless of weight."
        ),
    })


def make_tbs_empty_tbl_borders_element() -> None:
    """Edge case: <w:tblBorders/> with no child border elements.

    Per ISO 29500-1 §17.4.38: 'If this element is omitted, then this
    table shall have the borders specified by the associated table
    style.' When present but empty, all explicit borders are absent.

    2x2 table. tblBorders element is present but has NO children.
    Cells should have no borders (empty tblBorders means no borders
    are defined at the table level).
    """
    doc = Document()
    doc.add_paragraph("Empty tblBorders element (present but no children).")

    tbl = doc.add_table(rows=2, cols=2)
    for r in range(2):
        for c in range(2):
            tbl.cell(r, c).text = f"R{r}C{c}"

    tbl_el = tbl._tbl
    tblPr = tbl_el.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_el.insert(0, tblPr)

    # Remove any existing tblBorders python-docx may have added
    for existing in tblPr.findall(w("tblBorders")):
        tblPr.remove(existing)

    # Add empty tblBorders (present element, no children)
    empty_borders = make_element("w:tblBorders")
    tblPr.append(empty_borders)

    save_fixture("table-borders-spec", "empty-tbl-borders", doc, {
        "name": "empty-tbl-borders",
        "spec_ref": "ISO 29500-1 §17.4.38",
        "description": (
            "2x2 table with <w:tblBorders/> (present but empty, no child elements). "
            "No borders should be defined at the table level. "
            "Cells should have no inherited borders."
        ),
    })


if __name__ == "__main__":
    main()

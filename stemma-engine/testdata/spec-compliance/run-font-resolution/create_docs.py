# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for run font resolution & style cascade tests.

Tests: font resolution through themes, cs/rtl property activation,
cs font slot forcing, style-defined position/kern cascade.

Run:  uv run create_docs.py
"""

import json
from pathlib import Path

from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  run-font-resolution/{name}/")


def set_theme_fonts(doc, minor_latin=None, minor_cs=None, major_latin=None, major_cs=None):
    """Modify the theme fonts in the document's theme1.xml."""
    theme_rel = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme"
    doc_rel = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument"
    theme_part = doc.part.package.part_related_by(doc_rel).part_related_by(theme_rel)
    theme_xml = etree.fromstring(theme_part.blob)
    ns = {"a": A_NS}

    if minor_latin is not None:
        el = theme_xml.find(".//a:minorFont/a:latin", ns)
        if el is not None:
            el.set("typeface", minor_latin)
    if minor_cs is not None:
        el = theme_xml.find(".//a:minorFont/a:cs", ns)
        if el is not None:
            el.set("typeface", minor_cs)
    if major_latin is not None:
        el = theme_xml.find(".//a:majorFont/a:latin", ns)
        if el is not None:
            el.set("typeface", major_latin)
    if major_cs is not None:
        el = theme_xml.find(".//a:majorFont/a:cs", ns)
        if el is not None:
            el.set("typeface", major_cs)

    theme_part._blob = etree.tostring(theme_xml, xml_declaration=True, encoding="UTF-8", standalone=True)


# ── 1. Theme font resolves through cascade ────────────────────────────


def make_theme_font_cascade() -> None:
    """Paragraph style with asciiTheme=minorHAnsi; theme defines minorHAnsi=Calibri.

    ECMA-376 §17.3.2.26: theme font references in styles should resolve
    through the theme's font scheme.
    """
    doc = Document()

    # Set theme minor latin to "Calibri"
    set_theme_fonts(doc, minor_latin="Calibri")

    # Create a paragraph style with asciiTheme=minorHAnsi
    styles_el = doc.styles.element
    para_s = make_element("w:style", {"w:type": "paragraph", "w:styleId": "ThemeFontPara"})
    para_s.append(make_element("w:name", {"w:val": "Theme Font Para"}))
    rpr = make_element("w:rPr")
    rfonts = make_element("w:rFonts", {
        "w:asciiTheme": "minorHAnsi",
        "w:hAnsiTheme": "minorHAnsi",
    })
    rpr.append(rfonts)
    para_s.append(rpr)
    styles_el.append(para_s)

    # A paragraph using that style
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pPr.append(make_element("w:pStyle", {"w:val": "ThemeFontPara"}))

    p.add_run("Theme cascade text")

    save_fixture("theme-font-cascade", doc, {
        "name": "theme-font-cascade",
        "spec_ref": "ECMA-376 §17.3.2.26",
        "description": "Paragraph style with asciiTheme=minorHAnsi, "
                       "theme defines minorHAnsi=Calibri",
        "expected_behavior": "Run font_family should resolve to 'Calibri' from theme",
    })


# ── 2. cs/rtl activates complex script properties ────────────────────


def make_cs_rtl_activation() -> None:
    """Run with bCs + i + cs — cs should activate bCs (bold) and suppress i (non-CS).

    ECMA-376 §17.3.2.7/§17.3.2.30: cs=true causes bCs/iCs/szCs to replace
    b/i/sz in the formatting output.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Run: bCs=on, i=on, cs=on
    # Expected: bold (from bCs since cs active), NOT italic (i is non-CS)
    run = p.add_run("CS activated run")
    rpr = run._r.get_or_add_rPr()
    rpr.append(make_element("w:bCs"))
    rpr.append(make_element("w:i"))
    rpr.append(make_element("w:cs"))

    # Control run: same marks but no cs
    # Expected: NOT bold (bCs ignored without cs), italic (i applies)
    run2 = p.add_run("Non-CS control run")
    rpr2 = run2._r.get_or_add_rPr()
    rpr2.append(make_element("w:bCs"))
    rpr2.append(make_element("w:i"))

    save_fixture("cs-rtl-activation", doc, {
        "name": "cs-rtl-activation",
        "spec_ref": "ECMA-376 §17.3.2.7/§17.3.2.30",
        "description": "Run with bCs + i + cs: cs activates bCs (bold) "
                       "and suppresses i (non-CS italic)",
        "expected_behavior": "CS run: bold=on (from bCs), italic=off (i is non-CS). "
                             "Control run: bold=off, italic=on.",
    })


# ── 3. cs forces cs font slot ────────────────────────────────────────


def make_cs_forces_cs_font() -> None:
    """Run with ascii=Courier, cs=Arial, cs=true on English text.

    MS-OI29500 §17.3.2.26(b): when cs is set, the cs font is used for ALL
    characters regardless of Unicode range.
    """
    doc = Document()
    p = doc.add_paragraph()

    run = p.add_run("CS font override text")
    rpr = run._r.get_or_add_rPr()
    rfonts = make_element("w:rFonts", {
        "w:ascii": "Courier",
        "w:hAnsi": "Courier",
        "w:cs": "Arial",
    })
    rpr.append(rfonts)
    rpr.append(make_element("w:cs"))

    save_fixture("cs-forces-cs-font", doc, {
        "name": "cs-forces-cs-font",
        "spec_ref": "MS-OI29500 §17.3.2.26(b)",
        "description": "Run with ascii=Courier, cs=Arial, w:cs=true on English text",
        "expected_behavior": "font_family should be 'Arial' (cs font) not 'Courier'",
    })


# ── 4. Style-defined position cascade ────────────────────────────────


def make_style_position_cascade() -> None:
    """Character style defining position=24, applied to a run.

    parse_rpr_marks does NOT parse w:position from style definitions,
    so this property will not cascade from the character style.
    """
    doc = Document()
    styles_el = doc.styles.element

    # Character style with position=24
    char_s = make_element("w:style", {"w:type": "character", "w:styleId": "RaisedChar"})
    char_s.append(make_element("w:name", {"w:val": "Raised Char"}))
    rpr = make_element("w:rPr")
    rpr.append(make_element("w:position", {"w:val": "24"}))
    char_s.append(rpr)
    styles_el.append(char_s)

    p = doc.add_paragraph()
    run = p.add_run("Style position text")
    run_rpr = run._r.get_or_add_rPr()
    run_rpr.insert(0, make_element("w:rStyle", {"w:val": "RaisedChar"}))

    save_fixture("style-position-cascade", doc, {
        "name": "style-position-cascade",
        "spec_ref": "ISO 29500-1 §17.3.2.19",
        "description": "Character style with position=24 applied to a run",
        "expected_behavior": "Run position should be Some(24) from style cascade. "
                             "KNOWN GAP: parse_rpr_marks does not parse w:position.",
    })


# ── 5. Style-defined kern cascade ────────────────────────────────────


def make_style_kern_cascade() -> None:
    """Character style defining kern=28, applied to a run.

    parse_rpr_marks does NOT parse w:kern from style definitions,
    so this property will not cascade from the character style.
    """
    doc = Document()
    styles_el = doc.styles.element

    # Character style with kern=28
    char_s = make_element("w:style", {"w:type": "character", "w:styleId": "KernedChar"})
    char_s.append(make_element("w:name", {"w:val": "Kerned Char"}))
    rpr = make_element("w:rPr")
    rpr.append(make_element("w:kern", {"w:val": "28"}))
    char_s.append(rpr)
    styles_el.append(char_s)

    p = doc.add_paragraph()
    run = p.add_run("Style kern text")
    run_rpr = run._r.get_or_add_rPr()
    run_rpr.insert(0, make_element("w:rStyle", {"w:val": "KernedChar"}))

    save_fixture("style-kern-cascade", doc, {
        "name": "style-kern-cascade",
        "spec_ref": "ISO 29500-1 §17.3.2.19(a)",
        "description": "Character style with kern=28 applied to a run",
        "expected_behavior": "Run kern should be Some(28) from style cascade. "
                             "KNOWN GAP: parse_rpr_marks does not parse w:kern.",
    })


# ── 6. East Asian Times New Roman bypass ─────────────────────────────


def make_east_asian_tnr_bypass() -> None:
    """Run with ascii=Calibri, hAnsi=Calibri, eastAsia=Times New Roman.

    MS-OI29500 §2.1.88(d): If eastAsia font is "Times New Roman" and
    ascii==hAnsi, the ascii font is used for all characters including
    those that would normally use the eastAsia slot.
    """
    doc = Document()
    p = doc.add_paragraph()

    # Run with eastAsia="Times New Roman" and ascii=hAnsi="Calibri"
    run = p.add_run("East Asian TNR bypass text")
    rpr = run._r.get_or_add_rPr()
    rfonts = make_element("w:rFonts", {
        "w:ascii": "Calibri",
        "w:hAnsi": "Calibri",
        "w:eastAsia": "Times New Roman",
    })
    rpr.append(rfonts)

    save_fixture("east-asian-tnr-bypass", doc, {
        "name": "east-asian-tnr-bypass",
        "spec_ref": "MS-OI29500 §2.1.88(d)",
        "description": "Run with ascii=hAnsi=Calibri, eastAsia=Times New Roman",
        "expected_behavior": "font_family=Calibri, font_east_asia=Calibri "
                             "(TNR bypass: eastAsia replaced with ascii font)",
    })


if __name__ == "__main__":
    print("── Run Font Resolution fixtures ──")
    make_theme_font_cascade()
    make_cs_rtl_activation()
    make_cs_forces_cs_font()
    make_style_position_cascade()
    make_style_kern_cascade()
    make_east_asian_tnr_bypass()
    print("Done.")

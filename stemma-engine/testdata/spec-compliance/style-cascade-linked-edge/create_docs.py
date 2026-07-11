# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for linked-style edge-case tests.

Exercises edge cases around linked style resolution (ISO 29500-1 §17.7.4.6),
specifically:
  1. Linked char style's basedOn chain properties lost by overcorrection
  2. Default character style not applied to unstyled runs (§17.7.4.17)
  3. Linked para style as basedOn target — does link leak into child?
  4. Direct formatting override of linked char style properties
  5. Toggle XOR semantics between linked pair and direct formatting

Run:  python create_docs.py
"""

import json
from pathlib import Path
from lxml import etree

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  style-cascade-linked-edge/{name}/")


# =========================================================================
# Fixture 1: Linked char basedOn chain property loss
#
# The fix for Bug 2 (overlay with raw marks instead of resolved chain)
# went too far: properties from the linked char style's basedOn chain
# that the para style doesn't set are now lost.
#
# Per §17.7.4.6, the linked char style provides the run-level rPr for
# the paragraph style. That should include the char style's full
# resolved chain, not just its own raw marks. The correct behavior is:
# - Properties explicitly set by the para style's rPr should be preserved
# - Properties from the linked char style's resolved chain that don't
#   conflict with the para style's explicit properties should apply
# =========================================================================

def make_linked_based_on_loss():
    """
    BaseChar: character style, underline=on, color=0000FF, highlight=yellow
    LinkedChar: character style, basedOn BaseChar, bold=on only
      linked to LinkedPara
    LinkedPara: paragraph style, sz=32, font_family=Arial
      linked to LinkedChar

    Para 1 (LinkedPara, no explicit char style):
      Expected:
        - sz=32 from para rPr (para's explicit value preserved)
        - font_family=Arial from para rPr (para's explicit value preserved)
        - bold=on from LinkedChar's raw marks (char explicitly sets it)
        - underline=on from BaseChar via LinkedChar's basedOn chain
        - color=0000FF from BaseChar via LinkedChar's basedOn chain
        - highlight=yellow from BaseChar via LinkedChar's basedOn chain
      BUG: underline, color, highlight are lost because overlay uses raw marks

    Para 2 (LinkedPara, explicit LinkedChar on run):
      Expected: char style wins for all properties it resolves
        - sz=32 from para rPr? No: explicit char style wins for font_size
          if it has one. LinkedChar resolves to no sz (BaseChar doesn't set sz
          either in this fixture). So sz falls through to para.
          Wait — BaseChar doesn't set sz. So explicit LinkedChar resolves
          bold=on, underline=on, color=0000FF, highlight=yellow (from chain).
          For sz: char style doesn't have it, para style has 32.
        - bold=on from LinkedChar
        - underline=on from BaseChar via resolved chain
        - color=0000FF from BaseChar via resolved chain
        - highlight=yellow from BaseChar
    """
    doc = Document()
    styles_el = doc.styles.element

    # BaseChar: character style with underline, color, highlight
    style_base = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "BaseChar",
    })
    style_base.append(make_element("w:name", {"w:val": "Base Char"}))
    base_rpr = make_element("w:rPr")
    base_rpr.append(make_element("w:u", {"w:val": "single"}))
    base_rpr.append(make_element("w:color", {"w:val": "0000FF"}))
    base_rpr.append(make_element("w:highlight", {"w:val": "yellow"}))
    style_base.append(base_rpr)
    styles_el.append(style_base)

    # LinkedChar: character style, basedOn BaseChar, bold=on only
    style_char = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "LinkedChar",
    })
    style_char.append(make_element("w:name", {"w:val": "Linked Char"}))
    style_char.append(make_element("w:basedOn", {"w:val": "BaseChar"}))
    style_char.append(make_element("w:link", {"w:val": "LinkedPara"}))
    char_rpr = make_element("w:rPr")
    char_rpr.append(make_element("w:b"))
    style_char.append(char_rpr)
    styles_el.append(style_char)

    # LinkedPara: paragraph style, sz=32, font=Arial, linked to LinkedChar
    style_para = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "LinkedPara",
    })
    style_para.append(make_element("w:name", {"w:val": "Linked Para"}))
    style_para.append(make_element("w:link", {"w:val": "LinkedChar"}))
    para_rpr = make_element("w:rPr")
    para_rpr.append(make_element("w:sz", {"w:val": "32"}))
    para_rpr.append(make_element("w:rFonts", {"w:ascii": "Arial", "w:hAnsi": "Arial"}))
    style_para.append(para_rpr)
    styles_el.append(style_para)

    # Para 1: LinkedPara, no explicit char style on run
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "LinkedPara"}))
    p1.add_run("Linked basedOn loss test")

    # Para 2: LinkedPara + explicit LinkedChar on run
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    p2_ppr.append(make_element("w:pStyle", {"w:val": "LinkedPara"}))
    r2 = p2.add_run("Explicit linked char")
    r2_rpr = r2._element.get_or_add_rPr()
    r2_rpr.insert(0, make_element("w:rStyle", {"w:val": "LinkedChar"}))

    save_fixture("linked-based-on-loss", doc, {
        "spec": "ISO 29500-1 §17.7.4.6 + §17.7.4.3",
        "description": (
            "LinkedPara (sz=32, Arial) linked to LinkedChar (bold=on, "
            "basedOn BaseChar(underline=on, color=0000FF, highlight=yellow)). "
            "Tests that char style's basedOn properties reach runs via linked "
            "overlay, without overwriting para style's explicit properties."
        ),
        "expected": {
            "para_1": (
                "sz=32 (para), Arial (para), bold=ON (char raw), "
                "underline=ON (char basedOn), color=0000FF (char basedOn), "
                "highlight=yellow (char basedOn)"
            ),
            "para_2": (
                "explicit LinkedChar on run: bold=ON, underline=ON, "
                "color=0000FF, highlight=yellow (all from resolved char chain). "
                "sz=32 (para, since char chain doesn't set sz)."
            ),
        },
    })


# =========================================================================
# Fixture 2: Default character style (§17.7.4.17) — unstyled runs
#
# Unstyled runs should inherit from the default character style when
# it exists and is not "DefaultParagraphFont" (which Word ignores).
# =========================================================================

def make_default_char_style():
    """
    CustomDefaultChar: default character style, italic=on, sz=28, color=FF0000
    NormalPara: default paragraph style, no rPr

    Para 1 (NormalPara, unstyled run):
      Expected: italic=on, sz=28, color=FF0000 from default char style
      BUG: resolve() doesn't apply default char style when char_style_id=None

    Para 2 (NormalPara, run with explicit char style OverrideChar):
      OverrideChar: bold=on only (no sz, no color, no italic)
      Expected: bold=on (from OverrideChar), italic=on (from default char),
                sz=28 (from default char), color=FF0000 (from default char)
      The resolve() function merges default char underneath explicit char,
      so properties the explicit style doesn't set fall through.
    """
    doc = Document()
    styles_el = doc.styles.element

    # Remove any existing DefaultParagraphFont to avoid conflict
    for existing in list(styles_el.findall(qn("w:style"))):
        sid = existing.get(qn("w:styleId"))
        if sid == "DefaultParagraphFont":
            styles_el.remove(existing)

    # CustomDefaultChar: default character style
    style_default = make_element("w:style", {
        "w:type": "character",
        "w:default": "1",
        "w:styleId": "CustomDefaultChar",
    })
    style_default.append(make_element("w:name", {"w:val": "Custom Default Char"}))
    default_rpr = make_element("w:rPr")
    default_rpr.append(make_element("w:i"))
    default_rpr.append(make_element("w:sz", {"w:val": "28"}))
    default_rpr.append(make_element("w:color", {"w:val": "FF0000"}))
    style_default.append(default_rpr)
    styles_el.append(style_default)

    # OverrideChar: explicit character style, bold=on only
    style_override = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "OverrideChar",
    })
    style_override.append(make_element("w:name", {"w:val": "Override Char"}))
    override_rpr = make_element("w:rPr")
    override_rpr.append(make_element("w:b"))
    style_override.append(override_rpr)
    styles_el.append(style_override)

    # Para 1: unstyled run
    p1 = doc.add_paragraph()
    # Remove any auto pStyle
    p1_ppr = p1._element.find(qn("w:pPr"))
    if p1_ppr is not None:
        ps = p1_ppr.find(qn("w:pStyle"))
        if ps is not None:
            p1_ppr.remove(ps)
    p1.add_run("Unstyled default char")

    # Para 2: run with explicit OverrideChar
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.find(qn("w:pPr"))
    if p2_ppr is not None:
        ps = p2_ppr.find(qn("w:pStyle"))
        if ps is not None:
            p2_ppr.remove(ps)
    r2 = p2.add_run("Explicit override char")
    r2_rpr = r2._element.get_or_add_rPr()
    r2_rpr.insert(0, make_element("w:rStyle", {"w:val": "OverrideChar"}))

    save_fixture("default-char-unstyled", doc, {
        "spec": "ISO 29500-1 §17.7.4.17 + §17.7.2",
        "description": (
            "CustomDefaultChar (default char style, italic=on, sz=28, "
            "color=FF0000). Tests that unstyled runs inherit from the "
            "default character style, and that explicit char styles merge "
            "with the default char style underneath."
        ),
        "expected": {
            "para_1": "italic=ON, sz=28, color=FF0000 (from default char style)",
            "para_2": (
                "bold=ON (from OverrideChar), italic=ON (default char), "
                "sz=28 (default char), color=FF0000 (default char)"
            ),
        },
    })


# =========================================================================
# Fixture 3: Linked para style as basedOn target
#
# When StyleA basedOn LinkedPara (which is linked to LinkedChar),
# does the linked char overlay leak into StyleA's resolution?
# Per spec, the link applies to LinkedPara directly — StyleA should
# only get LinkedPara's basedOn chain rPr, not the linked overlay.
# =========================================================================

def make_linked_based_on_target():
    """
    BaseCharT: character style, underline=on
    LinkedCharT: character style, basedOn BaseCharT, italic=on
      linked to LinkedParaT
    LinkedParaT: paragraph style, bold=on, sz=28
      linked to LinkedCharT
    ChildPara: paragraph style, basedOn LinkedParaT, color=00FF00

    Para 1 (LinkedParaT, unstyled run):
      Expected: bold=on (para), sz=28 (para), italic=on (linked char raw),
                underline=on (linked char basedOn)
      NOTE: This may also show basedOn loss from fixture 1.

    Para 2 (ChildPara, unstyled run):
      Expected: bold=on (from LinkedParaT via basedOn), sz=28 (from LinkedParaT),
                color=00FF00 (from ChildPara's own rPr)
      The link from LinkedParaT should NOT leak into ChildPara's resolution.
      italic and underline from LinkedCharT should NOT appear on ChildPara runs.
    """
    doc = Document()
    styles_el = doc.styles.element

    # BaseCharT: character style with underline=on
    style_base = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "BaseCharT",
    })
    style_base.append(make_element("w:name", {"w:val": "Base Char T"}))
    base_rpr = make_element("w:rPr")
    base_rpr.append(make_element("w:u", {"w:val": "single"}))
    style_base.append(base_rpr)
    styles_el.append(style_base)

    # LinkedCharT: character style, basedOn BaseCharT, italic=on
    style_char = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "LinkedCharT",
    })
    style_char.append(make_element("w:name", {"w:val": "Linked Char T"}))
    style_char.append(make_element("w:basedOn", {"w:val": "BaseCharT"}))
    style_char.append(make_element("w:link", {"w:val": "LinkedParaT"}))
    char_rpr = make_element("w:rPr")
    char_rpr.append(make_element("w:i"))
    style_char.append(char_rpr)
    styles_el.append(style_char)

    # LinkedParaT: paragraph style, bold=on, sz=28, linked to LinkedCharT
    style_para = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "LinkedParaT",
    })
    style_para.append(make_element("w:name", {"w:val": "Linked Para T"}))
    style_para.append(make_element("w:link", {"w:val": "LinkedCharT"}))
    para_rpr = make_element("w:rPr")
    para_rpr.append(make_element("w:b"))
    para_rpr.append(make_element("w:sz", {"w:val": "28"}))
    style_para.append(para_rpr)
    styles_el.append(style_para)

    # ChildPara: paragraph style, basedOn LinkedParaT, adds color=00FF00
    style_child = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "ChildPara",
    })
    style_child.append(make_element("w:name", {"w:val": "Child Para"}))
    style_child.append(make_element("w:basedOn", {"w:val": "LinkedParaT"}))
    child_rpr = make_element("w:rPr")
    child_rpr.append(make_element("w:color", {"w:val": "00FF00"}))
    style_child.append(child_rpr)
    styles_el.append(style_child)

    # Para 1: LinkedParaT, unstyled run
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "LinkedParaT"}))
    p1.add_run("Linked parent direct")

    # Para 2: ChildPara, unstyled run
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    p2_ppr.append(make_element("w:pStyle", {"w:val": "ChildPara"}))
    p2.add_run("Child of linked parent")

    save_fixture("linked-based-on-target", doc, {
        "spec": "ISO 29500-1 §17.7.4.6 + §17.7.4.3",
        "description": (
            "LinkedParaT (bold=on, sz=28) linked to LinkedCharT (italic=on, "
            "basedOn BaseCharT(underline=on)). ChildPara basedOn LinkedParaT "
            "(color=00FF00). Tests whether linked overlay leaks into child "
            "styles via basedOn."
        ),
        "expected": {
            "para_1": (
                "bold=on (para), sz=28 (para), italic=on (linked char raw), "
                "underline=on (linked char basedOn via BaseCharT)"
            ),
            "para_2": (
                "bold=on (from LinkedParaT via basedOn), sz=28 (LinkedParaT), "
                "color=00FF00 (ChildPara). NO italic, NO underline — linked "
                "overlay from LinkedParaT should not leak into ChildPara."
            ),
        },
    })


# =========================================================================
# Fixture 4: Direct formatting override of linked char properties
#
# When a run in a linked paragraph style has direct formatting, the
# direct formatting should override the linked char style's properties.
# =========================================================================

def make_direct_override_linked():
    """
    SimpleChar: character style, bold=on, sz=24
      linked to SimplePara
    SimplePara: paragraph style, sz=28, color=FF0000
      linked to SimpleChar

    Para 1 (SimplePara, run with direct sz=40):
      Expected: sz=40 (direct), bold=on (linked char), color=FF0000 (para)

    Para 2 (SimplePara, run with direct bold=off):
      Expected: bold=off (direct overrides linked char's bold=on)

    Para 3 (SimplePara, run with direct color=00FF00):
      Expected: color=00FF00 (direct overrides para's color=FF0000),
                bold=on (linked char), sz=28 (para)
    """
    doc = Document()
    styles_el = doc.styles.element

    # SimpleChar: character style, bold=on, sz=24
    style_char = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "SimpleChar",
    })
    style_char.append(make_element("w:name", {"w:val": "Simple Char"}))
    style_char.append(make_element("w:link", {"w:val": "SimplePara"}))
    char_rpr = make_element("w:rPr")
    char_rpr.append(make_element("w:b"))
    char_rpr.append(make_element("w:sz", {"w:val": "24"}))
    style_char.append(char_rpr)
    styles_el.append(style_char)

    # SimplePara: paragraph style, sz=28, color=FF0000
    style_para = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "SimplePara",
    })
    style_para.append(make_element("w:name", {"w:val": "Simple Para"}))
    style_para.append(make_element("w:link", {"w:val": "SimpleChar"}))
    para_rpr = make_element("w:rPr")
    para_rpr.append(make_element("w:sz", {"w:val": "28"}))
    para_rpr.append(make_element("w:color", {"w:val": "FF0000"}))
    style_para.append(para_rpr)
    styles_el.append(style_para)

    # Para 1: SimplePara, run with direct sz=40
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "SimplePara"}))
    r1 = p1.add_run("Direct sz override")
    r1_rpr = r1._element.get_or_add_rPr()
    r1_rpr.append(make_element("w:sz", {"w:val": "40"}))

    # Para 2: SimplePara, run with direct bold=off
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    p2_ppr.append(make_element("w:pStyle", {"w:val": "SimplePara"}))
    r2 = p2.add_run("Direct bold off")
    r2_rpr = r2._element.get_or_add_rPr()
    r2_rpr.append(make_element("w:b", {"w:val": "0"}))

    # Para 3: SimplePara, run with direct color=00FF00
    p3 = doc.add_paragraph()
    p3_ppr = p3._element.get_or_add_pPr()
    p3_ppr.append(make_element("w:pStyle", {"w:val": "SimplePara"}))
    r3 = p3.add_run("Direct color override")
    r3_rpr = r3._element.get_or_add_rPr()
    r3_rpr.append(make_element("w:color", {"w:val": "00FF00"}))

    save_fixture("direct-override-linked", doc, {
        "spec": "ISO 29500-1 §17.7.2 + §17.7.4.6",
        "description": (
            "SimplePara (sz=28, color=FF0000) linked to SimpleChar (bold=on, sz=24). "
            "Tests that direct formatting on runs overrides linked char properties."
        ),
        "expected": {
            "para_1": "sz=40 (direct), bold=on (linked char), color=FF0000 (para)",
            "para_2": "bold=off (direct), sz=24 (linked char), color=FF0000 (para)",
            "para_3": "color=00FF00 (direct), bold=on (linked char), sz=24 (linked char)",
        },
    })


# =========================================================================
# Fixture 5: Linked char style's basedOn chain — font_family leak
#
# A variation of Fixture 1 but specifically testing font_family, which
# uses a different XML element (w:rFonts) and may have different parsing.
# =========================================================================

def make_linked_font_family_leak():
    """
    BaseFontChar: character style, font=Courier New, italic=on
    FontLinkedChar: character style, basedOn BaseFontChar, bold=on
      linked to FontLinkedPara
    FontLinkedPara: paragraph style, font=Arial, sz=28
      linked to FontLinkedChar

    Para 1 (FontLinkedPara, unstyled run):
      Expected:
        - font=Arial from para rPr (para explicitly sets it)
        - sz=28 from para rPr
        - bold=on from FontLinkedChar's raw marks
        - italic=on from BaseFontChar via char basedOn chain
      BUG potential: italic from BaseFontChar may be lost due to raw marks overlay

    Para 2 (no para style, FontLinkedChar as explicit char style):
      Expected:
        - font=Courier New from BaseFontChar (resolved char chain)
        - italic=on from BaseFontChar
        - bold=on from FontLinkedChar
    """
    doc = Document()
    styles_el = doc.styles.element

    # BaseFontChar: character style, font=Courier New, italic=on
    style_base = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "BaseFontChar",
    })
    style_base.append(make_element("w:name", {"w:val": "Base Font Char"}))
    base_rpr = make_element("w:rPr")
    base_rpr.append(make_element("w:rFonts", {
        "w:ascii": "Courier New",
        "w:hAnsi": "Courier New",
    }))
    base_rpr.append(make_element("w:i"))
    style_base.append(base_rpr)
    styles_el.append(style_base)

    # FontLinkedChar: character style, basedOn BaseFontChar, bold=on
    style_char = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "FontLinkedChar",
    })
    style_char.append(make_element("w:name", {"w:val": "Font Linked Char"}))
    style_char.append(make_element("w:basedOn", {"w:val": "BaseFontChar"}))
    style_char.append(make_element("w:link", {"w:val": "FontLinkedPara"}))
    char_rpr = make_element("w:rPr")
    char_rpr.append(make_element("w:b"))
    style_char.append(char_rpr)
    styles_el.append(style_char)

    # FontLinkedPara: paragraph style, font=Arial, sz=28
    style_para = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "FontLinkedPara",
    })
    style_para.append(make_element("w:name", {"w:val": "Font Linked Para"}))
    style_para.append(make_element("w:link", {"w:val": "FontLinkedChar"}))
    para_rpr = make_element("w:rPr")
    para_rpr.append(make_element("w:rFonts", {
        "w:ascii": "Arial",
        "w:hAnsi": "Arial",
    }))
    para_rpr.append(make_element("w:sz", {"w:val": "28"}))
    style_para.append(para_rpr)
    styles_el.append(style_para)

    # Para 1: FontLinkedPara, unstyled run
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "FontLinkedPara"}))
    p1.add_run("Font linked basedOn loss")

    # Para 2: no para style, explicit FontLinkedChar on run
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.find(qn("w:pPr"))
    if p2_ppr is not None:
        ps = p2_ppr.find(qn("w:pStyle"))
        if ps is not None:
            p2_ppr.remove(ps)
    r2 = p2.add_run("Explicit font linked char")
    r2_rpr = r2._element.get_or_add_rPr()
    r2_rpr.insert(0, make_element("w:rStyle", {"w:val": "FontLinkedChar"}))

    save_fixture("linked-font-family-leak", doc, {
        "spec": "ISO 29500-1 §17.7.4.6 + §17.7.4.3",
        "description": (
            "FontLinkedPara (font=Arial, sz=28) linked to FontLinkedChar "
            "(bold=on, basedOn BaseFontChar(font=Courier New, italic=on)). "
            "Tests that char basedOn chain properties (italic) reach runs, "
            "while para's explicit font is preserved."
        ),
        "expected": {
            "para_1": (
                "font=Arial (para), sz=28 (para), bold=on (char raw), "
                "italic=on (char basedOn) — italic from BaseFontChar "
                "should not be lost by raw marks overlay"
            ),
            "para_2": (
                "font=Courier New (char basedOn), italic=on (char basedOn), "
                "bold=on (char raw)"
            ),
        },
    })


# =========================================================================
# Fixture 6: Para basedOn chain + linked overlay interaction
#
# ParaGrand -> ParaParent -> LinkedPara (linked to LinkedChar)
# Does resolve_chain for ParaGrand walk through LinkedPara and pick up
# LinkedPara's rPr from the basedOn chain? Yes, but LinkedPara's resolved
# para_styles entry includes the linked overlay. When ParaGrand resolves
# via resolve_chain, it only gets LinkedPara's RAW rPr (resolve_chain
# uses raw_styles, not para_styles). So the linked overlay should NOT
# leak into ParaGrand. Let's verify.
# =========================================================================

def make_para_chain_linked():
    """
    LinkCharP: character style, strike=on, linked to LinkParaP
    LinkParaP: paragraph style, bold=on, sz=24, linked to LinkCharP
    MidPara: paragraph style, basedOn LinkParaP, italic=on
    GrandPara: paragraph style, basedOn MidPara, color=FF0000

    Para 1 (LinkParaP, unstyled run):
      Expected: bold=on (para), sz=24 (para), strike=on (linked char)

    Para 2 (MidPara, unstyled run):
      Expected: bold=on (from LinkParaP via basedOn), sz=24, italic=on (own)
      NO strike — MidPara has no link, so linked overlay from LinkParaP
      should not leak.

    Para 3 (GrandPara, unstyled run):
      Expected: bold=on (from chain), sz=24 (from chain), italic=on (from MidPara),
                color=FF0000 (own). NO strike.
    """
    doc = Document()
    styles_el = doc.styles.element

    # LinkCharP: character style, strike=on
    style_char = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "LinkCharP",
    })
    style_char.append(make_element("w:name", {"w:val": "Link Char P"}))
    style_char.append(make_element("w:link", {"w:val": "LinkParaP"}))
    char_rpr = make_element("w:rPr")
    char_rpr.append(make_element("w:strike"))
    style_char.append(char_rpr)
    styles_el.append(style_char)

    # LinkParaP: paragraph style, bold=on, sz=24, linked to LinkCharP
    style_para = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "LinkParaP",
    })
    style_para.append(make_element("w:name", {"w:val": "Link Para P"}))
    style_para.append(make_element("w:link", {"w:val": "LinkCharP"}))
    para_rpr = make_element("w:rPr")
    para_rpr.append(make_element("w:b"))
    para_rpr.append(make_element("w:sz", {"w:val": "24"}))
    style_para.append(para_rpr)
    styles_el.append(style_para)

    # MidPara: basedOn LinkParaP, italic=on
    style_mid = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "MidPara",
    })
    style_mid.append(make_element("w:name", {"w:val": "Mid Para"}))
    style_mid.append(make_element("w:basedOn", {"w:val": "LinkParaP"}))
    mid_rpr = make_element("w:rPr")
    mid_rpr.append(make_element("w:i"))
    style_mid.append(mid_rpr)
    styles_el.append(style_mid)

    # GrandPara: basedOn MidPara, color=FF0000
    style_grand = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "GrandPara",
    })
    style_grand.append(make_element("w:name", {"w:val": "Grand Para"}))
    style_grand.append(make_element("w:basedOn", {"w:val": "MidPara"}))
    grand_rpr = make_element("w:rPr")
    grand_rpr.append(make_element("w:color", {"w:val": "FF0000"}))
    style_grand.append(grand_rpr)
    styles_el.append(style_grand)

    # Para 1: LinkParaP
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "LinkParaP"}))
    p1.add_run("Direct linked para")

    # Para 2: MidPara
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    p2_ppr.append(make_element("w:pStyle", {"w:val": "MidPara"}))
    p2.add_run("Child of linked para")

    # Para 3: GrandPara
    p3 = doc.add_paragraph()
    p3_ppr = p3._element.get_or_add_pPr()
    p3_ppr.append(make_element("w:pStyle", {"w:val": "GrandPara"}))
    p3.add_run("Grandchild of linked para")

    save_fixture("para-chain-linked", doc, {
        "spec": "ISO 29500-1 §17.7.4.6 + §17.7.4.3",
        "description": (
            "LinkParaP (bold=on, sz=24) linked to LinkCharP (strike=on). "
            "MidPara basedOn LinkParaP (italic=on). GrandPara basedOn MidPara "
            "(color=FF0000). Tests that linked overlay does NOT leak into "
            "child styles via basedOn chain."
        ),
        "expected": {
            "para_1": "bold=on, sz=24, strike=on (linked char overlay)",
            "para_2": "bold=on, sz=24, italic=on. NO strike (link doesn't leak)",
            "para_3": "bold=on, sz=24, italic=on, color=FF0000. NO strike",
        },
    })


# =========================================================================
# Fixture 7: Linked style where char style explicitly sets property to
#             same value as para — no information loss scenario
# =========================================================================

def make_linked_char_explicit_match():
    """
    MatchChar: character style, bold=on, sz=28 (matches para), color=0000FF
      linked to MatchPara
    MatchPara: paragraph style, bold=on, sz=28, italic=on
      linked to MatchChar

    Para 1 (MatchPara, unstyled run):
      Expected: bold=on, sz=28, italic=on (from para), color=0000FF (from char raw)
      The char explicitly sets sz=28 and bold=on, same as para. These overlay
      para's values (same result). Char also sets color=0000FF which para doesn't.
      Para's italic=on should be preserved (char doesn't set italic).
    """
    doc = Document()
    styles_el = doc.styles.element

    # MatchChar: character style, bold=on, sz=28, color=0000FF
    style_char = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "MatchChar",
    })
    style_char.append(make_element("w:name", {"w:val": "Match Char"}))
    style_char.append(make_element("w:link", {"w:val": "MatchPara"}))
    char_rpr = make_element("w:rPr")
    char_rpr.append(make_element("w:b"))
    char_rpr.append(make_element("w:sz", {"w:val": "28"}))
    char_rpr.append(make_element("w:color", {"w:val": "0000FF"}))
    style_char.append(char_rpr)
    styles_el.append(style_char)

    # MatchPara: paragraph style, bold=on, sz=28, italic=on
    style_para = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "MatchPara",
    })
    style_para.append(make_element("w:name", {"w:val": "Match Para"}))
    style_para.append(make_element("w:link", {"w:val": "MatchChar"}))
    para_rpr = make_element("w:rPr")
    para_rpr.append(make_element("w:b"))
    para_rpr.append(make_element("w:sz", {"w:val": "28"}))
    para_rpr.append(make_element("w:i"))
    style_para.append(para_rpr)
    styles_el.append(style_para)

    # Para 1: MatchPara, unstyled run
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "MatchPara"}))
    p1.add_run("Linked match test")

    save_fixture("linked-char-explicit-match", doc, {
        "spec": "ISO 29500-1 §17.7.4.6",
        "description": (
            "MatchPara (bold=on, sz=28, italic=on) linked to MatchChar "
            "(bold=on, sz=28, color=0000FF). Tests that when char explicitly "
            "sets same values as para, no information loss occurs. Char's "
            "additional color property should apply. Para's italic should "
            "be preserved."
        ),
        "expected": {
            "para_1": "bold=on, sz=28, italic=on (para), color=0000FF (char)",
        },
    })


# =========================================================================
# Main
# =========================================================================

if __name__ == "__main__":
    print("Generating style-cascade-linked-edge fixtures:")
    make_linked_based_on_loss()
    make_default_char_style()
    make_linked_based_on_target()
    make_direct_override_linked()
    make_linked_font_family_leak()
    make_para_chain_linked()
    make_linked_char_explicit_match()
    print("Done.")

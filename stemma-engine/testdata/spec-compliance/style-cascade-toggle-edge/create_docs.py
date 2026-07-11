# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for toggle property resolution edge cases.

Targets:
  1. Linked style basedOn chain property loss (§17.7.4.6 over-correction)
  2. 4-level basedOn toggle chain with inherit gap (§17.7.3)
  3. docDefaults toggle ON + deep basedOn chain with OFF (§17.7.3 + §2.1.230a)
  4. bold_cs toggle with cs/rtl active (§17.3.2.1 + §17.3.2.38)
  5. Default char style with toggle + explicit char style without it (§17.7.4.17)
  6. dstrike (double-strike) toggle resolution (§17.7.3)

Run:  python create_docs.py
"""

import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  style-cascade-toggle-edge/{name}/")


# =========================================================================
# 1. Linked style basedOn chain property loss (§17.7.4.6)
#
# When a paragraph style is linked to a char style that has a basedOn chain,
# the Bug 2 fix uses linked_raw.marks (not the resolved chain) to avoid
# overwriting para style properties with inherited ones. But this also
# loses properties the char style inherits from its basedOn chain that
# the para style doesn't set.
#
# Setup:
# - BaseChar: character style, italic=on, color=FF0000
# - AccentChar: character style, basedOn BaseChar, bold=on (no italic, no color)
# - AccentPara: paragraph style, linked to AccentChar, rPr has sz=32 (no italic/bold/color)
#
# Expected:
# - Run in AccentPara: bold=ON (from AccentChar), italic=ON (inherited from
#   BaseChar through AccentChar's basedOn), sz=32 (from AccentPara's rPr),
#   color=FF0000 (inherited from BaseChar through AccentChar's basedOn).
#
# Bug: The linked overlay uses AccentChar's raw marks (bold=On, italic=Inherit)
# so italic=On from BaseChar is lost.
# =========================================================================

def make_linked_based_on_loss():
    doc = Document()
    styles_el = doc.styles.element

    # BaseChar: italic=on, color=FF0000
    style_base = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "BaseChar",
    })
    style_base.append(make_element("w:name", {"w:val": "Base Char"}))
    base_rpr = make_element("w:rPr")
    base_rpr.append(make_element("w:i"))
    base_rpr.append(make_element("w:color", {"w:val": "FF0000"}))
    style_base.append(base_rpr)
    styles_el.append(style_base)

    # AccentChar: basedOn BaseChar, bold=on (no italic, no color)
    style_accent = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "AccentChar",
    })
    style_accent.append(make_element("w:name", {"w:val": "Accent Char"}))
    style_accent.append(make_element("w:basedOn", {"w:val": "BaseChar"}))
    style_accent.append(make_element("w:link", {"w:val": "AccentPara"}))
    accent_rpr = make_element("w:rPr")
    accent_rpr.append(make_element("w:b"))
    style_accent.append(accent_rpr)
    styles_el.append(style_accent)

    # AccentPara: linked to AccentChar, rPr has sz=32 only
    style_para = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "AccentPara",
    })
    style_para.append(make_element("w:name", {"w:val": "Accent Para"}))
    style_para.append(make_element("w:link", {"w:val": "AccentChar"}))
    para_rpr = make_element("w:rPr")
    para_rpr.append(make_element("w:sz", {"w:val": "32"}))
    style_para.append(para_rpr)
    styles_el.append(style_para)

    # Para 1: AccentPara style, no explicit char style
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "AccentPara"}))
    p1.add_run("Linked basedOn test")

    # Para 2: AccentPara + explicit AccentChar char style
    # When AccentChar is applied explicitly, its fully resolved chain (including
    # BaseChar's italic) should apply.
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    p2_ppr.append(make_element("w:pStyle", {"w:val": "AccentPara"}))
    r2 = p2.add_run("Explicit char style test")
    r2_rpr = r2._element.get_or_add_rPr()
    r2_rpr.insert(0, make_element("w:rStyle", {"w:val": "AccentChar"}))

    save_fixture("linked-basedOn-loss", doc, {
        "spec": "ISO 29500-1 §17.7.4.6",
        "description": "Linked char style basedOn chain properties lost in linked overlay",
        "expected": {
            "para_1_no_char": {
                "bold": True,
                "italic": True,
                "font_size": 32,
                "color": "FF0000",
            },
            "para_2_explicit_char": {
                "bold": True,
                "italic": True,
                "font_size": 32,
                "color": "FF0000",
            },
        },
    })


# =========================================================================
# 2. 4-level basedOn chain with toggle gap (§17.7.3)
#
# Root(bold=on) → A(bold=off) → B(bold not set/inherit) → Leaf(bold=on)
#
# Within a single hierarchy level (basedOn chain), child overrides parent.
# B doesn't set bold, so it inherits from A (bold=off).
# Leaf sets bold=on, overriding B's inherited off.
#
# Also test a property (caps) set at root, skipping all intermediate levels.
# =========================================================================

def make_4level_toggle_gap():
    doc = Document()
    styles_el = doc.styles.element

    # Root: bold=on, caps=on
    style_root = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "Root4",
    })
    style_root.append(make_element("w:name", {"w:val": "Root 4"}))
    root_rpr = make_element("w:rPr")
    root_rpr.append(make_element("w:b"))
    root_rpr.append(make_element("w:caps"))
    style_root.append(root_rpr)
    styles_el.append(style_root)

    # A: basedOn Root, bold=off (explicit), caps not set
    style_a = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "LevelA4",
    })
    style_a.append(make_element("w:name", {"w:val": "Level A"}))
    style_a.append(make_element("w:basedOn", {"w:val": "Root4"}))
    a_rpr = make_element("w:rPr")
    a_rpr.append(make_element("w:b", {"w:val": "0"}))
    style_a.append(a_rpr)
    styles_el.append(style_a)

    # B: basedOn A, bold not set (inherit gap), caps not set
    style_b = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "LevelB4",
    })
    style_b.append(make_element("w:name", {"w:val": "Level B"}))
    style_b.append(make_element("w:basedOn", {"w:val": "LevelA4"}))
    # No rPr at all — pure inherit
    styles_el.append(style_b)

    # Leaf: basedOn B, bold=on (overrides the inherited off from A)
    style_leaf = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "Leaf4",
    })
    style_leaf.append(make_element("w:name", {"w:val": "Leaf 4"}))
    style_leaf.append(make_element("w:basedOn", {"w:val": "LevelB4"}))
    leaf_rpr = make_element("w:rPr")
    leaf_rpr.append(make_element("w:b"))
    style_leaf.append(leaf_rpr)
    styles_el.append(style_leaf)

    # Para 1: Leaf style — bold=on (leaf), caps=on (from root)
    p1 = doc.add_paragraph()
    r1 = p1.add_run("Leaf 4-level bold on")
    r1_rpr = r1._element.get_or_add_rPr()
    r1_rpr.insert(0, make_element("w:rStyle", {"w:val": "Leaf4"}))

    # Para 2: LevelB style — bold=off (inherited from A), caps=on (from root)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("B 4-level bold inherit gap")
    r2_rpr = r2._element.get_or_add_rPr()
    r2_rpr.insert(0, make_element("w:rStyle", {"w:val": "LevelB4"}))

    # Para 3: LevelA style — bold=off (explicit), caps=on (from root)
    p3 = doc.add_paragraph()
    r3 = p3.add_run("A 4-level bold off")
    r3_rpr = r3._element.get_or_add_rPr()
    r3_rpr.insert(0, make_element("w:rStyle", {"w:val": "LevelA4"}))

    save_fixture("4level-toggle-gap", doc, {
        "spec": "ISO 29500-1 §17.7.3 (first-value-encountered in 4-level chain)",
        "description": "4-level basedOn chain with toggle inherit gap",
        "expected": {
            "para_1_leaf": {"bold": True, "caps": True},
            "para_2_B": {"bold": False, "caps": True},
            "para_3_A": {"bold": False, "caps": True},
        },
    })


# =========================================================================
# 3. docDefaults caps=on + deep basedOn chain where leaf sets caps=off
#    (§17.7.3 + MS-OI29500 §2.1.230a)
#
# This combines the docDefaults override (Bug 1 scenario) with a deep
# basedOn chain. Leaf explicitly sets caps=off, which should override
# docDefaults caps=on.
# =========================================================================

def make_doc_defaults_deep_chain_toggle():
    doc = Document()
    styles_el = doc.styles.element

    # Set docDefaults: caps=on, strike=on
    doc_defaults = styles_el.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = make_element("w:docDefaults")
        styles_el.insert(0, doc_defaults)
    rpr_default_el = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default_el is None:
        rpr_default_el = make_element("w:rPrDefault")
        doc_defaults.insert(0, rpr_default_el)
    rpr = rpr_default_el.find(qn("w:rPr"))
    if rpr is None:
        rpr = make_element("w:rPr")
        rpr_default_el.append(rpr)
    # Clean existing
    for tag in ["caps", "strike"]:
        for existing in rpr.findall(qn(f"w:{tag}")):
            rpr.remove(existing)
    rpr.append(make_element("w:caps"))
    rpr.append(make_element("w:strike"))

    # ParaBase: paragraph style, no toggle overrides
    style_base = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "ParaBase",
    })
    style_base.append(make_element("w:name", {"w:val": "Para Base"}))
    styles_el.append(style_base)

    # ParaMid: basedOn ParaBase, caps=off
    style_mid = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "ParaMid",
    })
    style_mid.append(make_element("w:name", {"w:val": "Para Mid"}))
    style_mid.append(make_element("w:basedOn", {"w:val": "ParaBase"}))
    mid_rpr = make_element("w:rPr")
    mid_rpr.append(make_element("w:caps", {"w:val": "0"}))
    style_mid.append(mid_rpr)
    styles_el.append(style_mid)

    # ParaLeaf: basedOn ParaMid, strike=off
    style_leaf = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "ParaLeaf",
    })
    style_leaf.append(make_element("w:name", {"w:val": "Para Leaf"}))
    style_leaf.append(make_element("w:basedOn", {"w:val": "ParaMid"}))
    leaf_rpr = make_element("w:rPr")
    leaf_rpr.append(make_element("w:strike", {"w:val": "0"}))
    style_leaf.append(leaf_rpr)
    styles_el.append(style_leaf)

    # Para 1: ParaLeaf — caps=off (from ParaMid), strike=off (from ParaLeaf)
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "ParaLeaf"}))
    p1.add_run("Deep chain overrides docDefaults")

    # Para 2: ParaMid — caps=off, strike=on (from docDefaults, not overridden)
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    p2_ppr.append(make_element("w:pStyle", {"w:val": "ParaMid"}))
    p2.add_run("Mid chain caps off strike on")

    # Para 3: ParaBase — caps=on (docDefaults), strike=on (docDefaults)
    p3 = doc.add_paragraph()
    p3_ppr = p3._element.get_or_add_pPr()
    p3_ppr.append(make_element("w:pStyle", {"w:val": "ParaBase"}))
    p3.add_run("Base chain both from docDefaults")

    # Para 4: Unstyled — caps=on (docDefaults), strike=on (docDefaults)
    p4 = doc.add_paragraph()
    # Remove any auto pStyle
    p4_ppr = p4._element.find(qn("w:pPr"))
    if p4_ppr is not None:
        ps = p4_ppr.find(qn("w:pStyle"))
        if ps is not None:
            p4_ppr.remove(ps)
    p4.add_run("Unstyled caps and strike from docDefaults")

    save_fixture("doc-defaults-deep-toggle", doc, {
        "spec": "ISO 29500-1 §17.7.3 + MS-OI29500 §2.1.230a",
        "description": "docDefaults caps/strike ON + deep basedOn chain turning them OFF",
        "expected": {
            "para_1_leaf": {"caps": False, "strike": False},
            "para_2_mid": {"caps": False, "strike": True},
            "para_3_base": {"caps": True, "strike": True},
            "para_4_unstyled": {"caps": True, "strike": True},
        },
    })


# =========================================================================
# 4. bold_cs toggle with cs/rtl active (§17.3.2.1 + §17.3.2.38)
#
# When cs=on or rtl=on, bold is replaced by bCs and italic by iCs.
# If bCs is absent (Inherit), bold should become effectively OFF.
# This tests that the cs/rtl replacement works through the style cascade.
# =========================================================================

def make_bold_cs_toggle():
    doc = Document()
    styles_el = doc.styles.element

    # CsStyle: paragraph style, bold=on (Latin), bCs absent (no complex script bold)
    # cs=on forces use of bCs instead of b
    style_cs_nobcs = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "CsNoBcs",
    })
    style_cs_nobcs.append(make_element("w:name", {"w:val": "CS No BCs"}))
    cs_rpr = make_element("w:rPr")
    cs_rpr.append(make_element("w:b"))       # bold=on for Latin
    # bCs deliberately absent — no bold for complex script
    cs_rpr.append(make_element("w:cs"))       # cs=on
    style_cs_nobcs.append(cs_rpr)
    styles_el.append(style_cs_nobcs)

    # CsWithBcs: paragraph style, bold=off (Latin), bCs=on (complex script bold)
    style_cs_bcs = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "CsWithBcs",
    })
    style_cs_bcs.append(make_element("w:name", {"w:val": "CS With BCs"}))
    bcs_rpr = make_element("w:rPr")
    bcs_rpr.append(make_element("w:b", {"w:val": "0"}))   # bold=off for Latin
    bcs_rpr.append(make_element("w:bCs"))                   # bCs=on
    bcs_rpr.append(make_element("w:cs"))                    # cs=on
    style_cs_bcs.append(bcs_rpr)
    styles_el.append(style_cs_bcs)

    # RtlBold: paragraph style, bold=on (Latin), bCs=on, rtl=on
    style_rtl = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "RtlBold",
    })
    style_rtl.append(make_element("w:name", {"w:val": "RTL Bold"}))
    rtl_rpr = make_element("w:rPr")
    rtl_rpr.append(make_element("w:b"))      # bold=on for Latin
    rtl_rpr.append(make_element("w:bCs"))    # bCs=on for complex script
    rtl_rpr.append(make_element("w:rtl"))    # rtl=on
    style_rtl.append(rtl_rpr)
    styles_el.append(style_rtl)

    # Para 1: CsNoBcs — cs=on, bold=on (Latin) but bCs absent
    # With cs=on, bold should be replaced by bCs. bCs is Inherit → effectively not bold.
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "CsNoBcs"}))
    p1.add_run("CS active bold but no bCs")

    # Para 2: CsWithBcs — cs=on, bold=off but bCs=on
    # With cs=on, bold is replaced by bCs=on → should be bold.
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    p2_ppr.append(make_element("w:pStyle", {"w:val": "CsWithBcs"}))
    p2.add_run("CS active bCs on overrides bold off")

    # Para 3: RtlBold — rtl=on, both bold=on and bCs=on
    # With rtl=on, bold is replaced by bCs=on → bold.
    p3 = doc.add_paragraph()
    p3_ppr = p3._element.get_or_add_pPr()
    p3_ppr.append(make_element("w:pStyle", {"w:val": "RtlBold"}))
    p3.add_run("RTL bold with bCs")

    save_fixture("bold-cs-toggle", doc, {
        "spec": "ECMA-376 §17.3.2.1 + §17.3.2.38",
        "description": "bold_cs toggle replacement when cs/rtl active",
        "expected": {
            "para_1": "bold=OFF (cs=on, bCs absent → not bold)",
            "para_2": "bold=ON (cs=on, bCs=on → bold)",
            "para_3": "bold=ON (rtl=on, bCs=on → bold)",
        },
    })


# =========================================================================
# 5. Default char style with toggle property + explicit char style
#    (§17.7.4.17 + §17.7.2)
#
# A custom default char style (w:default="1") has italic=on.
# An explicit char style has bold=on but not italic.
# When both apply, the merge should give: bold=on (explicit), italic=on (default).
# =========================================================================

def make_default_char_style_toggle():
    doc = Document()
    styles_el = doc.styles.element

    # Remove existing DefaultParagraphFont if present
    for style in styles_el.findall(qn("w:style")):
        if style.get(qn("w:styleId")) == "DefaultParagraphFont":
            styles_el.remove(style)

    # CustomDefault: character style, w:default="1", italic=on, smallCaps=on
    # Using a non-standard styleId so it's not treated as special-ignored.
    style_default = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "CustomDefaultChar",
        "w:default": "1",
    })
    style_default.append(make_element("w:name", {"w:val": "Custom Default Char"}))
    default_rpr = make_element("w:rPr")
    default_rpr.append(make_element("w:i"))          # italic=on
    default_rpr.append(make_element("w:smallCaps"))   # smallCaps=on
    style_default.append(default_rpr)
    styles_el.append(style_default)

    # ExplicitChar: character style, bold=on (no italic, no smallCaps)
    style_explicit = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "ExplicitChar",
    })
    style_explicit.append(make_element("w:name", {"w:val": "Explicit Char"}))
    explicit_rpr = make_element("w:rPr")
    explicit_rpr.append(make_element("w:b"))
    style_explicit.append(explicit_rpr)
    styles_el.append(style_explicit)

    # Para 1: explicit ExplicitChar + default CustomDefaultChar merge
    # Expected: bold=on (explicit), italic=on (default), smallCaps=on (default)
    p1 = doc.add_paragraph()
    r1 = p1.add_run("Explicit plus default char")
    r1_rpr = r1._element.get_or_add_rPr()
    r1_rpr.insert(0, make_element("w:rStyle", {"w:val": "ExplicitChar"}))

    # Para 2: no explicit char style, only default CustomDefaultChar
    # Expected: italic=on, smallCaps=on (from default char style)
    p2 = doc.add_paragraph()
    p2.add_run("Default char style only")

    # Para 3: explicit ExplicitChar that sets italic=off + default char italic=on
    # Explicit italic=off should override default char italic=on.
    style_no_italic = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "NoItalicChar",
    })
    style_no_italic.append(make_element("w:name", {"w:val": "No Italic Char"}))
    ni_rpr = make_element("w:rPr")
    ni_rpr.append(make_element("w:i", {"w:val": "0"}))
    ni_rpr.append(make_element("w:b"))
    style_no_italic.append(ni_rpr)
    styles_el.append(style_no_italic)

    p3 = doc.add_paragraph()
    r3 = p3.add_run("Explicit italic off overrides default")
    r3_rpr = r3._element.get_or_add_rPr()
    r3_rpr.insert(0, make_element("w:rStyle", {"w:val": "NoItalicChar"}))

    save_fixture("default-char-style-toggle", doc, {
        "spec": "ISO 29500-1 §17.7.4.17 + §17.7.2",
        "description": "Default char style toggle merge with explicit char style",
        "expected": {
            "para_1": {"bold": True, "italic": True, "smallCaps": True},
            "para_2": {"bold": False, "italic": True, "smallCaps": True},
            "para_3": {"bold": True, "italic": False, "smallCaps": True},
        },
    })


# =========================================================================
# 6. dstrike (double-strike) toggle resolution (§17.7.3)
#
# Per ECMA-376, dstrike is listed as a toggle property. Our code resolves
# it with resolve_mark (non-toggle) instead of resolve_toggle_mark.
# While the two functions are currently identical, this is architecturally
# incorrect.
#
# This test verifies that dstrike cascades correctly through the style
# hierarchy: para style on + char style on = on (per reset semantics).
# =========================================================================

def make_dstrike_toggle():
    doc = Document()
    styles_el = doc.styles.element

    # DstrikePara: paragraph style, dstrike=on
    style_para = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "DstrikePara",
    })
    style_para.append(make_element("w:name", {"w:val": "Dstrike Para"}))
    para_rpr = make_element("w:rPr")
    para_rpr.append(make_element("w:dstrike"))
    style_para.append(para_rpr)
    styles_el.append(style_para)

    # DstrikeChar: character style, dstrike=on
    style_char = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "DstrikeChar",
    })
    style_char.append(make_element("w:name", {"w:val": "Dstrike Char"}))
    char_rpr = make_element("w:rPr")
    char_rpr.append(make_element("w:dstrike"))
    style_char.append(char_rpr)
    styles_el.append(style_char)

    # Para 1: DstrikePara + DstrikeChar → dstrike=on (both set it)
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "DstrikePara"}))
    r1 = p1.add_run("Both para and char dstrike")
    r1_rpr = r1._element.get_or_add_rPr()
    r1_rpr.insert(0, make_element("w:rStyle", {"w:val": "DstrikeChar"}))

    # Para 2: DstrikePara only → dstrike=on
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    p2_ppr.append(make_element("w:pStyle", {"w:val": "DstrikePara"}))
    p2.add_run("Para dstrike only")

    # Para 3: DstrikePara + direct dstrike=off → dstrike=off
    p3 = doc.add_paragraph()
    p3_ppr = p3._element.get_or_add_pPr()
    p3_ppr.append(make_element("w:pStyle", {"w:val": "DstrikePara"}))
    r3 = p3.add_run("Direct dstrike off")
    r3_rpr = r3._element.get_or_add_rPr()
    r3_rpr.append(make_element("w:dstrike", {"w:val": "0"}))

    save_fixture("dstrike-toggle", doc, {
        "spec": "ISO 29500-1 §17.7.3 (dstrike as toggle)",
        "description": "dstrike toggle resolution across hierarchy levels",
        "expected": {
            "para_1": "dstrike=ON (both para and char set it)",
            "para_2": "dstrike=ON (para only)",
            "para_3": "dstrike=OFF (direct off overrides para on)",
        },
    })


# =========================================================================
# 7. Vanish + emboss toggle through basedOn chain with docDefaults
#    (§17.7.3 less-tested toggle properties)
#
# Tests less-commonly-used toggle properties (vanish, emboss, imprint,
# outline, shadow) to verify they go through the same resolution path
# as bold/italic.
# =========================================================================

def make_rare_toggle_cascade():
    doc = Document()
    styles_el = doc.styles.element

    # docDefaults: emboss=on
    doc_defaults = styles_el.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = make_element("w:docDefaults")
        styles_el.insert(0, doc_defaults)
    rpr_default_el = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default_el is None:
        rpr_default_el = make_element("w:rPrDefault")
        doc_defaults.insert(0, rpr_default_el)
    rpr = rpr_default_el.find(qn("w:rPr"))
    if rpr is None:
        rpr = make_element("w:rPr")
        rpr_default_el.append(rpr)
    for tag in ["emboss"]:
        for existing in rpr.findall(qn(f"w:{tag}")):
            rpr.remove(existing)
    rpr.append(make_element("w:emboss"))

    # VanishStyle: paragraph style with vanish=on, emboss=off
    style_vanish = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "VanishStyle",
    })
    style_vanish.append(make_element("w:name", {"w:val": "Vanish Style"}))
    vanish_rpr = make_element("w:rPr")
    vanish_rpr.append(make_element("w:vanish"))
    vanish_rpr.append(make_element("w:emboss", {"w:val": "0"}))
    style_vanish.append(vanish_rpr)
    styles_el.append(style_vanish)

    # ShadowChar: character style with shadow=on, outline=on
    style_shadow = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "ShadowChar",
    })
    style_shadow.append(make_element("w:name", {"w:val": "Shadow Char"}))
    shadow_rpr = make_element("w:rPr")
    shadow_rpr.append(make_element("w:shadow"))
    shadow_rpr.append(make_element("w:outline"))
    style_shadow.append(shadow_rpr)
    styles_el.append(style_shadow)

    # Para 1: VanishStyle + ShadowChar
    # vanish=on (para), emboss=off (para overrides docDefaults), shadow=on (char), outline=on (char)
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    p1_ppr.append(make_element("w:pStyle", {"w:val": "VanishStyle"}))
    r1 = p1.add_run("Vanish shadow outline test")
    r1_rpr = r1._element.get_or_add_rPr()
    r1_rpr.insert(0, make_element("w:rStyle", {"w:val": "ShadowChar"}))

    # Para 2: VanishStyle only
    # vanish=on, emboss=off (overrides docDefaults)
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    p2_ppr.append(make_element("w:pStyle", {"w:val": "VanishStyle"}))
    p2.add_run("Vanish emboss off test")

    # Para 3: No style — emboss=on from docDefaults
    p3 = doc.add_paragraph()
    p3_ppr = p3._element.find(qn("w:pPr"))
    if p3_ppr is not None:
        ps = p3_ppr.find(qn("w:pStyle"))
        if ps is not None:
            p3_ppr.remove(ps)
    p3.add_run("Unstyled emboss from docDefaults")

    save_fixture("rare-toggle-cascade", doc, {
        "spec": "ISO 29500-1 §17.7.3 (rare toggle properties)",
        "description": "Vanish, emboss, shadow, outline toggle cascade",
        "expected": {
            "para_1": {
                "vanish": True,
                "emboss": False,
                "shadow": True,
                "outline": True,
            },
            "para_2": {"vanish": True, "emboss": False},
            "para_3": {"emboss": True},
        },
    })


# =========================================================================
# 8. Toggle val="false" explicit OFF in basedOn chain
#    (§17.7.3 — Off value propagation)
#
# Root: bold=on. Mid: bold=false (explicit off). Leaf: no bold.
# Leaf should inherit bold=off from Mid (not bold=on from Root).
# This is subtly different from val="0" — both should produce Off.
# =========================================================================

def make_val_false_propagation():
    doc = Document()
    styles_el = doc.styles.element

    # Root: bold=on using val="true", italic=on using val="1"
    style_root = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "ValRoot",
    })
    style_root.append(make_element("w:name", {"w:val": "Val Root"}))
    root_rpr = make_element("w:rPr")
    root_rpr.append(make_element("w:b", {"w:val": "true"}))
    root_rpr.append(make_element("w:i", {"w:val": "1"}))
    root_rpr.append(make_element("w:caps"))   # bare element, no val
    style_root.append(root_rpr)
    styles_el.append(style_root)

    # Mid: bold=false (using "false"), italic=off (using "off")
    style_mid = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "ValMid",
    })
    style_mid.append(make_element("w:name", {"w:val": "Val Mid"}))
    style_mid.append(make_element("w:basedOn", {"w:val": "ValRoot"}))
    mid_rpr = make_element("w:rPr")
    mid_rpr.append(make_element("w:b", {"w:val": "false"}))
    mid_rpr.append(make_element("w:i", {"w:val": "off"}))
    style_mid.append(mid_rpr)
    styles_el.append(style_mid)

    # Leaf: no bold/italic (inherits from Mid)
    style_leaf = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "ValLeaf",
    })
    style_leaf.append(make_element("w:name", {"w:val": "Val Leaf"}))
    style_leaf.append(make_element("w:basedOn", {"w:val": "ValMid"}))
    # No rPr — inherits everything
    styles_el.append(style_leaf)

    # Para 1: ValRoot — bold=on (val="true"), italic=on (val="1"), caps=on
    p1 = doc.add_paragraph()
    r1 = p1.add_run("Root val true and 1")
    r1_rpr = r1._element.get_or_add_rPr()
    r1_rpr.insert(0, make_element("w:rStyle", {"w:val": "ValRoot"}))

    # Para 2: ValMid — bold=off (val="false"), italic=off (val="off"), caps=on (from root)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Mid val false and off")
    r2_rpr = r2._element.get_or_add_rPr()
    r2_rpr.insert(0, make_element("w:rStyle", {"w:val": "ValMid"}))

    # Para 3: ValLeaf — inherits Mid: bold=off, italic=off, caps=on
    p3 = doc.add_paragraph()
    r3 = p3.add_run("Leaf inherits false and off")
    r3_rpr = r3._element.get_or_add_rPr()
    r3_rpr.insert(0, make_element("w:rStyle", {"w:val": "ValLeaf"}))

    save_fixture("val-false-propagation", doc, {
        "spec": "ISO 29500-1 §22.9.2.7 (ST_OnOff) + §17.7.3",
        "description": "Toggle value variants: true/1/absent=on, false/off/0=off",
        "expected": {
            "para_1_root": {"bold": True, "italic": True, "caps": True},
            "para_2_mid": {"bold": False, "italic": False, "caps": True},
            "para_3_leaf": {"bold": False, "italic": False, "caps": True},
        },
    })


# =========================================================================
# Main
# =========================================================================

if __name__ == "__main__":
    print("Generating style-cascade-toggle-edge fixtures:")
    make_linked_based_on_loss()
    make_4level_toggle_gap()
    make_doc_defaults_deep_chain_toggle()
    make_bold_cs_toggle()
    make_default_char_style_toggle()
    make_dstrike_toggle()
    make_rare_toggle_cascade()
    make_val_false_propagation()
    print("Done.")

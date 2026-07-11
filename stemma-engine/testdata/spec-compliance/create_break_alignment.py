# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for spec-compliance testing of break types (§17.3.3.1)
and paragraph alignment / ST_Jc completeness (§17.18.44).

Fixtures:
  - break-types/break-types/     — w:br with type and clear attributes
  - break-types/alignment-distribute/ — all ST_Jc values including distribute

Run:  uv run create_break_alignment.py
"""

import json
from pathlib import Path
from lxml import etree

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    """Create an OxmlElement with optional attributes."""
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(
    area: str,
    name: str,
    doc: DocxDocument,
    metadata: dict,
    filename: str = "input.docx",
) -> None:
    """Save a single-doc fixture (for parsing/model tests)."""
    out = ROOT / area / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / filename))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  {area}/{name}/")


# =========================================================================
# BREAK TYPES (ISO 29500-1 §17.3.3.1)
# =========================================================================

def make_break_types() -> None:
    """w:br with type attribute variations.

    ISO 29500-1 §17.3.3.1: The break element can have type="page",
    type="column", type="textWrapping" (default), and clear attribute.
    """
    doc = Document()

    # Paragraph 1: page break within a run
    p1 = doc.add_paragraph()
    r1a = p1.add_run("Before page break")
    # Inject <w:br w:type="page"/> into the run
    br_page = make_element("w:br", {"w:type": "page"})
    r1a._r.append(br_page)
    r1b = p1.add_run("After page break")

    # Paragraph 2: column break
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Before column break")
    br_col = make_element("w:br", {"w:type": "column"})
    r2._r.append(br_col)

    # Paragraph 3: plain line break (no type attribute = textWrapping default)
    p3 = doc.add_paragraph()
    r3 = p3.add_run("Line break")
    br_line = make_element("w:br")
    r3._r.append(br_line)

    # Paragraph 4: textWrapping break with clear="all"
    p4 = doc.add_paragraph()
    r4 = p4.add_run("Clear break")
    br_clear = make_element("w:br", {"w:type": "textWrapping", "w:clear": "all"})
    r4._r.append(br_clear)

    save_fixture("break-types", "break-types", doc, {
        "name": "break-types",
        "spec_ref": "ISO 29500-1 §17.3.3.1",
        "description": "Paragraphs with page, column, line, and clear break types",
        "expected_behavior": "Parser should distinguish page/column breaks from line breaks via HardBreakNode",
        "current_status": "PARTIAL — HardBreakNode has no break_type field; all breaks are identical",
    })


# =========================================================================
# ALIGNMENT / ST_Jc (ISO 29500-1 §17.18.44)
# =========================================================================

def make_alignment_distribute() -> None:
    """All ST_Jc values including distribute and thaiDistribute.

    ISO 29500-1 §17.18.44 defines: start, center, end, both, distribute,
    mediumKashida, highKashida, lowKashida, thaiDistribute, left, right.
    """
    doc = Document()

    alignments = [
        ("Left aligned", "left"),
        ("Center aligned", "center"),
        ("Right aligned", "right"),
        ("Justified", "both"),
        ("Distributed", "distribute"),
        ("Start aligned", "start"),
        ("End aligned", "end"),
        ("Thai distribute", "thaiDistribute"),
    ]

    for text, jc_val in alignments:
        p = doc.add_paragraph(text)
        # Inject w:jc directly into pPr
        pPr = p._p.get_or_add_pPr()
        # Remove any existing jc element
        for existing_jc in pPr.findall(qn("w:jc")):
            pPr.remove(existing_jc)
        jc = make_element("w:jc", {"w:val": jc_val})
        pPr.append(jc)

    save_fixture("break-types", "alignment-distribute", doc, {
        "name": "alignment-distribute",
        "spec_ref": "ISO 29500-1 §17.18.44",
        "description": "Paragraphs with all ST_Jc alignment values including distribute and thaiDistribute",
        "expected_behavior": "Parser should preserve distribute/thaiDistribute as distinct Alignment variants",
        "current_status": "PARTIAL — Alignment enum lacks Distribute/ThaiDistribute; these silently map to Left",
    })


def main() -> None:
    print("Generating break & alignment spec-compliance fixtures:")
    make_break_types()
    make_alignment_distribute()
    print("Done.")


if __name__ == "__main__":
    main()

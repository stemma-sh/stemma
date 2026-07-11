# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate before/after DOCX pairs for tracked change serialization tests.

Each subdirectory gets a before.docx and after.docx pair that, when
diffed-and-redlined, exercises specific tracked change serialization rules.
"""

import json
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from lxml import etree

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
COMMENTS_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def save_pair(
    name: str,
    before: DocxDocument,
    after: DocxDocument,
    metadata: dict,
) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    before.save(str(out / "before.docx"))
    after.save(str(out / "after.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  {name}/")


def make_element(tag: str, attrs: dict | None = None) -> etree._Element:
    """Create an element with w: namespace."""
    el = etree.SubElement(
        etree.Element("dummy"), qn(tag)
    )
    if attrs:
        for k, v in attrs.items():
            el.set(qn(k), v)
    return el


def _build_comments_xml(comments: list[dict]) -> bytes:
    """Build word/comments.xml content from a list of comment dicts."""
    root = etree.Element(
        f"{{{W}}}comments",
        nsmap={"w": W, "r": R},
    )

    for c in comments:
        attrs = {
            f"{{{W}}}id": c["id"],
            f"{{{W}}}author": c["author"],
            f"{{{W}}}date": c["date"],
            f"{{{W}}}initials": c.get("initials", c["author"][:2].upper()),
        }
        comment_el = etree.SubElement(root, f"{{{W}}}comment", attrib=attrs)

        p = etree.SubElement(comment_el, f"{{{W}}}p")
        r = etree.SubElement(p, f"{{{W}}}r")
        t = etree.SubElement(r, f"{{{W}}}t")
        t.text = c["text"]

    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _inject_comment_range(p_element, comment_id: str,
                          text_before: str, commented_text: str,
                          text_after: str) -> None:
    """Inject comment range markers into a paragraph element."""
    # Clear existing runs
    for child in list(p_element):
        p_element.remove(child)

    # Text before
    if text_before:
        run = etree.SubElement(p_element, qn("w:r"))
        t = etree.SubElement(run, qn("w:t"))
        t.text = text_before
        t.set(qn("xml:space"), "preserve")

    # commentRangeStart
    range_start = etree.SubElement(p_element, qn("w:commentRangeStart"))
    range_start.set(qn("w:id"), comment_id)

    # Commented text
    run = etree.SubElement(p_element, qn("w:r"))
    t = etree.SubElement(run, qn("w:t"))
    t.text = commented_text
    t.set(qn("xml:space"), "preserve")

    # commentRangeEnd
    range_end = etree.SubElement(p_element, qn("w:commentRangeEnd"))
    range_end.set(qn("w:id"), comment_id)

    # commentReference
    ref_run = etree.SubElement(p_element, qn("w:r"))
    ref = etree.SubElement(ref_run, qn("w:commentReference"))
    ref.set(qn("w:id"), comment_id)

    # Text after
    if text_after:
        run = etree.SubElement(p_element, qn("w:r"))
        t = etree.SubElement(run, qn("w:t"))
        t.text = text_after
        t.set(qn("xml:space"), "preserve")


def _add_comments_part(doc: DocxDocument, comments_xml: bytes):
    """Add word/comments.xml to the docx package and wire up the relationship."""
    import zipfile
    import io

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    out_buf = io.BytesIO()
    rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"

    with zipfile.ZipFile(buf, "r") as zin, zipfile.ZipFile(out_buf, "w") as zout:
        for item in zin.namelist():
            data = zin.read(item)

            if item == "word/_rels/document.xml.rels":
                rels_root = etree.fromstring(data)
                max_id = 0
                for rel in rels_root:
                    rid = rel.get("Id", "")
                    if rid.startswith("rId"):
                        try:
                            max_id = max(max_id, int(rid[3:]))
                        except ValueError:
                            pass
                new_rid = f"rId{max_id + 1}"
                rel_el = etree.SubElement(rels_root, f"{{{rels_ns}}}Relationship")
                rel_el.set("Id", new_rid)
                rel_el.set("Type", COMMENTS_REL_TYPE)
                rel_el.set("Target", "comments.xml")
                data = etree.tostring(
                    rels_root, xml_declaration=True, encoding="UTF-8", standalone=True
                )

            zout.writestr(item, data)

        zout.writestr("word/comments.xml", comments_xml)

    out_buf.seek(0)
    return out_buf


# ---------------------------------------------------------------------------
# 1. del-uses-deltext — Delete text to test w:delText usage
# ---------------------------------------------------------------------------

def make_del_uses_deltext() -> None:
    """Old: "Hello World", New: "Hello" — deletion of " World".

    ECMA-376 §17.3.3.7: Text inside w:del must use w:delText, not w:t.
    """
    before = Document()
    before.add_paragraph("Hello World")

    after = Document()
    after.add_paragraph("Hello")

    save_pair("del-uses-deltext", before, after, {
        "name": "del-uses-deltext",
        "spec_ref": "ECMA-376 §17.3.3.7",
        "description": "Delete ' World' from 'Hello World' — verifies w:delText in w:del",
    })


# ---------------------------------------------------------------------------
# 2. no-nested-ins-del — Multiple edits across the document
# ---------------------------------------------------------------------------

def make_no_nested_ins_del() -> None:
    """Multiple text changes to verify no nested w:ins or w:del containers.

    MS-OI29500 §2.1.330/§2.1.334: Word does not support nested w:ins inside
    w:ins or w:del inside w:del.
    """
    before = Document()
    before.add_paragraph("First line of text to modify.")
    before.add_paragraph("Second line stays the same.")
    before.add_paragraph("Third line with old content.")
    before.add_paragraph("Fourth line to be removed.")

    after = Document()
    after.add_paragraph("First line of updated text.")
    after.add_paragraph("Second line stays the same.")
    after.add_paragraph("Third line with new content.")
    # Fourth line removed

    save_pair("no-nested-ins-del", before, after, {
        "name": "no-nested-ins-del",
        "spec_ref": "MS-OI29500 §2.1.330/§2.1.334",
        "description": "Multiple edits to verify no nested w:ins/w:del containers",
    })


# ---------------------------------------------------------------------------
# 3. no-revisions-in-comments — Comments must not contain tracked changes
# ---------------------------------------------------------------------------

def make_no_revisions_in_comments() -> None:
    """Paragraph with a comment, text around the comment is changed.

    MS-OI29500 §2.1.313: Word does not support revisions inside comments.
    """
    # before: paragraph with a comment
    before_doc = Document()
    p = before_doc.add_paragraph("Text before comment. Commented text here. Text after comment.")
    _inject_comment_range(
        p._p,
        comment_id="1",
        text_before="Text before comment. ",
        commented_text="Commented text here.",
        text_after=" Text after comment.",
    )

    comments_xml = _build_comments_xml([{
        "id": "1",
        "author": "Test Author",
        "date": "2025-01-01T00:00:00Z",
        "text": "This is a test comment.",
    }])

    before_buf = _add_comments_part(before_doc, comments_xml)

    # after: same comment, but paragraph text changed
    after_doc = Document()
    p = after_doc.add_paragraph("Modified before comment. Commented text here. Modified after comment.")
    _inject_comment_range(
        p._p,
        comment_id="1",
        text_before="Modified before comment. ",
        commented_text="Commented text here.",
        text_after=" Modified after comment.",
    )
    after_buf = _add_comments_part(after_doc, comments_xml)

    out = ROOT / "no-revisions-in-comments"
    out.mkdir(parents=True, exist_ok=True)
    (out / "before.docx").write_bytes(before_buf.getvalue())
    (out / "after.docx").write_bytes(after_buf.getvalue())
    (out / "metadata.json").write_text(json.dumps({
        "name": "no-revisions-in-comments",
        "spec_ref": "MS-OI29500 §2.1.313",
        "description": "Comment with text changes around it — comments must not contain tracked changes",
    }, indent=2) + "\n")
    print("  no-revisions-in-comments/")


# ---------------------------------------------------------------------------
# 4. flat-tracked-changes — Multiple edits for flat structure check
# ---------------------------------------------------------------------------

def make_flat_tracked_changes() -> None:
    """Complex multi-edit scenario for flat tracked changes verification.

    Cross-cutting rule: All run-level tracked change containers must be flat.
    """
    before = Document()
    before.add_paragraph("Alpha bravo charlie delta echo foxtrot.")
    before.add_paragraph("Golf hotel india juliet kilo lima.")
    before.add_paragraph("Mike november oscar papa quebec romeo.")

    after = Document()
    after.add_paragraph("Alpha bravo CHANGED delta echo foxtrot.")
    after.add_paragraph("Golf hotel india juliet kilo lima.")
    after.add_paragraph("Mike november MODIFIED papa quebec UPDATED.")

    save_pair("flat-tracked-changes", before, after, {
        "name": "flat-tracked-changes",
        "spec_ref": "Cross-cutting tracked change flat structure rule",
        "description": "Multiple edits across document — no nested ins/del of same type",
    })


# ---------------------------------------------------------------------------
# 5. move-ranges-paired — Paragraph move (A, B, C) -> (A, C, B)
# ---------------------------------------------------------------------------

def make_move_ranges_paired() -> None:
    """Move paragraph B after C: (A, B, C) -> (A, C, B).

    ECMA-376 §17.13.5.23-28: Move range markers must be paired.
    """
    before = Document()
    before.add_paragraph("Paragraph A stays first.")
    before.add_paragraph("Paragraph B will be moved.")
    before.add_paragraph("Paragraph C stays third.")

    after = Document()
    after.add_paragraph("Paragraph A stays first.")
    after.add_paragraph("Paragraph C stays third.")
    after.add_paragraph("Paragraph B will be moved.")

    save_pair("move-ranges-paired", before, after, {
        "name": "move-ranges-paired",
        "spec_ref": "ECMA-376 §17.13.5.23-28",
        "description": "Paragraph move to verify moveFrom/moveTo range pairing",
    })


# ---------------------------------------------------------------------------
# 6. revision-ids-unique — General uniqueness check
# ---------------------------------------------------------------------------

def make_revision_ids_unique() -> None:
    """Multiple text changes to produce many annotation IDs.

    ECMA-376 §17.13.5: Every annotation id must be unique.
    """
    before = Document()
    before.add_paragraph("Line one with original words.")
    before.add_paragraph("Line two to be removed.")
    before.add_paragraph("Line three with old text.")
    before.add_paragraph("Line four unchanged.")
    before.add_paragraph("Line five with more old text.")

    after = Document()
    after.add_paragraph("Line one with changed words.")
    # Line two removed
    after.add_paragraph("Line three with new text.")
    after.add_paragraph("Line four unchanged.")
    after.add_paragraph("Line five with more new text.")
    after.add_paragraph("Line six freshly added.")

    save_pair("revision-ids-unique", before, after, {
        "name": "revision-ids-unique",
        "spec_ref": "ECMA-376 §17.13.5",
        "description": "Multiple edits to verify all annotation IDs are unique",
    })


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    print("Generating tracked-change-serialization test fixtures:")
    make_del_uses_deltext()
    make_no_nested_ins_del()
    make_no_revisions_in_comments()
    make_flat_tracked_changes()
    make_move_ranges_paired()
    make_revision_ids_unique()
    print("Done.")


if __name__ == "__main__":
    main()

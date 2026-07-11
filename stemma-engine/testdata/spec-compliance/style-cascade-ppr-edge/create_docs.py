# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for paragraph property cascade edge-case tests.

Targets known spec violations in style resolution:
  1. pageBreakBefore not resolved through style chain (§17.3.1.23)
  2. Paragraph shading not resolved through style chain (§17.3.1.31)
  3. keepNext w:val="0" treated as true (§17.3.1.15)
  4. keepLines w:val="0" treated as true (§17.3.1.14)
  5. pageBreakBefore w:val="0" treated as true (§17.3.1.23)
  6. Tab stops through deep basedOn chain with clears (§17.3.1.38)
  7. Borders through basedOn chain — whole-element replacement (§17.3.1.24)

Run:  uv run create_docs.py
"""

import json
from pathlib import Path
from lxml import etree

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def ensure_ppr(p) -> OxmlElement:
    """Return the w:pPr element for a paragraph, creating it if absent."""
    ppr = p._element.find(qn("w:pPr"))
    if ppr is None:
        ppr = make_element("w:pPr")
        p._element.insert(0, ppr)
    return ppr


def strip_default_style(p):
    """Remove the auto-generated pStyle from a paragraph."""
    ppr = p._element.find(qn("w:pPr"))
    if ppr is not None:
        ps = ppr.find(qn("w:pStyle"))
        if ps is not None:
            ppr.remove(ps)


def save_fixture(name: str, doc, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  style-cascade-ppr-edge/{name}/")


# =========================================================================
# Fixture 1: pageBreakBefore from style chain (§17.3.1.23)
#
# Style "BreakBefore" defines pageBreakBefore. A paragraph referencing
# that style should inherit page_break_before=true.
# BUG: page_break_before is never resolved through the style chain.
# =========================================================================

def make_page_break_before_from_style():
    doc = Document()
    styles_el = doc.styles.element

    # BreakBefore: paragraph style with pageBreakBefore
    style = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "BreakBefore",
    })
    style.append(make_element("w:name", {"w:val": "Break Before"}))
    ppr = make_element("w:pPr")
    ppr.append(make_element("w:pageBreakBefore"))
    style.append(ppr)
    styles_el.append(style)

    # Para 0: no style — page_break_before=false
    p0 = doc.add_paragraph("No style no break")
    strip_default_style(p0)

    # Para 1: BreakBefore style — should inherit page_break_before=true
    p1 = doc.add_paragraph("Break style applied")
    ensure_ppr(p1).append(make_element("w:pStyle", {"w:val": "BreakBefore"}))

    # Para 2: BreakBefore style + direct pageBreakBefore val="0" — explicit off
    p2 = doc.add_paragraph("Break style with direct off")
    p2_ppr = ensure_ppr(p2)
    p2_ppr.append(make_element("w:pStyle", {"w:val": "BreakBefore"}))
    p2_ppr.append(make_element("w:pageBreakBefore", {"w:val": "0"}))

    save_fixture("page-break-before-from-style", doc, {
        "spec": "ISO 29500-1 §17.3.1.23",
        "description": (
            "Style defines pageBreakBefore. Paragraph referencing that "
            "style should inherit page_break_before=true. Direct val=0 "
            "should override to false."
        ),
        "expected": {
            "para_0": "page_break_before=false (no style)",
            "para_1": "page_break_before=true (from style)",
            "para_2": "page_break_before=false (direct val=0 overrides style)",
        },
    })


# =========================================================================
# Fixture 2: Paragraph shading from style chain (§17.3.1.31)
#
# Style "Shaded" defines w:shd in pPr. A paragraph referencing that
# style should inherit the shading.
# BUG: paragraph shading is read from direct pPr only, never from styles.
# =========================================================================

def make_para_shading_from_style():
    doc = Document()
    styles_el = doc.styles.element

    # Shaded: paragraph style with shading fill=FFFF00 (yellow)
    style = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "Shaded",
    })
    style.append(make_element("w:name", {"w:val": "Shaded"}))
    ppr = make_element("w:pPr")
    shd = make_element("w:shd", {
        "w:val": "clear",
        "w:color": "auto",
        "w:fill": "FFFF00",
    })
    ppr.append(shd)
    style.append(ppr)
    styles_el.append(style)

    # ShadedChild basedOn Shaded (no shading of its own)
    child_style = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "ShadedChild",
    })
    child_style.append(make_element("w:name", {"w:val": "Shaded Child"}))
    child_style.append(make_element("w:basedOn", {"w:val": "Shaded"}))
    child_ppr = make_element("w:pPr")
    child_ppr.append(make_element("w:jc", {"w:val": "center"}))
    child_style.append(child_ppr)
    styles_el.append(child_style)

    # Para 0: no style — no shading
    p0 = doc.add_paragraph("No style no shading")
    strip_default_style(p0)

    # Para 1: Shaded style — should inherit shading fill=FFFF00
    p1 = doc.add_paragraph("Shaded style applied")
    ensure_ppr(p1).append(make_element("w:pStyle", {"w:val": "Shaded"}))

    # Para 2: ShadedChild style (basedOn Shaded) — should inherit shading
    p2 = doc.add_paragraph("Shaded child style applied")
    ensure_ppr(p2).append(make_element("w:pStyle", {"w:val": "ShadedChild"}))

    # Para 3: Shaded style + direct shading override (green)
    p3 = doc.add_paragraph("Shaded style with direct override")
    p3_ppr = ensure_ppr(p3)
    p3_ppr.append(make_element("w:pStyle", {"w:val": "Shaded"}))
    p3_ppr.append(make_element("w:shd", {
        "w:val": "clear",
        "w:color": "auto",
        "w:fill": "00FF00",
    }))

    save_fixture("para-shading-from-style", doc, {
        "spec": "ISO 29500-1 §17.3.1.31",
        "description": (
            "Style defines paragraph shading (w:shd fill=FFFF00). "
            "Paragraphs should inherit shading from style chain."
        ),
        "expected": {
            "para_0": "shading=None (no style)",
            "para_1": "shading.fill=FFFF00 (from style)",
            "para_2": "shading.fill=FFFF00 (inherited through basedOn)",
            "para_3": "shading.fill=00FF00 (direct overrides style)",
        },
    })


# =========================================================================
# Fixture 3: keepNext val="0" parsing (§17.3.1.15)
#
# BUG: extract_keep_next checks presence only (is_some), ignoring val="0".
# A paragraph with <w:keepNext w:val="0"/> gets keep_next=true.
# =========================================================================

def make_keep_next_val_false():
    doc = Document()
    styles_el = doc.styles.element

    # KeepStyle: paragraph style with keepNext=true
    style = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "KeepStyle",
    })
    style.append(make_element("w:name", {"w:val": "Keep Style"}))
    ppr = make_element("w:pPr")
    ppr.append(make_element("w:keepNext"))
    style.append(ppr)
    styles_el.append(style)

    # Para 0: no style, no keepNext — keep_next=false
    p0 = doc.add_paragraph("No keep next")
    strip_default_style(p0)

    # Para 1: direct keepNext (no val) — keep_next=true
    p1 = doc.add_paragraph("Direct keep next on")
    p1_ppr = ensure_ppr(p1)
    strip_default_style(p1)
    p1_ppr.append(make_element("w:keepNext"))

    # Para 2: direct keepNext val="0" — should be keep_next=false
    p2 = doc.add_paragraph("Direct keep next val zero")
    p2_ppr = ensure_ppr(p2)
    strip_default_style(p2)
    p2_ppr.append(make_element("w:keepNext", {"w:val": "0"}))

    # Para 3: KeepStyle + direct keepNext val="0" — should be false (direct off overrides style on)
    p3 = doc.add_paragraph("Style on direct off")
    p3_ppr = ensure_ppr(p3)
    p3_ppr.append(make_element("w:pStyle", {"w:val": "KeepStyle"}))
    p3_ppr.append(make_element("w:keepNext", {"w:val": "0"}))

    save_fixture("keep-next-val-false", doc, {
        "spec": "ISO 29500-1 §17.3.1.15",
        "description": (
            "Tests that w:keepNext w:val='0' is correctly parsed as false. "
            "Also tests that direct val=0 overrides style keepNext=true."
        ),
        "expected": {
            "para_0": "keep_next=false (no keepNext element)",
            "para_1": "keep_next=true (element present, no val = true)",
            "para_2": "keep_next=false (val=0 means false)",
            "para_3": "keep_next=false (direct val=0 overrides style's true)",
        },
    })


# =========================================================================
# Fixture 4: keepLines val="0" parsing (§17.3.1.14)
#
# Same bug as keepNext: extract_keep_lines ignores val="0".
# =========================================================================

def make_keep_lines_val_false():
    doc = Document()

    # Para 0: no keepLines — false
    p0 = doc.add_paragraph("No keep lines")
    strip_default_style(p0)

    # Para 1: direct keepLines (no val) — true
    p1 = doc.add_paragraph("Direct keep lines on")
    p1_ppr = ensure_ppr(p1)
    strip_default_style(p1)
    p1_ppr.append(make_element("w:keepLines"))

    # Para 2: direct keepLines val="0" — should be false
    p2 = doc.add_paragraph("Direct keep lines val zero")
    p2_ppr = ensure_ppr(p2)
    strip_default_style(p2)
    p2_ppr.append(make_element("w:keepLines", {"w:val": "0"}))

    save_fixture("keep-lines-val-false", doc, {
        "spec": "ISO 29500-1 §17.3.1.14",
        "description": (
            "Tests that w:keepLines w:val='0' is correctly parsed as false."
        ),
        "expected": {
            "para_0": "keep_lines=false",
            "para_1": "keep_lines=true",
            "para_2": "keep_lines=false (val=0)",
        },
    })


# =========================================================================
# Fixture 5: pageBreakBefore val="0" parsing (§17.3.1.23)
#
# Same bug: extract_page_break_before ignores val="0".
# =========================================================================

def make_page_break_before_val_false():
    doc = Document()

    # Para 0: no pageBreakBefore — false
    p0 = doc.add_paragraph("No page break before")
    strip_default_style(p0)

    # Para 1: direct pageBreakBefore (no val) — true
    p1 = doc.add_paragraph("Direct page break before on")
    p1_ppr = ensure_ppr(p1)
    strip_default_style(p1)
    p1_ppr.append(make_element("w:pageBreakBefore"))

    # Para 2: direct pageBreakBefore val="0" — should be false
    p2 = doc.add_paragraph("Direct page break before val zero")
    p2_ppr = ensure_ppr(p2)
    strip_default_style(p2)
    p2_ppr.append(make_element("w:pageBreakBefore", {"w:val": "0"}))

    save_fixture("page-break-before-val-false", doc, {
        "spec": "ISO 29500-1 §17.3.1.23",
        "description": (
            "Tests that w:pageBreakBefore w:val='0' is parsed as false."
        ),
        "expected": {
            "para_0": "page_break_before=false",
            "para_1": "page_break_before=true",
            "para_2": "page_break_before=false (val=0)",
        },
    })


# =========================================================================
# Fixture 6: Tab stops through deep basedOn chain with clears (§17.3.1.38)
#
# GrandparentTabs: tabs at 720(left), 1440(center), 2880(right)
# ParentTabs (basedOn GrandparentTabs): adds 4320(left), clears 1440
# ChildTabs (basedOn ParentTabs): clears 720, adds 5760(decimal)
#
# Expected effective tabs for ChildTabs:
#   2880(right from grandparent) + 4320(left from parent) + 5760(decimal from child)
#
# The tab stop overlay logic (overlay_tab_stops) seems correct per code
# review. This test documents the expected behavior so we can catch
# regressions.
# =========================================================================

def make_tab_stops_deep_chain():
    doc = Document()
    styles_el = doc.styles.element

    # GrandparentTabs: tabs at 720(left), 1440(center), 2880(right, dot leader)
    gp_style = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "GrandparentTabs",
    })
    gp_style.append(make_element("w:name", {"w:val": "Grandparent Tabs"}))
    gp_ppr = make_element("w:pPr")
    gp_tabs = make_element("w:tabs")
    gp_tabs.append(make_element("w:tab", {
        "w:val": "left", "w:pos": "720",
    }))
    gp_tabs.append(make_element("w:tab", {
        "w:val": "center", "w:pos": "1440",
    }))
    gp_tabs.append(make_element("w:tab", {
        "w:val": "right", "w:pos": "2880", "w:leader": "dot",
    }))
    gp_ppr.append(gp_tabs)
    gp_style.append(gp_ppr)
    styles_el.append(gp_style)

    # ParentTabs: basedOn GrandparentTabs, adds 4320(left), clears 1440
    p_style = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "ParentTabs",
    })
    p_style.append(make_element("w:name", {"w:val": "Parent Tabs"}))
    p_style.append(make_element("w:basedOn", {"w:val": "GrandparentTabs"}))
    p_ppr = make_element("w:pPr")
    p_tabs = make_element("w:tabs")
    p_tabs.append(make_element("w:tab", {
        "w:val": "clear", "w:pos": "1440",
    }))
    p_tabs.append(make_element("w:tab", {
        "w:val": "left", "w:pos": "4320",
    }))
    p_ppr.append(p_tabs)
    p_style.append(p_ppr)
    styles_el.append(p_style)

    # ChildTabs: basedOn ParentTabs, clears 720, adds 5760(decimal)
    c_style = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "ChildTabs",
    })
    c_style.append(make_element("w:name", {"w:val": "Child Tabs"}))
    c_style.append(make_element("w:basedOn", {"w:val": "ParentTabs"}))
    c_ppr = make_element("w:pPr")
    c_tabs = make_element("w:tabs")
    c_tabs.append(make_element("w:tab", {
        "w:val": "clear", "w:pos": "720",
    }))
    c_tabs.append(make_element("w:tab", {
        "w:val": "decimal", "w:pos": "5760",
    }))
    c_ppr.append(c_tabs)
    c_style.append(c_ppr)
    styles_el.append(c_style)

    # Para 0: GrandparentTabs — 3 tab stops
    p0 = doc.add_paragraph("Grandparent tabs para")
    ensure_ppr(p0).append(make_element("w:pStyle", {"w:val": "GrandparentTabs"}))

    # Para 1: ParentTabs — 720(left), 2880(right), 4320(left) = 3 stops
    p1 = doc.add_paragraph("Parent tabs para")
    ensure_ppr(p1).append(make_element("w:pStyle", {"w:val": "ParentTabs"}))

    # Para 2: ChildTabs — 2880(right), 4320(left), 5760(decimal) = 3 stops
    p2 = doc.add_paragraph("Child tabs para")
    ensure_ppr(p2).append(make_element("w:pStyle", {"w:val": "ChildTabs"}))

    save_fixture("tab-stops-deep-chain", doc, {
        "spec": "ISO 29500-1 §17.3.1.38",
        "description": (
            "Three-level style chain with tab stop adds and clears. "
            "Tests overlay_tab_stops through basedOn chain."
        ),
        "expected": {
            "para_0_grandparent": "tabs: 720(left), 1440(center), 2880(right/dot)",
            "para_1_parent": "tabs: 720(left), 2880(right/dot), 4320(left) [1440 cleared]",
            "para_2_child": "tabs: 2880(right/dot), 4320(left), 5760(decimal) [720 cleared]",
        },
    })


# =========================================================================
# Fixture 7: Borders through basedOn chain (§17.3.1.24)
#
# Parent style defines top + bottom borders.
# Child style (basedOn parent) defines only left border.
# Per OOXML: pBdr is whole-element. Child's pBdr replaces parent's entirely.
# So child should have only left border, NOT top+bottom from parent.
#
# This tests that whole-element replacement is correctly implemented.
# =========================================================================

def make_borders_based_on_chain():
    doc = Document()
    styles_el = doc.styles.element

    # BorderParent: top + bottom borders
    parent_style = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "BorderParent",
    })
    parent_style.append(make_element("w:name", {"w:val": "Border Parent"}))
    parent_ppr = make_element("w:pPr")
    pbdr = make_element("w:pBdr")
    top_border = make_element("w:top", {
        "w:val": "single", "w:sz": "4", "w:color": "FF0000",
    })
    bottom_border = make_element("w:bottom", {
        "w:val": "single", "w:sz": "4", "w:color": "0000FF",
    })
    pbdr.append(top_border)
    pbdr.append(bottom_border)
    parent_ppr.append(pbdr)
    parent_style.append(parent_ppr)
    styles_el.append(parent_style)

    # BorderChild: basedOn BorderParent, defines only left border
    child_style = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "BorderChild",
    })
    child_style.append(make_element("w:name", {"w:val": "Border Child"}))
    child_style.append(make_element("w:basedOn", {"w:val": "BorderParent"}))
    child_ppr = make_element("w:pPr")
    child_pbdr = make_element("w:pBdr")
    left_border = make_element("w:left", {
        "w:val": "single", "w:sz": "4", "w:color": "00FF00",
    })
    child_pbdr.append(left_border)
    child_ppr.append(child_pbdr)
    child_style.append(child_ppr)
    styles_el.append(child_style)

    # BorderInherit: basedOn BorderParent, no pBdr — inherits parent borders
    inherit_style = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "BorderInherit",
    })
    inherit_style.append(make_element("w:name", {"w:val": "Border Inherit"}))
    inherit_style.append(make_element("w:basedOn", {"w:val": "BorderParent"}))
    inherit_ppr = make_element("w:pPr")
    inherit_ppr.append(make_element("w:jc", {"w:val": "center"}))
    inherit_style.append(inherit_ppr)
    styles_el.append(inherit_style)

    # Para 0: BorderParent — top + bottom
    p0 = doc.add_paragraph("Parent borders top bottom")
    ensure_ppr(p0).append(make_element("w:pStyle", {"w:val": "BorderParent"}))

    # Para 1: BorderChild — only left (parent borders replaced)
    p1 = doc.add_paragraph("Child border left only")
    ensure_ppr(p1).append(make_element("w:pStyle", {"w:val": "BorderChild"}))

    # Para 2: BorderInherit — should inherit top + bottom from parent
    p2 = doc.add_paragraph("Inherit borders from parent")
    ensure_ppr(p2).append(make_element("w:pStyle", {"w:val": "BorderInherit"}))

    save_fixture("borders-based-on-chain", doc, {
        "spec": "ISO 29500-1 §17.3.1.24",
        "description": (
            "Tests paragraph border inheritance through basedOn chain. "
            "pBdr is whole-element replacement: child with pBdr replaces parent. "
            "Child without pBdr inherits parent borders."
        ),
        "expected": {
            "para_0_parent": "borders: top=FF0000, bottom=0000FF",
            "para_1_child": "borders: left=00FF00 only (top+bottom NOT inherited)",
            "para_2_inherit": "borders: top=FF0000, bottom=0000FF (inherited)",
        },
    })


# =========================================================================
# Fixture 8: pageBreakBefore through deep basedOn chain
#
# Root style: pageBreakBefore=true
# Child (basedOn Root): no pageBreakBefore (should inherit true)
# Grandchild (basedOn Child): pageBreakBefore=false (explicit off)
#
# BUG: since pageBreakBefore is not resolved through style chain, none
# of these paragraphs will get the style's value.
# =========================================================================

def make_page_break_deep_chain():
    doc = Document()
    styles_el = doc.styles.element

    # Root: pageBreakBefore=true
    root_style = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "BreakRoot",
    })
    root_style.append(make_element("w:name", {"w:val": "Break Root"}))
    root_ppr = make_element("w:pPr")
    root_ppr.append(make_element("w:pageBreakBefore"))
    root_style.append(root_ppr)
    styles_el.append(root_style)

    # Child: basedOn Root, no pageBreakBefore (inherits true)
    child_style = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "BreakChild",
    })
    child_style.append(make_element("w:name", {"w:val": "Break Child"}))
    child_style.append(make_element("w:basedOn", {"w:val": "BreakRoot"}))
    child_ppr = make_element("w:pPr")
    child_ppr.append(make_element("w:jc", {"w:val": "center"}))
    child_style.append(child_ppr)
    styles_el.append(child_style)

    # Grandchild: basedOn Child, pageBreakBefore=false (explicit off)
    gc_style = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "BreakGrandchild",
    })
    gc_style.append(make_element("w:name", {"w:val": "Break Grandchild"}))
    gc_style.append(make_element("w:basedOn", {"w:val": "BreakChild"}))
    gc_ppr = make_element("w:pPr")
    gc_ppr.append(make_element("w:pageBreakBefore", {"w:val": "0"}))
    gc_style.append(gc_ppr)
    styles_el.append(gc_style)

    # Para 0: BreakRoot — page_break_before=true
    p0 = doc.add_paragraph("Root break before")
    ensure_ppr(p0).append(make_element("w:pStyle", {"w:val": "BreakRoot"}))

    # Para 1: BreakChild — inherits page_break_before=true from root
    p1 = doc.add_paragraph("Child inherits break")
    ensure_ppr(p1).append(make_element("w:pStyle", {"w:val": "BreakChild"}))

    # Para 2: BreakGrandchild — explicit off cancels inherited true
    p2 = doc.add_paragraph("Grandchild break off")
    ensure_ppr(p2).append(make_element("w:pStyle", {"w:val": "BreakGrandchild"}))

    save_fixture("page-break-deep-chain", doc, {
        "spec": "ISO 29500-1 §17.3.1.23 + §17.7.4.3",
        "description": (
            "Three-level basedOn chain for pageBreakBefore. "
            "Root=true, child inherits true, grandchild explicitly sets false."
        ),
        "expected": {
            "para_0": "page_break_before=true (from root style)",
            "para_1": "page_break_before=true (inherited from root through child)",
            "para_2": "page_break_before=false (grandchild explicitly overrides to false)",
        },
    })


# =========================================================================
# Fixture 9: contextualSpacing explicit false roundtrip (§17.3.1.9)
#
# Tests that contextualSpacing=false (w:val="0") is correctly distinguished
# from absent. When a style defines contextualSpacing=true and direct
# formatting sets it to false, the false should win.
# =========================================================================

def make_contextual_spacing_false():
    doc = Document()
    styles_el = doc.styles.element

    # CtxStyle: paragraph style with contextualSpacing=true
    style = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "CtxStyle",
    })
    style.append(make_element("w:name", {"w:val": "Ctx Style"}))
    ppr = make_element("w:pPr")
    ppr.append(make_element("w:contextualSpacing"))  # no val = true
    style.append(ppr)
    styles_el.append(style)

    # Para 0: no style — contextual_spacing=None
    p0 = doc.add_paragraph("No style no ctx spacing")
    strip_default_style(p0)

    # Para 1: CtxStyle — should inherit contextual_spacing=Some(true)
    p1 = doc.add_paragraph("Ctx style true")
    ensure_ppr(p1).append(make_element("w:pStyle", {"w:val": "CtxStyle"}))

    # Para 2: CtxStyle + direct contextualSpacing val="0" — should be Some(false)
    p2 = doc.add_paragraph("Ctx style direct false")
    p2_ppr = ensure_ppr(p2)
    p2_ppr.append(make_element("w:pStyle", {"w:val": "CtxStyle"}))
    p2_ppr.append(make_element("w:contextualSpacing", {"w:val": "0"}))

    save_fixture("contextual-spacing-false", doc, {
        "spec": "ISO 29500-1 §17.3.1.9",
        "description": (
            "Tests contextualSpacing=false overriding style's true. "
            "Ensures val=0 is parsed as explicit false, not ignored."
        ),
        "expected": {
            "para_0": "contextual_spacing=None",
            "para_1": "contextual_spacing=Some(true) (from style)",
            "para_2": "contextual_spacing=Some(false) (direct val=0 overrides style)",
        },
    })


# =========================================================================
# Main
# =========================================================================

if __name__ == "__main__":
    print("Generating style-cascade-ppr-edge fixtures:")
    make_page_break_before_from_style()
    make_para_shading_from_style()
    make_keep_next_val_false()
    make_keep_lines_val_false()
    make_page_break_before_val_false()
    make_tab_stops_deep_chain()
    make_borders_based_on_chain()
    make_page_break_deep_chain()
    make_contextual_spacing_false()
    print("Done.")

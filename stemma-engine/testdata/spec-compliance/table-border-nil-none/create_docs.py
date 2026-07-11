# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for nil vs none border semantics (MS-OI29500 SS2.1.169).

Run:  cd stemma-engine/testdata/spec-compliance/table-border-nil-none && mise exec -- uv run create_docs.py
"""

import json
from pathlib import Path
from lxml import etree

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc: DocxDocument, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  table-border-nil-none/{name}/")


def make_table_with_borders(doc, rows=1, cols=2, cell_texts=None):
    """Create a table, return (tbl_element, tblPr_element)."""
    tbl = doc.add_table(rows=rows, cols=cols)
    if cell_texts:
        for r in range(rows):
            for c in range(cols):
                if r < len(cell_texts) and c < len(cell_texts[r]):
                    tbl.cell(r, c).text = cell_texts[r][c]
    tbl_el = tbl._tbl
    tblPr = tbl_el.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_el.insert(0, tblPr)
    return tbl_el, tblPr


def add_tbl_borders(tblPr, edges, val="single", sz="8", color="000000", space="0"):
    """Add tblBorders to tblPr."""
    tbl_borders = make_element("w:tblBorders")
    for edge in edges:
        border = make_element(f"w:{edge}", {
            "w:val": val, "w:sz": sz, "w:color": color, "w:space": space,
        })
        tbl_borders.append(border)
    tblPr.append(tbl_borders)


def add_tc_borders(tbl_el, row, col, edges_spec):
    """Add tcBorders to a specific cell. edges_spec is a list of
    (edge_name, val, sz, color, space) tuples."""
    tr = tbl_el.findall(w("tr"))[row]
    tc = tr.findall(w("tc"))[col]
    tcPr = tc.find(w("tcPr"))
    if tcPr is None:
        tcPr = make_element("w:tcPr")
        tc.insert(0, tcPr)
    tc_borders = make_element("w:tcBorders")
    for edge, val, sz, color, space in edges_spec:
        border = make_element(f"w:{edge}", {
            "w:val": val, "w:sz": sz, "w:color": color, "w:space": space,
        })
        tc_borders.append(border)
    tcPr.append(tc_borders)


def add_tbl_grid(tbl_el, col_widths):
    """Add or replace tblGrid with specific column widths."""
    existing = tbl_el.find(w("tblGrid"))
    if existing is not None:
        tbl_el.remove(existing)
    grid = make_element("w:tblGrid")
    for cw in col_widths:
        grid.append(make_element("w:gridCol", {"w:w": str(cw)}))
    tblPr = tbl_el.find(w("tblPr"))
    if tblPr is not None:
        idx = list(tbl_el).index(tblPr) + 1
        tbl_el.insert(idx, grid)
    else:
        tbl_el.insert(0, grid)


# =====================================================================
# Fixture 1: nil suppresses opposing table border
# =====================================================================

def make_nil_suppresses_opposing():
    """MS-OI29500 SS2.1.169: nil means "no border shall be displayed",
    suppressing even the opposing side's inherited table border.

    1x2 table with table-level borders (single sz=8 black on all edges).
    Cell (0,0): tcBorders with right = nil
    Cell (0,1): no tcBorders (inherits table borders)

    At the shared edge, cell (0,0) has explicit nil. Per spec, nil means
    "no border shall be displayed" — the insideV inherited by cell (0,1)
    should NOT display because cell (0,0) has nil.
    """
    doc = Document()
    doc.add_paragraph("Fixture 1: nil suppresses opposing table border.")

    tbl_el, tblPr = make_table_with_borders(doc, 1, 2, [
        ["NilRight", "InheritsTable"],
    ])

    add_tbl_borders(tblPr,
                    ["top", "bottom", "left", "right", "insideH", "insideV"],
                    val="single", sz="8", color="000000")

    add_tbl_grid(tbl_el, [3000, 3000])

    # Cell (0,0): right border = nil
    add_tc_borders(tbl_el, 0, 0, [
        ("right", "nil", "0", "auto", "0"),
    ])

    save_fixture("nil-suppresses-opposing", doc, {
        "name": "nil-suppresses-opposing",
        "spec_ref": "MS-OI29500 SS2.1.169",
        "description": (
            "1x2 table with table-level borders (single sz=8 black). "
            "Cell (0,0) has tcBorders right=nil. Cell (0,1) has no tcBorders. "
            "Per MS-OI29500 SS2.1.169, nil means 'no border shall be displayed', "
            "so the shared edge should have NO visible border."
        ),
    })


# =====================================================================
# Fixture 2: none allows opposing border to display
# =====================================================================

def make_none_allows_opposing():
    """MS-OI29500 SS2.1.169: none means "I have no border", so the
    opposing side's border wins and is displayed.

    1x2 table with table-level borders (single sz=8 black on all edges).
    Cell (0,0): tcBorders with right = none (val="none")
    Cell (0,1): tcBorders with left = single sz=12 red

    At the shared edge, cell (0,0) has none (= "I have no border").
    Cell (0,1) has a red border. Per spec, the red border should win.
    """
    doc = Document()
    doc.add_paragraph("Fixture 2: none allows opposing border to display.")

    tbl_el, tblPr = make_table_with_borders(doc, 1, 2, [
        ["NoneRight", "RedLeft"],
    ])

    add_tbl_borders(tblPr,
                    ["top", "bottom", "left", "right", "insideH", "insideV"],
                    val="single", sz="8", color="000000")

    add_tbl_grid(tbl_el, [3000, 3000])

    # Cell (0,0): right border = none
    add_tc_borders(tbl_el, 0, 0, [
        ("right", "none", "0", "auto", "0"),
    ])

    # Cell (0,1): left border = single sz=12 red
    add_tc_borders(tbl_el, 0, 1, [
        ("left", "single", "12", "FF0000", "0"),
    ])

    save_fixture("none-allows-opposing", doc, {
        "name": "none-allows-opposing",
        "spec_ref": "MS-OI29500 SS2.1.169",
        "description": (
            "1x2 table with table-level borders (single sz=8 black). "
            "Cell (0,0) has tcBorders right=none. Cell (0,1) has tcBorders "
            "left=single sz=12 red. Per MS-OI29500 SS2.1.169, none means "
            "'I have no border' so the opposing red border should display."
        ),
    })


# =====================================================================
# Fixture 3: nil beats thick border
# =====================================================================

def make_nil_beats_thick():
    """MS-OI29500 SS2.1.169: nil suppresses even a thick heavy border
    on the opposing side.

    1x2 table (no table-level borders).
    Cell (0,0): tcBorders with right = nil
    Cell (0,1): tcBorders with left = thick sz=24 red

    Per spec, nil means "no border shall be displayed" — it suppresses
    even the thick border on the opposing side.
    """
    doc = Document()
    doc.add_paragraph("Fixture 3: nil beats thick border.")

    tbl_el, tblPr = make_table_with_borders(doc, 1, 2, [
        ["NilRight", "ThickLeft"],
    ])

    add_tbl_grid(tbl_el, [3000, 3000])

    # Cell (0,0): right border = nil
    add_tc_borders(tbl_el, 0, 0, [
        ("right", "nil", "0", "auto", "0"),
    ])

    # Cell (0,1): left border = thick sz=24 red
    add_tc_borders(tbl_el, 0, 1, [
        ("left", "thick", "24", "FF0000", "0"),
    ])

    save_fixture("nil-beats-thick-border", doc, {
        "name": "nil-beats-thick-border",
        "spec_ref": "MS-OI29500 SS2.1.169",
        "description": (
            "1x2 table with no table-level borders. "
            "Cell (0,0) has tcBorders right=nil. Cell (0,1) has tcBorders "
            "left=thick sz=24 red. Per MS-OI29500 SS2.1.169, nil means "
            "'no border shall be displayed' and suppresses the thick border."
        ),
    })


# =====================================================================
# Fixture 4: none vs none — no border
# =====================================================================

def make_none_vs_none():
    """Both sides have none — no border displays at the shared edge.

    1x2 table (no table-level borders).
    Cell (0,0): tcBorders with right = none
    Cell (0,1): tcBorders with left = none

    Both sides say "I have no border" — no border displays.
    """
    doc = Document()
    doc.add_paragraph("Fixture 4: none vs none — no border.")

    tbl_el, tblPr = make_table_with_borders(doc, 1, 2, [
        ["NoneRight", "NoneLeft"],
    ])

    add_tbl_grid(tbl_el, [3000, 3000])

    # Cell (0,0): right border = none
    add_tc_borders(tbl_el, 0, 0, [
        ("right", "none", "0", "auto", "0"),
    ])

    # Cell (0,1): left border = none
    add_tc_borders(tbl_el, 0, 1, [
        ("left", "none", "0", "auto", "0"),
    ])

    save_fixture("none-vs-none-no-border", doc, {
        "name": "none-vs-none-no-border",
        "spec_ref": "MS-OI29500 SS2.1.169",
        "description": (
            "1x2 table with no table-level borders. "
            "Cell (0,0) has tcBorders right=none. Cell (0,1) has tcBorders "
            "left=none. Both sides say 'I have no border' — no border displays."
        ),
    })


# =====================================================================
# Main
# =====================================================================

def main():
    print("\n== Table Border Nil vs None Fixtures (MS-OI29500 SS2.1.169) ==")
    make_nil_suppresses_opposing()
    make_none_allows_opposing()
    make_nil_beats_thick()
    make_none_vs_none()
    print("\nDone.")


if __name__ == "__main__":
    main()

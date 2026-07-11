# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for spec violations found during the Pandoc audit.

Fixtures:
  - pandoc-audit/nested-table-in-cell/  — table cell whose last block is a nested table (§17.4.65)
  - pandoc-audit/numbering-merge-base/  — base doc with numbered list (for merge ordering test)
  - pandoc-audit/numbering-merge-target/ — target doc with different numbered list (for merge ordering test)

Run:  uv run create_pandoc_audit.py
"""

import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc, metadata: dict) -> None:
    out = ROOT / "pandoc-audit" / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  pandoc-audit/{name}/")


# =========================================================================
# 1) Nested table in cell — last block is w:tbl, not w:p (§17.4.65)
# =========================================================================


def make_nested_table_in_cell():
    """Create a table where one cell contains a nested table as its LAST element.

    ECMA-376 §17.4.65: "If a table cell does not include at least one
    block-level element, then this document shall be considered corrupt."

    The XSD (CT_Tc) requires EG_BlockLevelElts with minOccurs=1, which
    includes both w:p and w:tbl. However, Word requires the last child
    of w:tc to be a w:p — a nested table as the final element is invalid
    and causes corruption in Word 2007/2010.

    This fixture has a cell with: [paragraph, nested-table] — no trailing w:p.
    We manipulate the XML directly to remove the trailing paragraph that
    python-docx would normally add.
    """
    doc = Document()
    doc.add_paragraph("Document with nested table whose cell lacks trailing paragraph.")

    # Create the outer table: 1 row, 1 cell
    outer_table = doc.add_table(rows=1, cols=1)
    outer_cell = outer_table.rows[0].cells[0]

    # Add text to the cell first
    outer_cell.paragraphs[0].text = "Text before nested table"

    # Add a nested table inside the cell
    nested_table = outer_cell.add_table(rows=1, cols=1)
    nested_table.rows[0].cells[0].paragraphs[0].text = "Nested cell content"

    # python-docx adds a trailing paragraph after the nested table.
    # We need to remove it to create the violation.
    tc_el = outer_cell._element
    # Find all w:p elements after the nested w:tbl
    tbl_el = tc_el.findall(w("tbl"))[-1]
    # Remove any w:p that comes after the nested table
    found_tbl = False
    to_remove = []
    for child in tc_el:
        if child is tbl_el:
            found_tbl = True
            continue
        if found_tbl and child.tag == w("p"):
            to_remove.append(child)
    for el in to_remove:
        tc_el.remove(el)

    save_fixture("nested-table-in-cell", doc, {
        "description": "Table cell whose last block-level element is a nested w:tbl, not w:p",
        "spec_ref": "ECMA-376 §17.4.65 (CT_Tc content model)",
        "violation": "Last child of w:tc must be w:p; nested w:tbl as final child is invalid",
    })


# =========================================================================
# 2) Numbering merge — base and target with non-overlapping numbering IDs
# =========================================================================


def _rewrite_numbering_xml(docx_path: str, abstract_nums: list[dict], nums: list[dict]) -> None:
    """Rewrite the numbering.xml inside an existing DOCX to contain only the
    specified abstractNum and num definitions. Uses zipfile replace-in-place."""
    import zipfile, io, shutil, tempfile
    from xml.etree import ElementTree as ET

    nsmap = {
        "wpc": "http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas",
        "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
        "o": "urn:schemas-microsoft-com:office:office",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
        "v": "urn:schemas-microsoft-com:vml",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "w10": "urn:schemas-microsoft-com:office:word",
        "w": W,
        "wne": "http://schemas.microsoft.com/office/word/2006/wordml",
    }

    # Build a fresh numbering root
    root = ET.Element(f"{{{W}}}numbering")
    for prefix, uri in nsmap.items():
        root.set(f"xmlns:{prefix}", uri)

    for adef in abstract_nums:
        abs_el = ET.SubElement(root, f"{{{W}}}abstractNum")
        abs_el.set(f"{{{W}}}abstractNumId", str(adef["id"]))
        # multiLevelType
        mlt = ET.SubElement(abs_el, f"{{{W}}}multiLevelType")
        mlt.set(f"{{{W}}}val", "hybridMultilevel")
        # Single level 0
        lvl = ET.SubElement(abs_el, f"{{{W}}}lvl")
        lvl.set(f"{{{W}}}ilvl", "0")
        start = ET.SubElement(lvl, f"{{{W}}}start")
        start.set(f"{{{W}}}val", str(adef.get("start", 1)))
        fmt = ET.SubElement(lvl, f"{{{W}}}numFmt")
        fmt.set(f"{{{W}}}val", adef.get("fmt", "decimal"))
        txt = ET.SubElement(lvl, f"{{{W}}}lvlText")
        txt.set(f"{{{W}}}val", adef.get("text", "%1."))

    for ndef in nums:
        num_el = ET.SubElement(root, f"{{{W}}}num")
        num_el.set(f"{{{W}}}numId", str(ndef["numId"]))
        abs_ref = ET.SubElement(num_el, f"{{{W}}}abstractNumId")
        abs_ref.set(f"{{{W}}}val", str(ndef["abstractNumId"]))

    xml_bytes = ET.tostring(root, encoding="UTF-8", xml_declaration=True)

    # Replace numbering.xml in the zip
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp.name, "w") as zout:
        for item in zin.infolist():
            if item.filename == "word/numbering.xml":
                zout.writestr(item, xml_bytes)
            else:
                zout.writestr(item, zin.read(item.filename))
    shutil.move(tmp.name, docx_path)


def _set_paragraph_numpr(p_element, num_id: int, ilvl: int = 0) -> None:
    """Set w:numPr on a paragraph element's pPr."""
    pPr = p_element.find(w("pPr"))
    if pPr is None:
        pPr = make_element("w:pPr")
        p_element.insert(0, pPr)
    numPr = make_element("w:numPr")
    ilvl_el = make_element("w:ilvl", {"w:val": str(ilvl)})
    numId_el = make_element("w:numId", {"w:val": str(num_id)})
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)


def make_numbering_merge_base():
    """Base document with a numbered list using abstractNumId=10, numId=100.

    These IDs are chosen to not overlap with the target document's IDs
    (abstractNumId=20, numId=200), forcing the merge to copy new elements.

    The document has shared text with the target so the diff matches paragraphs
    and recognizes the numbering changes, keeping target's numId on modified paras.
    """
    doc = Document()
    doc.add_paragraph("Shared paragraph one.")

    for i in range(1, 4):
        p = doc.add_paragraph(f"List item {i}")
        _set_paragraph_numpr(p._element, num_id=100, ilvl=0)

    doc.add_paragraph("Shared paragraph two.")

    # Additional paragraph that will be "new" in the target
    # (absent here, present in target) to trigger insert with target numbering.
    doc.add_paragraph("End of document.")

    out = ROOT / "pandoc-audit" / "numbering-merge-base"
    out.mkdir(parents=True, exist_ok=True)
    docx_path = str(out / "input.docx")
    doc.save(docx_path)

    # Rewrite numbering.xml with our specific IDs
    _rewrite_numbering_xml(docx_path, [
        {"id": 10, "fmt": "decimal", "text": "%1.", "start": 1},
    ], [
        {"numId": 100, "abstractNumId": 10},
    ])

    (out / "metadata.json").write_text(json.dumps({
        "description": "Base document with numbering: abstractNumId=10, numId=100",
        "spec_ref": "ECMA-376 Annex A, CT_Numbering (xsd:sequence: abstractNum* then num*)",
    }, indent=2) + "\n")
    print(f"  pandoc-audit/numbering-merge-base/")


def make_numbering_merge_target():
    """Target document with numbered lists using BOTH base numbering (numId=100)
    and new numbering (abstractNumId=20, numId=200).

    The shared paragraphs keep numId=100 (same as base), while new paragraphs
    use numId=200 (requires merge). This forces merge_target_numbering to
    copy abstractNumId=20 and numId=200 into the base's numbering.xml, which
    already has existing num elements — triggering the interleaving violation.
    """
    doc = Document()
    doc.add_paragraph("Shared paragraph one.")

    # Same list items as base (matched by diff), keep base numbering
    for i in range(1, 4):
        p = doc.add_paragraph(f"List item {i}")
        _set_paragraph_numpr(p._element, num_id=100, ilvl=0)

    doc.add_paragraph("Shared paragraph two.")

    # NEW paragraphs with different numbering — these trigger the merge
    for i in range(1, 3):
        p = doc.add_paragraph(f"New numbered item {i}")
        _set_paragraph_numpr(p._element, num_id=200, ilvl=0)

    doc.add_paragraph("End of document.")

    out = ROOT / "pandoc-audit" / "numbering-merge-target"
    out.mkdir(parents=True, exist_ok=True)
    docx_path = str(out / "input.docx")
    doc.save(docx_path)

    # Rewrite numbering.xml with BOTH definitions: base's (10/100) for shared
    # paragraphs, plus new (20/200) for inserted paragraphs
    _rewrite_numbering_xml(docx_path, [
        {"id": 10, "fmt": "decimal", "text": "%1.", "start": 1},
        {"id": 20, "fmt": "decimal", "text": "(%1)", "start": 1},
    ], [
        {"numId": 100, "abstractNumId": 10},
        {"numId": 200, "abstractNumId": 20},
    ])

    (out / "metadata.json").write_text(json.dumps({
        "description": "Target document with numbering: abstractNumId=10+20, numId=100+200",
        "spec_ref": "ECMA-376 Annex A, CT_Numbering (xsd:sequence: abstractNum* then num*)",
    }, indent=2) + "\n")
    print(f"  pandoc-audit/numbering-merge-target/")


# =========================================================================
# Main
# =========================================================================


if __name__ == "__main__":
    print("Generating pandoc-audit fixtures:")
    make_nested_table_in_cell()
    make_numbering_merge_base()
    make_numbering_merge_target()
    print("Done.")

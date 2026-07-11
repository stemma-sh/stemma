# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for the numbering resolution chain audit.

Tests numbering resolution semantics from ISO 29500-1 §17.9:
  - §17.9.25: start default of 0 when element omitted
  - §17.9.28: suff (suffix) types — tab, space, nothing
  - §17.9.10: lvlRestart edge cases
  - §17.9.21/27: numStyleLink/styleLink chain
  - §17.9.8/26: lvlOverride with both startOverride and level replacement
  - §17.9.4: isLgl with current level non-decimal format
  - §17.9.8: Two lvlOverrides on different levels of same num

Run:  uv run create_docs.py
"""

import json
from pathlib import Path
from lxml import etree

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

# ── XML namespace helpers ────────────────────────────────────────────────

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def w(tag: str) -> str:
    """Return a fully-qualified wordprocessingml tag."""
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    """Create an OxmlElement with optional attributes."""
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


# ── Save helpers ─────────────────────────────────────────────────────────

def save_fixture(
    name: str,
    doc: DocxDocument,
    metadata: dict,
    filename: str = "input.docx",
) -> None:
    """Save a single-doc fixture."""
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / filename))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  numbering-audit/{name}/")


def _inject_numbering_xml(doc: DocxDocument, numbering_xml: str) -> None:
    """Replace the numbering.xml part with custom XML."""
    numbering_part = None
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        pass

    if numbering_part is None:
        doc.add_paragraph("dummy", style="List Bullet")
        body = doc.element.body
        last_p = body.findall(w("p"))[-1]
        body.remove(last_p)
        numbering_part = doc.part.numbering_part

    numbering_part._element = etree.fromstring(numbering_xml.encode("utf-8"))


def _inject_style(doc: DocxDocument, style_xml: str) -> None:
    """Append a custom style element to the styles part."""
    style_el = etree.fromstring(style_xml.encode("utf-8"))
    doc.styles.element.append(style_el)


def _add_numbered_para(doc: DocxDocument, text: str, num_id: str, ilvl: str = "0"):
    """Add a paragraph with numPr referencing the given numId/ilvl."""
    p = doc.add_paragraph(text)
    pPr = p._p.get_or_add_pPr()
    numPr = make_element("w:numPr")
    numPr.append(make_element("w:ilvl", {"w:val": ilvl}))
    numPr.append(make_element("w:numId", {"w:val": num_id}))
    pPr.append(numPr)
    return p


# =========================================================================
# 1. start default of 0 (§17.9.25)
# =========================================================================

def make_start_default_zero() -> None:
    """§17.9.25: When w:start is omitted, the starting value shall be 0.

    Fixture:
      - Level 0: decimal, NO w:start element (should default to 0)
      - Level 1: lowerLetter, w:start w:val="1" (explicit, control)
      - Three paragraphs at level 0

    Expected: first paragraph is "0.", second is "1.", third is "2."
    Bug: implementation defaults to 1, producing "1.", "2.", "3."
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="1">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerLetter"/>
      <w:lvlText w:val="%2)"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    _add_numbered_para(doc, "First item (start omitted, should be 0)", "1", "0")
    _add_numbered_para(doc, "Second item (should be 1)", "1", "0")
    _add_numbered_para(doc, "Third item (should be 2)", "1", "0")
    _add_numbered_para(doc, "Sub-item (explicit start=1, control)", "1", "1")

    save_fixture("start-default-zero", doc, {
        "name": "start-default-zero",
        "spec_ref": "ISO 29500-1 §17.9.25",
        "description": (
            "Level 0 has no w:start element. Per §17.9.25, the default "
            "starting value is 0, not 1."
        ),
        "expected_behavior": (
            "Level 0 items: '0.', '1.', '2.'. "
            "Level 1 item: 'a)' (explicit start=1)."
        ),
    })


# =========================================================================
# 2. suff (suffix) types (§17.9.28)
# =========================================================================

def make_suffix_types() -> None:
    """§17.9.28: suff controls separator between number and text.

    Fixture:
      - numId=1: level 0 with suff="tab" (default, explicit)
      - numId=2: level 0 with suff="space"
      - numId=3: level 0 with suff="nothing"
      - Two paragraphs per numId

    Expected:
      - numId=1: "1.\tBody" (tab separator)
      - numId=2: "1. Body" (space separator — note: rendered_text should use space)
      - numId=3: "1.Body" (no separator)
    Bug: implementation always uses tab separator, ignoring suff.
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:suff w:val="tab"/>
    </w:lvl>
  </w:abstractNum>
  <w:abstractNum w:abstractNumId="1">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:suff w:val="space"/>
    </w:lvl>
  </w:abstractNum>
  <w:abstractNum w:abstractNumId="2">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:suff w:val="nothing"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
  <w:num w:numId="2">
    <w:abstractNumId w:val="1"/>
  </w:num>
  <w:num w:numId="3">
    <w:abstractNumId w:val="2"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    _add_numbered_para(doc, "Tab suffix first", "1", "0")
    _add_numbered_para(doc, "Tab suffix second", "1", "0")
    _add_numbered_para(doc, "Space suffix first", "2", "0")
    _add_numbered_para(doc, "Space suffix second", "2", "0")
    _add_numbered_para(doc, "Nothing suffix first", "3", "0")
    _add_numbered_para(doc, "Nothing suffix second", "3", "0")

    save_fixture("suffix-types", doc, {
        "name": "suffix-types",
        "spec_ref": "ISO 29500-1 §17.9.28",
        "description": (
            "Three numbering definitions with different suff values: "
            "tab, space, nothing. Tests that the suffix between number "
            "text and paragraph body is correct."
        ),
        "expected_behavior": (
            "Tab: '1.\\tBody'. Space: '1. Body'. Nothing: '1.Body'. "
            "Default when suff omitted is tab."
        ),
    })


# =========================================================================
# 3. lvlRestart specific level (§17.9.10)
# =========================================================================

def make_lvl_restart_specific_level() -> None:
    """§17.9.10: lvlRestart with a specific level (not 0, not default).

    Fixture: 3-level list.
      - Level 0: decimal
      - Level 1: lowerLetter (default restart — restarts on level 0)
      - Level 2: lowerRoman, lvlRestart=1 (restart only when level 0 is hit,
        NOT when level 1 is hit)

    Sequence:
      1. / a) / b) / i) / ii) / c) / iii) / iv) / 2. / a) / i)

    The key test: when level 1 (c) appears after level 2 (ii), level 2
    should NOT restart because lvlRestart=1 means "only restart when level 0
    (1-based index 1 → 0-indexed 0) is hit."

    Expected: after "c)", level 2 continues at "iii)" not "i)".
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="1">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerLetter"/>
      <w:lvlText w:val="%2)"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="2">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerRoman"/>
      <w:lvlText w:val="%3)"/>
      <w:lvlJc w:val="left"/>
      <w:lvlRestart w:val="1"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # Group 1
    _add_numbered_para(doc, "Top level", "1", "0")          # 1.
    _add_numbered_para(doc, "Sub a", "1", "1")              # a)
    _add_numbered_para(doc, "Sub b", "1", "1")              # b)
    _add_numbered_para(doc, "Sub-sub first", "1", "2")      # i)
    _add_numbered_para(doc, "Sub-sub second", "1", "2")     # ii)
    # Back to level 1 — level 2 should NOT restart (lvlRestart=1 means only level 0 triggers)
    _add_numbered_para(doc, "Sub c", "1", "1")              # c)
    _add_numbered_para(doc, "Sub-sub continues", "1", "2")  # iii) -- NOT i)
    _add_numbered_para(doc, "Sub-sub continues", "1", "2")  # iv)  -- NOT ii)
    # Back to level 0 — NOW level 2 should restart
    _add_numbered_para(doc, "Top level 2", "1", "0")        # 2.
    _add_numbered_para(doc, "Sub a again", "1", "1")        # a) (restarted)
    _add_numbered_para(doc, "Sub-sub restarted", "1", "2")  # i) (restarted because level 0 was hit)

    save_fixture("lvl-restart-specific", doc, {
        "name": "lvl-restart-specific",
        "spec_ref": "ISO 29500-1 §17.9.10",
        "description": (
            "Level 2 has lvlRestart=1 (restart only when level 0 is used). "
            "Level 1 appearing does NOT reset level 2's counter."
        ),
        "expected_behavior": (
            "Sequence: 1. / a) / b) / i) / ii) / c) / iii) / iv) / 2. / a) / i). "
            "Level 2 continues through iii) and iv) after level 1 'c)' because "
            "lvlRestart=1 means only level 0 triggers restart."
        ),
    })


# =========================================================================
# 4. numStyleLink chain (§17.9.21)
# =========================================================================

def make_num_style_link() -> None:
    """§17.9.21: numStyleLink references a numbering style.

    Fixture:
      - abstractNum 0: has numStyleLink pointing to "OutlineNumbering" style
        (no levels defined — it's just a reference)
      - abstractNum 1: has styleLink pointing to "OutlineNumbering"
        (the actual level definitions are here)
      - Style "OutlineNumbering" (type=numbering) with numPr pointing to numId=2
      - numId=1 -> abstractNum 0 (the reference)
      - numId=2 -> abstractNum 1 (the actual definition)
      - Paragraphs use numId=1 — should follow the chain through the style

    Expected: paragraphs get "1.", "2." etc from the actual definition in abstractNum 1.
    Bug: numStyleLink is not followed; paragraphs using numId=1 get no numbering
    because abstractNum 0 has no levels.
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:multiLevelType w:val="multilevel"/>
    <w:numStyleLink w:val="OutlineNumbering"/>
  </w:abstractNum>
  <w:abstractNum w:abstractNumId="1">
    <w:multiLevelType w:val="multilevel"/>
    <w:styleLink w:val="OutlineNumbering"/>
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="1">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerLetter"/>
      <w:lvlText w:val="%2)"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
  <w:num w:numId="2">
    <w:abstractNumId w:val="1"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # Create numbering style "OutlineNumbering" with numPr -> numId=2
    style_xml = f"""\
<w:style w:type="numbering" w:styleId="OutlineNumbering" xmlns:w="{W}">
  <w:name w:val="Outline Numbering"/>
  <w:pPr>
    <w:numPr>
      <w:numId w:val="2"/>
    </w:numPr>
  </w:pPr>
</w:style>"""
    _inject_style(doc, style_xml)

    _add_numbered_para(doc, "First via numStyleLink chain", "1", "0")
    _add_numbered_para(doc, "Second via numStyleLink chain", "1", "0")
    _add_numbered_para(doc, "Sub-item via numStyleLink chain", "1", "1")

    save_fixture("num-style-link", doc, {
        "name": "num-style-link",
        "spec_ref": "ISO 29500-1 §17.9.21, §17.9.27",
        "description": (
            "abstractNum 0 has numStyleLink='OutlineNumbering'. "
            "Style 'OutlineNumbering' (type=numbering) has numPr -> numId=2. "
            "abstractNum 1 has styleLink='OutlineNumbering' with actual levels. "
            "Paragraphs using numId=1 should follow the chain."
        ),
        "expected_behavior": (
            "Paragraphs using numId=1: '1.', '2.', 'a)'. "
            "The numStyleLink chain is followed to get actual level definitions."
        ),
    })


# =========================================================================
# 5. lvlOverride with both startOverride AND level replacement (§17.9.8)
# =========================================================================

def make_lvl_override_both() -> None:
    """§17.9.8: lvlOverride can contain both startOverride and a full level.

    When both are present, the level replacement provides the formatting,
    and startOverride provides the starting value.

    Fixture:
      - abstractNum 0: level 0 decimal with lvlText="%1."
      - numId=1: uses abstract directly
      - numId=2: lvlOverride on level 0 with:
        - startOverride=10
        - full lvl replacement changing format to upperLetter with lvlText="(%1)"

    Expected:
      - numId=1: "1.", "2."
      - numId=2: "(J)", "(K)" — upperLetter starting at 10 (J is 10th letter)
    Bug potential: if startOverride is applied to the abstract definition's start
    rather than the replacement level's start, or if one overrides the other.
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
  <w:num w:numId="2">
    <w:abstractNumId w:val="0"/>
    <w:lvlOverride w:ilvl="0">
      <w:startOverride w:val="10"/>
      <w:lvl w:ilvl="0">
        <w:start w:val="1"/>
        <w:numFmt w:val="upperLetter"/>
        <w:lvlText w:val="(%1)"/>
        <w:lvlJc w:val="left"/>
      </w:lvl>
    </w:lvlOverride>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    _add_numbered_para(doc, "Original first", "1", "0")
    _add_numbered_para(doc, "Original second", "1", "0")
    _add_numbered_para(doc, "Override with both", "2", "0")
    _add_numbered_para(doc, "Override continues", "2", "0")

    save_fixture("lvl-override-both", doc, {
        "name": "lvl-override-both",
        "spec_ref": "ISO 29500-1 §17.9.8, §17.9.26",
        "description": (
            "lvlOverride with both startOverride=10 and a full level "
            "replacement (upperLetter format, lvlText='(%1)'). "
            "The format comes from the replacement, start from startOverride."
        ),
        "expected_behavior": (
            "numId=1: '1.', '2.'. "
            "numId=2: '(J)', '(K)' — upperLetter starting at 10."
        ),
    })


# =========================================================================
# 6. isLgl with current level non-decimal (§17.9.4)
# =========================================================================

def make_is_lgl_current_level() -> None:
    """§17.9.4: isLgl forces ALL %N references to decimal, including current level.

    Fixture:
      - Level 0: decimal (control)
      - Level 1: lowerLetter with isLgl + lvlText "%1.%2"
        Both %1 (decimal) and %2 (lowerLetter) should render as decimal.

    Expected: Level 1 renders as "1.1" not "1.a".
    This tests that isLgl affects the current level's own format too,
    not just the inherited levels.
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="1">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerLetter"/>
      <w:isLgl/>
      <w:lvlText w:val="%1.%2"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    _add_numbered_para(doc, "Top level", "1", "0")
    _add_numbered_para(doc, "Sub with isLgl (should be 1.1)", "1", "1")
    _add_numbered_para(doc, "Sub continues (should be 1.2)", "1", "1")

    save_fixture("is-lgl-current-level", doc, {
        "name": "is-lgl-current-level",
        "spec_ref": "ISO 29500-1 §17.9.4",
        "description": (
            "Level 1 uses lowerLetter format but has isLgl set. "
            "Both %1 and %2 in lvlText should render as decimal."
        ),
        "expected_behavior": (
            "Level 0: '1.' Level 1: '1.1', '1.2' — NOT '1.a', '1.b'."
        ),
    })


# =========================================================================
# 7. Multiple lvlOverrides on different levels (§17.9.8)
# =========================================================================

def make_multi_level_override() -> None:
    """§17.9.8: A num can have lvlOverrides on multiple levels.

    Fixture:
      - abstractNum 0: 3-level outline (decimal / lowerLetter / lowerRoman)
      - numId=1: uses abstract directly (control)
      - numId=2: lvlOverride on BOTH level 0 (startOverride=5) and
        level 1 (startOverride=3)

    Expected with numId=2:
      - Level 0: starts at 5 ("5.")
      - Level 1: starts at 3 ("c)")
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="1">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerLetter"/>
      <w:lvlText w:val="%2)"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="2">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerRoman"/>
      <w:lvlText w:val="%3)"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
  <w:num w:numId="2">
    <w:abstractNumId w:val="0"/>
    <w:lvlOverride w:ilvl="0">
      <w:startOverride w:val="5"/>
    </w:lvlOverride>
    <w:lvlOverride w:ilvl="1">
      <w:startOverride w:val="3"/>
    </w:lvlOverride>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    # Control: numId=1
    _add_numbered_para(doc, "Control level 0", "1", "0")
    _add_numbered_para(doc, "Control level 1", "1", "1")

    # Test: numId=2 with both overrides
    _add_numbered_para(doc, "Override level 0 (should be 5.)", "2", "0")
    _add_numbered_para(doc, "Override level 1 (should be c))", "2", "1")
    _add_numbered_para(doc, "Override level 1 cont (should be d))", "2", "1")
    _add_numbered_para(doc, "Override level 0 cont (should be 6.)", "2", "0")

    save_fixture("multi-level-override", doc, {
        "name": "multi-level-override",
        "spec_ref": "ISO 29500-1 §17.9.8",
        "description": (
            "A num with lvlOverride on both level 0 (startOverride=5) and "
            "level 1 (startOverride=3). Tests that multiple overrides on "
            "different levels work correctly."
        ),
        "expected_behavior": (
            "Control: '1.', 'a)'. "
            "Override: '5.', 'c)', 'd)', '6.'."
        ),
    })


# =========================================================================
# 8. isLgl with numFmt=none level (§17.9.4)
# =========================================================================

def make_is_lgl_none_preserved() -> None:
    """§17.9.4: isLgl forces decimal, but numFmt=none is preserved.

    MS-OI29500 §17.9.4b: levels with numFmt=none are NOT forced to decimal.

    Fixture:
      - Level 0: none format (produces empty text)
      - Level 1: lowerLetter (normal)
      - Level 2: decimal with isLgl + lvlText "%1.%2.%3"

    Expected for level 2: ".1.1" — level 0 is none (produces ""), level 1
    and level 2 are forced to decimal. So the result is "%1=%2=1.%3=1" → ".1.1".
    """
    doc = Document()

    numbering_xml = f"""\
<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{W}" xmlns:r="{R}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="none"/>
      <w:lvlText w:val=""/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="1">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerLetter"/>
      <w:lvlText w:val="%2)"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
    <w:lvl w:ilvl="2">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:isLgl/>
      <w:lvlText w:val="%1.%2.%3"/>
      <w:lvlJc w:val="left"/>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""

    _inject_numbering_xml(doc, numbering_xml)

    _add_numbered_para(doc, "Level 0 (none format)", "1", "0")
    _add_numbered_para(doc, "Level 1 (lowerLetter)", "1", "1")
    _add_numbered_para(doc, "Level 2 with isLgl", "1", "2")

    save_fixture("is-lgl-none-preserved", doc, {
        "name": "is-lgl-none-preserved",
        "spec_ref": "ISO 29500-1 §17.9.4, MS-OI29500 §17.9.4b",
        "description": (
            "Level 0 is numFmt=none. Level 2 has isLgl and lvlText '%1.%2.%3'. "
            "isLgl forces decimal on all levels EXCEPT numFmt=none, which is "
            "preserved as empty."
        ),
        "expected_behavior": (
            "Level 2: '.1.1' — %1 is empty (none preserved), "
            "%2 and %3 are decimal."
        ),
    })


# =========================================================================

def main() -> None:
    print("Generating numbering audit fixtures:")
    make_start_default_zero()
    make_suffix_types()
    make_lvl_restart_specific_level()
    make_num_style_link()
    make_lvl_override_both()
    make_is_lgl_current_level()
    make_multi_level_override()
    make_is_lgl_none_preserved()
    print("\nDone.")


if __name__ == "__main__":
    main()

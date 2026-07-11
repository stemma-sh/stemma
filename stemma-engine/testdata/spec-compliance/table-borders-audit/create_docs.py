# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for table border conflict resolution & conditional
formatting audit tests (ISO 29500-1 SS17.4.66, SS17.7.6).

Run:  mise exec -- python3 stemma-engine/testdata/spec-compliance/table-borders-audit/create_docs.py
"""

import json
from pathlib import Path
from lxml import etree

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc: DocxDocument, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  table-borders-audit/{name}/")


# =====================================================================
# Helper: build a 2x2 table with tblBorders and optional cell overrides
# =====================================================================

def make_table_with_borders(doc, rows=2, cols=2, cell_texts=None):
    """Create a table, return (tbl_element, tblPr_element)."""
    tbl = doc.add_table(rows=rows, cols=cols)
    if cell_texts:
        for r in range(rows):
            for c in range(cols):
                if r < len(cell_texts) and c < len(cell_texts[r]):
                    tbl.cell(r, c).text = cell_texts[r][c]
    tbl_el = tbl._tbl
    tblPr = tbl_el.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_el.insert(0, tblPr)
    return tbl_el, tblPr


def add_tbl_borders(tblPr, edges, val="single", sz="12", color="FF0000", space="0"):
    """Add tblBorders to tblPr."""
    tbl_borders = make_element("w:tblBorders")
    for edge in edges:
        border = make_element(f"w:{edge}", {
            "w:val": val, "w:sz": sz, "w:color": color, "w:space": space,
        })
        tbl_borders.append(border)
    tblPr.append(tbl_borders)


def add_tc_borders(tbl_el, row, col, edges, val="single", sz="4", color="0000FF", space="0"):
    """Add tcBorders to a specific cell."""
    tr = tbl_el.findall(w("tr"))[row]
    tc = tr.findall(w("tc"))[col]
    tcPr = tc.find(w("tcPr"))
    if tcPr is None:
        tcPr = make_element("w:tcPr")
        tc.insert(0, tcPr)
    tc_borders = make_element("w:tcBorders")
    for edge in edges:
        border = make_element(f"w:{edge}", {
            "w:val": val, "w:sz": sz, "w:color": color, "w:space": space,
        })
        tc_borders.append(border)
    tcPr.append(tc_borders)


def add_tbl_grid(tbl_el, col_widths):
    """Add or replace tblGrid with specific column widths."""
    existing = tbl_el.find(w("tblGrid"))
    if existing is not None:
        tbl_el.remove(existing)
    grid = make_element("w:tblGrid")
    for cw in col_widths:
        grid.append(make_element("w:gridCol", {"w:w": str(cw)}))
    # Insert after tblPr
    tblPr = tbl_el.find(w("tblPr"))
    if tblPr is not None:
        idx = list(tbl_el).index(tblPr) + 1
        tbl_el.insert(idx, grid)
    else:
        tbl_el.insert(0, grid)


# =====================================================================
# Fixture 1: Table borders as cell defaults (no tcBorders)
# =====================================================================

def make_table_border_fallback_no_cell_borders():
    """SS17.4.66: When a cell has no tcBorders, table-level tblBorders
    should be used as the cell's effective borders.

    Table has red single sz=12 borders on all 6 edges (top, bottom, left,
    right, insideH, insideV). NO cells have tcBorders. All cells should
    inherit the table-level borders.

    Cell (0,0) is top-left: top=table.top, bottom=table.insideH,
      left=table.left, right=table.insideV
    Cell (0,1) is top-right: top=table.top, bottom=table.insideH,
      left=table.insideV, right=table.right
    Cell (1,0) is bottom-left: top=table.insideH, bottom=table.bottom,
      left=table.left, right=table.insideV
    Cell (1,1) is bottom-right: top=table.insideH, bottom=table.bottom,
      left=table.insideV, right=table.right
    """
    doc = Document()
    doc.add_paragraph("Table with tblBorders only, no cell-level tcBorders.")

    tbl_el, tblPr = make_table_with_borders(doc, 2, 2, [
        ["R0C0", "R0C1"],
        ["R1C0", "R1C1"],
    ])

    add_tbl_borders(tblPr,
                    ["top", "bottom", "left", "right", "insideH", "insideV"],
                    val="single", sz="12", color="FF0000")

    save_fixture("border-fallback-no-cell-borders", doc, {
        "name": "border-fallback-no-cell-borders",
        "spec_ref": "ISO 29500-1 SS17.4.66",
        "description": (
            "Table with tblBorders on all 6 edges, NO tcBorders on any cell. "
            "Cells should inherit table borders as defaults."
        ),
        "expected": {
            "cell_0_0": "top=table.top, bottom=insideH, left=table.left, right=insideV",
            "cell_0_1": "top=table.top, bottom=insideH, left=insideV, right=table.right",
            "cell_1_0": "top=insideH, bottom=table.bottom, left=table.left, right=insideV",
            "cell_1_1": "top=insideH, bottom=table.bottom, left=insideV, right=table.right",
        },
    })


# =====================================================================
# Fixture 2: insideV used for interior vertical edges
# =====================================================================

def make_inside_v_interior_cells():
    """SS17.4.66 + SS17.4.38: Interior cells' left/right borders should
    come from insideV, not from tblBorders top-level left/right.

    Table: 2x3 with different borders for left/right (blue sz=12) vs
    insideV (green sz=4). The middle column (col 1) should have green
    borders on both left and right, not blue.
    """
    doc = Document()
    doc.add_paragraph("Table with different left/right vs insideV borders.")

    tbl_el, tblPr = make_table_with_borders(doc, 2, 3, [
        ["R0C0", "R0C1", "R0C2"],
        ["R1C0", "R1C1", "R1C2"],
    ])

    # Set distinct borders: outer=blue, insideV=green
    tbl_borders = make_element("w:tblBorders")
    for edge in ["top", "bottom"]:
        tbl_borders.append(make_element(f"w:{edge}", {
            "w:val": "single", "w:sz": "12", "w:color": "0000FF", "w:space": "0",
        }))
    tbl_borders.append(make_element("w:left", {
        "w:val": "single", "w:sz": "12", "w:color": "0000FF", "w:space": "0",
    }))
    tbl_borders.append(make_element("w:right", {
        "w:val": "single", "w:sz": "12", "w:color": "0000FF", "w:space": "0",
    }))
    tbl_borders.append(make_element("w:insideH", {
        "w:val": "single", "w:sz": "8", "w:color": "888888", "w:space": "0",
    }))
    tbl_borders.append(make_element("w:insideV", {
        "w:val": "single", "w:sz": "4", "w:color": "00FF00", "w:space": "0",
    }))
    tblPr.append(tbl_borders)

    save_fixture("inside-v-interior-cells", doc, {
        "name": "inside-v-interior-cells",
        "spec_ref": "ISO 29500-1 SS17.4.66, SS17.4.38",
        "description": (
            "Table with distinct left/right (blue sz=12) vs insideV (green sz=4). "
            "Interior cells should use insideV for their vertical edges."
        ),
    })


# =====================================================================
# Fixture 3: Cell border overrides nil (no border)
# =====================================================================

def make_cell_border_nil_override():
    """SS17.4.66 rule 1: If either conflicting border is nil/none,
    the opposing border is displayed.

    Table has red single sz=12 borders on all edges.
    Cell (0,0) sets its top border to nil. Per rule 1, the table's
    top border (red) should NOT display — the nil wins because it's
    more specific (cell wins over table per rule 2).
    """
    doc = Document()
    doc.add_paragraph("Cell (0,0) has nil top border overriding table border.")

    tbl_el, tblPr = make_table_with_borders(doc, 2, 2, [
        ["NilTop", "Normal"],
        ["Normal", "Normal"],
    ])

    add_tbl_borders(tblPr,
                    ["top", "bottom", "left", "right", "insideH", "insideV"],
                    val="single", sz="12", color="FF0000")

    # Cell (0,0): explicit nil top border
    add_tc_borders(tbl_el, 0, 0, ["top"], val="nil", sz="0", color="auto")

    save_fixture("cell-border-nil-override", doc, {
        "name": "cell-border-nil-override",
        "spec_ref": "ISO 29500-1 SS17.4.66 rule 1-2",
        "description": (
            "Cell (0,0) has nil top border. Per SS17.4.66, cell borders win "
            "over table borders (rule 2). The nil border suppresses the table border."
        ),
    })


# =====================================================================
# Fixture 4: Cell margin per-side cascade
# =====================================================================

def make_cell_margin_per_side_cascade():
    """SS17.4.68: tcMar overrides tblCellMar, but per-side.

    Table-level default: top=100, left=200, bottom=100, right=200 (twips).
    Cell (0,0): tcMar overrides ONLY left=400.
    Per spec, cell (0,0) effective margins should be: top=100 (from table),
    left=400 (from cell), bottom=100 (from table), right=200 (from table).
    """
    doc = Document()
    doc.add_paragraph("Cell margin per-side cascade test.")

    tbl_el, tblPr = make_table_with_borders(doc, 2, 2, [
        ["LeftOverride", "Default"],
        ["Default", "Default"],
    ])

    # Table-level cell margins: all sides 100tw except left/right 200tw
    tbl_cell_mar = make_element("w:tblCellMar")
    tbl_cell_mar.append(make_element("w:top", {"w:w": "100", "w:type": "dxa"}))
    tbl_cell_mar.append(make_element("w:start", {"w:w": "200", "w:type": "dxa"}))
    tbl_cell_mar.append(make_element("w:bottom", {"w:w": "100", "w:type": "dxa"}))
    tbl_cell_mar.append(make_element("w:end", {"w:w": "200", "w:type": "dxa"}))
    tblPr.append(tbl_cell_mar)

    # Cell (0,0): override only left margin
    tr = tbl_el.findall(w("tr"))[0]
    tc = tr.findall(w("tc"))[0]
    tcPr = tc.find(w("tcPr"))
    if tcPr is None:
        tcPr = make_element("w:tcPr")
        tc.insert(0, tcPr)
    tc_mar = make_element("w:tcMar")
    tc_mar.append(make_element("w:start", {"w:w": "400", "w:type": "dxa"}))
    tcPr.append(tc_mar)

    save_fixture("cell-margin-per-side-cascade", doc, {
        "name": "cell-margin-per-side-cascade",
        "spec_ref": "ISO 29500-1 SS17.4.68",
        "description": (
            "Table-level margins: top=100, left=200, bottom=100, right=200. "
            "Cell (0,0) overrides only left=400 via tcMar. "
            "Per SS17.4.68, tcMar overrides tblCellMar — but the spec says "
            "'shall override the table cell margins', meaning the entire set "
            "is replaced. So cell (0,0) should have left=400 and the other "
            "sides fall back to defaults (not from table-level)."
        ),
    })


# =====================================================================
# Fixture 5: Width type nil
# =====================================================================

def make_width_type_nil():
    """SS17.18.90 ST_TblWidth + MS-OI29500 SS17.4.2(a):
    Width type 'nil' means 'no width specified' — not zero width.
    Word ignores the element when type='nil'.

    Table with two rows:
    - Row 0, Cell 0: tcW w:type="nil" w:w="0"
    - Row 0, Cell 1: tcW w:type="dxa" w:w="3000"
    - Row 1, Cell 0: tcW w:type="auto" w:w="0"
    """
    doc = Document()
    doc.add_paragraph("Width type nil test.")

    tbl_el, tblPr = make_table_with_borders(doc, 2, 2, [
        ["NilWidth", "DxaWidth"],
        ["AutoWidth", "Normal"],
    ])

    add_tbl_grid(tbl_el, [3000, 3000])

    # Cell (0,0): type=nil
    tr0 = tbl_el.findall(w("tr"))[0]
    tc00 = tr0.findall(w("tc"))[0]
    tcPr = tc00.find(w("tcPr"))
    if tcPr is None:
        tcPr = make_element("w:tcPr")
        tc00.insert(0, tcPr)
    # Remove existing tcW if any
    existing_tcw = tcPr.find(w("tcW"))
    if existing_tcw is not None:
        tcPr.remove(existing_tcw)
    tcPr.append(make_element("w:tcW", {"w:w": "0", "w:type": "nil"}))

    # Cell (0,1): type=dxa w=3000
    tc01 = tr0.findall(w("tc"))[1]
    tcPr01 = tc01.find(w("tcPr"))
    if tcPr01 is None:
        tcPr01 = make_element("w:tcPr")
        tc01.insert(0, tcPr01)
    existing_tcw = tcPr01.find(w("tcW"))
    if existing_tcw is not None:
        tcPr01.remove(existing_tcw)
    tcPr01.append(make_element("w:tcW", {"w:w": "3000", "w:type": "dxa"}))

    save_fixture("width-type-nil", doc, {
        "name": "width-type-nil",
        "spec_ref": "ISO 29500-1 SS17.18.90, MS-OI29500 SS17.4.2(a)",
        "description": (
            "Cell (0,0) has tcW type=nil w=0: per MS-OI29500, this means "
            "'no width specified' and Word ignores the element entirely. "
            "Cell (0,1) has tcW type=dxa w=3000."
        ),
    })


# =====================================================================
# Fixture 6: vMerge continue without restart
# =====================================================================

def make_vmerge_continue_no_restart():
    """SS17.4.84: vMerge continue without a preceding restart.

    3x2 table where row 0 has no vMerge, row 1 cell 0 has
    vMerge=continue (orphan). Per spec, the document is non-conformant.
    Word treats it as a standalone cell.
    """
    doc = Document()
    doc.add_paragraph("vMerge continue without restart.")

    tbl_el, tblPr = make_table_with_borders(doc, 3, 2, [
        ["NoMerge", "B"],
        ["OrphanContinue", "C"],
        ["Normal", "D"],
    ])

    # Row 1, Cell 0: set vMerge=continue (orphan — no restart above)
    tr1 = tbl_el.findall(w("tr"))[1]
    tc10 = tr1.findall(w("tc"))[0]
    tcPr = tc10.find(w("tcPr"))
    if tcPr is None:
        tcPr = make_element("w:tcPr")
        tc10.insert(0, tcPr)
    v_merge = make_element("w:vMerge")
    # No w:val attribute → defaults to "continue"
    tcPr.append(v_merge)

    save_fixture("vmerge-continue-no-restart", doc, {
        "name": "vmerge-continue-no-restart",
        "spec_ref": "ISO 29500-1 SS17.4.84",
        "description": (
            "Row 1, Cell 0 has vMerge=continue but there is no restart "
            "above it. Per spec, the document is non-conformant. Word treats "
            "the cell as standalone (no merge). Our code should normalize "
            "this to VerticalMerge::None."
        ),
    })


# =====================================================================
# Fixture 7: gridSpan exceeds remaining grid columns
# =====================================================================

def make_gridspan_overflow():
    """SS17.4.17: gridSpan must not exceed grid column count.

    3-column grid (3 gridCols). Row 0: 3 normal cells (span 1 each).
    Row 1: cell 0 has gridSpan=1, cell 1 has gridSpan=5 (exceeds grid).
    Our code should clamp or reject.
    """
    doc = Document()
    doc.add_paragraph("gridSpan overflow test.")

    tbl = doc.add_table(rows=2, cols=3)
    tbl.cell(0, 0).text = "A"
    tbl.cell(0, 1).text = "B"
    tbl.cell(0, 2).text = "C"
    tbl.cell(1, 0).text = "D"

    tbl_el = tbl._tbl

    add_tbl_grid(tbl_el, [2000, 2000, 2000])

    # Row 1: keep cell 0, merge cells 1+2 but set gridSpan=5 (too big)
    tr1 = tbl_el.findall(w("tr"))[1]
    tcs = tr1.findall(w("tc"))

    # Remove the 3rd cell from row 1 (merge into 2nd)
    tr1.remove(tcs[2])

    # Set cell 1 gridSpan=5
    tc_1_1 = tcs[1]
    tcPr = tc_1_1.find(w("tcPr"))
    if tcPr is None:
        tcPr = make_element("w:tcPr")
        tc_1_1.insert(0, tcPr)
    existing_gs = tcPr.find(w("gridSpan"))
    if existing_gs is not None:
        tcPr.remove(existing_gs)
    tcPr.append(make_element("w:gridSpan", {"w:val": "5"}))

    # Also set tcW for the merged cell
    existing_tcw = tcPr.find(w("tcW"))
    if existing_tcw is not None:
        tcPr.remove(existing_tcw)
    tcPr.append(make_element("w:tcW", {"w:w": "4000", "w:type": "dxa"}))

    save_fixture("gridspan-overflow", doc, {
        "name": "gridspan-overflow",
        "spec_ref": "ISO 29500-1 SS17.4.17, MS-OI29500 SS17.4.17(a)",
        "description": (
            "3-column grid. Row 1 cell 1 has gridSpan=5 which exceeds the "
            "grid width. MS-OI29500 says gridSpan of 0 is treated as 1. "
            "Our code should clamp gridSpan to the grid column count."
        ),
    })


# =====================================================================
# Fixture 8: Conditional banding with lastRow exclusion
# =====================================================================

def make_banding_last_row_exclusion():
    """SS17.7.6: Band calculation should exclude lastRow from banding.

    4-row table with tblLook: firstRow=true, lastRow=true, noHBand=false.
    Table style has band1Horz (gray D9D9D9) and band2Horz (white FFFFFF).
    With firstRow and lastRow both enabled:
    - Row 0: firstRow conditional (not banded)
    - Row 1: data row 0 → band1Horz (gray)
    - Row 2: data row 1 → band2Horz (white) — but since lastRow, should NOT be banded
    - Row 3: lastRow conditional (not banded)

    In a 4-row table with firstRow+lastRow enabled, rows 1 and 2 are data rows.
    Row 1 = data_row_idx 0 (band1), Row 2 = data_row_idx 1 (band2).
    But Row 2 is also not the last row (row 3 is), so banding SHOULD apply to row 2.
    """
    doc = Document()

    # Create a custom table style with banding
    styles_part = doc.part.element.find(w("body")).getparent()
    styles_el = styles_part.find(w("styles"))
    if styles_el is None:
        # Try the styles part
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        styles_part_obj = doc.part.part_related_by(RT.STYLES)
        styles_el = styles_part_obj.element

    # Define a table style with banding
    style = make_element("w:style", {"w:type": "table", "w:styleId": "BandingTestStyle"})
    style.append(make_element("w:name", {"w:val": "Banding Test Style"}))
    style.append(make_element("w:basedOn", {"w:val": "TableNormal"}))

    # Base table properties with band size = 1
    style_tbl_pr = make_element("w:tblPr")
    style_tbl_pr.append(make_element("w:tblStyleRowBandSize", {"w:val": "1"}))
    tbl_borders = make_element("w:tblBorders")
    for edge in ["top", "bottom", "left", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{edge}", {
            "w:val": "single", "w:sz": "4", "w:color": "000000", "w:space": "0",
        }))
    style_tbl_pr.append(tbl_borders)
    style.append(style_tbl_pr)

    # firstRow conditional: red shading
    fr = make_element("w:tblStylePr", {"w:type": "firstRow"})
    fr_tc = make_element("w:tcPr")
    fr_tc.append(make_element("w:shd", {
        "w:val": "clear", "w:color": "auto", "w:fill": "FF0000",
    }))
    fr.append(fr_tc)
    style.append(fr)

    # lastRow conditional: blue shading
    lr = make_element("w:tblStylePr", {"w:type": "lastRow"})
    lr_tc = make_element("w:tcPr")
    lr_tc.append(make_element("w:shd", {
        "w:val": "clear", "w:color": "auto", "w:fill": "0000FF",
    }))
    lr.append(lr_tc)
    style.append(lr)

    # band1Horz: gray shading
    b1 = make_element("w:tblStylePr", {"w:type": "band1Horz"})
    b1_tc = make_element("w:tcPr")
    b1_tc.append(make_element("w:shd", {
        "w:val": "clear", "w:color": "auto", "w:fill": "D9D9D9",
    }))
    b1.append(b1_tc)
    style.append(b1)

    # band2Horz: light yellow shading
    b2 = make_element("w:tblStylePr", {"w:type": "band2Horz"})
    b2_tc = make_element("w:tcPr")
    b2_tc.append(make_element("w:shd", {
        "w:val": "clear", "w:color": "auto", "w:fill": "FFFFCC",
    }))
    b2.append(b2_tc)
    style.append(b2)

    styles_el.append(style)

    # Create 4-row table using this style
    tbl = doc.add_table(rows=4, cols=2)
    for r in range(4):
        for c in range(2):
            tbl.cell(r, c).text = f"R{r}C{c}"

    tbl_el = tbl._tbl
    tblPr = tbl_el.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_el.insert(0, tblPr)

    # Remove any tblLook/tblStyle python-docx may have added automatically
    for existing in tblPr.findall(w("tblLook")):
        tblPr.remove(existing)
    for existing in tblPr.findall(w("tblStyle")):
        tblPr.remove(existing)

    # Reference the style (insert at beginning per CT_TblPr sequence)
    tbl_style_el = make_element("w:tblStyle", {"w:val": "BandingTestStyle"})
    tblPr.insert(0, tbl_style_el)

    # tblLook: firstRow=true, lastRow=true, noHBand=false (enable banding)
    tblPr.append(make_element("w:tblLook", {
        "w:val": "00A0",
        "w:firstRow": "1",
        "w:lastRow": "1",
        "w:firstColumn": "0",
        "w:lastColumn": "0",
        "w:noHBand": "0",
        "w:noVBand": "1",
    }))

    save_fixture("banding-last-row-exclusion", doc, {
        "name": "banding-last-row-exclusion",
        "spec_ref": "ISO 29500-1 SS17.7.6",
        "description": (
            "4-row table with firstRow+lastRow+banding. "
            "Row 0: firstRow (red). Row 3: lastRow (blue). "
            "Rows 1-2: data rows with band1/band2 alternation. "
            "Tests that data_row_idx is correctly offset for firstRow."
        ),
    })


# =====================================================================
# Fixture 9: Adjacent cell border conflict (cell-to-cell)
# =====================================================================

def make_adjacent_cell_border_conflict():
    """SS17.4.66 rule 3: When two adjacent cells both have borders, the
    one with higher weight wins.

    2x2 table. Cell (0,0) right border: thick sz=24 red.
    Cell (0,1) left border: single sz=4 blue.
    The shared edge should display cell (0,0)'s border (higher weight).
    """
    doc = Document()
    doc.add_paragraph("Adjacent cell border conflict test.")

    tbl_el, tblPr = make_table_with_borders(doc, 1, 2, [
        ["ThickRight", "ThinLeft"],
    ])

    # No table-level borders (only cell borders)
    # Cell (0,0): thick right border
    add_tc_borders(tbl_el, 0, 0, ["right"], val="single", sz="24", color="FF0000")
    # Cell (0,1): thin left border
    add_tc_borders(tbl_el, 0, 1, ["left"], val="single", sz="4", color="0000FF")

    save_fixture("adjacent-cell-border-conflict", doc, {
        "name": "adjacent-cell-border-conflict",
        "spec_ref": "ISO 29500-1 SS17.4.66 rule 3",
        "description": (
            "Cell (0,0) has thick right border (sz=24 red). "
            "Cell (0,1) has thin left border (sz=4 blue). "
            "Per SS17.4.66 rule 3, the heavier border wins at the shared edge."
        ),
    })


# =====================================================================
# Fixture 10: tblLook legacy val bitmask
# =====================================================================

def make_tbllook_legacy_bitmask():
    """MS-OI29500 SS17.4.55(c): When no individual attributes are present,
    Word falls back to the w:val hex bitmask.

    val="04A0" means: firstRow(0x0020)=true, noHBand(0x0200)=true,
    noVBand(0x0400)=true, others=false.

    We create a table with ONLY w:val="04A0" and no individual attributes.
    """
    doc = Document()
    doc.add_paragraph("tblLook legacy val bitmask test.")

    # Create a table style to test against
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    styles_part_obj = doc.part.part_related_by(RT.STYLES)
    styles_el = styles_part_obj.element

    style = make_element("w:style", {"w:type": "table", "w:styleId": "LegacyLookStyle"})
    style.append(make_element("w:name", {"w:val": "Legacy Look Style"}))
    style.append(make_element("w:basedOn", {"w:val": "TableNormal"}))

    style_tbl_pr = make_element("w:tblPr")
    style_tbl_pr.append(make_element("w:tblStyleRowBandSize", {"w:val": "1"}))
    style.append(style_tbl_pr)

    # firstRow conditional: red shading
    fr = make_element("w:tblStylePr", {"w:type": "firstRow"})
    fr_tc = make_element("w:tcPr")
    fr_tc.append(make_element("w:shd", {
        "w:val": "clear", "w:color": "auto", "w:fill": "FF0000",
    }))
    fr.append(fr_tc)
    style.append(fr)

    # lastRow conditional: blue shading
    lr = make_element("w:tblStylePr", {"w:type": "lastRow"})
    lr_tc = make_element("w:tcPr")
    lr_tc.append(make_element("w:shd", {
        "w:val": "clear", "w:color": "auto", "w:fill": "0000FF",
    }))
    lr.append(lr_tc)
    style.append(lr)

    styles_el.append(style)

    tbl = doc.add_table(rows=3, cols=2)
    for r in range(3):
        for c in range(2):
            tbl.cell(r, c).text = f"R{r}C{c}"

    tbl_el = tbl._tbl
    tblPr = tbl_el.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_el.insert(0, tblPr)

    tblPr.append(make_element("w:tblStyle", {"w:val": "LegacyLookStyle"}))

    # Remove any tblLook python-docx may have added automatically
    for existing in tblPr.findall(w("tblLook")):
        tblPr.remove(existing)

    # ONLY use w:val bitmask, no individual attributes
    # 0x04A0 = 0100_1010_0000 (binary)
    # 0x0020 = bit 5 = firstRow -> (0x04A0 & 0x0020) = 0x0020 -> true
    # 0x0040 = bit 6 = lastRow -> (0x04A0 & 0x0040) = 0 -> false
    # 0x0080 = bit 7 = firstCol -> (0x04A0 & 0x0080) = 0x0080 -> true
    # 0x0100 = bit 8 = lastCol -> (0x04A0 & 0x0100) = 0 -> false
    # 0x0200 = bit 9 = noHBand -> (0x04A0 & 0x0200) = 0x0200 -> true
    # 0x0400 = bit 10 = noVBand -> (0x04A0 & 0x0400) = 0x0400 -> true
    tblPr.append(make_element("w:tblLook", {"w:val": "04A0"}))

    save_fixture("tbllook-legacy-bitmask", doc, {
        "name": "tbllook-legacy-bitmask",
        "spec_ref": "MS-OI29500 SS17.4.55(c)",
        "description": (
            "tblLook with only w:val='04A0' bitmask (no individual attributes). "
            "Expected: firstRow=true, lastRow=false, firstCol=true, "
            "lastCol=false, noHBand=false, noVBand=true."
        ),
    })


# =====================================================================
# Fixture 11: Conditional formatting border override
# =====================================================================

def make_conditional_border_override():
    """SS17.7.6: Conditional formatting borders should override table-level
    borders for matching cells.

    Table style has thin black borders. firstRow conditional has thick red
    borders. The first row cells should display thick red, not thin black.
    """
    doc = Document()
    doc.add_paragraph("Conditional formatting border override test.")

    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    styles_part_obj = doc.part.part_related_by(RT.STYLES)
    styles_el = styles_part_obj.element

    style = make_element("w:style", {"w:type": "table", "w:styleId": "CondBorderStyle"})
    style.append(make_element("w:name", {"w:val": "Cond Border Style"}))
    style.append(make_element("w:basedOn", {"w:val": "TableNormal"}))

    style_tbl_pr = make_element("w:tblPr")
    tbl_borders = make_element("w:tblBorders")
    for edge in ["top", "bottom", "left", "right", "insideH", "insideV"]:
        tbl_borders.append(make_element(f"w:{edge}", {
            "w:val": "single", "w:sz": "4", "w:color": "000000", "w:space": "0",
        }))
    style_tbl_pr.append(tbl_borders)
    style.append(style_tbl_pr)

    # firstRow: thick red borders
    fr = make_element("w:tblStylePr", {"w:type": "firstRow"})
    fr_tc = make_element("w:tcPr")
    fr_borders = make_element("w:tcBorders")
    for edge in ["top", "bottom", "left", "right"]:
        fr_borders.append(make_element(f"w:{edge}", {
            "w:val": "single", "w:sz": "12", "w:color": "FF0000", "w:space": "0",
        }))
    fr_tc.append(fr_borders)
    fr.append(fr_tc)
    style.append(fr)

    styles_el.append(style)

    tbl = doc.add_table(rows=2, cols=2)
    for r in range(2):
        for c in range(2):
            tbl.cell(r, c).text = f"R{r}C{c}"

    tbl_el = tbl._tbl
    tblPr = tbl_el.tblPr
    if tblPr is None:
        tblPr = make_element("w:tblPr")
        tbl_el.insert(0, tblPr)

    # Remove any tblLook/tblStyle python-docx may have added automatically
    for existing in tblPr.findall(w("tblLook")):
        tblPr.remove(existing)
    for existing in tblPr.findall(w("tblStyle")):
        tblPr.remove(existing)

    tbl_style_el = make_element("w:tblStyle", {"w:val": "CondBorderStyle"})
    tblPr.insert(0, tbl_style_el)
    tblPr.append(make_element("w:tblLook", {
        "w:val": "0020",
        "w:firstRow": "1",
        "w:lastRow": "0",
        "w:firstColumn": "0",
        "w:lastColumn": "0",
        "w:noHBand": "1",
        "w:noVBand": "1",
    }))

    save_fixture("conditional-border-override", doc, {
        "name": "conditional-border-override",
        "spec_ref": "ISO 29500-1 SS17.7.6",
        "description": (
            "Table style has thin black borders. firstRow conditional has "
            "thick red borders (sz=12). First row cells should have red borders."
        ),
    })


# =====================================================================
# Main
# =====================================================================

def main():
    print("\n== Table Borders Audit Fixtures ==")
    make_table_border_fallback_no_cell_borders()
    make_inside_v_interior_cells()
    make_cell_border_nil_override()
    make_cell_margin_per_side_cascade()
    make_width_type_nil()
    make_vmerge_continue_no_restart()
    make_gridspan_overflow()
    make_banding_last_row_exclusion()
    make_adjacent_cell_border_conflict()
    make_tbllook_legacy_bitmask()
    make_conditional_border_override()
    print("\nDone.")


if __name__ == "__main__":
    main()

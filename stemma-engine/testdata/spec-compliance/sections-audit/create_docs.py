# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for deep section-properties audit (ISO 29500-1 §17.6).

These fixtures probe section property features that are NOT covered by the
existing spec_sections.rs or spec_section_compat_audit.rs tests.

Run:  uv run create_docs.py
"""

import json
from pathlib import Path
from lxml import etree

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc: DocxDocument, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  sections-audit/{name}/")


# ═══════════════════════════════════════════════════════════════════════
# 1. titlePg — Title page flag (§17.6.18)
# ═══════════════════════════════════════════════════════════════════════

def make_title_page() -> None:
    """w:titlePg — enable distinct first-page header/footer.

    ISO 29500-1 §17.6.18: When present, the section uses a different
    header and footer for the first page.
    """
    doc = Document()

    doc.add_paragraph("Content on first page of section with titlePg enabled.")

    p = doc.add_paragraph("Last paragraph of section.")
    pPr = p._p.get_or_add_pPr()
    sect_pr = make_element("w:sectPr")
    sect_pr.append(make_element("w:titlePg"))
    sect_pr.append(make_element("w:pgSz", {"w:w": "12240", "w:h": "15840"}))
    pPr.append(sect_pr)

    doc.add_paragraph("Content after section break.")

    save_fixture("title-page", doc, {
        "name": "title-page",
        "spec_ref": "ISO 29500-1 §17.6.18",
        "description": "Section with w:titlePg element (distinct first-page header/footer)",
        "expected_behavior": "SectionProperties should expose titlePg boolean",
    })


# ═══════════════════════════════════════════════════════════════════════
# 2. bidi — Bidirectional section (§17.6.1)
# ═══════════════════════════════════════════════════════════════════════

def make_bidi_section() -> None:
    """w:bidi — right-to-left section.

    ISO 29500-1 §17.6.1: When present, the section is right-to-left.
    """
    doc = Document()

    doc.add_paragraph("Content in a bidirectional section.")

    p = doc.add_paragraph("Last paragraph.")
    pPr = p._p.get_or_add_pPr()
    sect_pr = make_element("w:sectPr")
    sect_pr.append(make_element("w:bidi"))
    sect_pr.append(make_element("w:pgSz", {"w:w": "12240", "w:h": "15840"}))
    pPr.append(sect_pr)

    doc.add_paragraph("Content after bidi section.")

    save_fixture("bidi-section", doc, {
        "name": "bidi-section",
        "spec_ref": "ISO 29500-1 §17.6.1",
        "description": "Section with w:bidi element (right-to-left layout)",
        "expected_behavior": "SectionProperties should expose bidi boolean",
    })


# ═══════════════════════════════════════════════════════════════════════
# 3. formProt — Form protection (§17.6.6)
# ═══════════════════════════════════════════════════════════════════════

def make_form_prot() -> None:
    """w:formProt — section-level form protection.

    ISO 29500-1 §17.6.6: When present, only form fields in this
    section can be edited.
    """
    doc = Document()

    doc.add_paragraph("Content in a form-protected section.")

    p = doc.add_paragraph("Last paragraph.")
    pPr = p._p.get_or_add_pPr()
    sect_pr = make_element("w:sectPr")
    sect_pr.append(make_element("w:formProt", {"w:val": "1"}))
    sect_pr.append(make_element("w:pgSz", {"w:w": "12240", "w:h": "15840"}))
    pPr.append(sect_pr)

    doc.add_paragraph("Content after form-protected section.")

    save_fixture("form-prot", doc, {
        "name": "form-prot",
        "spec_ref": "ISO 29500-1 §17.6.6",
        "description": "Section with w:formProt val=1 (form protection enabled)",
        "expected_behavior": "SectionProperties should expose formProt boolean",
    })


# ═══════════════════════════════════════════════════════════════════════
# 4. noEndnote — Suppress endnotes (§17.6.9)
# ═══════════════════════════════════════════════════════════════════════

def make_no_endnote() -> None:
    """w:noEndnote — suppress endnotes in this section.

    ISO 29500-1 §17.6.9: When present, endnotes are not displayed
    at the end of this section.
    """
    doc = Document()

    doc.add_paragraph("Content in a section with suppressed endnotes.")

    p = doc.add_paragraph("Last paragraph.")
    pPr = p._p.get_or_add_pPr()
    sect_pr = make_element("w:sectPr")
    sect_pr.append(make_element("w:noEndnote"))
    sect_pr.append(make_element("w:pgSz", {"w:w": "12240", "w:h": "15840"}))
    pPr.append(sect_pr)

    doc.add_paragraph("Content after noEndnote section.")

    save_fixture("no-endnote", doc, {
        "name": "no-endnote",
        "spec_ref": "ISO 29500-1 §17.6.9",
        "description": "Section with w:noEndnote (endnotes suppressed)",
        "expected_behavior": "SectionProperties should expose noEndnote boolean",
    })


# ═══════════════════════════════════════════════════════════════════════
# 5. pgSz w:code — Paper size code (§17.6.14)
# ═══════════════════════════════════════════════════════════════════════

def make_paper_size_code() -> None:
    """w:pgSz w:code — paper size code attribute.

    ISO 29500-1 §17.6.14: The code attribute on pgSz specifies a
    standard paper size (e.g., 1=Letter, 9=A4). This is in addition
    to the explicit w/h dimensions.
    """
    doc = Document()

    doc.add_paragraph("Content in a section with paper size code A4.")

    p = doc.add_paragraph("Last paragraph.")
    pPr = p._p.get_or_add_pPr()
    sect_pr = make_element("w:sectPr")
    # A4: 11906 x 16838 twips, code=9
    sect_pr.append(make_element("w:pgSz", {
        "w:w": "11906",
        "w:h": "16838",
        "w:code": "9",
    }))
    pPr.append(sect_pr)

    doc.add_paragraph("Content after A4 section.")

    save_fixture("paper-size-code", doc, {
        "name": "paper-size-code",
        "spec_ref": "ISO 29500-1 §17.6.14",
        "description": "Section with pgSz code=9 (A4 paper size)",
        "expected_behavior": "SectionProperties should expose paper size code",
    })


# ═══════════════════════════════════════════════════════════════════════
# 6. cols w:sep — Column separator line (§17.6.4)
# ═══════════════════════════════════════════════════════════════════════

def make_column_separator() -> None:
    """w:cols w:sep — column separator line.

    ISO 29500-1 §17.6.4: The sep attribute on cols draws a vertical
    line between columns.
    """
    doc = Document()

    doc.add_paragraph("Content in a section with column separator.")

    p = doc.add_paragraph("Last paragraph.")
    pPr = p._p.get_or_add_pPr()
    sect_pr = make_element("w:sectPr")
    sect_pr.append(make_element("w:cols", {
        "w:num": "2",
        "w:space": "720",
        "w:sep": "1",
    }))
    sect_pr.append(make_element("w:pgSz", {"w:w": "12240", "w:h": "15840"}))
    pPr.append(sect_pr)

    doc.add_paragraph("Content after columns section.")

    save_fixture("column-separator", doc, {
        "name": "column-separator",
        "spec_ref": "ISO 29500-1 §17.6.4",
        "description": "Section with 2 columns and w:sep=1 (separator line between columns)",
        "expected_behavior": "SectionProperties should expose column separator boolean",
    })


# ═══════════════════════════════════════════════════════════════════════
# 7. pgNumType chapStyle/chapSep — Chapter numbering (§17.6.12)
# ═══════════════════════════════════════════════════════════════════════

def make_chapter_page_numbering() -> None:
    """w:pgNumType chapStyle/chapSep — chapter-relative page numbering.

    ISO 29500-1 §17.6.12: The chapStyle and chapSep attributes on
    pgNumType enable chapter-relative page numbering (e.g., "2-3"
    for page 3 of chapter 2).
    """
    doc = Document()

    doc.add_paragraph("Content with chapter page numbering.")

    p = doc.add_paragraph("Last paragraph.")
    pPr = p._p.get_or_add_pPr()
    sect_pr = make_element("w:sectPr")
    sect_pr.append(make_element("w:pgNumType", {
        "w:fmt": "decimal",
        "w:start": "1",
        "w:chapStyle": "1",
        "w:chapSep": "hyphen",
    }))
    sect_pr.append(make_element("w:pgSz", {"w:w": "12240", "w:h": "15840"}))
    pPr.append(sect_pr)

    doc.add_paragraph("Content after chapter numbering section.")

    save_fixture("chapter-page-numbering", doc, {
        "name": "chapter-page-numbering",
        "spec_ref": "ISO 29500-1 §17.6.12",
        "description": "Section with pgNumType chapStyle=1, chapSep=hyphen (chapter-relative page numbering)",
        "expected_behavior": "PageNumberType should expose chapStyle and chapSep",
    })


# ═══════════════════════════════════════════════════════════════════════
# 8. oddPage/evenPage section types (§17.6.17)
# ═══════════════════════════════════════════════════════════════════════

def make_odd_even_page_sections() -> None:
    """oddPage and evenPage section break types.

    ISO 29500-1 §17.6.17: oddPage forces the next section onto
    an odd page; evenPage forces it onto an even page.
    """
    doc = Document()

    doc.add_paragraph("Content before oddPage break.")
    p1 = doc.add_paragraph("End of section 1.")
    pPr1 = p1._p.get_or_add_pPr()
    sect_pr1 = make_element("w:sectPr")
    sect_pr1.append(make_element("w:type", {"w:val": "oddPage"}))
    sect_pr1.append(make_element("w:pgSz", {"w:w": "12240", "w:h": "15840"}))
    pPr1.append(sect_pr1)

    doc.add_paragraph("Content before evenPage break.")
    p2 = doc.add_paragraph("End of section 2.")
    pPr2 = p2._p.get_or_add_pPr()
    sect_pr2 = make_element("w:sectPr")
    sect_pr2.append(make_element("w:type", {"w:val": "evenPage"}))
    sect_pr2.append(make_element("w:pgSz", {"w:w": "12240", "w:h": "15840"}))
    pPr2.append(sect_pr2)

    doc.add_paragraph("Content in final section.")

    save_fixture("odd-even-page-sections", doc, {
        "name": "odd-even-page-sections",
        "spec_ref": "ISO 29500-1 §17.6.17",
        "description": "Two sections: oddPage and evenPage break types",
        "expected_behavior": "section_type should be 'oddPage' and 'evenPage' respectively",
    })


# ═══════════════════════════════════════════════════════════════════════
# 9. Body-level section properties (§17.6)
# ═══════════════════════════════════════════════════════════════════════

def make_body_section_props() -> None:
    """Body-level w:sectPr — final section properties on w:body.

    ISO 29500-1 §17.6: The last section in a document is defined by
    w:body/w:sectPr (not by pPr/sectPr). This tests that our parser
    captures body-level section properties into CanonDoc.body_section_properties.
    """
    doc = Document()

    doc.add_paragraph("Content in the only section of the document.")
    doc.add_paragraph("Second paragraph.")

    # python-docx creates body-level sectPr automatically. We modify it.
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is None:
        sect_pr = make_element("w:sectPr")
        body.append(sect_pr)

    # Set custom margins on body-level sectPr
    pg_mar = make_element("w:pgMar", {
        "w:top": "2160",
        "w:bottom": "2160",
        "w:left": "1440",
        "w:right": "1440",
        "w:header": "720",
        "w:footer": "720",
        "w:gutter": "0",
    })
    # Remove existing pgMar if any
    existing_pg_mar = sect_pr.find(qn("w:pgMar"))
    if existing_pg_mar is not None:
        sect_pr.remove(existing_pg_mar)
    sect_pr.append(pg_mar)

    # Add vAlign to body-level section
    v_align = make_element("w:vAlign", {"w:val": "both"})
    sect_pr.append(v_align)

    save_fixture("body-section-props", doc, {
        "name": "body-section-props",
        "spec_ref": "ISO 29500-1 §17.6",
        "description": "Document with custom body-level sectPr (margins and vAlign on w:body/w:sectPr)",
        "expected_behavior": "CanonDoc.body_section_properties should have margins and vAlign='both'",
    })


# ═══════════════════════════════════════════════════════════════════════
# 10. cols equalWidth=false with col defs (§17.6.4)
# ═══════════════════════════════════════════════════════════════════════

def make_unequal_columns_mid_doc() -> None:
    """w:cols equalWidth=false — unequal column widths in mid-document section.

    ISO 29500-1 §17.6.4: When equalWidth is false, individual col
    elements define per-column widths and spacing.
    """
    doc = Document()

    doc.add_paragraph("Content in unequal-columns section.")

    p = doc.add_paragraph("Last paragraph.")
    pPr = p._p.get_or_add_pPr()
    sect_pr = make_element("w:sectPr")

    cols = make_element("w:cols", {
        "w:num": "3",
        "w:space": "0",
        "w:equalWidth": "0",
    })
    cols.append(make_element("w:col", {"w:w": "2400", "w:space": "480"}))
    cols.append(make_element("w:col", {"w:w": "4800", "w:space": "480"}))
    cols.append(make_element("w:col", {"w:w": "2400"}))  # no space = 0
    sect_pr.append(cols)
    sect_pr.append(make_element("w:pgSz", {"w:w": "12240", "w:h": "15840"}))
    pPr.append(sect_pr)

    doc.add_paragraph("Content after unequal columns.")

    save_fixture("unequal-columns", doc, {
        "name": "unequal-columns",
        "spec_ref": "ISO 29500-1 §17.6.4",
        "description": "Section with 3 unequal columns (2400+480, 4800+480, 2400+0)",
        "expected_behavior": "SectionProperties.column_defs should have 3 entries with correct widths",
    })


# ═══════════════════════════════════════════════════════════════════════
# 11. Combined section with many properties (integration)
# ═══════════════════════════════════════════════════════════════════════

def make_combined_section_props() -> None:
    """Section with many properties combined — integration test.

    Tests that multiple section properties coexist correctly in a
    single sectPr element without interfering with each other.
    """
    doc = Document()

    doc.add_paragraph("Content in section with many properties.")

    p = doc.add_paragraph("End of multi-property section.")
    pPr = p._p.get_or_add_pPr()
    sect_pr = make_element("w:sectPr")

    # Page size: A4 landscape
    sect_pr.append(make_element("w:pgSz", {
        "w:w": "16838",
        "w:h": "11906",
        "w:orient": "landscape",
    }))

    # Margins
    sect_pr.append(make_element("w:pgMar", {
        "w:top": "1440",
        "w:bottom": "1440",
        "w:left": "1800",
        "w:right": "1800",
        "w:header": "720",
        "w:footer": "720",
        "w:gutter": "0",
    }))

    # Section type
    sect_pr.append(make_element("w:type", {"w:val": "continuous"}))

    # 2 columns
    sect_pr.append(make_element("w:cols", {"w:num": "2", "w:space": "720"}))

    # Vertical alignment
    sect_pr.append(make_element("w:vAlign", {"w:val": "center"}))

    # Page number format
    sect_pr.append(make_element("w:pgNumType", {"w:fmt": "upperLetter", "w:start": "1"}))

    # Document grid
    sect_pr.append(make_element("w:docGrid", {"w:linePitch": "360"}))

    pPr.append(sect_pr)

    doc.add_paragraph("Content in next section.")

    save_fixture("combined-section-props", doc, {
        "name": "combined-section-props",
        "spec_ref": "ISO 29500-1 §17.6",
        "description": "Section combining landscape A4, margins, continuous type, 2 columns, vAlign=center, pgNumType, docGrid",
        "expected_behavior": "All properties should be independently accessible in SectionProperties",
    })


# ═══════════════════════════════════════════════════════════════════════
# Main
# ═══════════════════════════════════════════════════════════════════════

def main() -> None:
    print("── Sections Deep Audit ──")
    make_title_page()
    make_bidi_section()
    make_form_prot()
    make_no_endnote()
    make_paper_size_code()
    make_column_separator()
    make_chapter_page_numbering()
    make_odd_even_page_sections()
    make_body_section_props()
    make_unequal_columns_mid_doc()
    make_combined_section_props()


if __name__ == "__main__":
    main()

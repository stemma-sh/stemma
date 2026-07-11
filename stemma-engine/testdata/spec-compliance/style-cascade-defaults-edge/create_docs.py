# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for docDefaults / Normal / style resolution edge cases.

Targets edge cases in:
  1. docDefaults pPr alignment applied to unstyled paragraphs
  2. docDefaults pPr indentation applied to unstyled paragraphs
  3. docDefaults pPr spacing + style spacing per-field merge
  4. Normal style with basedOn chain
  5. Multiple paragraph styles with w:default="1"
  6. basedOn pointing to non-existent styleId
  7. Circular basedOn chain (A -> B -> A)
  8. Cross-type basedOn (paragraph style basedOn character style)

Run:  python create_docs.py
"""

import json
from pathlib import Path
from lxml import etree

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  style-cascade-defaults-edge/{name}/")


def remove_pstyle(para_element):
    """Remove any auto-added pStyle from a paragraph element."""
    ppr = para_element.find(qn("w:pPr"))
    if ppr is not None:
        ps = ppr.find(qn("w:pStyle"))
        if ps is not None:
            ppr.remove(ps)


def clear_default_styles(styles_el):
    """Remove default="1" from all existing paragraph styles to control which is default."""
    for style in styles_el.findall(qn("w:style")):
        if style.get(qn("w:type")) == "paragraph":
            if style.get(qn("w:default")) in ("1", "true"):
                style.attrib.pop(qn("w:default"), None)


def get_or_create_ppr_defaults(styles_el):
    """Get or create the docDefaults/pPrDefault/pPr element, reusing existing one."""
    doc_defaults = styles_el.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = make_element("w:docDefaults")
        styles_el.insert(0, doc_defaults)

    ppr_default_el = doc_defaults.find(qn("w:pPrDefault"))
    if ppr_default_el is None:
        ppr_default_el = make_element("w:pPrDefault")
        doc_defaults.append(ppr_default_el)

    ppr = ppr_default_el.find(qn("w:pPr"))
    if ppr is None:
        ppr = make_element("w:pPr")
        ppr_default_el.append(ppr)

    return ppr


def clear_ppr_defaults(styles_el):
    """Remove any existing pPrDefault content so we can set it fresh."""
    doc_defaults = styles_el.find(qn("w:docDefaults"))
    if doc_defaults is not None:
        ppr_default_el = doc_defaults.find(qn("w:pPrDefault"))
        if ppr_default_el is not None:
            doc_defaults.remove(ppr_default_el)


# =========================================================================
# Fixture 1: docDefaults pPr alignment
#
# §17.7.5: docDefaults/pPrDefault/pPr can set alignment (jc=center).
# An unstyled paragraph should get this alignment as its default.
# A paragraph with a style that sets a different alignment should override.
# =========================================================================

def make_docdefaults_ppr_alignment():
    doc = Document()
    styles_el = doc.styles.element

    # Remove default="1" from Normal so docDefaults is the only source
    clear_default_styles(styles_el)

    # Clear existing pPrDefault and create fresh with jc=center
    clear_ppr_defaults(styles_el)
    ppr = get_or_create_ppr_defaults(styles_el)
    ppr.append(make_element("w:jc", {"w:val": "center"}))

    # LeftAlignStyle: paragraph style with jc=left
    style_left = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "LeftAlignStyle",
    })
    style_left.append(make_element("w:name", {"w:val": "Left Align Style"}))
    left_ppr = make_element("w:pPr")
    left_ppr.append(make_element("w:jc", {"w:val": "left"}))
    style_left.append(left_ppr)
    styles_el.append(style_left)

    # Para 1: unstyled — should get center from docDefaults pPr
    p1 = doc.add_paragraph()
    remove_pstyle(p1._element)
    p1.add_run("Unstyled center from docDefaults")

    # Para 2: LeftAlignStyle — should get left from style
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    # Remove any existing pStyle
    ps = p2_ppr.find(qn("w:pStyle"))
    if ps is not None:
        p2_ppr.remove(ps)
    p2_ppr.append(make_element("w:pStyle", {"w:val": "LeftAlignStyle"}))
    p2.add_run("Left from style")

    # Para 3: unstyled with direct jc=right — direct should win
    p3 = doc.add_paragraph()
    remove_pstyle(p3._element)
    p3_ppr = p3._element.get_or_add_pPr()
    p3_ppr.append(make_element("w:jc", {"w:val": "right"}))
    p3.add_run("Direct right overrides docDefaults")

    save_fixture("docdefaults-ppr-alignment", doc, {
        "spec": "ISO 29500-1 §17.7.5",
        "description": (
            "docDefaults pPr sets jc=center. Tests that unstyled paragraphs "
            "get center alignment from docDefaults, and that styles/direct "
            "formatting properly override."
        ),
        "expected": {
            "para_1": "alignment=center (from docDefaults pPr)",
            "para_2": "alignment=left (from LeftAlignStyle)",
            "para_3": "alignment=right (from direct formatting)",
        },
    })


# =========================================================================
# Fixture 2: docDefaults pPr indentation
#
# §17.7.5: docDefaults/pPrDefault/pPr can set indentation.
# An unstyled paragraph should inherit indent from docDefaults.
# =========================================================================

def make_docdefaults_ppr_indent():
    doc = Document()
    styles_el = doc.styles.element

    clear_default_styles(styles_el)

    # Clear existing pPrDefault and create fresh with indent
    clear_ppr_defaults(styles_el)
    ppr = get_or_create_ppr_defaults(styles_el)
    ppr.append(make_element("w:ind", {"w:left": "720", "w:right": "360"}))

    # Para 1: unstyled — should get indent from docDefaults
    p1 = doc.add_paragraph()
    remove_pstyle(p1._element)
    p1.add_run("Unstyled indent from docDefaults")

    # Para 2: unstyled with direct left=0 — direct should win for left,
    #          right should still come from docDefaults
    p2 = doc.add_paragraph()
    remove_pstyle(p2._element)
    p2_ppr = p2._element.get_or_add_pPr()
    p2_ppr.append(make_element("w:ind", {"w:left": "0"}))
    p2.add_run("Direct left zero overrides docDefaults")

    save_fixture("docdefaults-ppr-indent", doc, {
        "spec": "ISO 29500-1 §17.7.5",
        "description": (
            "docDefaults pPr sets ind left=720 right=360. Tests that unstyled "
            "paragraphs get indent from docDefaults, and that direct formatting "
            "overrides per-attribute."
        ),
        "expected": {
            "para_1": "indent left=720 right=360 (from docDefaults pPr)",
            "para_2": "indent left=0 (direct) right=360 (from docDefaults pPr)",
        },
    })


# =========================================================================
# Fixture 3: docDefaults pPr spacing + style spacing per-field merge
#
# §17.7.5 + §17.7.2: docDefaults provides base spacing. A style that only
# sets spacing_after should not obliterate spacing_before from docDefaults.
# =========================================================================

def make_docdefaults_ppr_spacing_merge():
    doc = Document()
    styles_el = doc.styles.element

    clear_default_styles(styles_el)

    # Clear existing pPrDefault and create fresh with spacing
    clear_ppr_defaults(styles_el)
    ppr = get_or_create_ppr_defaults(styles_el)
    ppr.append(make_element("w:spacing", {
        "w:before": "240",
        "w:line": "276",
    }))

    # AfterOnlyStyle: paragraph style, sets only spacing after=120
    style_after = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "AfterOnlyStyle",
    })
    style_after.append(make_element("w:name", {"w:val": "After Only Style"}))
    after_ppr = make_element("w:pPr")
    after_ppr.append(make_element("w:spacing", {"w:after": "120"}))
    style_after.append(after_ppr)
    styles_el.append(style_after)

    # Para 1: unstyled — should get before=240, line=276 from docDefaults
    p1 = doc.add_paragraph()
    remove_pstyle(p1._element)
    p1.add_run("Unstyled spacing from docDefaults")

    # Para 2: AfterOnlyStyle — should get before=240 from docDefaults (per-field merge)
    #          + after=120 from style + line=276 from docDefaults
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    ps = p2_ppr.find(qn("w:pStyle"))
    if ps is not None:
        p2_ppr.remove(ps)
    p2_ppr.append(make_element("w:pStyle", {"w:val": "AfterOnlyStyle"}))
    p2.add_run("Style after only merges with docDefaults")

    save_fixture("docdefaults-ppr-spacing-merge", doc, {
        "spec": "ISO 29500-1 §17.7.5 + §17.7.2",
        "description": (
            "docDefaults pPr sets spacing before=240 line=276. Style sets only "
            "after=120. Tests per-field merge: style after + docDefaults before and line."
        ),
        "expected": {
            "para_1": "before=240, line=276 (from docDefaults pPr)",
            "para_2": "before=240 (docDefaults), after=120 (style), line=276 (docDefaults)",
        },
    })


# =========================================================================
# Fixture 4: Normal style with basedOn chain
#
# §17.7.4.3: Normal is typically a root style, but it's valid for Normal
# to have basedOn pointing to another paragraph style. Properties from
# the parent should be inherited.
# =========================================================================

def make_normal_with_based_on():
    doc = Document()
    styles_el = doc.styles.element

    clear_default_styles(styles_el)

    # CustomBase: paragraph style with font_size=24 and color=FF0000
    style_base = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "CustomBase",
    })
    style_base.append(make_element("w:name", {"w:val": "Custom Base"}))
    base_rpr = make_element("w:rPr")
    base_rpr.append(make_element("w:sz", {"w:val": "24"}))
    base_rpr.append(make_element("w:color", {"w:val": "FF0000"}))
    style_base.append(base_rpr)
    base_ppr = make_element("w:pPr")
    base_ppr.append(make_element("w:spacing", {"w:after": "120"}))
    style_base.append(base_ppr)
    styles_el.append(style_base)

    # Normal: default paragraph style, basedOn CustomBase, adds font_size=28
    style_normal = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "Normal",
        "w:default": "1",
    })
    style_normal.append(make_element("w:name", {"w:val": "Normal"}))
    style_normal.append(make_element("w:basedOn", {"w:val": "CustomBase"}))
    normal_rpr = make_element("w:rPr")
    normal_rpr.append(make_element("w:sz", {"w:val": "28"}))
    style_normal.append(normal_rpr)
    styles_el.append(style_normal)

    # Para 1: unstyled (implicit Normal) — should get font_size=28 from Normal,
    # color=FF0000 from CustomBase (inherited via basedOn)
    p1 = doc.add_paragraph()
    remove_pstyle(p1._element)
    p1.add_run("Implicit Normal with basedOn")

    # Para 2: explicit pStyle=Normal
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    ps = p2_ppr.find(qn("w:pStyle"))
    if ps is not None:
        p2_ppr.remove(ps)
    p2_ppr.append(make_element("w:pStyle", {"w:val": "Normal"}))
    p2.add_run("Explicit Normal with basedOn")

    save_fixture("normal-with-based-on", doc, {
        "spec": "ISO 29500-1 §17.7.4.3 + §17.7.4.17",
        "description": (
            "Normal style (default paragraph style) has basedOn=CustomBase. "
            "CustomBase sets sz=24, color=FF0000. Normal overrides sz=28. "
            "Tests that Normal's basedOn chain is resolved correctly."
        ),
        "expected": {
            "para_1": "font_size=28 (Normal), color=FF0000 (CustomBase via basedOn), spacing_after=120 (CustomBase)",
            "para_2": "same as para_1",
        },
    })


# =========================================================================
# Fixture 5: Multiple paragraph styles with default="1"
#
# §17.7.4.17: Only one style per type should have default="1". But when
# two exist, behavior is implementation-defined. We test that the code
# doesn't crash and produces deterministic results.
# =========================================================================

def make_multiple_default_para_styles():
    doc = Document()
    styles_el = doc.styles.element

    clear_default_styles(styles_el)

    # FirstDefault: paragraph style with default="1", color=FF0000, sz=20
    style_first = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "FirstDefault",
        "w:default": "1",
    })
    style_first.append(make_element("w:name", {"w:val": "First Default"}))
    first_rpr = make_element("w:rPr")
    first_rpr.append(make_element("w:color", {"w:val": "FF0000"}))
    first_rpr.append(make_element("w:sz", {"w:val": "20"}))
    style_first.append(first_rpr)
    styles_el.append(style_first)

    # SecondDefault: paragraph style with default="1", color=0000FF, sz=28
    style_second = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "SecondDefault",
        "w:default": "1",
    })
    style_second.append(make_element("w:name", {"w:val": "Second Default"}))
    second_rpr = make_element("w:rPr")
    second_rpr.append(make_element("w:color", {"w:val": "0000FF"}))
    second_rpr.append(make_element("w:sz", {"w:val": "28"}))
    style_second.append(second_rpr)
    styles_el.append(style_second)

    # Para 1: unstyled — which default wins?
    p1 = doc.add_paragraph()
    remove_pstyle(p1._element)
    p1.add_run("Unstyled with two defaults")

    save_fixture("multiple-default-para-styles", doc, {
        "spec": "ISO 29500-1 §17.7.4.17",
        "description": (
            "Two paragraph styles both have w:default='1'. Per spec only one "
            "should exist. Tests that the code handles this gracefully and "
            "doesn't crash."
        ),
        "expected": {
            "para_1": "Properties from whichever default wins (implementation-defined order)",
        },
    })


# =========================================================================
# Fixture 6: basedOn pointing to non-existent styleId
#
# A style references basedOn with a styleId that doesn't exist in the
# document. The code should handle this gracefully.
# =========================================================================

def make_based_on_nonexistent():
    doc = Document()
    styles_el = doc.styles.element

    clear_default_styles(styles_el)

    # OrphanStyle: paragraph style with basedOn pointing to non-existent
    style_orphan = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "OrphanStyle",
    })
    style_orphan.append(make_element("w:name", {"w:val": "Orphan Style"}))
    style_orphan.append(make_element("w:basedOn", {"w:val": "NonExistentParent"}))
    orphan_rpr = make_element("w:rPr")
    orphan_rpr.append(make_element("w:sz", {"w:val": "28"}))
    orphan_rpr.append(make_element("w:color", {"w:val": "00FF00"}))
    style_orphan.append(orphan_rpr)
    styles_el.append(style_orphan)

    # Para 1: OrphanStyle — should get OrphanStyle's own props (sz=28, color=green)
    # without crashing despite broken basedOn chain
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    ps = p1_ppr.find(qn("w:pStyle"))
    if ps is not None:
        p1_ppr.remove(ps)
    p1_ppr.append(make_element("w:pStyle", {"w:val": "OrphanStyle"}))
    p1.add_run("Orphan style with broken basedOn")

    save_fixture("based-on-nonexistent", doc, {
        "spec": "ISO 29500-1 §17.7.4.3",
        "description": (
            "Paragraph style has basedOn pointing to a non-existent styleId. "
            "Tests that the code doesn't crash and the style's own properties "
            "are still available."
        ),
        "expected": {
            "para_1": "font_size=28, color=00FF00 (OrphanStyle's own props, basedOn ignored)",
        },
    })


# =========================================================================
# Fixture 7: Circular basedOn chain (A -> B -> A)
#
# §17.7.4.3: basedOn should not create cycles. The implementation should
# detect cycles and stop walking the chain.
# =========================================================================

def make_circular_based_on():
    doc = Document()
    styles_el = doc.styles.element

    clear_default_styles(styles_el)

    # CycleA: basedOn CycleB, sz=24
    style_a = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "CycleA",
    })
    style_a.append(make_element("w:name", {"w:val": "Cycle A"}))
    style_a.append(make_element("w:basedOn", {"w:val": "CycleB"}))
    a_rpr = make_element("w:rPr")
    a_rpr.append(make_element("w:sz", {"w:val": "24"}))
    a_rpr.append(make_element("w:b"))
    style_a.append(a_rpr)
    styles_el.append(style_a)

    # CycleB: basedOn CycleA, color=0000FF
    style_b = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "CycleB",
    })
    style_b.append(make_element("w:name", {"w:val": "Cycle B"}))
    style_b.append(make_element("w:basedOn", {"w:val": "CycleA"}))
    b_rpr = make_element("w:rPr")
    b_rpr.append(make_element("w:color", {"w:val": "0000FF"}))
    style_b.append(b_rpr)
    styles_el.append(style_b)

    # Para 1: CycleA — should resolve without infinite loop, getting sz=24 + bold
    # from itself, and color=0000FF from CycleB (the parent before cycle detection)
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    ps = p1_ppr.find(qn("w:pStyle"))
    if ps is not None:
        p1_ppr.remove(ps)
    p1_ppr.append(make_element("w:pStyle", {"w:val": "CycleA"}))
    p1.add_run("CycleA style run")

    # Para 2: CycleB — should resolve without infinite loop, getting color from
    # itself, and sz+bold from CycleA (the parent before cycle detection)
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    ps = p2_ppr.find(qn("w:pStyle"))
    if ps is not None:
        p2_ppr.remove(ps)
    p2_ppr.append(make_element("w:pStyle", {"w:val": "CycleB"}))
    p2.add_run("CycleB style run")

    save_fixture("circular-based-on", doc, {
        "spec": "ISO 29500-1 §17.7.4.3",
        "description": (
            "CycleA basedOn CycleB, CycleB basedOn CycleA. Tests cycle "
            "detection: code should not infinite loop and should resolve "
            "properties from each style's own definitions."
        ),
        "expected": {
            "para_1": "font_size=24, bold=on (CycleA), color=0000FF (CycleB, before cycle stops)",
            "para_2": "color=0000FF (CycleB), font_size=24, bold=on (CycleA, before cycle stops)",
        },
    })


# =========================================================================
# Fixture 8: Cross-type basedOn (paragraph style basedOn character style)
#
# §17.7.4.3: basedOn must reference a style of the same type. A cross-type
# reference should be ignored (the chain stops at the type mismatch).
# =========================================================================

def make_cross_type_based_on():
    doc = Document()
    styles_el = doc.styles.element

    clear_default_styles(styles_el)

    # CharParent: character style with color=FF0000
    style_char = make_element("w:style", {
        "w:type": "character",
        "w:styleId": "CharParent",
    })
    style_char.append(make_element("w:name", {"w:val": "Char Parent"}))
    char_rpr = make_element("w:rPr")
    char_rpr.append(make_element("w:color", {"w:val": "FF0000"}))
    char_rpr.append(make_element("w:sz", {"w:val": "32"}))
    style_char.append(char_rpr)
    styles_el.append(style_char)

    # CrossTypePara: paragraph style with basedOn pointing to CharParent (wrong type!)
    style_cross = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "CrossTypePara",
    })
    style_cross.append(make_element("w:name", {"w:val": "Cross Type Para"}))
    style_cross.append(make_element("w:basedOn", {"w:val": "CharParent"}))
    cross_rpr = make_element("w:rPr")
    cross_rpr.append(make_element("w:b"))
    style_cross.append(cross_rpr)
    styles_el.append(style_cross)

    # Para 1: CrossTypePara — should get bold from its own rPr,
    # but NOT color=FF0000 or sz=32 from CharParent (cross-type stops chain)
    p1 = doc.add_paragraph()
    p1_ppr = p1._element.get_or_add_pPr()
    ps = p1_ppr.find(qn("w:pStyle"))
    if ps is not None:
        p1_ppr.remove(ps)
    p1_ppr.append(make_element("w:pStyle", {"w:val": "CrossTypePara"}))
    p1.add_run("Cross type basedOn para")

    save_fixture("cross-type-based-on", doc, {
        "spec": "ISO 29500-1 §17.7.4.3",
        "description": (
            "Paragraph style CrossTypePara has basedOn pointing to character "
            "style CharParent. Per §17.7.4.3, this cross-type reference should "
            "be ignored. The paragraph style should only use its own properties."
        ),
        "expected": {
            "para_1": "bold=on (CrossTypePara's own rPr), color=None (CharParent NOT inherited)",
        },
    })


# =========================================================================
# Fixture 9: docDefaults pPr contextualSpacing
#
# §17.7.5: docDefaults pPr can set contextualSpacing. This should apply
# to all paragraphs that don't override it.
# =========================================================================

def make_docdefaults_ppr_contextual_spacing():
    doc = Document()
    styles_el = doc.styles.element

    clear_default_styles(styles_el)

    # Clear existing pPrDefault and create fresh with contextualSpacing
    clear_ppr_defaults(styles_el)
    ppr = get_or_create_ppr_defaults(styles_el)
    ppr.append(make_element("w:contextualSpacing"))

    # NoCtxStyle: paragraph style that explicitly sets contextualSpacing=false
    style_noctx = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "NoCtxStyle",
    })
    style_noctx.append(make_element("w:name", {"w:val": "No Ctx Style"}))
    noctx_ppr = make_element("w:pPr")
    noctx_ppr.append(make_element("w:contextualSpacing", {"w:val": "0"}))
    style_noctx.append(noctx_ppr)
    styles_el.append(style_noctx)

    # Para 1: unstyled — should get contextualSpacing=true from docDefaults
    p1 = doc.add_paragraph()
    remove_pstyle(p1._element)
    p1.add_run("Unstyled contextual from docDefaults")

    # Para 2: NoCtxStyle — should get contextualSpacing=false from style
    p2 = doc.add_paragraph()
    p2_ppr = p2._element.get_or_add_pPr()
    ps = p2_ppr.find(qn("w:pStyle"))
    if ps is not None:
        p2_ppr.remove(ps)
    p2_ppr.append(make_element("w:pStyle", {"w:val": "NoCtxStyle"}))
    p2.add_run("NoCtxStyle overrides docDefaults")

    save_fixture("docdefaults-ppr-contextual", doc, {
        "spec": "ISO 29500-1 §17.7.5 + §17.3.1.9",
        "description": (
            "docDefaults pPr sets contextualSpacing=true. Tests that unstyled "
            "paragraphs get this value and that a style can override it."
        ),
        "expected": {
            "para_1": "contextualSpacing=true (from docDefaults pPr)",
            "para_2": "contextualSpacing=false (from NoCtxStyle, overrides docDefaults)",
        },
    })


# =========================================================================
# Fixture 10: Normal style with basedOn — paragraph property inheritance
#
# When Normal basedOn CustomBase, and CustomBase sets spacing_after=120,
# does the unstyled paragraph get spacing_after=120 through Normal's
# basedOn chain?
# =========================================================================

def make_normal_based_on_ppr():
    doc = Document()
    styles_el = doc.styles.element

    clear_default_styles(styles_el)

    # SpacingBase: paragraph style with spacing after=200, before=100
    style_base = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "SpacingBase",
    })
    style_base.append(make_element("w:name", {"w:val": "Spacing Base"}))
    base_ppr = make_element("w:pPr")
    base_ppr.append(make_element("w:spacing", {"w:after": "200", "w:before": "100"}))
    base_ppr.append(make_element("w:jc", {"w:val": "center"}))
    style_base.append(base_ppr)
    styles_el.append(style_base)

    # Normal: default para style, basedOn SpacingBase, overrides before=0
    style_normal = make_element("w:style", {
        "w:type": "paragraph",
        "w:styleId": "Normal",
        "w:default": "1",
    })
    style_normal.append(make_element("w:name", {"w:val": "Normal"}))
    style_normal.append(make_element("w:basedOn", {"w:val": "SpacingBase"}))
    normal_ppr = make_element("w:pPr")
    normal_ppr.append(make_element("w:spacing", {"w:before": "0"}))
    style_normal.append(normal_ppr)
    styles_el.append(style_normal)

    # Para 1: unstyled (implicit Normal) — should get:
    #   spacing_before=0 (Normal overrides SpacingBase)
    #   spacing_after=200 (inherited from SpacingBase)
    #   alignment=center (inherited from SpacingBase)
    p1 = doc.add_paragraph()
    remove_pstyle(p1._element)
    p1.add_run("Implicit Normal inherits from SpacingBase")

    save_fixture("normal-based-on-ppr", doc, {
        "spec": "ISO 29500-1 §17.7.4.3 + §17.7.4.17",
        "description": (
            "Normal basedOn SpacingBase. SpacingBase sets after=200, before=100, "
            "jc=center. Normal overrides before=0. Tests that paragraph property "
            "inheritance through Normal's basedOn chain works correctly."
        ),
        "expected": {
            "para_1": "before=0 (Normal), after=200 (SpacingBase), alignment=center (SpacingBase)",
        },
    })


# =========================================================================
# Main
# =========================================================================

if __name__ == "__main__":
    print("Generating style-cascade-defaults-edge fixtures:")
    make_docdefaults_ppr_alignment()
    make_docdefaults_ppr_indent()
    make_docdefaults_ppr_spacing_merge()
    make_normal_with_based_on()
    make_multiple_default_para_styles()
    make_based_on_nonexistent()
    make_circular_based_on()
    make_cross_type_based_on()
    make_docdefaults_ppr_contextual_spacing()
    make_normal_based_on_ppr()
    print("Done.")

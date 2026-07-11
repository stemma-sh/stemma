# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "lxml",
# ]
# ///
"""
Generate DOCX fixtures for table border weight ranking tests
(MS-OI29500 §2.1.169).

Run:  cd stemma-engine/testdata/spec-compliance/table-border-weight-ranking && mise exec -- uv run create_docs.py
"""

import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def w(tag: str) -> str:
    return f"{{{W}}}{tag}"


def make_element(tag: str, attribs: dict | None = None) -> OxmlElement:
    el = OxmlElement(tag)
    if attribs:
        for k, v in attribs.items():
            el.set(qn(k), v)
    return el


def save_fixture(name: str, doc, metadata: dict) -> None:
    out = ROOT / name
    out.mkdir(parents=True, exist_ok=True)
    doc.save(str(out / "input.docx"))
    (out / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")
    print(f"  table-border-weight-ranking/{name}/")


def add_tbl_grid(tbl_el, col_widths):
    """Add or replace tblGrid with specific column widths."""
    existing = tbl_el.find(w("tblGrid"))
    if existing is not None:
        tbl_el.remove(existing)
    grid = make_element("w:tblGrid")
    for cw in col_widths:
        grid.append(make_element("w:gridCol", {"w:w": str(cw)}))
    tblPr = tbl_el.find(w("tblPr"))
    if tblPr is not None:
        idx = list(tbl_el).index(tblPr) + 1
        tbl_el.insert(idx, grid)
    else:
        tbl_el.insert(0, grid)


def add_tc_borders(tbl_el, row, col, edges_spec):
    """Add tcBorders to a specific cell.

    edges_spec is a list of (edge_name, val, sz, color, space) tuples.
    """
    tr = tbl_el.findall(w("tr"))[row]
    tc = tr.findall(w("tc"))[col]
    tcPr = tc.find(w("tcPr"))
    if tcPr is None:
        tcPr = make_element("w:tcPr")
        tc.insert(0, tcPr)
    tc_borders = make_element("w:tcBorders")
    for edge, val, sz, color, space in edges_spec:
        border = make_element(f"w:{edge}", {
            "w:val": val, "w:sz": sz, "w:color": color, "w:space": space,
        })
        tc_borders.append(border)
    tcPr.append(tc_borders)


# =====================================================================
# Fixture 1: compound-beats-single
# MS-OI29500 §2.1.169: Compound borders have higher rank than single.
# =====================================================================

def make_compound_beats_single():
    """1-row, 3-column table.

    Cell (0,0) right: thinThickSmallGap sz=8 red (rank 7, weight=56)
    Cell (0,1) left:  single sz=8 blue (rank 1, weight=8)
    Cell (0,1) right: wave sz=8 green (rank 16, weight=128)
    Cell (0,2) left:  single sz=8 blue (rank 1, weight=8)

    At edge (0,0)|(0,1): thinThickSmallGap wins (56 > 8). Red.
    At edge (0,1)|(0,2): wave wins (128 > 8). Green.
    """
    doc = Document()
    doc.add_paragraph("Compound border styles beat single at equal size.")

    tbl = doc.add_table(rows=1, cols=3)
    tbl.cell(0, 0).text = "TTSG-Right"
    tbl.cell(0, 1).text = "Middle"
    tbl.cell(0, 2).text = "Right"

    tbl_el = tbl._tbl
    add_tbl_grid(tbl_el, [2000, 2000, 2000])

    # Cell (0,0): right = thinThickSmallGap sz=8 red
    add_tc_borders(tbl_el, 0, 0, [
        ("right", "thinThickSmallGap", "8", "FF0000", "0"),
    ])
    # Cell (0,1): left = single sz=8 blue, right = wave sz=8 green
    add_tc_borders(tbl_el, 0, 1, [
        ("left", "single", "8", "0000FF", "0"),
        ("right", "wave", "8", "00FF00", "0"),
    ])
    # Cell (0,2): left = single sz=8 blue
    add_tc_borders(tbl_el, 0, 2, [
        ("left", "single", "8", "0000FF", "0"),
    ])

    save_fixture("compound-beats-single", doc, {
        "name": "compound-beats-single",
        "spec_ref": "MS-OI29500 §2.1.169",
        "description": (
            "1x3 table. Cell (0,0) right=thinThickSmallGap sz=8 red (rank 7, weight=56). "
            "Cell (0,1) left=single sz=8 blue (rank 1, weight=8), "
            "right=wave sz=8 green (rank 16, weight=128). "
            "Cell (0,2) left=single sz=8 blue (rank 1, weight=8). "
            "At edge (0,0)|(0,1): red wins. At edge (0,1)|(0,2): green wins."
        ),
    })


# =====================================================================
# Fixture 2: dotted-dashed-always-weight-1
# MS-OI29500 §2.1.169: Dotted/dashed borders always get weight 1.
# =====================================================================

def make_dotted_dashed_always_weight_1():
    """1-row, 2-column table.

    Cell (0,0) right: dotted sz=24 red (weight=1 per spec, regardless of size)
    Cell (0,1) left:  single sz=4 blue (rank 1, weight=4)

    Per MS-OI29500 §2.1.169: "The borders with dotted and dashed styles
    shall be assigned the weight 1 regardless of the border width and number."
    Single sz=4 (weight=4) beats dotted (weight=1). Blue wins.
    """
    doc = Document()
    doc.add_paragraph("Dotted/dashed borders always get weight 1.")

    tbl = doc.add_table(rows=1, cols=2)
    tbl.cell(0, 0).text = "Dotted"
    tbl.cell(0, 1).text = "Single"

    tbl_el = tbl._tbl
    add_tbl_grid(tbl_el, [3000, 3000])

    # Cell (0,0): right = dotted sz=24 red
    add_tc_borders(tbl_el, 0, 0, [
        ("right", "dotted", "24", "FF0000", "0"),
    ])
    # Cell (0,1): left = single sz=4 blue
    add_tc_borders(tbl_el, 0, 1, [
        ("left", "single", "4", "0000FF", "0"),
    ])

    save_fixture("dotted-dashed-always-weight-1", doc, {
        "name": "dotted-dashed-always-weight-1",
        "spec_ref": "MS-OI29500 §2.1.169",
        "description": (
            "1x2 table. Cell (0,0) right=dotted sz=24 red (weight=1 per spec). "
            "Cell (0,1) left=single sz=4 blue (weight=4). "
            "Dotted always gets weight 1 regardless of size, so blue single wins."
        ),
    })


# =====================================================================
# Fixture 3: 3d-emboss-high-rank
# MS-OI29500 §2.1.169: threeDEmboss has rank 20.
# =====================================================================

def make_3d_emboss_high_rank():
    """1-row, 2-column table.

    Cell (0,0) right: threeDEmboss sz=4 red (rank 20, weight=80)
    Cell (0,1) left:  double sz=4 blue (rank 3, weight=12)

    threeDEmboss wins (80 > 12). Red.
    """
    doc = Document()
    doc.add_paragraph("threeDEmboss (rank 20) beats double (rank 3) at same size.")

    tbl = doc.add_table(rows=1, cols=2)
    tbl.cell(0, 0).text = "3DEmboss"
    tbl.cell(0, 1).text = "Double"

    tbl_el = tbl._tbl
    add_tbl_grid(tbl_el, [3000, 3000])

    # Cell (0,0): right = threeDEmboss sz=4 red
    add_tc_borders(tbl_el, 0, 0, [
        ("right", "threeDEmboss", "4", "FF0000", "0"),
    ])
    # Cell (0,1): left = double sz=4 blue
    add_tc_borders(tbl_el, 0, 1, [
        ("left", "double", "4", "0000FF", "0"),
    ])

    save_fixture("3d-emboss-high-rank", doc, {
        "name": "3d-emboss-high-rank",
        "spec_ref": "MS-OI29500 §2.1.169",
        "description": (
            "1x2 table. Cell (0,0) right=threeDEmboss sz=4 red (rank 20, weight=80). "
            "Cell (0,1) left=double sz=4 blue (rank 3, weight=12). "
            "threeDEmboss wins. Red."
        ),
    })


# =====================================================================
# Main
# =====================================================================

def main():
    print("\n== Table Border Weight Ranking Fixtures ==")
    make_compound_beats_single()
    make_dotted_dashed_always_weight_1()
    make_3d_emboss_high_rank()
    print("\nDone.")


if __name__ == "__main__":
    main()

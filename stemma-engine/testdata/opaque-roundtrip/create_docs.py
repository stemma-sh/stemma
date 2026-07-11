# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
#     "Pillow",
# ]
# ///
"""
Generate before/after DOCX pairs that exercise the opaque roundtrip path.

The documents contain elements we do NOT semantically model but preserve
as opaque blobs: SDTs (content controls), fldSimple fields, hyperlinks,
bookmarks, and inline images.  Text changes are made *around* these
elements so the diff pipeline must process paragraphs that contain them
without corrupting or losing the opaque content.

This tests the safety net described in the audit:
  "Opaque preservation strategy — Unrecognized elements are roundtripped
   as raw XML rather than parsed."
"""

import io
import json
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def _make_test_png() -> bytes:
    """Generate a minimal 4x4 red PNG using Pillow."""
    from PIL import Image as PILImage
    img = PILImage.new("RGB", (4, 4), color=(255, 0, 0))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


TEST_PNG = _make_test_png()

OUTPUT_DIR = Path(__file__).parent


# ---------------------------------------------------------------------------
# XML helpers — all use OxmlElement (lxml-based) for python-docx compat
# ---------------------------------------------------------------------------

def _child(parent, tag: str):
    """Create an OxmlElement, append to parent, return it."""
    el = OxmlElement(tag)
    parent.append(el)
    return el


def _run_with_text(parent, text: str, preserve_space: bool = False):
    """Append a w:r with w:t to parent."""
    r = _child(parent, "w:r")
    t = _child(r, "w:t")
    t.text = text
    if preserve_space:
        t.set(qn("xml:space"), "preserve")
    return r


# ---------------------------------------------------------------------------
# Opaque element injectors
# ---------------------------------------------------------------------------

def add_sdt_block(doc: DocxDocument, alias: str, content_text: str) -> None:
    """Inject a block-level SDT (content control) into the document body.

    SDTs are not modeled in our canonical IR — they are preserved as opaque
    block elements during roundtrip.
    """
    sdt = OxmlElement("w:sdt")

    sdt_pr = _child(sdt, "w:sdtPr")
    alias_el = _child(sdt_pr, "w:alias")
    alias_el.set(qn("w:val"), alias)
    tag_el = _child(sdt_pr, "w:tag")
    tag_el.set(qn("w:val"), alias.lower().replace(" ", "_"))

    sdt_content = _child(sdt, "w:sdtContent")
    p = _child(sdt_content, "w:p")
    _run_with_text(p, content_text)

    doc.element.body.append(sdt)


def add_field_paragraph(
    doc: DocxDocument,
    prefix: str,
    field_instr: str,
    field_result: str,
    suffix: str,
) -> None:
    """Add a paragraph with text, an inline fldSimple field, and more text.

    fldSimple is modeled as an opaque inline (Field kind) in our IR.
    The surrounding text should diff normally while the field survives.
    """
    p = OxmlElement("w:p")

    _run_with_text(p, prefix, preserve_space=True)

    fld = _child(p, "w:fldSimple")
    fld.set(qn("w:instr"), field_instr)
    _run_with_text(fld, field_result)

    _run_with_text(p, suffix, preserve_space=True)

    doc.element.body.append(p)


def add_hyperlink_paragraph(
    doc: DocxDocument,
    prefix: str,
    link_text: str,
    url: str,
    suffix: str,
) -> None:
    """Add a paragraph with an inline hyperlink surrounded by text.

    Hyperlinks are modeled as opaque inlines (Hyperlink kind).
    """
    p = OxmlElement("w:p")

    _run_with_text(p, prefix, preserve_space=True)

    # Register external hyperlink relationship
    r_id = doc.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = _child(p, "w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    hr = _child(hyperlink, "w:r")
    rpr = _child(hr, "w:rPr")
    rstyle = _child(rpr, "w:rStyle")
    rstyle.set(qn("w:val"), "Hyperlink")
    ht = _child(hr, "w:t")
    ht.text = link_text

    _run_with_text(p, suffix, preserve_space=True)

    doc.element.body.append(p)


def add_bookmark_paragraph(
    doc: DocxDocument,
    text_before: str,
    bookmarked_text: str,
    bookmark_name: str,
    text_after: str,
) -> None:
    """Add a paragraph with a bookmark wrapping some text.

    Bookmarks (bookmarkStart/bookmarkEnd) are preserved in the XML
    but not semantically modeled in our IR.
    """
    p = OxmlElement("w:p")

    _run_with_text(p, text_before, preserve_space=True)

    bm_start = _child(p, "w:bookmarkStart")
    bm_start.set(qn("w:id"), "0")
    bm_start.set(qn("w:name"), bookmark_name)

    _run_with_text(p, bookmarked_text)

    bm_end = _child(p, "w:bookmarkEnd")
    bm_end.set(qn("w:id"), "0")

    _run_with_text(p, text_after, preserve_space=True)

    doc.element.body.append(p)


def inject_sdt_into_cell(cell, alias: str, text: str) -> None:
    """Replace default cell content with an SDT content control."""
    cell_el = cell._element
    sdt = OxmlElement("w:sdt")
    sdt_pr = _child(sdt, "w:sdtPr")
    alias_el = _child(sdt_pr, "w:alias")
    alias_el.set(qn("w:val"), alias)
    sdt_content = _child(sdt, "w:sdtContent")
    sp = _child(sdt_content, "w:p")
    _run_with_text(sp, text)
    # Remove default empty paragraph, add SDT
    for child in list(cell_el):
        if child.tag == qn("w:p"):
            cell_el.remove(child)
    cell_el.append(sdt)


# ---------------------------------------------------------------------------
# Document builders
# ---------------------------------------------------------------------------

def build_before() -> DocxDocument:
    doc = Document()

    # 1. Plain paragraph (anchor — unchanged between versions)
    doc.add_paragraph(
        "This document tests that opaque XML elements survive the "
        "parse-diff-redline-export pipeline without corruption."
    )

    # 2. Paragraph with inline fldSimple field, text changes around it
    add_field_paragraph(
        doc,
        prefix="Contract dated ",
        field_instr=' DATE \\@ "MMMM d, yyyy" ',
        field_result="January 15, 2025",
        suffix=" is hereby executed.",
    )

    # 3. Paragraph with hyperlink, text changes around it
    add_hyperlink_paragraph(
        doc,
        prefix="For more information, see ",
        link_text="the official documentation",
        url="https://example.com/docs",
        suffix=" for complete details.",
    )

    # 4. Paragraph with bookmark
    add_bookmark_paragraph(
        doc,
        text_before="Section reference: see ",
        bookmarked_text="Clause 4.2",
        bookmark_name="_Ref_Clause_4_2",
        text_after=" for warranty terms.",
    )

    # 5. Block-level SDT (content control)
    add_sdt_block(doc, "Effective Date", "January 15, 2025")

    # 6. Paragraph with inline image (drawing element = opaque widget)
    p = doc.add_paragraph("Exhibit A shows the original system diagram: ")
    run = p.add_run()
    run.add_picture(io.BytesIO(TEST_PNG), width=914400)  # 1 inch
    p.add_run(" as referenced above.")

    # 7. Another plain paragraph that changes
    doc.add_paragraph(
        "The parties agree to the original terms and conditions "
        "set forth in the preceding sections of this agreement."
    )

    # 8. Table with an SDT inside a cell
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Party"
    table.cell(0, 1).text = "Signature"
    table.cell(1, 0).text = "Vendor Corp."
    inject_sdt_into_cell(table.cell(1, 1), "Signature", "[Sign here]")

    return doc


def build_after() -> DocxDocument:
    """Same structure, but with text changes around the opaque elements."""
    doc = Document()

    # 1. Plain paragraph — UNCHANGED (anchor)
    doc.add_paragraph(
        "This document tests that opaque XML elements survive the "
        "parse-diff-redline-export pipeline without corruption."
    )

    # 2. Field paragraph — prefix and suffix text changed
    add_field_paragraph(
        doc,
        prefix="Agreement effective ",  # was "Contract dated "
        field_instr=' DATE \\@ "MMMM d, yyyy" ',
        field_result="January 15, 2025",
        suffix=" is now in force.",  # was " is hereby executed."
    )

    # 3. Hyperlink paragraph — surrounding text changed
    add_hyperlink_paragraph(
        doc,
        prefix="Please refer to ",  # was "For more information, see "
        link_text="the official documentation",
        url="https://example.com/docs",
        suffix=" for full specifications.",  # was " for complete details."
    )

    # 4. Bookmark paragraph — surrounding text changed
    add_bookmark_paragraph(
        doc,
        text_before="Cross-reference: refer to ",  # was "Section reference: see "
        bookmarked_text="Clause 4.2",
        bookmark_name="_Ref_Clause_4_2",
        text_after=" regarding warranty obligations.",  # was " for warranty terms."
    )

    # 5. Block-level SDT — UNCHANGED (opaque block should survive)
    add_sdt_block(doc, "Effective Date", "January 15, 2025")

    # 6. Image paragraph — surrounding text changed, same image
    p = doc.add_paragraph("Exhibit A depicts the updated system diagram: ")
    run = p.add_run()
    run.add_picture(io.BytesIO(TEST_PNG), width=914400)
    p.add_run(" as shown above.")  # was " as referenced above."

    # 7. Changed paragraph
    doc.add_paragraph(
        "The parties agree to the revised terms and conditions "  # was "original"
        "set forth in the preceding sections of this agreement."
    )

    # 8. Table — cell text changed, SDT in other cell unchanged
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Party"
    table.cell(0, 1).text = "Signature"
    table.cell(1, 0).text = "Vendor Inc."  # was "Vendor Corp."
    inject_sdt_into_cell(table.cell(1, 1), "Signature", "[Sign here]")

    return doc


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    before = build_before()
    after = build_after()

    before.save(str(OUTPUT_DIR / "before.docx"))
    after.save(str(OUTPUT_DIR / "after.docx"))

    metadata = {
        "name": "opaque-roundtrip",
        "description": (
            "Opaque element roundtrip safety net: SDTs, fldSimple fields, "
            "hyperlinks, bookmarks, and inline images with text changes around them"
        ),
        "opaque_elements": [
            "w:sdt (block-level content control)",
            "w:fldSimple (inline field)",
            "w:hyperlink (inline hyperlink)",
            "w:bookmarkStart/End (bookmark pair)",
            "w:drawing (inline image)",
            "w:sdt in table cell",
        ],
        "changes": [
            "Text around fldSimple field changed",
            "Text around hyperlink changed",
            "Text around bookmark changed",
            "Text around inline image changed",
            "Plain paragraph text changed",
            "Table cell text changed (adjacent to SDT cell)",
            "Block-level SDT unchanged",
        ],
    }
    (OUTPUT_DIR / "metadata.json").write_text(json.dumps(metadata, indent=2) + "\n")

    print("Created opaque-roundtrip sample:")
    print(f"  {OUTPUT_DIR / 'before.docx'}")
    print(f"  {OUTPUT_DIR / 'after.docx'}")
    print(f"  {OUTPUT_DIR / 'metadata.json'}")


if __name__ == "__main__":
    main()

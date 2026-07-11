# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx",
# ]
# ///
"""
Create showcase documents for Stemma landing page.

These documents demonstrate critical changes that are hard to spot
in standard Word redlines but are immediately surfaced by Stemma.
"""

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def add_heading(doc: DocxDocument, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_paragraph(doc: DocxDocument, text: str, bold_first: bool = False) -> None:
    p = doc.add_paragraph()
    if bold_first and ". " in text:
        title, rest = text.split(". ", 1)
        run = p.add_run(title + ". ")
        run.bold = True
        p.add_run(rest)
    else:
        p.add_run(text)


def create_before_document() -> DocxDocument:
    """Create the original (favorable to buyer) contract."""
    doc = Document()

    # Title
    title = doc.add_heading("SOFTWARE LICENSE AGREEMENT", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph(
        "This Software License Agreement (\"Agreement\") is entered into as of "
        "January 15, 2025 (\"Effective Date\") by and between TechVendor Inc., "
        "a Delaware corporation (\"Vendor\"), and Customer Corp., a California "
        "corporation (\"Customer\")."
    )

    # Section 1: License Grant
    add_heading(doc, "1. LICENSE GRANT", level=1)
    add_paragraph(
        doc,
        "1.1 Grant. Subject to the terms of this Agreement and payment of the "
        "applicable fees, Vendor grants to Customer a non-exclusive, worldwide "
        "license to use the Software for Customer's internal business purposes."
    )
    add_paragraph(
        doc,
        "1.2 Restrictions. Customer shall not: (a) sublicense, sell, or transfer "
        "the Software; (b) modify or create derivative works; (c) reverse engineer "
        "or decompile the Software; or (d) use the Software for any unlawful purpose."
    )

    # Section 2: Fees (THE TRAP - $50,000 vs $500,000)
    add_heading(doc, "2. FEES AND PAYMENT", level=1)
    add_paragraph(
        doc,
        "2.1 License Fee. Customer shall pay Vendor an annual license fee of "
        "Fifty Thousand Dollars ($50,000) payable in advance on each anniversary "
        "of the Effective Date."
    )
    add_paragraph(
        doc,
        "2.2 Payment Terms. All payments are due within thirty (30) days of "
        "invoice. Late payments shall accrue interest at the rate of 1.5% per "
        "month or the maximum rate permitted by law, whichever is less."
    )

    # Section 3: Term and Termination (THE TRAP - 90 days vs 30 days)
    add_heading(doc, "3. TERM AND TERMINATION", level=1)
    add_paragraph(
        doc,
        "3.1 Term. This Agreement shall commence on the Effective Date and "
        "continue for an initial term of three (3) years, unless earlier "
        "terminated as provided herein."
    )
    add_paragraph(
        doc,
        "3.2 Termination for Convenience. Either party may terminate this "
        "Agreement for any reason upon ninety (90) days' prior written notice "
        "to the other party."
    )
    add_paragraph(
        doc,
        "3.3 Termination for Cause. Either party may terminate this Agreement "
        "immediately upon written notice if the other party materially breaches "
        "this Agreement and fails to cure such breach within thirty (30) days "
        "after receipt of written notice thereof."
    )

    # Section 4: Warranty (THE TRAP - removal of warranty)
    add_heading(doc, "4. WARRANTIES", level=1)
    add_paragraph(
        doc,
        "4.1 Performance Warranty. Vendor warrants that the Software will "
        "perform substantially in accordance with its documentation for a "
        "period of twelve (12) months from delivery."
    )
    add_paragraph(
        doc,
        "4.2 Authority. Each party represents and warrants that it has full "
        "power and authority to enter into this Agreement."
    )
    add_paragraph(
        doc,
        "4.3 Malicious Code. Vendor warrants that the Software, as delivered, "
        "will be free from viruses, malware, and other malicious code."
    )

    # Section 5: Indemnification (THE TRAP - carveout addition)
    add_heading(doc, "5. INDEMNIFICATION", level=1)
    add_paragraph(
        doc,
        "5.1 Vendor Indemnification. Vendor shall defend, indemnify, and hold "
        "harmless Customer and its officers, directors, employees, and agents "
        "from and against any and all claims, damages, losses, costs, and "
        "expenses (including reasonable attorneys' fees) arising out of or "
        "relating to any third-party claim that the Software infringes any "
        "intellectual property right."
    )
    add_paragraph(
        doc,
        "5.2 Customer Indemnification. Customer shall defend, indemnify, and "
        "hold harmless Vendor from and against any claims arising out of "
        "Customer's use of the Software in violation of this Agreement."
    )

    # Section 6: Limitation of Liability (THE TRAP - cap removal)
    add_heading(doc, "6. LIMITATION OF LIABILITY", level=1)
    add_paragraph(
        doc,
        "6.1 Liability Cap. IN NO EVENT SHALL VENDOR'S TOTAL AGGREGATE "
        "LIABILITY UNDER THIS AGREEMENT EXCEED THE AMOUNT OF FEES PAID BY "
        "CUSTOMER TO VENDOR DURING THE TWELVE (12) MONTHS PRECEDING THE "
        "CLAIM. This limitation shall apply to all claims, whether based on "
        "warranty, contract, tort, or any other legal theory."
    )
    add_paragraph(
        doc,
        "6.2 Exclusion. IN NO EVENT SHALL EITHER PARTY BE LIABLE FOR ANY "
        "INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES."
    )

    # Section 7: Support (THE TRAP - "best efforts" to "commercially reasonable")
    add_heading(doc, "7. SUPPORT AND MAINTENANCE", level=1)
    add_paragraph(
        doc,
        "7.1 Support Services. Vendor shall use best efforts to provide "
        "technical support during normal business hours (9 AM to 6 PM Eastern "
        "Time, Monday through Friday, excluding holidays)."
    )
    add_paragraph(
        doc,
        "7.2 Response Times. Vendor shall use best efforts to respond to "
        "critical issues within four (4) hours and non-critical issues within "
        "twenty-four (24) hours."
    )
    add_paragraph(
        doc,
        "7.3 Updates. Vendor shall provide all updates, patches, and bug fixes "
        "released during the term of this Agreement at no additional cost."
    )

    # Section 8: Confidentiality
    add_heading(doc, "8. CONFIDENTIALITY", level=1)
    add_paragraph(
        doc,
        "8.1 Confidential Information. Each party agrees to maintain the "
        "confidentiality of any proprietary or confidential information "
        "disclosed by the other party and to use such information only for "
        "purposes of this Agreement."
    )
    add_paragraph(
        doc,
        "8.2 Survival. The obligations of confidentiality shall survive "
        "termination of this Agreement for a period of three (3) years."
    )

    # Section 9: General
    add_heading(doc, "9. GENERAL PROVISIONS", level=1)
    add_paragraph(
        doc,
        "9.1 Governing Law. This Agreement shall be governed by and construed "
        "in accordance with the laws of the State of California, without regard "
        "to its conflict of laws principles."
    )
    add_paragraph(
        doc,
        "9.2 Dispute Resolution. Any dispute arising out of this Agreement "
        "shall be resolved by binding arbitration in San Francisco, California, "
        "in accordance with the rules of the American Arbitration Association."
    )
    add_paragraph(
        doc,
        "9.3 Entire Agreement. This Agreement constitutes the entire agreement "
        "between the parties and supersedes all prior agreements, understandings, "
        "and communications."
    )
    add_paragraph(
        doc,
        "9.4 Amendment. This Agreement may not be amended except by a written "
        "instrument signed by both parties."
    )
    add_paragraph(
        doc,
        "9.5 Assignment. Neither party may assign this Agreement without the "
        "prior written consent of the other party."
    )

    # Signature block
    doc.add_paragraph()
    doc.add_paragraph("IN WITNESS WHEREOF, the parties have executed this Agreement as of the Effective Date.")
    doc.add_paragraph()

    # Create signature table
    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "TECHVENDOR INC."
    table.cell(0, 1).text = "CUSTOMER CORP."
    table.cell(1, 0).text = "By: _________________________"
    table.cell(1, 1).text = "By: _________________________"
    table.cell(2, 0).text = "Name:"
    table.cell(2, 1).text = "Name:"
    table.cell(3, 0).text = "Title:"
    table.cell(3, 1).text = "Title:"

    return doc


def create_after_document() -> DocxDocument:
    """Create the revised contract with sneaky but critical changes."""
    doc = Document()

    # Title
    title = doc.add_heading("SOFTWARE LICENSE AGREEMENT", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph(
        "This Software License Agreement (\"Agreement\") is entered into as of "
        "January 15, 2025 (\"Effective Date\") by and between TechVendor Inc., "
        "a Delaware corporation (\"Vendor\"), and Customer Corp., a California "
        "corporation (\"Customer\")."
    )

    # Section 1: License Grant
    add_heading(doc, "1. LICENSE GRANT", level=1)
    add_paragraph(
        doc,
        "1.1 Grant. Subject to the terms of this Agreement and payment of the "
        "applicable fees, Vendor grants to Customer a non-exclusive, worldwide "
        "license to use the Software for Customer's internal business purposes."
    )
    add_paragraph(
        doc,
        "1.2 Restrictions. Customer shall not: (a) sublicense, sell, or transfer "
        "the Software; (b) modify or create derivative works; (c) reverse engineer "
        "or decompile the Software; or (d) use the Software for any unlawful purpose."
    )

    # Section 2: Fees - CHANGED: $50,000 → $500,000
    add_heading(doc, "2. FEES AND PAYMENT", level=1)
    add_paragraph(
        doc,
        "2.1 License Fee. Customer shall pay Vendor an annual license fee of "
        "Five Hundred Thousand Dollars ($500,000) payable in advance on each anniversary "
        "of the Effective Date."
    )
    add_paragraph(
        doc,
        "2.2 Payment Terms. All payments are due within thirty (30) days of "
        "invoice. Late payments shall accrue interest at the rate of 1.5% per "
        "month or the maximum rate permitted by law, whichever is less."
    )

    # Section 3: Term and Termination - CHANGED: 90 days → 30 days
    add_heading(doc, "3. TERM AND TERMINATION", level=1)
    add_paragraph(
        doc,
        "3.1 Term. This Agreement shall commence on the Effective Date and "
        "continue for an initial term of three (3) years, unless earlier "
        "terminated as provided herein."
    )
    add_paragraph(
        doc,
        "3.2 Termination for Convenience. Either party may terminate this "
        "Agreement for any reason upon thirty (30) days' prior written notice "  # CHANGED from 90
        "to the other party."
    )
    add_paragraph(
        doc,
        "3.3 Termination for Cause. Either party may terminate this Agreement "
        "immediately upon written notice if the other party materially breaches "
        "this Agreement and fails to cure such breach within thirty (30) days "
        "after receipt of written notice thereof."
    )

    # Section 4: Warranty - CHANGED: Removed malicious code warranty entirely
    add_heading(doc, "4. WARRANTIES", level=1)
    add_paragraph(
        doc,
        "4.1 Performance Warranty. Vendor warrants that the Software will "
        "perform substantially in accordance with its documentation for a "
        "period of twelve (12) months from delivery."
    )
    add_paragraph(
        doc,
        "4.2 Authority. Each party represents and warrants that it has full "
        "power and authority to enter into this Agreement."
    )
    # REMOVED: 4.3 Malicious Code warranty

    # Section 5: Indemnification - CHANGED: Added significant carveout
    add_heading(doc, "5. INDEMNIFICATION", level=1)
    add_paragraph(
        doc,
        "5.1 Vendor Indemnification. Vendor shall defend, indemnify, and hold "
        "harmless Customer and its officers, directors, employees, and agents "
        "from and against any and all claims, damages, losses, costs, and "
        "expenses (including reasonable attorneys' fees) arising out of or "
        "relating to any third-party claim that the Software infringes any "
        "intellectual property right, except where such infringement arises "  # ADDED carveout
        "from Customer's modification of the Software, combination with other "
        "products, or use outside the scope of this Agreement."
    )
    add_paragraph(
        doc,
        "5.2 Customer Indemnification. Customer shall defend, indemnify, and "
        "hold harmless Vendor from and against any claims arising out of "
        "Customer's use of the Software in violation of this Agreement."
    )

    # Section 6: Limitation of Liability - CHANGED: Removed liability cap!
    add_heading(doc, "6. LIMITATION OF LIABILITY", level=1)
    # REMOVED the liability cap clause entirely!
    add_paragraph(
        doc,
        "6.1 Exclusion. IN NO EVENT SHALL EITHER PARTY BE LIABLE FOR ANY "  # Was 6.2
        "INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR PUNITIVE DAMAGES."
    )

    # Section 7: Support - CHANGED: "best efforts" → "commercially reasonable efforts"
    add_heading(doc, "7. SUPPORT AND MAINTENANCE", level=1)
    add_paragraph(
        doc,
        "7.1 Support Services. Vendor shall use commercially reasonable efforts to provide "  # CHANGED
        "technical support during normal business hours (9 AM to 6 PM Eastern "
        "Time, Monday through Friday, excluding holidays)."
    )
    add_paragraph(
        doc,
        "7.2 Response Times. Vendor shall use commercially reasonable efforts to respond to "  # CHANGED
        "critical issues within four (4) hours and non-critical issues within "
        "twenty-four (24) hours."
    )
    add_paragraph(
        doc,
        "7.3 Updates. Vendor shall provide all updates, patches, and bug fixes "
        "released during the term of this Agreement at no additional cost."
    )

    # Section 8: Confidentiality
    add_heading(doc, "8. CONFIDENTIALITY", level=1)
    add_paragraph(
        doc,
        "8.1 Confidential Information. Each party agrees to maintain the "
        "confidentiality of any proprietary or confidential information "
        "disclosed by the other party and to use such information only for "
        "purposes of this Agreement."
    )
    add_paragraph(
        doc,
        "8.2 Survival. The obligations of confidentiality shall survive "
        "termination of this Agreement for a period of three (3) years."
    )

    # Section 9: General - CHANGED: California → Delaware, San Francisco → Wilmington
    add_heading(doc, "9. GENERAL PROVISIONS", level=1)
    add_paragraph(
        doc,
        "9.1 Governing Law. This Agreement shall be governed by and construed "
        "in accordance with the laws of the State of Delaware, without regard "  # CHANGED
        "to its conflict of laws principles."
    )
    add_paragraph(
        doc,
        "9.2 Dispute Resolution. Any dispute arising out of this Agreement "
        "shall be resolved by binding arbitration in Wilmington, Delaware, "  # CHANGED
        "in accordance with the rules of the American Arbitration Association."
    )
    add_paragraph(
        doc,
        "9.3 Entire Agreement. This Agreement constitutes the entire agreement "
        "between the parties and supersedes all prior agreements, understandings, "
        "and communications."
    )
    add_paragraph(
        doc,
        "9.4 Amendment. This Agreement may not be amended except by a written "
        "instrument signed by both parties."
    )
    add_paragraph(
        doc,
        "9.5 Assignment. Neither party may assign this Agreement without the "
        "prior written consent of the other party."
    )

    # Signature block
    doc.add_paragraph()
    doc.add_paragraph("IN WITNESS WHEREOF, the parties have executed this Agreement as of the Effective Date.")
    doc.add_paragraph()

    # Create signature table
    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "TECHVENDOR INC."
    table.cell(0, 1).text = "CUSTOMER CORP."
    table.cell(1, 0).text = "By: _________________________"
    table.cell(1, 1).text = "By: _________________________"
    table.cell(2, 0).text = "Name:"
    table.cell(2, 1).text = "Name:"
    table.cell(3, 0).text = "Title:"
    table.cell(3, 1).text = "Title:"

    return doc


def main() -> None:
    output_dir = Path(__file__).parent
    output_dir.mkdir(parents=True, exist_ok=True)

    # Create and save documents
    before_doc = create_before_document()
    before_doc.save(str(output_dir / "before.docx"))
    print(f"Created: {output_dir / 'before.docx'}")

    after_doc = create_after_document()
    after_doc.save(str(output_dir / "after.docx"))
    print(f"Created: {output_dir / 'after.docx'}")

    # Summary of changes for reference
    print("\n" + "=" * 60)
    print("CRITICAL CHANGES (hard to spot in Word redlines):")
    print("=" * 60)
    print("""
1. LICENSE FEE: $50,000 → $500,000 (10x increase!)
   Location: Section 2.1

2. TERMINATION NOTICE: 90 days → 30 days
   Location: Section 3.2
   Impact: 60 fewer days to transition off the software

3. MALICIOUS CODE WARRANTY: Entirely removed
   Location: Section 4.3 (deleted)
   Impact: No protection against malware in delivered software

4. INDEMNIFICATION CARVEOUT: Added exceptions
   Location: Section 5.1
   Impact: Vendor no longer liable for IP infringement in common scenarios

5. LIABILITY CAP: Entirely removed!
   Location: Section 6.1 (deleted)
   Impact: Unlimited liability exposure for Customer

6. SUPPORT STANDARD: "best efforts" → "commercially reasonable efforts"
   Location: Sections 7.1 and 7.2
   Impact: Weaker support commitment

7. JURISDICTION: California → Delaware, San Francisco → Wilmington
   Location: Sections 9.1 and 9.2
   Impact: Less favorable venue for Customer (California company)
""")


if __name__ == "__main__":
    main()
